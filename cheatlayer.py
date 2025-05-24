#!python3.10
# -*- coding: utf-8 -*-
import os
import signal
import yaml
import math
import math
from collections import deque
import threading
import time
import os
import time
import pyautogui
import firebase_admin
from firebase_admin import credentials, storage, db
from datetime import datetime
from pathlib import Path
global firebase_email
firebase_email = ""
# ---------------------------- Configuration ---------------------------- #

# -------------------------- Initialize Firebase ------------------------- #

def initialize_firebase():
    """
    Initializes the Firebase app with the provided service account.
    """
    if not os.path.exists(os.path.expanduser(SERVICE_ACCOUNT_KEY_PATH)):
        raise FileNotFoundError(f"Service account key file not found at '{SERVICE_ACCOUNT_KEY_PATH}'.")
    try:
        cred = credentials.Certificate(os.path.expanduser(SERVICE_ACCOUNT_KEY_PATH))
        firebase_admin.initialize_app(cred, {
            'storageBucket': STORAGE_BUCKET,
            'databaseURL': DATABASE_URL
        })
        print("Firebase initialized successfully.")
    except: 
        pass


def load_marketing_data(user_email: str):
    """
    Fetch marketing data from Realtime Database at path: agentsbase/<user_email>/rows.
    Returns a list of dict items with keys: searchTarget, voiceover, caption.
    """
    ref = db.reference(f'agentsbase/{user_email}/rows')
    rows = ref.get()
    if not rows:
        return []

    # If rows is a dict with integer keys, convert it to a list in ascending order of keys
    if isinstance(rows, dict):
        sorted_keys = sorted(rows.keys(), key=lambda x: int(x) if x.isdigit() else x)
        marketing_data = [rows[k] for k in sorted_keys]
        return marketing_data
    # If rows is already a list:
    elif isinstance(rows, list):
        return rows
    else:
        return []
# ---------------------------- Take Screenshot --------------------------- #
import time
import json
from seleniumwire import webdriver
from selenium.webdriver.chrome.options import Options as Options2
from youtube_transcript_api import YouTubeTranscriptApi
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
# If you're using Firebase Admin in Python, make sure you've initialized it:
#   import firebase_admin
#   from firebase_admin import credentials, db
#   cred = credentials.Certificate("path/to/your-serviceAccountKey.json")
#   firebase_admin.initialize_app(cred, {
#       'databaseURL': 'https://<your-db>.firebaseio.com'
#   })
import time
import json
from seleniumwire import webdriver
from selenium.webdriver.chrome.options import Options as Options2
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

# If you're using Firebase, make sure to initialize your Firebase app elsewhere in your code
# from firebase_admin import db

# Optional: For YouTube transcripts (unavailable for TikTok)
# from youtube_transcript_api import YouTubeTranscriptApi
import json
import time
import os

from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC



def run_youtube_scraper(user_email: str, url: str):
    import os
    import time
    import json
    import hashlib
    import subprocess

    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from firebase_admin import db
    from selenium.webdriver.chrome.options import Options as Options2

    # Libraries for downloading & transcribing
    import yt_dlp
    import speech_recognition as sr

    # Import AuthBackend to authenticate using the cookies.txt file
    from tiktok_uploader.auth import AuthBackend

    # ------------------------------------------------------------------------------
    # Helper functions for transcript caching
    # ------------------------------------------------------------------------------
    CACHE_FILENAME = "transcripts_cache.json"

    def load_transcript_cache():
        if os.path.exists(CACHE_FILENAME):
            with open(CACHE_FILENAME, "r", encoding="utf-8") as f:
                return json.load(f)
        return {}

    def save_transcript_cache(cache_data):
        with open(CACHE_FILENAME, "w", encoding="utf-8") as f:
            json.dump(cache_data, f, indent=2, ensure_ascii=False)

    transcript_cache = load_transcript_cache()

def run_youtube_scraper(user_email: str, url: str):
    import os
    import time
    import json
    import hashlib
    import subprocess
    import logging

    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    from firebase_admin import db
    from selenium.webdriver.chrome.options import Options as Options2

    # Libraries for downloading
    import yt_dlp
    
    # For Google + PocketSphinx speech recognition
    import speech_recognition as sr

    # For TikTok authentication via cookies.txt
    from tiktok_uploader.auth import AuthBackend

    # --------------------------------------------------------------------
    # Logging config (if you want logs written to a file)
    # --------------------------------------------------------------------
    logging.basicConfig(
        filename='transcription.log',
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s'
    )

    # --------------------------------------------------------------------
    # Custom exception if no transcription is available
    # --------------------------------------------------------------------
    class NoTranscriptionError(Exception):
        pass

    # --------------------------------------------------------------------
    # Caching transcripts to avoid re-downloading or re-transcribing
    # --------------------------------------------------------------------
    CACHE_FILENAME = "transcripts_cache.json"

    def load_transcript_cache():
        if os.path.exists(CACHE_FILENAME):
            with open(CACHE_FILENAME, "r", encoding="utf-8") as f:
                return json.load(f)
        return {}

    def save_transcript_cache(cache_data):
        with open(CACHE_FILENAME, "w", encoding="utf-8") as f:
            json.dump(cache_data, f, indent=2, ensure_ascii=False)

    transcript_cache = load_transcript_cache()

    # --------------------------------------------------------------------
    # extract_audio_to_text (Your snippet-based method)
    #  • Extracts a 16k mono WAV with ffmpeg
    #  • Tries Google Speech Recognition first
    #  • Fallback to PocketSphinx local recognition
    # --------------------------------------------------------------------
    def extract_audio_to_text(video_path):
        """
        Returns the transcribed text from the given video_path.
        Raises NoTranscriptionError if no transcription is found.
        """
        try:
            wav_path = "temp_audio.wav"
            command = [
                "ffmpeg",
                "-i", video_path,
                "-vn",
                "-acodec", "pcm_s16le",
                "-ar", "16000",
                "-ac", "1",
                wav_path,
                '-y'
            ]
            print("Extracting audio: " + ' '.join(command))
            subprocess.run(command, check=True)

            if os.path.getsize(wav_path) == 0:
                raise ValueError("Extracted audio file is empty.")

            recognizer = sr.Recognizer()
            with sr.AudioFile(wav_path) as source:
                audio = recognizer.record(source)

            text = ""
            # Try Google first
            try:
                text = recognizer.recognize_google(audio)
                logging.info("Google SR succeeded.")
            except sr.UnknownValueError:
                print("Google SR: audio unclear.")
            except sr.RequestError as e:
                print(f"Google SR request error: {e}")
            except e:
                print(f"Google SR request error: {e}")


            # Fallback to PocketSphinx
            if not text.strip():
                try:
                    text = recognizer.recognize_sphinx(audio)
                    logging.info("PocketSphinx SR succeeded.")
                except sr.UnknownValueError:
                    print("PocketSphinx: audio unclear.")
                except Exception as e:
                    print(f"PocketSphinx error: {e}")

            # Clean up
            os.remove(wav_path)

            if text.strip():
                return text
            else:
                raise NoTranscriptionError("No transcription found.")
        except NoTranscriptionError as nte:
            print(str(nte))
            raise
        except Exception as e:
            print(f"Transcription error: {str(e)}")
            return ""

    # --------------------------------------------------------------------
    # Download + Transcribe function with caching
    # --------------------------------------------------------------------
    def download_and_transcribe(video_url: str) -> str:
        """
        Downloads a video from the given URL using yt-dlp, then calls
        extract_audio_to_text(...) to retrieve the transcript. Caches
        the transcript by video_url so subsequent runs won’t redo it.
        """
        video_id = hashlib.md5(video_url.encode("utf-8")).hexdigest()
        print(video_id)
        print("dowload")
        # If we already have transcript in cache, skip
        if video_id in transcript_cache:
            print(transcript_cache[video_id])

        outtmpl = f"{video_id}.%(ext)s"
        ydl_opts = {
            "outtmpl": outtmpl,
            'overwrites': True,
            'verbose': True,
            'format': 'best',
            'cookiesfrombrowser': ('chrome',),
            'extract_flat': False,
            'cachedir': False,
            'cookiefile': 'youtube.txt',
            'geo_bypass': True,
            'noplaylist': True,
            'extractor_args': {
                'youtube': {
                    'player_client': ['web']
                }
            },
            'http_headers': {
                'User-Agent': 'Mozilla/5.0',
            },
            'socket_timeout': 15,
            'retries': 3,
        }
        try:
            with yt_dlp.YoutubeDL(ydl_opts) as ydl:
                info = ydl.extract_info(video_url, download=True)
                downloaded_ext = info.get('ext', 'mp4')
                downloaded_file = f"{video_id}.{downloaded_ext}"
        except Exception as e:
            print(f"ERROR: Could not download {video_url}: {e}")
            transcript_cache[video_id] = "Transcription not available."
            save_transcript_cache(transcript_cache)
            return transcript_cache[video_id]

        # Now transcribe
        try:
            transcript = extract_audio_to_text(downloaded_file)
            print("track")
            print(transcript)
            if not transcript:
                transcript = "Transcription not available."
        except NoTranscriptionError:
            transcript = "Transcription not available."
        except Exception as e:
            logging.error(f"Error transcribing {video_url}: {e}")
            transcript = "Transcription not available."

        # Cache
        transcript_cache[video_id] = transcript
        save_transcript_cache(transcript_cache)

        # Clean up the downloaded file
        if os.path.exists(downloaded_file):
            os.remove(downloaded_file)

        return transcript

    # --------------------------------------------------------------------------------
    # We accumulate both YouTube and TikTok scraped data in final_videos_data
    # --------------------------------------------------------------------------------
    final_videos_data = []

    chrome_options = Options2()
    chrome_options.add_argument("--password-store=basic")
    driver = webdriver.Chrome(options=chrome_options)

    # --------------------------------------------------------------------------------
    # YOUTUBE SCRAPING
    # --------------------------------------------------------------------------------
    try:
        driver.get(url)
        time.sleep(60)  # Let the page fully load (user logs in if needed)

        video_elements = driver.find_elements(By.XPATH, "//a[@href and contains(@href, 'watch?v=')]")
        video_links = set(elem.get_attribute("href") for elem in video_elements if elem.get_attribute("href"))

        for link in video_links:
            driver.get(link)
            time.sleep(4)

            # Title
            try:
                title_elem = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located(
                        (By.CSS_SELECTOR, "h1.title.style-scope.ytd-video-primary-info-renderer")
                    )
                )
                title = title_elem.text.strip()
            except:
                title = "Unknown Title"

            # Views
            try:
                views_elem = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, "//span[contains(text(), 'views')]"))
                )
                views_text = views_elem.text.strip()
                views_num = views_text.split()[0].replace(",", "")
                if not views_num.isdigit():
                    views_num = "0"
            except:
                views_num = "0"

            # Description
            try:
                desc_elem = driver.find_element(
                    By.CSS_SELECTOR, "yt-formatted-string.content.style-scope.ytd-video-secondary-info-renderer"
                )
                description = desc_elem.text.strip()
            except:
                description = ""

            # Transcription
            transcript = download_and_transcribe(link) or "Transcription not available."

            # Collect data
            final_videos_data.append({
                "title": title,
                "description": description,
                "views": views_num,
                "url": link,
                "prompt": transcript,
                "video": "repurposed videos from youtube to youtube with reactions"
            })

        # Save to Firebase
        user_id = user_email.replace(".", ",")
        try:
            ref = db.reference(f'final/{user_id}/')
            ref.set(final_videos_data)
        except Exception as e:
            logging.error(f"Error storing YouTube data in Firebase: {e}")

        # Save locally
        with open("youtube_video_data.json", "w", encoding="utf-8") as f:
            json.dump(final_videos_data, f, indent=2, ensure_ascii=False)

    except Exception as e:
        logging.error(f"Exception during YouTube scraping: {e}")

    # --------------------------------------------------------------------------------
    # TIKTOK SCRAPING
    # --------------------------------------------------------------------------------
    try:
        # 1) Authenticate with cookies.txt
        auth = AuthBackend(cookies="/home/openfdistance/cookies.txt")
        auth.authenticate_agent(driver)
        time.sleep(3)

        # 2) Navigate to TikTok profile
        driver.get("https://www.tiktok.com/profile")
        time.sleep(10)

        video_elements = driver.find_elements(By.XPATH, "//a[contains(@href, '/video/')]")

        for elem in video_elements:
            href = elem.get_attribute("href")
            if not href:
                continue

            # View count from <strong>
            try:
                strong_elem = elem.find_element(By.XPATH, ".//strong")
                views_text = strong_elem.text.strip()
            except:
                views_text = "0"

            # Title from <img alt="...">
            try:
                img_elem = elem.find_element(By.XPATH, ".//img")
                title = img_elem.get_attribute("alt") or "Unknown Title"
            except:
                title = "Unknown Title"

            # Download & transcribe
            transcript = download_and_transcribe(href) or "Transcription not available."

            final_videos_data.append({
                "title": title,
                "description": "",
                "views": views_text,
                "url": href,
                "prompt": transcript,
                "video": "repurposed videos from youtube to tiktok with reactions"
            })

        # Store to Firebase (YouTube + TikTok data)
        user_id = user_email.replace(".", ",")
        try:
            ref = db.reference(f'final/{user_id}/')
            ref.set(final_videos_data)
        except Exception as e:
            print(f"Error storing TikTok data in Firebase: {e}")

        # Save locally
        with open("tiktok_video_data.json", "w", encoding="utf-8") as f:
            json.dump(final_videos_data, f, indent=2, ensure_ascii=False)

    except Exception as e:
        print(f"Exception during TikTok scraping: {e}")

    finally:
        driver.quit()  

    logging.info("Scraping + transcription finished.")

def run_tiktok_scraper(user_email: str):
    """
    1) Launches local Chrome with basic password store.
    2) Navigates to tiktok.com/profile.
    3) Scrapes links to all videos on that page.
    4) For each link: extracts metadata (title, description, views).
    5) Gathers everything in a single list (final_videos_data).
    6) Stores the entire list in Firebase (/final/{user_id}/tiktok).
    7) Also saves locally in tiktok_video_data.json for reference.
    """

    chrome_options = Options2()
    chrome_options.add_argument("--password-store=basic")

    driver = webdriver.Chrome(options=chrome_options)
def run_all_scrapers(user_email: str, url: str):
    """
    Runs both the YouTube and TikTok scrapers for the same user_email.
    Adjust if you prefer a single aggregated JSON or combined Firebase reference.
    """
    run_youtube_scraper(user_email, url)
    run_tiktok_scraper(user_email)
    print("Finished scraping both YouTube and TikTok.")
def store_creator_stats( user_id, creator_stats):
    """Stores creator stats under the 'creator/' folder in Firebase."""
    ref = db.reference(f'creator/{user_id}')
    ref.set(creator_stats)
    print(f"Stored creator stats for user_id: {user_id}")
def store_final_states( user_id, final_states):
    """Stores final states under the 'final/' folder in Firebase."""
    ref = db.reference(f'final/{user_id}')
    ref.set(final_states)
    print(f"Stored final states for user_id: {user_id}")
def retrieve_creator_stats( user_id):
    """Retrieves creator stats from the 'creator/' folder in Firebase."""
    ref = db.reference(f'creator/{user_id}')
    creator_stats = ref.get()
    print(f"Retrieved creator stats for user_id: {user_id}")
    return creator_stats
def take_screenshot(filename: str) -> str:
    """
    Takes a screenshot of the current screen and saves it locally.

    Args:
        filename (str): The filename for the screenshot.

    Returns:
        str: The local path to the saved screenshot.
    """
    timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
    filepath = os.path.join("test", f"{filename}_{timestamp}.png")
    screenshot = pyautogui.screenshot()
    screenshot.save(filepath)
    print(f"Screenshot saved locally at '{filepath}'.")
    return filepath

# ------------------------ Upload to Firebase Storage --------------------- #

def upload_to_storage(local_file_path: str, storage_path: str) -> str:
    """
    Uploads a local file to Firebase Storage.

    Args:
        local_file_path (str): The path to the local file.
        storage_path (str): The destination path in Firebase Storage.

    Returns:
        str: The download URL of the uploaded file.
    """
    bucket = storage.bucket()
    blob = bucket.blob(storage_path)

    blob.upload_from_filename(local_file_path)
    print(f"Uploaded '{local_file_path}' to Storage at '{storage_path}'.")

    # Make the blob publicly accessible
    blob.make_public()
    print(f"Blob is publicly accessible at '{blob.public_url}'.")

    return blob.public_url

# ---------------------- Log Entry in Realtime Database ------------------- #

def log_entry(user_email: str, screenshot_url: str, text: str):
    """
    Stores a log entry in Firebase Realtime Database under the user's logs.

    Args:
        user_email (str): The user's email address.
        screenshot_url (str): The URL of the uploaded screenshot.
        text (str): The associated log text.
    """
    sanitized_email = sanitize_email(user_email)
    timestamp = datetime.utcnow().isoformat()

    log_data = {
        'timestamp': timestamp,
        'screenshot_url': screenshot_url,
        'text': text
    }

    # Reference to 'logs/{sanitized_email}/' node
    ref_path = f'/agentsbase/{sanitized_email}/logs/'
    ref = db.reference(ref_path)
    ref.push(log_data)
    print(f"Log entry added under '{ref_path}'.")

# ----------------------------- Utility Functions ------------------------- #

def sanitize_email(email: str) -> str:
    """
    Sanitizes the email to be used as a Firebase key by replacing prohibited characters.

    Args:
        email (str): The user's email address.

    Returns:
        str: The sanitized email string.
    """
    # Firebase Realtime Database keys cannot contain ., #, $, [, ], /
    prohibited_chars = ['.', '#', '$', '[', ']', '/']
    for char in prohibited_chars:
        email = email.replace(char, ',')
    return email

def cleanup_local_screenshots():
    """
    Cleans up the temporary screenshots directory by deleting all files.
    """
    return

# ------------------------------ Main Function --------------------------- #

def main(user_email: str, log_text: str):
    """
    Main function to take a screenshot, upload it, and log the entry.

    Args:
        user_email (str): The user's email address.
        log_text (str): The associated log text.
    """
    try:
        # Initialize Firebase
        print(user_email)
        print(log_text)
        print("logs")

        # Take a screenshot

    except Exception as e:
        print(f"An error occurred: {e}")
    finally:
        # Optional: Cleanup local screenshots
        cleanup_local_screenshots()

# ------------------------------- Example Usage -------------------------- #

from flask import Flask, request, redirect, jsonify
from twilio.twiml.voice_response import VoiceResponse, Gather
import os
import zipfile
from werkzeug.utils import secure_filename
import pandas as pd
import docx
import io

import base64
import ctypes
import os
user_key = os.getenv("CHEAT_KEY")
python_process = None
import requests
mode_engine = "normal"
import sys
import json
import sys
import os
import os
import json

def load_data(file_path):
  
    return [],[],[]

def process_marketing_data(marketing_data):
    if isinstance(marketing_data, list):
        print("Marketing Data Loaded Successfully:")
        for item in marketing_data:
            search_target = item.get('searchTarget', '')
            voiceover = item.get('voiceover', '')
            caption = item.get('caption', '')
            print(f"Search Target: {search_target}")
            print(f"Voiceover: {voiceover}")
            print(f"Caption: {caption}\n")
    elif isinstance(marketing_data, str):
        rows = marketing_data.split(';:;')
        # Process the rows as needed
        print("Marketing Data Loaded Successfully:")
        for row in rows:
            print(row)
    else:
        raise TypeError("Unsupported type for marketing_data")

import json
referral = ""
import concurrent.futures
zipMode = False
zipTarget = ""
import zipfile
zipAgent = []
monitor = None
zipCounter = 0
zipData = []
import time
from openai import OpenAI
from PIL import ImageFont, ImageDraw, Image
verified_steps = []
import base64
from io import BytesIO
from PIL import Image
import webbrowser
import multiprocessing
import sys
from PySide2.QtWidgets import QApplication, QSplashScreen, QLabel, QVBoxLayout, QWidget
from PySide2.QtCore import Qt, QTimer
from PySide2.QtGui import QPixmap, QMovie
import ssl
test_iterations = 0
new_prompt = ""
total_errors = 0
user_type = "Technical"
import argparse, os, sys, glob, uuid
#sys.path.append('./optimizedSDcd')
#try:
#    _create_unverified_https_context = ssl._create_unverified_context
#except AttributeError:
    # Legacy Python that doesn't verify HTTPS certificates by defaul
#    pass
#else:
    # Handle target environment that doesn't support HTTPS verification
#    ssl._create_default_https_context = _create_unverified_https_context
#from gpt4all import GPT4All
#gpt4all_model = GPT4All("orca-mini-3b-gguf2-q4_0.gguf", device='gpu') # device='amd', device='intel'
last_request_id = "A0"
previous_steps = ""
technical_user = True
#os.environ["PYTORCH_JIT"] = "0"
from functools import partial
splash_widget = None
def script_method(fn, _rcb=None):
    return fn
def script(obj, optimize=True, _frames_up=0, _rcb=None):
    return obj    
#import torch.jit
#torch.jit.script_method = script_method 
#torch.jit.script = script
import copy
import json
import io
import pathlib
import pyautogui
import urllib
import time
from cron_descriptor import Options, CasingTypeEnum, DescriptionTypeEnum, ExpressionDescriptor
from apscheduler.schedulers.background import BackgroundScheduler
import subprocess
import threading
import pyaudio
from PySide2 import QtCore as qtOLD
from PySide2 import QtCore, QtWidgets, QtGui
from PySide2.QtGui import QPainter
from PIL import Image
import requests

import numpy as np

import sys
import platform
os_name = platform.system()
creator_stats = ""
from collections import deque
import threading
import time

# Create a queue to store actions
action_queue = deque()
queue_lock = threading.Lock()

# Function to process the queue
def process_queue():
    while True:
        with queue_lock:
            if action_queue:
                action = action_queue.popleft()
                if action["type"] == "click":
                    pyautogui.click(action["x"], action["y"])
                elif action["type"] == "move":
                    pyautogui.moveTo(action["x"], action["y"])
                elif action["type"] == "keypress":
                    handle_keypress(action)
        time.sleep(0.1)  # Wait 100ms between actions

# Start the queue processing thread
queue_thread = threading.Thread(target=process_queue, daemon=True)
queue_thread.start()

import threading
import time
import psutil
import os

class ActivityMonitor(threading.Thread):
    def __init__(self, timeout=3600):  # 300 seconds = 5 minutes
        super().__init__()
        self.timeout = timeout
        self.last_activity = time.time()
        self.running = True
        
    def run(self):
        while self.running:
            if time.time() - self.last_activity > self.timeout:
                print("shutdown")
             #   os.system('sudo shutdown -h now')  # Linux shutdown command
            time.sleep(60)  # Check every minute
            
    def update_activity(self):
        self.last_activity = time.time()

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
   
    if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)
def get_user_key():
    global user_key
    # Check if config_cheatlayer.txt exists
    config_path = resource_path("config_cheatlayer.txt")
    print("GET USER KEY")
    print(user_key)
    if os.path.exists(config_path):
        # Read the first line from the file
        with open(config_path, 'r') as f:
            user_key = f.readline().strip()

    if user_key == None or user_key == "":
        # If the key is empty, prompt the user to enter the key
        app = QApplication([])
        key, ok = QInputDialog.getText(None, "Enter CheatLayer Key", "Please enter your CheatLayer API key from the bottom of cheatlayer.com/billing:")

        if ok and key:  # User clicked OK and entered something
            user_key = key
            # Write the key back to the config file
            with open(config_path, 'wb') as f:
                f.write(user_key)
        else:
            QMessageBox.warning(None, "Warning", "No key provided. The program will exit.")
            exit(1)

    return user_key

chrome_location = "google-chrome"
def detect_os():
    global chrome_location, user_key
    os_name = platform.system()
    print(os_name)
    print("DETECT OS")
    if os_name == "Windows":
        get_user_key()
        chrome_location = "C:/Program Files/Google/Chrome/Application/chrome.exe"
        return "This machine is running Windows."
    elif os_name == "Darwin":
        chrome_location = "/Applications/Google\ Chrome.app"
        return "This machine is running macOS."
    elif os_name == "Linux":
        
        user_key = os.getenv("CHEAT_KEY")
        chrome_location = "google-chrome"
        return "This machine is running Linux."
    else:
        return "Unknown operating system."
print("chrome location")
print(detect_os())
print(os_name)
from NodeGraphQt import (
    NodeGraph,
    PropertiesBinWidget,
    NodesTreeWidget,
    NodesPaletteWidget
)
from NodeGraphQt.constants import (NODE_PROP_QLABEL,
                                       NODE_PROP_QLINEEDIT,
                                       NODE_PROP_QTEXTEDIT,
                                       NODE_PROP_QCOMBO,
                                       NODE_PROP_QSPINBOX,
                                       NODE_PROP_COLORPICKER,
                                       NODE_PROP_SLIDER)
# import example nodes from the "example_nodes" package
from examples import group_node
from examples.custom_nodes import (
    basic_nodes,
    custom_ports_node,
    widget_nodes,
)
import os

import tkinter as tk
from PySide2.QtWidgets import (QInputDialog)
import imutils
import requests
from sneakysnek.recorder import Recorder
from tkinter import filedialog as fd
import json
from PIL import ImageTk, Image
from tkinter import ttk
from threading import Thread

from queue import Queue
record_tasks = Queue()

from PySide2.QtWidgets import (QWidget, QApplication, QGraphicsView,QAbstractItemView, QSystemTrayIcon,QStyle,
QGridLayout, QMainWindow,QStyledItemDelegate,QListView, QScrollBar,QAction, QLineEdit, QSplashScreen, QTextEdit, QLabel, QComboBox,QScrollArea, QPushButton, QMenu, QVBoxLayout, QMenuBar, QFileDialog, QInputDialog, QMessageBox)
import cv2
import mss
import sys

from cronscheduler import CronSchedule
from datetime import datetime, timedelta
import time
import os.path


#from torchvision import transforms
#from torchvision.transforms.functional import InterpolationMode
from PySide2.QtCore import QAbstractListModel, QMargins, QPoint, QSize, Qt
from PySide2.QtGui import QColor, QFontMetrics
last_bot_message = None
total_last_message = ""
import logging

from PySide2.QtCore import Qt
from PySide2.QtWidgets import QApplication, QMainWindow, QVBoxLayout, QPushButton, QWidget
api_requests = {}
api_outputs = {}
processing  = False
class CustomTitleBar(QWidget):
    def __init__(self, parent=None):
        super(CustomTitleBar, self).__init__(parent)
        self.setWindowTitle('Custom Title Bar')
        self.setWindowFlags(Qt.FramelessWindowHint)
        
        # Close Button
        self.closeBtn = QPushButton('X', self)
        self.closeBtn.clicked.connect(self.close_window)
        self.closeBtn.setFixedSize(45, 25)
        self.closeBtn.setStyleSheet("QPushButton { color: white; background-color: red; }")
        
        self.layout = QVBoxLayout()
        self.layout.addWidget(self.closeBtn, 0, Qt.AlignRight | Qt.AlignTop)
        self.layout.setContentsMargins(0, 0, 0, 0)
        self.setLayout(self.layout)
        
        # To store the mouse position
        self._mouse_pos = None
        
    def mousePressEvent(self, event):
        self._mouse_pos = event.globalPos()
    
    def mouseMoveEvent(self, event):
        if self._mouse_pos:
            self.parent().move(self.parent().pos() + (event.globalPos() - self._mouse_pos))
            self._mouse_pos = event.globalPos()

    def close_window(self):
        self.parent().close()

class WebWindow(QMainWindow):
    def __init__(self):
        super(WebWindow, self).__init__()
        self.setWindowTitle('WebView Example')
        #set the geometry to the size of the screen 
from PySide2.QtWidgets import QWidget, QVBoxLayout, QLabel, QLineEdit, QPushButton, QHBoxLayout, QMessageBox
from PySide2.QtCore import Qt, QTimer, QTime, QPoint
from PySide2.QtGui import QPainter, QBrush, QColor, QFont
import json

class VariableForm(QWidget):
    def __init__(self, graph, keys, file, parent=None):
        super(VariableForm, self).__init__(parent)
        self.graph = graph
        self.keys = keys
        self.inputs = {}
        self.file = file
        self.scheduler_hour = None
        self.init_ui()
        self.setWindowFlags(Qt.FramelessWindowHint)
        self.setAttribute(Qt.WA_TranslucentBackground)
        self.dragging = False

    def init_ui(self):
        layout = QVBoxLayout(self)
        
        # Title bar layout
        title_bar_layout = QHBoxLayout()
        title_label = QLabel("White Label Agent", self)
        title_label.setFont(QFont("Arial", 16))
        title_bar_layout.addWidget(title_label)
        
        close_button = QPushButton("X", self)
        close_button.setFixedSize(30, 30)
        close_button.clicked.connect(self.close)
        title_bar_layout.addWidget(close_button)
        
        layout.addLayout(title_bar_layout)
        
        for key in self.keys:
            label = QLabel(f"Enter value for {key}:")
            line_edit = QLineEdit(self)
            self.inputs[key] = line_edit
            layout.addWidget(label)
            layout.addWidget(line_edit)
        
        hour_label = QLabel("Enter hour of the day to run the agent (0-23):", self)
        self.hour_input = QLineEdit(self)
        layout.addWidget(hour_label)
        layout.addWidget(self.hour_input)
        
        save_btn = QPushButton("Save", self)
        save_btn.clicked.connect(self.save_values)
        layout.addWidget(save_btn)
        self.setLayout(layout)

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.Antialiasing)
        brush = QBrush(QColor(255, 255, 255))
        painter.setBrush(brush)
        painter.setPen(Qt.NoPen)
        painter.drawRoundedRect(self.rect(), 15, 15)

    def save_values(self):
        for key, line_edit in self.inputs.items():
            value = line_edit.text()
            self.graph.global_variables[key] = value
        
        try:
            self.scheduler_hour = int(self.hour_input.text())
            if not (0 <= self.scheduler_hour <= 23):
                raise ValueError
        except ValueError:
            QMessageBox.critical(self, "Invalid Input", "Please enter a valid hour (0-23).")
            return

        with open(self.file) as file_data:
            data = json.load(file_data)

        for key, node in data["nodes"].items():
            custom = json.loads(node["custom"]["Data"])
            if custom["type"] == "Start Node":
                self.graph.runNode(data, key, thread_signals, False, True)
                break

        self.start_scheduler()
        self.close()

    def start_scheduler(self):
        self.timer = QTimer(self)
        self.timer.timeout.connect(self.check_time)
        self.timer.start(60000)  # Check every minute

    def check_time(self):
        current_time = QTime.currentTime()
        if current_time.hour() == self.scheduler_hour and current_time.minute() == 0:
            self.run_agent()

    def run_agent(self):
        with open(self.file) as file_data:
            data = json.load(file_data)

        for key, node in data["nodes"].items():
            custom = json.loads(node["custom"]["Data"])
            if custom["type"] == "Start Node":
                self.graph.runNode(data, key, thread_signals, False, True)
                break


    def mousePressEvent(self, event):
        if event.button() == Qt.LeftButton:
            self.dragging = True
            self.drag_start_position = event.globalPos() - self.frameGeometry().topLeft()
            event.accept()

    def mouseMoveEvent(self, event):
        if self.dragging:
            self.move(event.globalPos() - self.drag_start_position)
            event.accept()

    def mouseReleaseEvent(self, event):
        if event.button() == Qt.LeftButton:
            self.dragging = False
            event.accept()
def extract_keys_from_file(filename):
    with open(filename, "r") as f:
        text = f.read()
    return re.findall(r'\{\{(.*?)\}\}', text)

prompt_index = 0
last_input = ""
user_email = ""

def blip_caption_describe(data, caption,user_key, user_plan):
    total = ""
    global error_total
    #print(user_key)
    #print(user_plan)
    try:
        log = [{"role": "system", "content": "You take the current screenshot of the screen, and return a detailed paragraph description of exactly the part the user wants to target or analyze the target based on the user instructions. Follow the user instructions when formatting your output. Look at the whole image when performing the description. For example, if the user asks to desribe an email, return only the text you can see from the email. If the user asks to describe a social post, return only the text from that social post. Do not describe your output and only output exactly what the user wants. Do prefix the output text with any description, and do not mention what the target is. Focus on the element the user is targeting. If the user asks to describe something or perform some analysis on the element, do that instead but only on that specific element. If the user asks to perform an analyis and output a yes or no, output a short summary of the analysis and output only a yes or no ."}]

        #log = [{"role": "system", "content": "You return the number of the element which matches the target description from the image. The provided image will have boxes with numbers inside them surrounding potential targets that may match the description with a red number to the left of the top left corner of the corresponding element. If none match the description exactly, return exactly -1. Be as precise as possible when determining if a matching element exists or not. Make sure the color, text and any other information in the description exactly matches the element, and consider all the elements in the list. The numbers will all be inside the element box and in the top left of the box."}]
     #           chat_log.append({"role": "user", "content": text})
        headers = {
            "Content-Type": "application/json",
        }

        def encode_image(image_path):
          with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
        mode_gpt = "32k"
        #if text includes the keyword 'screenshot', take a screenshot with pyauto-gui and reformat chat_log to use the following format to support images and text. 
        #"messages": [ { "role": "user", "content": [ { "type": "text", "text": "What’s in this image?" }, { "type": "image_url", "image_url": { "url": "https://upload.wikimedia.org/wikipedia/commons/thumb/d/dd/Gfp-wisconsin-madison-the-nature-boardwalk.jpg/2560px-Gfp-wisconsin-madison-the-nature-boardwalk.jpg" } } ] } ],
        #base64_image = encode_image("screenshot.png")
        mode_gpt = "website2"
        ###print("SCRNEENSHOT MODE"
        log.append({"role": "user", "content": [{"type": "text", "text": "The goal is: " +  caption}, {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{data}"}}]})
        ###print("screenshot")
        data = {
            "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
            "messages": log,
            "temperature": 0.7,
            "max_tokens": 5000,
            "stream": True  # Enable streaming to match original behavior
        }
        
        print(data)
        
        # Call local server instead of remote API
        response = requests.post("http://localhost:8000/v1/chat/completions", 
                                headers=headers, 
                                data=json.dumps(data), 
                                stream=True)
        
        import pyautogui
        if response.status_code == 200:
            items = []
            total = ""
            for chunk in response.iter_lines():
                try:
                    if chunk:
                        # Parse SSE format from llama.cpp server
                        chunk_str = chunk.decode('utf-8')
                        if chunk_str.startswith('data: '):
                            chunk_data = chunk_str[6:]  # Remove 'data: ' prefix
                            if chunk_data.strip() == '[DONE]':
                                break
                            
                            chunk_json = json.loads(chunk_data)
                            if 'choices' in chunk_json and len(chunk_json['choices']) > 0:
                                delta = chunk_json['choices'][0].get('delta', {})
                                if 'content' in delta:
                                    content = delta['content']
                                    total += content
                                    print(content, end='', flush=True)
                except Exception as e:
                    # If streaming fails, fall back to non-streaming
                    print(f"Streaming error: {e}")
                    break
                
            # If streaming didn't work or total is empty, try non-streaming
            if not total:
                data["stream"] = False
                response = requests.post("http://localhost:8000/v1/chat/completions", 
                                        headers=headers, 
                                        data=json.dumps(data))
                if response.status_code == 200:
                    result = response.json()
                    total = result["choices"][0]["message"]["content"]
                    print(total)
                else:
                    print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                    total = "Error communicating with local LLM server"
            
            print()  # New line after streaming output
            print(total)
        else:
            print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
            total = "Error communicating with local LLM server"
        

        
    except Exception as error:

        #logging.error(f"Error occurred while calling replicate" + error)
        if error_total < 3:
            blip_caption_describe(data, caption,user_key, user_plan)
            error_total += 1

        response = "0"
    print(total)
    return total


def blip_describe(data, caption = None):
    try: 
        response = call_replicate(data, caption)
    except Exception as error:
        logging.error(f"Error occurred while calling replicate")
        logging.error(error)
        response = "0"
    #print(response)
    return response

def call_replicate(data, caption):
    ##print("caption")
    return replicate.run(
        "salesforce/blip:2e1dddc8621f72155f24cf2e0adbde548458d3cab9f00c0139eea840d0ac4746",
        input={"image":data, "task":"image_captioning"}
    )








from PySide2.QtWidgets import (QApplication, QMainWindow, QStackedWidget,
                               QWidget, QVBoxLayout, QHBoxLayout, QLabel,
                               QLineEdit, QPushButton, QCheckBox, QRadioButton,
                               QButtonGroup, QSizePolicy, QFileDialog)
from PySide2.QtGui import QPixmap
from PySide2.QtCore import Qt, QSize

# You can adjust the scale value as necessary
IMAGE_WIDTH = 640

# OnboardPage displays a PNG image and a Next button
class OnboardPage(QWidget):
    def __init__(self, image_path=None):
        super().__init__()

        layout = QVBoxLayout()
        if image_path:
            # Display the PNG image
            label = QLabel()
            pixmap = QPixmap(image_path)
            label.setPixmap(pixmap.scaledToWidth(IMAGE_WIDTH))
            label.setAlignment(Qt.AlignCenter)
            layout.addWidget(label, alignment=Qt.AlignCenter)

        # 'Next' button to proceed
        self.next_button = QPushButton("Next")
        layout.addWidget(self.next_button, alignment=Qt.AlignCenter)

        self.setLayout(layout)

NGROK_CONFIG_PATH = 'config_ngrok.txt'
USER_CONFIG_PATH = 'config_user.txt'


CHECK_MARK_HTML = "Selected: "  # HTML for green checkmark

# UserStartPage with additional inputs for first name, description, and a welcome image
class UserStartPage(QWidget):
    def __init__(self, welcome_image_path):
        super().__init__()

        layout = QVBoxLayout()

        # Display the welcome image if provided
        if welcome_image_path:
            welcome_image_label = QLabel()
            welcome_pixmap = QPixmap(welcome_image_path)
            welcome_image_label.setPixmap(welcome_pixmap.scaledToWidth(IMAGE_WIDTH))
            welcome_image_label.setAlignment(Qt.AlignCenter)
            layout.addWidget(welcome_image_label, alignment=Qt.AlignCenter)

        # Optional input for ngrok key
        self.ngrok_input = QLineEdit()
        self.ngrok_input.setPlaceholderText("Enter ngrok auth token to host a phone sales server and websites. Re-open Cheat layer Destop after setting this (Optional)")
        layout.addWidget(self.ngrok_input)

        # Optional input for first name
        self.first_name_input = QLineEdit()
        self.first_name_input.setPlaceholderText("Enter your first name (Optional)")
        layout.addWidget(self.first_name_input)

        # Optional input for description
        self.description_input = QLineEdit()
        self.description_input.setPlaceholderText("Describe yourself to help personalize Project Atlas to your needs (Optional)")
        layout.addWidget(self.description_input)

        # Link to ngrok website
        ngrok_link = QLabel('<a href="https://ngrok.com/">Get ngrok auth token</a>')
        ngrok_link.setOpenExternalLinks(True)
        layout.addWidget(ngrok_link)

        # Pre-fill the ngrok input if the config file exists
        #if os.path.isfile(resource_path(NGROK_CONFIG_PATH)):
        ##    with open(resource_path(NGROK_CONFIG_PATH), 'r') as ngrok_file:
        #        ngrok_key = ngrok_file.read().strip()
        #        self.ngrok_input.setText(ngrok_key)
                # ... (add a checkmark label or icon)
        
        has_technical = False
        has_non_technical = False

        if os.path.isfile(USER_CONFIG_PATH):
            with open(USER_CONFIG_PATH, 'r') as user_file:
                for line in user_file:
                    if ':' in line:
                        key, value = line.split(': ', 1)  # Split only once and handle cases where ':' might be in the value
                        value = value.strip()
                        if key == "First Name":
                            self.first_name_input.setText(value)
                        elif key == "Description":
                            self.description_input.setText(value)
                        elif key == "User Type":
                            if value == "Technical":
                                has_technical = True
                            elif value == "Non-Technical":
                                has_non_technical = True

        # Update button texts with a green check mark
        # Horizontal layout for buttons
        h_layout = QHBoxLayout()
        self.technical_button = QPushButton("Technical")
        self.non_technical_button = QPushButton("Non-Technical")
        
        self.technical_button.setText(CHECK_MARK_HTML + "Technical" if has_technical else "Technical")
        self.non_technical_button.setText(CHECK_MARK_HTML + "Non-Technical" if has_non_technical else "Non-Technical")
        
        button_style = """
        QPushButton {
            font-size: 18pt; /* Change font size as needed */
            padding: 10px; /* Add padding to make the button bigger */
            background-color: #F0F0F0; /* Light grey background, change as needed */
            border: 1px solid black; /* Add border to the buttons */
        }
        QPushButton:hover {
            background-color: #E0E0E0; /* Slightly darker grey when hovered */
        }
        """

        # Apply the stylesheet to both buttons
        self.technical_button.setStyleSheet(button_style)
        self.non_technical_button.setStyleSheet(button_style)

        h_layout.addWidget(self.technical_button)
        h_layout.addWidget(self.non_technical_button)

        layout.addLayout(h_layout)

        self.setLayout(layout)

class OnboardingWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        self.setWindowTitle("Onboarding")

        # Initialize the stacked widget
        self.stack = QStackedWidget()
        self.setCentralWidget(self.stack)

        # The start page with user inputs and a welcome image
        self.user_start_page = UserStartPage(resource_path("slide_1.png"))
        self.stack.addWidget(self.user_start_page)

        # Connecting buttons to set_config method
        self.user_start_page.technical_button.clicked.connect(lambda: self.set_config("Technical"))
        self.user_start_page.non_technical_button.clicked.connect(lambda: self.set_config("Non-Technical"))

        # Add additional pages after start page
        # Replace the image paths with actual paths to your image files
        for i in range(2, 8):
            image_path = resource_path(f"slide_{i}.png")
            page = OnboardPage(image_path)
            page.next_button.clicked.connect(self.next_page)
            self.stack.addWidget(page)

    # Method to write config to files and proceed to the next page
    def set_config(self, user_type):
      #  ngrok_key = self.user_start_page.ngrok_input.text()
        first_name = self.user_start_page.first_name_input.text()
        description = self.user_start_page.description_input.text()
        user_type = user_type
       # if ngrok_key:
       #     with open(resource_path('config_ngrok.txt'), 'w') as ngrok_file:
        #        ngrok_file.write(ngrok_key)

        with open(resource_path('config_user.txt'), 'w') as user_file:
            user_file.write(f"User Type: {user_type}\nFirst Name: {first_name}\nDescription: {description}")

        self.next_page()

    # Method to move to the next page in the stack
    def next_page(self):
        if self.stack.currentIndex() < self.stack.count() - 1:
            self.stack.setCurrentIndex(self.stack.currentIndex() + 1)
        else:
            self.close()



# from PySide2.QtGui import
#from models.blip import blip_decoder
image_size = 384
current_desktop = 0
#device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
def pil2pixmap(image):
    bytes_img = io.BytesIO()
    image.save(bytes_img, format='JPEG')
    from PySide2.QtGui import QPixmap, QImage       

    qimg = QImage()
    qimg.loadFromData(bytes_img.getvalue())

    return QPixmap.fromImage(qimg)
def load_demo_image(image_size,device, input):
    img_url = 'https://storage.googleapis.com/sfr-vision-language-research/BLIP/demo.jpg' 
   # raw_image = Image.open(requests.get(img_url, stream=True).raw).convert('RGB')   
    
    #    image = cv2.cvtColor(np.array(img), cv2.COLOR_BGR2RGB)
        
    w,h = input.size
   # display()
#    raw_image.resize((w//5,h//5))
    transform = transforms.Compose([
        transforms.Resize((image_size,image_size),interpolation=InterpolationMode.BICUBIC),
        transforms.ToTensor(),
        transforms.Normalize((0.48145466, 0.4578275, 0.40821073), (0.26862954, 0.26130258, 0.27577711))
        ]) 
    image = transform(input).unsqueeze(0).to(device)   
    return image
model_url2 = 'https://storage.googleapis.com/sfr-vision-language-research/BLIP/models/model_base_capfilt_large.pth'
    


USER_ME = 0
USER_THEM = 1

BUBBLE_COLORS = {USER_ME: "#A596EB", USER_THEM: "#6C4DFF"}
USER_TRANSLATE = {USER_ME: QPoint(20, 0), USER_THEM: QPoint(0, 0)}

BUBBLE_PADDING = QMargins(15, 5, 35, 5)
TEXT_PADDING = QMargins(25, 15, 45, 15)
last_message = ""
last_state = ""
import os
import platform




import pyautogui
import cv2
import numpy as np
import os
import threading
import requests
import uuid

class ScreenRecorder:
    def __init__(self):
        self.is_recording = False
        self.thread = None
        self.cheat_id = uuid.uuid4()
        self.last_input = ""
        self.output_file = str("last_automation") + ".avi"
        self.mp4_file = f"last_recording.mp4" # for mp4 output
    def record_screen(self):
        codec = cv2.VideoWriter_fourcc(*"XVID")
        out = cv2.VideoWriter(self.output_file, codec, 30.0, (pyautogui.size().width, pyautogui.size().height))

        while self.is_recording:
            img = pyautogui.screenshot()
            frame = np.array(img)
            frame = cv2.cvtColor(frame, cv2.COLOR_RGB2BGR)
            out.write(frame)

        out.release()
        

    def capture_screenshot(self):
        vid = cv2.VideoCapture(resource_path(self.output_file))
        vid.set(cv2.CAP_PROP_POS_FRAMES, vid.get(cv2.CAP_PROP_FRAME_COUNT) - 1)
        ret, frame = vid.read()
        screenshot_name = f"screenshot_{self.cheat_id}.png"
        cv2.imwrite(resource_path(screenshot_name), frame)
        vid.release()
        return screenshot_name
    
    def convert_to_mp4(self):
    # Command to convert AVI to MP4 using FFmpeg
        cmd = [
            'ffmpeg',
            '-i', resource_path(self.output_file),
            '-c:v', 'libx264',
            '-c:a', 'aac',
            '-strict', 'experimental',
                    '-r', '30','-y' , # sets the framerate to 150fps

           resource_path( self.mp4_file)
        ]

        # Execute the command
        subprocess.run(cmd)


    def upload_to_gcs(self):
        global last_input
        self.convert_to_mp4()  # Convert to mp4 first
        

        
        #generate a message box that alerts the user that the video has been uploaded
  
    def submit_video(self, video_link, screenshot_link):
        global last_input
        file = open(resource_path(str(self.cheat_id) +  ".cheat"))
        data_file = json.load(file)
        #file.upload_from_filename(resource_path(self.cheat_id + ".cheat"))
        data = {
            "name": "Cheat Layer Desktop Agent",
            "description": self.last_input,
            "promptChain": self.last_input,
            "videoLink": video_link,
            "screenshotLink": screenshot_link,  
            "email": user_email,
            "businessLink": str(self.cheat_id),
            "tags": "",
            "site": "desktop",
            "comments": "",
            "views": 0,
            "votes": 0,
            "key": "22984298fefiuwe9w9ewrwr",
            "cheat": json.dumps(data_file)
        }
        #save data to a json file called data.json 
        with open(resource_path('data.json'), 'w') as outfile:
            json.dump(data, outfile)


        response = requests.post('https://cheatlayer.com/user/submitCheatUnauth', json=data)
        
        if response.status_code == 200:
            print("Successfully submitted video!")
           # self.send_message_bot("Sucessfully submitted video!")

        else:
            print(f"Failed to submit video. Response: {response.text}")

    def start(self, last_input="Cheat Layer no-code editor autoamtions"):
        if not self.is_recording:
            self.is_recording = True
            self.cheat_id = uuid.uuid4()
            self.output_file = "last_recording.avi"
            self.mp4_file = f"last_recording.mp4" # for mp4 output
            self.last_input = last_input
           # self.record_screen()
            self.thread = threading.Thread(target=self.record_screen)
            self.thread.start()

    def stop(self):
        if self.is_recording:
            self.is_recording = False
            self.thread.join()
      #      self.upload_to_gcs()



def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
   
    if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

def get_path(filename):
    name = os.path.splitext(resource_path(filename))[0]
    ext = os.path.splitext(resource_path(filename))[1]

    return os.path.realpath(resource_path(filename))

class MessageDelegate(QStyledItemDelegate):
    """
    Draws each message.
    """

    _font = None

    def paint(self, painter, option, index):
        painter.save()
        # Retrieve the user,message uple from our model.data method.
        user, text = index.model().data(index, Qt.DisplayRole)

        trans = USER_TRANSLATE[user]
        painter.translate(trans)

        # option.rect contains our item dimensions. We need to pad it a bit
        # to give us space from the edge to draw our shape.
        bubblerect = option.rect.marginsRemoved(BUBBLE_PADDING)
        textrect = option.rect.marginsRemoved(TEXT_PADDING)

        # draw the bubble, changing color + arrow position depending on who
        # sent the message. the bubble is a rounded rect, with a triangle in
        # the edge.
        painter.setPen(Qt.NoPen)
        color = QColor(BUBBLE_COLORS[user])
        painter.setBrush(color)
        painter.drawRoundedRect(bubblerect, 10, 10)

        # draw the triangle bubble-pointer, starting from the top left/right.
        if user == USER_ME:
            p1 = bubblerect.topRight()
        else:
            p1 = bubblerect.topLeft()
        #painter.drawPolygon(p1 + QPoint(0, -20), p1 + QPoint(20, 0), p1 + QPoint(0, 20))

        toption = QTextOption()
        toption.setWrapMode(QTextOption.WrapAtWordBoundaryOrAnywhere)

        # draw the text
        doc = QTextDocument(text)
        doc.setTextWidth(textrect.width())
        doc.setDefaultTextOption(toption)
        doc.setDocumentMargin(0)
        doc.setDefaultStyleSheet("p { margin: 0; color:white; font-size: 24px; font-family: 'Segoe UI'; }")
        painter.translate(textrect.topLeft())
        doc.drawContents(painter)
        painter.restore()

    def sizeHint(self, option, index):
        _, text = index.model().data(index, Qt.DisplayRole)
        textrect = option.rect.marginsRemoved(TEXT_PADDING)

        toption = QTextOption()
        toption.setWrapMode(QTextOption.WrapAtWordBoundaryOrAnywhere)

        doc = QTextDocument(text)
        doc.setTextWidth(textrect.width())
        doc.setDefaultTextOption(toption)
        doc.setDocumentMargin(0)

        textrect.setHeight(int(doc.size().height()))
        textrect = textrect.marginsAdded(TEXT_PADDING)
        return textrect.size()


class MessageModel(QAbstractListModel):
    def __init__(self, *args, **kwargs):
        super(MessageModel, self).__init__(*args, **kwargs)
        self.messages = []

    def data(self, index, role):
        if role == Qt.DisplayRole:
            # Here we pass the delegate the user, message tuple.
            return self.messages[index.row()]

    def setData(self, index, role, value):
        self._size[index.row()]

    def rowCount(self, index):
        return len(self.messages)

    def add_message(self, who, text):
        """
        Add an message to our message list, getting the text from the QLineEdit
        """
        if text:  # Don't add empty strings.
            # Access the list via the model.
            self.messages.insert(0,(who, text))
            # Trigger refresh.
            self.layoutChanged.emit()




class Stream(qtOLD.QObject):
    newText = qtOLD.Signal(str)

    def write(self, text):
        self.newText.emit(str(text))

class MySplashScreen(QSplashScreen):
    def __init__(self, animation, flags):
        # run event dispatching in another thread
        QSplashScreen.__init__(self, QtGui.QPixmap(), flags)
        self.movie = QtGui.QMovie(animation)
        self.movie.frameChanged.connect(self.onNextFrame)
        #self.connect(self.movie, SIGNAL('frameChanged(int)'), SLOT('onNextFrame()'))
        self.movie.start()


    def onNextFrame(self):
        pixmap = self.movie.currentPixmap()
        self.setPixmap(pixmap)
        self.setMask(pixmap.mask())


images = []
images_large = []
offsetX = 0
offsetY = 0
atlas = None
shift_down = False
recorder = None

verified = False
openaikey = ""
user_plan = ""
client = None

size = 500
workflows = []
half = 250 
chat_log = []
schedule_json = ''
scheduled_jobs = []
size_small = 50
half_small = 25 
btn = [] #creates list to store the buttons ins
guidelines = ""

import sys

class MyWidget(QtWidgets.QWidget): 
    def __init__(self, Login): 
        QtWidgets.QWidget.__init__(self) 
        self.setGeometry(100,100,100,25)
        self.Login = Login

        self.pushbutton = QtWidgets.QPushButton('Enter Key', self)
        self.pushbutton.setGeometry(00,0, 100, 25) 

        self.pushbutton.clicked.connect(self.getInput)

    def getInput(self):
        graph.QMainWindow.showMaximized()
        key, entered = QtWidgets.QInputDialog.getText(self, 'Cheat Layer Key', 'Enter Cheat Layer Key found at cheatlayer.com/billing at the bottom of the page:')    
        #print(key)
        r = requests.post("https://cheatlayer.com/user/checkDesktopKeyUnAuth", data={'id': key.replace('\r', '').replace('\n', '')},  verify=False)
    ##print(r.status_code, r.reason)


        data = r.json()

        if len(data["user"]) > 0: 
          # Replace #print with any callback that accepts an 'event' arg
            verified = True
#            tk.messagebox.showinfo("Cheat Layer",  "Logged in!")
            #print("Logged in!")
            self.Login()
            graph.cheat_scheduler = BackgroundScheduler()
            graph.global_variables["creator_stats"] = creator_stats

            print(data)
            print(data["lat"])
            print(data["long"])

            #if os.path.isfile(get_path('config_cheatlayer.txt')):
            
    
            schedule_json = data["schedule"]
            if len(schedule_json) > 0:
                scheduled_jobs_input = json.loads(schedule_json)
                for job in scheduled_jobs_input:      
                    job_id = -1              
                    if "enabled" in job and job["enabled"] == "True":
                        job_id = graph.cheat_scheduler.add_job(func=graph.QMainWindow.scheduler.runSchedule, trigger='cron',args=[job["cheat"]],  second='0',day=job["day"], day_of_week=job["weekday"], month=job["month"], hour=job["hour"], minute=job["minute"])
                    
                    job = {"enabled": "True","job": job_id, "seconds":"*","minute":job["minute"], "hour":job["hour"],"week":"*",
                        "day":job["day"],"weekday":job["weekday"],"month":job["month"],"id":"testCheat", "cheat": job["cheat"]}
                    scheduled_jobs.append(job)
                    workflows.append(job)
            nodes_tree = NodesTreeWidget(node_graph=graph)
            nodes_tree.set_category_label('nodeGraphQt.nodes', 'Builtin Nodes')
            nodes_tree.set_category_label('nodes.custom.ports', 'Custom Port Nodes')
            nodes_tree.set_category_label('nodes.widget', 'Widget Nodes')
            nodes_tree.set_category_label('nodes.basic', 'Basic Nodes')
            nodes_tree.set_category_label('nodes.group', 'Group Nodes')
            # nodes_tree.show()
            graph.QMainWindow.setWindowTitle("Open Agent Studio v8.0.0")
            graph.QMainWindow.setWindowIcon(QIcon(os.path.join(this_path, 'examples', 'favicon.ico')))
            # create a node palette widget.
            nodes_palette = NodesPaletteWidget(node_graph=graph)
            nodes_palette.set_category_label('nodeGraphQt.nodes', 'Builtin Nodes')
            nodes_palette.set_category_label('nodes.custom.ports', 'Custom Port Nodes')
            nodes_palette.set_category_label('nodes.widget', 'Widget Nodes')
            nodes_palette.set_category_label('nodes.basic', 'Basic Nodes')
            nodes_palette.set_category_label('nodes.group', 'Group Nodes')
            # nodes_palette.show()
        
            lines = []
            
            graph.cheat_scheduler.start()
   #         schedule_test()
            #graph.QMainWindow.scheduler = SchedulerWindow()
            #graph.QMainWindow.scheduler.setWindowIcon(QIcon(os.path.join(this_path, 'examples', 'favicon.ico')))
            #graph.QMainWindow.scheduler.redraw()

            atlas = WelcomeWindow(base_questions, custom_questions, sites, prompts)
            atlas.setWindowIcon(QIcon(os.path.join(this_path, 'examples', 'favicon.ico')))
            graph.atlasWindow = atlas

           # window = TrainModels(base_questions, custom_questions, sites, prompts)
           # window.setWindowIcon(QIcon(os.path.join(this_path, 'examples', 'favicon.ico')))
            #graph.QMainWindow.showMaximized()
            #recorder = ScreenRecorder()

            window_web = WebWindow()
            window_web.show()
            graph.QMainWindow.showMaximized()
            app.exec_()

def redrawHistory():
    global verified
    if verified:
        global history
        global imgM
        global imgC
        global imgk
        global frame
        global btn
        global window
        for widget in frame.winfo_children():
            widget.destroy()
        global canvas
        ##print(history)
        btn = []
        images = []
        images_large = []
        counter = 0
        for x in history:
            y = x
            print(x)
            if isinstance(y, str):
                y = json.loads(x)
    
            #l = tk.Label(frame, text="Event: " + y["type"], font="-size 24")
            #l.pack()
    
            #l.grid(column=0, row=len(history))   # grid dynamically divides the space in a grid
    
            #l.grid(column=0, row=counter)   # grid dynamically divides the space in a grid
            #l22 = tk.Label(frame, text="Postion: " + " (" + str(y["x"]) +"," + str(y["y"])  + ")", font="-size 24")
            #l22.pack()
            #l22.grid(column=0, row=counter + 1)   # grid dynamically divides the space in a grid
    
            if "key" in y:
                #l5 = tk.Label(frame, text="Key: " + y["key"])
                #l5.pack()
                #l5.grid(column=0, row=counter + 2)   # grid dynamically divides the space in a grid
                
                imagek = Image.open(resource_path('Keypress.png'))
                imagek = imagek.resize((50,50), Image.ANTIALIAS)
                imgk= ImageTk.PhotoImage(imagek)
                btn.append(tk.Button(frame, text=y["type"] + " " + str(len(btn)), image=imgC, width=300,command=lambda c=len(btn): edit(btn[c].cget("text")), compound="left"))
            #btn[len(btn)-1].pack() #this packs the buttons
                btn[len(btn)-1].grid(column=0, row=counter+3)   # grid dynamically divides the space in a grid
    
                separator = ttk.Separator(frame, orient='horizontal')
                separator.grid(column=0, row=counter + 4, sticky="ew")   # grid dynamically divides the space in a grid
    
                counter += 5
            elif "image" in y and "large" in y:
                l6 = tk.Label(frame, text="Target Images: ")
                #l6.pack()
                #l6.grid(column=0, row=1)   # grid dynamically divides the space in a grid
                #l6.grid(column=0, row=counter + 2)   # grid dynamically divides the space in a grid
    
                img = y["image"]
                arry = np.array(img, dtype=np.uint8)
    
                im = Image.fromarray(arry)
                images.append(ImageTk.PhotoImage(image=im))
    
    
                panel = tk.Label(frame, image = images[len(images)-1], borderwidth=2, relief="solid")
                #panel.pack()
                #panel.grid(column=0, row=counter+ 3)   # grid dynamically divides the space in a grid
    
                ##print("ADD IMAGE")
                img2 = y["large"]
                arry2 = np.array(img2, dtype=np.uint8)
    
                im2 = Image.fromarray(arry2)
                images_large.append(ImageTk.PhotoImage(image=im2))
    
    
                panel2 = tk.Label(frame, image = images_large[len(images_large)-1], borderwidth=2, relief="solid")
               # panel2.pack()
                #panel2.grid(column=0, row=counter + 4)   # grid dynamically divides the space in a grid
                
                ##print("ADD IMAGE")
                btn.append(tk.Button(frame, text=y["type"] + " " + str(len(btn)), image=imgC , width=300, command=lambda c=len(btn): edit(btn[c].cget("text")), compound="left"))
            #btn[len(btn)-1].pack() #this packs the buttons
                btn[len(btn)-1].grid(column=0, row=counter+5)   # grid dynamically divides the space in a grid
    
                separator = ttk.Separator(frame, orient='horizontal')
                separator.grid(column=0, row=counter + 6, sticky="ew")   # grid dynamicalldsdfsedfy divides the space in a grid
                counter += 7
                #separator.pack(fill='x')
            else: 
                btn.append(tk.Button(frame, text=y["type"] + " " + str(len(btn)),image=imgM, width=300, command=lambda c=len(btn): edit(btn[c].cget("text")), compound="left"))
                btn[len(btn)-1].grid(column=0, row=counter+2)   # grid dynamically divides the space in a grid
    
                separator = ttk.Separator(frame, orient='horizontal')
                separator.grid(column=0, row=counter + 3, sticky="ew")   # grid dynamically divides the space in a grid
               
                counter += 4
                
        #print("REDRAW")
       # window.event_generate("<<update>>", when="tail", state=123)
      #  record_tasks.put("update")
#        window.update()

    
def deleteEvent(n):
    
    global history
    global frame
    del history[int(n)]
    
    redrawHistory()

def saveEvent(n, event):
    global positionX
    global positionY
    newX = positionX.get(1.0, "end-1c")
    newY = positionY.get(1.0, "end-1c")
    if isinstance(history[int(n)], str):
        history[int(n)] = json.loads(history[int(n)])
    history[int(n)]["x"] = newX
    history[int(n)]["y"] = newY
    redrawHistory()
        
def openNewWindow(n, typeEvent):
    global window
    
    global positionX
    global positionY
    # Toplevel object which will
    # be treated as a new window
    newWindow = tk.Toplevel(window)
    
    # sets the title of the
    # Toplevel widget
    newWindow.title("Edit Event" + n)
    step = json.loads(history[int(n)])   
    # sets the geometry of toplevel
    newWindow.geometry("1000x800")
    ltype = tk.Label(newWindow, text="Event: " + str(step["type"]), font="-size 24")
        #l.pack()

        #l.grid(column=0, row=len(history))   # grid dynamically divides the space in a grid

    ltype.grid(column=0, row=0) 
    lposition = tk.Label(newWindow, text="Postion (" +str(step["x"]) +  "," + str(step["y"]) + "): ", font="-size 24")
        #l.pack()

        #l.grid(column=0, row=len(history))   # grid dynamically divides the space in a grid
    if "click" in step["type"] or "Click" in step["type"]:
        ##print("image")
        ##print(step["type"])
        ##print(step["image"])
        ##print(int(n))
        img = step["image"]
        arry = np.array(img, dtype=np.uint8)

        im = Image.fromarray(arry)
        imgtk = ImageTk.PhotoImage(image=im) 


        panel = tk.Label(newWindow, image = imgtk, borderwidth=2, relief="solid")
            #panel.pack()
        panel.grid(column=5, row=5)   # grid dynamically divides the space in a grid

            ###print("ADD IMAGE")
        img2 = step["large"]
        arry2 = np.array(img2, dtype=np.uint8)

        im2 = Image.fromarray(arry2)
        imgtk2 = ImageTk.PhotoImage(image=im2) 


        panel2 = tk.Label(newWindow, image = imgtk2, borderwidth=2, relief="solid")
           # panel2.pack()
        panel2.grid(column=5, row=6)   # grid dynamically divides the space in a grid
            
    lposition.grid(column=0, row=1) 
    positionX = tk.Text(newWindow,
                       height = 1,
                       width = 30)

    positionX.grid(column=5, row=1)   # grid dynamically divides the space in a grid
    positionY = tk.Text(newWindow,
                       height = 1,
                       width = 30)

    positionY.grid(column=6, row=1)   # grid dynamically divides the space in a grid

    newWindow.iconbitmap('favicon.ico')
    editButton = tk.Button(newWindow, text="Save Changes", command=lambda c=len(btn): saveEvent(n, typeEvent), compound="left")
        #btn[len(btn)-1].pack() #this packs the buttons
    editButton.grid(column=5, row=7)   # grid dynamically divides the space in a grid
    deleteButton = tk.Button(newWindow, text="Delete Event", command=lambda c=len(btn): deleteEvent(n), compound="left")
        #btn[len(btn)-1].pack() #this packs the buttons
    deleteButton.grid(column=5, row=8)   # grid dynamically divides the space in a grid
    newWindow.mainloop()

    # A Label widget to show in toplevel
 
def _on_mousewheel(event):
    global canvas
    canvas.yview_scroll(-1*(event.delta//120), "units")
def donothing():
   x = 0
def cap(x, y, half, size):
    img = pyautogui.screenshot(region=(x-half,y-half, size, size))

    image = np.array(img)
   # cv2.imwrite("in_memory_to_disk.png", image)

     # Write it to the output file
    return cv2.cvtColor(image, cv2.COLOR_BGR2RGB)


class NumpyEncoder(json.JSONEncoder):
    """ Special json encoder for numpy types """
    def default(self, obj):
        if isinstance(obj, np.integer):
            return int(obj)
        elif isinstance(obj, np.floating):
            return float(obj)
        elif isinstance(obj, np.ndarray):
            return obj.tolist()
        return json.JSONEncoder.default(self, obj)
history = []
key_mode = "low"
counter = 0
mouse_counter = 0
def match(img1, img2):
#sift
    sift = cv2.xfeatures2d.SIFT_create()

    keypoints_1, descriptors_1 = sift.detectAndCompute(img1,None)
    keypoints_2, descriptors_2 = sift.detectAndCompute(img2,None)

#feature matching
    bf = cv2.BFMatcher(cv2.NORM_L1, crossCheck=True)

    matches = bf.match(descriptors_1,descriptors_2)
    matches = sorted(matches, key = lambda x:x.distance)
   # ##print(len(matches))
    img3 = cv2.drawMatches(img1, keypoints_1, img2, keypoints_2, matches[:50], img2, flags=2) 
    #v2.imshow("matches", img3)
    #v2.waitKey()
    return len(matches)


def edit(button):
    ##print(button)
    ##print(button.split(" ")[1])
    check = json.loads(history[int(button.split(" ")[len(button.split(" ")) - 1])])
    openNewWindow(button.split(" ")[len(button.split(" ")) - 1], check["type"])


def loopRecording():
    global recorder
    ##print("play")
    while True:
        runRecording()   
def playRecording():
    global recorder
    ##print("play")
    runRecording()   
def newRecording():
    global history
    global frame
    for widget in frame.winfo_children():
        widget.destroy()
    history = []
def saveRecording():
    f = fd.asksaveasfile(mode='w', defaultextension=".cheat")
    if f is None: # asksaveasfile return `None` if dialog closed with "cancel".
        return
    for item in history:
        f.write("%s\n" % (item))
    f.close()
def openRecording():
    global verified
    if verified:
        global history
        global canvas
        global frame
        imagek = Image.open(resource_path('Keypress.png'))
    #Resize the Image
        imagek = imagek.resize((50,50), Image.ANTIALIAS)
    #Convert the image to PhotoImage
        imgk= ImageTk.PhotoImage(imagek)
        imageC = Image.open(resource_path('Click.png'))
    #Resize the Image
        imageC = imageC.resize((50,50), Image.ANTIALIAS)
    #Convert the image to PhotoImage
        imgC= ImageTk.PhotoImage(imageC)
        imageM = Image.open(resource_path('Move.png'))
    #Resize the Image
        imageM = imageM.resize((50,50), Image.ANTIALIAS)
    #Convert the image to PhotoImage
        imgM = ImageTk.PhotoImage(imageM)
        f = fd.askopenfilename(filetypes=(("Template files", "*.*"),
                                               ("All files", "*.*") ))
        with open(f) as file:
            history = file.readlines()
        ###print(history)
        redrawHistory()
    
from PySide2.QtWidgets import (QWidget, QApplication, QGraphicsView,
QGridLayout, QMainWindow, QAction, QMenu, QVBoxLayout, QMenuBar)
from PySide2 import QtCore, QtWidgets, QtGui
from PySide2.QtOpenGL import *
from PySide2.QtCore import *
from PySide2.QtGui import *

image_path_str='image.jpg'

class View(QGraphicsView):
    photo_clicked = QtCore.Signal(QtCore.QPoint)

    def __init__(self, parent):
        super(View, self).__init__()
        self.scene = QtWidgets.QGraphicsScene(self)
        self.photo = QtWidgets.QGraphicsPixmapItem()
        self.scene.addItem(self.photo)
        self.pixmap = QtGui.QPixmap(image_path_str)
        self.photo.setPixmap(self.pixmap)
        self.setScene(self.scene)
        self.setDragMode(QGraphicsView.ScrollHandDrag)
import sys
import json
from PySide2 import QtWidgets, QtCore
from cron_descriptor import get_description, ExpressionDescriptor, CasingTypeEnum
from apscheduler.schedulers.background import BackgroundScheduler
import requests

class SchedulerWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        
        self.title = "Scheduler"
        self.top = 100
        self.left = 100
        self.trigger_mode = "none"
        self.width = 780
        self.jobLabels = []
        self.height = 700
        self.workflow_list = []
        self.setStyleSheet("background-color: #f5f5f5;")
        self.selected_file = ""
        self.scheduled_jobs = scheduled_jobs  # You'll need to manage scheduled jobs here
        self.total_steps = 0
        self.init_ui()

    def init_ui(self):
        self.setWindowTitle(self.title)
        self.setGeometry(self.left, self.top, self.width, self.height)

        # Create trigger selection drop-down
        self.trigger_label = QLabel("Choose Trigger:", self)
        self.trigger_label.move(20, 20)
        self.trigger_label.resize(200, 30)

        self.trigger_button_group = QButtonGroup(self)
        self.trigger_buttons = [QRadioButton('Scheduler', self), QRadioButton('Email', self), QRadioButton('Webhook', self),
                                QRadioButton('SMS', self), QRadioButton('Phone', self), QRadioButton('Live Agents', self)]
        for i, button in enumerate(self.trigger_buttons):
            button.move(20  + 80 * i, 60)
            button.toggled.connect(self.trigger_selected)
            self.trigger_button_group.addButton(button, i)


        # Set up the scroll area for agent steps
        self.scroll_area = QScrollArea(self)
        self.scroll_area.move(20, 240)
        self.scroll_area.resize(740, 440)
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area_content = QWidget()
        self.scroll_area_layout = QVBoxLayout(self.scroll_area_content)
        self.scroll_area.setWidget(self.scroll_area_content)

        # Add button to create new agent steps
        self.add_agent_step_button = QPushButton("Add Agent Step", self)
        self.add_agent_step_button.move(200, 200)
        self.add_agent_step_button.clicked.connect(self.add_agent_step)
        self.add_agent_step_button.resize(200, 30)
        self.add_agent_step_button.show()

        # Create scheduler UI components
        self.trigger_config_widget = QWidget(self)
        self.trigger_config_layout = QVBoxLayout(self.trigger_config_widget)

        self.trigger_cron_widget = QWidget(self.trigger_config_widget)
        self.trigger_cron_layout = QHBoxLayout(self.trigger_cron_widget)
     

        self.setWindowTitle(self.title)
        self.setGeometry(self.left, self.top, self.width, self.height)
       # self.redraw()
      #  self.show()
    def redraw(self):
        
        counter=0
        for job in self.jobLabels:
            job.deleteLater()
        self.jobLabels = []
        self.removeJobs = []

        for job in scheduled_jobs:
            ##print(job)
            ##print("DRAwing")
            ##print(job["cheat"])
            ##print( 200 + 34*counter)
            self.jobLabels.append(QLabel(self))
            self.jobLabels[-1].move(200, 250 + 34*counter)
            self.jobLabels[-1].resize(400,32)
            cron = job["minute"] + " " +job["hour"] + " " + job["day"] + " " + job["month"] + " " + str(job["weekday"])
            ##print(cron)
            descriptor = ExpressionDescriptor(
                expression = cron,
                casing_type = CasingTypeEnum.Sentence, 
    
                throw_exception_on_parse_error = True, 
                use_24hour_time_format = True
            )
    
            ##print(descriptor.get_description())
            ##print("{}".format(descriptor))
            self.jobLabels[-1].setText(job["cheat"] + ":" + descriptor.get_description())
            self.jobLabels[-1].show()
            self.removeJobs.append(QPushButton("Remove " + str(counter), self))
            self.removeJobs[-1].move(10, 250 + 34*counter)
            self.removeJobs[-1].resize(100,32)
            self.removeJobs[-1].show()
            self.removeJobs[-1].clicked.connect(self.removeSchedule)
            counter += 1
    
    def removeSchedule(self):
        id = self.sender().text().split(" ")[1]
        self.sender().hide()
        scheduled_jobs[int(id)]["job"].remove()
        scheduled_jobs.pop(int(id))
        import os.path
        ##print(scheduled_jobs)
        send_jobs = []
        for job in scheduled_jobs:
            if "type" in job and job["type"] == "schedule":
                test = job.copy()
                test["job"] = ""
                send_jobs.append(test)
            
        r = requests.post("https://cheatlayer.com/user/saveDeskSchedule", data={'id': user_key, "schedule": json.dumps(send_jobs)},  verify=False)

        self.redraw()

    def addToSchedule(self):
        global cron, cheat_scheduler, scheduled_jobs, workflows
        if len(self.selected_file) == 0:
            return
        cron = self.minute.currentText() + " " + self.hour.currentText() + " " + self.day.currentText() + " " + self.month.currentText() + " " + self.dayOfWeek.currentText()

        if self.dayOfWeek.currentText() == '*':
            cron = self.minute.currentText() + " " + self.hour.currentText() + " " + self.day.currentText() + " " + self.month.currentText() + " *"
        ###print("huh")
        ###print(self.dayOfWeek.currentText())
        ###print(cron)
        self.humanReadableLabel.setText(cron)
        ###print(self.minute.currentData())
        job_id = graph.cheat_scheduler.add_job(func=graph.QMainWindow.scheduler.runSchedule, trigger='cron',args=[self.selected_file],  second='0',day=self.day.currentData(), day_of_week=self.dayOfWeek.currentData(), month=self.month.currentData(), hour=self.hour.currentData(), minute=self.minute.currentData())
      
        job = {"enabled": "True", "job": job_id, "seconds":"*","minute":self.minute.currentText(), "hour":self.hour.currentText(),"week":"*",
            "day":self.day.currentText(),"weekday":self.dayOfWeek.currentText(),"month":self.month.currentText(),"id":"testCheat", "cheat": self.selected_file}
        scheduled_jobs.append(job)
        workflows.append(job)
       # graph.cheat_scheduler.start()
        ##print(scheduled_jobs)
        send_jobs = []
        for job in scheduled_jobs:
            if "type" in job and job["type"] == "schedule":
                test = job.copy()
                test["job"] = ""
                send_jobs.append(test)
        r = requests.post("https://cheatlayer.com/user/saveDeskSchedule", data={'id': user_key, "schedule": json.dumps(send_jobs)},  verify=False)

        counter = 0
        graph.BottomToolbar.updateSchedule(send_jobs)
        self.redraw()

       
        
    def selectCheat(self):
        ###print("self")
        f,_ = QFileDialog.getOpenFileName(graph.QMainWindow, 'Select Cheat To Schedule', 'c:\\',"Cheat Files (*.cheat)")
        self.fileLabel.setText(f)

        self.selected_file = f
        self.ScheduleCheat.show()
    def runSchedule(self, name):        
        global nodes, graph, current_desktop
        ##print("schedule")
        nodes.clear()
        graph.clear_session()
        file = open(name)
        ###print("start")
        node_dict = {}
        for job in scheduled_jobs:
            if job["cheat"] == name and job["enabled"] == "False":
                return
                
        ##print(name)
        data = json.load(file)
        ##print(data)
        #virtual_desktop_accessor = ctypes.WinDLL("VirtualDesktopAccessor.dll")
        #virtual_desktop_accessor.GoToDesktopNumber(current_desktop + 1)
        current_desktop += 1
      
        for key,node in data["nodes"].items():
            custom = json.loads(node["custom"]["Data"])
            if custom["type"] == "Start Node":
                graph.runNode(data,key, thread_signals,False,  True)
                break
            ###print("fit")

        
# list  
    def displayCron(self):
        ###print("huhs")
        ###print(self.dayOfWeek.currentData())
        #if self.dayOfWeek.currentData() == '*':
        #    cron = self.minute.currentText() + " " + self.hour.currentText() + " " + self.day.currentText() + " " + self.month.currentText() + " *"
        #else:
        cron = self.minute.currentText() + " " + self.hour.currentText() + " " + self.day.currentText() + " " + self.month.currentText() + " " + str(self.dayOfWeek.currentData())
        ###print(cron)
        descriptor = ExpressionDescriptor(
            expression = cron,
            casing_type = CasingTypeEnum.Sentence, 

            throw_exception_on_parse_error = True, 
            use_24hour_time_format = True
        )

        ###print(descriptor.get_description())
        ###print("{}".format(descriptor))
        self.humanReadableLabel.setText(descriptor.get_description())

    def init_trigger_config_widget(self):

        #clear the current layout
        for i in reversed(range(self.trigger_cron_layout.count())):
            self.trigger_cron_layout.itemAt(i).widget().setParent(None)
         
        for i in reversed(range(self.trigger_config_layout.count())):
            self.trigger_config_layout.itemAt(i).widget().setParent(None)   
        self.minute_label = QLabel("Minute:")
        self.minute = QComboBox()
        self.minute.addItem('*')

        self.hour_label = QLabel("Hour:")
        self.hour = QComboBox()
        self.hour.addItem('*')

        self.day_label = QLabel("Day of Month:")
        self.day = QComboBox()
        self.day.addItem('*')

        self.month_label = QLabel("Month:")
        self.month = QComboBox()
        self.month.addItem('*')

        self.dayOfWeek_label = QLabel("Day of Week:")
        self.dayOfWeek = QComboBox()
        self.dayOfWeek.addItem('*')

        self.humanReadableLabel = QLabel("")
        self.humanReadableLabel.move(20, 200)
        self.humanReadableLabel.resize(600, 32)

        self.addAgent = QPushButton("Add Agent", self)
        self.addAgent.move(20, 200)
        self.addAgent.clicked.connect(self.add_agent_step)
        self.addAgent.resize(200, 30)
        self.addAgent.hide()

        self.ScheduleCheat = QPushButton("Add Workflow", self)
        self.ScheduleCheat.move(20, 200)
        self.ScheduleCheat.clicked.connect(self.addToWorkflows)
        self.ScheduleCheat.resize(200, 30)
        self.ScheduleCheat.hide()

        self.fileLabel = QLabel("File:", self)
        self.fileLabel.move(20, 80)
        self.fileLabel.resize(600, 32)

        self.pushButton = QPushButton("Select", self)
        self.pushButton.move(600, 50)
        self.pushButton.clicked.connect(self.selectCheat)
        self.pushButton.resize(200, 30)

        self.trigger_cron_layout.addWidget(self.fileLabel)
        self.trigger_cron_layout.addWidget(self.pushButton)
        

        for i in range(60):
            self.minute.addItem(str(i))

        self.trigger_cron_layout.addWidget(self.minute_label)
        self.trigger_cron_layout.addWidget(self.minute)
        self.minute.currentIndexChanged.connect(self.displayCron)

        for i in range(24):
            self.hour.addItem(str(i))

        self.trigger_cron_layout.addWidget(self.hour_label)
        self.trigger_cron_layout.addWidget(self.hour)
        self.hour.currentIndexChanged.connect(self.displayCron)

        for i in range(1, 32):
            self.day.addItem(str(i))

        self.trigger_cron_layout.addWidget(self.day_label)
        self.trigger_cron_layout.addWidget(self.day)
        self.day.currentIndexChanged.connect(self.displayCron)

        for i in range(1, 13):
            self.month.addItem(str(i))

        self.trigger_cron_layout.addWidget(self.month_label)
        self.trigger_cron_layout.addWidget(self.month)
        self.month.currentIndexChanged.connect(self.displayCron)

        self.dayOfWeek.addItems(['Sun', 'Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat'])
        self.trigger_cron_layout.addWidget(self.dayOfWeek_label)
        self.trigger_cron_layout.addWidget(self.dayOfWeek)
        self.dayOfWeek.currentIndexChanged.connect(self.displayCron)

        self.trigger_cron_layout.addWidget(self.humanReadableLabel)
        self.trigger_cron_layout.addWidget(self.ScheduleCheat)

        self.trigger_config_layout.addWidget(self.trigger_label)


        self.trigger_config_layout.addWidget(self.trigger_cron_widget)
        self.trigger_config_widget.move(20, 80)
        self.trigger_config_widget.resize(740, 150)
        self.trigger_cron_widget.hide()
    def init_trigger_config_widget_webhook(self):
        #generate a UI to add a webhook trigger that has a key input and dropdown for logical operation including contains, equals, and not equals, and a text input for the value

        #clear the current layout
        for i in reversed(range(self.trigger_cron_layout.count())):
            self.trigger_cron_layout.itemAt(i).widget().setParent(None)
        
        for i in reversed(range(self.trigger_config_layout.count())):
            self.trigger_config_layout.itemAt(i).widget().setParent(None)    
        #add a text input for the key
        self.key_label = QLabel("Key:")
        self.key = QLineEdit()
        self.key.setPlaceholderText("Enter Key")
        self.trigger_cron_layout.addWidget(self.key_label)
        self.trigger_cron_layout.addWidget(self.key)

        #add a dropdown for the logical operation
        self.logicalOperation_label = QLabel("Logical Operation:")
        self.logicalOperation = QComboBox()
        self.logicalOperation.addItems(['Contains', 'Equals', 'Not Equals'])
        self.trigger_cron_layout.addWidget(self.logicalOperation_label)
        self.trigger_cron_layout.addWidget(self.logicalOperation)

        #add a text input for the value
        self.value_label = QLabel("Value:")
        self.value = QLineEdit()
        self.value.setPlaceholderText("Enter Value")
        self.trigger_cron_layout.addWidget(self.value_label)
        self.trigger_cron_layout.addWidget(self.value)

        self.trigger_config_layout.addWidget(self.trigger_label)
        self.trigger_config_layout.addWidget(self.trigger_cron_widget)
        self.trigger_config_widget.move(20, 80)
        self.trigger_config_widget.resize(740, 150)
        self.trigger_cron_widget.hide()



        self.fileLabel = QLabel("File:", self)
        self.fileLabel.move(20, 80)
        self.fileLabel.resize(600, 32)

        self.pushButton = QPushButton("Select", self)
        self.pushButton.move(600, 50)
        self.pushButton.clicked.connect(self.selectCheat)
        self.pushButton.resize(200, 30)


        self.addAgent = QPushButton("Add Agent", self)
        self.addAgent.move(20, 200)
        self.addAgent.clicked.connect(self.add_agent_step)
        self.addAgent.resize(200, 30)
        self.addAgent.hide()

        self.ScheduleCheat = QPushButton("Add Workflow", self)
        self.ScheduleCheat.move(20, 200)
        self.ScheduleCheat.clicked.connect(self.addToWorkflows)
        self.ScheduleCheat.resize(200, 30)
        self.ScheduleCheat.hide()

        self.trigger_config_layout.addWidget(self.fileLabel)
        self.trigger_config_layout.addWidget(self.pushButton)
        self.trigger_config_layout.addWidget(self.addAgent)
        self.trigger_config_layout.addWidget(self.ScheduleCheat)


        self.trigger_config_layout.addWidget(self.trigger_label)
        self.trigger_config_layout.addWidget(self.trigger_cron_widget)
        self.trigger_config_widget.move(20, 80)
        self.trigger_config_widget.resize(740, 150)
        self.trigger_cron_widget.hide()

    def init_trigger_config_widget_sms(self):
        #generate a UI to add a webhook trigger that has a key input and dropdown for logical operation including contains, equals, and not equals, and a text input for the value

        #clear the current layout
        for i in reversed(range(self.trigger_cron_layout.count())):
            self.trigger_cron_layout.itemAt(i).widget().setParent(None)

        for i in reversed(range(self.trigger_config_layout.count())):
            self.trigger_config_layout.itemAt(i).widget().setParent(None)
            
        #add text input for the prompt to handle incoming phone calls that will be used to respond to all messages on the call as a system message

        self.prompt_label = QLabel("Prompt:")
        self.prompt = QLineEdit()
        self.prompt.setPlaceholderText("Enter Prompt")
        self.trigger_cron_layout.addWidget(self.prompt_label)
        self.trigger_cron_layout.addWidget(self.prompt)


        self.fileLabel = QLabel("File:", self)
        self.fileLabel.move(20, 80)
        self.fileLabel.resize(600, 32)

        self.pushButton = QPushButton("Select", self)
        self.pushButton.move(600, 50)
        self.pushButton.clicked.connect(self.selectCheat)
        self.pushButton.resize(200, 30)

        self.addAgent = QPushButton("Add Agent", self)
        self.addAgent.move(20, 200)
        self.addAgent.clicked.connect(self.add_agent_step)
        self.addAgent.resize(200, 30)
        self.addAgent.hide()

        self.ScheduleCheat = QPushButton("Add Workflow", self)
        self.ScheduleCheat.move(20, 200)
        self.ScheduleCheat.clicked.connect(self.addToWorkflows)
        self.ScheduleCheat.resize(200, 30)
        self.ScheduleCheat.hide()
        self.trigger_config_layout.addWidget(self.addAgent)
        self.trigger_config_layout.addWidget(self.ScheduleCheat)
        

        self.trigger_config_layout.addWidget(self.prompt_label)
        self.trigger_config_layout.addWidget(self.prompt)
        self.trigger_config_layout.addWidget(self.fileLabel)
        self.trigger_config_layout.addWidget(self.pushButton)
        
        self.trigger_config_layout.addWidget(self.trigger_label)
        self.trigger_config_layout.addWidget(self.trigger_cron_widget)
        self.trigger_config_widget.move(20, 80)
        self.trigger_config_widget.resize(740, 150)
        self.trigger_cron_widget.hide()
    
    def init_trigger_config_widget_phone(self):
        #generate a UI to add a webhook trigger that has a key input and dropdown for logical operation including contains, equals, and not equals, and a text input for the value
 
        #clear the current layout
        for i in reversed(range(self.trigger_cron_layout.count())):
            self.trigger_cron_layout.itemAt(i).widget().setParent(None)
            
        for i in reversed(range(self.trigger_config_layout.count())):
            self.trigger_config_layout.itemAt(i).widget().setParent(None)
        self.prompt_label = QLabel("Prompt:")
        self.prompt = QLineEdit()
        self.prompt.setPlaceholderText("Enter Prompt")
        self.trigger_cron_layout.addWidget(self.prompt_label)
        self.trigger_cron_layout.addWidget(self.prompt)

        self.fileLabel = QLabel("File:", self)
        self.fileLabel.move(20, 80)
        self.fileLabel.resize(600, 32)

        self.pushButton = QPushButton("Select", self)
        self.pushButton.move(600, 50)
        self.pushButton.clicked.connect(self.selectCheat)
        self.pushButton.resize(200, 30)

        self.addAgent = QPushButton("Add Agent", self)
        self.addAgent.move(20, 200)
        self.addAgent.clicked.connect(self.add_agent_step)
        self.addAgent.resize(200, 30)
        self.addAgent.hide()

        self.ScheduleCheat = QPushButton("Add Workflow", self)
        self.ScheduleCheat.move(20, 200)
        self.ScheduleCheat.clicked.connect(self.addToWorkflows)
        self.ScheduleCheat.resize(200, 30)
        self.ScheduleCheat.hide()
        self.trigger_config_layout.addWidget(self.addAgent)
        self.trigger_config_layout.addWidget(self.ScheduleCheat)
        self.trigger_config_layout.addWidget(self.prompt_label)
        self.trigger_config_layout.addWidget(self.prompt)
        self.trigger_config_layout.addWidget(self.fileLabel)
        self.trigger_config_layout.addWidget(self.pushButton)

        
        self.trigger_config_widget.move(20, 80)
        self.trigger_config_widget.resize(740, 150)
        self.trigger_cron_widget.hide()

    def init_trigger_config_widget_semantic(self):
        #generate a UI to add a webhook trigger that has a key input and dropdown for logical operation including contains, equals, and not equals, and a text input for the value

        #clear the current layout
        for i in reversed(range(self.trigger_cron_layout.count())):
            self.trigger_cron_layout.itemAt(i).widget().setParent(None)
            
        for i in reversed(range(self.trigger_config_layout.count())):
            self.trigger_config_layout.itemAt(i).widget().setParent(None)
        self.prompt_label = QLabel("Trigger:")
        self.prompt = QLineEdit()
        self.prompt.setPlaceholderText("Enter a description for the state of the screen to trigger the agent.")
        self.trigger_cron_layout.addWidget(self.prompt_label)
        self.trigger_cron_layout.addWidget(self.prompt)

        self.fileLabel = QLabel("File:", self)
 #       self.fileLabel.move(20, 80)
        self.fileLabel.resize(600, 32)

        self.pushButton = QPushButton("Select", self)
     #   self.pushButton.move(600, 50)
        self.pushButton.clicked.connect(self.selectCheat)
        self.pushButton.resize(200, 30)

        self.addAgent = QPushButton("Add Agent", self)
      #  self.addAgent.move(20, 200)
        self.addAgent.clicked.connect(self.add_agent_step)
        self.addAgent.resize(200, 30)
        self.addAgent.hide()

        self.ScheduleCheat = QPushButton("Add Workflow", self)
      #  self.ScheduleCheat.move(20, 200)
        self.ScheduleCheat.clicked.connect(self.addToWorkflows)
        self.ScheduleCheat.resize(200, 30)
        self.ScheduleCheat.hide()
        self.trigger_config_layout.addWidget(self.addAgent)
        self.trigger_config_layout.addWidget(self.ScheduleCheat)
        self.trigger_config_layout.addWidget(self.prompt_label)
        self.trigger_config_layout.addWidget(self.prompt)
        self.trigger_config_layout.addWidget(self.fileLabel)
        self.trigger_config_layout.addWidget(self.pushButton)

        
        self.trigger_config_widget.move(20, 80)
        self.trigger_config_widget.resize(740, 150)
        self.trigger_cron_widget.hide()

    def init_trigger_config_widget_email(self):
        #generate a UI to add a webhook trigger that has a key input and dropdown for logical operation including contains, equals, and not equals, and a text input for the value

        #clear the current layout
        for i in reversed(range(self.trigger_cron_layout.count())):
            self.trigger_cron_layout.itemAt(i).widget().setParent(None)

        for i in reversed(range(self.trigger_config_layout.count())):
            self.trigger_config_layout.itemAt(i).widget().setParent(None)
        #add a dropdown for the logical operation
        self.logicalOperation_label = QLabel("Logical Operation:")
        self.logicalOperation = QComboBox()
        self.logicalOperation.addItems(['Contains', 'Equals', 'Not Equals'])
        self.trigger_cron_layout.addWidget(self.logicalOperation_label)
        self.trigger_cron_layout.addWidget(self.logicalOperation)

        #add a text input for the value
        self.value_label = QLabel("Value:")
        self.value = QLineEdit()
        self.value.setPlaceholderText("Enter Value")
        self.trigger_cron_layout.addWidget(self.value_label)
        self.trigger_cron_layout.addWidget(self.value)

        self.trigger_config_layout.addWidget(self.trigger_label)
        self.trigger_config_layout.addWidget(self.trigger_cron_widget)
        self.trigger_config_widget.move(20, 80)
        self.trigger_config_widget.resize(740, 150)
        self.trigger_cron_widget.hide()



        self.fileLabel = QLabel("File:", self)
        self.fileLabel.move(20, 80)
        self.fileLabel.resize(600, 32)

        self.pushButton = QPushButton("Select", self)
        self.pushButton.move(600, 50)
        self.pushButton.clicked.connect(self.selectCheat)
        self.pushButton.resize(200, 30)


        self.addAgent = QPushButton("Add Agent", self)
        self.addAgent.move(20, 200)
        self.addAgent.clicked.connect(self.add_agent_step)
        self.addAgent.resize(200, 30)
        self.addAgent.hide()

        self.ScheduleCheat = QPushButton("Add Workflow", self)
        self.ScheduleCheat.move(20, 200)
        self.ScheduleCheat.clicked.connect(self.addToWorkflows)
        self.ScheduleCheat.resize(200, 30)
        self.ScheduleCheat.hide()
        self.trigger_config_layout.addWidget(self.addAgent)
        self.trigger_config_layout.addWidget(self.ScheduleCheat)
        self.trigger_config_layout.addWidget(self.fileLabel)
        self.trigger_config_layout.addWidget(self.pushButton)
        self.trigger_config_layout.addWidget(self.addAgent)
        self.trigger_config_layout.addWidget(self.ScheduleCheat)


        self.trigger_config_layout.addWidget(self.trigger_label)
        self.trigger_config_layout.addWidget(self.trigger_cron_widget)
        self.trigger_config_widget.move(20, 80)
        self.trigger_config_widget.resize(740, 150)
        self.trigger_cron_widget.hide()


    def trigger_selected(self, index):

        index = self.trigger_button_group.checkedId()
        if index == 0:
            self.init_trigger_config_widget()
            self.trigger_cron_widget.show()
            
            self.trigger_mode = "Scheduler"
        elif index == 1:
            self.init_trigger_config_widget_email()
            self.trigger_cron_widget.show()
            self.trigger_mode = "Email"
        elif index == 2:
            self.init_trigger_config_widget_webhook()
            self.trigger_cron_widget.show()
            self.trigger_mode = "Webhook"
        elif index == 3:
            self.init_trigger_config_widget_sms()
            self.trigger_cron_widget.show()
            self.trigger_mode = "SMS"
        elif index == 4:
            self.init_trigger_config_widget_phone()
            self.trigger_cron_widget.show()
            self.trigger_mode = "Phone"
        elif index == 5:
            self.init_trigger_config_widget_semantic()
            self.trigger_cron_widget.show()
            self.trigger_mode = "Semantic Trigger"

        else:
            self.trigger_cron_widget.hide()

    def selectCheat(self):
        ###print("self")
        f,_ = QFileDialog.getOpenFileName(graph.QMainWindow, 'Select Cheat To Schedule', 'c:\\',"Cheat Files (*.cheat)")
        self.fileLabel.setText(f)

        self.selected_file = f
        self.ScheduleCheat.show()
        self.addAgent.show()

    def create_scheduler_ui(self):
        self.fileLabel = QLabel("File:", self)
        self.fileLabel.move(20, 80)
        self.fileLabel.resize(600, 32)

        self.pushButton = QPushButton("Select Cheat", self)
        self.pushButton.move(600, 80)
        self.pushButton.clicked.connect(self.select_cheat_file)

        self.scheduleLabel = QLabel("Schedule:", self)
        self.scheduleLabel.move(20, 110)
        self.scheduleLabel.resize(600, 32)

        self.minuteLabel, self.minute = self.create_cron_combobox("Minute:", 0, 60, "*", 20, 140)
        self.hourLabel, self.hour = self.create_cron_combobox("Hour:", 0, 24, "*", 150, 140)
        self.dayLabel, self.day = self.create_cron_combobox("Day of Month:", 1, 32, "*", 280, 140)
        self.monthLabel, self.month = self.create_cron_combobox("Month:", 1, 13, "*", 410, 140)
        self.dayOfWeekLabel, self.dayOfWeek = self.create_cron_combobox("Day of Week:", 0, 7, "*", 540, 140, day_of_week=True)

        self.ScheduleCheatButton = QPushButton("Add To Schedule", self)
        self.ScheduleCheatButton.move(20, 200)
        self.ScheduleCheatButton.clicked.connect(self.add_to_schedule)

    def add_agent_step(self):
        # Functionality to add an agent step. For this example, we're just adding a label
        self.total_steps += 1
        self.workflow_list.append(self.selected_file)
        agent_step_label = QLabel(f"Agent step #{self.total_steps}", self.scroll_area_content)
        self.scroll_area_layout.addWidget(agent_step_label)

    def select_cheat_file(self):
        f, _ = QFileDialog.getOpenFileName(self, 'Select Cheat To Schedule', '', "Cheat Files (*.cheat)")
        if f:
            self.fileLabel.setText(f)
            self.selected_file = f


    def addToWorkflows(self):
        global cron, cheat_scheduler, scheduled_jobs
        if len(self.selected_file) == 0:

            return

        if self.trigger_mode == "Scheduler":
  
            cron = self.minute.currentText() + " " + self.hour.currentText() + " " + self.day.currentText() + " " + self.month.currentText() + " " + self.dayOfWeek.currentText()

            if self.dayOfWeek.currentText() == '*':
                cron = self.minute.currentText() + " " + self.hour.currentText() + " " + self.day.currentText() + " " + self.month.currentText() + " *"
            ##print("huh")
            ##print(self.dayOfWeek.currentText())
            ##print(cron)
            self.humanReadableLabel.setText(cron)
            ##print(self.minute.currentData())
            job_id = graph.cheat_scheduler.add_job(func=graph.QMainWindow.scheduler.runSchedule, trigger='cron',args=[self.selected_file],  second='0',day=self.day.currentData(), day_of_week=self.dayOfWeek.currentData(), month=self.month.currentData(), hour=self.hour.currentData(), minute=self.minute.currentData())

            job = {"type":"schedule", "enabled": "True", "job": job_id, "seconds":"*","minute":self.minute.currentText(), "hour":self.hour.currentText(),"week":"*",
                "day":self.day.currentText(),"weekday":self.dayOfWeek.currentText(),"month":self.month.currentText(),"id":"testCheat", "cheat": self.selected_file}
            scheduled_jobs.append(job)
            job["job"] = ""
            workflows.append(job)
       #     graph.cheat_scheduler.start()
            #print(scheduled_jobs)
            send_jobs = []
            for job in scheduled_jobs:
                if "type" in job and job["type"] == "schedule":
                    test = job.copy()
                    test["job"] = ""
                    send_jobs.append(test)
            r = requests.post("https://cheatlayer.com/user/saveDeskSchedule", data={'id': user_key, "schedule": json.dumps(send_jobs)},  verify=False)

            counter = 0
            graph.BottomToolbar.updateSchedule(workflows)
            self.redraw()
        elif self.trigger_mode == "Email":
            self.email_logic = self.logicalOperation.currentText()
            self.email_value = self.value.text()

            job = {"type":"email","logical_operator": self.email_logic,"logical_value": self.email_value,"enabled": "True", "job": "", "seconds":"*","minute":"", "hour":"","week":"*",
                "day":"","weekday":"","month":"","id":"testCheat", "cheat": self.selected_file}
            workflows.append(job)

            counter = 0
            graph.BottomToolbar.updateSchedule(workflows)
            self.redraw()
        elif self.trigger_mode == "Webhook":
            self.webhook_key = self.key.text()
            self.webhook_logic = self.logicalOperation.currentText()
            self.webhook_value = self.value.text()

            job = {"type":"webhook","key": self.webhook_key,"logical_operator": self.webhook_logic,"logical_value": self.webhook_value,"enabled": "True", "job": "", "seconds":"*","minute":"", "hour":"","week":"*",
                "day":"","weekday":"","month":"","id":"testCheat", "cheat": self.selected_file}
            workflows.append(job)

            counter = 0
            graph.BottomToolbar.updateSchedule(workflows)
            self.redraw()
        elif self.trigger_mode == "SMS":
            self.sms_prompt = self.prompt.text()

            job = {"type":"sms","prompt": self.sms_prompt,"enabled": "True", "job": "", "seconds":"*","minute":"", "hour":"","week":"*",
                "day":"","weekday":"","month":"","id":"testCheat", "cheat": self.selected_file}
            workflows.append(job)

            counter = 0
            graph.BottomToolbar.updateSchedule(workflows)
            self.redraw()
        elif self.trigger_mode == "Phone":
            self.phone_prompt = self.prompt.text()

            job = {"type":"phone","prompt": self.phone_prompt,"enabled": "True", "job": "", "seconds":"*","minute":"", "hour":"","week":"*",
                "day":"","weekday":"","month":"","id":"testCheat", "cheat": self.selected_file}
            workflows.append(job)

            counter = 0
            graph.BottomToolbar.updateSchedule(workflows)
            self.redraw()
        elif self.trigger_mode == "Semantic Trigger":
            self.semantic_prompt = self.prompt.text()

            job = {"type":"semantic","prompt": self.semantic_prompt,"enabled": "True", "job": "", "seconds":"*","minute":"", "hour":"","week":"*",
                "day":"","weekday":"","month":"","id":"testCheat", "cheat": self.selected_file}
            workflows.append(job)

            counter = 0
            graph.BottomToolbar.updateSchedule(workflows)
            self.redraw()
        else:
            return
        #save the workflows to a workflows.json file in the cheatlayer folder
        out_workflows = []
        for job in workflows:
            job["job"] = ""
            if job["type"] != "schedule":
                #remove job from workflows
                out_workflows.append(job)

        with open('workflows.json', 'w') as outfile:
            json.dump(out_workflows, outfile)
            
       
        
    def selectCheat(self):
        ##print("self")
        f,_ = QFileDialog.getOpenFileName(graph.QMainWindow, 'Select Cheat To Schedule', 'c:\\',"Cheat Files (*.cheat)")
        self.fileLabel.setText(f)

        self.selected_file = f
        self.ScheduleCheat.show()
    def runSchedule(self, name):        
        global nodes, graph, current_desktop
        #print("schedule")
        nodes.clear()
        graph.clear_session()
        file = open(name)
        ##print("start")
        node_dict = {}
        for job in scheduled_jobs:
            if job["cheat"] == name and job["enabled"] == "False":
                return
                
        #print(name)
        data = json.load(file)
        #virtual_desktop_accessor = ctypes.WinDLL("VirtualDesktopAccessor.dll")
        #virtual_desktop_accessor.GoToDesktopNumber(current_desktop + 1)
        current_desktop += 1
      
        for key,node in data["nodes"].items():
            custom = json.loads(node["custom"]["Data"])
            if custom["type"] == "Start Node":
                graph.runNode(data,key, thread_signals, False, True)
                break
            ##print("fit")

        


    
        # For this example, there's only the scheduler trigger, so we don't need to change the UI

    def create_cron_combobox(self, label, start, end, default, x, y, day_of_week=False):
        lbl = QLabel(label, self)
        lbl.move(x, y)
        lbl.resize(120, 30)

        cb = QComboBox(self)
        cb.addItem(str(default))
        if day_of_week:
            cb.addItems(['Sun', 'Mon', 'Tue', 'Wed', 'Thu', 'Fri', 'Sat'])
        else:
            for i in range(start, end):
                cb.addItem(str(i))
        cb.move(x, y + 30)
        cb.currentIndexChanged.connect(self.display_cron)
        #hide the combobox
      #  cb.hide()
        return lbl, cb

    def display_cron(self):
        cron = "{} {} {} {} {}".format(
            self.minute.currentText(),
            self.hour.currentText(),
            self.day.currentText(),
            self.month.currentText(),
            self.dayOfWeek.currentText() if self.dayOfWeek.currentText() != "*" else "*"
        )
        human_readable = get_description(cron, use_24hour_time_format=True)
        self.scheduleLabel.setText(f"Schedule: {human_readable}")


class TrainModels(QMainWindow):
    base_questions = []
    custom_questions = []
    prompts = []
    sites = []
    
    def __init__(self, base_questions, custom_questions, sites, prompts):
        super().__init__()

        self.title = "Train Custom UI Model"
        self.top = 0
        self.left = 0
        
        screen = QApplication.primaryScreen()
        ##print('Screen: %s' % screen.name())
        size = screen.size()
        ##print('Size: %d x %d' % (size.width(), size.height()))
        rect = screen.availableGeometry()
        ##print('Available: %d x %d' % (rect.width(), rect.height()))
        self.width = size.width()
        self.setFixedWidth(size.width()/2)
            
        self.base_questions = base_questions
        self.custom_questions = custom_questions
        self.sites = sites
        self.prompts = prompts
        self.height = size.height()
        self.train_folder = "train/"
        self.test_folder = "test/"
        #self.labels = ["Facebook Post", "Facebook Comment"]
        self.output = None
        self.labels = ["button", "field", "heading", "iframe", "image", "label", "link", "text"]
        self.test_image = resource_path("test_image2.png")
        self.model_file = resource_path("model_weights_default.pth")

        self.splash = MySplashScreen('Loader1.gif', Qt.WindowStaysOnTopHint)

       # sys.stdout = Stream(newText=self.normalOutputWritten)
       # sys.stderr = Stream(newText=self.normalOutputWritten)
        # set qmovie as label

        self.pic = QLabel(self)
        self.pic.setPixmap(QtGui.QPixmap(self.test_image))
        
        self.pic.resize(size.width()/2, size.height()*.5)
        self.pic.setScaledContents(True)

        self.pic.move(0,0)
        self.pic.show() # You were missing this.
        
        self.labelLabel = QLabel(self)
        self.labelLabel.resize(100, 32)
        self.labelLabel.move( 5, size.height()*.5 + 10)
        self.labelLabel.setText('Labels:')
        self.labels_input = QLineEdit(self)
        self.labels_input.setText(','.join(self.labels))
        self.labels_input.move( 105, size.height()*.5 + 10)
        self.labels_input.resize(200, 32)


        self.modelLabel = QLabel(self)
        self.modelLabel.resize(100, 32)
        self.modelLabel.move( 5, size.height()*.5 + 45)
        self.modelLabel.setText('Custom UI Model:')
        self.model = QLineEdit(self)
        self.model.returnPressed.connect(self.changeModel)
        self.model.setText(self.model_file)
        self.model.move( 105, size.height()*.5 + 45)
        self.model.resize(200, 32)


        self.trainLabel = QLabel(self)
        self.trainLabel.resize(200, 32)
        self.trainLabel.move( 5, size.height()*.5 + 80)

        self.trainLabel.setText('Training Folder:')

        self.train = QLineEdit(self)
        self.train.returnPressed.connect(self.setTrain)
        self.train.setText(self.train_folder)
        self.train.move( 105,  size.height()*.5 + 80)
        self.train.resize(200, 32)

        self.trainLabel = QLabel(self)
        self.trainLabel.resize(200, 32)
        self.trainLabel.move( 5, size.height()*.5 + 105)
        self.trainLabel.setText('Validation Folder:')

        self.test = QLineEdit(self)
        self.test.returnPressed.connect(self.setTest)
        self.test.setText(self.test_folder)
        self.test.move( 105, size.height()*.5 + 105)
        self.test.resize(200, 32)

        
        self.threshLabel = QLabel(self)
        self.threshLabel.resize(200, 32)
        self.threshLabel.move( 5, size.height()*.5 + 140)
        self.threshLabel.setText('Threshold:')

        self.thresh = QLineEdit(self)
        self.thresh.returnPressed.connect(self.setTest)
        self.thresh.setText(".1")
        self.thresh.move( 105, size.height()*.5 + 140)
        self.thresh.resize(200, 32)

  
        self.learningRate = QLabel(self)
        self.learningRate.resize(200, 32)
        self.learningRate.move( 5, size.height()*.5 + 175)
        self.learningRate.setText('Learning Rate:')

        self.rate = QLineEdit(self)
        self.rate.returnPressed.connect(self.setTest)
        self.rate.setText(".001")
        self.rate.move( 105, size.height()*.5 + 175)
        self.rate.resize(200, 32)

        
        self.learningRate = QLabel(self)
        self.learningRate.resize(200, 32)
        self.learningRate.move( 5, size.height()*.5 + 210)
        self.learningRate.setText('Epochs:')

        self.epochs = QLineEdit(self)
        self.epochs.setText("500")
        self.epochs.move( 105, size.height()*.5 + 210)
        self.epochs.resize(200, 32)

        self.runTraining = QPushButton("", self)
        self.runTraining.move( 310,size.height()/2 + 20)
        self.runTraining.clicked.connect(self.trainModel) 
        self.runTraining.setToolTip("<h3>Train Custom UI Model</h3>")

        self.runTraining.setIcon(QIcon(resource_path( "TrainCustomUIModel.png")))

        self.runTraining.setIconSize(QSize(size.width()/12, size.width()/12))
        self.runTraining.setFixedSize(QSize(size.width()/12, size.width()/12))

        self.runTest = QPushButton("", self)
        self.runTest.move( size.width()/12 + 310,size.height()/2 + 20)
        self.runTest.clicked.connect(self.testModel) 
        self.runTest.setToolTip("<h3>Test Custom UI Model</h3>")
        self.runTest.setIcon(QIcon( resource_path("TestModel.png")))

        self.runTest.setIconSize(QSize(size.width()/12, size.width()/12))
        self.runTest.setFixedSize(QSize(size.width()/12, size.width()/12))

        self.changeImage = QPushButton("", self)
        self.changeImage.move( 310 + size.width()/12*2,size.height()/2 + 20)
        self.changeImage.resize(300, 32)
        self.changeImage.clicked.connect(self.changeTestImage) 
        self.changeImage.setIcon(QIcon( resource_path("ChangeTestImage.png")))
        self.changeImage.setIconSize(QSize(size.width()/12, size.width()/12))
        self.changeImage.setFixedSize(QSize(size.width()/12, size.width()/12))
        self.changeImage.setToolTip("<h3>Change Test Image</h3>")


        self.main_window()
    def displayLoading(self):
        self.splash.show()
        app.processEvents()

        # Close SplashScreen after 2 seconds (2000 ms)
    def genQuestions(self):
        
        self.questions = []
        q_counter = 0
        
        screen = QApplication.primaryScreen()
        ##print('Screen: %s' % screen.name())
        size = screen.size()
        ##print('Size: %d x %d' % (size.width(), size.height()))
        rect = screen.availableGeometry()
        ##print('Available: %d x %d' % (rect.width(), rect.height()))
        self.width = size.width()
        for q in self.custom_questions:
            if q["site"].split('.')[0].lower() in self.sender().toolTip().lower():
                for question in q["questions"]:
                    self.questions.append(QPushButton(question.split(' (')[0], self))
                    self.questions[-1].setIcon(QIcon(os.path.join(this_path, 'icons', self.sender().toolTip().split(' (')[0] + '.ico')))
                   # self.questions[-1].move(20,200+len(self.questions)*100)/border: 1px solid black;

                    self.questions[-1].setStyleSheet("text-align:left;font-size:24px;")
                    self.questions[-1].setIconSize(QSize(size.width()/12, size.width()/12))
                    self.questions[-1].setFixedSize(QSize(size.width()/2,152))
                    self.questions[-1].show()
                    self.questions[-1].setToolTip("<h3>" + question + "</h3>")
                    self.questions[-1].clicked.connect(self.submitQ)
                    self.vbox.addWidget(self.questions[-1],  q_counter, 0 ) 
                    q_counter += 1

        for q in self.base_questions:
            self.questions.append(QPushButton("  " + q + self.sender().toolTip().split(' (')[0] + ".", self))
            self.questions[-1].setIcon(QIcon(os.path.join(this_path, 'icons', self.sender().toolTip().split(' (')[0] + '.ico')))
           # self.questions[-1].move(20,200+len(self.questions)*100)
            self.questions[-1].setStyleSheet("text-align:left;font-size:24px;")
            self.questions[-1].setIconSize(QSize(size.width()/12, size.width()/12))
            self.questions[-1].setFixedSize(QSize(size.width()/2,152))
            self.questions[-1].show()
            self.questions[-1].setToolTip("<h3>" + q + "</h3>")
            self.questions[-1].clicked.connect(self.submitQ)
            self.vbox.addWidget(self.questions[-1],  q_counter, 0 ) 
            q_counter += 1
        self.questions.append(QPushButton("", self))
        self.questions[-1].setIcon(QIcon(resource_path('backroot.png')))
        self.questions[-1].setIconSize(QSize(size.width()/12, size.width()/12))
        self.questions[-1].setFixedSize(QSize(size.width()/2,152))
        self.questions[-1].show()
        self.questions[-1].setToolTip("<h3>Back</h3>")
        self.questions[-1].clicked.connect(self.backRoot)
        self.vbox.addWidget(self.questions[-1],  q_counter, 0 )

        for x in self.atlas_array:
            x.hide()
    def backRoot(self):
        for x in self.questions:
            x.hide()
        for x in self.atlas_array:
            x.show()
    
    def submitQ(self):
        self.displayLoading()

        app.processEvents()

        for q in self.questions:
            app.processEvents()

            q.hide()
        time.sleep(2)
        for x in self.atlas_array:
            x.show()
            app.processEvents()
        time.sleep(2)

        self.line.setText(self.sender().text().split('(')[0])
        self.atlas()

    def changeTestImage(self):
        self.test_image = QFileDialog.getOpenFileName(self, 'Open File', '/home')[0]
        self.pic.setPixmap(QtGui.QPixmap(self.test_image))

    def __del__(self):
        # Restore sys.stdout
        sys.stdout = sys.__stdout__
    
    def normalOutputWritten(self, text):
        """Append text to the QTextEdit."""
        global graph
        # Maybe QTextEdit.append() works as well, but this is how I do it:
        cursor = self.consoleOutput.textCursor()
        cursor.movePosition(QtGui.QTextCursor.End)
        cursor.insertText(text)
        self.consoleOutput.setTextCursor(cursor)
        self.consoleOutput.ensureCursorVisible()
        cursor = graph.consoleOutput.textCursor()

        cursor.movePosition(QtGui.QTextCursor.End)
        cursor.insertText(text)
        graph.consoleOutput.setTextCursor(cursor)
        graph.consoleOutput.ensureCursorVisible()
        

    def trainer(self,image,  train, test, model_file , labels, output):
        global test_image
       # ##print(torch.cuda.is_available())
       # model = core.Model.load( "model_weights.pth",  ["links","buttons","sliders","dropdowns","switches","menus","textareas","textfields","checkboxes","radioboxes","images","text"])
     
        #custom_transforms = transforms.Compose([transforms.ToPILImage(),transforms.Resize(900), transforms.RandomHorizontalFlip(0.5), transforms.ColorJitter(saturation=0.2),transforms.ToTensor(),utils.normalize_transform(), ])
        
        utils.xml_to_csv(train, 'train_labels.csv')
        utils.xml_to_csv(test, 'val_labels.csv')
        import csv
       # with open('train_labels.csv', 'r') as file:
       #     # open the file in the write mode
       #    
       #     reader = csv.reader(file)
       #     counter = 0
       #     for row in reader:
       #         counter += 1
       #         if "xmin" != row[4] and (int(row[4]) >= int(row[6]) or int(row[5]) >= int(row[7])):notifi
       #             ##print(row[4], row[6], row[5], row[7])
       #             ##print(row[0].split(".jpg")[0] + ".xml")
       #             ##print(row[0])
       #             import os
       #             try:
       #                 os.remove('train/' + row[0].split(".jpg")[0] + ".xml")
       #             
       #                 os.remove( 'train/' + row[0])
       #             except:
       #                 ##print("error")
#
       # with open('val_labels.csv', 'r') as file:
       #     # open the file in the write mode
       #    
       #     reader = csv.reader(file)
       #     counter = 0
       #     for row in reader:
       #         counter += 1
       #         if "xmin" != row[4] and (int(row[4]) >= int(row[6]) or int(row[5]) >= int(row[7])):
       #             ##print(row[4], row[6], row[5], row[7])
       #             ##print(row[0].split(".jpg")[0] + ".xml")
       #             ##print(row[0])
       #             import os
       #             try:
       #                 os.remove('test/' + row[0].split(".jpg")[0] + ".xml")
       #             
       #                 os.remove('test/' + row[0])
       #             except:
       #                 ##print("error")
# Defin#e custom transforms to apply to your datasetcd 
#
# Pass in a CSV file instead of XML files for faster Dataset initialization speeds

     
        dataset = core.Dataset('train_labels.csv', train)

        valid_dataset = []
        valid_valid = []
        removes = []


        

        val_dataset = core.Dataset('val_labels.csv', test)  # Validation dataset for training
    
# Create your own DataLoader with custom option
# Create your own DataLoader with custom options

        loader=core.DataLoader(dataset, batch_size=2, shuffle=True)#L3
        # Use MobileNet instead of the default ResNet
        ##print(labels)
        output = core.Model(labels)
        losses = output.fit(loader, val_dataset, epochs=int(self.epochs.text()),  lr_step_size=5,learning_rate=float(self.rate.text()), verbose=True)

        #plt.plot(losses)  # Visualize loss throughout training
        #plt.show()
        #image = utils.read_image("test_image2.png") 
        #predictions = model.predict(image)
        #labels, boxes, scores = predictions

        output.save(model_file)  # Save model to a file

        #show_labeled_image(image, boxes, labels)
    def trainModel(self):
        import threading
        t0 = threading.Thread(target = self.trainer, args=(self.test_image, self.train.text(), self.test.text(), self.model.text(), self.labels_input.text().split(","), self.output))
        t0.start()
        #t0.join()
        #self.output.save(self.model_file) 
        #self.trainer()

    """
        Non-max Suppression Algorithm
        @param list  Object candidate bounding boxes
        @param list  Confidence score of bounding boxes
        @param float IoU threshold
        @return Rest boxes after nms operation
    """
    def nms(self, bounding_boxes, confidence_score, labels, threshold):
        # If no bounding boxes, return empty list
        if len(bounding_boxes) == 0:
            return [], []
    
        # Bounding boxes
        boxes = np.array(bounding_boxes)
    
        # coordinates of bounding boxes
        start_x = boxes[:, 0]
        start_y = boxes[:, 1]
        end_x = boxes[:, 2]
        end_y = boxes[:, 3]
    
        # Confidence scores of bounding boxes
        score = np.array(confidence_score)
    
        # Picked bounding boxes
        picked_boxes = []
        picked_score = []
        picked_labels = []
    
        # Compute areas of bounding boxes
        areas = (end_x - start_x + 1) * (end_y - start_y + 1)
    
        # Sort by confidence score of bounding boxes
        order = np.argsort(score)
    
        # Iterate bounding boxes
        while order.size > 0:
            # The index of largest confidence score
            index = order[-1]
    
            # Pick the bounding box with largest confidence score
            picked_boxes.append(index)
    
            # Compute ordinates of intersection-over-union(IOU)
            x1 = np.maximum(start_x[index], start_x[order[:-1]])
            x2 = np.minimum(end_x[index], end_x[order[:-1]])
            y1 = np.maximum(start_y[index], start_y[order[:-1]])
            y2 = np.minimum(end_y[index], end_y[order[:-1]])
    
            # Compute areas of intersection-over-union
            w = np.maximum(0.0, x2 - x1 + 1)
            h = np.maximum(0.0, y2 - y1 + 1)
            intersection = w * h
    
            # Compute the ratio between intersection and union
            ratio = intersection / (areas[index] + areas[order[:-1]] - intersection)
    
            left = np.where(ratio < threshold)
            order = order[left]
    
        return picked_boxes

    def soft_nms_pytorch(self, dets, labels, box_scores, sigma=0.5, thresh=0.001, cuda=0):
        """
        Build a pytorch implement of Soft NMS algorithm.
        # Augments
            dets:        boxes coordinate tensor (format:[y1, x1, y2, x2])
            box_scores:  box score tensors
            sigma:       variance of Gaussian function
            thresh:      score thresh
            cuda:        CUDA flag
        # Return
            the index of the selected boxes
        """
    
        # Indexes concatenate boxes with the last column
        N = dets.shape[0]
        if cuda:
            indexes = torch.arange(0, N, dtype=torch.float).cuda().view(N, 1)
        else:
            indexes = torch.arange(0, N, dtype=torch.float).view(N, 1)
        dets = torch.cat((dets, indexes), dim=1)
    
        # The order of boxes coordinate is [y1,x1,y2,x2]
        y1 = dets[:, 0]
        x1 = dets[:, 1]
        y2 = dets[:, 2]
        x2 = dets[:, 3]
        scores = box_scores
        areas = (x2 - x1 + 1) * (y2 - y1 + 1)
    
        for i in range(N):
            # intermediate parameters for later parameters exchange
            tscore = scores[i].clone()
            pos = i + 1
    
            if i != N - 1:
                maxscore, maxpos = torch.max(scores[pos:], dim=0)
                if tscore < maxscore:
                    dets[i], dets[maxpos.item() + i + 1] = dets[maxpos.item() + i + 1].clone(), dets[i].clone()
                    scores[i], scores[maxpos.item() + i + 1] = scores[maxpos.item() + i + 1].clone(), scores[i].clone()
                    areas[i], areas[maxpos + i + 1] = areas[maxpos + i + 1].clone(), areas[i].clone()
    
            # IoU calculate
            yy1 = np.maximum(dets[i, 0].to("cpu").numpy(), dets[pos:, 0].to("cpu").numpy())
            xx1 = np.maximum(dets[i, 1].to("cpu").numpy(), dets[pos:, 1].to("cpu").numpy())
            yy2 = np.minimum(dets[i, 2].to("cpu").numpy(), dets[pos:, 2].to("cpu").numpy())
            xx2 = np.minimum(dets[i, 3].to("cpu").numpy(), dets[pos:, 3].to("cpu").numpy())
            
            w = np.maximum(0.0, xx2 - xx1 + 1)
            h = np.maximum(0.0, yy2 - yy1 + 1)
            inter = torch.tensor(w * h).cuda() if cuda else torch.tensor(w * h)
            ovr = torch.div(inter, (areas[i] + areas[pos:] - inter))
    
            # Gaussian decay
            weight = torch.exp(-(ovr * ovr) / sigma)
            scores[pos:] = weight * scores[pos:]
        keep= []
        ##print(scores)
        for i in range(0, len(dets)):
            if scores[i] < thresh:
                dets[i].remove()
                labels[i].remove()
        # select the boxes and keep the corresponding indexes
        keep = dets
        labels_keep = labels
        return keep, labels_keep    
    def testModel(self):
        import torch 

       # self.output.save(self.model_file) 
        ##print(self.labels_input.text().split(","))
        model = core.Model.load(self.model.text(),  self.labels_input.text().split(","))
        image = utils.read_image(self.test_image) 
        predictions = model.predict(image)
        labels, boxes, scores = predictions
        if boxes.ndim == 1:
            boxes = boxes.view(1, 4)
        thresh=float(self.thresh.text())
        ##print(scores)
        filtered_indices=np.where(scores>thresh)
        filtered_scores=scores[filtered_indices]
        filtered_boxes=boxes[filtered_indices]
        num_list = filtered_indices[0].tolist()
        filtered_labels = [labels[i] for i in num_list]
       # show_labeled_image(image, filtered_boxes, filtered_labels)
        # Plot each box
        
        ##print(filtered_boxes)
        ##print(filtered_labels)        
        indices = self.nms(filtered_boxes, filtered_scores, filtered_labels, .05)
        filtered_boxes = filtered_boxes[indices]
        filtered_labels = [filtered_labels[i] for i in indices]

       # filtered_labels = filtered_labels[indices]
        for i in range(filtered_boxes.shape[0]):
            box = filtered_boxes[i]
            
            width, height = (box[2] - box[0]).item(), (box[3] - box[1]).item()
            initial_pos = (int(box[0].item()), int(box[1].item()))
            end_pos = (int(box[2].item()), int(box[3].item()))
            cv2.rectangle(image,initial_pos, end_pos, (0, 255, 0), 2)
            if filtered_labels is not None:
                cv2.putText(image, filtered_labels[i],initial_pos, cv2.FONT_HERSHEY_SIMPLEX, 1, (255, 0, 0), 2)
        height, width, channel = image.shape
        bytesPerLine = 3 * width
        qImg = QImage(image.data, width, height, bytesPerLine, QImage.Format_RGB888)
        self.pic.setPixmap(QPixmap(qImg))


    def setTest(self):
        self.test_folder = self.test.text()
    def setTrain(self):
        self.train_folder = self.train.text()
    def changeModel(self):
        self.model_file = self.model.text()
    def back(self):
        self.question1.hide()
        self.question2.hide()
        self.question3.hide()
        self.question4.hide()
        self.backButton.hide()
        self.API.show()
        self.Custom.show()
        self.Amazon.show()
        self.Google.show()
      #  self.speakButton.show()
        self.pushButton.show()
        self.line.show()
        self.nameLabel.show()
        self.backButton.hide()

    def submitQ1(self):
        self.question1.hide()
        self.question2.hide()
        self.question3.hide()
        self.question4.hide()
        self.backButton.hide()
        self.Amazon.show()
        self.Google.show()
        self.Custom.show()
        self.API.show()
    #    self.speakButton.show()
        self.line.setText(self.question1.text())
        self.atlas()

    def submitQ2(self):
        self.question1.hide()
        self.question2.hide()
        self.question3.hide()
        self.question4.hide()
        self.backButton.hide()
        
    #    self.speakButton.show()
        self.Amazon.show()
        self.Google.show()
        self.Custom.show()
        self.API.show()
        self.line.setText(self.question2.text())
        self.atlas()

    def submitQ3(self):

        self.question1.hide()
        self.question2.hide()
        self.question3.hide()
        self.question4.hide()
        self.backButton.hide()
        
     #   self.speakButton.show()
        self.Amazon.show()
        self.Google.show()
        self.Custom.show()
        self.API.show()
        self.line.setText(self.question3.text())
        self.atlas()

    def submitQ4(self):
        self.question1.hide()
        self.question2.hide()
        self.question3.hide()
        self.question4.hide()
        self.backButton.hide()
        
    #    self.speakButton.show()
        self.Amazon.show()
        self.Google.show()
        self.Custom.show()
        self.API.show()
        self.line.setText(self.question4.text())
        self.atlas()
      
    def speak(self):
        # get audio from the microphone                                                                       
        r = sr.Recognizer()                                                                                   
        with sr.Microphone() as source:                                                                       
            ##print("Speak:")       
            r.adjust_for_ambient_noise(source)                                                                            
            audio = r.listen(source)   
            ##print("Done")
        try:
            ##print(r.recognize_google(audio))
            self.line.setText(r.recognize_google(audio))
            ##print("You said " + r.recognize_google(audio))
            self.atlas()
        except sr.UnknownValueError:
            print("Could not understand audio")
        except sr.RequestError as e:
            print("Could not request results; {0}".format(e))
    def handleAmazon(self):
        self.Amazon.hide()
        self.Google.hide()
        self.Custom.hide()
        self.API.hide()
        
   #     self.speakButton.hide()
        self.backButton.show()
        self.question1.show()
        self.question2.show()
        self.question3.show()
        self.question4.show()


        #print("Amazon")
    def handleGoogle(self):
        print("Google")
    def getIcon(self,url):
        data = urllib.request.urlopen(url).read()
        icon = QtGui.QPixmap()
        icon.loadFromData(data)
        return icon

  
    def main_window(self):
        
        self.setWindowTitle(self.title)
        self.setGeometry(self.top, self.left, self.width, self.height)
        self.show()
from PySide2.QtWidgets import (QApplication, QVBoxLayout, QHBoxLayout, QRadioButton,QButtonGroup, QLineEdit, QListWidget, QListWidgetItem, QPushButton,
                               QGraphicsDropShadowEffect, QWidget, QDialog, QLabel)
from PySide2.QtCore import Qt, QPoint
from PySide2.QtGui import QColor, QRegion, QPainterPath, QPalette

class DraggableListWidget(QListWidget):
    def mousePressEvent(self, event):
        self.parent().mousePressEvent(event)
    
    def mouseMoveEvent(self, event):
        self.parent().mouseMoveEvent(event)

base_log = []
new_base = []
total_tokens = 0
code_counter = 0

# Mock data

user_plan = "free"
ComPortTikTok = []  # Mock data list
port = None  # Placeholder
ipBlock = False
automation_prompt = ""
faq_prompt = ""


class UserMessageItem(QListWidgetItem):
    def __init__(self, text):
        super().__init__()
        self.text = text

        # Create a custom widget for the list item
        self.widget = QWidget()
        layout = QHBoxLayout(self.widget)

        self.label = QLabel(self.text)

        self.label.setWordWrap(True)
        self.label.setTextInteractionFlags(Qt.TextSelectableByMouse)
        self.label.setStyleSheet("""
            background-color: qlineargradient(spread:pad, x1:0.773, y1:0, x2:0.251, y2:0, stop:0 #3D2C5C, stop:1 #5B01FF);
            border-radius: 16px;
            padding: 14px 18px;
            border: 1px black solid;
                      max-width: 300px;           
            color: #FFFFFF;
            font-size: 18px;
            line-height: 22px;
        """)
   
        layout.addWidget(self.label, 0, Qt.AlignRight)

        self.widget.setLayout(layout)
        self.setSizeHint(self.widget.sizeHint())

class UploadButton(QListWidgetItem):
    def __init__(self, upload_recording, parent=None):
        super().__init__(parent)

        # Create a custom widget for the list item
        self.widget = QWidget()
        self.layout = QHBoxLayout(self.widget)

        self.addButton = QPushButton("Share Agent", self.widget)
        # self.addButton.clicked.connect(upload_recording)  # Connect your upload_recording signal.
        self.addButton.clicked.connect(upload_recording)

        self.layout.addWidget(self.addButton, 0)  # Removed alignment, as it's unnecessary with stretch.

        # Add spacer after the button to push everything to the left.
        self.layout.addStretch(1)

        self.widget.setLayout(self.layout)

        # Calculate size hint based on the layout.
        self.setSizeHint(self.widget.sizeHint())
from PySide2.QtWidgets import QListWidgetItem, QLabel, QWidget, QHBoxLayout
from PySide2.QtCore import Qt
from PySide2.QtGui import QSyntaxHighlighter, QTextCharFormat, QFont
import re

class SimpleSyntaxHighlighter(QSyntaxHighlighter):
    def __init__(self, parent=None):
        super().__init__(parent)

        # Define styles for each syntax element
        self.keywordFormat = QTextCharFormat()
        self.keywordFormat.setForeground(QColor("#FFD700"))
        self.keywordFormat.setFontWeight(QFont.Bold)

        self.stringFormat = QTextCharFormat()
        self.stringFormat.setForeground(QColor("#00FF00"))

        self.commentFormat = QTextCharFormat()
        self.commentFormat.setForeground(QColor("#ADADAD"))
        self.commentFormat.setFontItalic(True)

        self.functionFormat = QTextCharFormat()
        self.functionFormat.setForeground(QColor("#0000FF"))

        # Pattern for keywords
        keywords = [
            "class", "def", "import", "as", "from", "return",
            "break", "continue", "if", "elif", "else",
            "try", "except", "raise", "with", "finally",
            "for", "while", "pass", "not", "in", "is", "and",
            "or", "lambda", "global", "nonlocal", "assert", "yield",
            "del", "True", "False", "None"
        ]
        self.keywordPattern = "\\b(" + "|".join(keywords) + ")\\b"

        # Pattern for function names
        self.functionPattern = "\\bdef\\s+(\\w+)"

        # Pattern for string literals
        self.stringPattern = r'"[^"\\]*(\\.[^"\\]*)*"|\'[^\'\\]*(\\.[^\'\\]*)*\''

        # Single-line comments
        self.commentPattern = r"#[^\n]*"

        self.rules = [
            (self.keywordPattern, self.keywordFormat),
            (self.functionPattern, self.functionFormat),
            (self.stringPattern, self.stringFormat),
            (self.commentPattern, self.commentFormat),
        ]

    def highlightBlock(self, text):
        for pattern, fmt in self.rules:
            for match in re.finditer(pattern, text):
                start = match.start(1) if "def" in pattern else match.start()
                length = match.end(1) - start if "def" in pattern else match.end() - start
                self.setFormat(start, length, fmt)

        # Apply the default state to the rest of the text
        self.setCurrentBlockState(0)

# Assuming SimpleSyntaxHighlighter is a previously defined class for syntax highlighting
# from your_import_module import SimpleSyntaxHighlighter
from PySide2.QtWidgets import QLabel, QWidget, QVBoxLayout
from PySide2.QtCore import Qt, Signal
from PySide2.QtGui import QPalette, QColor, QPainter, QLinearGradient

class ClickableLabel_runtime(QLabel):
    clicked = Signal()

    def __init__(self, text, parent=None):
        super().__init__(text, parent)
        self.full_text = text
        self.setWordWrap(True)
        self.is_code_visible = False

        self.overlay = QLabel("View Runtime Logs", self)
        self.overlay.setAlignment(Qt.AlignCenter)
        self.overlay.setStyleSheet("""
            color: white;
            font-weight: bold;
        """)

        self.update_style()

    def update_style(self):
        if not self.is_code_visible:
            self.setStyleSheet("""
                color: transparent;
                font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
                padding: 5px;
                border-radius: 5px;
                background-color: #000;
            """)
            self.overlay.show()
        else:
            self.setStyleSheet("""
                color: #FFF;
                font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
                padding: 5px;
                border-radius: 5px;
                background-color: #000;
            """)
            self.overlay.hide()
        self.setFixedHeight(self.sizeHint().height())

    def mousePressEvent(self, event):
        self.is_code_visible = not self.is_code_visible
        self.update_style()
        self.clicked.emit()
        super().mousePressEvent(event)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self.overlay.setGeometry(0, 0, self.width(), self.height())

    def paintEvent(self, event):
        super().paintEvent(event)
        if not self.is_code_visible:
            painter = QPainter(self)
            gradient = QLinearGradient(0, 0, 0, self.height())
            gradient.setColorAt(0.0, QColor(0, 0, 0, 180))
            gradient.setColorAt(1, QColor(0, 0, 0, 0))
            painter.fillRect(self.rect(), gradient)
class ClickableLabel(QLabel):
    clicked = Signal()

    def __init__(self, text, parent=None):
        super().__init__(text, parent)
        self.full_text = text
        self.setWordWrap(True)
        self.is_code_visible = False

        self.overlay = QLabel("View Code", self)
        self.overlay.setAlignment(Qt.AlignCenter)
        self.overlay.setStyleSheet("""
            color: white;
            font-weight: bold;
        """)

        self.update_style()

    def update_style(self):
        if not self.is_code_visible:
            self.setStyleSheet("""
                color: transparent;
                font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
                padding: 5px;
                border-radius: 5px;
                background-color: #000;
            """)
            self.overlay.show()
        else:
            self.setStyleSheet("""
                color: #FFF;
                font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
                padding: 5px;
                border-radius: 5px;
                background-color: #000;
            """)
            self.overlay.hide()
        self.setFixedHeight(self.sizeHint().height())

    def mousePressEvent(self, event):
        self.is_code_visible = not self.is_code_visible
        self.update_style()
        self.clicked.emit()
        super().mousePressEvent(event)

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self.overlay.setGeometry(0, 0, self.width(), self.height())

    def paintEvent(self, event):
        super().paintEvent(event)
        if not self.is_code_visible:
            painter = QPainter(self)
            gradient = QLinearGradient(0, 0, 0, self.height())
            gradient.setColorAt(0.0, QColor(0, 0, 0, 180))
            gradient.setColorAt(1, QColor(0, 0, 0, 0))
            painter.fillRect(self.rect(), gradient)
class ReceiverMessageItem(QListWidgetItem):
    def __init__(self, text):
        super().__init__()
        self.text = text

        # Create a custom widget for the list item
        self.widget = QWidget()

        self.layout = QHBoxLayout(self.widget)
        
        # Check for a code block
        if "bash" in self.text or "python" in self.text and user_type != "Technical":
            self.label = ClickableLabel(self.text)
            # Potentially apply a syntax highlighter
            # self.highlighter = SimpleSyntaxHighlighter(self.label)
        elif "runtime" in self.text and user_type != "Technical":
            self.label = ClickableLabel_runtime(self.text)
        elif "runtime" in self.text and user_type == "Technical":
            self.label = QLabel(self.text)
            self.label.setStyleSheet("""
                background-color: #000;
                border-radius: 16px;
                padding: 14px 18px;
                border: 1px solid black;
                max-width: 300px;
                color: #FFFFFF;
                font-size: 18px;
                line-height: 22px;
            """)

        elif "bash" in self.text or "python" in self.text and user_type == "Technical":
            self.label = QLabel(self.text)
            self.label.setStyleSheet("""
                background-color: #000;
                border-radius: 16px;
                padding: 14px 18px;
                border: 1px solid black;
                max-width: 300px;
                color: #FFFFFF;
                font-size: 18px;
                line-height: 22px;
            """)


        else:
            self.label = QLabel(self.text)
            self.label.setStyleSheet("""
                background: #FBFBFB;
                border-radius: 16px;
                padding: 14px 18px;
                border: 1px solid black;
                max-width: 300px;      
                color: black;
                font-size: 18px;
                line-height: 22px;
            """)
        
        self.label.setWordWrap(True)
        self.label.setTextInteractionFlags(Qt.TextSelectableByMouse)
        
        # Align label to the left
        self.label.setAlignment(Qt.AlignLeft)
        self.layout.addWidget(self.label, 0, Qt.AlignLeft)

        # Add a spacer after the label to fill the remaining space
        self.layout.addStretch(1)

        self.widget.setLayout(self.layout)
        self.setSizeHint(self.widget.sizeHint())

    def hideGradientOverlay(self, event):
        if self.label.styleSheet() == "":
            self.label.setStyleSheet("""
                background-color: linear-gradient(to bottom, rgba(0,0,0,0.3) 0%, rgba(0,0,0,0) 100%);
                color: #FFFFFF;
                font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
                padding: 5px;
                border-radius: 5px;
                max-width: 300px;
            """)
            self.label.setText("Click to view\n" + self.text)  # Add 'click to view' text
        else:
            self.label.setStyleSheet("""
                background: #FBFBFB;
                border-radius: 16px;
                padding: 14px 18px;
                border: 1px solid black;
                max-width: 300px;      
                color: black;
                font-size: 18px;
                line-height: 22px;
            """)
            self.label.setText(self.text)  # Restore original text

        super().mousePressEvent(event)
 #       self.setSizeHint(self.widget.sizeHint())
from contextlib import redirect_stdout
import io

full = io.StringIO()
import pandas as pd
from PyPDF2 import PdfReader
import docx
from io import BytesIO

from PySide2.QtCore import Signal as pyqtSignal

class WorkerSignals(QObject):
    progress = pyqtSignal(str)  # this signal emits a string
    update = pyqtSignal(str)  # this signal emits a string
    recording = pyqtSignal()  # this signal emits a string
    recording_start = pyqtSignal()  # this signal emits a string
    upload = pyqtSignal()  # this signal emits a string
    notifications = pyqtSignal(str, str,str,object, str)  # this signal emits a string
    scraper = pyqtSignal(str, str)
    process_nodes = pyqtSignal(str, object, object, str, str)
    open = pyqtSignal(str)  # this signal emits a string
    chat = pyqtSignal(str)  # this signal emits a string

    setkey = pyqtSignal(str)  # this signal emits a string
    hide = pyqtSignal()  # this signal emits a string
    show = pyqtSignal()  # this signal emits a string
    reconnectself = pyqtSignal()  # this signal emits a string
    run = pyqtSignal(object)
    semanticScrape = pyqtSignal(str)  # this signal emits a string
    API = pyqtSignal(object,str)  # this signal emits a string
class ChatApp(QDialog, QObject):
    messageSignal = pyqtSignal(str)

    def __init__(self, thread_signals, open_file):
        super().__init__(None, Qt.WindowStaysOnTopHint)
        self.slide_up = False
        self.is_dragging = False
        self.is_resizing = False
        self.drag_position = None
        self.resize_direction = None
        self.mode = "chat"
        self.total_errors = 0
        # Set up main window properties
        self.setGeometry(0, 0, 600, 100)
        self.setWindowTitle('Chat Interface')
        self.setWindowFlags(Qt.FramelessWindowHint | Qt.Dialog | Qt.WindowStaysOnTopHint)
   #     self.setAttribute(Qt.WA_TranslucentBackground)  # Make the background transparent
        self.setWindowOpacity(0.9)  # Set opacity to 90%
        self.open_file = open_file
        # Apply the shadow effect
        shadow = QGraphicsDropShadowEffect(self)
        shadow.setBlurRadius(90)
        shadow.setXOffset(0)
        shadow.setYOffset(0)
        shadow.setColor(QColor(0, 0, 0, 255))
        self.setGraphicsEffect(shadow)
        #calculate the center of the screen and move the window there using only .primaryScreen() and .size()
        screen = QApplication.primaryScreen()
        #print('Screen: %s' % screen.name())
        size = screen.size()
        #print('Size: %d x %d' % (size.width(), size.height()))
        rect = screen.availableGeometry()
        #print('Available: %d x %d' % (rect.width(), rect.height()))
        #move the window to the center of the screen
        self.move(size.width()/2 - 300, size.height()/2 - 350)
        # Create the main layout
        self.layout = QVBoxLayout(self)
        self.setContentsMargins(0, 0, 0, 0)  # This adds a margin of 10 pixels on all sides

        self.setLayout(self.layout)
        #set the layout background color to transparent


        self.monitor = ActivityMonitor(timeout=3600)

        self.monitor.start()
        self.monitor.update_activity()
        thread_signals.progress.connect(self.send_message_bot)
        thread_signals.recording.connect(self.stop_screen_recorder)
        thread_signals.upload.connect(self.upload_recording)
        #thread_signals.store_final.connect(self.store_final)
        #thread_signals.store_creator.connect(self.store_creator)
        #thread_signals.get_creator.connect(self.get_creator)
        
        thread_signals.process_nodes.connect(self.process_nodes)
        thread_signals.recording_start.connect(self.start_screen_recorder)
        thread_signals.open.connect(self.open_file)
        thread_signals.chat.connect(self.send_message)
        thread_signals.notifications.connect(self.notification)
        thread_signals.scraper.connect(self.scraper)
        thread_signals.hide.connect(self.hide_chat_display)
        thread_signals.show.connect(self.toggle_chat_display)
        thread_signals.setkey.connect(self.setkey)
        thread_signals.reconnectself.connect(self.reconnect)
        thread_signals.update.connect(self.send_update_bot)
        thread_signals.semanticScrape.connect(self.semanticScrape)
        thread_signals.API.connect(self.API)
        thread_signals.run.connect(self.run)
     #   self.thread_signals.update.connect(self.send_update_bot)

        # Initialize quick launch buttons
        

        self.chat_display = DraggableListWidget(self)
        self.layout.addWidget(self.chat_display)
        #hidfe the chat window initially 
        #self.chat_display.hide()
        # Horizontal layout for entry field and send button
        hlayout = QHBoxLayout()
        self.entry_field = QLineEdit(self)
        #self.send_button = QPushButton('Send', self)
        
        self.button = QPushButton('', self)
        self.button.setIcon(QIcon(resource_path('upload.png')))  # Set the icon here
        self.button.clicked.connect(self.openFileDialog)
        hlayout.addWidget(self.button)
        
        hlayout.addWidget(self.entry_field)
       # hlayout.addWidget(self.send_button)
        self.layout.addLayout(hlayout)
        self.textEdit = ""
        self.image_data = ""
        
        self.entry_field.setPlaceholderText("Ask Project Atlas anything or hit ESC to close...") # Placeholder text
        self.entry_field.textChanged.connect(self.update_quick_launch_menu)
        self.entry_field.installEventFilter(self)  # Install event filter

        # Connect send button to the functionality
   #     self.send_button.clicked.connect(self.send_message)
        self.entry_field.returnPressed.connect(self.send_message)

        self.setStyleSheet("""
                           
        QDialog {
    background-color: rgba(240,240,240, 250);
margin: 0px;
padding: 0px;
                                   border: 1px solid #ccc;

                                          }
        QVBoxLayout {
            border: 1px solid #ccc;
           margin: 0px;
padding: 0px;
    background-color: rgba(240,240,240, 250);

        }
                           DraggableListWidget {
                                   border: 1px solid #ccc;
margin: 0px;
padding: 0px;
                                   

                           }

                           QLineEdit {
                                      border: 1px solid #ccc;
                           min-width: 150px;
                            min-height: 64px;
                            border-radius: 16px;
                            border: 0px;
                            padding:5px;
                            font-size: 18px;
                           margin:0px;
                            line-height: 21px;
                           }
        QPushButton {
            background-color: qlineargradient(spread:pad, x1:0.773, y1:0, x2:0.251, y2:0, stop:0 #FFFFFF, stop:1 #FEFEFE);
            min-width: 30px;
            min-height: 48px;
            border-radius: 16px;
            border: 30px;
            width: 10px;
            height: 64px;
            font-size: 16px;
                           margin:5px;
            line-height: 18px;
            color: #000000;
            font-weight: bold;
        }
        QPushButton:hover {
            background-color: qlineargradient(spread:pad, x1:0.773, y1:0, x2:0.251, y2:0, stop:0 #5B01FF, stop:1 #3D2C5C);
            color: #FFFFFF;
        }
        """)
        self.chat_display.hide()  # Hide the chat display initially
        self.animation = QPropertyAnimation(self, b"geometry")
        self.animation.finished.connect(self.on_animation_finished)
        #self.create_quick_launch_buttons()

    def openFileDialog(self):
        options = QFileDialog.Options()
        fileName, _ = QFileDialog.getOpenFileName(self, "Select File", "", "All Files (*);;Text Files (*.txt);;CSV Files (*.csv);;Excel Files (*.xlsx);;Word Files (*.docx);;PDF Files (*.pdf)", options=options)
        if fileName:
            self.loadFile(fileName)
    def store_final(self, final_states):
        store_final_states( "user_id", final_states)

    def store_creator(self, creator_stats):
        store_creator_stats( "user_id", creator_stats)

    def get_creator(self):
        return retrieve_creator_stats("")
    
    def API(self, text = "test", data = "test"):
        global api_requests, api_outputs, graph
        print("API saved: ", data)
        print(text)
        api_requests[data] = "done"
        api_outputs[data] = text
        self.monitor.update_activity()

        #graph.save_session("form.cheat")
        print("done api")
    def run(self, data):
        global graph
        print(data)
        self.monitor.update_activity()

        print("save data")
        graph.openCheat(data["json_output"][0]["file"])
        
        for key, line_edit in data["json_output"][0].items():
            value = line_edit
            graph.global_variables[key] = value
        for key, line_edit in data.items():
            value = line_edit
            graph.global_variables[key] = value
            
        print("run agents")
        print(graph.global_variables)
        graph.playRecording()
    def loadFile(self, fileName):
        global zipMode
        file_ext = fileName.split('.')[-1].lower()
        data = resource_path(fileName)
        print(data)
        self.monitor.update_activity()

        print("FILE LOADED")

        self.textEdit = data
        if "zip" in file_ext:
            zipMode = True

#        self.image_data  = data_image
    def imageToDataUrl(self, fileName):
        mime_type, _ = mimetypes.guess_type(fileName)
        with open(fileName, 'rb') as file:
            encoded_string = base64.b64encode(file.read()).decode('utf-8')
        return f'data:{mime_type};base64,{encoded_string}'
    def eventFilter(self, source, event):
        if source == self.entry_field:
            if event.type() == QEvent.MouseButtonPress:
                self.mousePressEvent(event)
                return True
            elif event.type() == QEvent.MouseMove:
                self.mouseMoveEvent(event)
                return True
            elif event.type() == QEvent.MouseButtonRelease:
                self.mouseReleaseEvent(event)
                return True
        return super().eventFilter(source, event)
    def hide_chat_display(self):
        self.chat_display.hide()
        graph.notification_manager.hide()

    def semanticScrape(self, text):
        graph.sendMessageRTCAsync('semanticSearch:' + text)
    def create_quick_launch_buttons(self):
        """Initializes the quick-launch menu with placeholder buttons."""
        self.quick_launch_layout = QGridLayout()
        
        # Placeholder information for buttons
        self.buttons_info = [
            {"name": "action agent", "icon": resource_path("action_agent.png"), "cheat": "action recorder"},
            {"name": "Movie_Generator_0", "icon": resource_path("movie_agent.png"),"cheat": "Movie_Generator_0.cheat"},
            {"name": "TikTok_Marketing_Agent_0", "icon": resource_path("tiktok.png") ,"cheat": "TikTok_Marketing_Agent_0.cheat"},
        ]
        self.buttons = []
        for i, action in enumerate(self.buttons_info):
            btn = QPushButton()
            icon = QIcon(action["icon"])
            btn.setIcon(icon)
            btn.setIconSize(QSize(100, 100))  # Adjust icon size as needed
            btn.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)  # Optional: Make buttons expand
            btn.setObjectName(action["name"])
            btn.clicked.connect(self.send_cheat)
            self.quick_launch_layout.addWidget(btn, i // 3, i % 3)
            self.buttons.append(btn)

        self.layout.addLayout(self.quick_launch_layout)
        #give the layout a transparent background gray
    def send_cheat(self):
        global desktop_guidelines, previous_steps, last_input
        cheat = self.sender().objectName()
        #print(cheat)
        if "action agent" in cheat:
            self.chat_display.hide()
            graph.startRecording()
        else:
            self.chat_display.hide()

            cheat = cheat + ".cheat"
            #print(cheat)
            openCheat(cheat)

    def update_quick_launch_menu(self):
        """Filters the quick-launch buttons based on the text input."""
        search_text = self.entry_field.text().lower()
        found = False
#        for btn in self.buttons:
     #       btn.setVisible(search_text in btn.objectName().lower())
     #       if btn.isVisible():
     #           found = True
       
    def toggle_chat_display(self):
        graph.notification_manager.show()
        is_collapsed = self.chat_display.isHidden()
        if self.slide_up == False:
            app_rect = self.geometry()
            current_height = app_rect.height()
            is_collapsed = self.chat_display.isHidden()
    
                # If the chat_display is hidden, calculate the expanded height
                # and set geometry to show the display
            end_height = 700
            self.chat_display.show()
            self.slide_up = True
    
                   # Start the animation for expand/collapse
            start_rect = QRect(app_rect.x(), app_rect.y(), app_rect.width(), 0)
            end_rect = QRect(app_rect.x(), app_rect.y(), app_rect.width(), 700)
            
            self.animation.setStartValue(start_rect)
            self.animation.setEndValue(end_rect)
            self.animation.setDuration(2000)  # Duration in milliseconds
            self.animation.start()
        self.chat_display.show()
    def on_animation_finished(self):
        # Re-enable shadow effect after animation

        self.update()    
    def find_guidelines(self, input_text, guidelines):
        global desktop_guidelines, previous_steps, last_input
        #get the current time zone offset 
        offset = datetime.now().astimezone().utcoffset().total_seconds()
        #print(offset)
        total_guidelines = "The current time zone is " + str(offset) + " hours from UTC. When scheduling tasks, add the offset to the cron parameter to schedule tasks in the correct time zone."

      #  #print("all guidelines")
      #  #print(guidelines)
        #for guideline_obj in guidelines:
        #    for keyword in guideline_obj["keywords"]:
        #        if keyword.lower() in input_text.lower():
        #           # #print(keyword)
        #            total_guidelines += guideline_obj["guideline"]  # add newline for separation
        #            break  # if one keyword matches, no need to check the others for this guideline
        total_guidelines += new_prompt


        total_guidelines += r'''
            If the goal has already been accomplished, output only 'goal accomplished' with a short explanation, for example:
            Input: post a tweet about dogs. Previous completed steps: open twitter, click tweet button, type tweet. 
            Output:
            goal accomplished. It appears the tweet has been posted successfully based on the current state of the screen and previously provided steps.
            Do not assume the goal has been accomplished on just the previous steps alone, and use the saved data and the current screenshot to verify the goal has been accomplished. 
            If the user asks to scrape a link, always use browserScrape instead of semanticScrape, since links always require this. 
            If the user references the screen or screenshot, always use the semanticScrape action to scrape the screen.
            If the user asks to scrape something which is visible on the screen, for example the URL of the website on the screen, use semanticScrape. If the user asks to scrape something which is not visible on the screen, use browserScrape.
        Only generate automations using this json language and never generate code directly.

            Some use-cases require python to complete the tasks. In most cases, try to use the nodes above already mentioned like browserScrape, semanticScrape, open tab, click, keypress, and delay.
            If the use-case can only be accomplished using python code, you can ues a python action node. Format the python code correctly with newlines and tabs so it runs. Use the yt-dlp library for all youtube and tiktok automations. Install any relevant python libraries necessary first using a bash node, including yt-dlp. For example, if the user asks to download a tiktok video:
            Input: download this tiktok video: https://www.tiktok.com/@cheatlayer/video/7337267587287354625 and cut it into 3 parts and add a voiceover that says "cheatlayer" every 5 seconds
            Output:
            [start json]

            [{"type": "bash", "command": "pip install -U --pre \\"yt-dlp[default]\\""},
            {"type": "python", "code": "import yt_dlp\nydl = yt_dlp.YoutubeDL({'outtmpl': '%(id)s.%(ext)s'})\ninfo = ydl.extract_info('https://www.tiktok.com/@cheatlayer/video/7337267587287354625', download=False)\nurl = info['url']\nydl.download([url])"}]            
            Dont forget to include escape characters before the double quotes for yt-dlp
            Always use the python action node for complex tasks that require python code. Do not use the python action node for simple tasks that can be accomplished with the other nodes. All scraping tasks should be done with the browserScrape and semanticScrape nodes, and never use python for web scraping unless you need to like for tiktok videos.
            Always add print statements throughout the python code so it becomes easier to debug errors. 
            
            If the user supplies the location of a file, including PDF, doc, excel, text, or any file, use the add_data action node to load the file data into a variable. 
            If the user asks to scrape a list of links from a file, use the add_data action node to load the file data into a variable and then use the browserScrape action node to scrape the links.
            For example:
            Input: scrape all the links from this file C:\\Users\\Rohan\\Downloads\\templates\\31\\data.txt
            Output:
            [start json]
            [{"type": "add_data", "target": "links", "data": "C:\\Users\\Rohan\\Downloads\\templates\\31\\data.txt"},
                        {"type":"google_sheets_create", "URL": "sheet_URL", "Sheet_Name": "sheet_name" },

            {"type": "open tab", "target": "links", "actions": [{"type": "browserScrape", "target": "a", "data": "links", "description": "all links"}, {"type":"google_sheets_add_row", "data": ["links"] , "URL": "sheet_URL", "Sheet_Name": "sheet_name"  }]}
            ]
            [end json]

            If the user asks to summarize this pdf C:\\Users\\Rohan\\Downloads\\templates\\31\\data.pdf and email it to rohan@cheatlayer.com
            Input: summarize the pdf C:\\Users\\Rohan\\Downloads\\templates\\31\\data.pdf and email it to rohan@cheatlayer.com
            Output:
            [start json]
            [{"type": "add_data", "target": "pdf", "data": "C:\\Users\\Rohan\\Downloads\\templates\\31\\data.pdf"},
            {"type": "gpt4", "input": ["pdf"], "prompt": "Summarize the pdf", "data": "summary"},
            {"type": "email", "to": "rohan@cheatlayer.com", "subject": "Summary", "body": "Here is a summary of the pdf", "data": "summary"}]
            [end json]
            use the specific file location the user specifies, and do not use C:\\Users\\Rohan\\Downloads\\templates\\31\\data.pdf as the file location.


If the user asks to scrape data then email it, do not create a google sheet to scrape the data and email it directly. 
        If the user asks to generate charts or graphs to email, follow this pattern and assume there is a web_prompt variable already saved so do not generate it yourself:
        Input: analyze the input google sheet and generate a report to email to rohan@cheatlayer.com
        Output:
        [start json]
        [{"type": "google_sheets_read", "URL": "https://docs.google.com/spreadsheets/d/1uBTTkIcKM684UvD_qDucQ32rXNeS9Ip20RfatzX5KuI/edit?gid=0#gid=0", "Sheet_Name": "Sheet1", "data": "data_sheet"},
        {"type":"gpt4", "input": ["data_sheet"], "prompt": web_prompt, "data": "report"},
        
        {"type": "email", "to": "rohan@cheatlayer.com", "subject": "Report", "body": "Here is a report ", "data": "report"}]
        [end json]
       Always generate properly formatted JSON that can be run using the JSON.parse() function. Never concatenate strings in the generated JSON, and make sure all the keys and values are properly formatted and escaped. 
        The "input" parameter in the "gpt4" action is an array of the data you want to process. The "prompt" parameter is the prompt you want to use to generate the report. The "data" parameter is the variable name you want to save the output to. The "to" parameter in the "email" action is the email address you want to send the report to. The "subject" parameter is the subject of the email. The "body" parameter is the body of the email. The "data" parameter is the variable name of the data you want to send in the email.

        Always the web_prompt variable to generate the prompt for the gpt4 action in the use-cases that involve generating charts and graphs. 
  
  If the user asks to process information, you can use the 'gpt4' action like below:
            Input: analyze the input google sheet and generate a report to email to rohan@cheatlayer.com
            Output:
            

            [start json]
 [{"type": "google_sheets_read", "URL": "https://docs.google.com/spreadsheets/d/1uBTTkIcKM684UvD_qDucQ32rXNeS9Ip20RfatzX5KuI/edit?gid=0#gid=0", "Sheet_Name": "Sheet1", "data": "data_sheet"},
             {"type":"gpt4", "input": ["data_sheet"], "prompt": "Generate a report from the data in the sheet", "data": "report"},
            {"type": "email", "to": "rohan@cheatlayer.com", "subject": "Report", "body": "Here is a report ", "data": "report"}]
            [end json]

            If the user asks to summarize a page directly, use the semanticScrape action like below:
            Input: summarize the page and email it to rohan@cheatlayer.com
            Output:
            [start json]
            [ {"type": "semanticScrape", "target": "summarize all the articles on the page using their descriptions and titles. There are multiple articles across the whole page that are financial articles", "data": "summary"},
            {"type": "email", "to": "rohan@cheatlayer.com", "subject": "Summary", "body": "Here is a summary of the page ", "data": "summary"}]

If you are generating a GPT4 action to generate data for google sheets, output a comma seperated array like ["data1","data2"] without a prefix like json``` or any other prefix. Generate only directly the array.
Generate a single dimensional array and not a double array like [[]] for the google sheets data.

When extracting data from PDFs using gpt, use this pattern:

[{"type": "gpt4", "input": ["pdf"], "prompt": "Extract the name of the candidate from this pdf", "data": "extracted_name"},
{"type": "gpt4", "input": ["pdf"], "prompt": "Extract the years of experience for this candidate", "data": "extracted_years"},
{"type": "google_sheets_create", "URL": "sheet_URL", "Sheet_Name": "sheet_name"},
{"type": "google_sheets_add_row", "URL": "sheet_URL", "Sheet_Name": "sheet_name", "data": ["extracted_name", "extracted_years"]}]


Here is example python code to use to download youtube transcripts:
# importing the module
from youtube_transcript_api import YouTubeTranscriptApi

# retrieve the available transcripts
transcript_list = YouTubeTranscriptApi.list_transcripts('pxiP-HJLCx0')

# iterate over all available transcripts
for transcript in transcript_list:

    # the Transcript object provides metadata
    # properties
    print(
        transcript.video_id,
        transcript.language,
        transcript.language_code,
    
        # whether it has been manually created or
        # generated by YouTube
        transcript.is_generated,
        
        # whether this transcript can be translated
        # or not
        transcript.is_translatable,
        
        # a list of languages the transcript can be
        # translated to
        transcript.translation_languages,
    )

    # fetch the actual transcript data
    print(transcript.fetch())

    # translating the transcript will return another
    # transcript object
    print(transcript.translate('en').fetch())

# you can also directly filter for the language you are
# looking for, using the transcript list
transcript = transcript_list.find_transcript(['en'])

# or just filter for manually created transcripts
transcript = transcript_list.find_manually_created_transcript(['en'])

# importing modules
from youtube_transcript_api import YouTubeTranscriptApi

# using the srt variable with the list of dictionaries
# obtained by the .get_transcript() function
srt = YouTubeTranscriptApi.get_transcript("pxiP-HJLCx0")

# creating or overwriting a file "subtitles.txt" with
# the info inside the context manager
with open("subtitles.txt", "w") as f:

        # iterating through each element of list srt
    for i in srt:
        # writing each element of srt on a new line
        f.write("{}\n".format(i))

        To download youtube videos, please use the only the pytube library. Do NOT use the yt-dlp library to download youtube videos.

        To download titkok videos, you can use the yt_dlp library. For example:

        input: download this youtube video https://www.youtube.com/watch?v=G09W3NwwAVs
        output:
        [start json]
        [{"type": "bash", "command": "pip install pytube"},
        {"type": "python", "code": "from pytube import YouTube\nyt = YouTube('https://www.youtube.com/watch?v=G09W3NwwAVs')\nstream = yt.streams.first()\nstream.download()"}]
        [end json]

        input: download the transcript for this youtube video: https://www.youtube.com/watch?v=G09W3NwwAVs
        output:
        [start json]
        [{"type": "bash", "command": "pip install youtube_transcript_api"},
        {"type": "python", "code": "from youtube_transcript_api import YouTubeTranscriptApi\ntranscript = YouTubeTranscriptApi.get_transcript('G09W3NwwAVs')\nwith open('transcript.txt', 'w') as file:\n    for line in transcript:\n        file.write(line['text'] + '\\n')"}]
        [end json]

        For the click and semanticScrape actions, if the automation involves a browser, add a parameter for 'browser_mode': 'true' to the action. For example:
        Input: click the button that says 'next' on the website
        Output:
        [start json]
        [{"type": "click", "target": "next", "browser_mode": "true"}]
        [end json]


        input: search tiktok for videos related to ai automation and download the video with the most views. Then clip out the 10th to last frame from the video into clip.png.
        output:
        [start json]
        [{"type": "bash", "command": "pip install -U --pre \\"yt-dlp[default]\\""},
        {"type": "python", "code": "import yt_dlp\nydl = yt_dlp.YoutubeDL({'outtmpl': '%(id)s.%(ext)s'})\ninfo = ydl.extract_info('https://www.tiktok.com/@cheatlayer/video/7337267587287354625', download=False)\nurl = info['url']\nydl.download([url])"}]
        [end json]
        
        input: search tiktok for videos related to ai automation and download the video with the most views. Then clip out the 10th to last frame from the video into clip.png. Then open https://app.runwayml.com/video-tools/teams/rohanarun/ai-tools/generative-video in a new tab. Click "drop an image or click upload" and then type clip.png and hit enter. 
        output:
        [start json]
        [{"type": "bash", "command": "pip install -U --pre \\"yt-dlp[default]\\""},
        {"type": "python", "code": "import yt_dlp\nydl = yt_dlp.YoutubeDL({'outtmpl': '%(id)s.%(ext)s'})\ninfo = ydl.extract_info('https://www.tiktok.com/@cheatlayer/video/7337267587287354625', download=False)\nurl = info['url']\nydl.download([url])"},
        {"type": "open tab", "target": "https://app.runwayml.com/video-tools/teams/rohanarun/ai-tools/generative-video", "browser_mode": "true"},
        {"type": "click", "target": "drop an image or click upload", "browser_mode": "true"},
        {"type": "keypress", "target": "clip.png", "browser_mode": "true"}]
        [end json]


        '''
        return total_guidelines
    
    def reRunAutomation(self):
        self.send_message(last_message + " This automation failed and this is the current state of the screen:" + last_state)
    def verifyStep(self, node, goal_type , last_stored_data, thread_signals):
        global verified_steps, last_screenshot
        result = False
        print("verifying step")
        import pyautogui
        screenshot = pyautogui.screenshot()
        if goal_type == "general":
            result = True
        elif goal_type == "schedule":
            result = True
        elif goal_type == "browserScrape":
            if last_stored_data != "noscrape":
                result = True
        elif goal_type == "semanticScrape":
            if last_stored_data != "noscrape":
                result = True
        elif goal_type == "keypress":
             
            import pyautogui
            buf = io.BytesIO()
            screenshot.save(buf, format='JPEG')
            byte_im = buf.getvalue()
            encoded_jpeg = base64.b64encode(io.BytesIO(byte_im).read()).decode('utf-8')
            caption = blip_caption_describe(encoded_jpeg,"If the screen includes this text anywhere, say 'yes' only:" + last_stored_data + ". Otherwise, say exactly 'no' only", graph.user_key, graph.user_plan)
            print(caption)
            result = True
        elif goal_type == "python":
            try:
                print("verifying python")
                log = [{"role": "system", "content": "Take the python code and the output and determine if the code was executed successfully. Output only 'success' if the code was executed successfully and 'failure' if the code was not executed successfully."}]
                headers = {
                    "Content-Type": "application/json",
                }
                mode_gpt = "website2"
                log.append({"role": "user", "content": [{"type": "text", "text":  "Python code: " +node["code"] +  " Python output:" + json.dumps(last_stored_data)}]})
                data = {
                    "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
                    "messages": log,
                    "temperature": 0.7,
                    "max_tokens": 5000,
                    "stream": True  # Enable streaming to match original behavior
                }

                print(data)

                # Call local server instead of remote API
                response = requests.post("http://localhost:8000/v1/chat/completions", 
                                        headers=headers, 
                                        data=json.dumps(data), 
                                        stream=True)

                import pyautogui
                if response.status_code == 200:
                    items = []
                    total = ""
                    for chunk in response.iter_lines():
                        try:
                            if chunk:
                                # Parse SSE format from llama.cpp server
                                chunk_str = chunk.decode('utf-8')
                                if chunk_str.startswith('data: '):
                                    chunk_data = chunk_str[6:]  # Remove 'data: ' prefix
                                    if chunk_data.strip() == '[DONE]':
                                        break
                                    
                                    chunk_json = json.loads(chunk_data)
                                    if 'choices' in chunk_json and len(chunk_json['choices']) > 0:
                                        delta = chunk_json['choices'][0].get('delta', {})
                                        if 'content' in delta:
                                            content = delta['content']
                                            total += content
                                            print(content, end='', flush=True)
                        except Exception as e:
                            # If streaming fails, fall back to non-streaming
                            print(f"Streaming error: {e}")
                            break
                        
                    # If streaming didn't work or total is empty, try non-streaming
                    if not total:
                        data["stream"] = False
                        response = requests.post("http://localhost:8000/v1/chat/completions", 
                                                headers=headers, 
                                                data=json.dumps(data))
                        if response.status_code == 200:
                            result = response.json()
                            total = result["choices"][0]["message"]["content"]
                            print(total)
                        else:
                            print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                            total = "Error communicating with local LLM server"

                    print()  # New line after streaming output
                    print(total)
                else:
                    print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                    total = "Error communicating with local LLM server"
                    summary = total
                    print(summary)
                    if "success" in summary or "Success" in summary:
                        result = True
            except Exception as e:
                print(e)
          
            print("verifying keypress")
        elif goal_type == "click":
            print("verifying click")
            result = True
        elif goal_type == "google_sheets_read":
            if "https" in last_stored_data:
                result = True
        elif goal_type == "google_sheets_create":
            if "https" in last_stored_data:
                result = True
        elif goal_type == "google_sheets_add_row":
            if "https" in last_stored_data:
                result = True
        elif goal_type == "email":
            if "af3j2kdw234" in last_stored_data:
                result = True
        elif goal_type == "api":
            if "API response" in last_stored_data:
                result = True
        elif goal_type == "bash":
            result = True
        elif goal_type == "open tab":
            import pyautogui
            buf = io.BytesIO()
            screenshot.save(buf, format='JPEG')
            byte_im = buf.getvalue()
            encoded_jpeg = base64.b64encode(io.BytesIO(byte_im).read()).decode('utf-8')
            caption = blip_caption_describe(encoded_jpeg,"If the screenshot shows this website is open, say 'yes' only:" + last_stored_data + ". Otherwise, say exactly 'no' only", graph.user_key, graph.user_plan)
            print(caption)
            result = True
        else:
            result = True

        last_screenshot = screenshot
        print(result)
        print(goal_type)
        print(last_stored_data)
        if result == True:
            verified_steps.append({"goal_type": goal_type, "last_stored_data": last_stored_data, "accompished": "yes"})
            thread_signals.notifications.emit("Verifying Step ","Verified Step " + goal_type + ": " + last_stored_data, "icon.png", None, "No")
        else:   
            thread_signals.notifications.emit("Step Failed ","Step Failed " + goal_type + ": " + last_stored_data, "icon.png", None, "No")
        return result
    def process_nodes(self, total, json_output, thread_signals, request, request_id = "0"):
                    global test_iterations, last_input, window2, nodes, graph, graph_nodes, last_message, last_state, previous_steps, verified_steps, last_request_id, zipTarget, zipMode, zipData, zipCounter, zipAgent
                    agent_data = {}
                    block_counter = 0
                    print("PROCESSING NODES")
                    print(json_output)
                    test_iterations = 0
                    for node in json_output:
                        #print(node)
                        last_stored_data = None
                        
                        print("PROCESSING PROJECT ATLAS NODE")
                        print(node)
                        #print("this is a project atlas node")
                        if "type" in node and node["type"] == "delay":
                            import time
                            time.sleep(int(node["time"]))
                            
                            # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
                        if "type" in node and node["type"] == "keypress":
                            print("keypress")
                            #print("lower keypress regular")
                            #print(node["text"])\
                            import pyautogui
                            if node["prompt"] == "enter":
                                pyautogui.press('enter')
                            else:
                                out = node["prompt"]
                                print(graph.global_variables)
                                if "data" in node and node["data"] in graph.global_variables:
                                    out += graph.global_variables[node["data"]]
                                print(out)
                                pyautogui.write(out)
                                print("done")
                            last_stored_data = node["prompt"]
                           # thread_signals.notifications.emit("Keypress ", "Keypress: " + node["prompt"], "icon.png", None, "No")

                        #    thread_signals.notifications.emit("Keypress", node["text"], "icon.png", None, "No")
                            #clipboard.copy(node["text"])
                            #pyautogui.hotkey('ctrl', 'v')
                          #  graph.runNode(graph_nodes,nodes[len(nodes)-1].id, thread_signals, False)

                           # self.mode = "chat"
                        if "type" in node and  node["type"] == "click":
                            #print("process click")
                            x = {"type":"Left Mouse Click", "semanticTarget":node["target"], "x":100, "y":100}

                            #print("process click 1")
                            #print("added click")
                            try:
                                #print("clicking")
                                if "browser_mode" in node:
                                    if node["browser_mode"] == "true":
                                        print("browser mode")
                                        graph.browserClick(node["target"])
                                        time.sleep(15)
                                        last_stored_data = "the target was clicked: " + node["target"]
                                else:
                                    new_max = graph.semanticSearch(node["target"], 100,100)
                                #print(new_max)
                                #print("clicking")
                                    import pyautogui
                                    pyautogui.click(int(new_max['x']), int(new_max['y']))

                           #     graph.QMainWindow.showMinimized()
                                #print("run node click")
                           #     graph.runNode(graph_nodes,nodes[len(nodes)-1].id, thread_signals, False)

                                    self.mode = "chat"
                            except Exception as e:
                                print(e)
                        if "type" in node and node["type"] == "open tab":
                            print("open tab")

                            x = {"type":"Open"}
                            nodes.append(graph.create_node('nodes.widget.ImageNode', name="Open Program " + str(len(nodes)), data=x))#, color= "#FFFFFF"
                              #      nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Open Program " + str(len(nodes)), data=x))
                            program = chrome_location + " " + node["target"]
                            
                            print(program)
                            nodes[len(nodes)-1].create_property('program', program, widget_type=NODE_PROP_QLINEEDIT)
                            nodes[len(nodes)-1].create_property('arguments', "", widget_type=NODE_PROP_QLINEEDIT)
                            nodes[len(nodes)-1].create_property('Automated Input',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)
                        #   nodes[len(nodes)-1].create_property('Automated Lists',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                            nodes[len(nodes)-1].updateImage(resource_path("OpenProgram.png"))

                            nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
                            connectToLastNode(nodes[len(nodes)-1])
                            graph_nodes = graph.serialize_session()
                            import subprocess
                            subprocess.run(program.split(" "))
                            
                            last_stored_data = node["target"]
                            thread_signals.notifications.emit("Open Tab ", "Open Tab: " + last_stored_data, "icon.png", None, "No")

                #           
                           # thread_signals.notifications.emit("Open", program, "icon.png", None, "No")
                            #graph.runNode(graph_nodes,nodes[len(nodes)-1].id, thread_signals, False)
                            self.mode = "chat"
                            #print(program)
                            #print("open program")
                            #print(program)
                        if "type" in node and node["type"] == "google_sheets_create":
                            print("google sheets create")
                            try:

                                sheet_url = graph.sendMessageRTCAsync('google_sheets_create:')
                                
                                graph.global_variables[node["URL"]] = sheet_url["URL"] 
                                graph.global_variables[node["Sheet_Name"]] = sheet_url["Sheet_Name"]
                                print(sheet_url)
                                print("created google sheet")
                                last_stored_data = json.dumps(sheet_url)
                                thread_signals.notifications.emit("Create Google Sheets  ", "Create Google Sheets: " + last_stored_data, "icon.png", None, "No")

                                # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
                                for node in nodes:
                                    print(node.get_property("name") )
                                    if "Magic Scraper" in node.get_property("name"):
                                        graph.global_variables["Magic Scraper Output" + node.get_property("name").split("Magic Scraper ")[1]] = "none"
                                    if "GPT4" in node.get_property("name"):
                                        graph.global_variables["GPT4" + node.get_property("name").split("GPT4")[1]] = "none"
                        
                                x = {"type":"Gsheets"}
                                nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Google Sheets " + str(len(nodes)), data=x))
                                nodes[len(nodes)-1].create_property('URL', "none", widget_type=NODE_PROP_QLINEEDIT)
                                nodes[len(nodes)-1].create_property('Read/Write/Create', "Create",items=["Read", "Write", "Create"] ,  widget_type=NODE_PROP_QCOMBO)
                        
                                nodes[len(nodes)-1].create_property('sheet name', "none", widget_type=NODE_PROP_QLINEEDIT)
                                nodes[len(nodes)-1].create_property('Row 1',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 2',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 3',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 4',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 5',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 6',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 7',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 8',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 9',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 10',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                graph.global_variables["Gsheets"  + str(len(nodes))] = ""
                                nodes[len(nodes)-1].updateImage(resource_path("GoogleSheets.png"))
                        
                        
                                nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
                                if len(nodes) > 0:
                                    connectToLastNode(nodes[-1])
                                      
                                graph.auto_layout_nodes()
                                graph.fit_to_selection()
                            except Exception as e:
                                print(e)
                        if "type" in node and node["type"] == "general":
                            if "description" in node:
                                thread_signals.chat.emit( node["description"] + " here is the current global agent data to use for context:" + json.dumps(graph.global_variables))
                            if "goal" in node:
                                thread_signals.chat.emit( node["goal"] + " here is the current global agent data to use for context:" + json.dumps(graph.global_variables))
                                x = {"type":"general"}
                                nodes.append(graph.create_node('nodes.widget.ImageNode', name="Generalized Agent " + str(len(nodes)), data=x))#, color= "#FFFFFF"
                                nodes[len(nodes)-1].create_property('prompt',node["goal"], widget_type=NODE_PROP_QLINEEDIT)
                                nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
                                if len(nodes) > 0:
                                    connectToLastNode(nodes[-1])
                                nodes[len(nodes)-1].updateImage(resource_path("GeneralTrigger.png"))

                                graph.auto_layout_nodes()
                                graph.fit_to_selection()
                        if "type" in node and node["type"] == "open program":
                            print("open program")
                            x = {"type":"Open"}
                            nodes.append(graph.create_node('nodes.widget.ImageNode', name="Open Program " + str(len(nodes)), data=x))#, color= "#FFFFFF"
                              #      nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Open Program " + str(len(nodes)), data=x))
                            program = node["target"]
                            if "http" in program:
                                program = chrome_location + " " + node["target"]
                                
                            print(program)
                            nodes[len(nodes)-1].create_property('program', program, widget_type=NODE_PROP_QLINEEDIT)
                            nodes[len(nodes)-1].create_property('arguments', "", widget_type=NODE_PROP_QLINEEDIT)
                            nodes[len(nodes)-1].create_property('Automated Input',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)
                     #       nodes[len(nodes)-1].create_property('Automated Lists',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)

                            nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
                            connectToLastNode(nodes[len(nodes)-1])
                            graph_nodes = graph.serialize_session()
                            import subprocess
                            subprocess.run(program.split(" "))

                #    
                           # thread_signals.notifications.emit("Open", program, "icon.png", None, "No")
                            #graph.runNode(graph_nodes,nodes[len(nodes)-1].id, thread_signals, False)
                            self.mode = "chat"
                            #print(program)
                            #print("open program")
                            #print(program)
                        if "type" in node and node["type"] == "schedule":
                            cron = node["cron"]
                            #print("scheduler")
                            #print(cron)
                            starting_website = "https://cheatlayer.com"
                            x = {"type":"Scheduler", "cron":cron}
                            kk = 0  # Use the correct index value for your case
                            schedule = json_output

                            data = {
                                "cron": node["cron"],
                                "start": starting_website,
                                "script": json_output,
                                "goal": request,
                                "name": "scheduled_desktop_agent",
                                "proxy": "",
                                "cookies": "",
                                "location": "Local extension"
                            }

                            try:

                                sheet_url = graph.sendMessageRTCAsync('schedule:' + json.dumps(data))
                                
                            except Exception as e:
                                print(e)
                        if "type" in node and node["type"] == "google_sheets_add_row":
                            print("google sheets add row")
                            try:
                                row_data = []
                                for data in node["data"]:
                                    print(data)
                                    if data in graph.global_variables:
                                       row_data.append(graph.global_variables[data])
                                    else: 
                                       row_data.append(data)    
                                print(row_data)
                                print("sending row")
                                sheet_url = graph.sendMessageRTCAsync('google_sheets_add_row:' + graph.global_variables[node["URL"]] + ":;:" +  graph.global_variables[node["Sheet_Name"]]  + ":;:" + json.dumps(row_data))
                                print(sheet_url)
                                
                                last_stored_data = json.dumps('google_sheets_add_row:' + node["URL"] + ":;:" + json.dumps(row_data)) + "" + sheet_url
                                thread_signals.notifications.emit("Add Row Google Sheets  ", "Add Row Google Sheets: " + json.dumps(row_data), "icon.png", None, "No")
                                import csv
                                csv_file_name = graph.global_variables[node["URL"]].replace("https://", "").replace("/", "_") + ".csv"
                                
                                # Check if the CSV file exists
                                file_exists = os.path.isfile(csv_file_name)
                                
                                with open(csv_file_name, mode='a', newline='') as file:
                                    writer = csv.writer(file)
                                    
                                    # If the file is being created for the first time, add a header row
                                    if not file_exists:
                                        header = node["data"]
                                        writer.writerow(header)
                                    
                                    # Write the row data
                                    writer.writerow(row_data)
                                
                                # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
                                for node in nodes:
                                    print(node.get_property("name") )
                                    if "Magic Scraper" in node.get_property("name"):
                                        graph.global_variables["Magic Scraper Output" + node.get_property("name").split("Magic Scraper ")[1]] = "none"
                                    if "GPT4" in node.get_property("name"):
                                        graph.global_variables["GPT4" + node.get_property("name").split("GPT4")[1]] = "none"
                        
                                x = {"type":"Gsheets"}
                                nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Google Sheets " + str(len(nodes)), data=x))
                                nodes[len(nodes)-1].create_property('URL', graph.global_variables[node["URL"]], widget_type=NODE_PROP_QLINEEDIT)
                                nodes[len(nodes)-1].create_property('Read/Write/Create', "Write",items=["Read", "Write", "Create"] ,  widget_type=NODE_PROP_QCOMBO)
                        
                                nodes[len(nodes)-1].create_property('sheet name',  graph.global_variables[node["Sheet_Name"]] , widget_type=NODE_PROP_QLINEEDIT)
                                nodes[len(nodes)-1].create_property('Row 1',row_data[0],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 2',row_data[1%len(row_data)],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 3',row_data[2%len(row_data)],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 4',row_data[3%len(row_data)],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 5',row_data[4%len(row_data)],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 6',row_data[5%len(row_data)],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 7',row_data[6%len(row_data)],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 8',row_data[7%len(row_data)],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 9',row_data[8%len(row_data)],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 10',row_data[9%len(row_data)],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                graph.global_variables["Gsheets"  + str(len(nodes))] = ""
                                nodes[len(nodes)-1].updateImage(resource_path("GoogleSheets.png"))
                        
                        
                                nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
                                if len(nodes) > 0:
                                    connectToLastNode(nodes[-1])
                                      
                                graph.auto_layout_nodes()
                                graph.fit_to_selection()
                                print(f"Data written to CSV backup file: {csv_file_name}")
                            except Exception as e:
                                print(e)
                        if "type" in node and node["type"] == "python":
                           # thread_signals.notifications.emit("Delay", "Running python code.", "icon.png", None, "No")
                            print("python code")
                            import site
                            import sys
                            packages = site.getusersitepackages()
                            for package in packages:
                                sys.path.append(package)

                            import subprocess

                            def get_global_sitepackages():
                                try:
                                    
                                    subprocess.check_output(['export','LD_LIBRARY_PATH=/usr/local/lib:$LD_LIBRARY_PATH'], universal_newlines=True)
                                    result = subprocess.check_output(['python3', '-c', 'from distutils.sysconfig import get_python_lib; print(get_python_lib())'], universal_newlines=True)
                                    return result.strip()
                                except Exception as e:
                                    print(f"Error: {e}")
                                    print("please install python version 3.10 from the windows store")
                                    return None

                            global_sitepackages = get_global_sitepackages()
                            if global_sitepackages:
                                import sys
                                sys.path.append(global_sitepackages)
                            #bundled_dir = sys._MEIPASS
                            #sys.path.append(bundled_dir)




                            sys.path.append(resource_path(""))
                            print("finish summaries 2")

                            old_stdout = sys.stdout
                            sys.stdout = io.StringIO()
                            print("finish summaries 3")

                            try:
                                if python_process != None:
                                    python_process.terminate()
                                print("finish summaries 4")

# Start a proces            s to run the given code.
                                code = node["code"]
                                if "{{" in code:
                                    #replace any strings that include {{name}} with agent_memory["name"]
                                    print("REPLACING CODE")
                                    print(graph.global_variables)
                                    for key, value in graph.global_variables.items():
                                        if key in code:
                                            code = code.replace("{{" + key + "}}", "agent_memory['" + key + "']")
                                print(code)
                                output = error = None
                                print("running code")
                                exec(code, {'thread_signals':thread_signals, 'notification_manager':graph.notification_manager,'folder': "", 'send_complex_message': graph.send_complex_message, 'agent_memory': graph.global_variables ,'store_analytics': graph.store_analytics, 'user_key': graph.user_key, 'imageToVideo': graph.imageToVideo, 'genSyntheticVideo' : graph.genSyntheticVideo, 'genImage': graph.genImage, 'genVoice': graph.genVoice, 'gpt3Prompt': graph.gpt3Prompt, 'last_mouse_y':0,'last_mouse_x': 0, '__name__': '__main__','videoSearch': graph.videoSearch, 'semanticClick': graph.semanticClick, "keypress": graph.keypress , "semanticMove": graph.semanticMove, "click": graph.click, "move":graph.move, "getData": graph.getData, "sendData": graph.sendData} )
                                print("ran code")
                                output = sys.stdout.getvalue()

                                # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
                                x = {"type":"python"}
                                labels = ["button", "field", "heading", "iframe", "image", "label", "link", "text"]
                                nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Python Code " + str(len(nodes)), data=x))
                                nodes[len(nodes)-1].updateImage(resource_path("CustomCode.png"))
                        
                                nodes[len(nodes)-1].create_property('code', code, widget_type=NODE_PROP_QTEXTEDIT)
                                nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
                                if len(nodes) > 0:
                                    connectToLastNode(nodes[-1])
                                      
                                graph.auto_layout_nodes()
                                graph.fit_to_selection()
                                last_stored_data = output
                                thread_signals.notifications.emit("Python  ", "Python: " + last_stored_data, "icon.png", None, "No")

                                   # try:
#                                   #     self.notification_manager.add_notification("Analyzing Current State", "Analyzing the current state of the screen. Please wait..", "icon.png", lambda: print("Notification closed"))
                                 #   thread_signals.notifications.emit("Analyzing Current State", "Analyzing the current state of the screen. Please wait..", "icon.png")
                                  #  except Exception as e:
                                  ##      print(e)
                                    #    print("Error in code")
                            except Exception as err:
                                import traceback
                                error_class = err.__class__.__name__
                                print(f"An error occurred: {err}")
                                
                                def search_github(query):
                                    url = f"https://api.github.com/search/repositories?q={query}"
                                    response = requests.get(url)
                                    results = json.loads(response.text)
                                    if "items" not in results:
                                        return []
                                    else:
                                       return [f"{result.get('html_url')}" for result in results["items"]][::3]
    
                                def search_stackoverflow(query):
                                    url = f"https://api.stackexchange.com/search/advanced?site=stackoverflow.com&q={query}"
                                    response = requests.get(url)
                                    results = json.loads(response.text)
                                    if "items" not in results:
                                        return []
                                    else:
                                        return [f"{result.get('link')}" for result in results["items"]][::3]
    
                                def extract_content(url):
                                    response = requests.get(url)
                                    text = response.text
                                    return text
    
                                def summarize_text(text):
                                    summary = ""
                                    try:
                                        log = [{"role": "system", "content": "You summarize the website content based on the intended goal to only the most relevant content. Try to find implementation details for python libraries, like the bash install instructions and the python syntax. If there is no relevant content, output only 'none'"}]
    
                                        headers = {
                                            "Content-Type": "application/json",
                                        }
    
                                        mode_gpt = "website2"
                                        log.append({"role": "user", "content": [{"type": "text", "text":  "Goal: " +request +  " website text to summarize:" + text}]})
                                        data = {
                                            "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
                                            "messages": log,
                                            "temperature": 0.7,
                                            "max_tokens": 5000,
                                            "stream": True  # Enable streaming to match original behavior
                                        }

                                        print(data)

                                        # Call local server instead of remote API
                                        response = requests.post("http://localhost:8000/v1/chat/completions", 
                                                                headers=headers, 
                                                                data=json.dumps(data), 
                                                                stream=True)

                                        import pyautogui
                                        if response.status_code == 200:
                                            items = []
                                            total = ""
                                            for chunk in response.iter_lines():
                                                try:
                                                    if chunk:
                                                        # Parse SSE format from llama.cpp server
                                                        chunk_str = chunk.decode('utf-8')
                                                        if chunk_str.startswith('data: '):
                                                            chunk_data = chunk_str[6:]  # Remove 'data: ' prefix
                                                            if chunk_data.strip() == '[DONE]':
                                                                break
                                                            
                                                            chunk_json = json.loads(chunk_data)
                                                            if 'choices' in chunk_json and len(chunk_json['choices']) > 0:
                                                                delta = chunk_json['choices'][0].get('delta', {})
                                                                if 'content' in delta:
                                                                    content = delta['content']
                                                                    total += content
                                                                    print(content, end='', flush=True)
                                                except Exception as e:
                                                    # If streaming fails, fall back to non-streaming
                                                    print(f"Streaming error: {e}")
                                                    break
                                                
                                            # If streaming didn't work or total is empty, try non-streaming
                                            if not total:
                                                data["stream"] = False
                                                response = requests.post("http://localhost:8000/v1/chat/completions", 
                                                                        headers=headers, 
                                                                        data=json.dumps(data))
                                                if response.status_code == 200:
                                                    result = response.json()
                                                    total = result["choices"][0]["message"]["content"]
                                                    print(total)
                                                else:
                                                    print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                                                    total = "Error communicating with local LLM server"

                                            print()  # New line after streaming output
                                            print(total)
                                        else:
                                            print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                                            total = "Error communicating with local LLM server"
                                        summary = total
                                    except Exception as e:
                                        print(e)
                                    return summary
    
                                def process_url(url):
                                    time.sleep(1)
                                    content = extract_content(url)
                                    summary = summarize_text(content)
                                    return (url, summary)
    
                                def search_and_summarize(library_name):
                                    github_query = f"{library_name}"
                                    stackoverflow_query = f"{library_name}"
    
                                    with concurrent.futures.ThreadPoolExecutor() as executor:
                                        github_future = executor.submit(search_github, github_query)
                                        stackoverflow_future = executor.submit(search_stackoverflow, stackoverflow_query)
    
                                        github_results = github_future.result()
                                        stackoverflow_results = stackoverflow_future.result()
    
                                        github_summaries_future = executor.map(process_url, github_results)
                                        stackoverflow_summaries_future = executor.map(process_url, stackoverflow_results)
    
                                        summaries = {
                                            "github": list(github_summaries_future),
                                            "stackoverflow": list(stackoverflow_summaries_future)
                                        }
    
                                    return summaries
    
                                summaries = search_and_summarize(request.split("https://")[0].replace(" ", "%20"))
                                print("GitHub Summaries:")
                                summaries_context = ""
                                for url, summary in summaries['github']:
                                    print(f"URL: {url}\nSummary: {summary}\n")
                                    summaries_context += f"URL: {url}\nSummary: {summary}\n"
                                print("\nStackOverflow Summaries:")
                                for url, summary in summaries['stackoverflow']:
                                    print(f"URL: {url}\nSummary: {summary}\n")
                                    summaries_context += f"URL: {url}\nSummary: {summary}\n"
                                print("Finish summaries")
                                if self.total_errors < 4:
                                    #            thread_signals.chat.emit( "Here is the current goal and a screenshot of the active tab: "+ request + " If the goal has been accomplished based on all the previous steps and the current state of the screen, output only 'goal accomplished', otherwise output the json necessary to perform the next step to reach the goal. Here are the previously completed steps: " + previous_steps + " Here is the current data stored from previous steps:" + json.dumps(variables)[::50000] + ". If this data indicates this step has already been accoplished, do not perform it again and output only goal accomplished")
                                    print(f"A {error_class} error occurred while attempting this goal: " +  request +   ". Here is the previously generated steps and code that failed: " + previous_steps + ".  Please re-generate the JSON to fix it: {err}")
                                    thread_signals.chat.emit(f"A {error_class} error occurred while attempting this goal: " +  request +   ". Here is the previously generated steps and code that failed: " + previous_steps + " Here is some research that may help:" + summaries_context +  ".  Please re-generate the JSON to fix it: {err}")
                                    self.total_errors += 1
                                else:
                                    error = sys.stderr.getvalue()
                                    print(f"An error occurred: {err}")
                                    last_stored_data = "An error occurred running the python code: " + str(err) + ":" +str(error) + " Code: " + code
                            finally:
                                sys.stdout = old_stdout
                                print("finish summaries 5")
                                print(last_stored_data)
                                
                        if "type" in node and node["type"] == "browserScrape":
                            print("magic scraper")
                            print(node["target"])
                            print("gpt-4")
                            last_stored_data = "noscrape"
                            try:

                                browser_elements = graph.sendMessageRTCAsync('semanticSearch:' + node["target"])
                            except Exception as e:
                                print(e)
                            print("browser elements")
                            print(browser_elements)
                            description = node["description"]

                            x = {"type":"SemanticDescribe", "semanticTarget":"caption"}
                            nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Magic Scraper " + str(len(nodes)), data=x))


                            log = [{"role": "system", "content": "You are a helpfull assistant that takes the list of input elements and returns only the array of elements which matches the user's intent. Return a properly formatted array in a json like this:'{\"data\":[data1,data2,data3...]}'.  Only return the json directly and don't add quotes or a prefix like ```json."}]

                            print("sending to semantic search")
                            log.append({"role": "user", "content": "The intended target is " + description  + ". Generate a valid JSON only with double quotes for strings. Only return the json directly and don't add quotes or a prefix like ```json. The list of input elements is:" + json.dumps(browser_elements)[:50000]})
                            data = {
                                "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
                                "messages": log,
                                "temperature": 0.7,
                                "max_tokens": 5000,
                                "stream": True  # Enable streaming to match original behavior
                            }

                            print(data)

                            # Call local server instead of remote API
                            response = requests.post("http://localhost:8000/v1/chat/completions", 
                                                    headers=headers, 
                                                    data=json.dumps(data), 
                                                    stream=True)

                            import pyautogui
                            if response.status_code == 200:
                                items = []
                                total = ""
                                for chunk in response.iter_lines():
                                    try:
                                        if chunk:
                                            # Parse SSE format from llama.cpp server
                                            chunk_str = chunk.decode('utf-8')
                                            if chunk_str.startswith('data: '):
                                                chunk_data = chunk_str[6:]  # Remove 'data: ' prefix
                                                if chunk_data.strip() == '[DONE]':
                                                    break
                                                
                                                chunk_json = json.loads(chunk_data)
                                                if 'choices' in chunk_json and len(chunk_json['choices']) > 0:
                                                    delta = chunk_json['choices'][0].get('delta', {})
                                                    if 'content' in delta:
                                                        content = delta['content']
                                                        total += content
                                                        print(content, end='', flush=True)
                                    except Exception as e:
                                        # If streaming fails, fall back to non-streaming
                                        print(f"Streaming error: {e}")
                                        break
                                    
                                # If streaming didn't work or total is empty, try non-streaming
                                if not total:
                                    data["stream"] = False
                                    response = requests.post("http://localhost:8000/v1/chat/completions", 
                                                            headers=headers, 
                                                            data=json.dumps(data))
                                    if response.status_code == 200:
                                        result = response.json()
                                        total = result["choices"][0]["message"]["content"]
                                        print(total)
                                    else:
                                        print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                                        total = "Error communicating with local LLM server"

                                print()  # New line after streaming output
                                print(total)
                            else:
                                print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                                total = "Error communicating with local LLM server"

                                import re
                                def is_valid_json(json_string):
                                    try:
                                        json.loads(json_string)  # Try parsing the JSON
                                        return True  # Return True if parsing is successful
                                    except json.JSONDecodeError:
                                        return False  # Return False if an error occurs



                                def fix_json(json_string):
                                    # Check and fix missing quotes, brackets, or trailing commas
                                    if not is_valid_json(json_string):
                                        try_fixes = [
                                            json_string + ']}',             # Close with nested structures
                                            json_string + '}',              # Close with object
                                            json_string + ']',              # Close with array
                                            json_string.replace('}{', '},{') # Fix missing commas between objects
                                        ]
                                        for fix in try_fixes:
                                            if is_valid_json(fix):
                                                return fix
                                        return None
                                    return json_string
                                if is_valid_json(total):
                                    print("valid json")
                                else:
                                    print("invalid json")
                                    total = fix_json(total)
                                    print(total)
                        
                                last_stored_data = total
                                thread_signals.notifications.emit("Magic Scraper  ", "Magic Scraper " + last_stored_data, "icon.png", None, "No")

                                graph.global_variables[node["data"]] = json.loads(total)["data"]
                            # Use the in-memory bytes to call `blip_caption` directly
                            caption = "the state of the screen"
                            nodes[len(nodes)-1].create_property('Target In English', description, widget_type=NODE_PROP_QLINEEDIT)
                            nodes[len(nodes)-1].create_property('Automation Input',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["", "phone_transcript", "email_transcript","sms_transcript"] + [""],  widget_type=NODE_PROP_QCOMBO)

                            nodes[len(nodes)-1].create_property('Scrape Browser',"",items=["", "Links","Images", "Text", "Headings", "Audios", "Videos", "Textareas", "Tables", "Divs", "Spans", "Headings",  "Buttons" ],  widget_type=NODE_PROP_QCOMBO)
                            nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
                            if len(nodes) > 0:
                                connectToLastNode(nodes[-1])
                            graph.global_variables["Magic Scraper Output" + str(len(nodes) -1)] = "the state of the screen"
                            graph.global_variables["Magic Scraper All Outputs"] = "The full screen"

                            graph.global_variables["Browser Elements"  + str(len(nodes) -1)] = []
                            graph.global_variables["Browser Texts" + str(len(nodes) -1)] = ""

                            graph.global_variables["Links"  + str(len(nodes) -1)] = []
                            graph.global_variables["Link Texts"  + str(len(nodes) -1)] = ""
                            graph.auto_layout_nodes()
                            graph.fit_to_selection()
                        if "type" in node and node["type"] == "bash":
                            
                            import subprocess
                            #f = StringIO()
                        #    self.thread_signals.notifications.emit("Delay", "Running bash code.", "icon.png", None, "No")

                            #with redirect_stdout(f):
                            print("Running bash code please wait")
                            time.sleep(10)
                            print("SPLIT CODES")
                                #write the code to a file
                            for code in node["command"].split("\n"):
                                if len(code) > 0:
                                    print("CODE")
                                    print(code)
                                    print("SSTART BASH CODE")
                                    subprocess.call(code, shell=True)
                                    #os.system(code)
                                    time.sleep(1)
                                    print("BASH OUTPUT")    
                                    thread_signals.notifications.emit("Bash  ", "Bash: " + code, "icon.png", None, "No")
                                    last_stored_data = code
#                                                                    thread_signals.notifications.emit("Magic Scraper  ", "Magic Scraper " + last_stored_data, "icon.png", None, "No")

                                    #subprocess.call("chmod +x code.sh", shell=True)
                                    #subprocess.call("./code.sh", shell=True)
     #                               os.system("./" + resource_path("code.sh"))
#                            output = os.system(x["code"])
                            print("BASH OUTPUT END")
                            
                            # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
                            x = {"type":"bash"}
                            labels = ["button", "field", "heading", "iframe", "image", "label", "link", "text"]
                            nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Python Code " + str(len(nodes)), data=x))
                            nodes[len(nodes)-1].updateImage(resource_path("CustomCode.png"))
                    
                            nodes[len(nodes)-1].create_property('code', node["command"], widget_type=NODE_PROP_QTEXTEDIT)
                            nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
                            if len(nodes) > 0:
                                connectToLastNode(nodes[-1])
                                  
                            graph.auto_layout_nodes()
                            graph.fit_to_selection()
                            last_stored_data = output
                            #thread_signals.progress.emit(f.getvalue())
                        if "type" in node and node["type"] == "semanticScrape":
                            print("magic scraper")
                            import pyautogui
                            screenshot = pyautogui.screenshot()
                            last_stored_data = "noscrape"
                            buf = io.BytesIO()

                            rgb_screenshot = screenshot.convert("RGB")
                            rgb_screenshot.save(buf, format='JPEG')
                            byte_im = buf.getvalue()

                            encoded_jpeg = base64.b64encode(io.BytesIO(byte_im).read()).decode('utf-8')

                            description = node["target"]
                            caption = blip_caption_describe(encoded_jpeg,description, graph.user_key, graph.user_plan)
                            print(caption)
                            graph.global_variables[node["data"]] = caption
                            print(graph.global_variables)
                            thread_signals.notifications.emit("Magic Scraper ", caption, "icon.png", None, "No")
                            time.sleep(2)
                            last_stored_data = caption

                                                
                            x = {"type":"SemanticDescribe", "semanticTarget":"caption"}
                            
                            nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Magic Scraper " + str(len(nodes)), data=x))
                            nodes[len(nodes)-1].updateImage(resource_path("Describe.png"))
                    
                    
                            # Use the in-memory bytes to call `blip_caption` directly
                            caption = "the state of the screen"
                            nodes[len(nodes)-1].create_property('Target In English', caption, widget_type=NODE_PROP_QTEXTEDIT)
                            nodes[len(nodes)-1].create_property('Automation Input',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)
                    
                            nodes[len(nodes)-1].create_property('Scrape Browser',"",items=["","Links","Images", "Text", "Headings", "Audios", "Videos", "Textareas", "Tables", "Divs", "Spans", "Headings",  "Buttons" ],  widget_type=NODE_PROP_QCOMBO)
                            nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
                            nodes[len(nodes)-1].create_property('Add Node Runs To Prompt',"false",items=["false", "true" ],  widget_type=NODE_PROP_QCOMBO)
                    
                            if len(nodes) > 0:
                                connectToLastNode(nodes[-1])
                            graph.global_variables["Magic Scraper Output" + str(len(nodes) -1)] = "the state of the screen"
                            graph.global_variables["Magic Scraper All Outputs"] = "The full screen"
                    
                            graph.global_variables["Links" + str(len(nodes) -1)] = []
                            graph.global_variables["Browser Elements" + str(len(nodes) -1)] = []
                    
                            for node in nodes:
                                print(node.get_property("name") )
                                if "Google Sheets" in node.get_property("name"):
                                    old_node = node
                                    x = old_node.get_property("Data")
                                    sheet_url = old_node.get_property("URL")    
                                    sheet_name = old_node.get_property("sheet name")
                                    row1 = old_node.get_property("Row 1")
                                    row2 = old_node.get_property("Row 2")
                                    row3 = old_node.get_property("Row 3")
                    
                                    row4 = old_node.get_property("Row 4")
                                    Read_write = old_node.get_property("Read/Write")
                                    row5 = old_node.get_property("Row 5")
                                    row6 = old_node.get_property("Row 6")
                                    row7 = old_node.get_property("Row 7")
                                    row8 = old_node.get_property("Row 8")
                                    row9 = old_node.get_property("Row 9")
                                    row10 = old_node.get_property("Row 10")
                                    name = old_node.get_property("name")
                                    in_connections = []
                                    out_connections = []
                                    data = graph.serialize_session()
                                    old_connections_id = old_node.id
                                    
                    
                                    
                    
                    
                                    new_node = graph.create_node('nodes.widget.ImageNode',     name=name, data=x)
                    
                                    for connection in data["connections"]:
                                        if connection["in"][0] == old_connections_id:
                                            new_node.input(0).connect_to(graph.get_node_by_id(connection["out"][0]).output(0))
                                        if connection["out"][0] == old_connections_id:
                                            graph.get_node_by_id(connection["in"][0]).input(0).connect_to(new_node.output(0))
                                    graph.delete_node(old_node, False)
                                    #remove the node from the list of nodes
                                    nodes.remove(old_node)
                                    new_node.create_property('URL', sheet_url, widget_type=NODE_PROP_QLINEEDIT)
                                    new_node.create_property('Read/Write', Read_write,items=["Read", "Write"] ,  widget_type=NODE_PROP_QCOMBO)
                                    new_node.create_property('sheet name', sheet_name, widget_type=NODE_PROP_QLINEEDIT)
                                    new_node.create_property('Row 1',row1,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                    new_node.create_property('Row 2',row2,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                    new_node.create_property('Row 3',row3,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                    new_node.create_property('Row 4',row4,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                    new_node.create_property('Row 5',row5,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                    new_node.create_property('Row 6',row6,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                    new_node.create_property('Row 7',row7,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                    new_node.create_property('Row 8',row8,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                    new_node.create_property('Row 9',row9,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                    new_node.create_property('Row 10',row10,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                    graph.global_variables["Gsheets"  + str(len(nodes))] = ""
                                    nodes.append(new_node)
                    
                                    new_node.create_property('Data', json.dumps(x, cls=NumpyEncoder))
                        
                            
                            graph.auto_layout_nodes()
                            graph.fit_to_selection()
                        if "type" in node and node["type"] == "google_sheets_read":
                            print("google sheets read")
                            try:

                                sheet_data = graph.sendMessageRTCAsync('google_sheets_read:' + node["URL"] + ";:;" + node["Sheet_Name"])
                                print(sheet_data)
                                last_stored_data = sheet_data
                                graph.global_variables[node["data"]] = sheet_data
                                thread_signals.notifications.emit("Read Google Sheets  ", "Read Google Sheets: " + last_stored_data, "icon.png", None, "No")
     
                                # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
                                for node in nodes:
                                    print(node.get_property("name") )
                                    if "Magic Scraper" in node.get_property("name"):
                                        graph.global_variables["Magic Scraper Output" + node.get_property("name").split("Magic Scraper ")[1]] = "none"
                                    if "GPT4" in node.get_property("name"):
                                        graph.global_variables["GPT4" + node.get_property("name").split("GPT4")[1]] = "none"
                        
                                x = {"type":"Gsheets"}
                                nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Google Sheets " + str(len(nodes)), data=x))
                                nodes[len(nodes)-1].create_property('URL', node["URL"], widget_type=NODE_PROP_QLINEEDIT)
                                nodes[len(nodes)-1].create_property('Read/Write/Create', "Read",items=["Read", "Write", "Create"] ,  widget_type=NODE_PROP_QCOMBO)
                        
                                nodes[len(nodes)-1].create_property('sheet name',  node["Sheet_Name"] , widget_type=NODE_PROP_QLINEEDIT)
                                nodes[len(nodes)-1].create_property('Row 1',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 2',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 3',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 4',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 5',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 6',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 7',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 8',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 9',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                nodes[len(nodes)-1].create_property('Row 10',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                                graph.global_variables["Gsheets"  + str(len(nodes))] = ""
                                nodes[len(nodes)-1].updateImage(resource_path("GoogleSheets.png"))
                        
                        
                                nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
                                if len(nodes) > 0:
                                    connectToLastNode(nodes[-1])
                                      
                                graph.auto_layout_nodes()
                                graph.fit_to_selection()
                            except Exception as e:
                                print(e)
                        if "type" in node and node["type"] == "add_data":
                            print("add data")
                            #[{"type": "add_data", "target": "pdf", "data": "C:\Users\Rohan\Downloads\templates\31\data.pdf"},0
                            fileName = resource_path(node["data"])
                            zipTarget = node["target"]
                            data = ""
                            data_array = [] 
                            #load the file contents if it is a pdf, doc, excel, cvs, or text file
                            file_ext = fileName.split('.')[-1].lower()
                            if file_ext == 'txt':
                                with open(fileName, 'r') as file:
                                    data = file.read()
                            elif file_ext == 'csv':
                                data = pd.read_csv(fileName).to_string()
                            elif file_ext == 'xlsx':
                                data = pd.read_excel(fileName).to_string()
                            elif file_ext == 'docx':
                                doc = docx.Document(fileName)
                                data = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                            elif file_ext in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff']:
                                data_image = self.imageToDataUrl(fileName)
                            elif file_ext == 'pdf':
                                with open(fileName, 'rb') as file:
                                    reader = PdfReader(file)
                                    data = ''
                                    for page_num in range(len(reader.pages)):
                                        if reader.pages[page_num].extract_text() != None:
                                            data += reader.pages[page_num].extract_text()
                                            print(data)
                                    last_stored_data = data
                            #add support for .zip files and load them into an array of files in data
                            elif file_ext == 'zip':
                                zipMode = True
                                zipAgent = json_output
                                zipData = []
                                data_array = []
                                
                                fileName = resource_path(fileName)  # Replace with your ZIP file name
                                temp_dir = resource_path('temp')
                                
                                # Ensure the temp directory exists
                                if not os.path.exists(temp_dir):
                                    os.makedirs(temp_dir)
                                print("loading zip")
                                # Extract ZIP file
                                with zipfile.ZipFile(fileName, 'r') as zip_ref:
                                    zip_ref.extractall(temp_dir)

                                # Read the files inside the extracted directory
                                for file in os.listdir(temp_dir):
                                    fileName = resource_path(os.path.join(temp_dir, file))
                                    
                                    file_ext = fileName.split('.')[-1].lower()
                                    if file_ext == 'txt':
                                        with open(fileName, 'r') as file:
                                            data = file.read()
                                    elif file_ext == 'csv':
                                        data = pd.read_csv(fileName).to_string()
                                    elif file_ext == 'xlsx':
                                        data = pd.read_excel(fileName).to_string()
                                    elif file_ext == 'docx':
                                        doc = docx.Document(fileName)
                                        data = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                                    elif file_ext in ['png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff']:
                                        data_image = self.imageToDataUrl(fileName)
                                    elif file_ext == 'pdf':
                                        with open(fileName, 'rb') as file:
                                            reader = PdfReader(file)
                                            data = ''
                                            for page_num in range(len(reader.pages)):
                                                if reader.pages[page_num].extract_text() != None:
                                                    data += reader.pages[page_num].extract_text()
                                                    print(data)
                                            last_stored_data = data
                                    data_array.append(data)
                                print(data_array)
                                print("done loading")
                                # Assuming the first file is the one we need
                                data = data_array[zipCounter]
                                zipData = data_array
                                print("done loading")
                                print(zipData)

                                graph.global_variables["Download " + str(len(nodes))] = ""
                        
                                x = {"type":"getfiles"}
                                nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Download " + str(len(nodes)), data=x))
                                nodes[len(nodes)-1].create_property('URL or File', resource_path(node["data"]), widget_type=NODE_PROP_QLINEEDIT)
                                nodes[len(nodes)-1].updateImage(resource_path("Download.png"))
                        
                                nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
                                if len(nodes) > 0:
                                    connectToLastNode(nodes[-1])
                                      
                                graph.auto_layout_nodes()
                                graph.fit_to_selection()
                                print("done loading")
                                print(data)
                                # Process the JSON data if applicable
                                # Cleanup the temp directory
                            else:
                                data = "Unsupported file type."
                            graph.global_variables[node["target"]] = data
                            thread_signals.notifications.emit("Add Data  ", "Add Data: ", "icon.png", None, "No")
                        if "type" in node and node["type"] == "gpt4":
                            print("gpt-4")
                            
                            try:
                                print("verifying gpt4")
                                log = [{"role": "system", "content":  node["prompt"]}]
                                print(log)
                                headers = {"Content-Type": "application/json"}
                                mode_gpt = "website2"
                                input_prompt = ""
                                print(log)
                                if "input" in node:
                                    for data_node in node["input"]:
                                        if data_node in graph.global_variables:
                                            input_prompt += json.dumps(graph.global_variables[data_node])
                                    data = {
                                        "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
                                        "messages": log,
                                        "temperature": 0.7,
                                        "max_tokens": 5000,
                                        "stream": True  # Enable streaming to match original behavior
                                    }

                                    print(data)

                                    # Call local server instead of remote API
                                    response = requests.post("http://localhost:8000/v1/chat/completions", 
                                                            headers=headers, 
                                                            data=json.dumps(data), 
                                                            stream=True)

                                    import pyautogui
                                    if response.status_code == 200:
                                        items = []
                                        total = ""
                                        for chunk in response.iter_lines():
                                            try:
                                                if chunk:
                                                    # Parse SSE format from llama.cpp server
                                                    chunk_str = chunk.decode('utf-8')
                                                    if chunk_str.startswith('data: '):
                                                        chunk_data = chunk_str[6:]  # Remove 'data: ' prefix
                                                        if chunk_data.strip() == '[DONE]':
                                                            break
                                                        
                                                        chunk_json = json.loads(chunk_data)
                                                        if 'choices' in chunk_json and len(chunk_json['choices']) > 0:
                                                            delta = chunk_json['choices'][0].get('delta', {})
                                                            if 'content' in delta:
                                                                content = delta['content']
                                                                total += content
                                                                print(content, end='', flush=True)
                                            except Exception as e:
                                                # If streaming fails, fall back to non-streaming
                                                print(f"Streaming error: {e}")
                                                break
                                            
                                        # If streaming didn't work or total is empty, try non-streaming
                                        if not total:
                                            data["stream"] = False
                                            response = requests.post("http://localhost:8000/v1/chat/completions", 
                                                                    headers=headers, 
                                                                    data=json.dumps(data))
                                            if response.status_code == 200:
                                                result = response.json()
                                                total = result["choices"][0]["message"]["content"]
                                                print(total)
                                            else:
                                                print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                                                total = "Error communicating with local LLM server"

                                        print()  # New line after streaming output
                                        print(total)
                                    else:
                                        print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                                        total = "Error communicating with local LLM server"
                                    summary = total
                                    print(summary)
                                    print('done')
                                    last_stored_data = summary
                                    graph.global_variables[node["data"]] = summary
                                    thread_signals.notifications.emit("GPT-4  ", "GPT-4: " + last_stored_data, "icon.png", None, "No")

                                    #, graph
                                    # x = {"type":"Move Mouse", "data":"vaaddriables""x":500,"y":500}
                                    x = {"type":"GPT4"}
                                    graph.global_variables["GPT4" + str(len(nodes))] = ""

                            except Exception as e:
                                print(e)
                        if "type" in node and node["type"] == "email":
                            print("email")
    
                            # Mailgun API URL
                            api_url = f'https://cheatlayer.com/user/sendAgentEmail'

                            # Email details
                            from_email = 'agents@cheatlayer.com'  # Replace with your sending email address
                            to_email = node["to"]  # Replace with the recipient's email address
                            subject = node["subject"]
                            text = node["body"].replace("'", "\'")
                            # Path to the attachment file
                            # add folder to the file 
                            print(node["data"])
                            print(graph.global_variables)   
                            var = graph.global_variables[node["data"]]
                            # Prepare the data for the API request
                            data = {
                                'from': from_email,
                                'to': to_email,
                                'subject': subject,
                                'body': text + str(var),
                                'security':'af3j2kdw234'
                            }
                            last_stored_data = json.dumps(data)
                            # Prepare the attachment file
                            #check the file exists and is not none
                            files = {}

                        
                            print("SEND Cheat Layer")
                            print(data)
                            print(files)

                            # Send the email using the Mailgun API
                            response = requests.post(
                                api_url,
                                data=data
                            )

                            #, graph
                            # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
                            x = {"type":"Email"}
                            nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Gmail " + str(len(nodes)), data=x))
                            nodes[len(nodes)-1].updateImage(resource_path("Email.png"))

                            nodes[len(nodes)-1].create_property('to', node["to"] , widget_type=NODE_PROP_QLINEEDIT)
                            nodes[len(nodes)-1].create_property('to variable',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["", "phone_transcript", "email_transcript","sms_transcript"]  + ["","Loop Total Runs","Loop Node Runs"],  widget_type=NODE_PROP_QCOMBO)

                            nodes[len(nodes)-1].create_property('subject', node["subject"] , widget_type=NODE_PROP_QLINEEDIT)
                            nodes[len(nodes)-1].create_property('file', "none", widget_type=NODE_PROP_QLINEEDIT)

                            nodes[len(nodes)-1].create_property('body',node["body"] , widget_type=NODE_PROP_QLINEEDIT)
                            nodes[len(nodes)-1].create_property('Body variable',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["", "phone_transcript", "email_transcript","sms_transcript"]  + ["","Loop Total Runs","Loop Node Runs"],  widget_type=NODE_PROP_QCOMBO)
                            nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
                            if len(nodes) > 0:
                                connectToLastNode(nodes[-1])

                            graph.auto_layout_nodes()
                            graph.fit_to_selection()
                                                # Check the response from Mailgun
                            if response.status_code == 200:
                                print('Email sent successfully!')
                                
                                thread_signals.notifications.emit("Email  ", "Email: " + last_stored_data, "icon.png", None, "No")
                            else:
                                print(f'Failed to send email: {response.status_code} {response.text}')

                        if "type" in node and node["type"] == "api":
                            print(graph.global_variables)
                            data = ""
                            if "data" in node:
                                data = graph.global_variables[node["data"]] 
                                print("api data")
                            print(data)

                            url = node["URL"]
                            headers = node["headers"]
                            print(url)
                            print(headers)
                            json_data = node["body"]

                            if "messages" in node["body"] and "content" in node["body"]["messages"][1]:
                                if "DATA_" in node["body"]["messages"][1]["content"]:
                                    print("DATA")

                                    data_label = node["body"]["messages"][1]["content"].split("DATA_")[1].split("\"")[0]
                                    new_data = graph.global_variables[data_label]

                                    node["body"]["messages"][1]["content"] = json.dumps(new_data)

                            print(json_data)
                            print("RUNNING API")
                            json_data["data"] = data

                            try:
                                response = requests.post(url, headers=headers, data=json.dumps(json_data))
                                response_data = response.json()
                                print('Success:', response_data)

                                if "output" in node:
                                    if "choices" in response_data:
                                        graph.global_variables[node["output"]] = response_data["choices"][0]["message"]["content"]
                                    else:
                                        graph.global_variables[node["output"]] = json.dumps(response_data)

                            except requests.exceptions.RequestException as error:
                                print('Error:', error)
                            
                            x = {"type":"sendData"}
                            nodes.append(graph.create_node('nodes.widget.ImageNode', name="Send Data " + str(len(nodes)), data=x))#, color= "#FFFFFF"
                            nodes[len(nodes)-1].create_property('Body Key 1', "key 1", widget_type=NODE_PROP_QLINEEDIT)
                            nodes[len(nodes)-1].create_property('Body Value 1',"",  widget_type=NODE_PROP_QLINEEDIT)
                    
                            nodes[len(nodes)-1].create_property('Body Key 2', "key 2", widget_type=NODE_PROP_QLINEEDIT)
                            nodes[len(nodes)-1].create_property('Body Value 2',"",  widget_type=NODE_PROP_QLINEEDIT)
                            
                            nodes[len(nodes)-1].create_property('Body Key 3', "key 3", widget_type=NODE_PROP_QLINEEDIT)
                            nodes[len(nodes)-1].create_property('Body Value 3',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["","Total Runs", "Current Time", "Current Date", "Total Scraped"],  widget_type=NODE_PROP_QCOMBO)
                            
                            nodes[len(nodes)-1].create_property('Body Key 4', "key 4", widget_type=NODE_PROP_QLINEEDIT)
                            nodes[len(nodes)-1].create_property('Body Value 4',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["","Total Runs", "Current Time", "Current Date", "Total Scraped"],  widget_type=NODE_PROP_QCOMBO)
                            
                            nodes[len(nodes)-1].create_property('Body Key 5', "key 5", widget_type=NODE_PROP_QLINEEDIT)
                            nodes[len(nodes)-1].create_property('Body Value 5',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["","Total Runs", "Current Time", "Current Date", "Total Scraped"],  widget_type=NODE_PROP_QCOMBO)
                    
                            nodes[len(nodes)-1].create_property('Request', "", widget_type=NODE_PROP_QLINEEDIT)
                            nodes[len(nodes)-1].create_property('headers', node["headers"], widget_type=NODE_PROP_QLINEEDIT)

                            nodes[len(nodes)-1].create_property('URL', node["URL"], widget_type=NODE_PROP_QLINEEDIT)
                            nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
                            nodes[len(nodes)-1].updateImage(resource_path("SendData.png"))
                    
                            if len(nodes) > 0:
                                connectToLastNode(nodes[-1])
                    
                            graph.auto_layout_nodes()
                            graph.fit_to_selection()
                        last_stored_data = json.dumps(last_stored_data)
                        result = self.verifyStep(node, node["type"] , last_stored_data, thread_signals)
                        if result == True:
                            print("SUCCESS")
                            thread_signals.notifications.emit("Step " + str(block_counter), "This step has been completed successfully.", "icon.png", None, "No")
                        if result == False:
                            #regenerate and run the step again 
                            print("FAILED")
                            prompt = "here is are the list of previously generated steps: " + json.dumps(json_output) + " This step has failed: " + json.dumps(node) + " Please regenerate the automation to fix this step. Common mistakes includes not specifying the correct variables for passing data between nodes or not adding a to wait for websites to load. Common solutions include adjusting the variable names in the 'data' parameter and adding delays."
                            thread_signals.notifications.emit("Regenerating Step " + str(block_counter), "Regenerating Step " + str(block_counter), "icon.png", None, "No")

         
                    if len(last_request_id) > 3:
                        thread_signals.API.emit(verified_steps, str(last_request_id))
                    if False and "accomplished" not in total and test_iterations < 5:
                        print("FAILED? check if the goal has been accomplished")
                        #run a new request with the text that says you are an agent. this is the current state of the screen. What should we do next to accomplish the original goal?
                        #graph.chatInstance.hide_chat_display()
                        test_iterations += 1
                        last_input = request
                    #    graph.notification_manager.add_notification("Step " + block_counter, "Analyzing the current state of the screen. Please wait..", "icon.png", lambda: print("Notification closed"))
                      #  window2.show()
                        print(previous_steps)
                        print(graph.global_variables)
                        variables = graph.global_variables
                        variables["chat_log"] = ""
                        time.sleep(5)
                        print( "Here is the current goal and a screenshot of the active tab: "+request + " If the goal has been accomplished based on all the previous steps and the current state of the screen, output only 'goal accomplished', otherwise output the json necessary to perform the next step to reach the goal. Do not assume the goal has been accomplished on just the previous steps alone, and use the saved data and the current screenshot to verify the goal has been accomplished.  Here are the previously completed steps, : " + previous_steps + " Here is the current data stored from previous steps to verify their outputs:" + json.dumps(graph.global_variables)[::50000] + " Here are the previoulsy verified steps which were completed and their outputs" + json.dumps(verified_steps) + ". ")
                        thread_signals.chat.emit( "Here is the current goal and a screenshot of the active tab: "+ request + " If the goal has been accomplished based on all the previous steps and the current state of the screen, output only 'goal accomplished', otherwise output the json necessary to perform the next step to reach the goal. Here are the previously attempted steps: " + previous_steps + "  Here are the previoulsy verified steps which were completed: " + json.dumps(verified_steps) + ". If this list indicates all the attempted steps have been accomplished, do not perform any actions and output only 'goal accomplished'. The verification process will output accomplished: yes if the steps have been successful. Use these to determine if the goal has been accomplished.")
                    elif "accomplished" in total:   
                        test_iterations = 0
                        print("ACCOMPLISHED ATTEMPT")
                        print(last_request_id)
                        print(verified_steps)
                        time.sleep(2)
                        if len(last_request_id) > 3:
                            thread_signals.API.emit(verified_steps, str(last_request_id))

                        if zipMode:
                            #store th   e data in the global variables from zip files and run processNodes with zipAgent
                            print(zipData)
                            graph.global_variables[zipTarget] = zipData[zipCounter]
                            zipCounter += 1
                            for agent in zipAgent:
                                if agent["type"] == "google_sheets_create":
                                    print("google sheets create")
                                    zipAgent.remove(agent)
                            print(zipAgent)
                            print("ZIPMODE")
                            if zipCounter < len(zipData):
                                #run the next step
                                id_api = 0
                                self.process_nodes(request, zipAgent, thread_signals, request, str(id_api))
                            else:
                                zipMode = False
                                
                       # window2.show()
                       # #window2.show()
                    agent_data = verified_steps
                    return agent_data
    def processChunkThread(self, total,item, request, thread_signals, id_api = "0"):
        global test_iterations, last_input, window2, nodes, graph, graph_nodes, last_message, last_state, previous_steps, verified_steps, last_request_id, zipMode, zipAgent, zipData, zipCounter, python_process, last_stored_data
        print(total)
        print(request)
        print("processing output")
        print(last_input)
        graph.label.hide()
        graph.label2.hide()
        graph.label3.hide()
        #print(request)
        json_output = {}
        last_stored_data = ""
        split_blocks = total.split("[start")
        block_counter = 0
        print(split_blocks)
        for block in split_blocks:
            block_counter += 1
            print(block)
            print(len(block))
            if len(block) > 2:
               # time.sleep(3)
                

#                item[block_counter].label.setText(block)
               # thread_signals.notifications.emit("Step " + str(block_counter), total, resource_path("icon.png"))
                #if "start python" in block or "start bash" in block or "javascript" in block:
             #       thread_signals.progress.emit("Running Code: [start" + block)
              #      thread_signals.progress.emit("code output")
                if "end json" in block:
                    text = block.split("json]")[1].split("[end")[0].strip().rstrip()
                    #print(text)do
                    print(text)
                    print("processing step")
                    #print("PROCESSING PROJECT ATLAS")
                    def fix_json(input_str):
                        import re

                        # Remove trailing commas
                        try:
                            
                            input_str = re.sub(r',\s*([}\]])', r'\1', input_str)

                        # Check for missing quotes around keys
                            input_str = re.sub(r'([{,])(\s*)(\w+)(\s*):', r'\1"\3":', input_str)

                            # Try to load the JSON to see if it is valid now
                            json_data = json.loads(input_str)
                            return json.dumps(json_data, indent=2)
                        except json.JSONDecodeError as e:
                            return f"Error: {e}"
                    try:
                        print("try fix json")
                        json_output = json.loads(fix_json(text))
                    except Exception as e:
                        print(e)
                    print(json_output)
                    if "name" in json_output: 
                            #print("FOUND PREBUILT AGENT")
                            #print(graph.global_variables)
                            for key, value in json_output.items():
                                #print(key)
                                #print(value)
                                graph.global_variables[key] = value
                            #print(graph.global_variables)
                            #print("SAVED GLOBALS")
                            job_folder_name = "general"

                            file_name = json_output["name"]
                            file_path = file_name

                            with open(resource_path(file_path)) as file:
                                data = json.load(file)
                                #print("loaded data")
                                #print(data)
                                current_desktop = 0  # Assuming `current_desktop` is previously defined
                                ##self.thread_signals.recording_start.emit()
                                job_folder_path = os.path.join("general", job_folder_name)
                                for key, node in data["nodes"].items():
                                    #print(node)
                                    custom = json.loads(node["custom"]["Data"])
                                    if custom["type"] == "Start Node":
                                        #print("start node")
                                        # Pass job_folder_path to the runNode function:
                                        graph.runNode(data, key, thread_signals, False, True, folder=job_folder_path)
                                        break
                
                    else:
                        print("try first")
                        agent_data = self.process_nodes(total, json_output, thread_signals, request,  str(id_api))
                        print(agent_data)
                else:
                    print("2nd first")
                    if "accomplished" in total:   
                        test_iterations = 0
                        print("ACCOMPLISHED ATTEMPT")
                        print(last_request_id)
                        print(verified_steps)
                        time.sleep(2)
                        if len(last_request_id) > 3:
                            thread_signals.API.emit(verified_steps, str(last_request_id))

                        if zipMode:
                            print("RUNNING ZIP MODE")
                            print(zipTarget)
                            print(zipData)
                            print(zipCounter)
                            print(zipAgent)
                            #store th   e data in the global variables from zip files and run processNodes with zipAgent
                            if zipTarget in graph.global_variables and zipCounter < len(zipData):
                                graph.global_variables[zipTarget] = zipData[zipCounter]
                            zipCounter += 1
                            if zipCounter < len(zipData):
                                print("RUNNING ZIP")
                                print(zipData)
                                print(zipCounter)
                                print(zipAgent)
                                #run the next step
                                id_api = 0
                                self.process_nodes(request, zipAgent, thread_signals, request, str(id_api))
                            else:
                                zipMode = False
                                
                    #agent_data = self.process_nodes(total, [], thread_signals, request,  str(id_api))
            else:
                print("third first")

                #agent_data = self.process_nodes(total, [], thread_signals, request, str(id_api))
#        item.label.setTextFormat(QtCore.Qt.PlainText)
    def sendRequest(self, msg, thread_signals, system_log):
    #    #print("SENDING REQUEST")
       # #print(msg)
        system_log[0]["content"] = system_log[0]["content"].replace("Current Goal: .", "Current Goal: " + msg["data"]["lastInput"] + ".")
        import pyautogui
        screenshot = pyautogui.screenshot()
        buffer = BytesIO()
        
        rgb_screenshot = screenshot.convert("RGB")
        rgb_screenshot.save(buffer, format='JPEG')
        # Save the screenshot to the buffer in PNG format
        # Seek to the beginning of the buffer
        buffer.seek(0)
        # Read the buffer content into bytes
        image_bytes = buffer.getvalue()
        # Encode the bytes to base64
        base64_image = base64.b64encode(image_bytes)
        # If you need the base64 string, decode the bytes to a string
        base64_string = base64_image.decode('utf-8')
        #log.append({"role": "user", "content": [{"type": "text", "text": "Current goal: " + goal  + " . The screenshot represents the current state of the screen. What is the next step based on the screenshot? Do not generate more than 1 python code block for this step." }, {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_string}"}}]})
                    
        system_log[1]["content"][1]["image_url"]["url"] = f"data:image/jpeg;base64,{base64_string}" 
        global last_input, previous_steps
      #  #print(last_input)
        chat_log = []
        base_log = [] 
        new_base = []
        total_tokens = 0
        code_counter = 0
        thread_signals.progress.emit("Please wait up to 1 minute while I generate the automation for you.")
        thread_signals.progress.emit("You don't need to understand or run the generated code, since I will also execute it next.")
       # graph.chatInstance.toggle_chat_display()
        thread_signals.progress.emit("python")

        from datetime import datetime

# datetime object containing current date and time
        now = datetime.now()
        try:

            response = requests.post(
                "https://cheatlayer.com/user/saveChat",
                data={
                    "input": f"{msg['data']['lastInput']};:;" +
                            f"Desktop;:;" +
                            f"{now}"
                }
            )

            if response.status_code != 200:
                # Handle the error similar to the AJAX error callback
                error_data = {
                    'String': 'CollectMediaFromAccountError',
                    'Request': response.content,
                    'Status': response.status_code,
                    'AjaxError': response.reason
                }
                print(error_data)

            else:
               # print(response)
               print("success")
              #  dataobj = response.json()
               # if len(dataobj) > 1 and user_plan == "free":
               #     for entry in ComPortTikTok:
               #         if entry['sender']['tab']['id'] == port['sender']['tab']['id']:
                           # send_message("IPBLOCK", "data", "not supported", entry)
                #    ipBlock = True
        except Exception as e:
            print(e)
            print("error")
        for kk in range(len(msg['data']['chat_log']) - 1, -1, -1):
            chat_entry = msg['data']['chat_log'][kk]
            total_tokens += len(chat_entry["content"]) / 4
        
            if not "[end HTML]" in chat_entry["content"] and "[start HTML]" in chat_entry["content"]:
                for entry in ComPortTikTok:
                    if entry['sender']['tab']['id'] == port['sender']['tab']['id']:
                        print("HTML error")
                      #    send_message("chatOutput", "data", "There was an error in generating that response. Please try again.", entry)
            elif "<html" in chat_entry["content"] and code_counter < 1:
                code_counter += 1
                if total_tokens < 3600:
                    new_base.append(chat_entry)
            elif "<html" not in chat_entry["content"] and total_tokens < 3600:
                new_base.append(chat_entry)
        
        newinput = msg['data']['chat_log'][0]
        if not new_base:
            newinput["content"] = newinput["content"][:50000]
            new_base.append(newinput)
        
        new_base.reverse()
        base_log += new_base    

        
        headers = {
            "Content-Type": "application/json",
        }
        
        data = {
            "input": system_log,  # Use the correct fallback value for client.exampleInput
            "max_tokens": 5000,
            "id": user_key,
            "key": openaikey,
            "plan": user_plan,
            "mode":"free"
        }
     #   print("SENDING REQUEST 2")
     #   print(data)
        try:

               
            log = [{"role": "system", "content": """
                    
                    You are a professional software engineer who generates all the steps necessary to accomplish a given goal. You generate the steps in a JSON array of steps like shown in the example below based on the user requests. Think through your output step-by-step. First generate a draft of steps, then analyze the draft and improve it by generating a final output which adds more details and takes into consideration interfaces and any missed details. Generate the final draft after the [Final Draft] block of text. For each step, generate the expected output to enable testing the results. Include the language from a choice of either bash, python, and javascript. You have the ability to interleave code from bash, python, and javascript that runs in the browser to accomplish the goals. For example, if the user asks to download a youtube video and scrape the transcript, generate the following:
User: Download a youtube video and scrape the transcript
                    Assistant:
[{“Step Goal”: “Install youtube-dl”, “Step Language”: “bash”,“Step final prompt”: “”, “expected output”:”the console output indicates youtube-dl was installed or already installed”},
{“Step Goal”: “use python to download the youtube video”, “Step Language”: “python”,“Step final prompt”: “”, “expected output”:”console output says it was downloaded or already installed”},
{“Step Goal”: “Open the youtube video and scrape the transcript to a transcript.txt file”, “Step Language”: “javascript”,“Step final prompt”: “”, “expected output”: “the file transcript.txt exists in the downloads folder”}]
                    [Final Draft]
[{“Step Goal”: “Install youtube-dl”, “Step Language”: “bash”,“Step final prompt”: “”, “expected output”:”the console output indicates youtube-dl was installed or already installed”},
{“Step Goal”: “use python to import youtube-dl and download the youtube video”, “Step Language”: “python”,“Step final prompt”: “”, “expected output”:”console output says it was downloaded or already installed”},
{“Step Goal”: “Open the youtube video and scrape the transcript to a string and download that string to a transcript.txt file”, “Step Language”: “javascript”,“Step final prompt”: “”, “expected output”: “the file transcript.txt exists in the downloads folder”},
{“Step Goal”: “open the downloaded video using python”, “Step Language”: “python”,“Step final prompt”: “The console  output indicates it was opened”}]
                    
                    This is just an example and you should generate new steps for each request. Only use use a language if it is necessary, for example do not use Javascript if you can accomplish the goal entirely using bash and python. Always generate the steps to install the ncessary python librarie. 
                    """}]
            log.append({"role": "user", "content": msg['data']['lastInput'] + " File Data:" +  self.textEdit})
            last_input = msg['data']['lastInput']
            print(log)
            print("sent log")
            #data2 = {
            #     "input": log,  # Use the correct fallback value for client.exampleInput
            #     "max_tokens": 5000,
            #     "id": user_key,
            #     "key": "",
            #     "plan": user_plan,
            #     "mode":"compress"
            # }
#
            #response = requests.post("https://streaming-16k.vercel.app/api/request", headers=headers, data=json.dumps(data2), stream=True)
            #
            #if response.status_code == 200:
         #
            #    total = ""
            #    for chunk in response.iter_content(chunk_size=1024):
            #        total += chunk.decode('utf-8')
            #        print(chunk.decode('utf-8'))
            #    print("START STEPS")
            #   # steps = json.loads(total.split("[Final Draft]")[1])
            #    print(total.split("[Final Draft]")[1])
                #    print("STEPS")
            data = {
                "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
                "messages": log,
                "temperature": 0.7,
                "max_tokens": 5000,
                "stream": True  # Enable streaming to match original behavior
            }
            
            print(data)
            
            # Call local server instead of remote API
            response = requests.post("http://localhost:8000/v1/chat/completions", 
                                    headers=headers, 
                                    data=json.dumps(data), 
                                    stream=True)
            
            import pyautogui
            if response.status_code == 200:
                items = []
                total = ""
                for chunk in response.iter_lines():
                    try:
                        if chunk:
                            # Parse SSE format from llama.cpp server
                            chunk_str = chunk.decode('utf-8')
                            if chunk_str.startswith('data: '):
                                chunk_data = chunk_str[6:]  # Remove 'data: ' prefix
                                if chunk_data.strip() == '[DONE]':
                                    break
                                
                                chunk_json = json.loads(chunk_data)
                                if 'choices' in chunk_json and len(chunk_json['choices']) > 0:
                                    delta = chunk_json['choices'][0].get('delta', {})
                                    if 'content' in delta:
                                        content = delta['content']
                                        total += content
                                        print(content, end='', flush=True)
                    except Exception as e:
                        # If streaming fails, fall back to non-streaming
                        print(f"Streaming error: {e}")
                        break
                    
                # If streaming didn't work or total is empty, try non-streaming
                if not total:
                    data["stream"] = False
                    response = requests.post("http://localhost:8000/v1/chat/completions", 
                                            headers=headers, 
                                            data=json.dumps(data))
                    if response.status_code == 200:
                        result = response.json()
                        total = result["choices"][0]["message"]["content"]
                        print(total)
                    else:
                        print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                        total = "Error communicating with local LLM server"
                
                print()  # New line after streaming output
                print(total)
            else:
                print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                total = "Error communicating with local LLM server"
        
                chat_log.append({"role": "assistant", "content": total})
                total_steps = len(previous_steps.split('end step'))
                previous_steps += "[step f{total_steps}]"  + total + "[end step]"
                thread_signals.notifications.emit("Accomplished!", total, "icon.png", None, "No")
#                for step in total.split("[start"):
                
 #                   thread_signals.progress.emit("Running Code: [start " + step)
                # Iterate over the response in ch   unks
                # Iterate over the response in ch   unks
             #   recorder.start(last_input)

       #         x = threading.Thread(target=self.processChunkThread, args=[total, items, msg['data']['lastInput'],  thread_signals]).start()
                self.processChunkThread(total, items, msg['data']['lastInput'], thread_signals)
                #result = graph.verifyRequest(msg['data']['lastInput'])
                #result_text = "The request was verified using Atlas-1 with a probability of success of %.4f"%result
                #+ "% . If you would like Atlas-1 to re-attempt the automation with this extra knowledge, click the button below."
             
                #item2 = ReceiverMessageItem("The request was verified using Atlas-1 with a probability of success of %.4f"%result)
                #self.chat_display.addItem(item2)
                #self.chat_display.setItemWidget(item2, item2.widget)
                #self.chat_display.sortItems(reverse = True)
                #self.chat_display.scrollToBottom()  # Scroll to the botto
                ##add a button to the chat_display to re-run the automation
                #button = QPushButton("Re-run automation")
                #button.clicked.connect(self.reRunAutomation)
                #item3 = ReceiverMessageItem("")
                #self.chat_display.addItem(item3)
                #self.chat_display.setItemWidget(item3, button)
                #self.chat_display.sortItems(reverse = True)
                #self.chat_display.scrollToBottom()  # Scroll to the botto
               #  Replace the next block with the equivalent logic for your Python application.
          
               # If you need to send progress to other entities, use the appropriate method in Python.
                
                # Building the fine_tuning_update string
              #  fine_tuning_update = '{"messages": ['
              #  for log_entry in json.loads(input_value):  # Assuming 'input' is a JSON string.
              #      fine_tuning_update += json.dumps(log_entry) + ","
              #  fine_tuning_update += '{"role": "assistant", "content": "' + total + '"}]'
              #  print(fine_tuning_update)
                # Process the fine_tuning_update string as needed
                # 
      
        except Exception as e:
            # Handle the exception and add your logic here.
           # self.sendRequest(msg, thread_signals, system_log)
            print(e)
            """
            for kk in range(len(ComPortTikTok)):
                # Logic using ComPortTikTok and SendMessage in the context of an error
            """
    
    def addResponse(self, response, image):
        #print("add response")
        #print(response)
        #print("add response")
        ##print(image)
        #print("response")
        #print(response)
        item3 = UserMessageItem(response)
        #print("created")

        self.chat_display.addItem(item3)
        #print("added")

        self.chat_display.setItemWidget(item3, item3.widget)
        #print("setWidget")

        self.chat_display.scrollToBottom()  # Scroll to the bottom
        #print("finsihed")

    def upload_recording(self):
       # recorder.upload_to_gcs()
        #run the upload using a thread 
        x = threading.Thread(target=recorder.upload_to_gcs).start()
        self.send_message_bot("The automation is being uploaded to the community. Thank you for sharing your knowledge! Please wait up to 5 minutes.")
        
#        self.addButton.hide()
    def send_update_bot(self, input_text=None):
        global last_bot_message, total_last_message
        total_last_message += input_text
        #check if the height of the last message is greater than 300

        if last_bot_message != None and last_bot_message.label.height() > 300:
            
            
            self.send_message_bot('...', 'runtime')
            total_last_message = ""
        


        if last_bot_message != None and input_text.strip():
            #display only the last 100 characters
            last_bot_message.label.setText(total_last_message)
        
            last_bot_message.label.setStyleSheet("""
                background-color: #000000;
                color: #FFFFFF;
                font-family: 'Consolas', 'Monaco', 'Courier New', monospace;
                padding: 5px;
                border-radius: 5px;
                max-width: 300px;
            """)
            # Use a syntax highlighter if the text is code
            self.highlighter = SimpleSyntaxHighlighter(last_bot_message.label)

            last_bot_message.setSizeHint(last_bot_message.widget.sizeHint())

            self.chat_display.updateGeometry()
            self.chat_display.update()  # Force an update to the widget
            self.chat_display.scrollToBottom()

    def send_message_bot(self, input_text=None, mode = "chat"):
        global last_bot_message, total_last_message
        text = self.entry_field.text()
        
        if input_text:
            text = input_text
        last_message = text
        total_last_message = ""
        if text.strip():
            if mode == "code":
                last_bot_message = ReceiverMessageItem("python")
            elif mode == "runtime" or "node" in text or "http" in text:
                last_bot_message = ReceiverMessageItem("runtime")
            else:
                last_bot_message = ReceiverMessageItem(text)
            last_bot_message.setSizeHint(last_bot_message.widget.sizeHint())
            self.chat_display.addItem(last_bot_message)
            self.chat_display.setViewMode(QListView.ListMode)  # or QListView.IconMode

            self.chat_display.setItemWidget(last_bot_message, last_bot_message.widget)
            self.chat_display.doItemsLayout()

            # Now force the list widget to update its geometry
            self.chat_display.updateGeometry()
            self.chat_display.update()  # Force an update to the widget
            self.chat_display.scrollToBottom()  # Scroll to the bottom
            QTimer.singleShot(0, self.chat_display.scrollToBottom)


    # In some cases, you might want to call these as well:
    # self.chat_display.adjustSize()
    # QApplication.processEvents()  # Process any pending events that may be affecting 
            self.entry_field.clear()
    def stop_screen_recorder(self, input_text=None):
      #  recorder.stop()
         # Create a button
        addButton = UploadButton(self.upload_recording)
        self.chat_display.addItem(addButton)
        self.chat_display.setItemWidget(addButton, addButton.widget)
        self.chat_display.sortItems(reverse = True)
        self.chat_display.scrollToBottom()  # Scroll to the botto
        #self.layout.addWidget(self.addButton)
    def start_screen_recorder(self, input_text=None):
        print("Start screen recorder")
      #  recorder.start()


    def setkey(self, input_text=None):
        global openaikey, user_plan, user_key, guidelines, desktop_guidelines

        user_key = input_text
    def scraper(self, email, url):
        run_all_scrapers(email, url)
    def notification(self, title="notification",input_text="notification", image="icon.png", lambatest=lambda: print("Notification closed"), prompt=""):
        global openaikey, user_plan, user_key, guidelines, desktop_guidelines
        print("notification posted")
        self.monitor.update_activity()


        main(firebase_email, input_text)
        graph.notification_manager.add_notification(title, input_text, image, lambatest, prompt)
        old_text = graph.BottomToolbar.consoleOutput.toPlainText()
        graph.BottomToolbar.consoleOutput.setText(old_text + title + ": " + input_text)

    def reconnect(self, input_text=None):
        global openaikey, user_plan, user_key, guidelines, desktop_guidelines
        #print("reconnecting")
        graph.restartServer()


    def send_message(self, input_text=None):

        global last_message, chat_log, last_state, desktop_guidelines, user_key, total_errors, previous_steps
        text = self.entry_field.text()
        #slowly animate enlarging the window to 700 height in 2 seconds
        self.chat_display.show()
        #print("SENDING MESSAGE")
        #print(input_text)
      #  self.resize(self.width(), 700)
     #   self.move(self.x(), self.y() - 100)
        
        if input_text:
            text = input_text
        if "error" in text:
            total_errors += 1
            self.mode = "error"

        if total_errors > 3:
            total_errors = 0
            return
        last_message = input_text
        #print(user_key)
        if text.strip():
            item2 = UserMessageItem(text)
            self.toggle_chat_display()

            self.chat_display.addItem(item2)
            self.chat_display.setItemWidget(item2, item2.widget)
            self.chat_display.scrollToBottom()  # Scroll to the bottom
            self.entry_field.clear()
            #print(openaikey)
            if False:
                ##print("website generator")
                out_guidelines = self.find_guidelines(text, guidelines)
            
                msg_data_mode = "website" # replace this with actual value
                faq_prompt = "some value"   # replace this with actual value
                
                auto_prompt = ""
                max_tokens = None
                
                
                log = [{"role": "system", "content": out_guidelines}]
     #           chat_log.append({"role": "user", "content": text})
             #   graph.chatInstance.showMinimized()
                #print(log)
                #print("system log")
                def encode_image(image_path):
                  with open(image_path, "rb") as image_file:
                    return base64.b64encode(image_file.read()).decode('utf-8')
                mode_gpt = "32k"
                #if text includes the keyword 'screenshot', take a screenshot with pyauto-gui and reformat chat_log to use the following format to support images and text. 
                #"messages": [ { "role": "user", "content": [ { "type": "text", "text": "What’s in this image?" }, { "type": "image_url", "image_url": { "url": "https://upload.wikimedia.org/wikipedia/commons/thumb/d/dd/Gfp-wisconsin-madison-the-nature-boardwalk.jpg/2560px-Gfp-wisconsin-madison-the-nature-boardwalk.jpg" } } ] } ],
               
                #print("taking agent screenshot")
               # graph.chatInstance.showMinimized()
                screenshot = pyautogui.screenshot()
                buffer = BytesIO()
                # Save the screenshot to the buffer in PNG format
                screenshot.save(buffer, format="JPEG", quality=50)  # Here is where we save as JPEG
                # Seek to the beginning of the buffer
                buffer.seek(0)
                # Read the buffer content into bytes
                image_bytes = buffer.getvalue()
                # Encode the bytes to base64
                base64_image = base64.b64encode(image_bytes)
                # If you need the base64 string, decode the bytes to a string
                base64_string = base64_image.decode('utf-8')
                mode_gpt = "website2"
                ##print("SCRNEENSHOT MODE")
                goal = text
                if "accomplish this original goal:" in text:
                    goal = text.split("accomplish this original goal:")[1]
                log.append({"role": "user", "content": [{"type": "text", "text": "Current goal: " + goal  + " . The screenshot represents the current state of the screen. What is the next step based on the screenshot? Do not generate more than 1 python code block for this step. Here are the previously executed steps:" + previous_steps }, {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_string}"}}]})
                ##print("screenshot")
        
           #     graph.chatInstance.toggle_chat_display()

                base_log = [] # assuming you have this or fetch it from somewhere else
                log_process = log + chat_log
                # Function to filter out all but the last user message
                def keep_last_user_message(log_list):
                    user_messages = [message for message in log_list if message.get("role") == "user"]
                    if user_messages:
                        last_user_message = user_messages[-1]  # Get the last user message
                        # Filter out all but the last user message
                        filtered_log = [message for message in log_list if message.get("role") != "user" or message is last_user_message]
                    else:
                        filtered_log = log_list  # If there are no user messages, don't change the list
                    return filtered_log
                
                # Apply the function to the log_process
                log_process = keep_last_user_message(log_process)
                msg = {
                    'data': {
                        'lastInput': text,
                        'url': 'Desktop',
                        'createdAt': 'example_date',
                        'chat_log': log_process,
                        'mode': mode_gpt,
                        'agent_id': 'example_agent_id',
                        'task_id': 'example_task_id'
                    }
                }
                ##print("SENT MESSAGE")
                ##print(log_process)
                #print("SENT MESSAGE")
              #  #print(log_process)
               # window2.hide()
               # graph.QMainWindow.showMinimized()

                x = threading.Thread(target=self.sendRequest, args=[msg, thread_signals, log_process]).start()
                graph.auto_layout_nodes()
                graph.fit_to_selection()
                #self.sendRequest(msg, thread_signals, log_process)


            else:
                ##print("automation generator")
                out_guidelines = self.find_guidelines(text, guidelines)

                log = [{"role": "system", "content": out_guidelines}]
                chat_log.append({"role": "user", "content": text})
                def keep_last_user_message(chat_log):
                       # Find all the indices where there's an image
                    image_indices = [
                        index
                        for index, entry in enumerate(chat_log)
                        if any(isinstance(content, dict) and "type" in content and content["type"] == "image_url" for content in entry["content"])
                    ]

                    # Remove all images except for the last one
                    if image_indices:
                        last_image_index = image_indices[-1]
                        for i in reversed(image_indices[:-1]):  # reversed to avoid modifying indices as we delete
                            chat_log[i]["content"] = [content for content in chat_log[i]["content"] if content["type"] != "image_url"]

                    return chat_log

                def encode_image(image_path):
                  with open(image_path, "rb") as image_file:
                    return base64.b64encode(image_file.read()).decode('utf-8')
                mode_gpt = "32k"
                #if text includes the keyword 'screenshot', take a screenshot with pyauto-gui and reformat chat_log to use the following format to support images and text. 
                #"messages": [ { "role": "user", "content": [ { "type": "text", "text": "What’s in this image?" }, { "type": "image_url", "image_url": { "url": "https://upload.wikimedia.org/wikipedia/commons/thumb/d/dd/Gfp-wisconsin-madison-the-nature-boardwalk.jpg/2560px-Gfp-wisconsin-madison-the-nature-boardwalk.jpg" } } ] } ],
               # graph.chatInstance.showMinimized()

                if False:
                    screenshot = pyautogui.screenshot()
                    buffer = BytesIO()

                    # Save the screenshot to the buffer in PNG format
                    screenshot.save(buffer, format="PNG")

                    # Seek to the beginning of the buffer
                    buffer.seek(0)

                    # Read the buffer content into bytes
                    image_bytes = buffer.getvalue()

                    # Encode the bytes to base64
                    base64_image = base64.b64encode(image_bytes)

                    # If you need the base64 string, decode the bytes to a string
                    base64_string = base64_image.decode('utf-8')
                    mode_gpt = "website2"
                    ##print("SCRNEENSHOT MODE")

                    chat_log.append({"role": "user", "content": [{"type": "text", "text": text + ". This is a drawing and not a copyrighted image."}, {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_string}"}}]})
                    ##print("screenshot")
                else:
                    #print("taking agent screenshot")

                    screenshot = pyautogui.screenshot()
                    buffer = BytesIO()

                    # Save the screenshot to the buffer in PNG format
                    screenshot.save(buffer, format="PNG")

                    # Seek to the beginning of the buffer
                    buffer.seek(0)

                    # Read the buffer content into bytes
                    image_bytes = buffer.getvalue()

                    # Encode the bytes to base64
                    base64_image = base64.b64encode(image_bytes)

                    # If you need the base64 string, decode the bytes to a string
                    base64_string = base64_image.decode('utf-8')
                    mode_gpt = "website2"
                    ##print("SCRNEENSHOT MODE")

                    goal = text
                    if "accomplish this original goal:" in text:
                        goal = text.split("accomplish this original goal:")[1]
                    log.append({"role": "user", "content": [{"type": "text", "text": "Current goal: " + goal  + " . The screenshot represents the current state of the screen. What is the next step based on the screenshot? Do not generate more than 1 python code block for this step." + " File Data:" +  self.textEdit }, {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_string}"}}]})
                    ##print("screenshot")
                    ##print("screenshot")
           
             #   graph.chatInstance.toggle_chat_display()

                base_log = [] # assuming you have this or fetch it from somewhere else
                log_process = log + chat_log
                log_process = keep_last_user_message(log_process)

                msg = {
                    'data': {
                        'lastInput': text,
                        'url': 'Desktop',
                        'createdAt': 'example_date',
                        'chat_log': log_process,
                        'mode': mode_gpt,
                        'agent_id': 'example_agent_id',
                        'task_id': 'example_task_id'
                    }
                }
                ##print("SENT MESSAGE")
                ##print(log_process)
                #print("SENT MESSAGE")
      #          #print(log_process)
                window2.hide()
                #d()

                x = threading.Thread(target=self.sendRequest, args=[msg, thread_signals, log_process]).start()
                graph.auto_layout_nodes()
                graph.fit_to_selection()
               
              #  self.sendRequest(msg, thread_signals, log_process)

                
                
            
    def mousePressEvent(self, event):
        self.drag_position = event.globalPos() - self.frameGeometry().topLeft()
        padding = 10  # Defined in your layout
        edgeTolerance = 0

        xPos = event.pos().x()
        yPos = event.pos().y()

        # Adjusted for padding
        if xPos <= padding + edgeTolerance:
            if yPos <= padding + edgeTolerance:
                self.resize_direction = "top-left"
            elif self.height() - yPos <= padding + edgeTolerance:
                self.resize_direction = "bottom-left"
            else:
                self.resize_direction = "left"
        elif self.width() - xPos <= padding + edgeTolerance:
            if yPos <= padding + edgeTolerance:
                self.resize_direction = "top-right"
            elif self.height() - yPos <= padding + edgeTolerance:
                self.resize_direction = "bottom-right"
            else:
                self.resize_direction = "right"
        elif yPos <= padding + edgeTolerance:
            self.resize_direction = "top"
        elif self.height() - yPos <= padding + edgeTolerance:
            self.resize_direction = "bottom"
        else:
            self.resize_direction = None

        self.is_dragging = True


    def mouseMoveEvent(self, event):
        if self.is_dragging and not self.resize_direction:
            self.move(event.globalPos() - self.drag_position)
        elif self.resize_direction:
            if self.resize_direction.endswith("left"):
                dx = event.globalPos().x() - self.x()
                self.setGeometry(self.x() + dx, self.y(), self.width() - dx, self.height())
            if self.resize_direction.endswith("right"):
                self.resize(event.x(), self.height())
            if self.resize_direction.startswith("top"):
                dy = event.globalPos().y() - self.y()
                self.setGeometry(self.x(), self.y() + dy, self.width(), self.height() - dy)
            if self.resize_direction.startswith("bottom"):
                self.resize(self.width(), event.y())

    def mouseReleaseEvent(self, event):
        self.is_dragging = False
        self.resize_direction = None
        self.drag_position = None

class WelcomeWindow(QMainWindow):
    base_questions = []
    custom_questions = []
    prompts = []
    sites = []
    
    def onNextFrame(self):
        pixmap = self.movie.currentPixmap()
        self.pic.setPixmap(pixmap)

    def center(self):
        qr = self.frameGeometry()
        cp = QDesktopWidget().availableGeometry().center()
        qr.moveCenter(cp)
        self.move(qr.topLeft())

    def mousePressEvent(self, event):
        self.oldPos = event.globalPos()

    def mouseMoveEvent(self, event):
        delta = QPoint (event.globalPos() - self.oldPos)
        self.move(self.x() + delta.x(), self.y() + delta.y())
        self.oldPos = event.globalPos()
    def onTrayIconActivated(self, reason):
        if reason == QSystemTrayIcon.Trigger:
            self.show()
            self.tray_icon.hide()
    def __init__(self, base_questions, custom_questions, sites, prompts):
        super().__init__(None, Qt.WindowStaysOnTopHint)

        self.title = "Project Atlas"
        self.top = 0
        self.left = 0
        self.setWindowFlag(Qt.FramelessWindowHint)
        self.setStyleSheet("background-color: rgba(255, 255, 255, 250)")
        # Init QSystemTrayIcon
        self.tray_icon = QSystemTrayIcon(self)
        self.tray_icon.setIcon(QIcon( "favicon.ico"))

        '''
            Define and add steps to work with the system tray icon
            show - show window
            hide - hide window
            exit - exit from application
        '''
        show_action = QAction("Show", self)
        quit_action = QAction("Exit", self)
        hide_action = QAction("Hide", self)
        show_action.triggered.connect(self.show)
        hide_action.triggered.connect(self.hide)
        quit_action.triggered.connect(QApplication.instance().quit)
        tray_menu = QMenu()
        tray_menu.addAction(show_action)
        tray_menu.addAction(hide_action)
        tray_menu.addAction(quit_action)
        self.tray_icon.setContextMenu(tray_menu)
        self.tray_icon.show()

        self.setAttribute(Qt.WA_TranslucentBackground, True)
        self.setWindowOpacity(0.9)
        self.mode = "chat"
        self.input_queue = []
        screen = QApplication.primaryScreen()
        ###print('Screen: %s' % screen.name())
        size = screen.size()
        ###print('Size: %d x %d' % (size.width(), size.height()))
        rect = screen.availableGeometry()
        ###print('Available: %d x %d' % (rect.width(), rect.height()))
        self.width = size.width()
        self.setFixedWidth(500)
        self.setFixedHeight(300)

        self.openAIKey = ''
        self.messages = QListView(self)
       # self.messages.setVerticalScrollMode(QAbstractItemView.ScrollPerPixel)
        self.messages.resize(int(400), int(200))

#       
        self.messages.move(50,50)

        # Use our delegate to draw items in this view.
        self.messages.setItemDelegate(MessageDelegate())

        # scroll bar
        self.scroll_bar = QScrollBar(self)

        # setting style sheet to the scroll bar
        self.messages.setStyleSheet("QListView::item {font-family: Times New Roman;font-style: italic;font-size: 26pt;font-weight: bold;border-bottom-right-radius:10px;border-bottom-left-radius:10px;}; p {font-size:24px;font-color:  rgb(255, 255, 255);}")

        # setting vertical scroll bar to it
        self.messages.setVerticalScrollBar(self.scroll_bar)



        self.model = MessageModel()

        self.messages.setModel(self.model)
        self.messages.show()
        self.status_txt = QLabel(self)
        movie = QtGui.QMovie(resource_path("pa.gif"))
        self.status_txt.setMovie(movie)
        movie.start()
        self.model.add_message(USER_THEM, "Welcome to Project Atlas! I'm here to help you build any automation you need. Start by asking what you want in natural language. \n\n\nFor example, you can say \n'Write a post about motivational quotes to twitter'\n or 'Write a blog post about dogs'\n or 'Scrape amazon for tent prices and reviews'\n or 'Post an image about AI automation to the cheatlayer facebook group'") 

        self.status_txt.move(50,0)
        self.status_txt.show() # You were missing this.
        
        self.status_txt.setFixedSize(QSize(50,50))

        self.base_questions = base_questions
        self.custom_questions = custom_questions
        self.sites = sites
        self.prompts = prompts
        self.height = size.height()
        self.train_folder = "train/"
        self.test_folder = "test/"
        #self.labels = ["Facebook Post", "Facebook Comment"]
        self.output = None
        self.labels = ["button", "field", "heading", "iframe", "image", "label", "link", "text"]
        self.test_image = resource_path("test_image2.png")
        self.model_file = resource_path("model_weights_default.pth")

        self.pushButton = QPushButton("", self)
        self.pushButton.move(400, 0)
        self.pushButton.setToolTip("<h3>Generate Automation</h3>")
        self.pushButton.resize(32, 32)
        self.pushButton.setIcon(QIcon( resource_path("GenerateAutomation.png")))
        self.splash = MySplashScreen('Loader1.gif', Qt.WindowStaysOnTopHint)

        self.pushButton.setIconSize(QSize(50, 50))
        self.pushButton.setFixedSize(QSize(50,50))

        
        # set qmovie as label

#        self.speakButton = QPushButton("", self)
 #       self.speakButton.move(size.width()/2- 250, 0)
  #      self.speakButton.setToolTip("<h3>Speak</h3>")
   #     self.speakButton.resize(32, 32)

        self.line = QLineEdit(self)
        self.line.returnPressed.connect(self.atlas)
        self.line.setStyleSheet("border: 0px;")

        self.line.move(100, 0)
        self.line.resize(300, 50)
        self.font = self.line.font()      # lineedit current font
        self.font.setPointSize(24)               # change it's size
        self.line.setFont(self.font)      # set font


        self.pushButton.clicked.connect(self.atlas) 
      #  self.speakButton.clicked.connect(self.speak) 
        
     #   self.speakButton.setIcon(QIcon( "Speak.png"))
    #    self.speakButton.setIconSize(QSize(50, 50))
    #    self.speakButton.setFixedSize(QSize(50, 50))

        self.atlas_array = []
        self.atlas_icons = []
        site_counter = 0
        
        self.main_window()
    def keyPressEvent(self, e):  
        if e.key() == QtCore.Qt.Key_Escape:
            self.close()
        if e.key() == QtCore.Qt.Key_F11:
            if self.isMaximized():
                self.showNormal()
                
                self.hide()
               # self.tray_icon.showMessage(
               ##     "Tray Program",
               #     "Application was minimized to Tray",
               #     QSystemTrayIcon.Information,
               #     2000
               # )
            else:
                self.showMaximized()
    def closeEvent(self, event):
        if True:
            event.ignore()
            self.hide()
          #  self.tray_icon.showMessage(
          #      "Tray Program",
          #      "Application was minimized to Tray",
          #      QSystemTrayIcon.Information,
          #      2000
          #  )
    def displayLoading(self):
        self.splash.show()
        app.processEvents()

        # Close SplashScreen after 2 seconds (2000 ms)
    def genQuestions(self, text):
        
        self.questions = []
        q_counter = 0
        
        screen = QApplication.primaryScreen()
        ###print('Screen: %s' % screen.name())
        size = screen.size()
        ###print('Size: %d x %d' % (size.width(), size.height()))
        rect = screen.availableGeometry()
        ###print('Available: %d x %d' % (rect.width(), rect.height()))
        self.width = size.width()
        for q in self.custom_questions:
            if q["site"].split('.')[0].lower() in text:
                for question in q["questions"]:
                    self.questions.append(QPushButton(question.split(' (')[0], self))
                    self.questions[-1].setIcon(QIcon(os.path.join(this_path, 'icons',q["site"])))
                   # self.questions[-1].move(20,200+len(self.questions)*100)/border: 1px solid black;

                    self.questions[-1].setStyleSheet("text-align:left;font-size:24px;")
                    self.questions[-1].setIconSize(QSize(size.width()/12, size.width()/12))
                    self.questions[-1].setFixedSize(QSize(size.width()/2,152))
                    self.questions[-1].show()
                    self.questions[-1].setToolTip("<h3>" + question + "</h3>")
                    self.questions[-1].clicked.connect(self.submitQ)
                    self.vbox.addWidget(self.questions[-1],  q_counter, 0 ) 
                    q_counter += 1

        for q in self.base_questions:
            self.questions.append(QPushButton("  " + q + self.sender().toolTip().split(' (')[0] + ".", self))
            self.questions[-1].setIcon(QIcon(os.path.join(this_path, 'icons', self.sender().toolTip().split(' (')[0] + '.ico')))
           # self.questions[-1].move(20,200+len(self.questions)*100)
            self.questions[-1].setStyleSheet("text-align:left;font-size:24px;")
            self.questions[-1].setIconSize(QSize(size.width()/12, size.width()/12))
            self.questions[-1].setFixedSize(QSize(size.width()/2,152))
            self.questions[-1].show()
            self.questions[-1].setToolTip("<h3>" + q + "</h3>")
            self.questions[-1].clicked.connect(self.submitQ)
            self.vbox.addWidget(self.questions[-1],  q_counter, 0 ) 
            q_counter += 1
        self.questions.append(QPushButton("", self))
        self.questions[-1].setIcon(QIcon(resource_path('backroot.png')))
        self.questions[-1].setIconSize(QSize(size.width()/12, size.width()/12))
        self.questions[-1].setFixedSize(QSize(size.width()/2,152))
        self.questions[-1].show()
        self.questions[-1].setToolTip("<h3>Back</h3>")
        self.questions[-1].clicked.connect(self.backRoot)
        self.vbox.addWidget(self.questions[-1],  q_counter, 0 )

        for x in self.atlas_array:
            x.hide()
    def backRoot(self):
        for x in self.questions:
            x.hide()
        for x in self.atlas_array:
            x.show()
    
    def submitQ(self):
        self.displayLoading()

        app.processEvents()

        for q in self.questions:
            app.processEvents()

            q.hide()
        time.sleep(2)
        for x in self.atlas_array:
            x.show()
            app.processEvents()
        time.sleep(2)

        self.line.setText(self.sender().text().split('(')[0])
        self.atlas()

    def changeTestImage(self):
        self.test_image = QFileDialog.getOpenFileName(self, 'Open File', '/home')[0]
        self.pic.setPixmap(QtGui.QPixmap(self.test_image))

    def __del__(self):
        # Restore sys.stdout
        sys.stdout = sys.__stdout__
    
    def normalOutputWritten(self, text):
        """Append text to the QTextEdit."""
        global graph
        # Maybe QTextEdit.append() works as well, but this is how I do it:
        print("captured" +text)
        

    def trainer(self,image,  train, test, model_file , labels, output):
        global test_image
       # ###print(torch.cuda.is_available())
       # model = core.Model.load( "model_weights.pth",  ["links","buttons","sliders","dropdowns","switches","menus","textareas","textfields","checkboxes","radioboxes","images","text"])
     
        #custom_transforms = transforms.Compose([transforms.ToPILImage(),transforms.Resize(900), transforms.RandomHorizontalFlip(0.5), transforms.ColorJitter(saturation=0.2),transforms.ToTensor(),utils.normalize_transform(), ])
        
        utils.xml_to_csv(train, 'train_labels.csv')
        utils.xml_to_csv(test, 'val_labels.csv')
        import csv
       # with open('train_labels.csv', 'r') as file:
       #     # open the file in the write mode
       #    
       #     reader = csv.reader(file)
       #     counter = 0
       #     for row in reader:
       #         counter += 1
       #         if "xmin" != row[4] and (int(row[4]) >= int(row[6]) or int(row[5]) >= int(row[7])):
       #             ###print(row[4], row[6], row[5], row[7])
       #             ###print(row[0].split(".jpg")[0] + ".xml")
       #             ####print(row[0])
       #             import os
       #             try:
       #                 os.remove('train/' + row[0].split(".jpg")[0] + ".xml")
       #             
       #                 os.remove( 'train/' + row[0])
       #             except:
       #                 ####print("error")
#
       # with open('val_labels.csv', 'r') as file:
       #     # open the file in the write mode
       #    
       #     reader = csv.reader(file)
       #     counter = 0
       #     for row in reader:
       #         counter += 1
       #         if "xmin" != row[4] and (int(row[4]) >= int(row[6]) or int(row[5]) >= int(row[7])):
       #             ####print(row[4], row[6], row[5], row[7])
       #             ####print(row[0].split(".jpg")[0] + ".xml")
       #             ####print(row[0])
       #             import os
       #             try:
       #                 os.remove('test/' + row[0].split(".jpg")[0] + ".xml")
       #             
       #                 os.remove('test/' + row[0])
       #             except:
       #                 ####print("error")
# Defin#e custom transforms to apply to your datasetcd 
#
# Pass in a CSV file instead of XML files for faster Dataset initialization speeds

        dataset = core.Dataset('train_labels.csv', train)

        valid_dataset = []
        valid_valid = []
        removes = []


        

        val_dataset = core.Dataset('val_labels.csv', test)  # Validation dataset for training
    
# Create your own DataLoader with custom option
# Create your own DataLoader with custom options

        loader=core.DataLoader(dataset, batch_size=2, shuffle=True)#L3
        # Use MobileNet instead of the default ResNet
        ####print(labels)
        output = core.Model(labels)
        losses = output.fit(loader, val_dataset, epochs=int(self.epochs.text()),  lr_step_size=5,learning_rate=float(self.rate.text()), verbose=True)

        #plt.plot(losses)  # Visualize loss throughout training
        #plt.show()
        #image = utils.read_image("test_image2.png") 
        #predictions = model.predict(image)
        #labels, boxes, scores = predictions

        output.save(model_file)  # Save model to a file

        #show_labeled_image(image, boxes, labels)
    def trainModel(self):
        import threading
        t0 = threading.Thread(target = self.trainer, args=(self.test_image, self.train.text(), self.test.text(), self.model.text(), self.labels_input.text().split(","), self.output))
        t0.start()
        #t0.join()
        #self.output.save(self.model_file) 
        #self.trainer()

    """
        Non-max Suppression Algorithm
        @param list  Object candidate bounding boxes
        @param list  Confidence score of bounding boxes
        @param float IoU threshold
        @return Rest boxes after nms operation
    """
    def nms(self, bounding_boxes, confidence_score, labels, threshold):
        # If no bounding boxes, return empty list
        if len(bounding_boxes) == 0:
            return [], []
    
        # Bounding boxes
        boxes = np.array(bounding_boxes)
    
        # coordinates of bounding boxes
        start_x = boxes[:, 0]
        start_y = boxes[:, 1]
        end_x = boxes[:, 2]
        end_y = boxes[:, 3]
    
        # Confidence scores of bounding boxes
        score = np.array(confidence_score)
    
        # Picked bounding boxes
        picked_boxes = []
        picked_score = []
        picked_labels = []
    
        # Compute areas of bounding boxes
        areas = (end_x - start_x + 1) * (end_y - start_y + 1)
    
        # Sort by confidence score of bounding boxes
        order = np.argsort(score)
    
        # Iterate bounding boxes
        while order.size > 0:
            # The index of largest confidence score
            index = order[-1]
    
            # Pick the bounding box with largest confidence score
            picked_boxes.append(index)
    
            # Compute ordinates of intersection-over-union(IOU)
            x1 = np.maximum(start_x[index], start_x[order[:-1]])
            x2 = np.minimum(end_x[index], end_x[order[:-1]])
            y1 = np.maximum(start_y[index], start_y[order[:-1]])
            y2 = np.minimum(end_y[index], end_y[order[:-1]])
    
            # Compute areas of intersection-over-union
            w = np.maximum(0.0, x2 - x1 + 1)
            h = np.maximum(0.0, y2 - y1 + 1)
            intersection = w * h
    
            # Compute the ratio between intersection and union
            ratio = intersection / (areas[index] + areas[order[:-1]] - intersection)
    
            left = np.where(ratio < threshold)
            order = order[left]
    
        return picked_boxes

    def soft_nms_pytorch(self, dets, labels, box_scores, sigma=0.5, thresh=0.001, cuda=0):
        """
        Build a pytorch implement of Soft NMS algorithm.
        # Augments
            dets:        boxes coordinate tensor (format:[y1, x1, y2, x2])
            box_scores:  box score tensors
            sigma:       variance of Gaussian function
            thresh:      score thresh
            cuda:        CUDA flag
        # Return
            the index of the selected boxes
        """
    
        # Indexes concatenate boxes with the last column
        N = dets.shape[0]
        if cuda:
            indexes = torch.arange(0, N, dtype=torch.float).cuda().view(N, 1)
        else:
            indexes = torch.arange(0, N, dtype=torch.float).view(N, 1)
        dets = torch.cat((dets, indexes), dim=1)
    
        # The order of boxes coordinate is [y1,x1,y2,x2]
        y1 = dets[:, 0]
        x1 = dets[:, 1]
        y2 = dets[:, 2]
        x2 = dets[:, 3]
        scores = box_scores
        areas = (x2 - x1 + 1) * (y2 - y1 + 1)
    
        for i in range(N):
            # intermediate parameters for later parameters exchange
            tscore = scores[i].clone()
            pos = i + 1
    
            if i != N - 1:
                maxscore, maxpos = torch.max(scores[pos:], dim=0)
                if tscore < maxscore:
                    dets[i], dets[maxpos.item() + i + 1] = dets[maxpos.item() + i + 1].clone(), dets[i].clone()
                    scores[i], scores[maxpos.item() + i + 1] = scores[maxpos.item() + i + 1].clone(), scores[i].clone()
                    areas[i], areas[maxpos + i + 1] = areas[maxpos + i + 1].clone(), areas[i].clone()
    
            # IoU calculate
            yy1 = np.maximum(dets[i, 0].to("cpu").numpy(), dets[pos:, 0].to("cpu").numpy())
            xx1 = np.maximum(dets[i, 1].to("cpu").numpy(), dets[pos:, 1].to("cpu").numpy())
            yy2 = np.minimum(dets[i, 2].to("cpu").numpy(), dets[pos:, 2].to("cpu").numpy())
            xx2 = np.minimum(dets[i, 3].to("cpu").numpy(), dets[pos:, 3].to("cpu").numpy())
            
            w = np.maximum(0.0, xx2 - xx1 + 1)
            h = np.maximum(0.0, yy2 - yy1 + 1)
            inter = torch.tensor(w * h).cuda() if cuda else torch.tensor(w * h)
            ovr = torch.div(inter, (areas[i] + areas[pos:] - inter))
    
            # Gaussian decay
            weight = torch.exp(-(ovr * ovr) / sigma)
            scores[pos:] = weight * scores[pos:]
        keep= []
        ####print(scores)
        for i in range(0, len(dets)):
            if scores[i] < thresh:
                dets[i].remove()
                labels[i].remove()
        # select the boxes and keep the corresponding indexes
        keep = dets
        labels_keep = labels
        return keep, labels_keep    
    def testModel(self):
        import torch 

       # self.output.save(self.model_file) 
        ####print(self.labels_input.text().split(","))
        model = core.Model.load(self.model.text(),  self.labels_input.text().split(","))
        image = utils.read_image(self.test_image) 
        predictions = model.predict(image)
        labels, boxes, scores = predictions
        if boxes.ndim == 1:
            boxes = boxes.view(1, 4)
        thresh=float(self.thresh.text())
        ####print(scores)
        filtered_indices=np.where(scores>thresh)
        filtered_scores=scores[filtered_indices]
        filtered_boxes=boxes[filtered_indices]
        num_list = filtered_indices[0].tolist()
        filtered_labels = [labels[i] for i in num_list]
       # show_labeled_image(image, filtered_boxes, filtered_labels)
        # Plot each box
        
        ####print(filtered_boxes)
        ####print(filtered_labels)        
        indices = self.nms(filtered_boxes, filtered_scores, filtered_labels, .05)
        filtered_boxes = filtered_boxes[indices]
        filtered_labels = [filtered_labels[i] for i in indices]

       # filtered_labels = filtered_labels[indices]
        for i in range(filtered_boxes.shape[0]):
            box = filtered_boxes[i]
            
            width, height = (box[2] - box[0]).item(), (box[3] - box[1]).item()
            initial_pos = (int(box[0].item()), int(box[1].item()))
            end_pos = (int(box[2].item()), int(box[3].item()))
            cv2.rectangle(image,initial_pos, end_pos, (0, 255, 0), 2)
            if filtered_labels is not None:
                cv2.putText(image, filtered_labels[i],initial_pos, cv2.FONT_HERSHEY_SIMPLEX, 1, (255, 0, 0), 2)
        height, width, channel = image.shape
        bytesPerLine = 3 * width
        qImg = QImage(image.data, width, height, bytesPerLine, QImage.Format_RGB888)
        self.pic.setPixmap(QPixmap(qImg))


    def setTest(self):
        self.test_folder = self.test.text()
    def setTrain(self):
        self.train_folder = self.train.text()
    def changeModel(self):
        self.model_file = self.model.text()
    def back(self):
        self.question1.hide()
        self.question2.hide()
        self.question3.hide()
        self.question4.hide()
        self.backButton.hide()
        self.API.show()
        self.Custom.show()
        self.Amazon.show()
        self.Google.show()
    #    self.speakButton.show()
        self.pushButton.show()
        self.line.show()
        self.nameLabel.show()
        self.backButton.hide()

    def submitQ1(self):
        self.question1.hide()
        self.question2.hide()
        self.question3.hide()
        self.question4.hide()
        self.backButton.hide()
        self.Amazon.show()
        self.Google.show()
        self.Custom.show()
        self.API.show()
    #    self.speakButton.show()
        self.line.setText(self.question1.text())
        self.atlas()

    def submitQ2(self):
        self.question1.hide()
        self.question2.hide()
        self.question3.hide()
        self.question4.hide()
        self.backButton.hide()
        
    #    self.speakButton.show()
        self.Amazon.show()
        self.Google.show()
        self.Custom.show()
        self.API.show()
        self.line.setText(self.question2.text())
        self.atlas()

    def submitQ3(self):

        self.question1.hide()
        self.question2.hide()
        self.question3.hide()
        self.question4.hide()
        self.backButton.hide()
        
     #   self.speakButton.show()
        self.Amazon.show()
        self.Google.show()
        self.Custom.show()
        self.API.show()
        self.line.setText(self.question3.text())
        self.atlas()

    def submitQ4(self):
        self.question1.hide()
        self.question2.hide()
        self.question3.hide()
        self.question4.hide()
        self.backButton.hide()
        
     #   self.speakButton.show()
        self.Amazon.show()
        self.Google.show()
        self.Custom.show()
        self.API.show()
        self.line.setText(self.question4.text())
        self.atlas()
      

    def handleAmazon(self):
        self.Amazon.hide()
        self.Google.hide()
        self.Custom.hide()
        self.API.hide()
        
      #  self.speakButton.hide()
        self.backButton.show()
        self.question1.show()
        self.question2.show()
        self.question3.show()
        self.question4.show()


        ####print("Amazon")
    def handleGoogle(self):
        print("Google")
    def getIcon(self,url):
        data = urllib.request.urlopen(url).read()
        icon = QtGui.QPixmap()
        icon.loadFromData(data)
        return icon

    def atlas(self):
        ###print("OPENAI KEY")
        ####print(len(self.openAIKey))
        ####print(self.openAIKey)
        ####print(self.input_queue)
        global thread_signals
        if self.mode == "input":
            if len(self.line.text().strip().rstrip()) == 0:
                self.model.add_message(USER_THEM, "Please enter an input")
                return
            self.inputs.append({self.input_queue[0]:self.line.text().strip().rstrip().replace('\n','')})
            self.model.add_message(USER_ME, self.line.text())
            
            if "API_KEY" in self.input_queue[0]:
                self.openAIKey = self.line.text()
            
            self.line.setText('')
            if len(self.input_queue) > 0:
                self.input_queue.pop(0)
                if len(self.input_queue) > 0:

                    if "API" in self.input_queue[0]:
                        if len(self.openAIKey) > 0:
                            self.input_queue.pop(0)
                            self.inputs.append({self.input_queue[0]:self.openAIKey})
                           # self.model.add_message(USER_ME, self.openAIKey)
                            self.line.setText('')
                        else:
                            if "API" in self.input_queue[0]:
                                if len(self.openAIKey) > 0:
                                    self.inputs.append({self.input_queue[0]:self.openAIKey})
                                    self.model.add_message(USER_ME, self.openAIKey)
                                    self.line.setText('')
                                    self.input_queue.pop(0)

                                else:
                                    self.model.add_message(USER_THEM, "Please enter your openAI API Key:")                                
                            elif "SHEET_URL" in self.input_queue[0]:
                                self.model.add_message(USER_THEM, "Please enter your google sheets URL:")                   
                            elif "SHEET_NAME" in self.input_queue[0]:
                                self.model.add_message(USER_THEM, "Please enter your google sheets Sheet Name(usually Sheet1):")
                            else:
                                self.model.add_message(USER_THEM, "Please enter the input:" + self.input_queue[0])
                                
                    else:
                        if len(self.openAIKey) > 0 and "API" in self.input_queue[0]:
                            self.input_queue.pop(0)
                            self.inputs.append({self.input_queue[0]:self.openAIKey})
                            self.model.add_message(USER_ME, self.openAIKey)
                            self.line.setText('')
                        else:
                            if "API" in self.input_queue[0]:
                                self.model.add_message(USER_THEM, "Please enter your openAI API Key:")                                
                            elif "SHEET_URL" in self.input_queue[0]:
                                self.model.add_message(USER_THEM, "Please enter your google sheets URL:")
                                                    
                            elif "SHEET_NAME" in self.input_queue[0]:
                                self.model.add_message(USER_THEM, "Please enter your google sheets Sheet Name(usually Sheet1):")
                            else:
                                self.model.add_message(USER_THEM, "Please enter the input:" + self.input_queue[0])

                if len(self.input_queue) == 0:
                    self.mode = "chat"
                    textOutput = self.response["choices"][0]["text"]
                    for x in self.inputs:
                        for key, value in x.items():
                            textOutput = textOutput.replace(key, "\'" + value + "\'" )
                    output = json.loads(textOutput)
                    nodesList = output["nodes"]
                    schedule = ''
                    for node in nodesList:
                        app.processEvents()
    
                        if node["type"] == "runOnCheatLayer":
                            x = {"type":"cheatlayer", "code":node["script"], "url" : node["url"]}
                            nodes.append(graph.create_node('nodes.widget.ImageNode', name="Cheat Layer Extension " + str(len(nodes)), data=x))#, color= "#FFFFFF"
                            nodes[len(nodes)-1].create_property('code',node["script"], widget_type=NODE_PROP_QLINEEDIT)
                            nodes[len(nodes)-1].create_property('URL',node["url"], widget_type=NODE_PROP_QLINEEDIT)
                            nodes[len(nodes)-1].create_property('key',user_key, widget_type=NODE_PROP_QLINEEDIT)
                            nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
    
                            connectToLastNode(nodes[len(nodes)-1])
                            graph_nodes = graph.serialize_session()
                            #d()
                            graph.runNode(graph_nodes,nodes[len(nodes)-1].id, thread_signals)

                            self.model.add_message(USER_THEM, "Step completed!")
                            self.mode = "chat"
                            if "schedule" in output:
                                schedule = output["schedule"]
                            if "blogBasic" in node["script"]:
                                time.sleep(30)
                                self.model.add_message(USER_THEM, "Clicking publish")
                                pyautogui.click(resource_path('wordpress_publish.png'))    
                                time.sleep(5)
                                self.model.add_message(USER_THEM, "Clicking publish")
                                pyautogui.click(resource_path('wordpress_publish2.png'))    
                            if "postBasic" in node["script"] and "image" in node["script"]:
                                ###print('Dalle!')
                                prompt = node["script"].split(",")
                                ###print(prompt[0])
                                ###print("PROMTP")
                                key = self.openAIKey
                                ###print(prompt[1])
                                r = requests.post("https://cheatlayer.com/user/generateDalle", data={"id":user_key, "input":prompt[2].replace('[','').replace(']','').replace("'",'')}, verify=False)
                                ###print("PROMTP")
                                ###print(r.text)
                                out =r.json()
                                ###print(out["data"][0]["url"])           
                                ###print("RESPONSE")
                                img_data = requests.get(out["data"][0]["url"]).content
                                with open(resource_path('dalle.jpg'), 'wb') as handler:
                                    handler.write(img_data)
                                time.sleep(20)
                                self.model.add_message(USER_THEM, "Clicking photos")
                                pyautogui.click(resource_path('facebook_photos.png'))    
                                time.sleep(5)
                                pyautogui.write(resource_path('dalle.jpg'))
                                time.sleep(5)
                                pyautogui.press("enter")
                                time.sleep(5)
                                pyautogui.click(resource_path('facebook_post.png'))    
    
                                
    
        
                                #####print('Dalle!')
                                prompt = node["script"].split(",")
                                #####print("PROMTP")
                                            

                                key = self.openAIKey
                                if len(key) == 0:
                                    key = self.inputs[0]["value"]
                                #####print("THIS IS THE PROMPT WE USE")
                                r = requests.post("https://cheatlayer.com/user/generateDalle", data={"id":user_key, "input":prompt[2].replace('[','').replace(']','').replace("'",'')}, verify=False)
                                ###print(r.text)
                                out =r.json()
                                out =json.loads(r.text)

                                if "data" not in out:
                                    msg = QMessageBox()
                                    msg.setWindowTitle("Not Enough Cheat Cloud Tasks")
                                    msg.setText("You need more Cheat Cloud Tasks or there was an error with your account. Please upgrade at cheatlayer.com/billing")
                                    _ = msg.exec_()
                                    return
                                ######print(out["data"][0]["url"])           
                                ####print("RESPONSE")
                                img_data = requests.get(out["data"][0]["url"]).content
                                with open(resource_path('dalle.jpg'), 'wb') as handler:
                                    handler.write(img_data)
                                subprocess.call(["open", "-a", "Google Chrome"])
                                screen = QApplication.primaryScreen()
                                size = screen.size()
                                import pyautogui
                                pyautogui.move(size.width()/2,size.height()/2)
                                pyautogui.dragTo(size.width()/2,size.height()/2, button='left')
                                pyautogui.click(size.width()/2,size.height()/2)  
                                pyautogui.click(size.width()/2,size.height()/2)  
                                import pyautogui
                                from pynput.mouse import Button, Controller             
                                mouse = Controller()
                                pyautogui.moveTo(size.width()/2,size.height()/2)
                                mouse.click(Button.left)
                                time.sleep(8)
                                self.model.add_message(USER_THEM, "Clicking.")
                                import pyautogui
                                pyautogui.press('tab')  
                                time.sleep(3)
                                pyautogui.press('tab')  
                                time.sleep(3)
                                pyautogui.press("enter")
                                time.sleep(3)
                                pyautogui.press('down')
                                time.sleep(2)
                                pyautogui.press("enter")
                                
                                time.sleep(2)
                                pyautogui.press('tab')  
                                time.sleep(3)
                                pyautogui.press('tab')  
                                time.sleep(3)
                                pyautogui.press("enter")
                                time.sleep(3)
                                pyautogui.press('down')
                                time.sleep(2)
                                pyautogui.press("enter")
    #           for     node in nodes:
            #x = {"type":"codex", "prompt":"Scrape amazon for all the greenie prices.", "selected":""}
            #nodes.append(graph.create_node('nodes.widget.ImageNode', name="Codex Generator " + str(len(nodes)), data=x))#, color= "#FFFFFF"
            #nodes[len(nodes)-1].create_property('prompt',"Scrape amazon for all the greenie prices.", widget_type=NODE_PROP_QLINEEDIT)
                    if schedule != '':
                       ####print("scheduler final")
                       name, save = QFileDialog.getSaveFileName(graph.QMainWindow, 'Save Cheat',"", "CHEAT (*.cheat)")
                       graph.save_session(name)
                       job_id = graph.cheat_scheduler.add_job(func=graph.QMainWindow.scheduler.runSchedule, trigger='cron',args=[name],  second='0',day=schedule.split(" ")[2], day_of_week=schedule.split(" ")[4], month=schedule.split(" ")[3], hour=schedule.split(" ")[1], minute=schedule.split(" ")[0])
                       
                       job = {"enabled": "True","job": job_id, "seconds":"*","minute":schedule.split(" ")[0], "hour":schedule.split(" ")[1],"week":"*",
                           "day":schedule.split(" ")[2],"weekday":schedule.split(" ")[3],"month":schedule.split(" ")[4],"id":"testCheat", "cheat": name}
                       scheduled_jobs.append(job)
                       
                       workflows.append(job)
                       #graph.cheat_scheduler.shutdown()
                       #graph.cheat_scheduler.start()
                       self.model.add_message(USER_THEM, "Scheduled Jobs: " + str(scheduled_jobs))
                       graph.QMainWindow.scheduler.redraw()

                       self.model.add_message(USER_THEM, "I'll also add this cron format to your scheduler and schedule this automation for you: " + schedule)
                       ###print(scheduled_jobs)
                       send_jobs = []
                       for job in scheduled_jobs:
                            if "type" in job and job["type"] == "schedule":
                                test = job.copy()
                                test["job"] = ""
                                send_jobs.append(test)
                       r = requests.post("https://cheatlayer.com/user/saveDeskSchedule", data={'id': user_key, "schedule": json.dumps(send_jobs)},  verify=False)
                    graph.auto_layout_nodes()
                    graph.fit_to_selection()
        elif self.mode == "chat":
            self.displayLoading()
            prompt =  self.line.text().rstrip() + "\n"
            self.model.add_message(USER_ME, self.line.text())
            self.line.setText('')
            #self.line.setText("")
            schedule = ''
            if True:
                app.processEvents()
                prompt_input = self.prompts[0]["prompt"]
                for prompt_list in self.prompts:
                    for website in prompt_list["websites"]:
                        if website.lower() in prompt.lower():
                           # ####print(website)
                            prompt_input = prompt_list["prompt"]
                response = requests.post("https://cheatlayer.com/user/generateAtlas", data={'id': user_key, "input":prompt_input + prompt},  verify=False)
               # response = openai.Completion.create(
               #       model="text-davinci-002",
               #       prompt=prompt_input + prompt,
               #       temperature=0,
               #       max_tokens=256,
               #       top_p=1,
               #       frequency_penalty=0,
               #       presence_penalty=0,
               #       stop=["###"]
               #     )
                self.response = response.json()
                self.model.add_message(USER_THEM, "Ok I'll build that for you. Please wait a moment.")


                if "not supported" in self.response["choices"][0]["text"] or  "This is not a valid input" in self.response["choices"][0]["text"]:
                    self.model.add_message(USER_THEM, "This feature is not supported! Cheat Layer has been notified and you will get an email follow up when it's ready or if we have further updates.")
                    #response = requests.post("https://cheatlayer.com/user/generateGPTUser", data={'id': user_key, "input":prompt},  verify=False)

                    ####print("This feature is not supported! Cheat Layer has been notified and you will get an email follow up when it's ready or if we have further updates.")
                    #self.line.setText("not supported")
                    #time.sleep(2)
                    r = requests.post("https://cheatlayer.com/user/trainAtlas", data={'id': user_key, prompt: prompt},  verify=False)
    
                    self.splash.close()
                else:
                    if "You have run out of machine learning credits!" in self.response["choices"][0]["text"]:
                        self.model.add_message(USER_THEM, "You have run out of machine learning credits! Please upgrade on the billing page at cheatlayer.com/billing")
                    else:
                        if "inputs" in self.response["choices"][0]["text"]:
                            ###print(self.response["choices"][0]["text"])
                            inputs = self.response["choices"][0]["text"].split('"inputs": [')[1].split("]")[0].split(",")
                            output = json.loads(self.response["choices"][0]["text"])
                      #      ####print(output)
    
                          #  ####print("runOnCheatLayer")
                           # ####print(output)
                            if "schedule" in output:
                               ####print("scheduler0")
                               schedule = output["schedule"]
                            self.inputs = []
                            #time.sleep(2)
                            self.splash.close()
                            if len(output["inputs"]) > 0:
                                self.input_queue = output["inputs"]
        
                            if len(self.input_queue) > 0:
                                if "API" in self.input_queue[0]:
                                    if len(self.openAIKey) > 0:
                                        self.inputs.append({self.input_queue[0]:self.openAIKey})
                                        self.model.add_message(USER_ME, self.openAIKey)
                                        self.line.setText('')
                                        self.input_queue.pop(0)
    
                                    else:
                                        self.model.add_message(USER_THEM, "Please enter your openAI API Key:")                                
                                elif "SHEET_URL" in self.input_queue[0]:
                                    self.model.add_message(USER_THEM, "Please enter your google sheets URL:")
                                                        
                                elif "SHEET_NAME" in self.input_queue[0]:
                                    self.model.add_message(USER_THEM, "Please enter your google sheets Sheet Name(usually Sheet1):")
                                else:
                                    self.model.add_message(USER_THEM, "Please enter the input:" + self.input_queue[0])
                                self.mode = "input"
                            if len(self.input_queue) == 0:
                                self.mode = "chat"
                                textOutput = self.response["choices"][0]["text"]
                                for x in self.inputs:
                                    for key, value in x.items():
                                        textOutput = textOutput.replace(key, "\'" + value + "\'" )
                                output = json.loads(textOutput)
                                nodesList = output["nodes"]
                                for node in nodesList:
                                    app.processEvents()
                
                                    if node["type"] == "runOnCheatLayer":
                                        x = {"type":"cheatlayer", "code":node["script"], "url" : node["url"]}
                                        nodes.append(graph.create_node('nodes.widget.ImageNode', name="Cheat Layer Extension " + str(len(nodes)), data=x))#, color= "#FFFFFF"
                                        nodes[len(nodes)-1].create_property('code',node["script"], widget_type=NODE_PROP_QLINEEDIT)
                                        nodes[len(nodes)-1].create_property('URL',node["url"], widget_type=NODE_PROP_QLINEEDIT)
                                        nodes[len(nodes)-1].create_property('key',user_key, widget_type=NODE_PROP_QLINEEDIT)
                                        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
                
                                        connectToLastNode(nodes[len(nodes)-1])
                                        graph_nodes = graph.serialize_session()
                
                                        graph.runNode(graph_nodes,nodes[len(nodes)-1].id, thread_signals)
                         
                                        self.model.add_message(USER_THEM, "Step completed!")
                                        self.mode = "chat"
                                        
                                        if "schedule" in output:
                                           ####print("scheduler1")
                                           schedule = output["schedule"]
                                        if "blogBasic" in node["script"]:
                                            time.sleep(30)
                                            self.model.add_message(USER_THEM, "Clicking publish")
            
                                            pyautogui.click(resource_path('wordpress_publish.png'))    
                                            
                                            time.sleep(5)
                                            self.model.add_message(USER_THEM, "Clicking publish")
            
                                            pyautogui.click(resource_path('wordpress_publish2.png'))    
                                        if "postBasic" in node["script"] and "image" in node["script"]:
                                            ###print('Dalle!')
    
                                            prompt = node["script"].split(",")
                                            ###print(prompt[0])
                                            ###print("PROMTP")
                                            
    
                                            key = self.openAIKey
                                            
                                            ###print(prompt[1])
                                            r = requests.post("https://cheatlayer.com/user/generateDalle", data={"id":user_key, "input":prompt[2].replace('[','').replace(']','').replace("'",'')}, verify=False)
                                            ###print("PROMTP")
                                            ###print(r.text)
                                            out =r.json()
                                            if "data" not in out:
                                                self.model.add_message(USER_THEM, "Error: you have run out of Cheat Cloud tasks. Please upgrade at cheatlayer.com/billing")
                                                            
                                                msg = QMessageBox()
                                                msg.setWindowTitle("Not Enough Cheat Cloud Tasks")
                                                msg.setText("You need more Cheat Cloud Tasks. Please upgrade at cheatlayer.com/billing")
                                                _ = msg.exec_()
                                                return
                                               
                                            ###print(out["data"][0]["url"])           
                                            ###print("RESPONSE")
                                            img_data = requests.get(out["data"][0]["url"]).content
                                            with open(resource_path('dalle.jpg'), 'wb') as handler:
                                                handler.write(img_data)
                                            import pyautogui
                                            time.sleep(20)
                                            self.model.add_message(USER_THEM, "Uploading Dalle image")
            
                                            pyautogui.click(resource_path('facebook_photos.png'))    
                                            time.sleep(5)
                                            pyautogui.write(resource_path('dalle.jpg'))
            
                                            time.sleep(5)
                                            pyautogui.press("enter")
    
                                            time.sleep(5)
                                            pyautogui.click(resource_path('facebook_post.png'))    
    
                                
    
        
                                            ####print('Dalle!')
    
                                            prompt = node["script"].split(",")
                                            
    
                                if schedule != '':
                                   ####print("set SCHEDULE")
                                   name, save = QFileDialog.getSaveFileName(graph.QMainWindow, 'Save Cheat',"", "CHEAT (*.cheat)")
                                   graph.save_session(name)
                                   job_id = graph.cheat_scheduler.add_job(func=graph.QMainWindow.scheduler.runSchedule, trigger='cron',args=[name],  second='0',day=schedule.split(" ")[2], day_of_week=schedule.split(" ")[4], month=schedule.split(" ")[3], hour=schedule.split(" ")[1], minute=schedule.split(" ")[0])
                                   
                                   job = {"enabled": "True","job": job_id, "seconds":"*","minute":schedule.split(" ")[0], "hour":schedule.split(" ")[1],"week":"*",
                                       "day":schedule.split(" ")[2],"weekday":schedule.split(" ")[3],"month":schedule.split(" ")[4],"id":"testCheat", "cheat": name}
                                   scheduled_jobs.append(job)
                                   #graph.cheat_scheduler.shutdown()
                                   #graph.cheat_scheduler.start()
                                   self.model.add_message(USER_THEM, "Scheduled Jobs: " + str(scheduled_jobs))
                                   self.model.add_message(USER_THEM, "I'll also add this cron format to your scheduler and schedule this automation for you: " + schedule)
                                   ###print(scheduled_jobs)
                                   send_jobs = []
                                   for job in scheduled_jobs:
                                        if "type" in job and job["type"] == "schedule":
                                            test = job.copy()
                                            test["job"] = ""
                                            send_jobs.append(test)
                                   r = requests.post("https://cheatlayer.com/user/saveDeskSchedule", data={'id': user_key, "schedule": json.dumps(send_jobs)},  verify=False)

                                #graph.QMainWindow.scheduler.redraw()

                                graph.auto_layout_nodes()
                                graph.fit_to_selection()
                        else:
                            self.model.add_message(USER_THEM, "Error processing request by openAI. Please try again or use a different phrase.")
   
                        #graph.playRecording()
    
    
    def main_window(self):
        
        self.setWindowTitle(self.title)
        self.setGeometry(self.top, self.left, self.width, self.height)
        self.show()


class Window(QWidget):
    def __init__(self):
        super(Window, self).__init__()
        self.view = View(self)

        self.layout_contain_P1_P2 = QtWidgets.QGridLayout()
        self.checkbox_P1= QtWidgets.QCheckBox("P1",self)
        self.line_edit_P1_x = QtWidgets.QLineEdit(self)
        self.line_edit_P1_x.setReadOnly(True)
        self.line_edit_P1_y = QtWidgets.QLineEdit(self)
        self.line_edit_P1_y.setReadOnly(True)

        self.layout_contain_P1_P2.addWidget(self.checkbox_P1, 0, 0, Qt.AlignLeft)

        self.grid_layout_P1_x_y = QtWidgets.QGridLayout()
        self.grid_layout_P1_x_y.addWidget(self.line_edit_P1_x, 1, 0, Qt.AlignLeft)
        self.grid_layout_P1_x_y.addWidget(self.line_edit_P1_y, 2, 0, Qt.AlignLeft)

        self.layout_contain_P1_P2.addLayout(self.grid_layout_P1_x_y, 0, 1, 1, 1)
        self.checkbox_P2 = QtWidgets.QCheckBox("P2",self)
        self.line_edit_P2_x = QtWidgets.QLineEdit(self)
        self.line_edit_P2_x.setReadOnly(True)
        self.line_edit_P2_y = QtWidgets.QLineEdit(self)
        self.line_edit_P2_y.setReadOnly(True)

        self.layout_contain_P1_P2.addWidget(self.checkbox_P2, 1, 0, Qt.AlignLeft)
        self.grid_layout_P2_x_y = QtWidgets.QGridLayout()

        self.grid_layout_P2_x_y.addWidget(self.line_edit_P2_x, 0, 0, Qt.AlignLeft)
        self.grid_layout_P2_x_y.addWidget(self.line_edit_P2_y, 1, 0, Qt.AlignLeft)

        self.layout_contain_P1_P2.addLayout(self.grid_layout_P2_x_y, 1, 1, Qt.AlignLeft)

        self.combo_box1 = QtWidgets.QComboBox(self)
        self.combo_box1.addItem("measurements set 1")
        self.combo_box1.addItem("measurements set 1")

        self.combo_box2 = QtWidgets.QComboBox(self)
        self.combo_box2.addItem("P1-P2")
        self.combo_box2.addItem("P3-P4")

        self.vertical1= QtWidgets.QVBoxLayout()

        # self.vertical1.addWidget(self.menubar)
        self.vertical1.addWidget(self.combo_box1)
        self.vertical1.addWidget(self.combo_box2)
        self.vertical1.addLayout(self.layout_contain_P1_P2)

        self.vertical2= QtWidgets.QVBoxLayout()
        self.vertical2.addWidget(self.view)

        self.horizontal= QtWidgets.QHBoxLayout()
        self.horizontal.addLayout(self.vertical1)
        self.horizontal.addLayout(self.vertical2)

        self.setLayout(self.horizontal)
        self.setWindowTitle("Image viewer")
        self.setGeometry(200, 200, 1000, 800)
import sys
from PySide2.QtWidgets import QApplication, QMainWindow, QStackedLayout, QPushButton, QWidget, QGridLayout, QVBoxLayout, QLabel, QDialog
from PySide2.QtGui import QDesktopServices

from PySide2.QtMultimedia import QMediaPlayer, QMediaContent
from PySide2.QtMultimediaWidgets import QVideoWidget
from PySide2.QtCore import Qt, QUrl, QSize
from PySide2.QtGui import QIcon, QPixmap

class DetailsDialog(QDialog):
    def __init__(self, automation_info):
        super().__init__()

        layout = QVBoxLayout()


        self.Instance = vlc.Instance()
        self.media_player = self.Instance.media_player_new()
        self.VideoFrame = QWidget(self)
        self.VideoFrame.setFixedSize(720, 300)  # Set video frame size
        self.media_player.set_hwnd(self.VideoFrame.winId())

        this_path = os.path.dirname(os.path.abspath(__file__))
        
        video_path = os.path.join(this_path, 'examples', automation_info['video_path'])
        download("https://cheatlayer.com/agents/" + automation_info['video_path'], video_path)

        self.media = self.Instance.media_new(video_path)
        self.media_player.set_media(self.media)

        play_path = os.path.join(this_path, 'examples', 'play.png')
        pause_path = os.path.join(this_path, 'examples', 'pause.png')

        self.play_icon = QIcon(play_path)
        self.pause_icon = QIcon(pause_path)

        container = QWidget()
        container_layout = QVBoxLayout()
        container.setLayout(container_layout)
        container_layout.addWidget(self.VideoFrame)
        container_layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self.VideoFrame)

        layout.addWidget(container)
        

        self.play_pause_button = QPushButton(self)
        self.play_pause_button.setIcon(self.play_icon)
        self.play_pause_button.clicked.connect(self.toggle_play_pause)
        self.play_pause_button.setFixedSize(50, 50)
        self.play_pause_button.setStyleSheet("background-color: transparent; border: none;")
        self.play_pause_button.setParent(container)
        self.play_pause_button.raise_()
        #self.play_pause_button.move(50, 300)

        # Description
        description_label = QLabel(automation_info['description'])
        layout.addWidget(description_label)
        self.media_player.play()
        # Link
        link_button = QPushButton("More Details")
        link_button.clicked.connect(lambda: QDesktopServices.openUrl(QUrl(automation_info['link'])))
        layout.addWidget(link_button)
        self.setFixedHeight(600)
        # Deploy
        deploy_button = QPushButton("Deploy Automation")
        deploy_button.clicked.connect(partial(self.deployAutomation, automation_info))
        layout.addWidget(deploy_button)

        self.setLayout(layout)

    def deployAutomation(self, automation):
        global graph
        # Placeholder for deploying automation
        graph.clear_session()
        download(automation['automation_path'], automation['automation_path'].split("agents/")[1])
        openFileCheat(automation['automation_path'].split("agents/")[1])
        graph.auto_layout_nodes()
        graph.fit_to_selection()
        graph.playRecording()
        ###print("Deploying automation...")
        self.close()

    def closeEvent(self, event):  # Handle the close event to stop the video
        self.media_player.stop()
        event.accept()
        
    def toggle_play_pause(self):
        if self.media_player.get_state() == vlc.State.Playing:
            self.media_player.pause()
            self.play_pause_button.setIcon(self.play_icon)
        else:
            self.media_player.play()
            self.play_pause_button.setIcon(self.pause_icon)
    def resizeEvent(self, event):
        super().resizeEvent(event)
        #self.play_pause_button.move(10, 300)
        self.play_pause_button.raise_()
        self.play_pause_button.repaint()
import os
import sys
this_path = os.path.dirname(os.path.abspath(__file__))

# Add the directory containing your bundled VLC libraries to the system's PATH
#pipos.add_dll_directory(os.path.join(os.path.dirname(os.path.abspath(__file__)), this_path))


def download(url, filename):
    """
    Downloads an MP4 file from the given URL to the specified filename.
    
    Args:
    - url (str): URL of the MP4 file.
    - filename (str): Local file name to save the MP4 as.
    
    Returns:
    None.
    """
    ##print("Downloading from {} to {}".format(url, filename))
    response = requests.get(url, stream=True)
    
    # Ensure the request was successful.
    response.raise_for_status()
    
    # Save the video to the local file.
    with open(filename, 'wb') as file:
        for chunk in response.iter_content(chunk_size=8192):
            file.write(chunk)
            
    ##print("Download complete.")

if __name__ == '__main__':
    # handle SIGINT to make the app terminate on CTRL+C
    multiprocessing.freeze_support()

    ref_point = []
    cropping = False
    signal.signal(signal.SIGINT, signal.SIG_DFL)
    nodes = []
    QtCore.QCoreApplication.setAttribute(QtCore.Qt.AA_EnableHighDpiScaling)
    app = QtWidgets.QApplication([])
    def normalOutputWritten(text):
        """Append text to the QTextEdit."""    # Maybe QTextEdit.append() works as well, but this is how I do it:
        print("captured" +text)
    def hello():
        print("hello")
    def schedule_test():
        sch = CronSchedule() 
        job = {'seconds':'*/1','minute':'*', 'hour':'*','week':'*',
            'day':'*','weekday':'*','month':'*','id':'testCheat'}
        sch.addJob(drawHistory,job,[history])
        #sch.start()
    

    def connectToLastNode(node_in):
        global nodes, graph, thread_signals
        graph_nodes = graph.serialize_session()
        thread_signals.notifications.emit("Added Node","Node Added", "icon.png", None, "No")
       
        for key, node in graph_nodes["nodes"].items():
            found_in = False
            edit_node = graph.get_node_by_id(key)
           # edit_node.set_property('color', (240,240,240,255))
          #  edit_node.set_property('text_color', (0,0,0,255))

            found_out = False
            if "connections" not in graph_nodes:
                node_in.input(0).connect_to(nodes[0].output(0))
            else:
                for connections in graph_nodes["connections"]:
                    #print(connections["out"])
                    if key == connections["out"][0]:
                        ##print("Found In")
                        found_out = True
                        ##print(connections["out"][1])
                        ##print(connections["out"][0])
                    
                    if key == connections["in"][0]:
                        ##print("Found out")
                        found_in = True
                        ##print(connections["out"][1])
                        ##print(connections["out"][0])
                if found_in == True and found_out == False:
                    
                    
                    node_in.input(0).connect_to(graph.get_node_by_id(key).output(0))

        graph.auto_layout_nodes()

    def shape_selection_semantic(event, x, y, flags, param):
      # grab references to the global variables
      global ref_point, cropping, screenshot, global_variables
    
      # if the left mouse button was clicked, record the starting
      # (x, y) coordinates and indicate that cropping is being
      # performed
      if event == cv2.EVENT_LBUTTONDOWN:
        ref_point = [(x, y)]
        cropping = True
    
      # check to see if the left mouse button was released
      elif event == cv2.EVENT_LBUTTONUP:
        # record the ending (x, y) coordinates and indicate that
        # the cropping operation is finished
        ref_point.append((x, y))
        cropping = False
    
        # draw a rectangle around the region of interest
        cv2.rectangle(screenshot, ref_point[0], ref_point[1], (0, 255, 0), 2)
        cv2.imshow("OCR", screenshot)
        x = {"type":"OCR"}
        nodes.append(graph.create_node('nodes.widget.ImageNode', name="OCR " + str(len(nodes)), data=x))#, color= "#FFFFFF"
        nodes[len(nodes)-1].create_property('Bounding Box X1', ref_point[0][0], widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Bounding Box Y1', ref_point[0][1], widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Bounding Box X2', ref_point[1][0], widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Bounding Box Y2', ref_point[1][1], widget_type=NODE_PROP_QLINEEDIT)

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        this_path = os.path.dirname(os.path.abspath(__file__))
        icon = os.path.join(this_path, 'examples', 'OCRScraper.png')
   
        thresh = 255 - cv2.threshold(cv2.cvtColor(screenshot, cv2.COLOR_BGR2GRAY), 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
            
        xp,yp,w,h = ref_point[0][0], ref_point[0][1], ref_point[1][0]-ref_point[0][0], ref_point[1][1]-ref_point[0][1]  
        ROI = thresh[ref_point[0][1]:ref_point[1][1],ref_point[0][0]:ref_point[1][0]]
        import pytesseract

        data = pytesseract.image_to_string(ROI, lang='eng',config='--psm 6')
        ###print("restults")
        ###print(data)
        graph.global_variables["OCR" + str(len(graph.global_variables))] = (data)
        cv2.imshow("thresh",thresh)
        cv2.imshow("ROI",ROI)

        

        nodes[len(nodes)-1].set_icon(icon)
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])

        graph.auto_layout_nodes()
            
        cv2.waitKey(1)

    # crate a backdrop node and wrap it around
    # "custom port node" and "group node".
    # fit node selection to the viewer.
        graph.fit_to_selection()
        time.sleep(1)
        
        graph.QMainWindow.showMaximized()

    def shape_selection(event, x, y, flags, param):
      # grab references to the global variables
      global ref_point, cropping, screenshot, global_variables
    
      # if the left mouse button was clicked, record the starting
      # (x, y) coordinates and indicate that cropping is being
      # performed
      if event == cv2.EVENT_LBUTTONDOWN:
        ref_point = [(x, y)]
        cropping = True
    
      # check to see if the left mouse button was released
      elif event == cv2.EVENT_LBUTTONUP:
        # record the ending (x, y) coordinates and indicate that
        # the cropping operation is finished
        ref_point.append((x, y))
        cropping = False
    
        # draw a rectangle around the region of interest
        cv2.rectangle(screenshot, ref_point[0], ref_point[1], (0, 255, 0), 2)
        cv2.imshow("OCR", screenshot)
        x = {"type":"OCR"}
        nodes.append(graph.create_node('nodes.widget.ImageNode', name="OCR " + str(len(nodes)), data=x))#, color= "#FFFFFF"
        nodes[len(nodes)-1].create_property('Bounding Box X1', ref_point[0][0], widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Bounding Box Y1', ref_point[0][1], widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Bounding Box X2', ref_point[1][0], widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Bounding Box Y2', ref_point[1][1], widget_type=NODE_PROP_QLINEEDIT)

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        this_path = os.path.dirname(os.path.abspath(__file__))
        icon = os.path.join(this_path, 'examples', 'OCRScraper.png')
   
        thresh = 255 - cv2.threshold(cv2.cvtColor(screenshot, cv2.COLOR_BGR2GRAY), 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
            
        xp,yp,w,h = ref_point[0][0], ref_point[0][1], ref_point[1][0]-ref_point[0][0], ref_point[1][1]-ref_point[0][1]  
        ROI = thresh[ref_point[0][1]:ref_point[1][1],ref_point[0][0]:ref_point[1][0]]
        import pytesseract

        data = pytesseract.image_to_string(ROI, lang='eng',config='--psm 6')
        ###print("restults")
        ###print(data)
        graph.global_variables["OCR" + str(len(graph.global_variables))] = (data)
        cv2.imshow("thresh",thresh)
        cv2.imshow("ROI",ROI)

        

        nodes[len(nodes)-1].set_icon(icon)
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])

        graph.auto_layout_nodes()
            
        cv2.waitKey(1)

    # crate a backdrop node and wrap it around
    # "custom port node" and "group node".
    # fit node selection to the viewer.
        graph.fit_to_selection()
        time.sleep(1)
        
        graph.QMainWindow.showMaximized()

    def click_coordinatesScrape(event,x1,y,flags,params):
      # grab references to the global variables
      global ref_point, screenshot, global_variables    
      
      if event == cv2.EVENT_LBUTTONDOWN:
        ref_point = [(x1-100, y-100),(x1+100,y+100)]
        cut_image = pyautogui.screenshot( region=(x1, y, 200, 200))
        # draw a rectangle around the region of interest
        ##print(cut_image)
       # cv2.imwrite(str(len(nodes) + 1) + "_click.png", cut_image)
       # cv2.waitKey(1)
       # cropped = load_demo_image(image_size=image_size, device=device, input = cut_image)
        cut_image.save(r"crop.png")

        #caption = blip_describe(open("crop.png", "rb"))
        time.sleep(3)
        
        x = {"type":"semanticScrape"}
        labels = ["button", "field", "heading", "iframe", "image", "label", "link", "text"]
        
        graph.global_variables["Scrape" + str(len(nodes))] = ""
        graph.global_variables["Scrape_all" + str(len(nodes))] = ""
        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Semantic Scrape " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('Semantic Description', "text", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Threshold', 0.8, widget_type=NODE_PROP_QLINEEDIT)

        nodes[len(nodes)-1].create_property('Mode',"All",items=["First","All","Last", "Loop Total Runs","Loop Node Runs"],  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Target',"span",items=["span","link text","link destination","article", "div", "Twitter posts"],  widget_type=NODE_PROP_QCOMBO)

        nodes[len(nodes)-1].create_property('X', x1, widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Y', y, widget_type=NODE_PROP_QLINEEDIT)


        for i in range(0,10):
            graph.global_variables["Scrape_" + str(i) + "_" + str(len(graph.global_variables))] = ""

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
              
        graph.auto_layout_nodes()
        graph.fit_to_selection()

        cv2.waitKey(1)
        graph.fit_to_selection()
        time.sleep(1)
        
        graph.QMainWindow.showMaximized()

    def click_coordinates(event,x,y,flags,params):
      # grab references to the global variables
      global ref_point, screenshot, global_variables    
      #print('add clicks')
      #print(event)
      #print(x)
      #print(y)
      if event == cv2.EVENT_LBUTTONDOWN:
        ref_point = [(x-100, y-100),(x+100,y+100)]
        raw_image = pyautogui.screenshot()
       # cut_image  = raw_image.crop((float(element['x']) - element['width']/2, float(element['y']) - element['height']/2, float(element['x']) + element['width']/2, float(element['y']) + element['height']/2))
        # draw a rectangle around the region of interest
        ###print(cut_image)
       # cv2.imwrite(str(len(nodes) + 1) + "_click.png", cut_image)
       # cv2.waitKey(1)
       #draw a red rectangle around the region of interest using draw 
        draw = ImageDraw.Draw(raw_image)
        draw.rectangle([x-100, y-20, x+100, y+20], outline ="red", width=10)
      
       # cropped = load_demo_image(image_size=image_size, device=device, input = cut_image) "Loop Total Runs","Loop Node Runs"
        raw_image.save(resource_path("crop.png"))
                # Use gpt3 to generate an appropriate response
        headers = {
                        "Content-Type": "application/json",
                    }
        cv2.waitKey(1)
        time.sleep(1)
        
        log = [{"role": "system", "content": "You describe the element highlighted by a red bounding box in the input image with as many details as possible and as concisely as possible in 5-10 short words and as few words as possible. If there is text in the image, say exactly what the text says. Only describe the element within the RED bounding box, but you can explain where it is in relation to other objects if you need to be more discerning, for example if there are multiple similar elements. Do not add extra filler words, and try to be as descriptive as you can in only 2 sentences. The images are all UI components like buttons or inputs. Describe only the UI element within the red bounding box, and nothing outside it."}]
     #           chat_log.append({"role": "user", "content": text})

        def encode_image(image_path):
          with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
        mode_gpt = "32k"
        #if text includes the keyword 'screenshot', take a screenshot with pyauto-gui and reformat chat_log to use the following format to support images and text. 
        #"messages": [ { "role": "user", "content": [ { "type": "text", "text": "What’s in this image?" }, { "type": "image_url", "image_url": { "url": "https://upload.wikimedia.org/wikipedia/commons/thumb/d/dd/Gfp-wisconsin-madison-the-nature-boardwalk.jpg/2560px-Gfp-wisconsin-madison-the-nature-boardwalk.jpg" } } ] } ],
        raw_image.save(resource_path(str(len(nodes)) + "_click.png"))
        base64_image = encode_image(resource_path(str(len(nodes)) + "_click.png"))
        mode_gpt = "website2"
        ###print("SCRNEENSHOT MODE"
        log.append({"role": "user", "content": [{"type": "text", "text":  "Describe the element within the red bounding box with as much details as possible in only 5-10 short words and as few words as possible.  If there is text in the bounding box, say exactly what the text says. Only describe the element within the RED bounding box, but you can explain where it is in relation to other objects if you need to be more discerning, for example if there are multiple similar elements."}, {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}]})
        ###print("screenshot")
        data = {
            "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
            "messages": log,
            "temperature": 0.7,
            "max_tokens": 5000,
            "stream": True  # Enable streaming to match original behavior
        }
        
        print(data)
        
        # Call local server instead of remote API
        response = requests.post("http://localhost:8000/v1/chat/completions", 
                                headers=headers, 
                                data=json.dumps(data), 
                                stream=True)
        
        import pyautogui
        if response.status_code == 200:
            items = []
            total = ""
            for chunk in response.iter_lines():
                try:
                    if chunk:
                        # Parse SSE format from llama.cpp server
                        chunk_str = chunk.decode('utf-8')
                        if chunk_str.startswith('data: '):
                            chunk_data = chunk_str[6:]  # Remove 'data: ' prefix
                            if chunk_data.strip() == '[DONE]':
                                break
                            
                            chunk_json = json.loads(chunk_data)
                            if 'choices' in chunk_json and len(chunk_json['choices']) > 0:
                                delta = chunk_json['choices'][0].get('delta', {})
                                if 'content' in delta:
                                    content = delta['content']
                                    total += content
                                    print(content, end='', flush=True)
                except Exception as e:
                    # If streaming fails, fall back to non-streaming
                    print(f"Streaming error: {e}")
                    break
                
            # If streaming didn't work or total is empty, try non-streaming
            if not total:
                data["stream"] = False
                response = requests.post("http://localhost:8000/v1/chat/completions", 
                                        headers=headers, 
                                        data=json.dumps(data))
                if response.status_code == 200:
                    result = response.json()
                    total = result["choices"][0]["message"]["content"]
                    print(total)
                else:
                    print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                    total = "Error communicating with local LLM server"
            
            print()  # New line after streaming output
            print(total)
        else:
            print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
            total = "Error communicating with local LLM server"
        
        caption = total
        time.sleep(3)
        x = {"type":"Left Mouse Click", "semanticTarget":"target", "x":100, "y":100}
        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="CLICK " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('Automation Input',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)

        nodes[len(nodes)-1].create_property('X',0,  widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Y',0,  widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Click Type',"Single Left Click",items=["Single Left Click", "Single Right Click","Double Click", "Drag"],  widget_type=NODE_PROP_QCOMBO)

        nodes[len(nodes)-1].create_property('Mode',"Browser",items=["Desktop","Browser","First","Last", "Loop Total Runs","Loop Node Runs"],  widget_type=NODE_PROP_QCOMBO)

        nodes[len(nodes)-1].create_property('Target In English', caption, widget_type=NODE_PROP_QLINEEDIT)

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        this_path = os.path.dirname(os.path.abspath(__file__))
        icon = os.path.join(this_path, 'examples', 'OCRScraper.png')
      #  global_variables["CLICK_" + str(len(global_variables))]
        this_path = os.path.dirname(os.path.abspath(__file__))
        icon = os.path.join(this_path, 'examples', 'Move.png')
        ####print(icon)
        ####print(resource_path(str(len(nodes)) + "_click.png"))
        nodes[len(nodes)-1].set_icon( "Click.png")
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
        node_id = nodes[-1].id
        
        image_node = graph.get_node_by_id(node_id)
        image_node.updateImage(resource_path(str(len(nodes)-1) + "_click.png"))
        graph.auto_layout_nodes()            
     
        graph.fit_to_selection()
        
        graph.viewer().clear_key_state()
        graph.viewer().clearFocus()
        graph.QMainWindow.showMaximized()

    def addMove():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
        x = {"type":"Move Mouse"}
        thread_signals.notifications.emit("Added Node","Move Node Added", "icon.png", None, "No")

        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Move " + str(len(nodes)), data=x))#, color= "#FFFFFF"
        nodes[len(nodes)-1].create_property('X Coordinate',0,  widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Y Coordinate',0,  widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Type', "MoveMouse",    widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Data', json.dumps(x,   cls=NumpyEncoder),widget_type=NODE_PROP_QLINEEDIT)
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
        nodes[len(nodes)-1].updateImage(resource_path("Move.png"))
        graph.auto_layout_nodes()
        graph.fit_to_selection()
   
    def addClick():
        global nodes, screenshot, graph
        
        x = {"type":"Left Mouse Click", "semanticTarget":"target", "x":100, "y":100}
        thread_signals.notifications.emit("Added Node","Magic Clicker Node Added", "icon.png", None, "No")

        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="CLICK " + str(len(nodes)), data=x))
        
        nodes[len(nodes)-1].create_property('Target In English', "target", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Automation Input',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].updateImage(resource_path("Click.png"))
        nodes[len(nodes)-1].create_property('Click Type',"Single Left Click",items=["Single Left Click", "Single Right Click","Double Click", "Drag"],  widget_type=NODE_PROP_QCOMBO)

        nodes[len(nodes)-1].create_property('Mode',"Desktop",items=["Desktop","Browser","First","Last", "Loop Total Runs","Loop Node Runs"],  widget_type=NODE_PROP_QCOMBO)



        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        this_path = os.path.dirname(os.path.abspath(__file__))
        icon = os.path.join(this_path, 'examples', 'OCRScraper.png')
      #  global_variables["CLICK_" + str(len(global_variables))]
        this_path = os.path.dirname(os.path.abspath(__file__))
        icon = os.path.join(this_path, 'examples', 'Move.png')
        ####print(icon)
        ####print(resource_path(str(len(nodes)) + "_click.png"))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
        node_id = nodes[-1].id
        
        image_node = graph.get_node_by_id(node_id)
        nodes[len(nodes)-1].updateImage(resource_path("Click.png"))
        graph.auto_layout_nodes()            
     
        graph.fit_to_selection()
        
        graph.viewer().clear_key_state()
        graph.viewer().clearFocus()
        graph.QMainWindow.showMaximized()


    def addDescribe():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
        x = {"type":"SemanticDescribe", "semanticTarget":"caption"}
        thread_signals.notifications.emit("Added Node","Magic Scraper Node Added", "icon.png", None, "No")

        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Magic Scraper " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].updateImage(resource_path("Describe.png"))


        list_vars = []
        for node in nodes:
            print(node.get_property("name") )
            list_vars.append(node.get_property("name").replace(" ", ""))

        # Use the in-memory bytes to call `blip_caption` directly
        caption = "the state of the screen"
        nodes[len(nodes)-1].create_property('Target In English', caption, widget_type=NODE_PROP_QTEXTEDIT)
        nodes[len(nodes)-1].create_property('Automation Input',"",items=list(graph.global_variables.keys()) + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)

        nodes[len(nodes)-1].create_property('Scrape Browser',"",items=["","Links","Images", "Text", "Headings", "Audios", "Videos", "Textareas", "Tables", "Divs", "Spans", "Headings",  "Buttons" ],  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        nodes[len(nodes)-1].create_property('Add Node Runs To Prompt',"false",items=["false", "true" ],  widget_type=NODE_PROP_QCOMBO)

        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
        graph.global_variables["Magic Scraper Output" + str(len(nodes) -1)] = "the state of the screen"
        graph.global_variables["Magic Scraper All Outputs"] = "The full screen"

        graph.global_variables["Links" + str(len(nodes) -1)] = []
        graph.global_variables["Browser Elements" + str(len(nodes) -1)] = []

        for node in nodes:
            print(node.get_property("name") )
            if "Google Sheets" in node.get_property("name"):
                old_node = node
                x = old_node.get_property("Data")
                sheet_url = old_node.get_property("URL")    
                sheet_name = old_node.get_property("sheet name")
                row1 = old_node.get_property("Row 1")
                row2 = old_node.get_property("Row 2")
                row3 = old_node.get_property("Row 3")

                row4 = old_node.get_property("Row 4")
                Read_write = old_node.get_property("Read/Write")
                row5 = old_node.get_property("Row 5")
                row6 = old_node.get_property("Row 6")
                row7 = old_node.get_property("Row 7")
                row8 = old_node.get_property("Row 8")
                row9 = old_node.get_property("Row 9")
                row10 = old_node.get_property("Row 10")
                name = old_node.get_property("name")
                in_connections = []
                out_connections = []
                data = graph.serialize_session()
                old_connections_id = old_node.id
                

                


                new_node = graph.create_node('nodes.widget.ImageNode',     name=name, data=x)

                for connection in data["connections"]:
                    if connection["in"][0] == old_connections_id:
                        new_node.input(0).connect_to(graph.get_node_by_id(connection["out"][0]).output(0))
                    if connection["out"][0] == old_connections_id:
                        graph.get_node_by_id(connection["in"][0]).input(0).connect_to(new_node.output(0))
                graph.delete_node(old_node, False)
                #remove the node from the list of nodes
                nodes.remove(old_node)
                new_node.create_property('URL', sheet_url, widget_type=NODE_PROP_QLINEEDIT)
                new_node.create_property('Read/Write', Read_write,items=["Read", "Write"] ,  widget_type=NODE_PROP_QCOMBO)
                new_node.create_property('sheet name', sheet_name, widget_type=NODE_PROP_QLINEEDIT)
                new_node.create_property('Row 1',row1,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                new_node.create_property('Row 2',row2,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                new_node.create_property('Row 3',row3,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                new_node.create_property('Row 4',row4,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                new_node.create_property('Row 5',row5,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                new_node.create_property('Row 6',row6,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                new_node.create_property('Row 7',row7,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                new_node.create_property('Row 8',row8,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                new_node.create_property('Row 9',row9,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                new_node.create_property('Row 10',row10,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                graph.global_variables["Gsheets"  + str(len(nodes))] = ""
                nodes.append(new_node)

                new_node.create_property('Data', json.dumps(x, cls=NumpyEncoder))
      
        
        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addCustomCode():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
        x = {"type":"python"}
        labels = ["button", "field", "heading", "iframe", "image", "label", "link", "text"]
        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Python Code " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].updateImage(resource_path("CustomCode.png"))
        nodes[len(nodes)-1].create_property('prompt', "open twitter and generate a post about ai automation", widget_type=NODE_PROP_QTEXTEDIT)

        nodes[len(nodes)-1].create_property('code', "print('hello world')", widget_type=NODE_PROP_QTEXTEDIT)
        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
              
        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addScrape():
        global nodes, graph
        global nodes, screenshot, graph
        #d() # minimize main window to select an element
        #msg = QMessageBox()
        #msg.setWindowTitle("Semantic Scrape")
        #msg.setText("In the next screenshot of the screen, please click the middle of the text you want to scrape next in the automation. It will be saved for subsequent nodes to use.")
        #_ = msg.exec_()
        time.sleep(.25) # to get rid of shadow
        cv2.namedWindow("CLICK")    # create new window
        cv2.setMouseCallback("CLICK", click_coordinatesScrape) # get click coordinaates
        img = pyautogui.screenshot()
        screenshot = cv2.cvtColor(np.array(img), cv2.COLOR_BGR2RGB)
        cv2.imshow("CLICK", screenshot)

        cv2.setWindowProperty("CLICK", cv2.WND_PROP_TOPMOST, 1)
        cv2.waitKey(1)
        # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
    def addEmail():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
        x = {"type":"Email"}
        
        list_vars = []
        for node in nodes:
            print(node.get_property("name") )
            list_vars.append(node.get_property("name").replace(" ", ""))

        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Gmail " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].updateImage(resource_path("Email.png"))

        nodes[len(nodes)-1].create_property('to', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('to variable',"",items=list(graph.global_variables.keys()) + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript" , "user_email"] + ["", "phone_transcript", "email_transcript","sms_transcript"]  + ["","Loop Total Runs","Loop Node Runs"],  widget_type=NODE_PROP_QCOMBO)

        nodes[len(nodes)-1].create_property('subject', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('file', "none", widget_type=NODE_PROP_QLINEEDIT)

        nodes[len(nodes)-1].create_property('body', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Body variable',"",items=list(graph.global_variables.keys()) + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["", "phone_transcript", "email_transcript","sms_transcript"]  + ["","Loop Total Runs","Loop Node Runs"],  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
              
        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addGetFiles():
        global nodes, graph
        graph.global_variables["Download " + str(len(nodes))] = ""

        x = {"type":"getfiles"}
        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Download " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('URL or File', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].updateImage(resource_path("Download.png"))

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
              
        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addRiku():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
        x = {"type":"Riku"}
        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Riku.ai Generator " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('Account Holder', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('API Key', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Prompt ID', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Input 1', "",items=list(global_variables.keys()) ,  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Input 2', "",items=list(global_variables.keys()) ,  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Input 3', "",items=list(global_variables.keys()) ,  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Input 4', "",items=list(global_variables.keys()) ,  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Input 5', "",items=list(global_variables.keys()) ,  widget_type=NODE_PROP_QCOMBO)
        
        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
              
        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addMath():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
        x = {"type":"Math"}
        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Math " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('action', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('value', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('input',"",items=list(global_variables.keys()) ,  widget_type=NODE_PROP_QCOMBO)

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
              
        graph.auto_layout_nodes()
        graph.fit_to_selection()
    import json
    import unittest

    def parse_multiple_escapes(possibly_escaped, max_attempts=3):
        """
        Attempt to parse 'possibly_escaped' as JSON up to 'max_attempts' times.
        Return a Python object if successful; raise ValueError on failure.
        """

        # We'll store the current state in 'parsed'
        parsed = possibly_escaped

        for attempt in range(max_attempts):
            # If it's already a dict, nothing more to do
            if isinstance(parsed, dict):
                return parsed

            # If it's a string, try json.loads
            if isinstance(parsed, str):
                try:
                    parsed = json.loads(parsed)
                except json.JSONDecodeError:
                    # As soon as we can't parse, break (or raise)
                    raise ValueError(
                        f"Failed to parse JSON on attempt {attempt+1}: {parsed}"
                    )
            else:
                # It's not a string or dict, so we can't parse further
                raise ValueError(
                    f"Unsupported type after {attempt} parse attempts: {type(parsed)}"
                )

        # After max_attempts, if we still don't have a dict:
        if not isinstance(parsed, dict):
            raise ValueError(
                "Could not parse string into a dictionary after "
                f"{max_attempts} attempts. Final type is {type(parsed)}."
            )
        return parsed

    def addGsheets():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
        for node in nodes:
            print(node.get_property("name") )
            if "Magic Scraper" in node.get_property("name"):
                graph.global_variables["Magic Scraper Output" + node.get_property("name").split("Magic Scraper ")[1]] = "none"
            if "GPT4" in node.get_property("name"):
                graph.global_variables["GPT4" + node.get_property("name").split("GPT4")[1]] = "none"

        list_vars = []
        for node in nodes:
            print(node.get_property("name") )
            list_vars.append(node.get_property("name").replace(" ", ""))
        x = {"type":"Gsheets"}
        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Google Sheets " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('URL', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Read/Write/Create', "Read",items=["Read", "Read Rows", "Write", "Create"] ,  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Cell Row', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Cell Column', "none", widget_type=NODE_PROP_QLINEEDIT)

        nodes[len(nodes)-1].create_property('sheet name', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Row 1',"",items=list(graph.global_variables.keys()) + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Row 2',"",items=list(graph.global_variables.keys()) + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Row 3',"",items=list(graph.global_variables.keys()) + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Row 4',"",items=list(graph.global_variables.keys()) + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Row 5',"",items=list(graph.global_variables.keys()) + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Row 6',"",items=list(graph.global_variables.keys()) + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Row 7',"",items=list(graph.global_variables.keys()) + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Row 8',"",items=list(graph.global_variables.keys()) + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Row 9',"",items=list(graph.global_variables.keys()) + list_vars +   ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Row 10',"",items=list(graph.global_variables.keys()) + list_vars  + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
        graph.global_variables["Gsheets"  + str(len(nodes))] = ""
        nodes[len(nodes)-1].updateImage(resource_path("GoogleSheets.png"))


        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
              
        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addOpen():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
        x = {"type":"Open"}
        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Open Program " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('program', chrome_location + " --password-store=basic https://cheatlayer.com", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].updateImage(resource_path("OpenProgram.png"))

        list_vars = []
        for node in nodes:
            print(node.get_property("name") )
            list_vars.append(node.get_property("name").replace(" ", ""))
      #  nodes[len(nodes)-1].create_property('arguments', "", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Automated Input',"",items=list(graph.global_variables.keys()) + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)
 #       nodes[len(nodes)-1].create_property('Automated Lists',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
              
        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addDelay():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
        x = {"type":"Delay"}
        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Wait " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('seconds', "5", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].updateImage(resource_path("Wait.png"))

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
              
        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addLlama():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"vaaddriables""x":500,"y":500}
        x = {"type":"Llama"}
        graph.global_variables["GPT4" + str(len(nodes))] = ""

        nodes.append(graph.create_node('nodes.widget.TextInputNode',     name="Llama " + str(len(nodes)), data=x))
        
        nodes[len(nodes)-1].create_property('Automation Input',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Webhook Input', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Tracking Tag',"None",items=["None", "324" + user_key[:5] ],  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('type output',"false",items=["false", "true" ],  widget_type=NODE_PROP_QCOMBO)

        graph.global_variables["GPT4" + str(len(graph.global_variables))] = ""

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
     
        graph.auto_layout_nodes()
        graph.fit_to_selection()

    def addWordpress():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"vaaddriables""x":500,"y":500}
        x = {"type":"tracker"}

        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Creator " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('Video Description',"",  widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Prompts Used', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].updateImage(resource_path("Extension.png"))


        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
     
        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addTiktok():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"vaaddriables""x":500,"y":500}
        x = {"type":"tracker"}

        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Creator " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('Video Description',"",  widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Prompts Used', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].updateImage(resource_path("Extension.png"))

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
     
        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addYoutube():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"vaaddriables""x":500,"y":500}
        x = {"type":"tracker"}

        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Creator " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('Video Description',"",  widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Prompts Used', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].updateImage(resource_path("Extension.png"))


        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
     
        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addGoogleScholar():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"vaaddriables""x":500,"y":500}
        x = {"type":"tracker"}

        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Creator " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('Video Description',"",  widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Prompts Used', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].updateImage(resource_path("Extension.png"))


        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
     
        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addGoogleNews():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"vaaddriables""x":500,"y":500}
        x = {"type":"tracker"}

        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Creator " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('Video Description',"",  widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Prompts Used', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].updateImage(resource_path("Extension.png"))


        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
     
        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addTracker():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"vaaddriables""x":500,"y":500}
        x = {"type":"tracker"}

        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Creator " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('Video Description',"",  widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Prompts Used', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].updateImage(resource_path("Extension.png"))


        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
     
        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addWebcam():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"vaaddriables""x":500,"y":500}
        x = {"type":"webcam"}
        graph.global_variables["webcam" + str(len(nodes))] = ""

        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Webcam " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('prompt',"",  widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Webhook Input', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Tracking Tag',"None",items=["None", "324" + user_key[:5] ],  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('type output',"false",items=["false", "true" ],  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].updateImage(resource_path("webcam.png"))

        graph.global_variables["webcam" + str(len(graph.global_variables))] = ""

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
     
        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addGPT4():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"vaaddriables""x":500,"y":500}
        x = {"type":"GPT4"}
        graph.global_variables["GPT4" + str(len(nodes))] = ""

        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="GPT4 " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('input', "enter prompt here", widget_type=NODE_PROP_QTEXTEDIT)
        
        list_vars = []
        for node in nodes:
            print(node.get_property("name") )
            list_vars.append(node.get_property("name").replace(" ", ""))

        nodes[len(nodes)-1].create_property('Automation Input',"",items=list(graph.global_variables.keys()) + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)

        nodes[len(nodes)-1].create_property('Webhook Input', "none", widget_type=NODE_PROP_QLINEEDIT)
     #   nodes[len(nodes)-1].create_property('Tracking Tag',"None",items=["None", "324" + user_key[:5] ],  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('type output',"false",items=["false", "true" ],  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Add Node Runs To Prompt',"false",items=["false", "true" ],  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].updateImage(resource_path("GPT4.png"))

        graph.global_variables["GPT4" + str(len(graph.global_variables))] = ""

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
     
        for node in nodes:
            print(node.get_property("name") )
            if "Google Sheets" in node.get_property("name"):
                old_node = node
                x = old_node.get_property("Data")
                sheet_url = old_node.get_property("URL")    
                sheet_name = old_node.get_property("sheet name")
                row1 = old_node.get_property("Row 1")
                row2 = old_node.get_property("Row 2")
                row3 = old_node.get_property("Row 3")

                row4 = old_node.get_property("Row 4")
                Read_write = old_node.get_property("Read/Write")
                row5 = old_node.get_property("Row 5")
                row6 = old_node.get_property("Row 6")
                row7 = old_node.get_property("Row 7")
                row8 = old_node.get_property("Row 8")
                row9 = old_node.get_property("Row 9")
                row10 = old_node.get_property("Row 10")
                name = old_node.get_property("name")
                in_connections = []
                out_connections = []
                data = graph.serialize_session()
                old_connections_id = old_node.id
                

                


                new_node = graph.create_node('nodes.widget.ImageNode',     name=name, data=x)

                for connection in data["connections"]:
                    if connection["in"][0] == old_connections_id:
                        new_node.input(0).connect_to(graph.get_node_by_id(connection["out"][0]).output(0))
                    if connection["out"][0] == old_connections_id:
                        graph.get_node_by_id(connection["in"][0]).input(0).connect_to(new_node.output(0))
                graph.delete_node(old_node, False)
                
                list_vars = []
                for node in nodes:
                    print(node.get_property("name") )
                    list_vars.append(node.get_property("name").replace(" ", ""))

                #remove the node from the list of nodes
                nodes.remove(old_node)
                new_node.create_property('URL', sheet_url, widget_type=NODE_PROP_QLINEEDIT)
                new_node.create_property('Read/Write', Read_write,items=["Read", "Write"] ,  widget_type=NODE_PROP_QCOMBO)
                new_node.create_property('sheet name', sheet_name, widget_type=NODE_PROP_QLINEEDIT)
                new_node.create_property('Row 1',row1,items=list(graph.global_variables.keys()) + list_vars +  ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                new_node.create_property('Row 2',row2,items=list(graph.global_variables.keys())  + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                new_node.create_property('Row 3',row3,items=list(graph.global_variables.keys())  + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                new_node.create_property('Row 4',row4,items=list(graph.global_variables.keys())  + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                new_node.create_property('Row 5',row5,items=list(graph.global_variables.keys())  + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                new_node.create_property('Row 6',row6,items=list(graph.global_variables.keys())  + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                new_node.create_property('Row 7',row7,items=list(graph.global_variables.keys())  + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                new_node.create_property('Row 8',row8,items=list(graph.global_variables.keys())  + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                new_node.create_property('Row 9',row9,items=list(graph.global_variables.keys())  + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                new_node.create_property('Row 10',row10,items=list(graph.global_variables.keys())  + list_vars + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                graph.global_variables["Gsheets"  + str(len(nodes))] = ""
                nodes.append(new_node)

                new_node.create_property('Data', json.dumps(x, cls=NumpyEncoder))
      
        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addElevenLabs():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"vaaddriables""x":500,"y":500}
        x = {"type":"ElevenLabs"}
        graph.global_variables["ElevenLabs" + str(len(nodes))] = ""

        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Eleven Labs Voice " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('Automation Input',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Webhook Input', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('prompt', "none", widget_type=NODE_PROP_QLINEEDIT)

        nodes[len(nodes)-1].create_property('File Name', "voice.mp3", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Play Audio',"true",items=["false", "true" ],  widget_type=NODE_PROP_QCOMBO)

        graph.global_variables["ElevenLabs" + str(len(graph.global_variables))] = ""
        nodes[len(nodes)-1].updateImage(resource_path("Voice.png"))

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
     
        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addSynthesia():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"vaaddriables""x":500,"y":500}
        x = {"type":"Synthesia"}
        graph.global_variables["Synthesia" + str(len(nodes))] = ""

        nodes.append(graph.create_node('nodes.widget.TextInputNode',     name="Synthesia " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('Automation Input',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Webhook Input', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Avatar', "anna_costume1_cameraA", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Background', "green_screen", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('API Key', "green_screen", widget_type=NODE_PROP_QLINEEDIT)

        graph.global_variables["Synthesia" + str(len(graph.global_variables))] = ""

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
     
        graph.auto_layout_nodes()
        graph.fit_to_selection()

    def getRecording():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"vaaddriables""x":500,"y":500}
        x = {"type":"getRecording"}
        graph.global_variables["Recording" + str(len(nodes))] = ""

        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Recording " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('Audio File', "", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Time', "", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Screenshot File', "", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].updateImage(resource_path("record.png"))

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))

        if len(nodes) > 0:
            connectToLastNode(nodes[-1])

        graph.auto_layout_nodes()
        graph.fit_to_selection()

    def AddStableVideo():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
        x = {"type":"StableVideo"}
        labels = ["button", "field", "heading", "iframe", "image", "label", "link", "text"]
        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Synthetic Video " + str(len(nodes)), data=x))


        nodes[len(nodes)-1].create_property('prompt', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('motion bucket', 127, widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('file', "synthetic.mp4", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Automation Input',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)

        nodes[len(nodes)-1].create_property('frames',"",items= ["25_frames_with_svd_xt","14_frames_with_svd"],  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('resolution',"",items=["maintain_aspect_ratio","crop_to_16_9", "use_image_dimensions"],  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].updateImage(resource_path("Synthetic_Video.png"))

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
              
        graph.auto_layout_nodes()
        graph.fit_to_selection()
                        
    def AddStableDiffusion():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
        x = {"type":"StableDiffusion"}
        labels = ["button", "field", "heading", "iframe", "image", "label", "link", "text"]
        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Dalle 3 " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('mode',"Text To Image",items=["Text To Image"],  widget_type=NODE_PROP_QCOMBO)

        nodes[len(nodes)-1].create_property('Automation Input',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)

        nodes[len(nodes)-1].create_property('prompt', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
              
        graph.auto_layout_nodes()
        graph.fit_to_selection()
                        
    def addAIDetector():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
        x = {"type":"AIDetector"}
        labels = ["button", "field", "heading", "iframe", "image", "label", "link", "text"]
        selectors = ["OCR", "Object Detection","Index","First", "Random"]
        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="AIDetector " + str(len(nodes)), data=x))

        nodes[len(nodes)-1].create_property('UImodel', resource_path("model_weights_default.pth"), widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('labels', ','.join(labels), widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('chosen_label',labels[0],items=labels,  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('key',user_key, widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('threshold', 0.157, widget_type=NODE_PROP_QLINEEDIT)

        nodes[len(nodes)-1].create_property('selector',selectors[0],items=selectors,  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('value', "", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('action', "", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Google Sheet Output', "none", widget_type=NODE_PROP_QLINEEDIT)

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))

        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
              
        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addAIScrape():
        global nodes, graph
        # x = {"type":"Move Mouse", "data":"variables""x":500,"y":500}
        x = {"type":"AIClick"}
        labels = ["button", "field", "heading", "iframe", "image", "label", "link", "text"]
        selectors = ["OCR", "Object Detection","Index","First", "Random"]
        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="AIDetector " + str(len(nodes)), data=x))

        nodes[len(nodes)-1].create_property('UImodel', resource_path("cheat_model.pth"), widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('labels',labels[0],items=labels,  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('labels',labels[0],items=labels,  widget_type=NODE_PROP_QCOMBO)

        nodes[len(nodes)-1].create_property('selector', x["selector"], widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('value', x["value"], widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])

        graph.auto_layout_nodes()
        graph.fit_to_selection()

    def getData():
        global nodes
        x = {"type":"getData"}
        nodes.append(graph.create_node('nodes.widget.ImageNode', name="GetData " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('URL', "", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Payload', "", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Cookies', "", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Headers', "", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Type', "getData", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])

        graph.auto_layout_nodes()
        graph.fit_to_selection()
        graph.global_variables["GetData" + str(len(nodes) -1)] = ""

    def download_file():
        global nodes
        default_download = "FILE_NAME_HERE"
        x={"type":"download"}
        nodes.append(graph.create_node('nodes.widget.ImageNode', name="Download " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('Download Location', x["download_location"],
            widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('URL or File', x["download_location"],
            widget_type=NODE_PROP_QLINEEDIT)

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])

        graph.auto_layout_nodes()
        graph.fit_to_selection()
        
    def getScreenshot():
        global nodes, screenshot
        default_download = "screenshot_image.png"
        x={"type":"screenshot"}
        nodes.append(graph.create_node('nodes.widget.ImageNode', name="Screenshot " + str(len(nodes)), data=x))
        nodes[len(nodes)-1].create_property('Download Location', "",
            widget_type=NODE_PROP_QLINEEDIT)

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])

        graph.auto_layout_nodes()
        graph.fit_to_selection()

    def addScroll():
        global nodes
        x = {"type":"scroll"}
        nodes.append(graph.create_node('nodes.widget.ImageNode', name="Scroll " + str(len(nodes)), data=x))#, color= "#FFFFFF"
        nodes[len(nodes)-1].create_property('Distance',100, widget_type=NODE_PROP_QLINEEDIT)

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
        nodes[len(nodes)-1].updateImage(resource_path("Scroll.png"))

        graph.auto_layout_nodes()
        graph.fit_to_selection()

    def addSendData():
        global nodes
        x = {"type":"sendData"}
        nodes.append(graph.create_node('nodes.widget.ImageNode', name="Send Data " + str(len(nodes)), data=x))#, color= "#FFFFFF"
        nodes[len(nodes)-1].create_property('Body Key 1', "key 1", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Body Value 1',"",  widget_type=NODE_PROP_QLINEEDIT)

        nodes[len(nodes)-1].create_property('Body Key 2', "key 2", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Body Value 2',"",  widget_type=NODE_PROP_QLINEEDIT)
        
        nodes[len(nodes)-1].create_property('Body Key 3', "key 3", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Body Value 3',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["","Total Runs", "Current Time", "Current Date", "Total Scraped"],  widget_type=NODE_PROP_QCOMBO)
        
        nodes[len(nodes)-1].create_property('Body Key 4', "key 4", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Body Value 4',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["","Total Runs", "Current Time", "Current Date", "Total Scraped"],  widget_type=NODE_PROP_QCOMBO)
        
        nodes[len(nodes)-1].create_property('Body Key 5', "key 5", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Body Value 5',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["","Total Runs", "Current Time", "Current Date", "Total Scraped"],  widget_type=NODE_PROP_QCOMBO)

        nodes[len(nodes)-1].create_property('Request', "", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('URL', "https://cheatlayer.com/triggers/extension", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        nodes[len(nodes)-1].updateImage(resource_path("SendData.png"))

        if len(nodes) > 0:
            connectToLastNode(nodes[-1])

        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addIfElse():
        global nodes
        x = {"type":"IfElse"}
        nodes.append(graph.create_node('nodes.basic.BasicNodeB', name="If Else " + str(len(nodes)), data=x))#, color= "#FFFFFF"
        nodes[len(nodes)-1].create_property('Variables',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["", "Total Run","Current Time", "Current Date", "Total Scraped"],  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('operator',"includes",items=["includes","equals","greater than","less than","regex match"],  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('condition', "none", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])

        graph.auto_layout_nodes()
        graph.fit_to_selection()

    def addKeypress():
        global nodes
        x = {"type":"keypress_manual"}
        nodes.append(graph.create_node('nodes.widget.ImageNode', name="Keypress " + str(len(nodes)), data=x))#, color= "#FFFFFF"
        nodes[len(nodes)-1].create_property('String',"test", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Saved Values', "None",items=list(["None", "Current Directory"]) ,  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('GPT-4 Mode', "False",items=list(["False", "True"]) ,  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Automation Input',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].updateImage(resource_path("Keypress.png"))

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])

        graph.auto_layout_nodes()
        graph.fit_to_selection()
    
    def openPhone():
        start_ngrok()

    def addGeneral():
        global nodes
        x = {"type":"general"}
        nodes.append(graph.create_node('nodes.widget.ImageNode', name="Generalized Agent " + str(len(nodes)), data=x))#, color= "#FFFFFF"
        nodes[len(nodes)-1].create_property('prompt',"Search Google for dogs.", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
        nodes[len(nodes)-1].updateImage(resource_path("GeneralTrigger.png"))

        graph.auto_layout_nodes()
        graph.fit_to_selection()

    def runOnCheatLayer():
        global nodes
        x = {"type":"cheatlayer"}
        nodes.append(graph.create_node('nodes.widget.ImageNode', name="Cheat Layer Extension " + str(len(nodes)), data=x))#, color= "#FFFFFF"
        nodes[len(nodes)-1].create_property('code',"alert('test')", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('URL',"https://cheatlayer.com", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])

        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addPrint():
        global nodes
        x = {"type":"print"}
        nodes.append(graph.create_node('nodes.widget.ImageNode', name="Print " + str(len(nodes)), data=x))#, color= "#FFFFFF"
        nodes[len(nodes)-1].create_property('Variables',"",items=list(global_variables.keys()) ,  widget_type=NODE_PROP_QCOMBO)
        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])
        nodes[len(nodes)-1].updateImage(resource_path("PrintData.png"))

        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def runSchedule(self, name):
        print(name)
    def trainModels(self):
        global jobs, graph
        #print("Scehdule")
        graph.QMainWindow.UILayer = ModelWindow()
    def scheduleCheat():
        global jobs, graph
        #print("Scehdule")
        graph.QMainWindow.scheduler = SchedulerWindow()
#        graph.QMainWindow.scheduler.redraw()
        #f,_ = QFileDialog.getOpenFileName(graph.QMainWindow, 'Select Cheat To Schedule', 'c:\\',"Cheat Files (*.cheat)")

        #sch = CronSchedule() 
        #job = {'seconds':'*/1','minute':'*', 'hour':'*','week':'*',
         #   'day':'*','weekday':'*','month':'*',
        #    'startDate':datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        #    'endDate':None}
        #jobs.append(job)
        #sch.addJob(runSchedule  ,job,[f])
    def launchAtlas():
        global window
        #window = WelcomeWindow()
        #window.setWindowIcon(QIcon(os.path.join(this_path, 'examples', 'favicon.ico')))

        window = TrainModels(base_questions, custom_questions, sites, prompts)
        window.setWindowIcon(QIcon(os.path.join(this_path, 'examples', 'favicon.ico')))
        window.show()
    def openAtlas():
        global atlas, window2
        #window = WelcomeWindow()
        #window.setWindowIcon(QIcon(os.path.join(this_path, 'examples', 'favicon.ico')))

        window2.show()
    def runSchedule( name):        
        global nodes, graph
        #print("SCHEDULE")
        ##print(name)
        nodes.clear()
        graph.clear_session()
        file = open(name)
        data = json.load(file)
        #print(data)
        for key,node in data["nodes"].items():

            ##print(node)
            if "Start" in node["name"]:
                nodes.append(graph.create_node(node["type_"], name=node["name"]))#, color= "#FFFFFF"
                nodes[len(nodes)-1].create_property('Initial Program', node["custom"]["Initial Program"], widget_type=NODE_PROP_QLINEEDIT)
                this_path = os.path.dirname(os.path.abspath(__file__))
                icon = os.path.join(this_path, 'examples', 'Move.png')
                nodes[len(nodes)-1].set_icon(icon)
                #graph.auto_layout_nodes()
            else:
                x = json.loads(node["custom"]["Data"])
                nodes.append(graph.create_node(node["type_"], name=node["name"], data=x))#, color= "#FFFFFF"
                this_path = os.path.dirname(os.path.abspath(__file__))
                icon = os.path.join(this_path, 'examples', 'Move.png')
                if "Move" in x["type"]:
                    icon = os.path.join(this_path, 'examples', 'Move.png')
                    nodes[len(nodes)-1].create_property('X Coordinate', x["x"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Y Coordinate', x["y"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Type', "Move", widget_type=NODE_PROP_QLINEEDIT)
    
                if "Click" in x["type"]:
                    icon = os.path.join(this_path, 'examples', 'Click.png')
                    nodes[len(nodes)-1].create_property('X Coordinate', x["x"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Y Coordinate', x["y"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Type', "Click", widget_type=NODE_PROP_QLINEEDIT)
    
                if "keypress" in x["type"]:
                    icon = os.path.join(this_path, 'examples', 'Keypress.png')
                    nodes[len(nodes)-1].create_property('Key', x["key"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Type', "Keypress", widget_type=NODE_PROP_QLINEEDIT)
                if "cheatlayer" in x["type"]:
                    #x = {"type":"cheatlayer", "code":node["script"], "url" : node["url"]}
                   # nodes.append(graph.create_node('nodes.widget.ImageNode', name="Cheat Layer Extension " + str(len(nodes)), data=x))#, color= "#FFFFFF"
                    nodes[len(nodes)-1].create_property('code',x["code"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('URL',x["url"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('key',user_key, widget_type=NODE_PROP_QLINEEDIT)

                nodes[len(nodes)-1].create_property('Data', json.dumps(x), widget_type=NODE_PROP_QLINEEDIT)
    
                nodes[len(nodes)-1].set_icon(icon)
        graph.auto_layout_nodes()
         
            # crate a backdrop node and wrap it around
            # "custom port node" and "group node".
            # fit node selection to the viewer.
        graph.fit_to_selection()
        graph.playRecording()

    def newCheat():
        global nodes, graph
        graph.history = []
        nodes = []
        graph.clear_session()
        nodes.append(graph.create_node('nodes.basic.BasicNodeA', name="Start Node " + str(len(nodes)), data={"type":"Start Node", "x": 0, "y": 0, "Application":"google-chrome https://cheatlayer.com"}))#, color= "#FFFFFF"
        nodes[len(nodes)-1].create_property('Initial Program', chrome_location + " --password-store=basic https://cheatlayer.com", widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Data',  json.dumps({"type":"Start Node", "x": 0, "y": 0, "Application":"google-chrome https://cheatlayer.com"}, cls=NumpyEncoder), widget_type=NODE_PROP_QLINEEDIT)
        this_path = os.path.dirname(os.path.abspath(__file__))
        icon = os.path.join(this_path, 'examples', 'Move.png')
        nodes[len(nodes)-1].set_icon(icon)
        graph.auto_layout_nodes()
        graph.fit_to_selection()

    def openFileCheat(file):
        global nodes, graph
  
        file = open(file)
        data = json.load(file)
        node_dict = {}
        for key,node in data["nodes"].items():

            if "Start" in node["name"]:
                nodes.append(graph.create_node(node["type_"], name=node["name"]))#, color= "#FFFFFF"
                nodes[len(nodes)-1].create_property('Initial Program', node["custom"]["Initial Program"], widget_type=NODE_PROP_QLINEEDIT)
                nodes[len(nodes)-1].create_property('Data',  json.dumps({"type":"Start Node", "x": 0, "y": 0, "Application":node["custom"]["Initial Program"]}, cls=NumpyEncoder), widget_type=NODE_PROP_QLINEEDIT)

                this_path = os.path.dirname(os.path.abspath(__file__))
                icon = os.path.join(this_path, 'examples', 'Move.png')
                nodes[len(nodes)-1].set_icon(icon)
                graph.auto_layout_nodes()
            else:
                x = json.loads(node["custom"]["Data"])
                if "IfElse" in x["type"]:
                    nodes.append(graph.create_node('nodes.basic.BasicNodeB', name=node["name"], data=x))
                else:
                    nodes.append(graph.create_node(node["type_"], name=node["name"], data=x))#, color= "#FFFFFF"
                this_path = os.path.dirname(os.path.abspath(__file__))
                icon = os.path.join(this_path, 'examples', 'Move.png')
                if "Move" in x["type"]:
                    icon = os.path.join(this_path, 'examples', 'Move.png')
                    nodes[len(nodes)-1].create_property('X Coordinate', node["custom"]["X Coordinate"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Y Coordinate', node["custom"]["Y Coordinate"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Type', "Move", widget_type=NODE_PROP_QLINEEDIT)
                if "bash" in x["type"]:
                    #icon = os.path.join(this_path, 'examples', 'Move.png')
                    nodes[len(nodes)-1].create_property('code', node["custom"]["code"], widget_type=NODE_PROP_QTEXTEDIT)
                if "python" in x["type"]:
                    #icon = os.path.join(this_path, 'examples', 'Move.png')
                    nodes[len(nodes)-1].create_property('code', node["custom"]["code"], widget_type=NODE_PROP_QTEXTEDIT)
                    
                if "Click" in x["type"]:
                    icon = os.path.join(this_path, 'examples', 'Click.png')
                    nodes[len(nodes)-1].create_property('X Coordinate', node["custom"]["X Coordinate"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Y Coordinate',  node["custom"]["Y Coordinate"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Type', "Click", widget_type=NODE_PROP_QLINEEDIT)
                if "Open" in x["type"]:        
                    nodes[len(nodes)-1].create_property('program',  node["custom"]["program"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('arguments',  node["custom"]["arguments"], widget_type=NODE_PROP_QLINEEDIT)

                if "Riku" in x["type"]:
                    nodes[len(nodes)-1].create_property('Account Holder', node["custom"]["Account Holder"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('API Key', node["custom"]["API Key"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Prompt ID', node["custom"]["Prompt ID"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Input 1', node["custom"]["Input 1"],items=list(global_variables.keys()) ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Input 2',  node["custom"]["Input 2"],items=list(global_variables.keys()) ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Input 3',  node["custom"]["Input 3"],items=list(global_variables.keys()) ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Input 4',  node["custom"]["Input 4"],items=list(global_variables.keys()) ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Input 5',  node["custom"]["Input 5"],items=list(global_variables.keys()) ,  widget_type=NODE_PROP_QCOMBO)
        
                if "AIDetector" in x["type"]:
                    labels = ["button", "field", "heading", "iframe", "image", "label", "link", "text"]
                    selectors = ["OCR", "Object Detection","Index","First", "Random"]
            
                    nodes[len(nodes)-1].create_property('UImodel', node["custom"]["UImodel"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('labels', node["custom"]["labels"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('chosen_label',node["custom"]["chosen_label"],items=labels,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('key',node["custom"]["key"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('threshold', node["custom"]["threshold"], widget_type=NODE_PROP_QLINEEDIT)
            
                    nodes[len(nodes)-1].create_property('selector',node["custom"]["selector"],items=selectors,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('value', node["custom"]["value"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('action', node["custom"]["action"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Google Sheet Output', node["custom"]["Google Sheet Output"], widget_type=NODE_PROP_QLINEEDIT)

                if "StableDiffusion" in x["type"]:
                    nodes[len(nodes)-1].create_property('mode',node["custom"]["mode"],items=["Text To Image"],  widget_type=NODE_PROP_QCOMBO)
            
                    nodes[len(nodes)-1].create_property('prompt', node["custom"]["prompt"], widget_type=NODE_PROP_QLINEEDIT)
            
                    nodes[len(nodes)-1].create_property('turbo',node["custom"]["turbo"],items=["False","True"],  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('ddim_steps', node["custom"]["ddim_steps"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('seed', node["custom"]["seed"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('tile', 256, widget_type=NODE_PROP_QLINEEDIT)

                    nodes[len(nodes)-1].create_property('variable',node["custom"]["variable"],items=list(global_variables.keys())  + ["Loop Total Runs","Loop Node Runs"],  widget_type=NODE_PROP_QCOMBO)
            
                    nodes[len(nodes)-1].create_property('input image',node["custom"]["input image"],items=["Loop Total Runs","Loop Node Runs"],  widget_type=NODE_PROP_QCOMBO)
            
                    nodes[len(nodes)-1].create_property('height_value', node["custom"]["height_value"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('width_value', node["custom"]["width_value"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('scale', node["custom"]["scale"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('sampler',node["custom"]["sampler"],items=["ddim","ddim_a","euler","euler_a","plms"],  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('precision',node["custom"]["precision"],items=["full","autocast"],  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('iterations', node["custom"]["iterations"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('samples', node["custom"]["samples"], widget_type=NODE_PROP_QLINEEDIT)

                    nodes[len(nodes)-1].create_property('device',node["custom"]["device"],items=["cuda","cpu"],  widget_type=NODE_PROP_QCOMBO)
                if "IfElse" in x["type"]:
                    icon = os.path.join(this_path, 'examples', 'Click.png')
                    nodes[len(nodes)-1].create_property('Variables',node["custom"]["Variables"],items=list(global_variables.keys())  + ["", "Total Run","Current Time", "Current Date", "Total Scraped"],  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('operator', node["custom"]["operator"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('condition', node["custom"]["condition"], widget_type=NODE_PROP_QLINEEDIT)
                if "sendData" in x["type"]:
                    nodes[len(nodes)-1].create_property('Body Key 1', "key 1", widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Body Value 1',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["", "Total Run","Current Time", "Current Date", "Total Scraped"],  widget_type=NODE_PROP_QCOMBO)
               
                    nodes[len(nodes)-1].create_property('Body Key 2', "key 2", widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Body Value 2',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["", "Total Run","Current Time", "Current Date", "Total Scraped"],  widget_type=NODE_PROP_QCOMBO)
                    
                    nodes[len(nodes)-1].create_property('Body Key 3', "key 3", widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Body Value 3',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["", "Total Run","Current Time", "Current Date", "Total Scraped"],  widget_type=NODE_PROP_QCOMBO)
                    
                    nodes[len(nodes)-1].create_property('Body Key 4', "key 4", widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Body Value 4',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["", "Total Run","Current Time", "Current Date", "Total Scraped"],  widget_type=NODE_PROP_QCOMBO)
                    
                    nodes[len(nodes)-1].create_property('Body Key 5', "key 5", widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Body Value 5',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] + ["", "Total Run","Current Time", "Current Date", "Total Scraped"],  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Request', json.dumps(x, cls=NumpyEncoder), widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('URL', "https://cheatlayer.com/triggers/extension", widget_type=NODE_PROP_QLINEEDIT)
                if "print" in x["type"]:
                    nodes[len(nodes)-1].create_property('Variables',node["custom"]["Variables"],items=list(global_variables.keys()) ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
                if "OCR" in x["type"]:
                    nodes[len(nodes)-1].create_property('Bounding Box X1', node["custom"]["Bounding Box X1"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Bounding Box Y1', node["custom"]["Bounding Box Y1"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Bounding Box X2', node["custom"]["Bounding Box X2"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Bounding Box Y2', node["custom"]["Bounding Box Y2"], widget_type=NODE_PROP_QLINEEDIT)
                if "scroll" in x["type"]:
                    nodes[len(nodes)-1].create_property('Distance',node["custom"]["Distance"], widget_type=NODE_PROP_QLINEEDIT)

                if "getData" in x["type"]:
                    nodes[len(nodes)-1].create_property('URL', node["custom"]["URL"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Payload', node["custom"]["Payload"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Cookies', node["custom"]["Cookies"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Headers', node["custom"]["Headers"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Type', "getData", widget_type=NODE_PROP_QLINEEDIT)
                if "screenshot" in x["type"]:       
                   nodes[len(nodes)-1].create_property('Download Location', node["custom"]["Download Location"],widget_type=NODE_PROP_QLINEEDIT)
                if "getfiles" in x["type"]:
                    nodes[len(nodes)-1].create_property('Location', node["custom"]["Location"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Type', node["custom"]["Type"], widget_type=NODE_PROP_QLINEEDIT)
                if "download" in x["type"]:
                    nodes[len(nodes)-1].create_property('Download Location', node["custom"]["Download Location"],
                        widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('URL', node["custom"]["URL"],
                        widget_type=NODE_PROP_QLINEEDIT)
                if "keypress" in x["type"]:
                    icon = os.path.join(this_path, 'examples', 'Keypress.png')
                    
                    nodes[len(nodes)-1].create_property('String', node["custom"]["String"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('GPT-4 Mode',  node["custom"]["GPT-4 Mode"],items=list(["False", "True"]) ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Automation Input', node["custom"]["Automation Input"],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)

                    nodes[len(nodes)-1].create_property('Saved Values', node["custom"]["Saved Values"],items=["None","Current Directory"],  widget_type=NODE_PROP_QCOMBO)

                if "cheatlayer" in x["type"]:
                    #x = {"type":"cheatlayer", "code":node["script"], "url" : node["url"]}
                  #  nodes.append(graph.create_node('nodes.widget.ImageNode', name="Cheat Layer Extension " + str(len(nodes)), data=x))#, color= "#FFFFFF"
                    nodes[len(nodes)-1].create_property('code',node["custom"]["code"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('URL',node["custom"]["URL"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('key',user_key, widget_type=NODE_PROP_QLINEEDIT)
                #nodes[len(nodes)-1].model.set_property("id",key)
                nodes[len(nodes)-1].set_pos(node["pos"][0],node["pos"][1] )

                nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
            node_dict[key] = nodes[len(nodes)-1]
            ##print(key)
            nodes[len(nodes)-1].set_icon(icon)
            #graph.auto_layout_nodes()
            graph.clear_selection()
         
            # crate a backdrop node and wrap it around
            # "custom port node" and "group node".
            # fit node selection to the viewer.
            graph.fit_to_selection()
#          Iterating through the jso

        for connections in data["connections"]:
            ##print(connections)
            if connections["in"][1] == "in A":
                if connections["out"][1] == "out A":
                    node_dict[connections["in"][0]].input(0).connect_to(node_dict[connections["out"][0]].output(0))
                if connections["out"][1] == "out B":
                    node_dict[connections["in"][0]].input(0).connect_to(node_dict[connections["out"][0]].output(1))
            
            if connections["in"][1] == "in B":
                if connections["out"][1] == "out A":
                    node_dict[connections["in"][0]].input(1).connect_to(node_dict[connections["out"][0]].output(0))
                if connections["out"][1] == "out B":
                    node_dict[connections["in"][0]].input(1).connect_to(node_dict[connections["out"][0]].output(1))
    def openCheat(file2 = None):
        global nodes
        nodes.clear()
        graph.clear_session()
                 
        graph.auto_layout_nodes()
        graph.clear_selection()
         
            # crate a backdrop node and wrap it around
            # "custom port node" and "group node".
            # fit node selection to the viewer.
        #print(file2)
        #print("OPEN")
        graph.fit_to_selection()
        graph.last_cheat = file2
        if file2 == None or type(file2) == bool:
            f,_ = QFileDialog.getOpenFileName(graph.QMainWindow, 'Open file', 'c:\\',"")
        else:
            f = resource_path(file2)
        file = open(f)
        data = json.load(file)
        node_dict = {}
        for key,node in data["nodes"].items():
            print(node["custom"])
            if "Start" in node["name"]:
                nodes.append(graph.create_node(node["type_"], name=node["name"]))#, color= "#FFFFFF"
                nodes[len(nodes)-1].create_property('Initial Program', node["custom"]["Initial Program"], widget_type=NODE_PROP_QLINEEDIT)
                nodes[len(nodes)-1].create_property('Data',  json.dumps({"type":"Start Node", "x": 0, "y": 0, "Application":node["custom"]["Initial Program"]}, cls=NumpyEncoder), widget_type=NODE_PROP_QLINEEDIT)

                this_path = os.path.dirname(os.path.abspath(__file__))
                icon = os.path.join(this_path, 'examples', 'Move.png')
                nodes[len(nodes)-1].set_icon(icon)
                graph.auto_layout_nodes()
            else:
                xdata = parse_multiple_escapes(node["custom"]["Data"])
                print(xdata)
                print(xdata["type"])
                if "IfElse" == xdata["type"]:
                    nodes.append(graph.create_node('nodes.basic.BasicNodeB', name=node["name"], data=xdata))
                else:
                    nodes.append(graph.create_node('nodes.widget.ImageNode', name=node["name"], data=xdata))#, color= "#FFFFFF"
                this_path = os.path.dirname(os.path.abspath(__file__))
                icon = os.path.join(this_path, 'examples', 'Move.png')
                if "Move" in xdata["type"]:
                    icon = os.path.join(this_path, 'examples', 'Move.png')
                    nodes[len(nodes)-1].create_property('X Coordinate', node["custom"]["X Coordinate"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Y Coordinate', node["custom"]["Y Coordinate"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Type', "Move", widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].updateImage(resource_path("Move.png"))
                if "tracker" in xdata["type"]:
                    

                    nodes[len(nodes)-1].create_property('Video Description',node["custom"]["Video Description"],  widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Prompts Used', node["custom"]["Prompts Used"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].updateImage(resource_path("creator.png"))

                if "Gsheets" in xdata["type"]:
                    nodes[len(nodes)-1].create_property('URL', node["custom"]["URL"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('key',node["custom"]["key"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].updateImage(resource_path("GoogleSheets.png"))
                    if "Read/Write/Create" in  node["custom"]:
                        nodes[len(nodes)-1].create_property('Read/Write/Create', node["custom"]["Read/Write/Create"],items=["Read", "Read Rows", "Write", "Create"] ,  widget_type=NODE_PROP_QCOMBO)
                    if "Cell Row" in node["custom"]:
                        nodes[len(nodes)-1].create_property('Cell Row', node["custom"]["Cell Row"], widget_type=NODE_PROP_QLINEEDIT)
                    if "Cell Column" in node["custom"]:
                        nodes[len(nodes)-1].create_property('Cell Column', node["custom"]["Cell Column"], widget_type=NODE_PROP_QLINEEDIT)


                    nodes[len(nodes)-1].create_property('sheet name', node["custom"]["sheet name"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Row 1',node["custom"]["Row 1"],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Row 2',node["custom"]["Row 2"],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Row 3',node["custom"]["Row 3"],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Row 4',node["custom"]["Row 4"],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Row 5',node["custom"]["Row 5"],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Row 6',node["custom"]["Row 6"],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Row 7',node["custom"]["Row 7"],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Row 8',node["custom"]["Row 8"],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Row 9',node["custom"]["Row 9"],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Row 10',node["custom"]["Row 10"],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"] ,  widget_type=NODE_PROP_QCOMBO)
                if "Email" in xdata["type"]:
                    nodes[len(nodes)-1].create_property('file', node["custom"]["file"] , widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].updateImage(resource_path("IfElse.png"))

                    if "Body variable" in node["custom"]:
                        nodes[len(nodes)-1].create_property('Body variable',node["custom"]["Body variable"] ,items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"]  + ["Loop Total Runs","Loop Node Runs", node["custom"]["Body variable"]],  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('to', node["custom"]["to"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('subject', node["custom"]["subject"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('body', node["custom"]["body"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('key',user_key, widget_type=NODE_PROP_QLINEEDIT)
                if "semanticScrape" in xdata["type"]:
                    nodes[len(nodes)-1].create_property('Semantic Description', node["custom"]["Semantic Description"], widget_type=NODE_PROP_QLINEEDIT)
                    #nodes[len(nodes)-1].create_property('Mode',node["custom"]["Mode"],items=["First","All","Last"],  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Mode',node["custom"]["Mode"],items=["First","All","Last", "Loop Total Runs","Loop Node Runs"],  widget_type=NODE_PROP_QCOMBO)
                    if "Target" in node["custom"]:
                        nodes[len(nodes)-1].create_property('Target',node["custom"]["Target"],items=["span","link text","link destination", "article", "Twitter posts"],  widget_type=NODE_PROP_QCOMBO)
#                    nodes[len(nodes)-1].create_property('Target',node["custom"]["Target"],items=["span","link text","link destination", "article", "Twitter posts"],  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].updateImage(resource_path("Describe.png"))

                    nodes[len(nodes)-1].create_property('X', node["custom"]["X"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Y', node["custom"]["Y"], widget_type=NODE_PROP_QLINEEDIT)
                if "SemanticDescribe" in xdata["type"]:

                    
                    nodes[len(nodes)-1].updateImage(resource_path("Describe.png"))



                    # Use the in-memory bytes to call `blip_caption` directly
                    caption = "the full screen"
                    if "Description" in node["custom"]:
                        nodes[len(nodes)-1].create_property('Description', node["custom"]["Description"], widget_type=NODE_PROP_QLINEEDIT)
                    if "Target In English" in node["custom"]:
                        nodes[len(nodes)-1].create_property('Target In English', node["custom"]["Target In English"], widget_type=NODE_PROP_QLINEEDIT)
                    # find the screen dimensions and use that 
                    if "Scrape Browser" in node["custom"]:
                        nodes[len(nodes)-1].create_property('Scrape Browser', node["custom"]["Scrape Browser"], items=["","Links","Images", "Text", "Headings", "Audios", "Videos", "Textareas", "Tables", "Divs", "Spans", "Headings",  "Buttons" ],  widget_type=NODE_PROP_QCOMBO)
                    if "Scrape Links" in node["custom"]:
                        nodes[len(nodes)-1].create_property('Scrape Links', node["custom"]["Scrape Links"],items=["True","False"],  widget_type=NODE_PROP_QCOMBO)
                    if "Automation Input" in node["custom"]:
                        nodes[len(nodes)-1].create_property('Automation Input', node["custom"]["Automation Input"],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)
            
                    graph.global_variables["Magic Scraper Output" + str(len(nodes) -1)] = "The full screen"
                    graph.global_variables["Magic Scraper All Outputs"] = "The full screen"
                    
                    graph.global_variables["Links"  + str(len(nodes) -1)] = []
                    graph.global_variables["Link Texts"  + str(len(nodes) -1)] = []


                if "Delay" in xdata["type"]:
                    nodes[len(nodes)-1].create_property('seconds', node["custom"]["seconds"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].updateImage(resource_path("Wait.png"))

                if "bash" in xdata["type"]:
                    #icon = os.path.join(this_path, 'examples', 'Move.png')
                    nodes[len(nodes)-1].create_property('code', node["custom"]["code"], widget_type=NODE_PROP_QTEXTEDIT)
                    nodes[len(nodes)-1].updateImage(resource_path("CustomCode.png"))

                if "python" in xdata["type"]:
                    #icon = os.path.join(this_path, 'examples', 'Move.png')
                    # Search for dynamic tags in the code
                    dynamic_tags = re.findall(r'{{(.*?)}}',node["custom"]["code"])
                    nodes[len(nodes)-1].updateImage(resource_path("CustomCode.png"))

                    # Assuming 'nodes' is a list of node objects and the last node is the one we're working with

                    for tag in dynamic_tags:
                        property_name = tag.strip()
                        # Creates a property for each tag found. Using the tag as property name and the same tag as the value for now.
                        try:
                            nodes[len(nodes)-1].create_property(property_name, '{{' + property_name + '}}',  widget_type=NODE_PROP_QLINEEDIT)
                        except:
                            pass
                    nodes[len(nodes)-1].create_property('code', node["custom"]["code"], widget_type=NODE_PROP_QTEXTEDIT)
                if "GPT4" in xdata["type"]:
                    #nodes[len(nodes)-1].set_property('input',  node["custom"]["input"])
                    nodes[len(nodes)-1].create_property('input', node["custom"]["input"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].updateImage(resource_path("GPT4.png"))
                    if "type output" in node["custom"]:
                        nodes[len(nodes)-1].create_property('type output',  node["custom"]["type output"], widget_type=NODE_PROP_QLINEEDIT)
                            
                    if "Automation Input" in node["custom"]:
                        nodes[len(nodes)-1].create_property('Automation Input',  node["custom"]["Automation Input"], widget_type=NODE_PROP_QLINEEDIT)
                            
                    if "Webhook Input" in node["custom"]:          
                        nodes[len(nodes)-1].create_property('Webhook Input',  node["custom"]["Webhook Input"], widget_type=NODE_PROP_QLINEEDIT)
                    if "Add Node Runs To Prompt" in node["custom"]:  
                        nodes[len(nodes)-1].create_property('Add Node Runs To Prompt',node["custom"]["Add Node Runs To Prompt"],items=["false", "true" ],  widget_type=NODE_PROP_QCOMBO)

                if "StableVideo" in xdata["type"]:
                    nodes[len(nodes)-1].create_property('prompt', node["custom"]["prompt"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('motion bucket',  node["custom"]["motion bucket"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('file',  node["custom"]["file"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Automation Input', node["custom"]["Automation Input"],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)

                    nodes[len(nodes)-1].create_property('frames', node["custom"]["frames"],items= ["25_frames_with_svd_xt","14_frames_with_svd"],  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('resolution', node["custom"]["resolution"],items=["maintain_aspect_ratio","crop_to_16_9", "use_image_dimensions"],  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].updateImage(resource_path("Synthetic_Video.png"))

                if "Click" in xdata["type"]:
                    icon = os.path.join(this_path, 'examples', 'Click.png')
                    nodes[len(nodes)-1].updateImage(resource_path("Click.png"))

                    if "Click Type" in node["custom"]:
                        nodes[len(nodes)-1].create_property('Click Type', node["custom"]["Click Type"],items=["Single Left Click", "Single Right Click","Double Click", "Drag"],  widget_type=NODE_PROP_QCOMBO)
 #                   nodes[len(nodes)-1].create_property('Click Type', node["custom"]["Click Type"],items=["Single Left Click", "Single Right Click","Double Click", "Drag"],  widget_type=NODE_PROP_QCOMBO)
                    if "Target In English" in node["custom"]:

                        nodes[len(nodes)-1].create_property('Target In English', node["custom"]["Target In English"], widget_type=NODE_PROP_QLINEEDIT)
                    if "semanticTarget" in node["custom"]:
                        nodes[len(nodes)-1].create_property('semanticTarget', node["custom"]["semanticTarget"], widget_type=NODE_PROP_QLINEEDIT)
                    if "Mode" in node["custom"]:
                        nodes[len(nodes)-1].create_property('Mode',node["custom"]["Mode"],items=["Desktop", "Browser", "First","Last", "Loop Total Runs","Loop Node Runs"],  widget_type=NODE_PROP_QCOMBO)

                    nodes[len(nodes)-1].create_property('Type', "Click", widget_type=NODE_PROP_QLINEEDIT)
                if "Open" in xdata["type"]:        
                    nodes[len(nodes)-1].create_property('program',  node["custom"]["program"], widget_type=NODE_PROP_QLINEEDIT)
                    if "arguments" in node["custom"]:
                        nodes[len(nodes)-1].create_property('arguments',  node["custom"]["arguments"], widget_type=NODE_PROP_QLINEEDIT)
                    if "automated mode" in node["custom"]:
                        nodes[len(nodes)-1].create_property('automated mode',node["custom"]["automated mode"],items=["Loop Total Runs","Loop Node Runs"],  widget_type=NODE_PROP_QCOMBO)
                        nodes[len(nodes)-1].create_property('automated URL',node["custom"]["automated URL"],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"]  + [node["custom"]["automated URL"]],  widget_type=NODE_PROP_QCOMBO)
                    if "Automated Input" in node["custom"]:
                        array_keys = []

                        nodes[len(nodes)-1].create_property('Automated Input',  node["custom"]["Automated Input"],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"]  + [node["custom"]["Automated Input"]],  widget_type=NODE_PROP_QCOMBO)
                   #     nodes[len(nodes)-1].create_property('Automated Lists',  node["custom"]["Automated Lists"],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"]  + [node["custom"]["Automated Lists"]],  widget_type=NODE_PROP_QCOMBO)

                if "Riku" in xdata["type"]:
                    nodes[len(nodes)-1].create_property('Account Holder', node["custom"]["Account Holder"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('API Key', node["custom"]["API Key"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Prompt ID', node["custom"]["Prompt ID"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Input 1', node["custom"]["Input 1"],items=list(global_variables.keys()) ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Input 2',  node["custom"]["Input 2"],items=list(global_variables.keys()) ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Input 3',  node["custom"]["Input 3"],items=list(global_variables.keys()) ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Input 4',  node["custom"]["Input 4"],items=list(global_variables.keys()) ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Input 5',  node["custom"]["Input 5"],items=list(global_variables.keys()) ,  widget_type=NODE_PROP_QCOMBO)
        
                if "AIDetector" in xdata["type"]:
                    labels = ["button", "field", "heading", "iframe", "image", "label", "link", "text"]
                    selectors = ["OCR", "Object Detection","Index","First", "Random"]
            
                    nodes[len(nodes)-1].create_property('UImodel', node["custom"]["UImodel"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('labels', node["custom"]["labels"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('chosen_label',node["custom"]["chosen_label"],items=labels,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('key',node["custom"]["key"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('threshold', node["custom"]["threshold"], widget_type=NODE_PROP_QLINEEDIT)
            
                    nodes[len(nodes)-1].create_property('selector',node["custom"]["selector"],items=selectors,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('value', node["custom"]["value"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('action', node["custom"]["action"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Google Sheet Output', node["custom"]["Google Sheet Output"], widget_type=NODE_PROP_QLINEEDIT)

                if "StableDiffusion" in xdata["type"]:
                    nodes[len(nodes)-1].create_property('mode',node["custom"]["mode"],items=["Text To Image"],  widget_type=NODE_PROP_QCOMBO)
            
                    nodes[len(nodes)-1].create_property('prompt', node["custom"]["prompt"], widget_type=NODE_PROP_QLINEEDIT)
                    if "turbo" in node["custom"]:
                        nodes[len(nodes)-1].create_property('turbo',node["custom"]["turbo"],items=["False","True"],  widget_type=NODE_PROP_QCOMBO)
                    if "variable" in node["custom"]:
                        nodes[len(nodes)-1].create_property('variable',node["custom"]["variable"],items=list(global_variables.keys())  + ["Loop Total Runs","Loop Node Runs"],  widget_type=NODE_PROP_QCOMBO)
            
                    nodes[len(nodes)-1].create_property('input image',node["custom"]["input image"],items=["Loop Total Runs","Loop Node Runs"],  widget_type=NODE_PROP_QCOMBO)
            

                if "IfElse" in xdata["type"]:
                    icon = os.path.join(this_path, 'examples', 'Click.png')

                    nodes[len(nodes)-1].create_property('Variables',node["custom"]["Variables"],items=list(global_variables.keys())  + [node["custom"]["Variables"], "Total Run","Current Time", "Current Date", "Total Scraped"],  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('operator', node["custom"]["operator"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('condition', node["custom"]["condition"], widget_type=NODE_PROP_QLINEEDIT)
                if "sendData" in xdata["type"]:
                    nodes[len(nodes)-1].create_property('Body Key 1', node["custom"]["Body Key 1"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Body Value 1', node["custom"]["Body Value 1"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].updateImage(resource_path("SendData.png"))

                    nodes[len(nodes)-1].create_property('Body Key 2', node["custom"]["Body Key 2"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Body Value 2', node["custom"]["Body Value 2"], widget_type=NODE_PROP_QLINEEDIT)
                    
                    nodes[len(nodes)-1].create_property('Body Key 3', node["custom"]["Body Key 3"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Body Value 3', node["custom"]["Body Value 3"], widget_type=NODE_PROP_QLINEEDIT)
                    
                    nodes[len(nodes)-1].create_property('Body Key 4', node["custom"]["Body Key 4"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Body Value 4', node["custom"]["Body Value 4"], widget_type=NODE_PROP_QLINEEDIT)
                    
                    nodes[len(nodes)-1].create_property('Body Key 5', node["custom"]["Body Key 5"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Body Value 5', node["custom"]["Body Value 5"], widget_type=NODE_PROP_QLINEEDIT)
                    if "Headers" in node["custom"]:
                        nodes[len(nodes)-1].create_property('Headers', node["custom"]["Headers"], widget_type=NODE_PROP_QLINEEDIT)
                 
                    nodes[len(nodes)-1].create_property('Request', json.dumps(x, cls=NumpyEncoder), widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('URL', node["custom"]["URL"], widget_type=NODE_PROP_QLINEEDIT)
                if "print" in xdata["type"]:
                    nodes[len(nodes)-1].create_property('Variables',node["custom"]["Variables"],items=list(global_variables.keys()) ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Data', json.dumps(xdata, cls=NumpyEncoder))

                    nodes[len(nodes)-1].updateImage(resource_path("PrintData.png"))

                if "OCR" in xdata["type"]:
                    nodes[len(nodes)-1].create_property('Bounding Box X1', node["custom"]["Bounding Box X1"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Bounding Box Y1', node["custom"]["Bounding Box Y1"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Bounding Box X2', node["custom"]["Bounding Box X2"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Bounding Box Y2', node["custom"]["Bounding Box Y2"], widget_type=NODE_PROP_QLINEEDIT)
                if "scroll" in xdata["type"]:
                    nodes[len(nodes)-1].create_property('Distance',node["custom"]["Distance"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].updateImage(resource_path("Scroll.png"))
                if "getData" in xdata["type"]:
                    nodes[len(nodes)-1].create_property('URL', node["custom"]["URL"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Payload', node["custom"]["Payload"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Cookies', node["custom"]["Cookies"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Headers', node["custom"]["Headers"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Type', "getData", widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].updateImage(resource_path("Download.png"))

                if "screenshot" in xdata["type"]:       
                   nodes[len(nodes)-1].create_property('Download Location', node["custom"]["Download Location"],widget_type=NODE_PROP_QLINEEDIT)
                if "getfiles" in xdata["type"]:
                    nodes[len(nodes)-1].create_property('URL or File', node["custom"]["URL or File"], widget_type=NODE_PROP_QLINEEDIT)
#                    nodes[len(nodes)-1].create_property('Type', node["custom"]["Type"], widget_type=NODE_PROP_QLINEEDIT)
                if "download" in xdata["type"]:
                    nodes[len(nodes)-1].create_property('Download Location', node["custom"]["Download Location"],
                        widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('URL', node["custom"]["URL"],
                        widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].updateImage(resource_path("Download.png"))

                if "keypress" in xdata["type"] or "keypress_manual" in xdata["type"]:
                    icon = os.path.join(this_path, 'examples', 'Keypress.png')
                    
                    nodes[len(nodes)-1].create_property('String', node["custom"]["String"], widget_type=NODE_PROP_QLINEEDIT)
                    if "GPT-4 Mode" in node["custom"]:
                        nodes[len(nodes)-1].create_property('GPT-4 Mode',  node["custom"]["GPT-4 Mode"],items=list(["False", "True"]) ,  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Automation Input', node["custom"]["Automation Input"],items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)

                    nodes[len(nodes)-1].create_property('Saved Values', node["custom"]["Saved Values"],items=["None","Current Directory"],  widget_type=NODE_PROP_QCOMBO)

                    nodes[len(nodes)-1].updateImage(resource_path("Keypress.png"))

                if "cheatlayer" in xdata["type"]:
                    #x = {"type":"cheatlayer", "code":node["script"], "url" : node["url"]}
                  #  nodes.append(graph.create_node('nodes.widget.ImageNode', name="Cheat Layer Extension " + str(len(nodes)), data=x))#, color= "#FFFFFF"
                    nodes[len(nodes)-1].create_property('code',node["custom"]["code"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('URL',node["custom"]["URL"], widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('key',user_key, widget_type=NODE_PROP_QLINEEDIT)
                #nodes[len(nodes)-1].model.set_property("id",key)
                nodes[len(nodes)-1].set_pos(node["pos"][0],node["pos"][1] )

                nodes[len(nodes)-1].create_property('Data', json.dumps(xdata, cls=NumpyEncoder))
            node_dict[key] = nodes[len(nodes)-1]
            ##print(key)
            nodes[len(nodes)-1].set_icon(icon)
            #graph.auto_layout_nodes()
            graph.clear_selection()

                        
            #if len(nodes) > 1:
            #    nodes[-1].input(0).connect_to(nodes[-2].output(0))
            # crate a backdrop node and wrap it around
            # "custom port node" and "group node".
            # fit node selection to the viewer.
            graph.fit_to_selection()
#          Iterating through the jso
        if "connections" in data:
            for connections in data["connections"]:
                ##print(connections)
                if connections["in"][1] == "in A":
                    if connections["out"][1] == "out A":
                        node_dict[connections["in"][0]].input(0).connect_to(node_dict[connections["out"][0]].output(0))
                    if connections["out"][1] == "out B":
                        node_dict[connections["in"][0]].input(0).connect_to(node_dict[connections["out"][0]].output(1))
                    if connections["out"][1] == "out":
                        node_dict[connections["in"][0]].input(0).connect_to(node_dict[connections["out"][0]].output(0))

                if connections["in"][1] == "in B":
                    if connections["out"][1] == "out A":
                        node_dict[connections["in"][0]].input(1).connect_to(node_dict[connections["out"][0]].output(0))
                    if connections["out"][1] == "out B":
                        node_dict[connections["in"][0]].input(1).connect_to(node_dict[connections["out"][0]].output(1))
                    if connections["out"][1] == "out":
                        node_dict[connections["in"][0]].input(1).connect_to(node_dict[connections["out"][0]].output(0))
# list  
                if connections["in"][1] == "in":
                    if connections["out"][1] == "out A":
                        node_dict[connections["in"][0]].input(0).connect_to(node_dict[connections["out"][0]].output(0))
                    if connections["out"][1] == "out B":
                        node_dict[connections["in"][0]].input(0).connect_to(node_dict[connections["out"][0]].output(1))
                    if connections["out"][1] == "out":
                        node_dict[connections["in"][0]].input(0).connect_to(node_dict[connections["out"][0]].output(0)) 
# list
    def addFFmpeg():
        global nodes
        x = {"type":"scroll"}
        nodes.append(graph.create_node('nodes.widget.ImageNode', name="Scroll " + str(len(nodes)), data=x))#, color= "#FFFFFF"
        nodes[len(nodes)-1].create_property('Distance',100, widget_type=NODE_PROP_QLINEEDIT)

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        if len(nodes) > 0:
            connectToLastNode(nodes[-1])

        graph.auto_layout_nodes()
        graph.fit_to_selection()
    def addOCR():
        global nodes, screenshot, graph
        #d()
        time.sleep(2)
        cv2.namedWindow("OCR")
        cv2.setMouseCallback("OCR", shape_selection)
        img = pyautogui.screenshot()
        screenshot = cv2.cvtColor(np.array(img), cv2.COLOR_BGR2RGB)
        cv2.imshow("OCR", screenshot)
        cv2.setWindowProperty("OCR", cv2.WND_PROP_TOPMOST, 1)

        cv2.waitKey(1)

        #nodes.append(graph.create_node('nodes.widget.ImageNode', name="OCR", data={}))#, color= "#FFFFFF"
    def addOCRSemantic():
        global nodes, screenshot, graph
        #d()
        time.sleep(2)
        cv2.namedWindow("OCR")
        cv2.setMouseCallback("OCR", shape_selection_semantic)
        img = pyautopytongui.screenshot()
        screenshot = cv2.cvtColor(np.array(img), cv2.COLOR_BGR2RGB)
        cv2.imshow("OCR", screenshot)
        cv2.setWindowProperty("OCR", cv2.WND_PROP_TOPMOST, 1)

        cv2.waitKey(1)
    def openGraph():
        global nodes, graph
        graph.QMainWindow.show()
        graph.QMainWindow.showMaximized()
        #nodes.append(graph.create_node('nodes.widget.ImageNode', name="OCR", data={}))#, color= "#FFFFFF"
    def drawHistory(history):
        global nodes
        move_sequence = []  # To store a sequence of move actions
        keypress_sequence = []  # To store a sequence of keypress actions
        #print("RECORDING HISTORY")

        #print(history)
        
        #print("RECORDING HISTORY")
        for x in history:
            #print(x)
            this_path = os.path.dirname(os.path.abspath(__file__))
            icon_path = {'Move': 'Move.png', 'Click': 'Click.png', 'keypress': 'Keypress.png'}
    
            this_path = os.path.dirname(os.path.abspath(__file__))
            icon = os.path.join(this_path, 'examples', 'Move.png')
            if "Move" in x["type"]:
                # When a move action is encountered, add it to the move sequence list
                move_sequence.append(x)
                continue  # Skip the rest of the loop to start processing the next action

            #generate code for compound keypress sequences
            if "keypress" in x["type"]:
                x["type"] = "keypress_manual"
                x["event"] = str(x["event"])
                #print("ADDING KEYPRESS")
                keypress_sequence.append(x)
                continue

            if keypress_sequence:
                # Create a compound keypress node with the accumulated keypress actions
                compound_keypress_node = graph.create_node('nodes.widget.ImageNode', name='Compound Keypress', data=keypress_sequence)
                compound_keypress_node.create_property('recording', str([keypress['key'] for keypress in keypress_sequence]), widget_type=NODE_PROP_QLINEEDIT)
                #look for KEY_INSERT in the keypress sequence. If it exists, add a magic scraper node first

                compound_keypress_node.create_property('events', str([str(keypress['event']) for keypress in keypress_sequence]), widget_type=NODE_PROP_QLINEEDIT)

                compound_keypress_icon = os.path.join(this_path, 'examples', 'Keypress.png')
                compound_keypress_node.create_property('GPT-4 Mode', "False",items=list(["False", "True"]) ,  widget_type=NODE_PROP_QCOMBO)

                for keypress in keypress_sequence:
                    if 'KEY_INSERT' in keypress['key']:
                        nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Magic Scraper " + str(len(nodes)), data=x))
                        # Use the in-memory bytes to call `blip_caption` directly
                        caption = "the state of the screen"
                        nodes[len(nodes)-1].create_property('Target In English', caption, widget_type=NODE_PROP_QLINEEDIT)
                        nodes[len(nodes)-1].create_property('Automation Input',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)
                        nodes[len(nodes)-1].create_property('Scrape Links',"False",items=["True","False"],  widget_type=NODE_PROP_QCOMBO)
                        graph.global_variables["Magic Scraper Output"  +  str(len(nodes)-1)] = ""
                        graph.global_variables["Magic Scraper All Outputs"] = "The full screen"

                        graph.global_variables["Links"  + str(len(nodes) -1)] = []
                        graph.global_variables["Link Texts"  + str(len(nodes) -1)] = []
                        test = {"type":"SemanticDescribe", "semanticTarget":"caption"}
                        nodes[len(nodes)-1].create_property('Data', json.dumps(test, cls=NumpyEncoder), widget_type=NODE_PROP_QLINEEDIT)

                        if len(nodes) > 0:
                            nodes[-1].input(0).connect_to(nodes[-2].output(0))
                        break
                compound_keypress_node.create_property('Automation Input',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)

                compound_keypress_node.create_property('String', "", widget_type=NODE_PROP_QLINEEDIT)
                compound_keypress_node.create_property('Type', "keypress", widget_type=NODE_PROP_QLINEEDIT)
                compound_keypress_node.set_icon(compound_keypress_icon)
                compound_keypress_node.create_property('Data', {"type":"keypress"}, widget_type=NODE_PROP_QLINEEDIT)

                nodes.append(compound_keypress_node)
                keypress_sequence.clear()
                if len(nodes) > 0:
                    nodes[-1].input(0).connect_to(nodes[-2].output(0))
            # If the current action is not a move or move_sequence is not empty, process the accumulated move actions
            if move_sequence:
                # Create a compound move node with the accumulated move actions
                compound_move_node = graph.create_node('nodes.widget.ImageNode', name='Compound Move', data=move_sequence)
                compound_move_node.create_property('recording', str([(move['x'], move['y']) for move in move_sequence]), widget_type=NODE_PROP_QLINEEDIT)
                compound_move_icon = os.path.join(this_path, 'examples', 'Move.png')
                
                compound_move_node.create_property('X Coordinate', 0, widget_type=NODE_PROP_QLINEEDIT)
                compound_move_node.create_property('Y Coordinate', 0, widget_type=NODE_PROP_QLINEEDIT)
                compound_move_node.create_property('Type', "Move Mouse", widget_type=NODE_PROP_QLINEEDIT)
                compound_move_node.set_icon(compound_move_icon)
                compound_move_node.create_property('Data', {"type":"Move Mouse"}, widget_type=NODE_PROP_QLINEEDIT)
                compound_move_node.updateImage(resource_path("Move.png"))

                nodes.append(compound_move_node)
                if len(nodes) > 0:
                    nodes[-1].input(0).connect_to(nodes[-2].output(0))
                move_sequence.clear()  # Clear the move_sequence list for the next compound move
           

            if "Click" in x["type"]:
              #check if the last node is also a click and make it a doulbe click instead dof adding a new click
              if nodes != None and len(nodes) > 0 and "Click" in nodes[-1].get_property('Type'):
                nodes[-1].set_property('Click Type', "Double Click")
                continue
              icon = os.path.join(this_path, 'examples', 'Click.png')
              
              image = x["image"]
# infer on a      local image
                
              image_size = 384
  
              headers = {
                              "Content-Type": "application/json",
                          }
              log = [{"role": "system", "content": "You describe the element highlighted by a red bounding box in the input image with as many details as possible and as concisely as possible in 5-10 short words and as few words as possible. If there is text in the image, say exactly what the text says. Only describe the element within the RED bounding box, but you can explain where it is in relation to other objects if you need to be more discerning, for example if there are multiple similar elements. Do not add extra filler words, and try to be as descriptive as you can in only 2 sentences. The images are all UI components like buttons or inputs. Describe only the UI element within the red bounding box, and nothing outside it."}]
     #                    chat_log.append({"role": "user", "content": text})
  
              def encode_image(image_path):
                with open(image_path, "rb") as image_file:
                  return base64.b64encode(image_file.read()).decode('utf-8')
              mode_gpt = "32k"
              #if text includes the keyword 'screenshot', take a screenshot with pyauto-gui and reformat chat_log to use the following format to support images and text. 
              #"messages": [ { "role": "user", "content": [ { "type": "text", "text": "What’s in this image?" }, { "type": "image_url", "image_url": { "url": "https://upload.wikimedia.org/wikipedia/commons/thumb/d/dd/Gfp-wisconsin-madison-the-nature-boardwalk.jpg/2560px-Gfp-wisconsin-madison-the-nature-boardwalk.jpg" } } ] } ],
              buffer = BytesIO()
              # Save the screenshot to the buffer in PNG format
              image.save(buffer, format="PNG")
              # Seek to the beginning of the buffer
              buffer.seek(0)
              # Read the buffer content into bytes
              image_bytes = buffer.getvalue()
              # Encode the bytes to base64
              base64_image = base64.b64encode(image_bytes)
              # If you need the base64 string, decode the bytes to a string
              base64_string = base64_image.decode('utf-8')
              mode_gpt = "website2"
              ##print("SCRNEENSHOT MODE"
              log.append({"role": "user", "content": [{"type": "text", "text":  "Describe the element within the red bounding box with as much details as possible in only 5-10 short words and as few words as possible.  If there is text in the bounding box, say exactly what the text says. Only describe the element within the RED bounding box, but you can explain where it is in relation to other objects if you need to be more discerning, for example if there are multiple similar elements."}, {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_string}"}}]})
                    ##print("screenshot")
              data = {
                  "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
                  "messages": log,
                  "temperature": 0.7,
                  "max_tokens": 5000,
                  "stream": True  # Enable streaming to match original behavior
              }
              
              print(data)
              
              # Call local server instead of remote API
              response = requests.post("http://localhost:8000/v1/chat/completions", 
                                      headers=headers, 
                                      data=json.dumps(data), 
                                      stream=True)
              
              import pyautogui
              if response.status_code == 200:
                  items = []
                  total = ""
                  for chunk in response.iter_lines():
                      try:
                          if chunk:
                              # Parse SSE format from llama.cpp server
                              chunk_str = chunk.decode('utf-8')
                              if chunk_str.startswith('data: '):
                                  chunk_data = chunk_str[6:]  # Remove 'data: ' prefix
                                  if chunk_data.strip() == '[DONE]':
                                      break
                                  
                                  chunk_json = json.loads(chunk_data)
                                  if 'choices' in chunk_json and len(chunk_json['choices']) > 0:
                                      delta = chunk_json['choices'][0].get('delta', {})
                                      if 'content' in delta:
                                          content = delta['content']
                                          total += content
                                          print(content, end='', flush=True)
                      except Exception as e:
                          # If streaming fails, fall back to non-streaming
                          print(f"Streaming error: {e}")
                          break
                      
                  # If streaming didn't work or total is empty, try non-streaming
                  if not total:
                      data["stream"] = False
                      response = requests.post("http://localhost:8000/v1/chat/completions", 
                                              headers=headers, 
                                              data=json.dumps(data))
                      if response.status_code == 200:
                          result = response.json()
                          total = result["choices"][0]["message"]["content"]
                          print(total)
                      else:
                          print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                          total = "Error communicating with local LLM server"
                  
                  print()  # New line after streaming output
                  print(total)
              else:
                  print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                  total = "Error communicating with local LLM server"
              
                  print(total)
                  ##print(response.text)
                  ##print(response.json())asdasd
                  caption = total
                  #print(x)
                  click_node = graph.create_node('nodes.widget.ImageNode', name="Click ", data=x)#, color= "#FFFFFF"
                  #print("CLICK NODE start set")
                  click_node.create_property('Target In English', caption, widget_type=NODE_PROP_QLINEEDIT)
                  click_node.create_property('Click Type',"Single Left Click",items=["Single Left Click", "Single Right Click","Double Click", "Drag"],  widget_type=NODE_PROP_QCOMBO)
                  click_node.create_property('Mode',"Browser",items=["Browser","First","Last", "Loop Total Runs","Loop Node Runs"],  widget_type=NODE_PROP_QCOMBO)

                  click_node.create_property('X', x["Bounding Box X1"], widget_type=NODE_PROP_QLINEEDIT)
                  click_node.create_property('Y', x["Bounding Box Y1"], widget_type=NODE_PROP_QLINEEDIT)
                  click_node.create_property('Type', "Click", widget_type=NODE_PROP_QLINEEDIT)
            #         Image.fromarray(np.array(image)).save(resource_path(str(len(nodes)) + "_click.png")
                   #               )
             #        time.sleep(1)
                  ###print(resource_path(str(len(nodes)) + "_click.png"))
                  #print("CLICK NODE finish set")
                  click_node.set_icon( "Click.png")

                  click_node.create_property('Data', {"type":"Left Mouse Click"}, widget_type=NODE_PROP_QLINEEDIT) 
                  click_node.updateImage(resource_path("Click.png"))

                  nodes.append(click_node)
                  image.save(resource_path(str(len(nodes) -1) + "_click.png"))
                  
                  node_id = nodes[-1].id
                  #image_node = graph.get_node_by_id(node_id)
                  #image_node.updateImage(resource_path(str(len(nodes)-1) + "_click.png"))
                  if len(nodes) > 0:
                    nodes[-1].input(0).connect_to(nodes[-2].output(0))
                  
                  #print("CLICK NODE finish")
                    #connectToLastNode(nodes[-1])
        if move_sequence:
            compound_move_node = graph.create_node('nodes.widget.ImageNode', name='Compound Move', data=move_sequence)
            compound_move_node.create_property('recording', str([(move['x'], move['y']) for move in move_sequence]), widget_type=NODE_PROP_QLINEEDIT)
            compound_move_icon = os.path.join(this_path, 'examples', 'Move.png')
            
            compound_move_node.create_property('X Coordinate', 0, widget_type=NODE_PROP_QLINEEDIT)
            compound_move_node.create_property('Y Coordinate', 0, widget_type=NODE_PROP_QLINEEDIT)
            compound_move_node.create_property('Type', "Move Mouse", widget_type=NODE_PROP_QLINEEDIT)
            compound_move_node.set_icon(compound_move_icon)
            nodes.append(compound_move_node)
            compound_move_node.create_property('Data', {"type":"Move Mouse"}, widget_type=NODE_PROP_QLINEEDIT)

            if len(nodes) > 0:
                nodes[-1].input(0).connect_to(nodes[-2].output(0))
        if keypress_sequence:
            compound_keypress_node = graph.create_node('nodes.widget.ImageNode', name='Compound Keypress', data=keypress_sequence)
            compound_keypress_node.create_property('recording', str([keypress['key'] for keypress in keypress_sequence]), widget_type=NODE_PROP_QLINEEDIT)
            compound_keypress_icon = os.path.join(this_path, 'examples', 'Keypress.png')
            compound_keypress_node.create_property('events', str([str(keypress['event']) for keypress in keypress_sequence]), widget_type=NODE_PROP_QLINEEDIT)
            for keypress in keypress_sequence:
                if 'KEY_INSERT' in keypress['key']:
                    nodes.append(graph.create_node('nodes.widget.ImageNode',     name="Magic Scraper " + str(len(nodes)), data=x))
                    # Use the in-memory bytes to call `blip_caption` directly
                    caption = "the state of the screen"
                    nodes[len(nodes)-1].create_property('Target In English', caption, widget_type=NODE_PROP_QLINEEDIT)
                    nodes[len(nodes)-1].create_property('Automation Input',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)
                    nodes[len(nodes)-1].create_property('Scrape Links',"False",items=["True","False"],  widget_type=NODE_PROP_QCOMBO)
                        
                    graph.global_variables["Magic Scraper Output"  +  str(len(nodes)-1)] = ""
                    graph.global_variables["Magic Scraper All Outputs"] = "The full screen"

                    graph.global_variables["Links"  + str(len(nodes) -1)] = []
                    graph.global_variables["Link Texts"  + str(len(nodes) -1)] = []
                    test = {"type":"SemanticDescribe", "semanticTarget":"caption"}
                    nodes[len(nodes)-1].create_property('Data', json.dumps(test, cls=NumpyEncoder), widget_type=NODE_PROP_QLINEEDIT)
                    if len(nodes) > 0:
                        nodes[-1].input(0).connect_to(nodes[-2].output(0))
                    break
            compound_keypress_node.create_property('String', "", widget_type=NODE_PROP_QLINEEDIT)
            compound_keypress_node.create_property('GPT-4 Mode', "False",items=list(["False", "True"]) ,  widget_type=NODE_PROP_QCOMBO)
            compound_keypress_node.create_property('Automation Input',"",items=list(graph.global_variables.keys()) + ["", "phone_transcript", "email_transcript","sms_transcript"],  widget_type=NODE_PROP_QCOMBO)

            compound_keypress_node.create_property('Type', "Keypress", widget_type=NODE_PROP_QLINEEDIT)
            compound_keypress_node.set_icon(icon)
            compound_keypress_node.set_icon(compound_keypress_icon)
            nodes.append(compound_keypress_node)
            compound_keypress_node.create_property('Data', {"type":"keypress"}, widget_type=NODE_PROP_QLINEEDIT)

            if len(nodes) > 0:
                nodes[-1].input(0).connect_to(nodes[-2].output(0))
    
        #print("RECORDING HISTORY finished")
        
    # Custom builtin widgets sasasd
    # create graph controller.
    
    def is_python_installed(version):
       
        return True
    
    def download_and_install_python(version, installer_url):
        installer_name = f"python-{version}-amd64.exe"
        installer_path = os.path.join(os.getcwd(), installer_name)
        
        #print(f"Downloading Python {version} installer...")
        urllib.request.urlretrieve(installer_url, installer_name)
        
        #print("Starting Python installer...")
        # Using the '/quiet' flag installs Python without user interaction - use with caution!
        subprocess.run([installer_path, "/quiet", "InstallAllUsers=1", "PrependPath=1"])
        
        #print("Python has been installed.")
    
    
    python_version = "3.10"
    python_installer_url = "https://www.python.org/ftp/python/3.10.0/python-3.10.0-amd64.exe"
    
  
    verified = False
    global thread_signals
    
    thread_signals= WorkerSignals()

    window2 = ChatApp(thread_signals, openCheat)
    blip_decoder = None
    gpt4_system_prompts = []
    #load workflows from the workfflows.json file 
    if os.path.isfile(get_path('workflows.json')):
        with open(get_path('workflows.json'), 'r') as file:
            workflows = json.load(file)
    prompt_index = 0

    for workflow in workflows:

        if workflow["type"] == "phone":
            with open(workflow["cheat"], 'r') as cheat_file:
                data = json.load(cheat_file)
                for key, node in data["nodes"].items():
                    if "GPT4" in node.get("custom", {}).get("Data", {}):
                        gpt4_system_prompts.append(node["custom"]["input"])
        if workflow["type"] == "webhook":
            #print("webhook")
            print(workflow)
            #graph.global_variables[
    key = 'None'
    ngrok_key = ""
    RoboflowKey = "Y0kBgAMpHpaJPgyhHB1i"
    #load the ngrok key and rob#
    ##print(ngrok_key)
    ##print("USER KEY -3")
    ##print(RoboflowKey)
    ###print(key)

    #start the server with the ngrok key

    
    graph = NodeGraph(addTracker, referral, openPhone, user_email, addGeneral,addWebcam, addElevenLabs, addSynthesia, addLlama, getRecording, AddStableVideo, RoboflowKey, workflows, addScrape, scheduled_jobs, atlas, user_key, drawHistory, verified, addOCR, addPrint, addScroll, addSendData, addIfElse, addKeypress, addMove, addClick,getData,getScreenshot,download_file, openCheat, scheduleCheat, addAIDetector, runOnCheatLayer, trainModels, launchAtlas, AddStableDiffusion, addCustomCode,addEmail,addRiku,addMath,addGsheets,addOpen,addDelay, newCheat, addGetFiles, openAtlas, ChatApp.addResponse, window2, blip_decoder, addDescribe, addOCRSemantic, openaikey, user_plan, thread_signals, addGPT4)
    graph.set_acyclic(False)
    QApplication.setOverrideCursor(QCursor(Qt.CrossCursor))
    

    ##print("USER KEY -5")
    ##print(ngrok_key)
    #start_ngrok(ngrok_key)


    # registered example nodes.
    graph.register_nodes([
        basic_nodes.BasicNodeA,
        basic_nodes.BasicNodeB,
        custom_ports_node.CustomPortsNode,
        group_node.MyGroupNode,
        widget_nodes.DropdownMenuNode,
        widget_nodes.TextInputNode,
        widget_nodes.CheckboxNode,
        widget_nodes.ImageNode
    ])

    # show the node graph widget.
    graph_widget = graph.QMainWindow
    graph_widget.resize(1920, 1080)
    graph_widget.show()
    global_variables = {}
    graph_widget2 = graph.widget
    graph_widget2.resize(1920, 1080)
    graph_widget2.show()

    # auto layout nodes.

    nodes.append(graph.create_node('nodes.basic.BasicNodeA', name="Start Node " + str(len(nodes)), data={"type":"Start Node", "x": 0, "y": 0, "Application":"google-chrome https://cheatlayer.com"}))#, color= "#FFFFFF"
    
    nodes[len(nodes)-1].create_property('Initial Program', chrome_location + " https://cheatlayer.com", widget_type=NODE_PROP_QLINEEDIT)
 #   nodes[len(nodes)-1].create_property('Copy This To Intial Program To Open Chrome', , widget_type=NODE_PROP_QLINEEDIT)

    nodes[len(nodes)-1].create_property('Data',  json.dumps({"type":"Start Node", "x": 0, "y": 0, "Application":"google-chrome https://cheatlayer.com"}, cls=NumpyEncoder), widget_type=NODE_PROP_QLINEEDIT)
    this_path = os.path.dirname(os.path.abspath(__file__))
    icon = os.path.join(this_path, 'examples', 'Move.png')
    nodes[len(nodes)-1].set_icon(icon)
    graph.auto_layout_nodes()

    # crate a backdrop node and wrap it around
    # "custom port node" and "group node".
    # fit node selection to the viewer.
    graph.fit_to_selection()
    # Custom builtin widgets from NodeGraphQt
    # ---------------------------------------sdfasdasgdasdfasdf

    # create a node properties bin widget.
    properties_bin = PropertiesBinWidget(node_graph=graph, parent=graph.QMainWindow)
    properties_bin.setWindowFlags(QtCore.Qt.Tool | QtCore.Qt.Dialog)

    # example show the node properties bin widget when a node is double clicked.
    def display_properties_bin(node):
        if not properties_bin.isVisible():
            properties_bin.show()


    # wire function to "node_double_clicked" signal.
    graph.node_double_clicked.connect(display_properties_bin)
    
    ###print(key)
    def Login():
        global graph, recorder
        
    graph.cheat_scheduler = BackgroundScheduler()

    
# Example condition check function
    def is_key_available():
        global user_key, splash_widget
        # Here you would have code that checks if the key is actually available.
        # For demonstration, let's just simulate it by changing the global variable after a few seconds.
        if user_key != "None":
            #close the splash widget
            splash_widget.close()

    # Simulated function to change the key state (for demonstration purposes)
    def change_key_state():
        global user_key
        user_key = "Available"

    # Your user_key placeholder


    # Create a splash screen with a layout
    splash_widget = QWidget()
    splash_widget.setWindowFlags(Qt.SplashScreen | Qt.FramelessWindowHint)
    splash_widget.setAttribute(Qt.WA_TranslucentBackground)

    layout = QVBoxLayout()
 #   webbrowser.open('https://https://cheatlayer.com/beta.html', new=2)

    # Add GIF animation to the layout
    movie = QMovie("ga.gif")
    movie_label = QLabel()
    movie_label.setMovie(movie)
    layout.addWidget(movie_label)
    movie.start()

    # Add text label beneath the GIF
    text_label = QLabel("Connecting to Cheat Layer Chrome Extension. Please open your browser and make sure the extension is installed.")
    text_label.setAlignment(Qt.AlignCenter)
    layout.addWidget(text_label)

    # Set layout and display the splash
    splash_widget.setLayout(layout)
    splash_widget.show()

    # QTimer to periodically check the condition
    check_timer = QTimer()
    check_timer.timeout.connect(is_key_available)
    check_timer.start(1000)  # Check every second

    #if os.path.isfile(get_path(resource_path('config_cheatlayer.
    while user_key == "None" or user_key == "":
        ##print(user_key)
        time.sleep(2)
        graph.QMainWindow.showMaximized()
        key, entered = QtWidgets.QInputDialog.getText(graph.QMainWindow, 'Cheat Layer Key', 'Enter Cheat Layer Key found at the bottom of the billing page at cheatlayer.com/billing')  
        user_key = key
       # msg = QMessageBox()
       # msg.setWindowTitle("Connecting To Chrome Extension")
       # msg.setText("Please open the browser to allow Cheat Layer Extension to connect to the Desktop version. Re-install the extension locally and remove it from all other devices. Make sure the browser with Cheat Layer is open when opening Desktop.")
       # _ = msg.exec_()
    graph.user_key = user_key 
    graph.job_runner.user_key = user_key
    graph.user_email = user_email
    ##print(user_key)
    ##print({'id': user_key.replace('\r', '').replace('\n', '')})
    if user_key == 'None':
        graph.QMainWindow.showMaximized()
        key, entered = QtWidgets.QInputDialog.getText(graph.QMainWindow, 'Cheat Layer Key', 'Enter Cheat Layer Key found at cheatlayer.com/billing at the bottom of the page:')  
        user_key = key
        ####print(key)
        r = requests.post("https://cheatlayer.com/user/checkDesktopKeyUnAuth", data={'id': user_key.replace('\r', '').replace('\n', '')},  verify=False)
    #####print(r.status_code, r.reason)
        data = r.json()
        
        if "base_questions" not in data:
            
            msg = QMessageBox()
            msg.setWindowTitle("Invalid Key")
            msg.setText("Connection error or this account does not have access to Cheat Layer Desktop Beta, which requires a subscription. If you have access, please remove and re-download the latest version of Cheat Layer Desktop and make sure the extension is installed in a new chrome profile with the browser open when you open Cheat Layer Desktop. You may need to re-open it due to network issues. Upgrade on the Billing page and contact support@cheatlayer.com for help if you continue to have trouble after re-installing both Desktop and the Extension after logging in at cheatlayer.com. Make sure to first log in at cheatlayer.com before re-installing.")
            _ = msg.exec_()
            ##print("Free trial expired or invalid key. Please check your key and billing page at cheatlayer.com")
        else:
            
            base_questions = []
            custom_questions = []
            sites = []
            prompts = []
            user_plan = data['plan']
            graph.user_plan = user_plan
            graph.job_runner.user_key = user_key
            guidelines = data['guidelines']
            openaikey = data['openaikey']
            creator_stats = data["creator_stats"]
            graph.global_variables["creator_stats"] = creator_stats

            client = OpenAI(api_key=openaikey)
            twilio_auth_token = data['auth']
            twilio_account_sid = data['sid']
            twilio_phone_number = data['phone']
            print(data)
            print(data["lat"])
            print(data["long"])
            graph.lat = data["lat"]
            graph.long = data["long"]
            #gpt4_system_prompts
            agents = data['agents']
            desktop_guidelines = data['desktop_guidelines']
            base_questions = data["base_questions"]
            custom_questions = data["custom_questions"]
            sites = data["sites"]
            prompts = data["prompts"]
            user_email = data["user"]
            graph.user_email = user_email
            new_prompt = data["new_prompt"]
            schedule_json = []
            graph.cheat_scheduler = BackgroundScheduler()
            check_version = data["version"]
            python_version = sys.version
            if "3.10" not in python_version :
                msg = QMessageBox()
                msg.setWindowTitle("Update Python")
                msg.setText("Please update to Python 3.10 to use Cheat Layer Desktop Beta. Please download the latest version at cheatlayer.com")
                _ = msg.exec_()
                ##print("Please update to Python 3.10 to use Cheat Layer Desktop Beta. Please download the latest version at cheatlayer.com")
            if check_version != "8.0.0":
                msg = QMessageBox()
                msg.setWindowTitle("Update Available")
                msg.setText("A new version of Cheat Layer is available. Current version(8.0.0). Please download the latest version at cheatlayer.com:" + check_version)
                _ = msg.exec_()
                ##print("A new version of Cheat Layer is available. Please download the latest version at cheatlayer.com")

            else:
            
           # graph.QMainWindow.scheduler = SchedulerWindow()
          #  graph.QMainWindow.scheduler.setWindowIcon(QIcon(os.path.join(this_path, 'examples', 'favicon.ico')))
                if len(schedule_json) > 0:
                    scheduled_jobs_input = json.loads(schedule_json)
                    for job in scheduled_jobs_input:  
                        job_id = -1                  
                        if "enabled" in job and job["enabled"] == "True":
    
                            job_id = graph.cheat_scheduler.add_job(func=graph.QMainWindow.scheduler.runSchedule, trigger='cron',args=[job["cheat"]],  second='0',day=job["day"], day_of_week=job["weekday"], month=job["month"], hour=job["hour"], minute=job["minute"])

                        job = {"enabled": "True","job": job_id, "seconds":"*","minute":job["minute"], "hour":job["hour"],"week":"*",
                            "day":job["day"],"weekday":job["weekday"],"month":job["month"],"id":"testCheat", "cheat": job["cheat"]}
                        scheduled_jobs.append(job)
                if True: 
                  # Replace #print with any callback that accepts an 'event' arg
                    verified = True

#                    tk.messagebox.showinfo("Cheat Layer",  "Logged in!")
                    ##print("Logged in!")
                   # self.Login()
    # ...   
                    graph.updateSchedule(scheduled_jobs)



                   #if os.path.isfile(get_path('config_cheatlayer.txt')):
                 
                    # create a nodes tree widget.
                    nodes_tree = NodesTreeWidget(node_graph=graph)
                    nodes_tree.set_category_label('nodeGraphQt.nodes', 'Builtin Nodes')
                    nodes_tree.set_category_label('nodes.custom.ports', 'Custom Port Nodes')
                    nodes_tree.set_category_label('nodes.widget', 'Widget Nodes')
                    nodes_tree.set_category_label('nodes.basic', 'Basic Nodes')
                    nodes_tree.set_category_label('nodes.group', 'Group Nodes')
                    # nodes_tree.show()
                    graph.QMainWindow.setWindowTitle("Open Agent Studio 8.0.0")
           #         graph.QMainWindow.setWindowIcon(QIcon(os.path.join(this_path, 'examples', 'favicon.ico')))
                    # create a node palette widget.
                    nodes_palette = NodesPaletteWidget(node_graph=graph)
                    nodes_palette.set_category_label('nodeGraphQt.nodes', 'Builtin Nodes')
                    nodes_palette.set_category_label('nodes.custom.ports', 'Custom Port Nodes')
                    nodes_palette.set_category_label('nodes.widget', 'Widget Nodes')
                    nodes_palette.set_category_label('nodes.basic', 'Basic Nodes')
                    nodes_palette.set_category_label('nodes.group', 'Group Nodes')
                    # nodes_palette.show()


                    lines = []

                    graph.cheat_scheduler.start()
   #                 schedule_test()
                 #   graph.QMainWindow.scheduler = SchedulerWindow()
                 #   graph.QMainWindow.scheduler.setWindowIcon(QIcon(os.path.join(this_path, 'examples', 'favicon.ico')))

                #    window = TrainModels(base_questions, custom_questions, sites, prompts)
                #    window.setWindowIcon(QIcon(os.path.join(this_path, 'examples', 'favicon.ico')))

                #    atlas = WelcomeWindow(base_questions, custom_questions, sites, prompts)
                #    atlas.setWindowIcon(QIcon(os.path.join(this_path, 'examples', 'favicon.ico')))
                    graph.atlasWindow = atlas

                    #window = WelcomeWindow(base_questions, custom_questions, sites, prompts)
                    #window.setWindowIcon(QIcon(os.path.join(this_path, 'examples', 'favicon.ico')))
                 #   graph.QMainWindow.showMaximized()
                   # graph.QMainWindow.scheduler.redraw()
                 #   recorder = ScreenRecorder()
                 #   window2.show()
                    graph.QMainWindow.showMaximized()          
                       # window_web = WebWindow()


                   # window_web.show()
                    USER_CONFIG_PATH = 'config_user.txt'  # Adjust this path as necessary

                    #onboarding_window = OnboardingWindow()
                        # Set the window to stay on top and full screen
                        
                   # onboarding_window.setWindowFlags(onboarding_window.windowFlags() | Qt.WindowStaysOnTopHint)
                   # if ngrok_key == "":
                   #     onboarding_window.show()
       
  #                             
                                   # quickstar = QuickStartMenu()
                           #         quickstar.show()
                    app.exec_()
    else:
        print(user_key)
        r = requests.post("https://cheatlayer.com/user/checkDesktopKeyUnAuth", data={'id': user_key.replace('\r', '').replace('\n', '')},  verify=False)
    ###print(r.status_code, r.reason)
        
        data = r.json()
        
        print(data)
        if len(data["user"]) > 0 and "free trial expired" in data["user"] and False: 
      
            msg = QMessageBox()
            msg.setWindowTitle("Invalid Key")
            msg.setText("This account does not have access to Cheat Layer Desktop Beta or the key was entered incorrectly. Please Upgrade on the Billing page and contact support@cheatlayer.com for help accessing the waitlist.")
            _ = msg.exec_()
            ##print("Free trial expired or invalid key. Please check your key and billing page at cheatlayer.com")

        if True: 
          # Replace #print with any callback that accepts an 'event' arg
            verified = True
            if True:
                Login()
                # create a nodes tree widget.
                nodes_tree = NodesTreeWidget(node_graph=graph)
                nodes_tree.set_category_label('nodeGraphQt.nodes', 'Builtin Nodes')
                nodes_tree.set_category_label('nodes.custom.ports', 'Custom Port Nodes')
                nodes_tree.set_category_label('nodes.widget', 'Widget Nodes')
                nodes_tree.set_category_label('nodes.basic', 'Basic Nodes')
                nodes_tree.set_category_label('nodes.group', 'Group Nodes')
                # nodes_tree.show()
                graph.QMainWindow.setWindowTitle("Open Agent Studio 8.0.0")
                #graph.QMainWindow.setWindowIcon(QIcon(os.path.join(this_path, 'examples', 'favicon.ico')))
                # create a node palette widget.
                nodes_palette = NodesPaletteWidget(node_graph=graph)
                nodes_palette.set_category_label('nodeGraphQt.nodes', 'Builtin Nodes')
                nodes_palette.set_category_label('nodes.custom.ports', 'Custom Port Nodes')
                nodes_palette.set_category_label('nodes.widget', 'Widget Nodes')
                nodes_palette.set_category_label('nodes.basic', 'Basic Nodes')
                nodes_palette.set_category_label('nodes.group', 'Group Nodes')
                # nodes_palette.show()

                graph.cheat_scheduler = BackgroundScheduler()
                lines = []

                graph.QMainWindow.scheduler = SchedulerWindow()
                graph.QMainWindow.scheduler.setWindowIcon(QIcon(os.path.join(this_path, 'examples', 'favicon.ico')))

                ##d()
                palette = QPalette()
                app.setStyle("Fusion")

      #          quickstar = QuickStartMenu()
       #         quickstar.show()
                palette.setColor(QPalette.Window, QColor(255, 255, 255))
                palette.setColor(QPalette.WindowText, Qt.black)
                palette.setColor(QPalette.Base, QColor(255, 255, 255))
                palette.setColor(QPalette.AlternateBase, QColor(255, 255, 255))
                palette.setColor(QPalette.ToolTipBase, Qt.white)
                palette.setColor(QPalette.ToolTipText, Qt.black)
                palette.setColor(QPalette.Text, Qt.black)
                palette.setColor(QPalette.Button, QColor(255, 255, 255))
                palette.setColor(QPalette.ButtonText, Qt.black)
                palette.setColor(QPalette.BrightText, Qt.white)
                palette.setColor(QPalette.Link, QColor(42, 130, 218))
                palette.setColor(QPalette.Highlight, QColor(42, 130, 218))
                palette.setColor(QPalette.HighlightedText, Qt.white)
                app.setPalette(palette)
             #   recorder = ScreenRecorder()
             #   window2.show()
              #  webbrowser.open('https://cheatlayer.com/dashboard', new=2)
                USER_CONFIG_PATH = 'config_user.txt'  # Adjust this path as necessary
            
                openCheat(resource_path("repurpose_tiktok_agent.cheat"))

                #onboarding_window = OnboardingWindow()
                        # Set the window to stay on top and full screen
                
                # Create the icon
                icon = QIcon(resource_path("icon.png"))

                # Create the tray
                tray = QSystemTrayIcon()
                tray.setIcon(icon)
                tray.setVisible(True)

                # Create the menu
                menu = QMenu()
                action = QAction("Project Atlas")
                action.triggered.connect(graph.openAtlas )

                menu.addAction(action)

                action2 = QAction("No-Code Edtior")
                action2.triggered.connect(openGraph)

                menu.addAction(action2)
                # Add a Quit option to the menu.
                quit = QAction("Quit")
                quit.triggered.connect(QApplication.instance().quit)
                menu.addAction(quit)

                #onboarding_window.setWindowFlags(onboarding_window.windowFlags() | Qt.WindowStaysOnTopHint)sdfsdf
                #if ngrok_key == "":
                if mode_engine == "executable":
                    
                    file = open("form.cheat")
                    ##print("start")
                    node_dict = {}
                
                    keys = extract_keys_from_file("form.cheat")  # Customize the source file name if needed
                    form = VariableForm(graph, keys, "form.cheat")
                    form.show()

                    #print(name)
                    data = json.load(file)
                    #virtual_desktop_accessor = ctypes.WinDLL("VirtualDesktopAccessor.dll")
                    #virtual_desktop_accessor.GoToDesktopNumber(current_desktop + 1)
                    current_desktop += 1

                    #for key,node in data["nodes"].items():
                    #    custom = json.loads(node["custom"]["Data"])
                    #    if custom["type"] == "Start Node":
                    #        graph.runNode(data,key, thread_signals, False, True)
                    #        break
                else:
                            
                # Add the menu to the tray
                    tray.setContextMenu(menu)
                    graph.QMainWindow.showMaximized()
                    
                    graph.viewer().clear_key_state()
                    graph.viewer().clearFocus()
                    graph.QMainWindow.showMaximized()
#                    onboarding_window.show()
  #              window_web = WebWindow()dsfdf
  #              window_web.show()
## [{'genera    ted_text': 'two birds are standing next to each other '}]
          #      window.hide()


        # You can nw use `marketing_data` and `cookies` as needed
        # For exampe:
        # handle_cookies(cookies)
                if len(sys.argv) > 1:
                    first_arg = sys.argv[1]
                    print("First argument:", first_arg)
                    graph.playRecording()
                graph.QMainWindow.showMaximized()
                subprocess.call(["python3.10","-m","pip","install","yt-dlp"])
              #  subprocess.call(["nohup", "python3.10","agent_server.py", "&"])
                agent_server = subprocess.Popen(["python3.10","agent_server.py"])
                app.exec_()
