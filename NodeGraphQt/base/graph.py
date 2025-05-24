#!/usr/bin/python
# -*- coding: utf-8 -*-
import copy
import youtubesearchpython 
import json
import io
from gtts import gTTS
import ffmpeg
from urllib.parse import urlparse

from seleniumwire import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

headers = {
    "Content-Type": "application/json",
}
import os
import pathlib
import yt_dlp
from TikTokApi import TikTokApi
import asyncio
from TikTokApi import TikTokApi
import asyncio
import os
import io
import time
import base64
import json
import requests
import pyautogui
import firebase_admin
from firebase_admin import credentials
from firebase_admin import db
import platform

def load_data(file_path):
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"The file '{file_path}' does not exist.")

    with open(file_path, 'r', encoding='utf-8') as file:
        try:
            data = json.load(file)
        except json.JSONDecodeError as e:
            raise json.JSONDecodeError(f"Invalid JSON format: {e.msg}", e.doc, e.pos)

    try:
        marketing_data = data['marketingData']
        cookies = data['cookies']
        email = data["email"]
    except KeyError as e:
        raise KeyError(f"Missing key in JSON data: {e}")

    return marketing_data, cookies, email

def normalize_url(url: str) -> str:
    """
    Convert a TikTok URL like:
      "https://www.tiktok.com/@cheatlayer/video/123456"
    into a standard form:
      "tiktok.com/@cheatlayer/video/123456"
    """
    parsed = urlparse(url.strip())
    # netloc might be "www.tiktok.com" or "tiktok.com" or "..."
    netloc = parsed.netloc.lower()
    # remove leading "www." if present
    if netloc.startswith("www."):
        netloc = netloc[4:]
    # combine netloc + path
    normalized = netloc + parsed.path
    return normalized

def process_marketing_data(marketing_data):
    if isinstance(marketing_data, list):
        print("Marketing Data Loaded Successfully:")
        for item in marketing_data:
            search_target = item.get('searchTarget', '')
            voiceover = item.get('voiceover', '')
            caption = item.get('caption', '')
            print(f"Search Target: {search_target}")
            print(f"Voiceover: {voiceover}")
            print(f"Caption: {caption}\n")
    elif isinstance(marketing_data, str):
        rows = marketing_data.split(';:;')
        # Process the rows as needed
        print("Marketing Data Loaded Successfully:")
        for row in rows:
            print(row)
    else:
        raise TypeError("Unsupported type for marketing_data")

from bs4 import BeautifulSoup
import requests
import re

def get_tiktok_views(url):
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/128.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Referer': 'https://www.tiktok.com/',
            'Cookie': 'tt_webid_v2=1234567890'  # You might need to update this
        }

        # Ensure URL is properly formatted
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url

        # Make request
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        content = response.text

        # Try different patterns to find view count
        patterns = [
            r'"playCount":(\d+)',
            r'play_count["\s:]+(\d+)',
            r'video-views[^>]*>([\d,\.]+)',
            r'>[^<]*?(\d+(?:\.\d+)?[KMB]?) views<'
        ]

        for pattern in patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                view_text = match.group(1)
                
                # Handle K, M, B suffixes
                if 'K' in view_text.upper():
                    return int(float(view_text.upper().replace('K', '')) * 1000)
                elif 'M' in view_text.upper():
                    return int(float(view_text.upper().replace('M', '')) * 1000000)
                elif 'B' in view_text.upper():
                    return int(float(view_text.upper().replace('B', '')) * 1000000000)
                else:
                    return int(''.join(filter(str.isdigit, view_text)))

        return None

    except requests.RequestException as e:
        print(f"Network Error: {str(e)}")
        return None
    except Exception as e:
        print(f"Error: {str(e)}")
        return None

def format_views(views):
    """Format view count to be more readable"""
    if views is None:
        return "Unable to get view count"
    if views >= 1000000000:
        return f"{views/1000000000:.1f}B views"
    elif views >= 1000000:
        return f"{views/1000000:.1f}M views"
    elif views >= 1000:
        return f"{views/1000:.1f}K views"
    else:
        return f"{views} views"

import moviepy
import yt_dlp
import moviepy.editor  # PyInstaller will detect this
dummy_variable = moviepy.editor.VideoFileClip  # Explicit usage ensures bundling
run_counter = 0
import os
from gtts import gTTS
import subprocess
import urllib
import sounddevice as sd
import soundfile as sf
import numpy as np
import os
import wave
import pyaudio
import wave
import time
from tkinter import filedialog as fd
from sneakysnek.recorder import Recorder
import math
import pytesseract
from PIL import ImageFont, ImageDraw, Image
import clipboard
import base64
import sys
import io

from contextlib import redirect_stdout, redirect_stderr

class StreamingStringIO(io.StringIO):
    def __init__(self, thread_signals):
        super().__init__()
        self.thread_signals = thread_signals
        self.buffer = ""

    def write(self, text):
        super().write(text)
        # Add your streaming callback 
        self.buffer += text
        
        # Only emit when we have a complete line or substantial buffer
        if '\n' in text or len(self.buffer) > 80:
            # Use Qt's signal mechanism to safely emit from any thread
            self.thread_signals.notifications.emit("Delay", self.buffer, "icon.png", None, "No")
            self.buffer = ""
        
        # For example, you could send it to a websocket, print it, etc.
        sys.stdout.write(text)
        sys.stdout.flush()

from io import BytesIO
from PIL import Image
import argparse
import cv2
import glob
import ctypes, time, shlex, subprocess
import asyncio
import multiprocessing
import ctypes
import os
import pyttsx3
import sys
node_runs = {}
import time
import json
import urllib.request

from tiktok_uploader.upload import upload_video

os_name = platform.system()
print("Hold right click for 2 seconds then release to end the recording for mouse")
print("Click 'ESC' to end the recording for keyboard")
print("Both are needed to finish recording")
user_key = ""

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)

def get_user_key():
    global user_key
    # Check if config_cheatlayer.txt exists
    config_path = resource_path("config_cheatlayer.txt")
    print("GET USER KEY")
    print(user_key)
    if os.path.exists(config_path):
        # Read the first line from the file
        with open(config_path, 'r') as f:
            user_key = f.readline().strip()

    if user_key == None or user_key == "":
        # If the key is empty, prompt the user to enter the key
        app = QApplication([])
        key, ok = QInputDialog.getText(None, "Enter CheatLayer Key", "Please enter your CheatLayer API key from the bottom of cheatlayer.com/billing:")

        if ok and key:  # User clicked OK and entered something
            user_key = key
            # Write the key back to the config file
            with open(config_path, 'w') as f:
                f.write(user_key)
        else:
            QMessageBox.warning(None, "Warning", "No key provided. The program will exit.")
            exit(1)

    return user_key

# Platform-specific configurations
def get_chrome_location():
    system = platform.system()
    if system == "Windows":
        return "C:/Program Files/Google/Chrome/Application/chrome.exe"
    elif system == "Darwin":  # macOS
        return "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"
    else:  # Linux
        return "google-chrome"

def get_font():
    system = platform.system()
    if system == "Windows":
        return "Arial"
    elif system == "Darwin":  # macOS
        return "Helvetica"
    else:  # Linux
        return "DejaVuSans"

chrome_location = get_chrome_location()
font = get_font()

def detect_os():
    global chrome_location, user_key
    os_name = platform.system()
    print(os_name)
    print("DETECT OS")
    if os_name == "Windows":
        get_user_key()
        chrome_location = "C:/Program Files/Google/Chrome/Application/chrome.exe"
        return "This machine is running Windows."
    elif os_name == "Darwin":
        chrome_location = "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome"
        return "This machine is running macOS."
    elif os_name == "Linux":
        user_key = os.getenv("CHEAT_KEY")
        chrome_location = "google-chrome"
        return "This machine is running Linux."
    else:
        return "Unknown operating system."

print("chrome location")
print(detect_os())
print(os_name)
recording = [] 
count = 0
agent_memory = {}
python_process = None

# Define a function that will execute your code.
def execute_code(code, context):
    exec(code, context)

block_counter = 0
total_errors = 0

# Function to create and start a new process, then return it.
def start_process(x, context):
    # Create the process object.
    process = multiprocessing.Process(target=execute_code, args=(x["code"], context))
    # Start the process.
    process.start()
    return process

import webbrowser
import os
import sys

# Step 2: Define your custom function
def my_custom_write(s):
    return  # Add stars before and after each message as an example modification

class bcolors:
    HEADER = '\033[95m'
    OKBLUE = '\033[94m'
    OKCYAN = '\033[96m'
    OKGREEN = '\033[92m'
    WARNING = '\033[93m'
    FAIL = '\033[91m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'

import websockets
from aiohttp import web
import threading
nodes = []
agent_approval = 0
test_iterations = 0
websocket_state = "waiting"
ws = web.WebSocketResponse()
websocket_data = ""
from aiohttp import web
from PySide2.QtCore import Signal as pyqtSignal
from PySide2 import QtCore as qtOLD

# Rest of the code remains the same

#!/usr/bin/python
# -*- coding: utf-8 -*-
import copy
import youtubesearchpython 
import json
import io
from gtts import gTTS
import ffmpeg
from urllib.parse import urlparse

from seleniumwire import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

import os
import pathlib
import yt_dlp
from TikTokApi import TikTokApi
import asyncio
from TikTokApi import TikTokApi
from TikTokApi import TikTokApi
import asyncio
import os
import io
import time
import base64
import json
import requests
import pyautogui
import firebase_admin
from firebase_admin import credentials
from firebase_admin import db

def load_data(file_path):
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"The file '{file_path}' does not exist.")

    with open(file_path, 'r', encoding='utf-8') as file:
        try:
            data = json.load(file)
        except json.JSONDecodeError as e:
            raise json.JSONDecodeError(f"Invalid JSON format: {e.msg}", e.doc, e.pos)

    try:
        marketing_data = data['marketingData']
        cookies = data['cookies']
        email = data["email"]
    except KeyError as e:
        raise KeyError(f"Missing key in JSON data: {e}")

    return marketing_data, cookies, email
def normalize_url(url: str) -> str:
    """
    Convert a TikTok URL like:
      "https://www.tiktok.com/@cheatlayer/video/123456"
    into a standard form:
      "tiktok.com/@cheatlayer/video/123456"
    """
    parsed = urlparse(url.strip())
    # netloc might be "www.tiktok.com" or "tiktok.com" or "..."
    netloc = parsed.netloc.lower()
    # remove leading "www." if present
    if netloc.startswith("www."):
        netloc = netloc[4:]
    # combine netloc + path
    normalized = netloc + parsed.path
    return normalized

def process_marketing_data(marketing_data):
    if isinstance(marketing_data, list):
        print("Marketing Data Loaded Successfully:")
        for item in marketing_data:
            search_target = item.get('searchTarget', '')
            voiceover = item.get('voiceover', '')
            caption = item.get('caption', '')
            print(f"Search Target: {search_target}")
            print(f"Voiceover: {voiceover}")
            print(f"Caption: {caption}\n")
    elif isinstance(marketing_data, str):
        rows = marketing_data.split(';:;')
        # Process the rows as needed
        print("Marketing Data Loaded Successfully:")
        for row in rows:
            print(row)
    else:
        raise TypeError("Unsupported type for marketing_data")
from bs4 import BeautifulSoup

import requests
import re
def get_tiktok_views(url):
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/128.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.5',
            'Referer': 'https://www.tiktok.com/',
            'Cookie': 'tt_webid_v2=1234567890'  # You might need to update this
        }

        # Ensure URL is properly formatted
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url

        # Make request
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        content = response.text

        # Try different patterns to find view count
        patterns = [
            r'"playCount":(\d+)',
            r'play_count["\s:]+(\d+)',
            r'video-views[^>]*>([\d,\.]+)',
            r'>[^<]*?(\d+(?:\.\d+)?[KMB]?) views<'
        ]

        for pattern in patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                view_text = match.group(1)
                
                # Handle K, M, B suffixes
                if 'K' in view_text.upper():
                    return int(float(view_text.upper().replace('K', '')) * 1000)
                elif 'M' in view_text.upper():
                    return int(float(view_text.upper().replace('M', '')) * 1000000)
                elif 'B' in view_text.upper():
                    return int(float(view_text.upper().replace('B', '')) * 1000000000)
                else:
                    return int(''.join(filter(str.isdigit, view_text)))

        return None

    except requests.RequestException as e:
        print(f"Network Error: {str(e)}")
        return None
    except Exception as e:
        print(f"Error: {str(e)}")
        return None

def format_views(views):
    """Format view count to be more readable"""
    if views is None:
        return "Unable to get view count"
    if views >= 1000000000:
        return f"{views/1000000000:.1f}B views"
    elif views >= 1000000:
        return f"{views/1000000:.1f}M views"
    elif views >= 1000:
        return f"{views/1000:.1f}K views"
    else:
        return f"{views} views"

import moviepy
import platform
import yt_dlp
import moviepy.editor  # PyInstaller will detect this
dummy_variable = moviepy.editor.VideoFileClip  # Explicit usage ensures bundling
run_counter = 0
import os
from gtts import gTTS
import subprocess
#from moviepy.editor import VideoFileClip, AudioFileClip
import urllib
import sounddevice as sd
import soundfile as sf
import numpy as np
import os
import wave
import pyaudio
import wave
import time
from tkinter import filedialog as fd
from sneakysnek.recorder import Recorder
import math
import pytesseract
from PIL import ImageFont, ImageDraw, Image
import clipboard
import base64
import sys
import io

from contextlib import redirect_stdout, redirect_stderr

class StreamingStringIO(io.StringIO):
    def __init__(self, thread_signals):
        super().__init__()
        self.thread_signals = thread_signals
        self.buffer = ""

    def write(self, text):
        super().write(text)
        # Add your streaming callback 


        self.buffer += text
        
        # Only emit when we have a complete line or substantial buffer
        if '\n' in text or len(self.buffer) > 80:
            # Use Qt's signal mechanism to safely emit from any thread
            self.thread_signals.notifications.emit("Delay", self.buffer, "icon.png", None, "No")

            self.buffer = ""
        # For example, you could send it to a websocket, print it, etc.
        sys.stdout.write(text)
        sys.stdout.flush()


from io import BytesIO
from PIL import Image
import argparse
import cv2
import glob
import ctypes, time, shlex, subprocess
import asyncio
import multiprocessing
import ctypes
import os
import pyttsx3
import sys
node_runs = {}
import time
import json

import urllib.request

from tiktok_uploader.upload import upload_video

os_name = ""
print("Hold right click for 2 seconds then release to end the recording for mouse")
print("Click 'ESC' to end the recording for keyboard")
print("Both are needed to finish recording")
user_key = ""
def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
   
    if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)
def get_user_key():
    global user_key
    # Check if config_cheatlayer.txt exists
    config_path = resource_path("config_cheatlayer.txt")
    print("GET USER KEY")
    print(user_key)
    if os.path.exists(config_path):
        # Read the first line from the file
        with open(config_path, 'r') as f:
            user_key = f.readline().strip()

    if user_key == None or user_key == "":
        # If the key is empty, prompt the user to enter the key
        app = QApplication([])
        key, ok = QInputDialog.getText(None, "Enter CheatLayer Key", "Please enter your CheatLayer API key from the bottom of cheatlayer.com/billing:")

        if ok and key:  # User clicked OK and entered something
            user_key = key
            # Write the key back to the config file
            with open(config_path, 'wb') as f:
                f.write(user_key)
        else:
            QMessageBox.warning(None, "Warning", "No key provided. The program will exit.")
            exit(1)

    return user_key
font = ""
chrome_location = "google-chrome"
def detect_os():
    global chrome_location, user_key
    os_name = platform.system()
    print(os_name)
    print("DETECT OS")
    if os_name == "Windows":
        get_user_key()
        chrome_location = "C:/Program Files/Google/Chrome/Application/chrome.exe"
        return "This machine is running Windows."
    elif os_name == "Darwin":
        chrome_location = "/Applications/Google\ Chrome.app"
        return "This machine is running macOS."
    elif os_name == "Linux":
        
        user_key = os.getenv("CHEAT_KEY")
        chrome_location = "google-chrome"
        return "This machine is running Linux."
    else:
        return "Unknown operating system."
print("chrome location")
print(detect_os())
print(os_name)
recording = [] 
count = 0
agent_memory = {}
python_process = None
# Define a function that will execute your code.
def execute_code(code, context):
    exec(code, context)
block_counter = 0
total_errors = 0

# Function to create and start a new process, then return it.
def start_process(x, context):
    # Create the process object.
    process = multiprocessing.Process(target=execute_code, args=(x["code"], context))
    # Start the process.
    process.start()
    return process
import webbrowser
import os
import sys
# Step 2: Define your custom function
def my_custom_write(s):
    return# Add stars before and after each message as an example modification
class bcolors:
    HEADER = '\033[95m'
    OKBLUE = '\033[94m'
    OKCYAN = '\033[96m'
    OKGREEN = '\033[92m'
    WARNING = '\033[93m'
    FAIL = '\033[91m'
    ENDC = '\033[0m'
    BOLD = '\033[1m'
    UNDERLINE = '\033[4m'
#import asyncio
import websockets
#import asyncio
from aiohttp import web
import threading
nodes = []
agent_approval = 0

test_iterations = 0

websocket_state = "waiting"
ws = web.WebSocketResponse()
websocket_data = ""
from aiohttp import web
from PySide2.QtCore import Signal as pyqtSignal
from PySide2 import QtCore as qtOLD

class ScreenshotTaker(threading.Thread):
    def __init__(self,workflows, user_key, user_plan, graph, thread_signals, interval=60):
        super().__init__()
        self.interval = interval  # Set the interval to 60 seconds
        self.user_key = user_key
        self.user_plan = user_plan
        self.graph = graph
        self.prev_screenshot = None

        self.workflows = workflows
        self.thread_signals = thread_signals

    def run(self):
        while True:
            try:
              #  screenshot = self.take_screenshot()
                self.process_screenshot()
                time.sleep(5)
            except Exception as e:
                print(e)
                time.sleep(5)
    def regenCode(self, system_prompt, code):
        headers = {"Content-Type": "application/json",}
        print("run GENERETAE PYTHON node")
        print("widget")
        prompt = ''
        prompt = system_prompt + " code: " + code + "generate working python3.10 code only without quotes or a prefix and explanations around the code.  The code is run using the exec function directly with the following functions that are defined already so you don't need to implement them. Do not implement any of these functions: exec(code,{'resource_path': resource_path, 'moviepy': moviepy, 'thread_signals':thread_signals, 'notification_manager':self.notification_manager,'folder': folder, 'send_complex_message': self.send_complex_message, 'agent_memory': self.global_variables ,'store_analytics': self.store_analytics, 'user_key': self.user_key, 'imageToVideo': self.imageToVideo, 'genSyntheticVideo' : self.genSyntheticVideo, 'genImage': self.genImage, 'genVoice': self.genVoice, 'gpt3Prompt': self.gpt3Prompt, 'last_mouse_y':last_mouse_y,'last_mouse_x': last_mouse_x, '__name__': '__main__','videoSearch': self.videoSearch, 'semanticClick': self.semanticClick, 'keypress': self.keypress , 'semanticMove': self.semanticMove, 'click': self.click, 'move':self.move, 'getData': self.getData, 'sendData': self.sendData} )" 
        log = [{"role": "system", "content": """You are a helpfull assistant that only generates working python code without placeholders, comments, quotes, prefixes that can be run directly. Use selenium with the local browser to use the local cookies for logged in services. Use this code to get the local chrome cookies: def get_chrome_data_dir(self): system = platform.system() if system == "Windows": return os.path.join(os.environ['LOCALAPPDATA'], 'Google', 'Chrome', 'User Data') elif system == "Darwin":  # macOS return os.path.expanduser('~/Library/Application Support/Google/Chrome') else:  # Linux return os.path.expanduser('~/.config/google-chrome') def setup_driver(self):try: options = webdriver.ChromeOptions() # Add the user data directory user_data_dir = self.get_chrome_data_dir() options.add_argument(f"user-data-dir={user_data_dir}") options.add_argument("--profile-directory=Default") # Add additional options options.add_argument("--start-maximized") options.add_argument("--disable-notifications") options.add_argument("--disable-popup-blocking") self.driver = webdriver.Chrome(options=options) print("Chrome driver initialized successfully") except Exception as e: print(f"Error initializing Chrome driver: {str(e)}") print("If Chrome is already running, please close all Chrome windows and try again.") raise """}]
        print(prompt)
        Automation_input = ""
        Webhook_input = ""
        log.append({"role": "user", "content": prompt})
        data = {
            "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
            "messages": log,
            "temperature": 0.7,
            "max_tokens": 5000,
            "stream": True  # Enable streaming to match original behavior
        }
        
        print(data)
        
        # Call local server instead of remote API
        response = requests.post("http://localhost:8000/v1/chat/completions", 
                                headers=headers, 
                                data=json.dumps(data), 
                                stream=True)
        
        import pyautogui
        if response.status_code == 200:
            items = []
            total = ""
            for chunk in response.iter_lines():
                try:
                    if chunk:
                        # Parse SSE format from llama.cpp server
                        chunk_str = chunk.decode('utf-8')
                        if chunk_str.startswith('data: '):
                            chunk_data = chunk_str[6:]  # Remove 'data: ' prefix
                            if chunk_data.strip() == '[DONE]':
                                break
                            
                            chunk_json = json.loads(chunk_data)
                            if 'choices' in chunk_json and len(chunk_json['choices']) > 0:
                                delta = chunk_json['choices'][0].get('delta', {})
                                if 'content' in delta:
                                    content = delta['content']
                                    total += content
                                    print(content, end='', flush=True)
                except Exception as e:
                    # If streaming fails, fall back to non-streaming
                    print(f"Streaming error: {e}")
                    break
                
            # If streaming didn't work or total is empty, try non-streaming
            if not total:
                data["stream"] = False
                response = requests.post("http://localhost:8000/v1/chat/completions", 
                                        headers=headers, 
                                        data=json.dumps(data))
                if response.status_code == 200:
                    result = response.json()
                    total = result["choices"][0]["message"]["content"]
                    print(total)
                else:
                    print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                    total = "Error communicating with local LLM server"
            
            print()  # New line after streaming output
            print(total)
        else:
            print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
            total = "Error communicating with local LLM server"
            code = total
            return code
    def gpt3Prompt(self, system_prompt, description):
    # Setup local llama.cpp server if it's not running
        if not hasattr(self, 'llama_server_process') or self.llama_server_process is None:
            self.setup_local_llama_server()

        log = [{"role": "system", "content": system_prompt}]
        log.append({"role": "user", "content": description})

        headers = {
            "Content-Type": "application/json",
        }

        data = {
            "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
            "messages": log,
            "temperature": 0.7,
            "max_tokens": 5000,
        }

        # Call local server instead of remote API
        response = requests.post("http://localhost:8000/v1/chat/completions", 
                                 headers=headers, 
                                 data=json.dumps(data))

        if response.status_code == 200:
            result = response.json()
            total = result["choices"][0]["message"]["content"]
            print(total)
            return total
        else:
            print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
            return "Error communicating with local LLM server"

    def setup_local_llama_server(self):
        """
        Download the Gemma-3 model if needed and start a local llama.cpp server with vision capabilities
        """
        print("Setting up local llama.cpp server with Gemma-3 4B-IT model...")

        # Clone llama.cpp repository
        if not os.path.exists("llama.cpp"):
            subprocess.run(["git", "clone", "https://github.com/ggerganov/llama.cpp.git"], check=True)

        # Build llama.cpp with vision support
        if not os.path.exists("llama.cpp/server"):
            os.chdir("llama.cpp")
            subprocess.run(["make", "server", "LLAMA_BLAS=ON", "LLAMA_BLAS_VENDOR=OpenBLAS", f"-j{os.cpu_count()}"], check=True)
            os.chdir("..")

        # Create models directory if it doesn't exist
        os.makedirs("models/gemma3", exist_ok=True)

        # Check if model and projector exist, download if needed
        model_path = "models/gemma3/gemma-3-4b-it-q4_k_xl.gguf"
        if not os.path.exists(model_path):
            print("Downloading Gemma-3 4B-IT model (this may take a while)...")
            subprocess.run([
                "wget", "-O", model_path,
                "https://huggingface.co/bartowski/google_gemma-3-4b-it-GGUF/resolve/main/gemma-3-4b-it-Q4_K_XL.gguf"
            ], check=True)

        projector_path = "models/gemma3/mmproj-gemma-3-4b-it-f16.gguf"
        if not os.path.exists(projector_path):
            print("Downloading multimodal projector (this may take a while)...")
            subprocess.run([
                "wget", "-O", projector_path,
                "https://huggingface.co/bartowski/google_gemma-3-4b-it-GGUF/resolve/main/mmproj-gemma-3-4b-it-f16.gguf"
            ], check=True)

        # Check if server is already running
        try:
            response = requests.get("http://localhost:8000/v1/models")
            if response.status_code == 200:
                print("Local llama.cpp server is already running")
                return
        except:
            pass  # Server not running, we'll start it
        
        # Start the server
        print("Starting llama.cpp server with Gemma-3 4B-IT model...")

        # Determine number of CPU threads
        cpu_threads = os.cpu_count()
        if cpu_threads > 4:
            cpu_threads = cpu_threads - 2  # Leave some CPU for other tasks

        self.llama_server_process = subprocess.Popen([
            "llama.cpp/server",
            "-m", model_path,
            "--mmproj", projector_path,
            "--host", "0.0.0.0",
            "--port", "8000",
            "-t", str(cpu_threads),
            "-c", "4096",
            "--temperature", "0.7"
        ])

        # Give server time to start up
        time.sleep(5)
        print("Local llama.cpp server is running with Gemma-3 4B-IT!")

    @staticmethod
    def take_screenshot():
        screenshot = pyautogui.screenshot()
        return screenshot
    def runNodes(self, cheat_file, prompt):
        print("RUN nodes")
        print(cheat_file)
        screenshot = pyautogui.screenshot()
        
        # Convert current screenshot to PIL Image (pyautogui's screenshot is already a PIL Image object)
        current_screenshot = screenshot
        print("Is this a screenshot of " + prompt + "?")
        current_screenshot.save(resource_path("crop.png"))
        itm_output = blip_caption2_2(open(resource_path("crop.png"), "rb"),   prompt + ". Answer only yes or no." )
        print(itm_output)
        score = float(itm_output.split("probability of ")[1].split(".\n")[0])
        if score > .99:
            with open(cheat_file) as file:
                data = json.load(file)
                current_desktop = 0  # Assuming `current_desktop` is previously defined
                self.thread_signals.recording_start.emit
                for key, node in data["nodes"].items():
                    custom = json.loads(node["custom"]["Data"])
                    if custom["type"] == "Start Node":
                        # Pass job_folder_path to the runNode function:
                        self.graph.runNode(data, key, self.thread_signals, False, True, folder="default")
                        break
        else:
            print("No")
            self.thread_signals.notifications.emit("No", "The state is no longer:" + prompt, "icon.png", lambda: None, "No")

    def OpenNodes(self, cheat_file):
        print("RUN nodes")
        print(cheat_file)
        if True:
            with open(cheat_file) as file:
                data = json.load(file)
                current_desktop = 0  # Assuming `current_desktop` is previously defined
                self.thread_signals.recording_start.emit
                for key, node in data["nodes"].items():
                    custom = json.loads(node["custom"]["Data"])
                    if custom["type"] == "Start Node":
                        # Pass job_folder_path to the runNode function:
                        self.graph.runNode(data, key, self.thread_signals, False, True, folder="default")
                        break
        else:
            print("No")
            self.thread_signals.notifications.emit("No", "The state is no longer:" + prompt, "icon.png", lambda: None, "No")

    def pixel_difference(self, image1, image2, threshold=10):
        # This function assumes both inputs are PIL Image objects
        if image1 is None or image2 is None:
        # If either image is None, we can't compute a difference
            return float('inf')  # Return infinity to indicate no valid comparison can be made
    
        np_img1 = np.array(image1)
        np_img2 = np.array(image2)
        
        # Calculate absolute difference
        diff = np.abs(np_img1 - np_img2)
        
        # Threshold differences to filter out minor changes
        significant_diff = diff > threshold
        
        # Count the number of significantly different pixels
        num_diff_pixels = np.sum(significant_diff)
        
        return num_diff_pixels

    def process_screenshot(self):
        screen_width = pyautogui.size()[0]
        screen_height = pyautogui.size()[1]
        screenshot = pyautogui.screenshot(region=(0,0,screen_width-300,screen_height))
        
        # Convert current screenshot to PIL Image (pyautogui's screenshot is already a PIL Image object)
        current_screenshot = screenshot
        # Only proceed if there is no previous screenshot or the current screenshot is significantly different

        if self.prev_screenshot is None or self.pixel_difference(self.prev_screenshot, current_screenshot) > 507847:  # Adjust the threshold based on your needs
            for workflow in self.workflows:
                if workflow["type"] == "semantic":
                    print(workflow)
                    cheat_file = workflow["cheat"]
                    print("Is this a screenshot of " + workflow["prompt"] + "?")
                    current_screenshot.save(resource_path("crop.png"))
                    itm_output = blip_caption2_2(open(resource_path("crop.png"), "rb"),  workflow["prompt"] + ". Answer only yes or no.")

                    score = float(itm_output.split("probability of ")[1].split(".\n")[0])
                    if score > .95:
                        print("Triggered Agent")    
                        icon = "icon.png"
                        if "Gmail" in workflow["cheat"] or "gmail" in workflow["cheat"]:
                            icon = "gmail.png"
                        if "Linkedin" in workflow["cheat"] or "linkedin" in workflow["cheat"]:
                            icon = "linkedin.png"
                        self.thread_signals.notifications.emit(workflow["cheat"].split("/")[-1].split(".")[0].replace("_"," "), workflow["cheat"].split("/")[-1].split(".")[0].replace("_"," "), icon, lambda: self.runNodes(cheat_file, workflow["prompt"]), workflow["prompt"])
        
        # Update the previous screenshot for the next comparison
        self.prev_screenshot = current_screenshot
def start_screenshot_taker():
    screenshot_taker_thread = ScreenshotTaker()
    screenshot_taker_thread.daemon = True  # Set as a daemon so it will be killed once the main thread is dead.
    screenshot_taker_thread.start()
    return screenshot_taker_thread

class Stream(qtOLD.QObject):
    newText = qtOLD.Signal(str)

    def write(self, text):
        self.newText.emit(str(text))

class Server:
    def __init__(self, trigger, setkey, data_graph, runNode, thread_signals, sendMessageRTCAsync, graph):
        self.websocket_state = "waiting"
        self.websocket_data = ""
        self.trigger = trigger
        self.thread_signals = thread_signals
        self.setkey = setkey
        self.runNode = runNode
        self.data_graph = data_graph
        self.reconnect_needed = False
        self.sendMessageRTCAsync =  sendMessageRTCAsync
        self.graph = graph


    async def handle_websocket(self, request):
        global ws
        ws = web.WebSocketResponse()
        await ws.prepare(request)
        try:
            async for msg in ws:
                if msg.type == web.WSMsgType.TEXT:
                    if msg.data == 'close':
                        await ws.close()
                    else:
                        await self.process_message(msg)
                elif msg.type == web.WSMsgType.ERROR:
                    print(f'WebSocket connection closed with exception {ws.exception()}')
        except Exception as e:
            print(e)
            print('websocket connection closed')
            self.reconnect_needed = True
        finally:
            return ws

    async def reconnect_websocket(self, app):
        while True:
            if self.reconnect_needed:

                try:
                    print("Attempting to reconnect...")
                    self.thread_signals.reconnectself.emit()
                    self.reconnect_needed = False
                    print("WebSocket reconnected")
                except Exception as e:
                    print(e)
                    print("WebSocket reconnection failed. Retrying in 5 seconds...")
                    await asyncio.sleep(5)
            else:
                await asyncio.sleep(1)

    def verifyStep(self, current_goal, goal_type , last_stored_data):
        
        global verified_steps, last_screenshot
        result = False

        screenshot = pyautogui.screenshot()
        if goal_type == "general":
            result = True
        elif goal_type == "schedule":
            result = True
            
        elif goal_type == "keypress":
            #process screenshot to verify

            print("verifying keypress")
        else:
            result = True

        last_screenshot = screenshot

        return result
    async def process_message(self, msg):
        global workflows, agent_approval, websocket_data
        print(msg)
        print('WebSocket received: ' + msg.data)
        self.websocket_state = "received"
        if "Extension Connected" not in msg.data:
            websocket_data = msg.data
        if "Extension Connected" in msg.data and "undefined" in msg.data:
            print("Please re-install the chrome extension and make sure you are logged in at cheatlayer.com")
        elif "Extension Connected" in msg.data and len(msg.data.split(":")[1]) == 0:
            print("Extension Connected NO KEY")
        elif "Extension Connected" in msg.data and len(msg.data.split(":")[1]) > 2:
            print("Extension Connected")
            
            self.setkey(msg.data.split(":")[1])
            await ws.send_str( "DATAGRAPH:" + json.dumps(self.data_graph))
        if "trigger_agent" in msg.data:

            print("TRIGGER AGENT")
            print(msg.data)
            self.websocket_state = "trigger"
            json_job = json.loads(msg.data.split("trigger_agent:")[1])
            print(json_job)
            data = json_job
            current_desktop = 0  # Assuming `current_desktop` is previously defined
            self.thread_signals.recording_start.emit()
            try: 
                job_folder_name = str(randint(0, 1000000000000))

                for key, node in data["nodes"].items():
                    custom = json.loads(node["custom"]["Data"])
                    if custom["type"] == "Start Node":
                        # Pass job_folder_path to the runNode function:
                        self.runNode(data, key, self.thread_signals, False, True, folder=job_folder_name)
                        break  
            except Exception as e:
                print(e)
        elif "scheduled_desktop_agent" in msg.data:
            print("SCHEDULED AGENT")
            print(msg.data)
            self.websocket_state = "trigger"
            json_output = json.loads(msg.data)["script"]
            for node in json_output:
                print(node)
                last_stored_data = None
                print("this is a project atlas node")
                if "type" in node and node["type"] == "keypress":
                    x = {"type":"keypress_manual", "key":node["text"]}

                    import pyautogui
                    import pydirectinput
                    
                    print("lower keypress regular")
                    print(node["text"])
                    if node["text"] == "enter":
                        pydirectinput.press('enter')
                    else:
                        pydirectinput.write(node["text"])
                #    thread_signals.notifications.emit("Keypress", node["text"], "icon.png", None, "No")
                    #clipboard.copy(node["text"])
                    #pyautogui.hotkey('ctrl', 'v')
                  #  graph.runNode(graph_nodes,nodes[len(nodes)-1].id, thread_signals, False)
                    self.mode = "chat"
                if "type" in node and  node["type"] == "click":
                    print("process click")
                    x = {"type":"Left Mouse Click", "semanticTarget":node["target"], "x":100, "y":100}
                
                    print("added click")
                    try:
                        print("clicking")
                        new_max = graph.semanticSearch(node["target"], 100,100)
                        print(new_max)
                        print("clicking")
                        import pyautogui
                        pyautogui.click(int(new_max['x']), int(new_max['y']))
                   #     graph.QMainWindow.showMinimized()
                        print("run node click")
                   #     graph.runNode(graph_nodes,nodes[len(nodes)-1].id, thread_signals, False)
                        self.mode = "chat"
                    except Exception as e:
                        print(e)
                if "type" in node and node["type"] == "open tab":
                    print("open tab")
                    x = {"type":"Open"}
                
                    import subprocess
                    program = program 

                    print(self.global_variables)
                    import subprocess
               #     self.QMainWindow.showMinimized()
                    subprocess.call("pkill -f chrome/chrome", shell=True)
                    try:
                        
                            if "google-chrome" in program:

                                subprocess.Popen( program.replace("\\n","").split(" ") )
                            else:
                                subprocess.Popen(program.replace("\\n","").split(" "))
                                subprocess.Popen(program.split(" "))
                    except:
                        pass
#   
                   # thread_signals.noti
                if "type" in node and node["type"] == "google_sheets_create":
                    print("google sheets create")
                    try:
                        sheet_url = self.sendMessageRTCAsync('google_sheets_create:')
                        
                        self.graph.global_variables[node["URL"]] = sheet_url 
                        print(sheet_url)
                        last_stored_data = sheet_url
                    except Exception as e:
                        print(e)
                if "type" in node and node["type"] == "general":
                    self.thread_signals.chat.emit( node["description"])
                if "type" in node and node["type"] == "open program":
                    print("open program")
                
            #   
                    print(self.global_variables)
                    import subprocess
               #     self.QMainWindow.showMinimized()
                    subprocess.call("pkill -f chrome/chrome", shell=True)
                    program = node["target"]
                    if "http" in program:
                        program = node["target"].replace("\\n","")
                   # self.thread_signals.notifications.emit("Open", program, "icon.png", None, "No")
                    #graph.runNode(graph_nodes,nodes[len(nodes)-1].id, self.thread_signals, False)
                    

                    program =  program 
                    try:
                        
                            if "google-chrome" in program:

                                subprocess.Popen( program.replace("\\n","").split(" ") )
                            else:
                                subprocess.Popen(program.replace("\\n","").split(" "))
                                subprocess.Popen(program.split(" "))
                    except:
                        pass
                if "type" in node and node["type"] == "google_sheets_add_row":
                    print("google sheets add row")
                    try:
                        sheet_url = self.sendMessageRTCAsync('google_sheets_add_row:' + node["URL"] + ":;:" + json.dumps(self.graph.global_variables[node["data"]]))
                        print(sheet_url)
                        
                        last_stored_data = json.dumps('google_sheets_add_row:' + node["URL"] + ":;:" + json.dumps(self.graph.global_variables[node["data"]]))
                    except Exception as e:
                        print(e)
                if "type" in node and node["type"] == "python":
                    self.thread_signals.notifications.emit("Delay", "Running python code.", "icon.png", None, "No")
                    import site
                    import sys
                    packages = site.getusersitepackages()
                    for package in packages:
                        sys.path.append(package)
                    import subprocess
                    def get_global_sitepackages():
                        try:
                            subprocess.check_output(['python3.10','-m', 'pip','install', 'yt-dlp'], shell=True,universal_newlines=True)
                            import subprocess

# Run the export command in a shell
                            subprocess.check_output(
                                'export LD_LIBRARY_PATH=/usr/local/lib:$LD_LIBRARY_PATH',
                                shell=True,
                                universal_newlines=True
                            )

                            result = subprocess.check_output(['python3', '-c', 'from distutils.sysconfig import get_python_lib; print(get_python_lib())'], universal_newlines=True)
                            return result.strip()
                        except Exception as e:
                            print(f"Error: {e}")
                            print("please install python version 3.10 from the windows store")
                            return None
                    global_sitepackages = get_global_sitepackages()
                    if global_sitepackages:
                        import sys
                        sys.path.append(global_sitepackages)
                    #bundled_dir = sys._MEIPASS
                    #sys.path.append(bundled_dir)
                    sys.path.append(resource_path(""))
                    try:
                        if python_process != None:
                            python_process.terminate()
# Start a proce           s to run the given code.
                        code = node["code"]
                        if "{{" in code:
                            #replace any strings that include {{name}} with agent_memory["name"]
                            print("REPLACING CODE")
                            print(self.graph.global_variables)
                            for key, value in self.graph.global_variables.items():
                                if key in code:
                                    code = code.replace("{{" + key + "}}", "agent_memory['" + key + "']")
                        import moviepy  # Ensure this import is at the top of your script
                        exec(code,{'regenCode': graph.regenCode, 'moviepy': moviepy, 'thread_signals':self.thread_signals, 'notification_manager':graph.notification_manager,'folder': "", 'send_complex_message': graph.send_complex_message, 'agent_memory': self.graph.global_variables ,'store_analytics': graph.store_analytics, 'user_key': graph.user_key, 'imageToVideo': graph.imageToVideo, 'genSyntheticVideo' : graph.genSyntheticVideo, 'genImage': graph.genImage, 'genVoice': graph.genVoice, 'gpt3Prompt': graph.gpt3Prompt, 'last_mouse_y':0,'last_mouse_x': 0, '__name__': '__main__','videoSearch': graph.videoSearch, 'semanticClick': graph.semanticClick, "keypress": graph.keypress , "semanticMove": graph.semanticMove, "click": graph.click, "move":graph.move, "getData": graph.getData, "sendData": graph.sendData} )
                           # try:
#                                 #     self.notification_manager.add_notification("Analyzing Current State", "Analyzing the current state of the screen. Please wait..", "icon.png", lambda: print("Notification closed"))
                         #   self.thread_signals.notifications.emit("Analyzing Current State", "Analyzing the current state of the screen. Please wait..", "icon.png")
                          #  except Exception as e:
                          ##      print(e)
                            #    print("Error in code")
                    except Exception as err:
                        import traceback
                        error_class = err.__class__.__name__
                        print(f"An error occurred: {err}")
                        if self.total_errors < 4:
                           # self.thread_signals.chat.emit(f"A {error_class} error occurred.  Please re-generate the automation to fix it: {err}")
                            self.total_errors += 1
                if "type" in node and node["type"] == "browserScrape":
                    print("magic scraper")
                    print(node["target"])
                    print("gpt-4")
                    browser_elements = []
                    try:
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + node["target"])
                   
                        print("browser elements")
                        print(browser_elements)
                        description = node["description"]
                        x = {"type":"SemanticDescribe", "semanticTarget":"caption"}
                        log = [{"role": "system", "content": "You are a helpfull assistant that takes the list of input elements and returns only the array of elements which matches the user's intent. Return a properly formatted array in a json like this:'{\"data\":[data1,data2,data3...]}'.  Only return the json directly and don't add quotes or a prefix like ```json."}]
                        print("sending to semantic search")
                        log.append({"role": "user", "content": "The intended target is " + description  + ". Generate a valid JSON only with double quotes for strings. Only return the json directly and don't add quotes or a prefix like ```json. The list of input elements is:" + json.dumps(browser_elements)[:50000]})
                        data = {
                            "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
                            "messages": log,
                            "temperature": 0.7,
                            "max_tokens": 5000,
                            "stream": True  # Enable streaming to match original behavior
                        }

                        print(data)

                        # Call local server instead of remote API
                        response = requests.post("http://localhost:8000/v1/chat/completions", 
                                                headers=headers, 
                                                data=json.dumps(data), 
                                                stream=True)

                        import pyautogui
                        if response.status_code == 200:
                            items = []
                            total = ""
                            for chunk in response.iter_lines():
                                try:
                                    if chunk:
                                        # Parse SSE format from llama.cpp server
                                        chunk_str = chunk.decode('utf-8')
                                        if chunk_str.startswith('data: '):
                                            chunk_data = chunk_str[6:]  # Remove 'data: ' prefix
                                            if chunk_data.strip() == '[DONE]':
                                                break
                                            
                                            chunk_json = json.loads(chunk_data)
                                            if 'choices' in chunk_json and len(chunk_json['choices']) > 0:
                                                delta = chunk_json['choices'][0].get('delta', {})
                                                if 'content' in delta:
                                                    content = delta['content']
                                                    total += content
                                                    print(content, end='', flush=True)
                                except Exception as e:
                                    # If streaming fails, fall back to non-streaming
                                    print(f"Streaming error: {e}")
                                    break
                                
                            # If streaming didn't work or total is empty, try non-streaming
                            if not total:
                                data["stream"] = False
                                response = requests.post("http://localhost:8000/v1/chat/completions", 
                                                        headers=headers, 
                                                        data=json.dumps(data))
                                if response.status_code == 200:
                                    result = response.json()
                                    total = result["choices"][0]["message"]["content"]
                                    print(total)
                                else:
                                    print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                                    total = "Error communicating with local LLM server"

                            print()  # New line after streaming output
                            print(total)
                        else:
                            print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                            total = "Error communicating with local LLM server"
                            import re
                            def is_valid_json(json_string):
                                try:
                                    json.loads(json_string)  # Try parsing the JSON
                                    return True  # Return True if parsing is successful
                                except json.JSONDecodeError:
                                    return False  # Return False if an error occurs
                            def fix_json(json_string):
                                # Check and fix missing quotes, brackets, or trailing commas
                                if not is_valid_json(json_string):
                                    try_fixes = [
                                        json_string + ']}',             # Close with nested structures
                                        json_string + '}',              # Close with object
                                        json_string + ']',              # Close with array
                                        json_string.replace('}{', '},{') # Fix missing commas between objects
                                    ]
                                    for fix in try_fixes:
                                        if is_valid_json(fix):
                                            return fix
                                    return None
                                return json_string
                            if is_valid_json(total):
                                print("valid json")
                            else:
                                print("invalid json")
                                total = fix_json(total)
                                print(total)

                            last_stored_data = total
                            self.graph.global_variables[node["data"]] = json.loads(total)["data"]
                        # Use the in-memory bytes to call `blip_caption` directly
                    except Exception as e:
                        print(e)
                if "type" in node and node["type"] == "semanticScrape":
                    print("magic scraper")
                    import pyautogui
                    screenshot = pyautogui.screenshot()
                    buf = io.BytesIO()
                    screenshot.save(buf, format='JPEG')
                    byte_im = buf.getvalue()
                    encoded_jpeg = base64.b64encode(io.BytesIO(byte_im).read()).decode('utf-8')
                    description = node["target"]
                    caption = blip_caption_describe(encoded_jpeg,description, self.graph.user_key, self.graph.user_plan)
                    print(caption)
                    self.graph.global_variables[node["data"]] = caption
                    print(self.graph.global_variables)
                    self.thread_signals.notifications.emit("Magic Scraper ", caption, "icon.png", None, "No")
                    time.sleep(2)
                    last_stored_data = caption
            result = self.verifyStep(node, node["type"] , last_stored_data)
            if result == True:
                print("SUCCESS")
                self.thread_signals.notifications.emit("Step " + str(block_counter), "This step has been completed successfully.", "icon.png", None, "No")
            if result == False:
                #regenerate and run the step again 
                print("FAILED")
                prompt = "here is are the list of previously generated steps: " + json.dumps(json_output) + " This step has failed: " + json.dumps(node) + " Please regenerate the automation to fix this step. Common mistakes includes not specifying the correct variables for passing data between nodes or not adding a to wait for websites to load. Common solutions include adjusting the variable names in the 'data' parameter and adding delays."
                self.thread_signals.notifications.emit("Regenerating Step " + str(block_counter), "Regenerating Step " + str(block_counter), "icon.png", None, "No")
        elif "type" in msg.data and "{" in msg.data:
            print("TRIGGER")
            print(msg.data)
            

            self.websocket_state = "trigger"
            json_job = json.loads(msg.data)
            print(json_job)
            self.trigger(json_job)
            if "yes" in msg.data.lower():
                agent_approval = 1
            elif "no" in msg.data.lower():
                agent_approval = -1
        if "start" in msg.data:
            obj = json.loads(msg.data)
            if "start" in obj and obj["start"] == "desktop":
                self.trigger(obj)
        if "screenshots" in msg.data:
            print("Screenshots")
            websocket_data = msg.data
            print(msg.data)
        elif "Extension Connected" not in msg.data:
            websocket_data = msg.data
    def start(self):
        app = web.Application()
        app.router.add_route('GET', '/websocket', self.handle_websocket, name='websocket')

        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)

        runner = web.AppRunner(app)
        print("Starting server")
        loop.run_until_complete(runner.setup())
        site = web.TCPSite(runner, 'localhost', 8080)
        loop.run_until_complete(site.start())

        # Start the reconnection task
        loop.create_task(self.reconnect_websocket(app))

        loop.run_forever()

# Run the server in a separate thread
websocket_data = ""
async def sendMessageRTC(message):
    global websocket_state, websocket_data
    websocket_state = "sent"
    print("Sending Message")
    try:
        wait_counter = 0
        await ws.send_str( "test")

        await ws.send_str( message)
        while websocket_state == "sent":
            print("Waiting for response")
            wait_counter += 1
            print(wait_counter)
            print(websocket_data)
            if wait_counter > 50 or len(websocket_data) > 0:
                print("Timeout")
                break
            print(websocket_state)
            await asyncio.sleep(0.1)
        print("Received Message")
        print(websocket_data)
        return websocket_data
    except Exception as e:

        print(e)
        return ""
# Step 3: Assign your custom function to sys.stdout.write
backup = sys.stdout.write
sys.stdout.write = my_custom_write


rf = None#
project = None#rf.workspace().project("ui-components-detection")
UIModel =None# project.version(2).model
#import QSizePolicy
sys.stdout.write = backup
last_mouse_x = -1
last_mouse_y = -1
import concurrent.futures

from PySide2.QtWidgets import QSizePolicy
import threading
import os
import argparse, os, re
import numpy as np
from random import randint
from omegaconf import OmegaConf
from PIL import Image
from itertools import islice
import time
#import torch
from contextlib import contextmanager, nullcontext
#pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
import tkinter as tk
import gc
import replicate
import subprocess
#from detecto import core, utils
from NodeGraphQt.constants import (NODE_PROP_QLABEL,
                                       NODE_PROP_QLINEEDIT,
                                       NODE_PROP_QTEXTEDIT,
                                       NODE_PROP_QCOMBO,
                                       NODE_PROP_QSPINBOX,
                                       NODE_PROP_COLORPICKER,
                                       NODE_PROP_SLIDER)
# import example nodes from the "example_nodes" package
import imutils
import requests
import pyautogui
import json
from PIL import ImageTk, Image
from tkinter import ttk
from threading import Thread

from queue import Queue
record_tasks = Queue()

import mss



import os.path
import platform
import re
import math
from Qt import QtCore, QtWidgets, QtGui

import sys
import os
import gc
from IPython import display
from IPython.display import Image as ipyimg
from PIL import Image
from numpy import asarray
import time
import numpy as np
from datetime import datetime
import warnings
warnings.filterwarnings("ignore", category=UserWarning)

from NodeGraphQt.base.commands import (NodeAddedCmd,
                                       NodeRemovedCmd,
                                       NodeMovedCmd,
                                       PortConnectedCmd)
from NodeGraphQt.base.factory import NodeFactory
from NodeGraphQt.base.menu import NodeGraphMenu, NodesMenu
from NodeGraphQt.base.model import NodeGraphModel
from NodeGraphQt.base.node import NodeObject
from NodeGraphQt.base.port import Port
from NodeGraphQt.constants import (
    NODE_LAYOUT_DIRECTION, NODE_LAYOUT_HORIZONTAL, NODE_LAYOUT_VERTICAL,
    PIPE_LAYOUT_CURVED, PIPE_LAYOUT_STRAIGHT, PIPE_LAYOUT_ANGLE,
    URI_SCHEME, URN_SCHEME,
    IN_PORT, OUT_PORT,
    VIEWER_GRID_LINES,VIEWER_GRID_DOTS
)
from NodeGraphQt.nodes.backdrop_node import BackdropNode
from NodeGraphQt.nodes.base_node import BaseNode
from NodeGraphQt.nodes.group_node import GroupNode
from NodeGraphQt.nodes.port_node import PortInputNode, PortOutputNode
from NodeGraphQt.widgets.node_graph import NodeGraphWidget, SubGraphWidget
from NodeGraphQt.widgets.viewer import NodeViewer
from NodeGraphQt.widgets.viewer_nav import NodeNavigationWidget
from Qt import QtCore, QtWidgets, QtGui
import base64

from PySide2.QtWidgets import (QWidget, QPushButton, QHBoxLayout, QListWidgetItem, QToolBar, QScrollArea, QApplication, QTextEdit, QLabel, QGraphicsView,
QGridLayout, QMainWindow, QAction, QMenu, QVBoxLayout, QMenuBar, QFileDialog, QInputDialog, QMessageBox)
from PySide2.QtGui import QKeySequence
from PySide2.QtWidgets import QApplication, QWidget, QLabel, QVBoxLayout, QPushButton
from PySide2.QtGui import QPixmap, QIcon
from PySide2.QtCore import Qt, QSize, QTimer

from PySide2.QtWidgets import QApplication, QWidget, QVBoxLayout, QLabel, QPushButton, QHBoxLayout
from PySide2.QtCore import Qt, QTimer
from PySide2.QtCore import Qt, Signal, QObject
from PySide2.QtWidgets import QWidget, QLabel, QVBoxLayout, QApplication
from PySide2.QtCore import Qt, Signal
from PySide2.QtGui import QPixmap

class NotificationWidget(QWidget):
    closed = Signal(QWidget)

    def __init__(self, title, message, image_path, automation_callback=None, prompt="the screen", *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.setFixedSize(300, 160)  # Set the fixed size for the widget
        self.setWindowFlags(Qt.Tool | Qt.WindowStaysOnTopHint | Qt.FramelessWindowHint)
        self.setAttribute(Qt.WA_TranslucentBackground)

        self.automation_callback = automation_callback
        self.prompt = prompt
        cut_message = "    " + message[:150] + (message[150:] and '...')
        # Text Label Setup
        self.text_label = QLabel(cut_message, self)
        self.text_label.setAlignment(Qt.AlignCenter)
        self.text_label.setGeometry(0, 0, 300, 150)
        #add word-wrap and elipses after 3 lines max
        self.text_label.setStyleSheet("""
        background-color: rgba(255, 255, 255, 255);
        border-radius: 16px;
        padding: 10px;
                                      
        color: black;
        font-size: 16px;
        """)
                #add word-wrap and elipses after 3 lines max

        
        self.text_label.setWordWrap(True)

        # Image Icon Setup
        self.image_label = QLabel(self)
        self.image_label.setStyleSheet("background: transparent;")
        pixmap = QPixmap(image_path)
        if not pixmap.isNull():
            self.image_label.setPixmap(pixmap.scaled(25, 25, Qt.KeepAspectRatio, Qt.SmoothTransformation))
        self.image_label.setGeometry(10, 10, 25, 25)

        # Play Button Setup
        if self.automation_callback is not None:
            
            self.play_button = QPushButton(self)
            self.play_button.setText("►")
            self.play_button.setGeometry(160, 70, 30, 30)
            self.play_button.setStyleSheet("QPushButton { font-size: 14px; }")
            self.play_button.clicked.connect(self.execute_automation_callback)

        # Modify cursor
        self.setCursor(Qt.PointingHandCursor)
        
        # Initiate and set up QTimer for auto-close functionality
        self.close_timer = QTimer(self)
        self.close_timer.timeout.connect(self.auto_close)
        self.close_timer.setSingleShot(True)  # Ensure the timer runs only once
        self.close_timer.start(5000)  # 5000 milliseconds = 5 seconds

    def execute_automation_callback(self):
        if self.automation_callback:
            thread = threading.Thread(target=self.automation_callback)
            thread.start()

    def mousePressEvent(self, event):
        self.hide()
        self.close()

    def closeEvent(self, event):
        self.closed.emit(self)
        super().closeEvent(event)

    def auto_close(self):
        """Automatically close the notification widget."""
        if not self.isHidden():  # Check if the widget is still visible
            self.hide()
            self.close()
class NotificationManager(QObject):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.notifications = []
        self.are_notifications_visible = True
        self.last_notification = None
    def show(self):
        self.show_notifications()
    def hide(self):
        self.hide_notifications()
    def add_notification(self, title, message, image_path, automation_callback=None, prompt = " the screen"):

        notification = NotificationWidget(title, message, image_path, automation_callback, prompt)
        notification.closed.connect(self.remove_notification)
        notification.show()
        self.notifications.insert(0, notification)
        if len(self.notifications) > 1:
            to_remove = self.notifications.pop()
            to_remove.close()
        self.last_notification = message

        self.organize_notifications()

    def organize_notifications(self):
        desktop_rect = QApplication.desktop().availableGeometry()
        screen_right = desktop_rect.right()
        screen_top = desktop_rect.top()
        for index, notification in enumerate(self.notifications):
            notification.move(screen_right - notification.width(), screen_top + index * notification.height())
#            notification.layout.update()

    def remove_notification(self, notification_widget):
        notification_widget.close()
        self.organize_notifications()

    def hide_notifications(self):
        for notification in self.notifications:
            notification.hide()
        self.are_notifications_visible = False

    def show_notifications(self):
        for notification in self.notifications:
            notification.show()
        self.are_notifications_visible = True

    def toggle_notifications_visibility(self):
        if self.are_notifications_visible:
            self.hide_notifications()
        else:
            self.show_notifications()

class ReceiverMessageItem(QListWidgetItem):
    def __init__(self, text):
        super().__init__()
        self.text = text

        # Create a custom widget for the list item
        self.widget = QWidget()

        self.layout = QHBoxLayout(self.widget)

        self.label = QLabel(self.text)
        self.label.setWordWrap(True)
        self.label.setFixedHeight(300)
        self.label.setWordWrap(True)
        #float the label to the left
        self.label.setAlignment(QtCore.Qt.AlignLeft)
        self.label.setStyleSheet("""F
            background: #FBFBFB;
            border-radius: 16px;
            padding: 14px 18px;
            border: 1px black solid;

            height:1000px;
            color: black;
            font-size: 18px;
            line-height: 22px;
        """)
        self.layout.addWidget(QLabel(), 1)  # Spacer
        self.layout.addWidget(self.label, 0, QtCore.Qt.AlignLeft)

        self.widget.setLayout(self.layout)
        self.setSizeHint(self.widget.sizeHint())
from contextlib import redirect_stdout
import io

full = io.StringIO()



def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """

    if getattr(sys, 'frozen', False) and hasattr(sys, '_MEIPASS'):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.abspath(".")

    return os.path.join(base_path, relative_path)
import os
import platform

#fimport torch
#from torchvision import transforms
#from torchvisio


def pil2pixmap(image):
    bytes_img = io.BytesIO()
    image.save(bytes_img, format='PNG')
    from PySide2.QtGui import QPixmap, QImage

    qimg = QImage()
    qimg.loadFromData(bytes_img.getvalue())

    return QPixmap.fromImage(qimg)
def load_demo_image(image_size,device, input):
    img_url = 'https://storage.googleapis.com/sfr-vision-language-research/BLIP/demo.jpg'
   # raw_image = Image.open(requests.get(img_url, stream=True).raw).convert('RGB')

    #    image = cv2.cvtColor(np.array(img), cv2.COLOR_BGR2RGB)

    w,h = input.size
   # display()
#    raw_image.resize((w//5,h//5))
    transform = transforms.Compose([
        transforms.Resize((image_size,image_size),interpolation=InterpolationMode.BICUBIC),
        transforms.ToTensor(),
        transforms.Normalize((0.48145466, 0.4578275, 0.40821073), (0.26862954, 0.26130258, 0.27577711))
        ])
    image = transform(input).unsqueeze(0).to(device)
    return image
from io import StringIO
from contextlib import redirect_stdout
import youtube_dl

#from models.blip_itm import blip_itm

image_size = 384
#image = pyautogui.screenshot(region=(0,0,1920,1080))
#raw_image = image.crop((0,0, 400,500))
#image = load_demo_image(image_size=image_size, device=device, input = raw_image)

model_url = 'https://storage.googleapis.com/sfr-vision-language-research/BLIP/models/model_base_capfilt_large.pth'

import requests

API_URL = "https://api-inference.huggingface.co/models/Salesforce/blip-image-captioning-large"

import logging

error_total = 0

def blip_caption_check(data, caption,user_key, user_plan):
    total = ""
    global error_total
    print(user_key)
    print(user_plan)
    try:
        log = [{"role": "system", "content": "You take the current screenshot of the screen, and output either 'yes' or 'no' if the image exactly matches the description."}]

        #log = [{"role": "system", "content": "You return the number of the element which matches the target description from the image. The provided image will have boxes with numbers inside them surrounding potential targets that may match the description with a red number to the left of the top left corner of the corresponding element. If none match the description exactly, return exactly -1. Be as precise as possible when determining if a matching element exists or not. Make sure the color, text and any other information in the description exactly matches the element, and consider all the elements in the list. The numbers will all be inside the element box and in the top left of the box."}]
     #           chat_log.append({"role": "user", "content": text})
        headers = {
            "Content-Type": "application/json",
        }
        def encode_image(image_path):
          with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
        mode_gpt = "32k"
        #if text includes the keyword 'screenshot', take a screenshot with pyauto-gui and reformat chat_log to use the following format to support images and text. 
        #"messages": [ { "role": "user", "content": [ { "type": "text", "text": "What’s in this image?" }, { "type": "image_url", "image_url": { "url": "https://upload.wikimedia.org/wikipedia/commons/thumb/d/dd/Gfp-wisconsin-madison-the-nature-boardwalk.jpg/2560px-Gfp-wisconsin-madison-the-nature-boardwalk.jpg" } } ] } ],
        #base64_image = encode_image("screenshot.png")
        mode_gpt = "website2"
        ##print("SCRNEENSHOT MODE"
        log.append({"role": "user", "content": [{"type": "text", "text":  caption}, {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{data}"}}]})
        ##print("screenshot")
        data = {
            "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
            "messages": log,
            "temperature": 0.7,
            "max_tokens": 5000,
            "stream": True  # Enable streaming to match original behavior
        }

        print(data)

        # Call local server instead of remote API
        response = requests.post("http://localhost:8000/v1/chat/completions", 
                                headers=headers, 
                                data=json.dumps(data), 
                                stream=True)

        import pyautogui
        if response.status_code == 200:
            items = []
            total = ""
            for chunk in response.iter_lines():
                try:
                    if chunk:
                        # Parse SSE format from llama.cpp server
                        chunk_str = chunk.decode('utf-8')
                        if chunk_str.startswith('data: '):
                            chunk_data = chunk_str[6:]  # Remove 'data: ' prefix
                            if chunk_data.strip() == '[DONE]':
                                break
                            
                            chunk_json = json.loads(chunk_data)
                            if 'choices' in chunk_json and len(chunk_json['choices']) > 0:
                                delta = chunk_json['choices'][0].get('delta', {})
                                if 'content' in delta:
                                    content = delta['content']
                                    total += content
                                    print(content, end='', flush=True)
                except Exception as e:
                    # If streaming fails, fall back to non-streaming
                    print(f"Streaming error: {e}")
                    break
                
            # If streaming didn't work or total is empty, try non-streaming
            if not total:
                data["stream"] = False
                response = requests.post("http://localhost:8000/v1/chat/completions", 
                                        headers=headers, 
                                        data=json.dumps(data))
                if response.status_code == 200:
                    result = response.json()
                    total = result["choices"][0]["message"]["content"]
                    print(total)
                else:
                    print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                    total = "Error communicating with local LLM server"

            print()  # New line after streaming output
            print(total)
        else:
            print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
            total = "Error communicating with local LLM server"
            

        
    except Exception as error:

        #logging.error(f"Error occurred while calling replicate" + error)
        if error_total < 3:
            blip_caption_check(data, caption,user_key, user_plan)
            error_total += 1

        response = "0"
    print(total)
    return total

# Convert OpenAI vision format to llama.cpp format
def convert_message_for_llama(message):
    """Convert OpenAI vision API format to llama.cpp format"""
    if isinstance(message.get('content'), list):
        # Extract text and image from the content array
        text_content = ""
        image_data = None
        for item in message['content']:
            if item['type'] == 'text':
                text_content += item['text']
            elif item['type'] == 'image_url':
                # Extract base64 data from data URL
                image_url = item['image_url']['url']
                if image_url.startswith('data:image/'):
                    # Remove the data:image/jpeg;base64, prefix
                    image_data = image_url.split(',', 1)[1]
        return {
            "role": message['role'],
            "content": text_content,
            "images": [image_data] if image_data else []
        }
    else:
        # Regular text message
        return message
    
llama_server_process = None


def setup_local_llama_server(self):
    """
    Download the Gemma-3 model if needed and start a local llama.cpp server with vision capabilities
    """
    print("Setting up local llama.cpp server with Gemma-3 4B-IT model...")
    global llama_server_process
    # Clone llama.cpp repository
    if not os.path.exists("llama.cpp"):
        subprocess.run(["git", "clone", "https://github.com/ggerganov/llama.cpp.git"], check=True)
    # Build llama.cpp with vision support
    if not os.path.exists("llama.cpp/server"):
        os.chdir("llama.cpp")
        subprocess.run(["make", "server", "LLAMA_BLAS=ON", "LLAMA_BLAS_VENDOR=OpenBLAS", f"-j{os.cpu_count()}"], check=True)
        os.chdir("..")
    # Create models directory if it doesn't exist
    os.makedirs("models/gemma3", exist_ok=True)
    # Check if model and projector exist, download if needed
    model_path = "models/gemma3/gemma-3-4b-it-q4_k_xl.gguf"
    if not os.path.exists(model_path):
        print("Downloading Gemma-3 4B-IT model (this may take a while)...")
        subprocess.run([
            "wget", "-O", model_path,
            "https://huggingface.co/bartowski/google_gemma-3-4b-it-GGUF/resolve/main/gemma-3-4b-it-Q4_K_XL.gguf"
        ], check=True)
    projector_path = "models/gemma3/mmproj-gemma-3-4b-it-f16.gguf"
    if not os.path.exists(projector_path):
        print("Downloading multimodal projector (this may take a while)...")
        subprocess.run([
            "wget", "-O", projector_path,
            "https://huggingface.co/bartowski/google_gemma-3-4b-it-GGUF/resolve/main/mmproj-gemma-3-4b-it-f16.gguf"
        ], check=True)
    # Check if server is already running
    try:
        response = requests.get("http://localhost:8000/v1/models")
        if response.status_code == 200:
            print("Local llama.cpp server is already running")
            return
    except:
        pass  # Server not running, we'll start it
    
    # Start the server
    print("Starting llama.cpp server with Gemma-3 4B-IT model...")
    # Determine number of CPU threads
    cpu_threads = os.cpu_count()
    if cpu_threads > 4:
        cpu_threads = cpu_threads - 2  # Leave some CPU for other tasks
    llama_server_process = subprocess.Popen([
        "llama.cpp/server",
        "-m", model_path,
        "--mmproj", projector_path,
        "--host", "0.0.0.0",
        "--port", "8000",
        "-t", str(cpu_threads),
        "-c", "4096"
    ])
    # Give server time to start up
    time.sleep(5)
    print("Local llama.cpp server is running with Gemma-3 4B-IT!")
def blip_caption_describe(data, caption,user_key, user_plan):
    total = ""
    global error_total, llama_server_process
    print(user_key)
    print(user_plan)
    
    if llama_server_process is None:
        setup_local_llama_server()
    
    try:
        log = [{"role": "system", "content": "You take the current screenshot of the screen, and return a detailed description of exactly the part the user wants to target. For example, if the user asks to desribe an email, return only the text you can see from the email. If the user asks to describe a social post, return only the text from that social post. Do not describe your output and only output exactly what the user wants. Do prefix the output text with any description, and do not mention what the target is. ."}]

        #log = [{"role": "system", "content": "You return the number of the element which matches the target description from the image. The provided image will have boxes with numbers inside them surrounding potential targets that may match the description with a red number to the left of the top left corner of the corresponding element. If none match the description exactly, return exactly -1. Be as precise as possible when determining if a matching element exists or not. Make sure the color, text and any other information in the description exactly matches the element, and consider all the elements in the list. The numbers will all be inside the element box and in the top left of the box."}]
     #           chat_log.append({"role": "user", "content": text})
        headers = {
            "Content-Type": "application/json",
        }
        def encode_image(image_path):
          with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
        mode_gpt = "32k"
        #if text includes the keyword 'screenshot', take a screenshot with pyauto-gui and reformat chat_log to use the following format to support images and text. 
        #"messages": [ { "role": "user", "content": [ { "type": "text", "text": "What’s in this image?" }, { "type": "image_url", "image_url": { "url": "https://upload.wikimedia.org/wikipedia/commons/thumb/d/dd/Gfp-wisconsin-madison-the-nature-boardwalk.jpg/2560px-Gfp-wisconsin-madison-the-nature-boardwalk.jpg" } } ] } ],
        #base64_image = encode_image("screenshot.png")
        mode_gpt = "website2"
        ##print("SCRNEENSHOT MODE"
        log.append({"role": "user", "content": [{"type": "text", "text":  caption}, {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{data}"}}]})
        ##print("screenshot")
        
        # Convert OpenAI vision format to llama.cpp format
        def convert_message_for_llama(message):
            """Convert OpenAI vision API format to llama.cpp format"""
            if isinstance(message.get('content'), list):
                # Extract text and image from the content array
                text_content = ""
                image_data = None
                for item in message['content']:
                    if item['type'] == 'text':
                        text_content += item['text']
                    elif item['type'] == 'image_url':
                        # Extract base64 data from data URL
                        image_url = item['image_url']['url']
                        if image_url.startswith('data:image/'):
                            # Remove the data:image/jpeg;base64, prefix
                            image_data = image_url.split(',', 1)[1]
                return {
                    "role": message['role'],
                    "content": text_content,
                    "images": [image_data] if image_data else []
                }
            else:
                # Regular text message
                return message
        # Convert all messages to llama.cpp format
        converted_log = [convert_message_for_llama(msg) for msg in log]
        data = {
            "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
            "messages": converted_log,
            "temperature": 0.7,
            "max_tokens": 5000,
            "stream": True  # Enable streaming to match original behavior
        }

        print(data)

        # Call local server instead of remote API
        response = requests.post("http://localhost:8000/v1/chat/completions", 
                                headers=headers, 
                                data=json.dumps(data), 
                                stream=True)

        import pyautogui
        if response.status_code == 200:
            items = []
            total = ""
            for chunk in response.iter_lines():
                try:
                    if chunk:
                        # Parse SSE format from llama.cpp server
                        chunk_str = chunk.decode('utf-8')
                        if chunk_str.startswith('data: '):
                            chunk_data = chunk_str[6:]  # Remove 'data: ' prefix
                            if chunk_data.strip() == '[DONE]':
                                break
                            
                            chunk_json = json.loads(chunk_data)
                            if 'choices' in chunk_json and len(chunk_json['choices']) > 0:
                                delta = chunk_json['choices'][0].get('delta', {})
                                if 'content' in delta:
                                    content = delta['content']
                                    total += content
                                    print(content, end='', flush=True)
                except Exception as e:
                    # If streaming fails, fall back to non-streaming
                    print(f"Streaming error: {e}")
                    break
                
            # If streaming didn't work or total is empty, try non-streaming
            if not total:
                data["stream"] = False
                response = requests.post("http://localhost:8000/v1/chat/completions", 
                                        headers=headers, 
                                        data=json.dumps(data))
                if response.status_code == 200:
                    result = response.json()
                    total = result["choices"][0]["message"]["content"]
                    print(total)
                else:
                    print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                    total = "Error communicating with local LLM server"

            print()  # New line after streaming output
            print(total)
        else:
            print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
            total = "Error communicating with local LLM server"

        
    except Exception as error:
        return "error"
        response = "0"
    print(total)
    return total
def blip_caption(data, caption):
    try:
        response = call_replicate(data, caption)
    except Exception as error:
#        logging.error(f"Error occurred while calling replicate" + error)
        print("error")
        print(error)
        response = "0"
    print("output")
    print(response)
    return response

def blip_caption_new(data, caption,user_key, user_plan , data2):
    total = ""
    print(user_key)
    print(user_plan)
    
    if llama_server_process is None:
        setup_local_llama_server()
    
    try:
        
        log = [{"role": "system", "content": "You are provided two images of a screen with UI elements. In the first image, determine the locations of all the possible target UI elements that might possibly match the description given by the user. In the 2nd image, there are red indices inside each UI element with a red bounding box around the element. The red indices are in the middle or right hand side of each cooresponding element. Select all possible similar elements to this description and return a list of all possible elements that might possibly match the description. Use these indices to communicate which elements on the screenshot matches the description by returning an array of the indices that closely match the description.  If the description asks to exclude anything, make sure to never select that index. Output an array of elements like [2,3,42] that cloest match the description. If there are multiple similar elements, include them all. You return the numbers of the UI elements which matches the description exactly. Choose the numbers that is closest to the elements in pixels that you want. Use these numbers to refer to the element and return only the numbers if the target element exists in the image. If the user provides a location, pay attention to this when matching the description and eliminate any elements which don't appear in this location. If the description does not exist in the image at all, always only return the number -1. Return only a single array of numbers and do not explain or return text in the output. Ignore the notifications in the top right hand corner of the screen, since they will repeat the current step and will often include the words of the current step. Try to return at least 2 indices if possible. Never select a notification in the top right hand corner of the screen."}]

        #log = [{"role": "system", "content": "You return the number of the element which matches the target description from the image. The provided image will have boxes with numbers inside them surrounding potential targets that may match the description with a red number to the left of the top left corner of the corresponding element. If none match the description exactly, return exactly -1. Be as precise as possible when determining if a matching element exists or not. Make sure the color, text and any other information in the description exactly matches the element, and consider all the elements in the list. The numbers will all be inside the element box and in the top left of the box."}]
     #           chat_log.append({"role": "user", "content": text})

        def encode_image(image_path):
          with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
        mode_gpt = "32k"
        #if text includes the keyword 'screenshot', take a screenshot with pyauto-gui and reformat chat_log to use the following format to support images and text. 
        #"messages": [ { "role": "user", "content": [ { "type": "text", "text": "What’s in this image?" }, { "type": "image_url", "image_url": { "url": "https://upload.wikimedia.org/wikipedia/commons/thumb/d/dd/Gfp-wisconsin-madison-the-nature-boardwalk.jpg/2560px-Gfp-wisconsin-madison-the-nature-boardwalk.jpg" } } ] } ],
        #base64_image = encode_image("screenshot.png")
        mode_gpt = "click"
        ##print("SCRNEENSHOT MODE"
        log.append({"role": "user", "content": [{"type": "text", "text":  "First, find all the target elements in this image that might posissbly match the description. Ignore the notifications in the top right hand corner of the screen, since they will repeat the current step and will often include the words of the current step. Never select a notification in the top right hand corner of the screen.  If there are multiple similar elements, generate the full list of indices in the array, for example [21,412,1].  Here is the target description:" + caption}, {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{data2}"}}]})

        log.append({"role": "user", "content": [{"type": "text", "text":  "Next, use this image with bounding boxes and indices to return the indices of the elements which might match the provided description. The red indices are in the middle or right hand side of each cooresponding element. Select all possible similar elements to this description from the previous image and this description:" + caption  + ".  Output an array of elements like [2,3,42] that cloest match the description. If the description asks to exclude anything, make sure to never select that index.  Ignore the size of the numbers and the size has no importance. Instead, find the indices of the elements, text or images, which closely match this description best.  Ignore the notifications in the top right hand corner of the screen, since they will repeat the current step and will often include the words of the current step. Output an array of elements like [2,3,42] that closest match the description.  If the description asks to exclude anything, make sure to never select that index. If there are multiple similar elements, include them all.  Never select a notification in the top right hand corner of the screen. The red numbers are inside the cooresponding elements red bounding box. Return only a single array of indices and do not explain or return text in the output. Use the clean image to determine which elements to target based on my description, and the image with red numbers to determine the number of the element. The numbers will always be INSIDE the red bounding box of the coorsponding element, so find the closest bounding box first, then determine the number inside that bounding box.  Turn only either -1 if you can't find the elements or the number of the elements that matches exactly this description exactly. Choose the targets that might possibly matches this description exactly:" + caption},  {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{data}"}}]})
                ##print("screenshot")
                
        converted_log = [convert_message_for_llama(msg) for msg in log]
        data = {
            "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
            "messages": converted_log,
            "temperature": 0.7,
            "max_tokens": 5000,
            "stream": True  # Enable streaming to match original behavior
        }
        
        print(data)
        
        # Call local server instead of remote API
        response = requests.post("http://localhost:8000/v1/chat/completions", 
                                headers=headers, 
                                data=json.dumps(data), 
                                stream=True)
        
        import pyautogui
        if response.status_code == 200:
            items = []
            total = ""
            for chunk in response.iter_lines():
                try:
                    if chunk:
                        # Parse SSE format from llama.cpp server
                        chunk_str = chunk.decode('utf-8')
                        if chunk_str.startswith('data: '):
                            chunk_data = chunk_str[6:]  # Remove 'data: ' prefix
                            if chunk_data.strip() == '[DONE]':
                                break
                            
                            chunk_json = json.loads(chunk_data)
                            if 'choices' in chunk_json and len(chunk_json['choices']) > 0:
                                delta = chunk_json['choices'][0].get('delta', {})
                                if 'content' in delta:
                                    content = delta['content']
                                    total += content
                                    print(content, end='', flush=True)
                except Exception as e:
                    # If streaming fails, fall back to non-streaming
                    print(f"Streaming error: {e}")
                    break
                
            # If streaming didn't work or total is empty, try non-streaming
            if not total:
                data["stream"] = False
                response = requests.post("http://localhost:8000/v1/chat/completions", 
                                        headers=headers, 
                                        data=json.dumps(data))
                if response.status_code == 200:
                    result = response.json()
                    total = result["choices"][0]["message"]["content"]
                    print(total)
                else:
                    print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                    total = "Error communicating with local LLM server"
            
            print()  # New line after streaming output
            print(total)
        else:
            print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
            total = "Error communicating with local LLM server"
        
    except Exception as error:
        blip_caption_new(data, caption,user_key, user_plan, data2)
        response = "0"
    print(total)
    if len(total) > 1:
        return json.loads(total)
    else:
        return [0]
def blip_caption2(data, caption = ""):
    try:
        response = call_replicate2(data, caption)
    except Exception as error:
        logging.error(f"Error occurred while calling replicate" + error)
        response = "0"
    print(response)
    return response

def blip_caption2_2(data, caption = ""):
    try:
        response = call_replicate2_2(data, caption)
    except Exception as error:
        print(error)
        response = "0"
    print(response)
    return response



# Load a model
def UImodel(image_path):
    import json
    import requests
    # Run inference on an image
    url = "https://api.ultralytics.com/v1/predict/kVhczpBn9wqjExiUOYBU"
    headers = {"x-api-key": "ULTRALYTICSKEY"}
    data = {"size": 640, "confidence": 0.01, "iou": 0.2}
    results_out = []
    with open(image_path, "rb") as f:
        response = requests.post(url, headers=headers, data=data, files={"image": f}, verify=False)
        print(response.json())
        response_list = response.json()
        print(response_list["images"])
        print(response_list["images"][0])
        for box in response_list["images"][0]["results"]:
            results_out.append(box)
    return results_out
def call_replicate3(data, resolution = "maintain_aspect_ratio", frames = "25_frames_with_svd_xt", motion_bucket_id = 124):
    print("running replciate call")
    print(data)
    print(resolution)
    print(frames)
    return replicate.run(
        "stability-ai/stable-video-diffusion:3f0457e4619daac51203dedb472816fd4af51f3149fa7a9e0b5ffcf1b8172438",
        input={"input_image":data, "cond_aug": 0.05, "decoding_t": 7, "video_length": "25_frames_with_svd_xt", "sizing_strategy": "maintain_aspect_ratio",resolution: motion_bucket_id, "frames_per_second":12}
    )
def call_replicate2_2(data, caption):
    print("caption 2")
    caption_text = caption

    return replicate.run(
    "salesforce/blip:2e1dddc8621f72155f24cf2e0adbde548458d3cab9f00c0139eea840d0ac4746",
        input={"image":data, "caption":caption_text,         "task": "image_text_matching"}
    )
def call_replicate2(data, caption):
    print("caption")
    caption_text = " What does that text in this image say?"

    return replicate.run(
        "andreasjansson/blip-2:9109553e37d266369f2750e407ab95649c63eb8e13f13b1f3983ff0feb2f9ef7",
        input={"image":data, "question":caption_text, "caption":False}
    )
def call_replicate(data, caption):
    print("caption")
    return replicate.run(
        "salesforce/blip:2e1dddc8621f72155f24cf2e0adbde548458d3cab9f00c0139eea840d0ac4746",
        input={"image":data, "caption":caption, "task":"image_text_matching"}
    )

USER_ME = 0
USER_THEM = 1
def get_path(filename):
    name = os.path.splitext(filename)[0]
    ext = os.path.splitext(filename)[1]

    return os.path.realpath(filename)

class ChatApp(QWidget):
    def __init__(self):
        super().__init__()

        # Set up main window properties
        self.setGeometry(200, 600, 600, 600)
        self.setWindowTitle('Chat Interface')
        self.setWindowFlags(Qt.FramelessWindowHint)

        # Layouts and Widgets
        self.layout = QVBoxLayout(self)

        self.messages = QTextEdit(self)
        self.messages.setReadOnly(True)
        self.layout.addWidget(self.messages)

        self.input_layout = QVBoxLayout()

        self.input_text = QLineEdit(self)
        self.input_text.setPlaceholderText("Ask Project Atlas anything...")
        self.input_text.textChanged.connect(self.search_function)
        self.input_layout.addWidget(self.input_text)

        self.search_results = QListWidget(self)
        self.search_results.hide()  # hidden by default
        self.input_layout.addWidget(self.search_results)

        self.layout.addLayout(self.input_layout)

        self.input_text.returnPressed.connect(self.send_message)

        self.dragPos = None

    def search_function(self):
        # Example search functionality (simplified)
        # In a real-world scenario, you might want to integrate more complex logic.
        query = self.input_text.text().lower()
        if query:
            # Example: filter only those messages that contain the query.
            # You can replace this logic with more sophisticated search.
            filtered = [msg for msg in self.messages.toPlainText().split('\n') if query in msg.lower()]
            self.search_results.clear()
            self.search_results.addItems(filtered)
            self.search_results.show()
        else:
            self.search_results.hide()

    def send_message(self):
        msg = self.input_text.text()
        self.messages.append(msg)
        self.input_text.clear()
        self.search_results.hide()

    def mousePressEvent(self, event):
        self.dragPos = event.globalPos()

    def mouseMoveEvent(self, event):
        if event.buttons() == Qt.LeftButton:
            self.move(self.pos() + event.globalPos() - self.dragPos)
            self.dragPos = event.globalPos()
            event.accept()
from cron_descriptor import Options, CasingTypeEnum, DescriptionTypeEnum, ExpressionDescriptor

class JobWidget(QWidget):
    def __init__(self, job_data, schedule, graph, updateSchedule):
        super(JobWidget, self).__init__()
        self.job_data = job_data
        self.is_enabled = True  # Track the state of the switch
        self.graph = graph       
        self.schedule = schedule
        self.updateSchedule = updateSchedule

        # Create a vertical layout
        main_layout = QVBoxLayout()

        # Create a horizontal layout for the first line
        first_line_layout = QHBoxLayout()

        # Switch Button
        self.switch_button = QPushButton()
        self.switch_button.setCheckable(True)
        self.switch_button.setChecked(self.is_enabled)
        self.switch_button.setFixedSize(QSize(20, 20))
        self.switch_button.clicked.connect(self.toggle_switch)
        self.update_switch_color()
        first_line_layout.addWidget(self.switch_button)

        file_name_label = QLabel(job_data['cheat'].split("/")[-1].split(".")[0].replace("_", " "))
        first_line_layout.addWidget(file_name_label)

        main_layout.addLayout(first_line_layout)

        job = job_data
        cron = job["minute"] + " " +job["hour"] + " " + job["day"] + " " + job["month"] + " " + str(job["weekday"])
        descriptor = ExpressionDescriptor(
            expression = cron,
            casing_type = CasingTypeEnum.Sentence, 

            throw_exception_on_parse_error = True, 
            use_24hour_time_format = True
        )
        # Cron Schedule
        cron_label = QLabel(cron)
        #set max width for cron_label to 50 pixels
        cron_label.setMaximumWidth(50)
        #layout.addWidget(cron_label)

        # Human Readable Cron Format
        type_label = QLabel(job_data["type"])

        if job_data["type"] == "schedule":        

            readable_cron_label = QLabel(descriptor.get_description())
            readable_cron_label.setMaximumWidth(100)
      #  first_line_layout.addWidget(readable_cron_label)
        self.setFixedHeight(50)
        # Trash Icon Button for Removal
        remove_job_layout = QHBoxLayout()
        self.clickAction = QPushButton(self)
        self.clickAction.setIcon(QIcon(resource_path("trash.png")))
        self.clickAction.clicked.connect(self.remove_job)
        self.clickAction.setIconSize(QSize(10, 10))
        self.clickAction.setFixedSize(QSize(10, 10))
        if job_data["type"] == "schedule":        
           remove_job_layout.addWidget(readable_cron_label)
        else:
            remove_job_layout.addWidget(type_label)

        remove_job_layout.addWidget(self.clickAction)
        main_layout.addLayout(remove_job_layout)

        self.setLayout(main_layout)
    def remove_job(self):
        # Add code to remove the job associated with this widget
        # For example, you can emit a signal to handle the removal action
        
        #self.schedule[int(self.job_data["id"])]["job"].remove()
        #self.schedule.pop(int(self.job_data["id"]))
        for job in self.schedule:
            if job["id"] == self.job_data["id"]:
                self.schedule.remove(job)
        print(self.schedule)
        print("NEW SCHEDULE")
        import os.path
        print(self.schedule)
        send_jobs = []
        for job in self.schedule:
            test = job.copy()
            test["job"] = ""
            send_jobs.append(test)
        print(send_jobs)
        print(self.graph.user_key)
        r = requests.post("https://cheatlayer.com/user/saveDeskSchedule", data={'id': self.graph.user_key, "schedule": json.dumps(send_jobs)},  verify=False)
        print(r)
        print(r.text)
        print(r.status_code)
        
       # self.parent().parent().updateSchedule(self.schedule)
        self.updateSchedule(self.schedule)

    def toggle_switch(self):
        self.is_enabled = not self.is_enabled
        self.update_switch_color()
        # Add logic to enable or disable your job here

    def update_switch_color(self):
        if self.is_enabled:
            self.switch_button.setStyleSheet("background-color: green")
        else:
            self.switch_button.setStyleSheet("background-color: red")
import os
import json
from PySide2.QtWidgets import QPushButton, QToolBar, QScrollArea, QWidget, QVBoxLayout
from PySide2.QtGui import QIcon
from PySide2.QtCore import Qt, QUrl, QSize
from PySide2.QtGui import QIcon, QPixmap, QPainter

import queue
import threading
import json
from time import sleep

class JobRunner:
    def __init__(self, graph, thread_signals, user_key):
        self.webhook_jobs = [] # This should be populated with the jobs
        self.job_queue = self.load_queue()
        self.is_running = True  # Control flag to stop the thread
        self.graph = graph
        self.total_jobs = 0
        self.thread_signals = thread_signals
        self.worker_thread = threading.Thread(target=self.process_jobs)
        self.worker_thread.start()
        self.user_key = user_key
    def load_queue(self):
        return queue.Queue()
    def stop(self):
        self.is_running = False
        self.save_queue()

        self.worker_thread.join()

    def add_job(self, job) :
        print("ADDING JOB")
        print(job)
        if "parallel" in job and self.total_jobs < 20:
            print("PARALLEL")
            self.total_jobs += 1
            #run execute job in a separate thread
            job_folder_name = str(randint(0, 1000000000000))

            logger = logging.getLogger('thread-%s' % job_folder_name)
            logger.setLevel(logging.DEBUG)

            # create a file handler writing to a file named after the thread
            file_handler = logging.FileHandler('thread-%s.log' % job_folder_name)

            # create a custom formatter and register it for the file handler
            formatter = logging.Formatter('(%(threadName)-10s) %(message)s')
            file_handler.setFormatter(formatter)

            # register the file handler for the thread-specific logger
            logger.addHandler(file_handler)
            print("STARTING THREAD")
            thread = threading.Thread(target=self.execute_job, args=(job, job_folder_name, logger))
            thread.start()
      #      self.execute_job(job)
        else:
            self.job_queue.put(job)
            self.save_queue()
    
    def save_queue(self):
        with open('job_queue.json', 'w') as file:
            # A simple JSON dump of the list of jobs
            json.dump(list(self.job_queue.queue), file)
    def process_jobs(self):
        while self.is_running:
            try:
                job = self.job_queue.get(timeout=1)  # Timeout to check the flag periodically
                
                self.total_jobs += 1
                #run execute job in a separate thread
                job_folder_name = str(randint(0, 1000000000000))

                logger = logging.getLogger('thread-%s' % job_folder_name)
                logger.setLevel(logging.DEBUG)

                # create a file handler writing to a file named after the thread
                file_handler = logging.FileHandler('thread-%s.log' % job_folder_name)

                # create a custom formatter and register it for the file handler
                formatter = logging.Formatter('(%(threadName)-10s) %(message)s')
                file_handler.setFormatter(formatter)

                # register the file handler for the thread-specific logger
                logger.addHandler(file_handler)
                self.execute_job(job, job_folder_name, logger)
                self.job_queue.task_done()
                self.total_jobs -= 1
                self.save_queue()  # Save the queue after each job is processed

            except queue.Empty:
                continue
            except Exception as e:
                print(f"Error occurred: {e}")
    
    def execute_job(self, job, job_folder_name, logger):
        logger.info("Executing webhook")
        logger.info(job)
        logger.info("schedule")
        logger.info(self.graph.global_variables)
        # The rest of the job execution logic should be here

            # Create a directory specific for the job
        #generate a random ID for the folder name
    
        job_folder_path = os.path.join(os.getcwd(), job_folder_name)
        if not os.path.exists(job_folder_path):
            os.makedirs(job_folder_path)

        file_name = job['name'] + ".cheat"
        file_path = file_name

        with open(file_path) as file:
            data = json.load(file)
            current_desktop = 0  # Assuming `current_desktop` is previously defined
            self.thread_signals.recording_start.emit()

            for key, node in data["nodes"].items():
                custom = json.loads(node["custom"]["Data"])
                if custom["type"] == "Start Node":
                    # Pass job_folder_path to the runNode function:
                    logger.info("RUNNING NODE")
                    self.graph.runNode(data, key, self.thread_signals, False, True, folder=job_folder_path)
                    break
        logger.info("Finished webhook")
        self.thread_signals.recording.emit()
        #if global data includes a row parameter, send a webhook to update agent_tables data
        #open logs.txt and save it to self.graph.global_variables["logs"]

        with open(resource_path("logs.txt"), 'r') as file:
            logs = file.read()
            print(logs)
            print("LOGS SENT")
            self.graph.global_variables["logs"] = logs

        if "row" in self.graph.global_variables:
      # Replace 'your_api_key' with your actual Mailgun API key
                api_key = 'MAILGUNKEY'

                # Replace 'your_mailgun_domain' with your actual Mailgun domain

                # Mailgun API URL
                api_url = f'https://api.mailgun.net/v3/mg.cheatlayer.com/messages'

                # Email details
                from_email = 'agents@cheatlayer.com'  # Replace with your sending email address
                to_email = "desktop_" + self.graph.global_variables["key"] + "@mg.parsercheat.com"  # Replace with the recipient's email address
                subject = "desktop_agent"
                text = "desktop_agent"
                # Path to the attachment file
                file = ""
                self.graph.global_variables["key"] = ""
                self.graph.global_variables["URL"] = ""
                print(self.user_key)
                if job['name'] == "News_To_Blog_Generator_1":
                    file = "gpt4_blog_post.txt"
                
                    attachment_file_path = resource_path( os.path.join(job_folder_name,file))
                    file_text = ""
                    with open(attachment_file_path, 'r') as file:
                        file_text = file.read()
                    self.graph.global_variables["Blog text"] = file_text.strip() 

                    print("SEND Cheat Layer AGENT LOGS")
       
                    # add folder to the file 
                # Prepare the data for the API request

                data = {
                    'from': from_email,
                    'to': to_email,
                    'subject': self.graph.global_variables["row"] + ":" + self.graph.global_variables["column"] + ";" + self.graph.global_variables["sheet_id"] + ";" + self.graph.global_variables["tab"],
                    'text': json.dumps(self.graph.global_variables),
                }
                print(data)
                print(self.graph.global_variables)
                # Send the email using the Mailgun API
                if job['name'] == "News_To_Blog_Generator_1":

                    response = requests.post(
                        api_url,
                        auth=('api', api_key),
                        data=data
                    )
                else:
                    
                    response = requests.post(
                        api_url,
                        auth=('api', api_key),
                        data=data )

                # Check the response from Mailgun
                if response.status_code == 200:
                    print('Email sent successfully!')
                else:
                    print(f'Failed to send email: {response.status_code} {response.text}')
    def run(self, workflow):
        #for job in self.webhook_jobs:
        #    if workflow['cheat'] == job['script']:
        self.add_job(workflow)

   
class BottomToolbar(QToolBar):
    def __init__(self, schedule, scheduleCheat, user_key, openCheat, graph, thread_signals, workflows):
        QToolBar.__init__(self)
        self.setFloatable(False)
        self.setMovable(False)
        self.user_key = user_key
        self.scroll_widget = QWidget(self)
        self.scroll_layout = QVBoxLayout()
        self.scroll_widget.setLayout(self.scroll_layout)
        self.scheduleCheat = scheduleCheat
        self.openCheat = openCheat
        self.workflows = workflows

        self.thread_signals  = thread_signals
        self.scroll_layout.setContentsMargins(0, 0, 0, 0)
        self.scroll_layout.setSpacing(0)
        self.graph = graph
        total_height = 150
        for job in workflows:
            total_height += 200
            job_widget = JobWidget(job, workflows, graph, self.updateSchedule)
            self.scroll_layout.addWidget(job_widget)

        self.scroll_area = QScrollArea()
        self.scroll_area.setWidget(self.scroll_widget)
        self.addWidget(self.scroll_area)
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setFixedHeight(total_height)
        self.scroll_area.setFixedWidth(150)
        self.scroll_area.setContentsMargins(0, 0, 0, 0)
    # Look for .cheat files and add them to the toolbar in a separate layout
        self.icon_layout = QHBoxLayout()
        self.scroll_layout.addLayout(self.icon_layout)
        self.consoleOutput = QTextEdit()
        self.consoleOutput.setReadOnly(True)
        self.consoleOutput.setFixedHeight(400)
        self.consoleOutput.setFixedWidth(150)
        self.addWidget(self.consoleOutput)
                # Look for .cheat files and add them to the toolbar
        counter = 0

        CACHE_FILE = 'favicon_cache.json'

        def load_cache():
            if os.path.exists(CACHE_FILE):
                with open(CACHE_FILE, 'r') as f:
                    return json.load(f)
            return {}

        def save_cache(cache):
            with open(CACHE_FILE, 'w') as f:
                json.dump(cache, f)

        # Load the favicon cache
        favicon_cache = load_cache()

        def get_favicon(url, counter, local_path_function):
            icon_url = url + '/favicon.ico'
            print("start req")

            if icon_url in favicon_cache:
                print("Favicon cached")
                return favicon_cache[icon_url], counter

            response = requests.get(icon_url)
            print("end req") 
            print(response.status_code)
            if response.status_code == 200:
                local_path = local_path_function(str(counter) + 'favicon.ico')
                counter += 1
                try: 
                    with open(local_path, 'wb') as out_file:
                        out_file.write(response.content)
                    # Cache the local path
                    favicon_cache[icon_url] = local_path
                    # Save the updated cache
                    save_cache(favicon_cache)
                    return local_path, counter
                except:
                    return None, counter
            else:
                return None, counter

        data_graph = []
        counter = 0
        for file in os.listdir():
            if os.path.isfile(file) and file.endswith('.cheat'):
                with open(file, 'r') as f:
                    print("FILE")
                    print(file)
                    content = f.read()
                    icon = QIcon(resource_path("cheat.png"))
                    matches = re.findall(r'(\S*?\.com)', content)
                    try:
                        # Parse the JSON content from the file.
                        content = json.loads(content)
                        print("FILE")
                        print(file)  
                        # Add the content to data_graph
                        content["file"] = file

                        if len(matches) == 0:
                            matches.append("https://cheatlayer.com")
                        sent_URL = ""
                        for index, url in enumerate(set(matches)):
                            print(url)
                            icon_url = 'https://www.google.com/s2/favicons?domain=' + url
                            if "URL" not in icon_url and "git+" not in icon_url:
                                sent_URL = icon_url
                            print("UR")
                            local_path_function = resource_path # Or whatever your local path function is
                            favicon_path, counter = get_favicon(icon_url, counter, local_path_function)
                            if favicon_path is not None:
                                icon = QIcon(favicon_path)
                            print("DONE")
                        content["icon"] = sent_URL
                        data_graph.append(content)

                        def make_button(file, icon):
                            button = QPushButton()
                            button.clicked.connect(lambda: self.loadCheatFile(resource_path(file)))
                            button.setIcon(icon)
                            button.setFixedWidth(150)
                            file_name = file.replace('.cheat', '').replace('_', ' ').title()[:20]
                            padded_file_name = file_name.ljust(20, ' ')
                            # align the image on the button to the left
                            button.setStyleSheet("text-align:left")

                            button.setText(" " + padded_file_name)
                            self.addWidget(button)
                        make_button(file, icon)
                    except json.JSONDecodeError:
                        print(f"Error: Failed to parse JSON content in the file {file}")
                    except Exception as e:
                        print(f"unexpected error reading {file}: {e}")

        # Save the updated cache at the end of the script
        save_cache(favicon_cache)
        
    
        print(data_graph)
        print("DATA GRAPH")
#        loop = asyncio.new_event_loop()
    #    asyncio.set_event_loop(loop)
   #     browser_elements = loop.run_until_complete(sendMessageRTC('DATAGRAPH:' + json.dumps(data_graph)))
    #    loop.close()
#        print(browser_elements)
        print("fished loop")
        graph.data_graph = data_graph
        graph.server.data_graph = data_graph
       # sys.stdout = Stream(newText=self.normalOutputWritten)
        #sys.stderr = Stream(newText=self.normalOutputWritten)
        
        # Create QScrollArea
        self.scroll_area.horizontalScrollBar().setStyleSheet("""
       QScrollBar {height:0px;}
    """)
        self.scroll_area.verticalScrollBar().setStyleSheet("""
       QScrollBar {height:0px;}
    """)
        
    
    def normalOutputWritten(self, text):
        """Append text to the QTextEdit."""
        global graph
        # Maybe QTextEdit.append() works as well, but this is how I do it:
        cursor = self.consoleOutput.textCursor()
        cursor.movePosition(QtGui.QTextCursor.End)
        cursor.insertText(text)
        self.thread_signals.update.emit(text)
        self.consoleOutput.setTextCursor(cursor)
        self.consoleOutput.ensureCursorVisible()


    def loadCheatFile(self, filename):
        print("LOAD CHEAT FILE")
        print(filename)
        self.thread_signals.open.emit(filename)
        #self.openCheat(filename)
        # This function should define what happens when a .cheat button is clicked
    def updateSchedule(self, workflow):
        for i in reversed(range(self.scroll_layout.count())): 
            widget = self.scroll_layout.itemAt(i).widget()
            if widget is not None:  # Check if the widget exists before calling setParentapproval
                widget.setParent(None)    

        print("update schedule")
        print(workflow)

        total_height = 50
        for job in workflow:
            total_height += 50
            job_widget = JobWidget(job, workflow, self.graph, self.updateSchedule)
            self.scroll_layout.addWidget(job_widget)

        self.scroll_area.setFixedHeight(total_height)
        self.scroll_area.setFixedWidth(150)
        self.scroll_area.setContentsMargins(0, 0, 0, 0)
        self.scroll_layout.update()
        self.scroll_area.setWidget(self.scroll_widget)

        self.scroll_area.update()
        self.scroll_area.repaint()


class LeftToolbar(QToolBar):
    def __init__(self, addTracker, openPhone , addGeneral, addWebcam, addElevenLabs, addSynthesia, addLlama, getRecording, AddStableVideo, addScrape, addOCRSemantic, addDescribe, addOCR, addPrint, addScroll, addSendData, addIfElse, addKeypress, addMove, addClick,getData,getScreenshot,download_file, openCheat, scheduleCheat, addAIDetector, runOnCheatLayer, trainModel, launchAtlas, AddStableDiffusion, addCustomCode, addEmail,addRiku,addMath,addGsheets,addOpen,addDelay, newCheat, addGetFiles, AddGPT4):
        QToolBar.__init__(self)
        self.setFloatable(False)
        self.setMovable(False)

        self.scroll_widget = QWidget(self)
        self.scroll_layout = QVBoxLayout()
        self.scroll_widget.setLayout(self.scroll_layout)

        # Add your toolbar widgets here
        # Add your toolbar widgets here

        self.Navigation = QLabel(self)
        self.Navigation.resize(200, 28)
        self.Navigation.setText('Navigation Control')
        self.scroll_layout.addWidget(self.Navigation)
        self.Navigation.setFont(QtGui.QFont('Calibri', 8))

        self.clickAction = QPushButton(self)
        self.clickAction.setIcon(QtGui.QIcon(resource_path("Click.png")))
        self.clickAction.clicked.connect(addClick)
        self.clickAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.clickAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.clickAction)


        self.clickAction = QPushButton(self)
        self.clickAction.setIcon(QtGui.QIcon(resource_path("Describe.png")))
        self.clickAction.clicked.connect(addDescribe)
        self.clickAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.clickAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.clickAction)

       # self.approvalAction = QPushButton(self)
       # self.approvalAction.setIcon(QtGui.QIcon(resource_path("approval.png")))
      #  self.approvalAction.clicked.connect(addDescribe)
     #   self.approvalAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
      #  self.approvalAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
      #  self.scroll_layout.addWidget(self.approvalAction)

      #  self.scrapeAction = QPushButton(self)
     #   self.scrapeAction.setIcon(QtGui.QIcon(resource_path("Scrape.png")))
     #   self.scrapeAction.clicked.connect(addScrape)
     #   self.scrapeAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
     #   self.scrapeAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
     #   self.scroll_layout.addWidget(self.scrapeAction)

        self.keypressAction = QPushButton(self)
        self.keypressAction.setIcon(QtGui.QIcon(resource_path("Keypress.png")))
        self.keypressAction.clicked.connect(addKeypress)
        self.keypressAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.keypressAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.keypressAction)

        self.scrollAction = QPushButton(self)
        self.scrollAction.setIcon(QtGui.QIcon(resource_path("Scroll.png")))
        self.scrollAction.clicked.connect(addScroll)
        self.scrollAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.scrollAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.scrollAction)



        self.moveAction = QPushButton(self)
        self.moveAction.setIcon(QtGui.QIcon(resource_path("Move.png")))
        self.moveAction.clicked.connect(addMove)
        self.moveAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.moveAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.moveAction)



        self.openAction = QPushButton(self)
        self.openAction.setIcon(QtGui.QIcon(resource_path("OpenProgram.png")))
        self.openAction.clicked.connect(addOpen)
        self.openAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.openAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.openAction)


        self.threshLabel = QLabel(self)
        self.threshLabel.resize(200, 32)
        self.threshLabel.setText('Machine Learning')
        self.scroll_layout.addWidget(self.threshLabel)
        self.threshLabel.setFont(QtGui.QFont('Calibri', 8))



        self.gpt4 = QPushButton(self)
        self.gpt4.setIcon(QtGui.QIcon(resource_path("GPT4.png")))
        self.gpt4.clicked.connect(AddGPT4)
        self.gpt4.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.gpt4.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.gpt4)

        

        self.webcam = QPushButton(self)
        self.webcam.setIcon(QtGui.QIcon(resource_path("webcam.png")))
        self.webcam.clicked.connect(addWebcam)
        self.webcam.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.webcam.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.webcam)

        
        self.stableDiffusion = QPushButton(self)
        self.stableDiffusion.setIcon(QtGui.QIcon(resource_path("StableDiffusion.png")))
        self.stableDiffusion.clicked.connect(AddStableDiffusion)
        self.stableDiffusion.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.stableDiffusion.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.stableDiffusion)
        
 #       self.synthesia = QPushButton(self)
 #       self.synthesia.setIcon(QtGui.QIcon(resource_path("synthesia.png")))
 #       self.synthesia.clicked.connect(addSynthesia)
#        self.synthesia.setIconSize(QtCore.QSize(70*1.5,70*1.5))
#        self.synthesia.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
 #       self.scroll_layout.addWidget(self.synthesia)
        

        
        self.elevenlabs = QPushButton(self)
        self.elevenlabs.setIcon(QtGui.QIcon(resource_path("Voice.png")))
        self.elevenlabs.clicked.connect(addElevenLabs)
        self.elevenlabs.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.elevenlabs.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.elevenlabs)


        self.stableVideo = QPushButton(self)
        self.stableVideo.setIcon(QtGui.QIcon(resource_path("Synthetic_Video.png")))
        self.stableVideo.clicked.connect(AddStableVideo)
        self.stableVideo.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.stableVideo.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.stableVideo)

        self.Control = QLabel(self)
        self.Control.resize(200, 32)
        self.Control.setText('Automation Control')
        self.scroll_layout.addWidget(self.Control)

        self.Control.setFont(QtGui.QFont('Calibri', 8))

        self.genAction = QPushButton(self)
        self.genAction.setIcon(QtGui.QIcon(resource_path("GeneralTrigger.png")))
        self.genAction.clicked.connect(addGeneral)
        self.genAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.genAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.genAction)


        #self.cheatAction = QPushButton(self)
        #self.cheatAction.setIcon(QtGui.QIcon(resource_path("ExtensionTrigger.png")))
        #self.cheatAction.clicked.connect(runOnCheatLayer)
        #self.cheatAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        #self.cheatAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        #self.scroll_layout.addWidget(self.cheatAction)

        self.ifAction = QPushButton(self)
        self.ifAction.setIcon(QtGui.QIcon(resource_path("IfElse.png")))
        self.ifAction.clicked.connect(addIfElse)
        self.ifAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.ifAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.ifAction)


        self.customAction = QPushButton(self)
        self.customAction.setIcon(QtGui.QIcon(resource_path("CustomCode.png")))
        self.customAction.clicked.connect(addCustomCode)
        self.customAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.customAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.customAction)


        #self.mathAction = QPushButton(self)
        #self.mathAction.setIcon(QtGui.QIcon(resource_path("Math.png")))
        #self.mathAction.clicked.connect(addMath)
        #self.mathAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        #self.mathAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        #self.scroll_layout.addWidget(self.mathAction)



        self.delayAction = QPushButton(self)
        self.delayAction.setIcon(QtGui.QIcon(resource_path("Wait.png")))
        self.delayAction.clicked.connect(addDelay)
        self.delayAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.delayAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.delayAction)

        # Position the ToolBar on the left side of the Main Windo



      #  self.ExampleWidget1 = QPushButton(self)
     #   self.ExampleWidget1.setIcon(QtGui.QIcon(resource_path("OCRScraper.png")))
    #    self.ExampleWidget1.clicked.connect(addOCR)
     #   self.ExampleWidget1.setIconSize(QtCore.QSize(70*1.5,70*1.5))
     #   self.ExampleWidget1.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
      #  self.scroll_layout.addWidget(self.ExampleWidget1)



        self.Input = QLabel(self)
        self.Input.resize(200, 32)
        self.Input.setText('Data Management')
        self.scroll_layout.addWidget(self.Input)

        self.Input.setFont(QtGui.QFont('Calibri', 8))



        self.trackerAction = QPushButton(self)
        self.trackerAction.setIcon(QtGui.QIcon(resource_path("Extension.png")))
        self.trackerAction.clicked.connect(addTracker)
        self.trackerAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.trackerAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.trackerAction)
        
        self.getAction = QPushButton(self)
        self.getAction.setIcon(QtGui.QIcon(resource_path("GetData.png")))
        self.getAction.clicked.connect(getData)
        self.getAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.getAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.getAction)




        self.gsheetsAction = QPushButton(self)
        self.gsheetsAction.setIcon(QtGui.QIcon(resource_path("GoogleSheets.png")))
        self.gsheetsAction.clicked.connect(addGsheets)
        self.gsheetsAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.gsheetsAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.gsheetsAction)





        self.requestAction = QPushButton(self)
        self.requestAction.setIcon(QtGui.QIcon(resource_path("SendData.png")))
        self.requestAction.clicked.connect(addSendData)
        self.requestAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.requestAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.requestAction)



        self.screenshotAction = QPushButton(self)
        self.screenshotAction.setIcon(QtGui.QIcon(resource_path("record.png")))
        self.screenshotAction.clicked.connect(getRecording)
        self.screenshotAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.screenshotAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.screenshotAction)

        self.emailAction = QPushButton(self)
        self.emailAction.setIcon(QtGui.QIcon(resource_path("Email.png")))
        self.emailAction.clicked.connect(addEmail)
        self.emailAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.emailAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.emailAction)





        self.getFilesAction = QPushButton(self)
        self.getFilesAction.setIcon(QtGui.QIcon(resource_path("Download.png")))
        self.getFilesAction.clicked.connect(addGetFiles)
        self.getFilesAction.setIconSize(QtCore.QSize(70*1.5,70*1.5))
        self.getFilesAction.setFixedSize(QtCore.QSize(70*1.5,70*1.5))
        self.scroll_layout.addWidget(self.getFilesAction)




        # Create QScrollArea
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidget(self.scroll_widget)
        self.addWidget(self.scroll_area)

        self.scroll_area.horizontalScrollBar().setStyleSheet("""
       QScrollBar {height:0px;}
    """)
        self.scroll_area.verticalScrollBar().setStyleSheet("""
       QScrollBar {height:0px;}
    """)

class CustomMessageBox(QMessageBox):
    def __init__(self, *__args):
        QMessageBox.__init__(self)
        self.timeout = 0
        self.autoclose = False
        self.currentTime = 0

    def showEvent(self, QShowEvent):
        self.currentTime = 0
        if self.autoclose:
            self.startTimer(1000)

    def timerEvent(self, *args, **kwargs):
        self.currentTime += 1
        if self.currentTime >= self.timeout:
            self.done(0)

    @staticmethod
    def showWithTimeout(timeoutSeconds, message, title, icon=QMessageBox.Information, buttons=QMessageBox.Ok):
        w = CustomMessageBox()
        w.autoclose = True
        w.timeout = timeoutSeconds
        w.setText(message)
        w.setWindowTitle(title)
        w.setIcon(icon)
        w.setStandardButtons(buttons)
        w.exec_()

class NumpyEncoder(json.JSONEncoder):
    """ Special json encoder for numpy types """
    def default(self, obj):
        if isinstance(obj, np.integer):
            return int(obj)
        elif isinstance(obj, np.floating):
            return float(obj)
        elif isinstance(obj, np.ndarray):
            return obj.tolist()
        return json.JSONEncoder.default(self, obj)

class NodeGraph(QtCore.QObject):
    """
    The ``NodeGraph`` class is the main controller for managing all nodes
    and the node graph.

    Inherited from: :class:`PySide2.QtCore.QObject`

    .. image:: _images/graph.png
        :width: 60%
    """

    node_created = QtCore.Signal(NodeObject)
    """
    Signal triggered when a node is created in the node graph.

    :parameters: :class:`NodeGraphQt.NodeObject`
    :emits: created node
    """
    nodes_deleted = QtCore.Signal(list)
    """
    Signal triggered when nodes have been deleted from the node graph.

    :parameters: list[str]
    :emits: list of deleted node ids.fmf
    """
    node_selected = QtCore.Signal(NodeObject)
    """
    Signal triggered when a node is clicked with the LMB.

    :parameters: :class:`NodeGraphQt.NodeObject`
    :emits: selected node
    """
    node_selection_changed = QtCore.Signal(list, list)
    """
    Signal triggered when the node selection has changed.

    :parameters: list[:class:`NodeGraphQt.NodeObject`],
                 list[:class:`NodeGraphQt.NodeObject`]
    :emits: selected node, deselected nodes.
    """
    node_double_clicked = QtCore.Signal(NodeObject)
    """
    Signal triggered when a node is double clicked and emits the node.

    :parameters: :class:`NodeGraphQt.NodeObject`
    :emits: selected node
    """
    port_connected = QtCore.Signal(Port, Port)
    """
    Signal triggered when a node port has been connected.

    :parameters: :class:`NodeGraphQt.Port`, :class:`NodeGraphQt.Port`
    :emits: input port, output port
    """
    port_disconnected = QtCore.Signal(Port, Port)
    """
    Signal triggered when a node port has been disconnected.

    :parameters: :class:`NodeGraphQt.Port`, :class:`NodeGraphQt.Port`
    :emits: input port, output port
    """
    property_changed = QtCore.Signal(NodeObject, str, object)
    """
    Signal is triggered when a property has changed on a node.

    :parameters: :class:`NodeGraphQt.BaseNode`, str, object
    :emits: triggered node, property name, property value
    """
    data_dropped = QtCore.Signal(QtCore.QMimeData, QtCore.QPoint)
    """
    Signal is triggered when data has been dropped to the graph.

    :parameters: :class:`PySide2.QtCore.QMimeData`, :class:`PySide2.QtCore.QPoint`
    :emits: mime data, node graph position
    """
    session_changed = QtCore.Signal(str)
    """
    Signal is triggered when session has been changed.

    :parameters: :str
    :emits: new session path
    """

    def connectToLastNode(self,node_in):
        global nodes, graph
        graph_nodes = self.serialize_session()
        for key, node in graph_nodes["nodes"].items():
            found_in = False
            print(node)
            
            edit_node = self.get_node_by_id(key)
            edit_node.set_property('color', (240,240,240,255))
            edit_node.set_property('text_color', (0,0,0,255))
            print(key)
            found_out = False
            if "connections" not in graph_nodes:
                node_in.input(0).connect_to(nodes[0].output(0))
            else:
                for connections in graph_nodes["connections"]:
                    print(connections["out"])
                    if key == connections["out"][0]:
                        print("Found In")
                        found_out = True
                        print(connections["out"][1])
                        print(connections["out"][0])

                    if key == connections["in"][0]:
                        print("Found out")
                        found_in = True
                        print(connections["out"][1])
                        print(connections["out"][0])
                if found_in == True and found_out == False:
                    print(node)
                    print("Building Automation Graph. Please wait..")

                    node_in.input(0).connect_to(graph.get_node_by_id(key).output(0))
    def loopRecording(self):
        while True:
            self.runRecording()
    """
        Non-max Suppression Algorithm
        @param list  Object candidate bounding boxes
        @param list  Confidence score of bounding boxes
        @param float IoU threshold
        @return Rest boxes after nms operation
    """
    def nms(self, bounding_boxes, confidence_score, labels, threshold):
        # If no bounding boxes, return empty list
        if len(bounding_boxes) == 0:
            return [], []

        # Bounding boxes
        boxes = np.array(bounding_boxes)

        # coordinates of bounding boxes
        start_x = boxes[:, 0]
        start_y = boxes[:, 1]
        end_x = boxes[:, 2]
        end_y = boxes[:, 3]

        # Confidence scores of bounding boxes
        score = np.array(confidence_score)

        # Picked bounding boxes
        picked_boxes = []
        picked_score = []
        picked_labels = []

        # Compute areas of bounding boxes
        areas = (end_x - start_x + 1) * (end_y - start_y + 1)

        # Sort by confidence score of bounding boxes
        order = np.argsort(score)

        # Iterate bounding boxes
        while order.size > 0:
            # The index of largest confidence score
            index = order[-1]

            # Pick the bounding box with largest confidence score
            picked_boxes.append(index)

            # Compute ordinates of intersection-over-union(IOU)
            x1 = np.maximum(start_x[index], start_x[order[:-1]])
            x2 = np.minimum(end_x[index], end_x[order[:-1]])
            y1 = np.maximum(start_y[index], start_y[order[:-1]])
            y2 = np.minimum(end_y[index], end_y[order[:-1]])

            # Compute areas of intersection-over-union
            w = np.maximum(0.0, x2 - x1 + 1)
            h = np.maximum(0.0, y2 - y1 + 1)
            intersection = w * h

            # Compute the ratio between intersection and union
            ratio = intersection / (areas[index] + areas[order[:-1]] - intersection)

            left = np.where(ratio < threshold)
            order = order[left]

        return picked_boxes

    def StableDiffusion(self, opt_in):


            print(self.global_variables)
            print(globals)
            prompt = opt_in.replace("\\","").replace("/","").replace(":","").replace(".","").replace("\"","").replace(",","").replace("\n","")
            print("Inputs:")
            self.global_variables["Total Stable Diffusion Runs"] += 1
            self.global_variables["user"] = self.user_email
            print(prompt)
            def isEnglish(s):
                try:
                    s.encode(encoding='utf-8').decode('ascii')
                except UnicodeDecodeError:
                    return False
                else:
                    return True

      
    
            url = self.genImage(prompt)
            if "favion.ico" in url:
                return
            response = requests.get(url)


            with open(get_path(resource_path( "generated.png")), 'wb') as f:
                f.write(response.content)



    def StableDiffusionImg2Img(self, opt_in, processing):
            os.system("pkill -f chrome/chrome")
            os.system("taskkill /im firefox.exe")

            processing = True
            def chunk(it, size):
                it = iter(it)
                return iter(lambda: tuple(islice(it, size)), ())




            ckpt = resource_path("model.ckpt")

            class Struct:
                def __init__(self, **entries):
                    self.__dict__.update(entries)
            opt = Struct(**opt_in)
            #opt = {"prompt":'Austrian alps', "outdir":'outputs/img2img-samples', "init_img":'.\\tree.png', "skip_grid":False, "skip_save":False, "ddim_steps":50, "ddim_eta":0.0, "n_iter":2, "H":512, "W":512, "strength":0.8, "n_samples":5, "n_rows":0, "scale":7.5, "from_file":None, "seed":None, "device":'cuda', "unet_bs":1, "turbo":False, "precision":'autocast', "format":'png', "sampler":'ddim'}
            print(opt)
            tic = time.time()
            os.makedirs(opt.outdir, exist_ok=True)
            outpath = opt.outdir
            grid_count = len(os.listdir(outpath)) - 1


            toc = time.time()

            time_taken = (toc - tic) / 60.0
            print(
                (
                    "Samples finished in {0:.2f} minutes and exported to "
                    + sample_path
                    + "\n Seeds used = "
                    + seeds[:-1]
                ).format(time_taken)
            )


    def sendMessageRTCAsync(self, message):
        print("send message rtc async")
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        browser_elements = loop.run_until_complete(sendMessageRTC(message))
        loop.close()
        print(browser_elements)
        print("fished loop")
        if browser_elements == "" :
            print("Connect your extension or re-install to scrape.")
           
            return {"screenshots":[]}
        else:
            return json.loads(browser_elements)
    def browserClick(self, target):
        print("BROWSER CLICK")
        browser_elements = self.sendMessageRTCAsync('semanticClick:' + target)
        print(browser_elements)
        print("fished loop sent")
    def semanticSearch(self, description, x = 0, y = 0):

        browser_max = 0
        if x > 10 and y > 10:
            import subprocess, pyautogui
            pyautogui.moveTo(float(x), float(y) + 100)
           # pyautogui.moveTo(float(x), float(y))
       # time.sleep(5)
        browser_new_max = {}
      #  self.chatInstance.showMinimized()
        print("RUNNING SEMANTIC SEARCH")
      #  self.notification_manager.toggle_notifications_visibility()
        screen_width, screen_height = pyautogui.size()

        # Assume you want to capture the screenshot from (0, 20) to the bottom right corner
        # of the screen, leaving 20 pixels from the top. You have to calculate the height
        # of the region accordingly.
        region = (0, 0, screen_width, screen_height)

        # Capture the screenshot using the calculated region
        import time
        time.sleep(4)
        raw_image = pyautogui.screenshot()    #   self.notification_manager.toggle_notifications_visibility()
        box = (x -200/2, y - 200/2, x + 200/2, y + 200/2)
        crop_image_click = raw_image.crop(box)


        buf = io.BytesIO()
    
        try:
            rgb_screenshot = crop_image_click.convert("RGB")
            rgb_screenshot.save(buf, format='JPEG')
            byte_im = buf.getvalue()
            encoded_jpeg = base64.b64encode(io.BytesIO(byte_im).read()).decode('utf-8')
            raw_image.save(resource_path("screen.png"))
            old_elements = UImodel(resource_path("screen.png"))
            screen_width = screen_width
            screen_height = screen_height

            print(old_elements)
            elements = {"predictions": []}
            elements["predictions"] = old_elements
            for element in elements["predictions"]:
                element["x"] = (element['box']["x2"] - element['box']["x1"])/2 + element['box']["x1"]
                element["y"] = (element['box']["y2"] - element['box']["y1"])/2 + element['box']["y1"]
                element["width"] = (element['box']["x2"] - element['box']["x1"])
                element["height"] = (element['box']["y2"] - element['box']["y1"])
            print(elements)
            print("RUNNING SEMANTIC SEARCH")
            #UIModel.predict("screen.png", confidence=1, overlap=0).save("prediction.jpg")
            #set self.label3 to the prediction.jpg image
       #     pixmap = QPixmap()
       #     pixmap.load("prediction.jpg")
       #     resized_pixmap = pixmap.scaled(240, 360, Qt.KeepAspectRatio, Qt.SmoothTransformation)
# Attach     the resized pixmap to the label
        #self.label3.setPixmap(resized_pixmap)
       # self.label3.show()
        #self.label3.move(0, 400)
        #add elements to the array that include boxes of varying sizes around x and y that resemble buttons or other elements
          #add elements to the array that include boxes of varying sizes around x and y that resemble buttons or other elements

        # Creating a combined image with the width equal to the widest crop
            
            for index, element in enumerate(elements["predictions"]):
                # Crop the image
                box = (element["x"] - element["width"]/2, element["y"] - element["height"]/2, element["x"] + element["width"]/2, element["y"] + element["height"]/2)
                crop_image = raw_image.crop(box)
                element['cropped'] = crop_image
                element['browser'] = 'false'
            # Assuming raw_image is the original image you want to draw on

            buf2 = io.BytesIO()
            rgb_screenshot = raw_image.convert("RGB")
            clean_image = raw_image
            rgb_screenshot.save(buf2, format='JPEG')
            byte_im2 = buf2.getvalue()
            rgb_screenshot.save("prediction3.png")
            #save the image to disk using the x and y position in the name like screen_x_y.png
            #generate a file name based on the position first

    
            #self.label.setFixedSize(480, 1080)
            encoded_jpeg2 = base64.b64encode(io.BytesIO(byte_im2).read()).decode('utf-8')
            draw = ImageDraw.Draw(raw_image)
            font =  ImageFont.load_default() # You may need to change the font path according to your OS or font availability


            # Loop through all the elements to draw the index number next to the box

            from collections import namedtuple

            # Define a simple Rectangle class for easier handling of rectangles
            Rectangle = namedtuple("Rectangle", "x1 y1 x2 y2")
            def rectangles_overlap(rect1, rect2):
                """Return True if two rectangles overlap with up to 10 pixels overlap allowed."""
                overlap_allowance = -10
                return not (rect1.x2 + overlap_allowance < rect2.x1 or 
                            rect1.x1 > rect2.x2 + overlap_allowance or
                            rect1.y2 + overlap_allowance < rect2.y1 or 
                            rect1.y1 > rect2.y2 + overlap_allowance)

            # Keep track of already drawn rectangles to check for overlaps
            drawn_rectangles = []
            drawn_texts_positions = [] # This list will hold positions of texts that have been drawn

            for index, element in enumerate(elements["predictions"]):
                new_rectangle = Rectangle(
                    element["x"] - element["width"] / 2,
                    element["y"] - element["height"] / 2,
                    element["x"] + element["width"] / 2,
                    element["y"] + element["height"] / 2
                )

                # No operation related to drawn_rectangles shown, assuming you have that logic handled separately.

                # Calculate the center of the rectangle to position text
                rec_width = new_rectangle.x2 - new_rectangle.x1
                rec_height = new_rectangle.y2 - new_rectangle.y1

                min_dim = int(min(rec_height, rec_width)/2)

                rect_center_x = (new_rectangle.x2)  - int(min_dim/2)
                rect_center_y = (new_rectangle.y2)  - int(min_dim/2)

                text_position = (rect_center_x- 28, rect_center_y -28)

                # Check if the text position is sufficiently far from previously drawn texts
                too_close = False
                for pos in drawn_texts_positions:
                    if abs(pos[0] - text_position[0]) < 28 and abs(pos[1] - text_position[1]) < 28:
                        too_close = True
                        break
                #draw a rectangle around the element and fill it with white color

                if not too_close:
                    #prevent overlapping rectangles
                    too_close = False
                    for drawn_rectangle in drawn_rectangles:
                        if rectangles_overlap(new_rectangle, drawn_rectangle):
                            too_close = True
                            break
                    if not too_close:
                        draw.rectangle([new_rectangle.x1, new_rectangle.y1, new_rectangle.x2, new_rectangle.y2], outline=(255, 0, 0), width=2)
                        drawn_rectangles.append(new_rectangle)

                    # Draw the text if it's not too close to another
                        #adjust the font size to fill the rectangle height and width
                        rec_height = new_rectangle.y2 - new_rectangle.y1
                        rec_width = new_rectangle.x2 - new_rectangle.x1
                        min_dim = min(rec_height, rec_width)
                        #make sure the minimum font size is 56
                        if min_dim < 24:
                            min_dim = 24
                        # Assuming other parts of your code initialize these variables correctly
                        # image = Image.new(...)
                        font =  ImageFont.load_default()  # adjust the path to the font if necessary
                        # Calculate the size of the text to be drawn
                        text = str(index + 13)

# Use th    e fo        nt object to calculate text size
                        text_width = min_dim/2
                        if len(text) > 2:
                            text_width = min_dim
                        text_height = min_dim/2


                        # Adjust the text position to center the text in the rectangle
                        text_x = rect_center_x - text_width / 2
                        text_y = rect_center_y - text_height / 2
                        text_position = (text_x, text_y)

                        # Calculate the position for the white background based on the text size
                        padding = 1  # adjust the padding around the text
                        bg_x0 = text_x - padding
                        bg_y0 = text_y - padding
                        bg_x1 = text_x + text_width + padding
                        bg_y1 = text_y + text_height + padding

                        # Draw the white background
                        draw.rectangle([bg_x0, bg_y0, bg_x1, bg_y1], fill="white")

                        # Now draw the centered text over the white background
                        draw.text(text_position, text, fill=(255, 0, 0), font=font)
                # For a complete illustration, one would keep track of drawn rectangles
                # This line of code was missing from the initial snippet but should be there
                #draw a rectangle around the element

               # drawn_rectangles.append(new_rectangle)
                # Once all the annotations are drawn, save the annotated image
            print("Semantic Click")
            raw_image.save(resource_path("prediction2.png"))

            # Assume 'image' is a PIL Image object
            buf = io.BytesIO()
            rgb_screenshot = raw_image.convert("RGB")
            rgb_screenshot.save(buf, format='JPEG')
            byte_im = buf.getvalue()
            #save the image to disk using the x and y position in the name like screen_x_y.png
            #generate a file name based on the position first

            # Instead of saving and reloading the image from disk, use the in-memory bytes directly
            pixmap = QPixmap()
            pixmap.loadFromData(byte_im, 'JPEG')
            self.label.setPixmap(pixmap)
            #self.label.setFixedSize(480, 1080)
            encoded_jpeg = base64.b64encode(io.BytesIO(byte_im).read()).decode('utf-8')
            # Use the in-memory bytes to call `blip_caption` directly
                            #replace any strings that include {{name}} with self.global_variables["name"] in body 

            print("COMPLETED SEMANTIC SEARCH")
            print(self.global_variables)
            for key, value in self.global_variables.items():
                print(self.global_variables[key])
                print(key)
                if key in description:
                    description = description.replace("{{" + key + "}}", str(value))
            try:
                itm_output = blip_caption_new(encoded_jpeg, description, self.user_key, self.user_plan, encoded_jpeg2)
            except:
                itm_output = [0]
            #itm_output is an array of indices of the elements that match the description. loop through them, crop images for each element and send them to the model
            if len(itm_output) == 1:
                return elements['predictions'][int(itm_output[0])-13]
            else:
                log = [{"role": "system", "content": "You are provided multiple images of possible elements that closely match the description already. Your task is to select the image which most closely matches the description from the list of images based on all the qualities in the description. Output only the index of the image that most closely matches the description. The user will provide an index number for each image. Output only a single integer index number without a prefix or explanation. Do not putput any letters."}]
                for indicex in itm_output:
                    print(indicex)
                    print(elements["predictions"][int(indicex)-13])
                    element = elements["predictions"][int(indicex)-13]
                    box = (element["x"] - element["width"]/2, element["y"] - element["height"]/2, element["x"] + element["width"]/2, element["y"] + element["height"]/2)
                    crop_image = clean_image.crop(box)
                    buf_index = io.BytesIO()
                    crop_image.save(buf_index, format='JPEG')
                    byte_im_index = buf_index.getvalue()
                    data_index = base64.b64encode(io.BytesIO(byte_im_index).read()).decode('utf-8')
                    #save the image to inspect it 
                    crop_image.save(str(indicex) + "index.png")
                    log.append({"role": "user", "content": [{"type": "text", "text": "Element description:" + description +  "This is the smaller image showing only the element without surrounding context. The index for this image is:" + str(indicex)}, {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{data_index}"}}]})
                    #generate an image that is slightly larger than the element append it to log
                    box_big = (element["x"] - element["width"]/2 - 200, element["y"] - element["height"]/2 - 200, element["x"] + element["width"]/2 + 200, element["y"] + element["height"]/2 + 200)
                    crop_image_big = clean_image.crop(box_big)
                    buf_index_big = io.BytesIO()
                    crop_image_big.save(buf_index_big, format='JPEG')
                    byte_im_index_big = buf_index_big.getvalue()
                    data_index_big = base64.b64encode(io.BytesIO(byte_im_index_big).read()).decode('utf-8')
                    log.append({"role": "user", "content": [{"type": "text", "text": "Element description:" + description +  "This is a large image that shows the surrounding context for the element. Use this to diffferentiate between multiple similar elements based on their location and surrounding context. For example, for a grid of elements, choose the index that most closely exactly matches the cell. The index for this image is:" + str(indicex)}, {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{data_index_big}"}}]})
                print(log)
                data = {
                    "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
                    "messages": log,
                    "temperature": 0.7,
                    "max_tokens": 5000,
                    "stream": True  # Enable streaming to match original behavior
                }

                print(data)

                # Call local server instead of remote API
                response = requests.post("http://localhost:8000/v1/chat/completions", 
                                        headers=headers, 
                                        data=json.dumps(data), 
                                        stream=True)

                import pyautogui
                if response.status_code == 200:
                    items = []
                    total = ""
                    for chunk in response.iter_lines():
                        try:
                            if chunk:
                                # Parse SSE format from llama.cpp server
                                chunk_str = chunk.decode('utf-8')
                                if chunk_str.startswith('data: '):
                                    chunk_data = chunk_str[6:]  # Remove 'data: ' prefix
                                    if chunk_data.strip() == '[DONE]':
                                        break
                                    
                                    chunk_json = json.loads(chunk_data)
                                    if 'choices' in chunk_json and len(chunk_json['choices']) > 0:
                                        delta = chunk_json['choices'][0].get('delta', {})
                                        if 'content' in delta:
                                            content = delta['content']
                                            total += content
                                            print(content, end='', flush=True)
                        except Exception as e:
                            # If streaming fails, fall back to non-streaming
                            print(f"Streaming error: {e}")
                            break
                        
                    # If streaming didn't work or total is empty, try non-streaming
                    if not total:
                        data["stream"] = False
                        response = requests.post("http://localhost:8000/v1/chat/completions", 
                                                headers=headers, 
                                                data=json.dumps(data))
                        if response.status_code == 200:
                            result = response.json()
                            total = result["choices"][0]["message"]["content"]
                            print(total)
                        else:
                            print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                            total = "Error communicating with local LLM server"

                    print()  # New line after streaming output
                    print(total)
                else:
                    print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                    total = "Error communicating with local LLM server"

                if total == "":
                    total = itm_output[0]
                print("COMPLETED SEMANTIC SEARCH2")
                print(total)
                new_max = elements['predictions'][int(total)-13]
       #         model_url = 'https://storage.googleapis.com/sfr-vision-language-research/BLIP/models/model_base_capfilt_large.pth'
       #         image_size = 384
#       
       #         scores = []
       #         self.label.show()
       #         self.label2.show()
       #         def process_element(element):
       #             gc.collect()
       #             score_data = {"score": 0}
       #             if True:
       #                 caption = description
       #                 print('loop text: %s' %caption)
       #                 print('completed chat')
       #                 image = raw_image.crop((float(element['x']) - element['width']/2, float(element['y']) - element['height']/2, float(element['x']) + element['width']/2, float(element['y']) + element['height']/2))
       #                 #generate a unique string id for each element based on its coordinates
       #                  #this is used to save the cropped image to disk
       #                 id_element = str(element['x']) + str(element['y']) + str(element['width']) + str(element['height'])
#       
       #                 qpixmap = pil2pixmap(image)
#       
       #                 # Assume 'image' is a PIL Image object
       #                 buf = io.BytesIO()
       #                 image.save(buf, format='JPEG')
       #                 byte_im = buf.getvalue()
       #                 #save the image to disk using the x and y position in the name like screen_x_y.png
       #                 #generate a file name based on the position first
#       
       #                 
       #                 # Instead of saving and reloading the image from disk, use the in-memory bytes directly
       #                 pixmap = QPixmap()
       #                 pixmap.loadFromData(byte_im, 'JPEG')
       #                 self.label.setPixmap(pixmap)
       #                 self.label.setFixedSize(element['width'], element['height'])
       #                 encoded_jpeg = base64.b64encode(io.BytesIO(byte_im).read()).decode('utf-8')
#       
       #                 # Use the in-memory bytes to call `blip_caption` directly
       #                 itm_output = blip_caption(encoded_jpeg, caption, self.user_key, self.user_plan)
       #                 print(itm_output)
       #                 if len(itm_output) > 0:
       #                     score = float(itm_output)
       #                 else:
       #                     score = 0
       #                 print('completed output 3')
       #                 print(itm_output)
       #                 self.label2.setText("Semantic Step:" + caption + " Probability:" + str(score))
       #                 self.label2.setFixedSize(500,48)
       #                 print('The image and text is matched with a probability of %.4f'%score)
       #                 
       #                 file_name = str(element['x']) + "_" + str(element['y'])  + "_" + str(int(score*100))+ ".png"
       #                 
       #                 image.save(file_name, format="JPEG")
       #                 score_data = {"score": float(score), "x": element['x'], "y": element['y'], "width": element['width'], "height": element['height'], "cropped": pixmap, "browser": "false"}
       #             return score_data
#       
       #         # Function to split a list into chunks
       #         def chunks(lst, n):
       #             """Yield successive n-sized chunks from lst."""
       #             for i in range(0, len(lst), n):
       #                 yield lst[i:i + n]
#       
       #         # Process chunks of elements
       #         all_scores = []
       #         found = False
       #         total_runs = 0
       #         for chunk in chunks(elements['predictions'], 1):
       #             if found == True and total_runs > 5:
       #                  break
       #             with concurrent.futures.ThreadPoolExecutor() as executor:
       #                 scores_chunk = list(executor.map(process_element, chunk))
       #                 if found == False:
       #                     all_scores.extend(scores_chunk)
       #                 total_runs += 1
       #                 for score in scores_chunk:
       #                     print(score)
       #                     if float(score.get("score", 0)) > 0.99:
       #                         print("Semantic Click FOUND MATCH")
       #                         found = True
       #         if False:
       #             browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'a, button, input')
       #             print(browser_elements)
       #             print("fished loop sent")
       #             images_data_urls = browser_elements["screenshots"]  # Parse JSON into list of data URLs
       #             print(images_data_urls)
       #             print(len(images_data_urls))
       #             images = []
    #       
       #             # Loop through list of data URLs
       #             for data_url in images_data_urls:
       #                 # Strip header from data URL ('data:image/png;base64,')
       #                 if data_url != None and "data:image/png;base64," in data_url.img:
       #                     print(data_url)
       #                     base64_image = data_url.img.split(',')[1]
       #                     # Decode base64
       #                     decoded_image = base64.b64decode(base64_image)
       #                     # Open image and append to list
       #                     temp_image = Image.open(io.BytesIO(decoded_image))
       #                     # Convert transparent pixels to white
       #                     temp_image = temp_image.convert("RGBA")
       #                     images.append({"img":temp_image, "css": data_url.css})
    #       
       #                 else:
       #                     images.append(None)
       #                     
       #             # Now you have a list of PIL Image objects to process...
       #             browser_scores = []
       #             index = 0
       #             for image in images:
       #                 if image != None and image["img"] != None:
       #                     qpixmap = pil2pixmap(image["img"])
    #       
       #                     image["img"].save(resource_path(str(index) + "browser.png"))
       #                     pixmap = qpixmap
       #                     self.label.setPixmap(pixmap)
       #                     #set the label size based on the image width and height 
       #                     self.label.setFixedSize(image["img"].width, image["img"].height)
#      #                      self.label.setFixedSize(element['width'], element['height'])
       #                     print('completed browser output 1')
       #                     buf = io.BytesIO()
       #                     image["img"].save(buf, format='PNG')
       #                     byte_im = buf.getvalue()
       #                     itm_output = blip_caption(open(resource_path(str(index) + "browser.png"), "rb"), description)
       #                     print(itm_output)
       #                     score = float(itm_output.split("probability of ")[1].split(".\n")[0])
       #                     print('completed browser output 3')
       #                     print(itm_output)
       #                     self.label2.setText("Semantic Step browser:" + description + " Probability:" + str(score))
       #                     self.label2.setFixedSize(500,48)
       #                     print('The browser image and text is matched with a probability of %.4f'%score)
       #                     score_data = {"score": float(score), "x": -1, "y": -1, "width": -1 , "height": -1, "cropped": open(resource_path(str(index) + "browser.png"), "rb"), "index": index, "browser": "true"}
       #                     browser_scores.append(score_data)
    #       
       #                 index += 1
    #       
       #                 # Process image...
     # #                  image.show()
    #       
       #             #append browser scores to all scores
       #             all_scores.extend(browser_scores)
       #             for score in browser_scores:
       #                 itm_score = score.get("score", 0)
       #                 if itm_score > browser_max:
       #                     print("Semantic Click VALID")
       #                     browser_max = itm_score
       #                     browser_new_max = score
       #                     print(score)
       #             print(browser_new_max)
       #             print(browser_new_max["index"])
       #             print("browser max")
    #       
#       
#       
       #         #sort scores by highest probability
       #         all_scores.sort(key=lambda x: x["score"], reverse=True)
       #         # Continue with the rest of your code
       #         max_score = 0
      #          print("done append")
      #          new_max = {}
      #          print(all_scores)
      #          for score in all_scores:
      #              itm_score = score.get("score", 0)
      #              if itm_score > max_score:
      #                  print("Semantic Click VALID SEARCH")
      #                  max_score = itm_score
      #                  new_max = score
      #                  print(score)
      #          if max_score < browser_max:
      #              print("Browser Click VALID")
      #              new_max = browser_new_max
      #          # add 100 to the y coordinage of new_max
      #          
      #          new_max["y"] += 100
                return new_max
        except:
            return [0]


    def semanticSearchScrape(self, description, x = 0, y = 0, thresh=.95):
        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'span')
        print(browser_elements)
        print("fished loop sent")
        images_data_urls = browser_elements["screenshots"]  # Parse JSON into list of data URLs
        print(images_data_urls)
        print(len(images_data_urls))
        images = []

        # Loop through list of data URLs
        for data_url in images_data_urls:
            # Strip header from data URL ('data:image/png;base64,')
            if data_url != None and "data:image/png;base64," in data_url.img:
                print(data_url)
                base64_image = data_url.img.split(',')[1]
                # Decode base64
                decoded_image = base64.b64decode(base64_image)
                # Open image and append to list
                temp_image = Image.open(io.BytesIO(decoded_image))
                # Convert transparent pixels to white
                temp_image = temp_image.convert("RGBA")
                images.append({"img":temp_image, "css": data_url.css})

            else:
                images.append(None)
                
        # Now you have a list of PIL Image objects to process...
        browser_scores = []
        index = 0
        for image in images:
            if image != None and image["img"] != None:
                qpixmap = pil2pixmap(image["img"])

                image["img"].save(resource_path(str(index) + "browser.png"))
                pixmap = qpixmap
                self.label.setPixmap(pixmap)
                #set the label size based on the image width and height 
                self.label.setFixedSize(image["img"].width, image["img"].height)
#                self.label.setFixedSize(element['width'], element['height'])
                print('completed browser output 1')
                buf = io.BytesIO()
                image["img"].save(buf, format='PNG')
                byte_im = buf.getvalue()
                encoded_string = ""
                with open(resource_path(str(index) + "browser.jpeg"), "rb") as image_file:
                    encoded_string = base64.b64encode(image_file.read())

                itm_output = blip_caption(encoded_string, description, self.user_key, self.user_plan)
                print(itm_output)
                score = float(itm_output.split("probability of ")[1].split(".\n")[0])
                print('completed browser output 3')
                print(itm_output)
                self.label2.setText("Semantic Step browser:" + description + " Probability:" + str(score))
                self.label2.setFixedSize(1280,48)
                print('The browser image and text is matched with a probability of %.4f'%score)
                score_data = {"score": float(score), "x": -1, "y": -1, "width": -1 , "height": -1, "cropped": open(resource_path(str(index) + "browser.png"), "rb"), "index": index, "browser": "true"}
                browser_scores.append(score_data)

            index += 1

            # Process image...
     #       image.show()

        browser_max = 0
        browser_new_max = {}
        total_matches = []
        for score in browser_scores:
            itm_score = score.get("score", 0)
            if itm_score > browser_max:
                print("Semantic scrape VALID")
                browser_max = itm_score
                browser_new_max = score
                print(score)
            if itm_score > 0.95:
                browser_matches.append(score)
        print(browser_new_max)
        print(browser_new_max["index"])
        print("browser max")
     #   self.chatInstance.showMinimized()
        screen_width = pyautogui.size()[0]
        screen_height = pyautogui.size()[1]
        raw_image = pyautogui.screenshot(region=(0,0,screen_width,screen_height))
     #   self.chatInstance.toggle_chat_display()

        raw_image.save(r"screen.png")
        elements = UIModel.predict("screen.png", confidence=1, overlap=1).json()
        UIModel.predict("screen.png", confidence=1, overlap=1).save("prediction.jpg")
        #add elements to the array that include boxes of varying sizes around x and y that resemble buttons or other elements
          #add elements to the array that include boxes of varying sizes around x and y that resemble buttons or other elements
        if x != 0 and y != 0:
            elements['predictions'].insert(0,{"x": x, "y": y, "width": 100, "height": 100, "class": "Button"})
            elements['predictions'].insert(0,{"x": x, "y": y, "width": 200, "height": 200, "class": "Button"})
            elements['predictions'].insert(0,{"x": x, "y": y, "width": 300, "height": 300, "class": "Button"})
            elements['predictions'].insert(0,{"x": x, "y": y, "width": 50, "height": 50, "class": "Button"})
            elements['predictions'].insert(0,{"x": x, "y": y, "width": 200, "height": 50, "class": "Button"})
            elements['predictions'].insert(0,{"x": x, "y": y, "width": 200, "height": 100, "class": "Button"})
            elements['predictions'].insert(0,{"x": x, "y": y, "width": 300, "height": 300, "class": "Button"})
            elements['predictions'].insert(0,{"x": x, "y": y, "width": 400, "height": 400, "class": "Button"})
            elements['predictions'].insert(0,{"x": x, "y": y, "width": 500, "height": 200, "class": "Button"})
            
            elements['predictions'].insert(0,{"x": x, "y": y, "width": 50, "height": 50, "class": "Button"})
            elements['predictions'].insert(0,{"x": x, "y": y, "width": 200, "height": 50, "class": "Button"})
            elements['predictions'].insert(0,{"x": x, "y": y, "width": 200, "height": 50, "class": "Button"})
            elements['predictions'].insert(0,{"x": x, "y": y, "width": 300, "height": 50, "class": "Button"})
            elements['predictions'].insert(0,{"x": x, "y": y, "width": 400, "height": 50, "class": "Button"})
            elements['predictions'].insert(0,{"x": x, "y": y, "width": 500, "height": 200, "class": "Button"})
            elements['predictions'].insert(0,{"x": x, "y": y, "width": 50, "height": 200, "class": "Button"})
            elements['predictions'].insert(0,{"x": x, "y": y, "width": 200, "height": 200, "class": "Button"})
            elements['predictions'].insert(0,{"x": x, "y": y, "width": 200, "height": 200, "class": "Button"})
            elements['predictions'].insert(0,{"x": x, "y": y, "width": 300, "height": 200, "class": "Button"})
            elements['predictions'].insert(0,{"x": x, "y": y, "width": 400, "height": 200, "class": "Button"})

        print("Semantic Click")
        print(elements)

        model_url = 'https://storage.googleapis.com/sfr-vision-language-research/BLIP/models/model_base_capfilt_large.pth'
        image_size = 384

        scores = []
        self.label.show()
        self.label2.show()
        def process_element(element):
            gc.collect()
            score_data = {}
            if element["class"] != "TextView":
                caption = description
                print('loop text: %s' %caption)
                print('completed chat')
                image = raw_image.crop((float(element['x']) - element['width']/2, float(element['y']) - element['height']/2, float(element['x']) + element['width']/2, float(element['y']) + element['height']/2))
                #generate a unique string id for each element based on its coordinates
                 #this is used to save the cropped image to disk
                id_element = str(element['x']) + str(element['y']) + str(element['width']) + str(element['height'])

                qpixmap = pil2pixmap(image)

                image.save(resource_path(id_element + "crop.png"))
                pixmap = qpixmap
                self.label.setPixmap(pixmap)
                self.label.setFixedSize(element['width'], element['height'])
                print('completed output 1')
                buf = io.BytesIO()
                image.save(buf, format='JPEG')
                byte_im = buf.getvalue()
                
                encoded_string = ""
                with open(resource_path(id_element + "crop.png"), "rb") as image_file:
                    encoded_string = base64.b64encode(image_file.read())

                itm_output = blip_caption(encoded_string, caption, self.user_key, self.user_plan)
                print(itm_output)
                score = float(itm_output.split("probability of ")[1].split(".\n")[0])
                print('completed output 3')
                print(itm_output)
                self.label2.setText("Semantic Step:" + caption + " Probability:" + str(score))
                self.label2.setFixedSize(1280,48)
                print('The image and text is matched with a probability of %.4f'%score)
                score_data = {"score": float(score), "x": element['x'], "y": element['y'], "width": element['width'], "height": element['height'], "cropped": open(resource_path(id_element + "crop.png"), "rb"), "browser": "false"}
            return score_data

        # Function to split a list into chunks
        def chunks(lst, n):
            """Yield successive n-sized chunks from lst."""
            for i in range(0, len(lst), n):
                yield lst[i:i + n]

        # Process chunks of elements
        all_scores = []
        found = False
        total_runs = 0
        for chunk in chunks(elements['predictions'], 5):
            if found == True and total_runs > 5:
                 break
            with concurrent.futures.ThreadPoolExecutor() as executor:
                scores_chunk = list(executor.map(process_element, chunk))
                if found == False:
                    all_scores.extend(scores_chunk)
                total_runs += 1
                for score in scores_chunk:
                    print(score)
                    if float(score.get("score", 0)) > 0.99:
                        print("Semantic Click FOUND MATCH")
                        found = True

        # Continue with the rest of your code
        max_score = 0
        print("done append")
        new_max = {}
        for score in all_scores:
            itm_score = score.get("score", 0)
            if itm_score > max_score:
                print("Semantic scrape VALID")
                max_score = itm_score
                new_max = score
                print(score)
            if itm_score > thresh:
                total_matches.append(score)
        total_matches.append(new_max)

        return total_matches

    def verifyRequest(self, description):
        print("Verify Request")
        print(description)
      #  self.chatInstance.showMinimized()
        raw_image = pyautogui.screenshot()
# infer on a local image
        raw_image.save(r"screen.png")
        elements = UIModel.predict("screen.png", confidence=1, overlap=50).json()
        UIModel.predict("screen.png", confidence=3, overlap=1).save("prediction.jpg")

        print("Semantic Click")
        print(elements)
        scores = []

        model_url = 'https://storage.googleapis.com/sfr-vision-language-research/BLIP/models/model_base_capfilt_large.pth'
        image_size = 384
        # Load the image into a Q

          # print(element['x'] + " " + element['y'])
        caption = description
        print('loop text: %s' %caption)
            # Iterate over the response in ch   unks
        print('completed chat')
        image = raw_image
        cropped = load_demo_image(image_size=image_size, device=device, input = image)
        itm_output = model(cropped,caption,match_head='itm')

        print(itm_output)
        print('completed output')
    #    itm_score = torch.nn.functional.softmax(itm_output,dim=1)[:,1]
        #self.chatWindow(self.chatInstance,"Semantic Step:" + caption + " Probability:" +  '%.4f'%itm_score, image)
        #Image is a PIL image. convert image from pillow format to QPixmap
        # Add the QLabel to the QVBoxLayout

        # Show the QWidget
        print('The image and text is matched with a probability of %.4f'%itm_score)
        return itm_score


    def videoSearch(self, search):

        #convert this curl request to a python request curl -H "Authorization: tG0dqF7btj8DvBUgST6FJPCT6c7IkS7fVhjpLuL2kE7cIS6LYVFnm8wm"  "https://api.pexels.com/videos/search?query=nature&per_page=1"
        headers = {"Authorization": "tG0dqF7btj8DvBUgST6FJPCT6c7IkS7fVhjpLuL2kE7cIS6LYVFnm8wm"}
        response = requests.get("https://api.pexels.com/videos/search?query=" + search.replace(" ","%20") +"&per_page=1", headers=headers)
        print(response.json())
        print("VIDEO RESONSE")
        return response.json()
    def setup_local_llama_server(self):
        """
        Download the SmolVLM-500M model if needed and start a local llama.cpp server with vision capabilities
        """
        print("Setting up local llama.cpp server with SmolVLM-500M-Instruct model...")

        # Clone llama.cpp repository
        if not os.path.exists("llama.cpp"):
            subprocess.run(["git", "clone", "https://github.com/ggerganov/llama.cpp.git"], check=True)

        # Build llama.cpp with cmake instead of make
        if not os.path.exists("llama.cpp/build/bin/llama-server"):
            os.chdir("llama.cpp")

            # Create build directory if it doesn't exist
            if not os.path.exists("build"):
                os.makedirs("build")

            # Run cmake
            os.chdir("build")
            subprocess.run(["cmake", "..", "-DLLAMA_BLAS=ON", "-DLLAMA_BLAS_VENDOR=OpenBLAS"], check=True)

            # Run make or equivalent build command
            subprocess.run(["cmake", "--build", ".", "--config", "Release"], check=True)

            # Return to original directory
            os.chdir("../..")

        # Create models directory for SmolVLM
        os.makedirs("llama.cpp/models/smolvlm", exist_ok=True)

        # Check if model exists and verify it's valid, re-download if corrupted
        model_path = "llama.cpp/models/smolvlm/smolvlm-instruct-q4_k_m.gguf"

        def is_valid_gguf(filepath):
            """Check if GGUF file has valid magic header"""
            try:
                with open(filepath, 'rb') as f:
                    magic = f.read(4)
                    return magic == b'GGUF'
            except:
                return False

        if not os.path.exists(model_path) or not is_valid_gguf(model_path):
            if os.path.exists(model_path):
                print("Model file appears corrupted, re-downloading...")
                os.remove(model_path)

            print("Downloading SmolVLM-500M-Instruct model (this may take a while)...")
            # Try alternative working model since SmolVLM URLs might be unstable
            subprocess.run([
                "wget", "-O", model_path,
                "https://huggingface.co/microsoft/Phi-3-mini-4k-instruct-gguf/resolve/main/Phi-3-mini-4k-instruct-q4.gguf"
            ], check=True)

            # Verify download
            if not is_valid_gguf(model_path):
                raise Exception("Downloaded model file is corrupted")

        # For now, let's skip the projector and use text-only mode
        # projector_path = "llama.cpp/models/smolvlm/smolvlm-instruct-mmproj-f16.gguf"
        # if not os.path.exists(projector_path):
        #     print("Downloading SmolVLM multimodal projector (this may take a while)...")
        #     subprocess.run([
        #         "wget", "-O", projector_path,
        #         "https://huggingface.co/ggml-org/SmolVLM-500M-Instruct-GGUF/resolve/main/smolvlm-instruct-mmproj-f16.gguf"
        #     ], check=True)

        # Check if server is already running
        try:
            response = requests.get("http://localhost:8000/v1/models", timeout=5)
            if response.status_code == 200:
                print("Local llama.cpp server is already running")
                return
        except:
            pass # Server not running, we'll start it
        
        # Start the server
        print("Starting llama.cpp server with Phi-3-mini model...")

        # Determine number of CPU threads
        cpu_threads = os.cpu_count()
        if cpu_threads > 4:
            cpu_threads = cpu_threads - 2 # Leave some CPU for other tasks

        # Use the correct path to the server executable
        server_path = os.path.abspath("llama.cpp/build/bin/llama-server")

        # Start without multimodal projector for now (text-only mode)
        self.llama_server_process = subprocess.Popen([
            server_path,
            "--model", os.path.abspath(model_path),
            # "--mmproj", os.path.abspath(projector_path),  # Commented out for now
            "--host", "0.0.0.0",
            "--port", "8000",
            "--threads", str(cpu_threads),
            "--ctx-size", "2048",
            "--n-gpu-layers", "0"
        ])

        # Give server time to start up and verify it's running
        print("Waiting for server to start...")
        for i in range(30):
            try:
                response = requests.get("http://localhost:8000/v1/models", timeout=2)
                if response.status_code == 200:
                    print("Local llama.cpp server is running with Phi-3-mini!")
                    return
            except:
                pass
            time.sleep(2)

        print("Warning: Server may not have started properly. Check the process manually.")
    def gpt3Prompt(self, system_prompt, description):
    # Setup local llama.cpp server if it's not running
        if not hasattr(self, 'llama_server_process') or self.llama_server_process is None:
            self.setup_local_llama_server()
    
        log = [{"role": "system", "content": system_prompt}]
        log.append({"role": "user", "content": description})
    
        headers = {
            "Content-Type": "application/json",
        }
    
        data = {
            "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
            "messages": log,
            "temperature": 0.7,
            "max_tokens": 5000,
        }
        
        # Call local server instead of remote API
        response = requests.post("http://localhost:8000/v1/chat/completions", 
                                 headers=headers, 
                                 data=json.dumps(data))
    
        if response.status_code == 200:
            result = response.json()
            total = result["choices"][0]["message"]["content"]
            print(total)
            return total
        else:
            print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
            return "Error communicating with local LLM server"
    
    
    def upload_image(self, file_path):
        """
        Uploads an image to the specified endpoint.

        Args:
        - url (str): The URL of the endpoint where the image will be uploaded.
        - file_path (str): The path to the image file to be uploaded.

        Returns:
        - response: Server's response to the HTTP POST request.
        """
        # Open the file in binary mode
        with open(file_path, 'rb') as file:
            files = {'file': (file_path, file, 'multipart/form-data')}
            if "mp4" in file_path:
                response = requests.post("https://cheatlayer.com/user/upload", files=files, data={"id": self.user_key, "cheat": self.last_cheat.replace(".",""), "mp4": "true"}, verify=False)
            else:
                response = requests.post("https://cheatlayer.com/user/upload", files=files, data={"id": self.user_key, "cheat": self.last_cheat.replace(".",""), "mp4": "false"}, verify=False)

        return response

    def imageToVideo(self, image_URL, file_name = "synthetic.mp4"):
        response_URL = call_replicate3(image_URL)
        #download and open the .mp4
        response = requests.get(response_URL)
        with open(resource_path(file_name), 'wb') as f:
            f.write(response.content)
        #open the mp4
        #os.system("start " + resource_path(file_name))
        self.upload_image(resource_path(file_name))
        print("FINISED SYNTHETIC VIDEO GENERATOR")
        return response_URL
        #send the image to the server
    def genSyntheticVideo(self, description, file_name = "synthetic.mp4"):
        if description == "":
            description = "A video of a person walking in a park"
        response = requests.post("https://cheatlayer.com/user/generateDalleU", data={'id': self.user_key, "input":description},  verify=False)

        print("RUNNING SYNTHETIC VIDEO GENERATOR 2")
        print(response.json())
        if response.status_code != 200:
            print(response.json())

            raise ValueError("Failed to start image generation. Status code:", response.json())
    
        data = response.json()
        URL = data["data"][0]["url"]
        #"sdGenerationJob":{"generationId":
        #download and convert the image to base64
        response = requests.get(URL)
        img = Image.open(io.BytesIO(response.content))
        img.save(resource_path("image.png"))
        self.upload_image(resource_path("image.png"))
        print("RUNNING SYNTHETIC VIDEO GENERATOR 2")
        response_URL = call_replicate3(URL)
        #download and open the .mp4
        response = requests.get(response_URL)
        with open(resource_path(file_name), 'wb') as f:
            f.write(response.content)
        #open the mp4
        #save the URL to a text file
        with open(resource_path("logs.txt"), 'w') as f:
            f.write(response_URL)
        self.upload_image(resource_path(file_name))

        #os.system("start " + resource_path(file_name))
        self.global_variables["synthetic_video"] = response_URL
        print("FINISED SYNTHETIC VIDEO GENERATOR PYTHON")
        return response_URL
        #send the image to the server

    def genImage(self, request):
        API_URL = "https://cheatlayer.com/user/generateDalleU"

        headers = {
            'content-type': 'application/json'
        }
        print(request)
        data = {
            "id": self.user_key,
            "height": 512,
            "modelId": "6bef9f1b-29cb-40c7-b9df-32b51c1f67d3",
            "prompt": request,
            "width": 512
        }
    
        # Make the POST request to generate the image
        response = requests.post("https://cheatlayer.com/user/generateDalleU", data={'id': self.user_key, "input":request},  verify=False)

      #  response = requests.post(API_URL, headers=headers, data=json.dumps(data))
        print(response)
        if response.status_code != 200:
            print(response.json())

            raise ValueError("Failed to start image generation. Status code:", response.json())
    
        data = response.json()
        print(data)
        #download the image and update the custom node widget image 
        URL = "https://cheatlayer.com/favicon.ico"
        if "data" in data and "url" in data["data"][0]:
            URL = data["data"][0]["url"]
    #        URL = data["data"][0]["url"]

        response = requests.get(URL)
        img = Image.open(io.BytesIO(response.content))
        img.save(resource_path("image.png"))
        #update the node widget image 
        
        self.upload_image(resource_path("image.png"))


        #"sdGenerationJob":{"generationId":
        return data["data"][0]["url"]
    def check_duplicates(duplicates_file, entry):
        with open(duplicates_file, 'r+') as file:
            existing_entries = file.read().splitlines()
            if entry in existing_entries:
                return True
            else:
                file.write(entry + '\n')
                return False
            
    def store_analytics(self, json_data):
        filename = 'analytics.json'
        file_data = []

        # Check if the file exists and read its contents if it does
        if os.path.isfile(filename):
            with open(filename, 'r') as file:
                try:
                    file_data = json.load(file)
                except json.JSONDecodeError:
                    print('File is not a valid JSON. Starting a fresh file.')

        # Verify that file_data is a list
        if type(file_data) is not list:
            print('The existing data is not a list. Starting a fresh list.')
            file_data = []

        # Append the new JSON data
        file_data.append(json_data)

        # Write back to the file
        with open(filename, 'w') as file:
            json.dump(file_data, file, indent=4)

        print(f'JSON data has been appended to {filename}')
    def genVoice(self,input_data, filename = "out.mp3"):
        url = "https://cheatlayer.com/user/genVoiceUnauth"
        headers = {
            "Content-Type": "application/json"
        }
        data = json.dumps({"input": input_data, 'id': self.user_key})

        response = requests.post(url, headers=headers, data=data)

        if response.status_code == 200:
            with open(resource_path(filename), 'wb') as f:
                f.write(response.content)
            print(f"Audio saved as '{filename}'")
        else:
            print("Error:", response.status_code, response.text)
    def videoAPI(self, description):
        print("Video API")
    def semanticClick(self, description, mode= "First"):
        global websocket_data
        self.chatInstance.showMinimized()
        print("SEMANTIC CLICK")
        self.thread_signals.hide.emit()

        self.label.hide()
        self.label2.hide()
        #self.notification_manager.add_notification("Semantic Click", "Searching the screen for the UI element which matches the intent of the next step: " + description, "Click.png", lambda: print("Notification closed"))
      #  self.notification_manager.toggle_notifications_visibility()
        first_score = 0
        if first_score < 0.95:
            matches = self.semanticSearch(description, last_mouse_x, last_mouse_y)
            print(matches)
            print("semantic search result")
            new_max = matches
     
            
            
            print(new_max)
            if 'cropped' in new_max:
              #  new_max['cropped'].save(resource_path("combined_image.png"))
                self.label.setPixmap(pil2pixmap(new_max['cropped']))
                self.label2.setText("Winning Semantic Step:" + description + " Probability:1.0")
                #save the pixmap to disk
              
                self.label.setFixedSize(new_max['width'], new_max['height'])
          #  self.label2.setFixedSize(500,48)
           # time.sleep(5)
          #  self.label.hide()
         #   self.label2.hide()
            print("winning click")
            if new_max["browser"] == "false":
                print("trying to click")
                pyautogui.click(new_max['x'], new_max['y'])
                

            else:
                self.sendMessageRTC("semanticClick:" + str(new_max['index']))
                websocket_data = ""
            #pyautogui.click(new_max['x'], new_max['y'])
            #break
        else:
            print("winning click")
            pyautogui.click(last_mouse_x, last_mouse_y)
        self.chatInstance.toggle_chat_display()
        self.thread_signals.show.emit()

                #break
# visualize your prediction
# model.predict("your_image.jpg", confidence=40, overlap=30).save("prediction.jpg")

# infer on an image hosted elsewhere
# print(model.predict("URL_OF_YOUR_IMAGE", hosted=True, confidence=40, overlap=30).json())
#        screen = self.QMainWindow.primaryScreen()
 #       print('Screen: %s' % screen.name())
 #       size = screen.size()
 #       print('Size: %d x %d' % (size.width(), size.height()))
 #       print("Semantic Click")
        print(description)
    def semanticMove(self, description):
        print("Semantic Move")
        print(description)
    def move(self, x, y):
        pyautogui.moveTo(x, y)
    def click(self, x, y):
        pyautogui.click(x, y)
    def keypress(self, key):
       # time.sleep(2)
        x = {"key": key}
     #   self.notification_manager.add_notification("Keypress", "Automating keyboard to type " + key, "Keypress.png", lambda: print("Notification closed"))

        if True:
            if "resource_path" in str(x["key"]).lower():
                pyautogui.write(resource_path(''), interval=0.05)
            if "backspace" in str(x["key"]).lower():
                pyautogui.press('backspace')
            elif "space" in str(x["key"]).lower():
                pyautogui.write(" ", interval=0.05)
            elif "period" in str(x["key"]).lower():
                pyautogui.write(".", interval=0.05)
            elif "slash" in str(x["key"]).lower():
                pyautogui.write("/", interval=0.05)
            elif "return" in str(x["key"]).lower():
                pyautogui.press('enter')
            elif "shift" in str(x["key"]).lower():
                pyautogui.keyDown("shift")
                self.shift_down = True
            elif "ctrl" in str(x["key"]).lower():
                pyautogui.keyDown("ctrl")
                self.ctrl_down = True
            elif "alt" in str(x["key"]).lower():
                pyautogui.keyDown("alt")
                self.alt_down = True
            elif "tab" in str(x["key"]).lower():
                pyautogui.keyDown("tab")
            else:
                if self.shift_down:  
                    clipboard.copy(str(x["key"]).upper())
                    pyautogui.hotkey('ctrl', 'v')
                   # pyautogui.write(str(x["key"]).upper(), interval=0.01)
                else:
                    clipboard.copy(str(x["key"]).upper())
                    pyautogui.hotkey('ctrl', 'v')
                   # pyautogui.write(str(x["key"]).lower(), interval=0.01)
    def sendData(self, url, data):
        print("Send Data")
        print(data)
    def getData(self, url):
        print("Get Data")
        print(url)
        
        from urllib.parse import urlparse

    def store_creator_stats(self, user_id, creator_stats):
        """
        Stores creator stats under 'creator/{user_id}' in Firebase,
        removing duplicates by normalizing the 'url' field.
        """
        # 1. Build a dict keyed by normalized URL so duplicates overwrite
        unique_by_url = {}
        for item in creator_stats:
            raw_url = item.get("url", "")
            if not raw_url:
                continue  # skip if there's no URL at all
            
            norm = normalize_url(raw_url)
            # This overwrites if we see the same normalized url again
            unique_by_url[norm] = item
    
        # 2. Convert dict back to a list
        deduped_stats = list(unique_by_url.values())
    
        # 3. Store in Firebase
        ref = db.reference(f'creator/{user_id}')
        ref.set(deduped_stats)
    
        print(
            f"Stored creator stats for user_id={user_id} "
            f"(deduped {len(creator_stats)} -> {len(deduped_stats)})"
        )
    def scroll_to_load_all(self, driver, max_scrolls=5, wait_time=2):
        """
        Scrolls down the page multiple times to load more videos.
        
        - max_scrolls: max number of repeated scroll attempts.
        - wait_time: how many seconds to wait after each scroll before checking the new height.
        """
        last_height = driver.execute_script("return document.body.scrollHeight")

        for _ in range(max_scrolls):
            # Scroll to the bottom
            driver.execute_script("window.scrollTo(0, document.body.scrollHeight)")
            time.sleep(wait_time)

            # Check new height
            new_height = driver.execute_script("return document.body.scrollHeight")
            if new_height == last_height:
                # No change in height => no more content loaded
                print("No more new videos loaded. Stopping scroll.")
                break
            last_height = new_height

    def store_final_states(self, user_id, final_states, page_url):
        """
        1. Always go to 'https://www.tiktok.com'.
        2. Click [data-e2e="nav-profile"], refresh, scroll to load all videos.
        3. Extract the video ID from each <a> in the loaded feed.
        4. Match the video ID to final_states, parse <strong data-e2e="video-views"> for each.
        5. Store updated results in Firebase at 'final/{user_id}'.
        """

        # Hardcode the main TikTok page
      #  page_url = "https://www.tiktok.com/profile"

        options = webdriver.ChromeOptions()
        # Example settings for Chrome:
        options.add_argument("--password-store=basic")
        options.add_argument("--start-maximized")
        options.add_argument("--disable-notifications")
        options.add_argument("--disable-popup-blocking")
        # options.add_argument("--headless")  # Use headless mode if desired

        driver = webdriver.Chrome(options=options)
        try:
            # --- 1. Navigate to TikTok
            driver.get(page_url)
            wait = WebDriverWait(driver, 10)

            # --- 2. Click the nav-profile link
            #nav_profile_selector = (By.CSS_SELECTOR, '[data-e2e="nav-profile"]')
            #wait.until(EC.element_to_be_clickable(nav_profile_selector)).click()
            #time.sleep(2)

            # Refresh + scroll to ensure all videos load
            driver.refresh()
            time.sleep(2)
            self.scroll_to_load_all(driver, max_scrolls=5, wait_time=2)

            # Wait for at least one video-views element
            wait.until(
                EC.presence_of_element_located(
                    (By.CSS_SELECTOR, "strong[data-e2e='video-views']")
                )
            )

            # --- 3. Build a map of {video_id -> link_element} from all <a> found
            link_elements = driver.find_elements(By.CSS_SELECTOR, "a[href*='/video/']")
            link_map = {}
            for link in link_elements:
                href = link.get_attribute("href")
                # Extract video ID via regex
                VIDEO_ID_PATTERN = re.compile(r"/video/(\d+)")
                match = VIDEO_ID_PATTERN.search(href)
                if match:
                    vid_id = match.group(1)
                    link_map[vid_id] = link

            # --- 4. Update final_states
            # Build a dict for quick access
            final_dict = {item["url"]: item for item in final_states}

            for url, item in final_dict.items():
                # Extract the video ID from the final_state's url
                VIDEO_ID_PATTERN = re.compile(r"/video/(\d+)")
                match_url = VIDEO_ID_PATTERN.search(url)
                if not match_url:
                    continue

                item_vid_id = match_url.group(1)
                if item_vid_id in link_map:
                    # Find <strong data-e2e="video-views"> in this link
                    link = link_map[item_vid_id]
                    try:
                        strong_tag = link.find_element(
                            By.CSS_SELECTOR, "strong[data-e2e='video-views']"
                        )
                        view_count_text = strong_tag.text  # e.g., "160"
                        views = int(view_count_text.replace(",", ""))
                        item["views"] = views
                    except Exception:
                        item["views"] = 0

            # Convert dict back to list
            updated_final_states = list(final_dict.values())

            # --- 5. Store updated states in Firebase
            ref = db.reference(f'final/{user_id}')
            ref.set(updated_final_states)
            print(
                f"Updated final states for user_id={user_id} with Selenium-based view counts."
            )

        finally:
            driver.quit()

    def retrieve_creator_stats(self, user_id):
        """Retrieves creator stats from the 'creator/' folder in Firebase."""
        ref = db.reference(f'creator/{user_id}')
        creator_stats = ref.get()
        print(f"Retrieved creator stats for user_id: {user_id}")
        return creator_stats
    
    def runNode(self, graph, node_id, thread_signals, direct = True, manual = False, folder = "default"):
        #self.QMainWindow.showMinimized()
        global last_mouse_x, last_mouse_y, run_counter, agent_approval, python_process, total_errors, agent_memory, websocket_data, node_runs
        import numpy as np
        #
     #   sys.stdout = Stream(newText=self.normalOutputWritten)
     #   sys.stderr = Stream(newText=self.normalOutputWritten)
        #return sys.stdout to their original values
        #return sys.stderr

        print("graph run node working")
#        print(graph["nodes"][node_id]["custom"]["Data"])
        path = "A"
        if self.stop == True:
            self.stop = False
            return
        if "Data" not in graph["nodes"][node_id]["custom"]:
            for key, node in graph["nodes"].items():
                if node_id == key:
                    if "connections" in graph:
                        for connections in graph["connections"]:
                            if key == connections["out"][0] and path in connections["out"][1]:
                                try:
                                    print("new node no data")
                                    print(connections["in"][0])
                                    self.runNode(graph, connections["in"][0], thread_signals, direct, manual, folder)
                                except Exception as e:
                                    import traceback
                                    print(traceback.format_exc())
        else:
            x = graph["nodes"][node_id]["custom"]["Data"]
            print("loading no data")
            if isinstance(x, str):
                 x = json.loads(graph["nodes"][node_id]["custom"]["Data"])
                #pnt(y)s
            print(x)
            print("loaded")
            print(node_runs)

            if graph["nodes"][node_id]["name"] in node_runs:
                node_runs[graph["nodes"][node_id]["name"]] += 1
            else:
                node_runs[graph["nodes"][node_id]["name"]] = 1
            for key, value in x.items()  :
                if key == "image":
                    x["image"] = np.asarray(x["image"],dtype = "uint8")

            print("TYPE")
            print(x["type"])
            if x["type"] == "IfElse":
                self.thread_signals.notifications.emit("Delay", "Running if/else conditional.", "icon.png", None, "No")

                if graph["nodes"][node_id]["custom"]["operator"] == "equals":

                    if self.global_variables[graph["nodes"][node_id]["custom"]["Variables"]] == graph["nodes"][node_id]["custom"]["condition"]:
                        path = "A"
                    else:
                        path = "B"
                if graph["nodes"][node_id]["custom"]["operator"] == "great than":
                    if self.global_variables[graph["nodes"][node_id]["custom"]["Variables"]] > graph["nodes"][node_id]["custom"]["condition"]:
                        path = "A"
                    else:
                        path = "B"
                if graph["nodes"][node_id]["custom"]["operator"] == "regex match":
                    if re.match(graph["nodes"][node_id]["custom"]["condition"], self.global_variables[graph["nodes"][node_id]["custom"]["Variables"]]):
                        path = "A"
                    else:
                        path = "B"

                if graph["nodes"][node_id]["custom"]["operator"] == "less than":
                    if self.global_variables[graph["nodes"][node_id]["custom"]["Variables"]] < graph["nodes"][node_id]["custom"]["condition"]:
                        path = "A"
                    else:
                        path = "B"
                if graph["nodes"][node_id]["custom"]["operator"] == "includes":
                    print("includes")
                    print(self.global_variables[graph["nodes"][node_id]["custom"]["Variables"]])
                    print(graph["nodes"][node_id]["custom"]["condition"])

                    if graph["nodes"][node_id]["custom"]["condition"] in self.global_variables[graph["nodes"][node_id]["custom"]["Variables"]]:
                        path = "A"
                    else:
                        path = "B"
                    print(path)
                    print("PATH")
            if x["type"] == "python":
                #f = StringIO()
                #with redirect_stdout(f):
                self.thread_signals.notifications.emit("Delay", "Running python code.", "icon.png", None, "No")

                import subprocess
               # self.QMainWindow.showMinimized()
                subprocess.call("pkill -f chrome/chrome", shell=True)
                import site
                import sys
                packages = site.getusersitepackages()
                for package in packages:
                    sys.path.append(package)

                import subprocess

                def get_global_sitepackages():
                    try:
                        subprocess.check_output('python3.10 -m pip install yt-dlp', shell=True,universal_newlines=True)


                        subprocess.check_output(
                                'export LD_LIBRARY_PATH=/usr/local/lib:$LD_LIBRARY_PATH',
                                shell=True,
                                universal_newlines=True
                            )                                    
                        result = subprocess.check_output(['python3', '-c', 'from distutils.sysconfig import get_python_lib; print(get_python_lib())'], universal_newlines=True)
                        return result.strip()
                    except Exception as e:
                        print(f"Error: {e}")
                        print("please install python version 3.10 from the windows store")
                        return None

                global_sitepackages = get_global_sitepackages()
                if global_sitepackages:
                    import sys
                    sys.path.append(global_sitepackages)
                #bundled_dir = sys._MEIPASS
                #sys.path.append(bundled_dir)
                sys.path.append(resource_path(""))
                try:
                    if python_process != None:
                        python_process.terminate()

# Start a process to run the given code.
                    x["code"] = graph["nodes"][node_id]["custom"]["code"]
                    if "Flask" in x["code"] or "flask" in x["code"] or "Flask" in x["code"] or "flask" in x["code"]:
                        python_process = start_process(x, {'__name__': '__main__'})
                        #open chrome to 127.0.0.1:3000 using subprocess
                        #subprocess.call("start chrome
                        subprocess.call(chrome_location + " http://127.0.0.1:3000" )
                    else:

                        code = graph["nodes"][node_id]["custom"]["code"]



                        if "prompt" in graph["nodes"][node_id]["custom"] and len(graph["nodes"][node_id]["custom"]["prompt"]) > 2:  
                            headers = {
                                "Content-Type": "application/json",
                            }
                            image_node = self.get_node_by_id(node_id)
                            print(image_node)
            #  p            rint the prorties 
                            print("run GENERETAE PYTHON node")

                            print("widget")
                            prompt = ''
                            prompt = graph["nodes"][node_id]["custom"]["prompt"] + " " + code + "generate working python3.10 code only without quotes or a prefix and explanations around the code.  The code is run using the exec function directly with the following functions that are defined already so you don't need to implement them. Do not implement any of these functions: exec(code,{'resource_path': resource_path, 'moviepy': moviepy, 'thread_signals':thread_signals, 'notification_manager':self.notification_manager,'folder': folder, 'send_complex_message': self.send_complex_message, 'agent_memory': self.global_variables ,'store_analytics': self.store_analytics, 'user_key': self.user_key, 'imageToVideo': self.imageToVideo, 'genSyntheticVideo' : self.genSyntheticVideo, 'genImage': self.genImage, 'genVoice': self.genVoice, 'gpt3Prompt': self.gpt3Prompt, 'last_mouse_y':last_mouse_y,'last_mouse_x': last_mouse_x, '__name__': '__main__','videoSearch': self.videoSearch, 'semanticClick': self.semanticClick, 'keypress': self.keypress , 'semanticMove': self.semanticMove, 'click': self.click, 'move':self.move, 'getData': self.getData, 'sendData': self.sendData} )" 
                            log = [{"role": "system", "content": """You are a helpfull assistant that only generates working python code without placeholders, comments, quotes, prefixes that can be run directly. Use selenium with the local browser to use the local cookies for logged in services. Use this code to get the local chrome cookies: def get_chrome_data_dir(self): system = platform.system() if system == "Windows": return os.path.join(os.environ['LOCALAPPDATA'], 'Google', 'Chrome', 'User Data') elif system == "Darwin":  # macOS return os.path.expanduser('~/Library/Application Support/Google/Chrome') else:  # Linux return os.path.expanduser('~/.config/google-chrome') def setup_driver(self):try: options = webdriver.ChromeOptions() # Add the user data directory user_data_dir = self.get_chrome_data_dir() options.add_argument(f"user-data-dir={user_data_dir}") options.add_argument("--profile-directory=Default") # Add additional options options.add_argument("--start-maximized") options.add_argument("--disable-notifications") options.add_argument("--disable-popup-blocking") self.driver = webdriver.Chrome(options=options) print("Chrome driver initialized successfully") except Exception as e: print(f"Error initializing Chrome driver: {str(e)}") print("If Chrome is already running, please close all Chrome windows and try again.") raise """}]
                            print(prompt)
                            Automation_input = ""
                            Webhook_input = ""

                            log.append({"role": "user", "content": prompt + Automation_input + Webhook_input})
                            data = {
                                "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
                                "messages": log,
                                "temperature": 0.7,
                                "max_tokens": 5000,
                                "stream": True  # Enable streaming to match original behavior
                            }

                            print(data)

                            # Call local server instead of remote API
                            response = requests.post("http://localhost:8000/v1/chat/completions", 
                                                    headers=headers, 
                                                    data=json.dumps(data), 
                                                    stream=True)

                            import pyautogui
                            if response.status_code == 200:
                                items = []
                                total = ""
                                for chunk in response.iter_lines():
                                    try:
                                        if chunk:
                                            # Parse SSE format from llama.cpp server
                                            chunk_str = chunk.decode('utf-8')
                                            if chunk_str.startswith('data: '):
                                                chunk_data = chunk_str[6:]  # Remove 'data: ' prefix
                                                if chunk_data.strip() == '[DONE]':
                                                    break
                                                
                                                chunk_json = json.loads(chunk_data)
                                                if 'choices' in chunk_json and len(chunk_json['choices']) > 0:
                                                    delta = chunk_json['choices'][0].get('delta', {})
                                                    if 'content' in delta:
                                                        content = delta['content']
                                                        total += content
                                                        print(content, end='', flush=True)
                                    except Exception as e:
                                        # If streaming fails, fall back to non-streaming
                                        print(f"Streaming error: {e}")
                                        break
                                    
                                # If streaming didn't work or total is empty, try non-streaming
                                if not total:
                                    data["stream"] = False
                                    response = requests.post("http://localhost:8000/v1/chat/completions", 
                                                            headers=headers, 
                                                            data=json.dumps(data))
                                    if response.status_code == 200:
                                        result = response.json()
                                        total = result["choices"][0]["message"]["content"]
                                        print(total)
                                    else:
                                        print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                                        total = "Error communicating with local LLM server"

                                print()  # New line after streaming output
                                print(total)
                            else:
                                print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                                total = "Error communicating with local LLM server"
                                code = total
                            print(self.global_variables)
                            if node_id in graph["nodes"] and "custom" in graph["nodes"][node_id]:
                                for key, value in graph["nodes"][node_id]["custom"].items():
                                    if "{{" not in value:
                                        self.global_variables[key] = value
                            for key, value in self.global_variables.items():
                                if key in code:
                                    code = code.replace("{{" + key + "}}", "agent_memory['" + key + "']")
                            if "{{folder}}" in code:
                                code = code.replace("{{folder}}", "'" + folder + "'")

                        if "#!/bin/bash" in code:
                            
                            for code in graph["nodes"][node_id]["custom"]["code"].split("\n"):
                                if len(code) > 0:
                                    print("CODE")
                                    print(code)
                                    print("SSTART BASH CODE")
                                    subprocess.call(code, shell=True)
                                    #os.system(code)
                                    import time
                                    time.sleep(1)
                                    print("BASH OUTPUT")    

#                                    
                                    #subprocess.call("chmod +x code.sh", shell=True)
                                    #subprocess.call("./code.sh", shell=True)
     #                               os.system("./" + resource_path("code.sh"))
#                            output = os.system(x["code"])
                            print("BASH OUTPUT END")
                        #thread_signals.progress.emit(f.getvalue())
                        else:
                            import moviepy
                            exec(code,{"resource_path": resource_path, 'moviepy': moviepy, 'thread_signals':thread_signals, 'notification_manager':self.notification_manager,'folder': folder, 'send_complex_message': self.send_complex_message, 'agent_memory': self.global_variables ,'store_analytics': self.store_analytics, 'user_key': self.user_key, 'imageToVideo': self.imageToVideo, 'genSyntheticVideo' : self.genSyntheticVideo, 'genImage': self.genImage, 'genVoice': self.genVoice, 'gpt3Prompt': self.gpt3Prompt, 'last_mouse_y':last_mouse_y,'last_mouse_x': last_mouse_x, '__name__': '__main__','videoSearch': self.videoSearch, 'semanticClick': self.semanticClick, "keypress": self.keypress , "semanticMove": self.semanticMove, "click": self.click, "move":self.move, "getData": self.getData, "sendData": self.sendData} )
                       # try:
#                       #     self.notification_manager.add_notification("Analyzing Current State", "Analyzing the current state of the screen. Please wait..", "icon.png", lambda: print("Notification closed"))
                     #   thread_signals.notifications.emit("Analyzing Current State", "Analyzing the current state of the screen. Please wait..", "icon.png")
                      #  except Exception as e:
                      ##      print(e)
                        #    print("Error in code")
                except Exception as err:
                    import traceback
                    error_class = err.__class__.__name__
                    print(f"An error occurred: {err}")
                    if self.total_errors < 4 and manual == False:
                        thread_signals.chat.emit(f"A {error_class} error occurred.  Please re-generate the automation to fix it: {err}")
                        self.total_errors += 1
                

            if x["type"] == "bash":
                import subprocess
                #f = StringIO()
                self.thread_signals.notifications.emit("Delay", "Running bash code.", "icon.png", None, "No")
                import time
                #with redirect_stdout(f):
                print("Running bash code please wait")
                time.sleep(10)
                print("SPLIT CODES")
                print(x["code"].split("\n"))
                    #write the code to a file
                for code in graph["nodes"][node_id]["custom"]["code"].split("\n"):
                    if len(code) > 0:
                        print("CODE")
                        print(code)
                        print("SSTART BASH CODE")
                        subprocess.call(code, shell=True)
                        #os.system(code)
                        time.sleep(1)
                        print("BASH OUTPUT")    
                        
#                        
                        #subprocess.call("chmod +x code.sh", shell=True)
                        #subprocess.call("./code.sh", shell=True)
     #                   os.system("./" + resource_path("code.sh"))
#                output = os.system(x["code"])
                print("BASH OUTPUT END")
                #thread_signals.progress.emit(f.getvalue())

                #os.system(x["code"])

            if x["type"] == "Start Node":
                print("start node!")
                self.global_variables["Total Runs"] += 1
                init_program = graph["nodes"][node_id]["custom"]["Initial Program"]
                self.global_variables["user"] = self.user_email
                print(init_program)
                print(self.global_variables)
                import subprocess
               # self.QMainWindow.showMinimized()
                subprocess.call("pkill -f chrome/chrome", shell=True)

                #replace any strings that include {{name}} with self.global_variables["name"] in body 
                for key, value in self.global_variables.items():
                    print(self.global_variables[key])
                    print(key)
                    if key in init_program:
                        init_program = init_program.replace("{{" + key + "}}", str(value))
                
                if len(init_program) > 0 :
                    try:

                      #  self.thread_signals.notifications.emit("Delay", "Opening initial program.", "icon.png", None, "No")

                        # Open the file based on the extension
                        init_program =  init_program
                        print(init_program.replace("\\n","").split(" "))
                        if "google-chrome" in init_program:
                            subprocess.Popen( init_program.replace("\\n","").split(" "))
                        else:
                            subprocess.Popen(init_program.replace("\\n","").split(" "))
                        import time
                        time.sleep(5)
                    except FileNotFoundError:
                         print("That Start Node program didn't work, try again")
            if x["type"] == "getfiles":
                
                #check if URL is a file or URL and download it if it's a URL or load the file from a pdf, doc, text, or csv
                if True:
                    #open the file and read the contents
                    #check if the file is a word doc, then load it appropriately
                    if ".docx" in graph["nodes"][node_id]["custom"]["URL or File"]:
                        import docx
                        doc = docx.Document(graph["nodes"][node_id]["custom"]["URL or File"])
                        data = ""
                        for para in doc.paragraphs:
                            data += para.text
                        self.global_variables["Download " + graph["nodes"][node_id]["name"].split("Download ")[1]] = data
                    elif "pdf" in graph["nodes"][node_id]["custom"]["URL or File"]:
                        import PyPDF2
                        pdfFileObj = open(graph["nodes"][node_id]["custom"]["URL or File"], 'rb')
                        pdfReader = PyPDF2.PdfReader(pdfFileObj)
                        data = ""
                        for pageObj in pdfReader.pages:
                            data += pageObj.extract_text()
                        self.global_variables["Download " + graph["nodes"][node_id]["name"].split("Download ")[1]] = data

                    elif ".csv" in graph["nodes"][node_id]["custom"]["URL or File"]:
                        import csv
                        with open(graph["nodes"][node_id]["custom"]["URL or File"], newline='') as csvfile:
                            data = ""
                            spamreader = csv.reader(csvfile, delimiter=' ', quotechar='|')
                            for row in spamreader:
                                data += ', '.join(row)
                        self.global_variables["Download " + graph["nodes"][node_id]["name"].split("Download ")[1]] = data

                    elif ".txt" in graph["nodes"][node_id]["custom"]["URL or File"]:
                        with open(graph["nodes"][node_id]["custom"]["URL or File"], "r") as file:
                            data = file.read()
                            self.global_variables["Download " + graph["nodes"][node_id]["name"].split("Download ")[1]] = data

                    elif "xlsx" in graph["nodes"][node_id]["custom"]["URL or File"]:
                        import openpyxl
                        wb = openpyxl.load_workbook(graph["nodes"][node_id]["custom"]["URL or File"])
                        sheet = wb.active
                        data = ""
                        for row in sheet.iter_rows(values_only=True):
                            data += str(row)
                        self.global_variables["Download " + graph["nodes"][node_id]["name"].split("Download ")[1]] = data
                    else:
                        with open(graph["nodes"][node_id]["custom"]["URL or File"], "r") as file:
                            data = file.read()
                            self.global_variables["Download " + graph["nodes"][node_id]["name"].split("Download ")[1]] = data
                elif "docs.google" in graph["nodes"][node_id]["custom"]["URL or File"]:
                    print("google sheets")
                    data = self.sendMessageRTCAsync('google_sheets_read:' + graph["nodes"][node_id]["custom"]["URL or File"])
                    print("get data output")
                    print(data)

                    self.global_variables["Download " + graph["nodes"][node_id]["name"].split("Download ")[1]] = data
                elif len(graph["nodes"][node_id]["custom"]["Headers"]) != 0:
                    response = requests.get(url=graph["nodes"][node_id]["custom"]["URL or File"],headers=json.loads(graph["nodes"][node_id]["custom"]["Headers"]))
                    self.global_variables["Download " + graph["nodes"][node_id]["name"].split("Download ")[1]] = response.text
                    print(self.global_variables)
                else:
                    response = requests.get(graph["nodes"][node_id]["custom"]["URL or File"])
                    open(graph["nodes"][node_id]["custom"]["Download Location"], "wb").write(response.content)
                    #store the download data in the global variables
                    self.global_variables["Download " + graph["nodes"][node_id]["name"].split("Download ")[1]] = response.content

                print(self.global_variables)
            if x["type"] == "CustomCode":

                print("SEND Cheat Layer")
                print(graph["nodes"][node_id]["custom"]["URL"])
                print(graph["nodes"][node_id]["custom"]["code"])
                r = requests.post("https://cheatlayer.com/triggers/extension", data={"key":graph["nodes"][node_id]["custom"]["key"]  ,"script":graph["nodes"][node_id]["custom"]["code"],"start":graph["nodes"][node_id]["custom"]["URL"]})
            if x["type"] == "Email":
                var = ""
                self.thread_signals.notifications.emit("Delay", "Sent Email", "icon.png", None, "No")

                if  len(graph["nodes"][node_id]["custom"]["Body variable"]) > 0:
                     var = self.global_variables[graph["nodes"][node_id]["custom"]["Body variable"]]
                body = graph["nodes"][node_id]["custom"]["body"]
                to = graph["nodes"][node_id]["custom"]["to"]
                if "to variable" in graph["nodes"][node_id]["custom"] and len(graph["nodes"][node_id]["custom"]["to variable"]) > 0:
                    to = self.global_variables[graph["nodes"][node_id]["custom"]["to variable"]].replace("\\n","")
                subject = graph["nodes"][node_id]["custom"]["subject"]
                
                if  len(graph["nodes"][node_id]["custom"]["body"]) > 0:
                    #replace any strings that include {{name}} with self.global_variables["name"] in body 
                    for key, value in self.global_variables.items():
                        print(self.global_variables[key])
                        print(key)
                        if key in body:
                            body = body.replace("{{" + key + "}}", str(value))
                if len(graph["nodes"][node_id]["custom"]["subject"]) > 0:
                    for key, value in self.global_variables.items():
                        if key in subject:
                            subject = subject.replace("{{" + key + "}}", str(value))
                if len(graph["nodes"][node_id]["custom"]["to"]) > 0:
                    for key, value in self.global_variables.items():
                        if key in to:
                            to = to.replace("{{" + key + "}}", str(value))
                print("SEND Cheat Layer")
                
                # Replace 'your_api_key' with your actual Mailgun API key

                # Replace 'your_mailgun_domain' with your actual Mailgun domain

                # Mailgun API URL
                api_url = f'https://cheatlayer.com/user/sendAgentEmail'

                # Email details
                from_email = 'agents@cheatlayer.com'  # Replace with your sending email address
                to_email = to  # Replace with the recipient's email address
                subject = subject
                text = body
                # Path to the attachment file
                attachment_file_path = resource_path( os.path.join(folder, graph["nodes"][node_id]["custom"]["file"]))
                # add folder to the file 
                # Prepare the data for the API request
                data = {
                    'from': from_email,
                    'to': to_email,
                    'subject': subject,
                    'body': text + str(var),
                    'security':'af3j2kdw234'
                }

                # Prepare the attachment file
                #check the file exists and is not none
                files = {}

                if os.path.exists(attachment_file_path) and graph["nodes"][node_id]["custom"]["file"] != "" and graph["nodes"][node_id]["custom"]["file"] != "none":
                    files = {
                        'attachment': (graph["nodes"][node_id]["custom"]["file"], open(attachment_file_path, 'rb').read()),
                    }
                print("SEND Cheat Layer")
                print(data)
                print(files)

                # Send the email using the Mailgun API
                response = requests.post(
                    api_url,
                    data=data
                )

                # Check the response from Mailgun
                if response.status_code == 200:
                    print('Email sent successfully!')
                else:
                    print(f'Failed to send email: {response.status_code} {response.text}')
               # r = requests.post("https://cheatlayer.com/triggers/extension", data={"key":graph["nodes"][node_id]["custom"]["key"]  ,"script":"sendGmail(`" + to + "`,`" + subject + "`,`" + body + var + "`)","start":""})

            if x["type"] == "Riku":
                data = { "Name": graph["nodes"][node_id]["custom"]["Account Holder"] , "Secret": graph["nodes"][node_id]["custom"]["API Key"] , "Prompt ID": graph["nodes"][node_id]["custom"]["Prompt ID"] , "n": 1,"Input 1": graph["nodes"][node_id]["custom"]["Input 1"] ,"Input 2": graph["nodes"][node_id]["custom"]["Input 2"],"Input 3": graph["nodes"][node_id]["custom"]["Input 3"],"Input 4": graph["nodes"][node_id]["custom"]["Input 4"],"Input 5": graph["nodes"][node_id]["custom"]["Input 5"]}
                r = requests.post("https://prompts.riku.ai/webhook/run", data=data)
                if "Outputs" in r.json()[0]:
                    self.global_variables["riku_" + str(len(self.global_variables))] = (r.json()[0]["Outputs"][0]["Generated"])
            if x["type"] == "Math":
                if graph["nodes"][node_id]["custom"]["action"] == "divide":
                    self.global_variables["math_" + str(len(self.global_variables))] = (graph["nodes"][node_id]["custom"]["input"]/graph["nodes"][node_id]["custom"]["value"])

                if graph["nodes"][node_id]["custom"]["action"] == "multiply":
                    self.global_variables["math_" + str(len(self.global_variables))] = (graph["nodes"][node_id]["custom"]["input"]*graph["nodes"][node_id]["custom"]["value"])

                if graph["nodes"][node_id]["custom"]["action"] == "add":
                    self.global_variables["math_" + str(len(self.global_variables))] = (graph["nodes"][node_id]["custom"]["input"]+graph["nodes"][node_id]["custom"]["value"])

                if graph["nodes"][node_id]["custom"]["action"] == "subtract":
                    self.global_variables["math_" + str(len(self.global_variables))] = (graph["nodes"][node_id]["custom"]["input"]-graph["nodes"][node_id]["custom"]["value"])

            if x["type"] == "Open":
                if "automated mode" in graph["nodes"][node_id] and graph["nodes"][node_id]["custom"]["automated mode"] == "Loop Total Runs":
                    runs = int(self.global_variables["Total Runs"])
                    self.global_variables["Total Runs"] += 1
                    URL = self.global_variables[graph["nodes"][node_id]["custom"]["automated URL"]][runs%len(self.global_variables[graph["nodes"][node_id]["custom"]["automated URL"]])]
                    import subprocess
                    #replace any strings that include {{name}} with self.global_variables["name"] in body 
                    for key, value in self.global_variables.items():
                        print(self.global_variables[key])
                        print(key)
                        if key in URL:
                            URL = URL.replace("{{" + key + "}}", str(value))
                    try:
                        subprocess.Popen(["google-chrome", URL])
                    except:
                        pass
                    self.thread_signals.notifications.emit("Delay", "Opening " +"google-chrome " + URL, "icon.png", None, "No")

 #                   webbrowser.open(URL)
                elif "Automated Input" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Automated Input"] != '':
                    import subprocess
                    print(self.global_variables[graph["nodes"][node_id]["custom"]["Automated Input"]])
                    if type(self.global_variables[graph["nodes"][node_id]["custom"]["Automated Input"]]) == list:
                        list_progs = self.global_variables[graph["nodes"][node_id]["custom"]["Automated Input"]]
                        runs = int(self.global_variables["Total Runs"])
                        print(list_progs)
                        print(runs)
                        self.global_variables["Total Runs"] += 1
                        program = graph["nodes"][node_id]["custom"]["program"].replace("\\n","") + " " + list_progs[runs%len(list_progs)].replace("\\n","")
                        program =  program
                        try:
                            if "google-chrome" in program:

                                subprocess.Popen( program.replace("\\n","").split(" ") )
                            else:
                                subprocess.Popen(program.replace("\\n","").split(" "))
                                subprocess.Popen(program.split(" "))
                        except:
                            pass
                    else:

                        # Get the file path from your graph
                        file_path = graph["nodes"][node_id]["custom"]["program"].replace("\\n","") + self.global_variables[graph["nodes"][node_id]["custom"]["Automated Input"]].replace("\\n","")
                        
                        # Get the file extension
                        file_extension = os.path.splitext(file_path)[1]
                        print(file_extension)
                        # Open the file based on the extension
                        try:
                            subprocess.Popen(file_path.replace("\\n","").split(" "))
                        except:
                            pass                    
                elif "Automated Lists" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Automated Lists"] != '':
                    import subprocess
                    list_progs = self.global_variables[graph["nodes"][node_id]["custom"]["Automated Lists"]]
                    runs = int(self.global_variables["Total Runs"])
                    print(list_progs)
                    print(runs)
                    self.global_variables["Total Runs"] += 1
                    program = graph["nodes"][node_id]["custom"]["program"] + " " + list_progs[runs%len(list_progs)].replace("\\n","")
                    program =  program 
                    try:
                        
                            if "google-chrome" in program:

                                subprocess.Popen( program.replace("\\n","").split(" ") )
                            else:
                                subprocess.Popen(program.replace("\\n","").split(" "))
                                subprocess.Popen(program.split(" "))
                    except:
                        pass
                else:
                    program = graph["nodes"][node_id]["custom"]["program"] 
                    import subprocess
                    #replace any strings that include {{name}} with self.global_variables["name"] in body 
                    for key, value in self.global_variables.items():
                        print(self.global_variables[key])
                        print(key)
                        if key in program:
                            program = program.replace("{{" + key + "}}", str(value))

                # Open the file based on the extension
                    program =  program
                    try:
                        
                            if "google-chrome" in program:

                                subprocess.Popen( program.replace("\\n","").split(" ") )
                            else:
                                subprocess.Popen(program.replace("\\n","").split(" "))
                                subprocess.Popen(program.split(" "))
                    except:
                        pass                               
            if x["type"] == "Gsheets":
                
                if "Read/Write" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Read/Write"] == "Read":
                    print("google sheets")
                    data = self.sendMessageRTCAsync('google_sheets_read:' + graph["nodes"][node_id]["custom"]["URL"] + ":;:" + graph['nodes'][node_id]['custom']["sheet name"])
                    print("get data output")
                    print(data)

                    self.global_variables["Gsheets"  +  graph["nodes"][node_id]["name"].split("Google Sheets ")[1]] = data
                elif "Read/Write" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Read/Write"] == "Read Rows":
                    print("google sheets")
                    data = self.sendMessageRTCAsync('google_sheets_read:' + graph["nodes"][node_id]["custom"]["URL"] + ":;:" + graph['nodes'][node_id]['custom']["sheet name"])
                    print("get data output")
                    print(data)

                    self.global_variables["Gsheets"  +  graph["nodes"][node_id]["name"].split("Google Sheets ")[1]] = data
                elif "Cell Row" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Cell Row"] != "none":


                    print(self.global_variables)
                    self.thread_signals.notifications.emit("Delay", "Sent data to Google Sheets", "icon.png", None, "No")
                    print(self.user_key)
                    print("gsheets key")
                    r = requests.post("https://cheatlayer.com/triggers/extension", data={
                        "key": self.user_key,
                        "script":json.dumps( {"type": "google_sheets_update_cell","url": graph['nodes'][node_id]['custom']["URL"], "name": graph['nodes'][node_id]['custom']["sheet name"] , "data": [self.global_variables.get(graph['nodes'][node_id]['custom'].get('Row 1', ''), '') ,  self.global_variables.get(graph['nodes'][node_id]['custom'].get('Row 2', ''), '') ,  self.global_variables.get(graph['nodes'][node_id]['custom'].get('Row 3', ''), ''),  self.global_variables.get(graph['nodes'][node_id]['custom'].get('Row 4', ''), '') , self.global_variables.get(graph['nodes'][node_id]['custom'].get('Row 5', ''), ''), self.global_variables.get(graph['nodes'][node_id]['custom'].get('Row 6', ''), ''), self.global_variables.get(graph['nodes'][node_id]['custom'].get('Row 7', ''), ''), self.global_variables.get(graph['nodes'][node_id]['custom'].get('Row 8', ''), ''), self.global_variables.get(graph['nodes'][node_id]['custom'].get('Row 9', ''), ''), self.global_variables.get(graph['nodes'][node_id]['custom'].get('Row 10', ''), '')]}),
                        "start": ""
                    })
                    print('SEND DATA TO GOOGLE SHEETS')

                else:
                    print(self.global_variables)
                    self.thread_signals.notifications.emit("Delay", "Sent data to Google Sheets", "icon.png", None, "No")
                    print(self.user_key)
                    print("gsheets key")
                    r = requests.post("https://cheatlayer.com/triggers/extension", data={
                        "key": self.user_key,
                        "script":json.dumps( {"type": "google_sheets_add_row","url": graph['nodes'][node_id]['custom']["URL"], "name": graph['nodes'][node_id]['custom']["sheet name"] , "data": [self.global_variables.get(graph['nodes'][node_id]['custom'].get('Row 1', ''), '') ,  self.global_variables.get(graph['nodes'][node_id]['custom'].get('Row 2', ''), '') ,  self.global_variables.get(graph['nodes'][node_id]['custom'].get('Row 3', ''), ''),  self.global_variables.get(graph['nodes'][node_id]['custom'].get('Row 4', ''), '') , self.global_variables.get(graph['nodes'][node_id]['custom'].get('Row 5', ''), ''), self.global_variables.get(graph['nodes'][node_id]['custom'].get('Row 6', ''), ''), self.global_variables.get(graph['nodes'][node_id]['custom'].get('Row 7', ''), ''), self.global_variables.get(graph['nodes'][node_id]['custom'].get('Row 8', ''), ''), self.global_variables.get(graph['nodes'][node_id]['custom'].get('Row 9', ''), ''), self.global_variables.get(graph['nodes'][node_id]['custom'].get('Row 10', ''), '')]}),
                        "start": ""
                    })
                    print('SEND DATA TO GOOGLE SHEETS')

            if x["type"] == "Delay":
                self.thread_signals.notifications.emit("Delay", "Delay for " + graph["nodes"][node_id]["custom"]["seconds"] + " seconds" , "icon.png", None, "No")
                import time
                time.sleep(int(graph["nodes"][node_id]["custom"]["seconds"]))

            if x["type"] == "StableDiffusion":
                import threading

                prompt = ''
                if "prompt" in graph["nodes"][node_id]["custom"]:
                    prompt = graph["nodes"][node_id]["custom"]["prompt"]
                if "Automation Input" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Automation Input"] != '':
                    prompt += self.global_variables[graph["nodes"][node_id]["custom"]["Automation Input"]]
                self.StableDiffusion(prompt)
                image_node = self.get_node_by_id(node_id)
                image_node.updateImage(resource_path("generated.png"))
                  #  graph["nodes"][node_id].updateImage("generated.png")
#                    t0 = threading.Thread(target = self.StableDiffusion, kwargs={"opt_in":self.opt})
#                    t0.start()

            if x["type"] == "AIDetector":
                f = "cheat_model.pth"
                time.sleep(5)
                image = pyautogui.screenshot()
                import numpy as np
                image = np.array(image)
                print(graph["nodes"][node_id]["custom"]["labels"])
                print(graph["nodes"][node_id]["custom"]["UImodel"])
                model = core.Model.load(graph["nodes"][node_id]["custom"]["UImodel"], graph["nodes"][node_id]["custom"]["labels"].split(","))
                #image = utils.read_image(self.test_image)
                predictions = model.predict(image)
                labels, boxes, scores = predictions

                print(labels)
                print(boxes)
                print(scores)
                if boxes.ndim == 1:
                    boxes = boxes.view(1, 4)

                # Plot each box

                if boxes.ndim == 1:
                    boxes = boxes.view(1, 4)
                thresh=float(graph["nodes"][node_id]["custom"]["threshold"])
                print(scores)
                filtered_indices=np.where(scores>thresh)
                filtered_scores=scores[filtered_indices]
                filtered_boxes=boxes[filtered_indices]
                num_list = filtered_indices[0].tolist()
                filtered_labels = [labels[i] for i in num_list]
               # show_labeled_image(image, filtered_boxes, filtered_labels)
                # Plot each box

                print(filtered_boxes)
                print(filtered_labels)
                indices = self.nms(filtered_boxes, filtered_scores, filtered_labels, .05)
                boxes = filtered_boxes[indices]

                for i in range(boxes.shape[0]):
                    box = boxes[i]
                    width, height = (box[2] - box[0]).item(), (box[3] - box[1]).item()
                    initial_pos = (int(box[0].item()), int(box[1].item()))
                    end_pos = (int(box[2].item()), int(box[3].item()))
                    cv2.rectangle(image,initial_pos, end_pos, (0, 255, 0), 2)

                screenshot = cv2.cvtColor(np.array(image), cv2.COLOR_BGR2RGB)
                cv2.imshow("CLICK", screenshot)
                cv2.setWindowProperty("CLICK", cv2.WND_PROP_TOPMOST, 1)
                cv2.waitKey(1)

                for i in range(boxes.shape[0]):
                    box = boxes[i]
                    width, height = (box[2] - box[0]).item(), (box[3] - box[1]).item()
                    initial_pos = (int(box[0].item()), int(box[1].item()))
                    end_pos = (int(box[2].item()), int(box[3].item()))
                    cv2.rectangle(image,initial_pos, end_pos, (0, 255, 0), 2)
                    if graph["nodes"][node_id]["custom"]["selector"] == "OCR":
                        thresh = 255 - cv2.threshold(cv2.cvtColor(image, cv2.COLOR_BGR2GRAY), 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]

                      #  xp,yp,w,h = x["bounding_box"][0][0], x["bounding_box"][0][1], x["bounding_box"][1][0]-x["bounding_box"][0][0], x["bounding_box"][1][1]-x["bounding_box"][0][1]
                        ROI = thresh[int(box[1].item()):int(box[3].item()),int(box[0].item()):int(box[2].item())]
                        data = pytesseract.image_to_string(ROI, lang='eng',config='--psm 6')
                        print(data)
                        r = requests.post("https://cheatlayer.com/triggers/extension", data={"key":graph["nodes"][node_id]["custom"]["key"]  ,"script":"google_sheets_add_row('" +graph["nodes"][node_id]["custom"]["Google Sheet Output"] + "', 'Sheet1',[`" + data + "`])","start":""})
                        if graph["nodes"][node_id]["custom"]["value"] in data:
                            if graph["nodes"][node_id]["custom"]["action"] == "click":
                                pyautogui.click(initial_pos[0] + width/2, initial_pos[1] + height/2)
                            print("Found AI  DETECTOR")
                height, width, channel = image.shape
                bytesPerLine = 3 * width

            if x["type"] == "ElevenLabs":
                        
                headers = {
                    "Content-Type": "application/json",
                }
                image_node = self.get_node_by_id(node_id)
                print(image_node)
            #  print the prorties 
                print("properties")

                print("widget")

                print(image_node.widgets())
                #widgets is an ord
                #help me find specific data in image_node by listing the object properties
                print("widget")

                print(dir(image_node))

                print(graph["nodes"][node_id])
                print(graph["nodes"][node_id]["custom"])
                prompt = graph["nodes"][node_id]["custom"]["prompt"]
                print(self.global_variables)
                self.thread_signals.notifications.emit("ElevenLabs", "Generating Eleven Labs voice", "icon.png", None, "No")

                if "{{" in prompt:
                    #replace any strings that include {{name}} with self.global_variables["name"]
                    for key, value in self.global_variables.items():
                        if key in prompt:
                            prompt = prompt.replace("{{" + key + "}}", self.global_variables[key])

                log = [{"role": "system", "content": "You are a helpfull assistant that only generates exactly the text the user requests. Do not explain the text. SImply generate exactly what the user requests. For example, if a user asks to generate a social post about ai automation, you would write 'Our business keeps growing faster thanks to AI automation #AI #aiautomation'"}]

                Automation_input = ""
                Webhook_input = ""
                if "Automation Input" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Automation Input"] != "None" and len(graph["nodes"][node_id]["custom"]["Automation Input"]) > 0:
                    Automation_input = self.global_variables[graph["nodes"][node_id]["custom"]["Automation Input"]]
                
                url = "https://cheatlayer.com/user/genVoiceUnauth"
                headers = {
                    "Content-Type": "application/json"
                }
                data = json.dumps({"input": prompt + Automation_input + Webhook_input, 'id': self.user_key})
        
                response = requests.post(url, headers=headers, data=data)
                filename = graph["nodes"][node_id]["custom"]["File Name"]
                if response.status_code == 200:
                    with open(resource_path(filename), 'wb') as f:
                        f.write(response.content)
                    print(f"Audio saved as '{filename}'")
                else:
                    print("Error:", response.status_code, response.text)
                if graph["nodes"][node_id]["custom"]["Play Audio"] == "true":
                    #play the audio using subprocess
                    import subprocess
                    subprocess.call(resource_path(filename), shell=True)
#                    for step in total.split("[start"):
                
            if x["type"] == "ElevenLabs":
                        
                headers = {
                    "Content-Type": "application/json",
                }
                image_node = self.get_node_by_id(node_id)
                print(image_node)
            #  print the prorties 
                print("properties")

                print("widget")

                #help me find specific data in image_node by listing the object properties
                print("widget")

                print(dir(image_node))
                self.thread_signals.notifications.emit("ElevenLabs", "Generating Eleven Labs voice", "icon.png", None, "No")

                print(graph["nodes"][node_id])
                print(graph["nodes"][node_id]["custom"])
                prompt = first_widget.get_value()
                print(self.global_variables)
                if "{{" in prompt:
                    #replace any strings that include {{name}} with self.global_variables["name"]
                    for key, value in self.global_variables.items():
                        if key in prompt:
                            prompt = prompt.replace("{{" + key + "}}", self.global_variables[key])

                log = [{"role": "system", "content": "You are a helpfull assistant that only generates exactly the text the user requests. Do not explain the text. SImply generate exactly what the user requests. For example, if a user asks to generate a social post about ai automation, you would write 'Our business keeps growing faster thanks to AI automation #AI #aiautomation'"}]

                Automation_input = ""
                Webhook_input = ""
                if "Automation Input" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Automation Input"] != "None" and len(graph["nodes"][node_id]["custom"]["Automation Input"]) > 0:
                    Automation_input = self.global_variables[graph["nodes"][node_id]["custom"]["Automation Input"]]
                
                url = "https://cheatlayer.com/user/genVoiceUnauth"
                headers = {
                    "Content-Type": "application/json"
                }
                data = json.dumps({"input": prompt + Automation_input + Webhook_input, 'id': self.user_key})
        
                response = requests.post(url, headers=headers, data=data)
                filename = graph["nodes"][node_id]["custom"]["File Name"]
                if response.status_code == 200:
                    with open(resource_path(filename), 'wb') as f:
                        f.write(response.content)
                    print(f"Audio saved as '{filename}'")
                else:
                    print("Error:", response.status_code, response.text)
                if graph["nodes"][node_id]["custom"]["Play Audio"] == "true":
                    #play the audio using subprocess
                    import subprocess
                    subprocess.call(resource_path(filename), shell=True)
#                    for step in total.split("[start"):
                
            if x["type"] == "Llama":
                from llama_cpp import Llama

                # Define the file name and the URL to download the file if it does not exist
                file_name = "llama-2-7b-chat.Q5_K_S.gguf"
                file_url = "https://storage.googleapis.com/cheatlayer/llama-2-7b-chat.Q5_K_S.gguf"

                def download_file(url, path):
                    """Download a file from a specified URL to a specified path"""
                    response = requests.get(url)
                    if response.status_code == 200:  # HTTP Status OK
                        with open(path, "wb") as file:
                            file.write(response.content)
                        print(f"File '{path}' downloaded successfully.")
                    else:
                        print(f"Failed to download file. HTTP Status Code: {response.status_code}")

                # Check if the file exists
                if not os.path.exists(file_name):
                    print(f"File '{file_name}' not found. Starting download...")
                    download_file(file_url, file_name)
                else:
                    print(f"File '{file_name}' already exists.")
                headers = {
                    "Content-Type": "application/json",
                }
                image_node = self.get_node_by_id(node_id)
                print(image_node)
            #  print the prorties 
                print("properties")

                print("widget")

                print(graph["nodes"][node_id])
                print(graph["nodes"][node_id]["custom"])
                prompt = first_widget.get_value()
                print(self.global_variables)
                if "{{" in prompt:
                    #replace any strings that include {{name}} with self.global_variables["name"]
                    for key, value in self.global_variables.items():
                        if key in prompt:
                            prompt = prompt.replace("{{" + key + "}}", self.global_variables[key])

                log = [{"role": "system", "content": "You are a helpfull assistant that only generates exactly the text the user requests. Do not explain the text. SImply generate exactly what the user requests. For example, if a user asks to generate a social post about ai automation, you would write 'Our business keeps growing faster thanks to AI automation #AI #aiautomation'"}]

                Automation_input = ""
                Webhook_input = ""
                if "Automation Input" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Automation Input"] != "None" and len(graph["nodes"][node_id]["custom"]["Automation Input"]) > 0:
                    Automation_input = self.global_variables[graph["nodes"][node_id]["custom"]["Automation Input"]]

                log.append({"role": "user", "content": prompt + Automation_input + Webhook_input})

                data = {
                    "input": log,  # Use the correct fallback value for client.exampleInput
                    "max_tokens": 5000,
                    "id": self.user_key,
                    "key": "",
                    "plan": self.user_plan,
                    "mode":"free"
                }
                print(data)

                llm = Llama(model_path="llama-2-7b-chat.Q5_K_S.gguf",chat_format="llama-2")
                completion = llm.create_chat_completion(messages = log )
                total = completion["choices"][0]["message"]["content"]
                print(total)
                if "Tracking Tag" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Tracking Tag"] != "None":
                    total += " #244" + self.global_variables["key"][:5]
                self.global_variables["Llama"  + graph["nodes"][node_id]["name"].split("Llama ")[1]] = total
                if "type output" in graph["nodes"][node_id]["custom"]:
                    if graph["nodes"][node_id]["custom"]["type output"] == "true":
                        pyautogui.write(total, interval=0.05)
                else:
                    pyautogui.write(total, interval=0.05)
                print(self.global_variables)
                self.thread_signals.notifications.emit("Llama", total, "icon.png", None, "No")

                  # first_widget._label = total
                  # first_widget.setTitle(total)
                
                 #   image_node.set_label(total)
                print("Llama output")
#                    for step in total.split("[start"):

            # Find and modify any GPT-4 or language model sections to use the local server
            if x["type"] == "GPT4":
                headers = {
                    "Content-Type": "application/json",
                }
                
                prompt = ''
                prompt = graph["nodes"][node_id]["custom"]["input"] 
                
                # Process any variables
                if "{{" in prompt:
                    for key, value in self.global_variables.items():
                        if key in prompt:
                            prompt = prompt.replace("{{" + key + "}}", str(value))[:500]
                
                # Get any additional inputs
                Automation_input = ""
                if "Automation Input" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Automation Input"] != "None" and len(graph["nodes"][node_id]["custom"]["Automation Input"]) > 0:
                    Automation_input = str(self.global_variables[graph["nodes"][node_id]["custom"]["Automation Input"]])
                
                # Call our local function instead of making the remote API request
                system_prompt = "You are a helpful assistant that only generates exactly the text the user requests. Do not explain the text. Simply generate exactly what the user requests."
                total = self.gpt3Prompt(system_prompt, prompt + Automation_input)
                
                # Process the results the same way
                if "Tracking Tag" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Tracking Tag"] != "None":
                    total += " #244" + self.global_variables["key"][:5]
                    
                self.global_variables["GPT4" + graph["nodes"][node_id]["name"].split("GPT4 ")[1]] = total
                
                if "type output" in graph["nodes"][node_id]["custom"]:
                    if graph["nodes"][node_id]["custom"]["type output"] == "true":
                        pyautogui.write(total, interval=0.05)
                else:
                    pyautogui.write(total, interval=0.05)
                    
                self.thread_signals.notifications.emit("GPT-4", total, "icon.png", None, "No")
#                    for step in total.split("[start"):
            if x["type"] == "webcam":
                print("webcam")
                cap = cv2.VideoCapture(0)
                ret, frame = cap.read()
                cv2.imwrite(resource_path("webcam_output.jpg"), frame)
                cap.release()
                
                image_node = self.get_node_by_id(node_id)
                image_node.updateImage(resource_path("webcam_output.jpg"))
                print("webcam")
                #load the image as an np array
                
                encoded_string = ""
                with open(resource_path("webcam_output.jpg"), "rb") as image_file:
                    encoded_string = base64.b64encode(image_file.read()).decode('utf-8')
                caption = blip_caption_describe(encoded_string, graph["nodes"][node_id]["custom"]["prompt"], self.user_key, self.user_plan)
                print(caption)
                self.global_variables["Webcam"  +  graph["nodes"][node_id]["name"].split("Webcam ")[1]] = caption
                print(self.global_variables)
                self.thread_signals.notifications.emit("Webcam Describe", caption, "icon.png", None, "No")

            if x["type"] == "keypress_manual":
                print("keypress")
                print(graph["nodes"][node_id]["custom"]["String"])
                x["key"] = graph["nodes"][node_id]["custom"]["String"]
                x["saved"] = graph["nodes"][node_id]["custom"]["Saved Values"]
                self.thread_signals.notifications.emit("Keypress", x["key"], "icon.png", None, "No")
                print("keypress run")
                import pyautogui
                if True:
                    if "Current Directory" in str(x["saved"]):
                        pyautogui.write(resource_path(''), interval=0.05)
                    if "resource_path" in str(x["key"]).lower():
                        pyautogui.write(resource_path(''), interval=0.05)
                    if "backspace" in str(x["key"]).lower():
                        pyautogui.press('backspace')
                    elif "space" in str(x["key"]).lower():
                        pyautogui.write(" ", interval=0.05)
                    elif "period" in str(x["key"]).lower():
                        pyautogui.write(".", interval=0.05)
                    elif "slash" in str(x["key"]).lower():
                        pyautogui.write("/", interval=0.05)
                    elif "return" in str(x["key"]).lower():
                        pyautogui.press('enter')
                    elif "shift" in str(x["key"]).lower():
                        pyautogui.keyDown("shift")
                        self.shift_down = True
                    elif "ctrl" in str(x["key"]).lower():
                        pyautogui.keyDown("ctrl")
                        self.ctrl_down = True
                    elif "alt" in str(x["key"]).lower():
                        pyautogui.keyDown("alt")
                        self.alt_down = True
                    elif "tab" in str(x["key"]).lower():
                        pyautogui.keyDown("tab")
                    elif "enter" in str(x["key"]).lower():
                        pyautogui.press('enter')
                    elif "esc" in str(x["key"]).lower():
                        pyautogui.press('escape')


                    else:
                        if self.shift_down:
                            if "{{folder}}" in str(x["key"]):
                                pyautogui.write(folder, interval=0.01)
                                pyautogui.write("\\", interval=0.01)
                                pyautogui.write(str(x["key"]).upper().replace("{{FOLDER}}", ""), interval=0.01)
                                
                            else:
                                if "GPT-4 Mode" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["GPT-4 Mode"] == "True":
                                    print("GPT-4 Mode")
                                    prompt = x["key"]
                                    log = [{"role": "system", "content": "You are a helpfull assistant that only generates exactly the text the user requests. Do not explain the text. SImply generate exactly what the user requests. For example, if a user asks to generate a social post about ai automation, you would write 'Our business keeps growing faster thanks to AI automation #AI #aiautomation'"}]

                                    Automation_input = ""
                                    Webhook_input = ""
                                    if "Automation Input" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Automation Input"] != "None" and len(graph["nodes"][node_id]["custom"]["Automation Input"]) > 0:
                                        Automation_input = self.global_variables[graph["nodes"][node_id]["custom"]["Automation Input"]]

                                    log.append({"role": "user", "content": prompt.replace("space", " ") + Automation_input + Webhook_input})
                                    data = {
                                        "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
                                        "messages": log,
                                        "temperature": 0.7,
                                        "max_tokens": 5000,
                                        "stream": True  # Enable streaming to match original behavior
                                    }
                                    
                                    print(data)
                                    
                                    # Call local server instead of remote API
                                    response = requests.post("http://localhost:8000/v1/chat/completions", 
                                                            headers=headers, 
                                                            data=json.dumps(data), 
                                                            stream=True)
                                    
                                    import pyautogui
                                    if response.status_code == 200:
                                        items = []
                                        total = ""
                                        for chunk in response.iter_lines():
                                            try:
                                                if chunk:
                                                    # Parse SSE format from llama.cpp server
                                                    chunk_str = chunk.decode('utf-8')
                                                    if chunk_str.startswith('data: '):
                                                        chunk_data = chunk_str[6:]  # Remove 'data: ' prefix
                                                        if chunk_data.strip() == '[DONE]':
                                                            break
                                                        
                                                        chunk_json = json.loads(chunk_data)
                                                        if 'choices' in chunk_json and len(chunk_json['choices']) > 0:
                                                            delta = chunk_json['choices'][0].get('delta', {})
                                                            if 'content' in delta:
                                                                content = delta['content']
                                                                total += content
                                                                print(content, end='', flush=True)
                                            except Exception as e:
                                                # If streaming fails, fall back to non-streaming
                                                print(f"Streaming error: {e}")
                                                break
                                            
                                        # If streaming didn't work or total is empty, try non-streaming
                                        if not total:
                                            data["stream"] = False
                                            response = requests.post("http://localhost:8000/v1/chat/completions", 
                                                                    headers=headers, 
                                                                    data=json.dumps(data))
                                            if response.status_code == 200:
                                                result = response.json()
                                                total = result["choices"][0]["message"]["content"]
                                                print(total)
                                            else:
                                                print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                                                total = "Error communicating with local LLM server"
                                        
                                        print()  # New line after streaming output
                                        print(total)
                                    else:
                                        print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                                        total = "Error communicating with local LLM server"
                                        if "Tracking Tag" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Tracking Tag"] != "None":
                                            total += " #244" + self.global_variables["key"][:5]
                                        self.global_variables["GPT4"  + graph["nodes"][node_id]["name"].split("GPT4 ")[1]] = total
                                        clipboard.copy(total)
                                        pyautogui.hotkey('ctrl', 'v')

                                        self.thread_signals.notifications.emit("GPT-4", total, "icon.png", None, "No")
                                        print(self.global_variables)


                                else:
                                    clipboard.copy(str(x["key"]).upper())
                                    pyautogui.hotkey('ctrl', 'v')
                                    #pyautogui.write(str(x["key"]).upper(), interval=0.01)                        
                        else:
                            print("lower keypress")
                            if "{{folder}}" in str(x["key"]):
                                pyautogui.write(folder, interval=0.01)
                                pyautogui.write("\\", interval=0.01)
                                pyautogui.write(str(x["key"]).lower().replace("{{folder}}", ""), interval=0.01)
                                
                            else:
                                if "GPT-4 Mode" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["GPT-4 Mode"] == "True":
                                    print("GPT-4 Mode")
                                    prompt = x["key"]
                                    log = [{"role": "system", "content": "You are a helpfull assistant that only generates exactly the text the user requests. Do not explain the text. SImply generate exactly what the user requests. For example, if a user asks to generate a social post about ai automation, you would write 'Our business keeps growing faster thanks to AI automation #AI #aiautomation'"}]

                                    Automation_input = ""
                                    Webhook_input = ""
                                    if "Automation Input" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Automation Input"] != "None" and len(graph["nodes"][node_id]["custom"]["Automation Input"]) > 0:
                                        Automation_input = self.global_variables[graph["nodes"][node_id]["custom"]["Automation Input"]]

                                    log.append({"role": "user", "content": prompt.replace("space", " ") + Automation_input + Webhook_input})
                                    data = {
                                        "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
                                        "messages": log,
                                        "temperature": 0.7,
                                        "max_tokens": 5000,
                                        "stream": True  # Enable streaming to match original behavior
                                    }

                                    print(data)

                                    # Call local server instead of remote API
                                    response = requests.post("http://localhost:8000/v1/chat/completions", 
                                                            headers=headers, 
                                                            data=json.dumps(data), 
                                                            stream=True)

                                    import pyautogui
                                    if response.status_code == 200:
                                        items = []
                                        total = ""
                                        for chunk in response.iter_lines():
                                            try:
                                                if chunk:
                                                    # Parse SSE format from llama.cpp server
                                                    chunk_str = chunk.decode('utf-8')
                                                    if chunk_str.startswith('data: '):
                                                        chunk_data = chunk_str[6:]  # Remove 'data: ' prefix
                                                        if chunk_data.strip() == '[DONE]':
                                                            break
                                                        
                                                        chunk_json = json.loads(chunk_data)
                                                        if 'choices' in chunk_json and len(chunk_json['choices']) > 0:
                                                            delta = chunk_json['choices'][0].get('delta', {})
                                                            if 'content' in delta:
                                                                content = delta['content']
                                                                total += content
                                                                print(content, end='', flush=True)
                                            except Exception as e:
                                                # If streaming fails, fall back to non-streaming
                                                print(f"Streaming error: {e}")
                                                break
                                            
                                        # If streaming didn't work or total is empty, try non-streaming
                                        if not total:
                                            data["stream"] = False
                                            response = requests.post("http://localhost:8000/v1/chat/completions", 
                                                                    headers=headers, 
                                                                    data=json.dumps(data))
                                            if response.status_code == 200:
                                                result = response.json()
                                                total = result["choices"][0]["message"]["content"]
                                                print(total)
                                            else:
                                                print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                                                total = "Error communicating with local LLM server"

                                        print()  # New line after streaming output
                                        print(total)
                                    else:
                                        print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                                        total = "Error communicating with local LLM server"
                                        if "Tracking Tag" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Tracking Tag"] != "None":
                                            total += " #244" + self.global_variables["key"][:5]
                                        self.global_variables["GPT4"  + graph["nodes"][node_id]["name"].split("GPT4 ")[1]] = total
                                        clipboard.copy(total)
                                        pyautogui.hotkey('ctrl', 'v')

                                        self.thread_signals.notifications.emit("GPT-4", total, "icon.png", None, "No")
                                        print(self.global_variables)


                                else:

                                    import pyautogui
                                    print("lower keypress regular")
                                    print(graph["nodes"][node_id]["custom"]["String"])

                                    pyautogui.typewrite(graph["nodes"][node_id]["custom"]["String"])

                                    print('finished keypress')
                
            if x["type"] == "sendData":
                print("SEND DATA")
                outdata = {}
                if graph["nodes"][node_id]["custom"]["Body Value 1"] in self.global_variables:
                    outdata[graph["nodes"][node_id]["custom"]["Body Key 1"]] = self.global_variables[graph["nodes"][node_id]["custom"]["Body Value 1"]]
                if graph["nodes"][node_id]["custom"]["Body Value 2"] in self.global_variables:
                    outdata[graph["nodes"][node_id]["custom"]["Body Key 2"]] = self.global_variables[graph["nodes"][node_id]["custom"]["Body Value 2"]]
                if graph["nodes"][node_id]["custom"]["Body Value 3"] in self.global_variables:
                   outdata[graph["nodes"][node_id]["custom"]["Body Key 3"]] = self.global_variables[graph["nodes"][node_id]["custom"]["Body Value 3"]]
                if graph["nodes"][node_id]["custom"]["Body Value 4"] in self.global_variables:
                    outdata[graph["nodes"][node_id]["custom"]["Body Key 4"]] = self.global_variables[graph["nodes"][node_id]["custom"]["Body Value 4"]]
                if graph["nodes"][node_id]["custom"]["Body Value 5"] in self.global_variables:
                    outdata[graph["nodes"][node_id]["custom"]["Body Key 5"]] = self.global_variables[graph["nodes"][node_id]["custom"]["Body Value 5"]]
                print(outdata)
                #loop through all the keys, and check if they have {{}} in them, if so, replace them with the global variable
                for key, value in outdata.items():
                    if "{{" in value:
                        #replace any strings that include {{name}} with self.global_variables["name"]
                        for key2, value2 in self.global_variables.items():
                            if key2 in value:
                                outdata[key] = outdata[key].replace("{{" + key2 + "}}", self.global_variables[key2])
                print(outdata)
                if "Headers" in graph["nodes"][node_id]["custom"]:
                    r = requests.post(graph["nodes"][node_id]["custom"]["URL"], data=outdata,  headers=json.loads(graph["nodes"][node_id]["custom"]["Headers"]))
                else:
                    r = requests.post(graph["nodes"][node_id]["custom"]["URL"], data=outdata)
                print(r.text)

            if x["type"] == "print":
                vars = graph["nodes"][node_id]["custom"]["Variables"]
                print(vars.split("_")[1])
                print(self.global_variables[vars])
                #CustomMessageBox.showWithTimeout(3, "Print Variable",self.global_variables[int(vars.split("_")[1])], icon=QMessageBox.Warning)

            if x["type"] == "OCR":
                time.sleep(5)
                print("OCR")
                #cv2.namedWindow("OCR")
                img = pyautogui.screenshot()
                screenshot = cv2.cvtColor(np.array(img), cv2.COLOR_BGR2RGB)
                print("start OCR")

            # draw a rectangle around the region of interest

                thresh = 255 - cv2.threshold(cv2.cvtColor(screenshot, cv2.COLOR_BGR2GRAY), 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]
                print("start OCR")
               # xp,yp,w,h = graph["nodes"][node_id]["custom"]["bounding_box"][0][0], graph["nodes"][node_id]["custom"]["bounding_box"][0][1], graph["nodes"][node_id]["custom"]["bounding_box"][1][0]-graph["nodes"][node_id]["custom"]["bounding_box"][0][0], graph["nodes"][node_id]["custom"]["bounding_box"][1][1]-graph["nodes"][node_id]["custom"]["bounding_box"][0][1]
                ROI = thresh[graph["nodes"][node_id]["custom"]["Bounding Box Y1"]:graph["nodes"][node_id]["custom"]["Bounding Box Y2"],graph["nodes"][node_id]["custom"]["Bounding Box X1"]:graph["nodes"][node_id]["custom"]["Bounding Box X2"]]
                print("start OCR")

                data = pytesseract.image_to_string(ROI, lang='eng',config='--psm 6')
                print("restults")
                print(data)
                self.global_variables["OCR" + str(len(self.global_variables))] = data
                cv2.imshow("thresh",thresh)
                cv2.imshow("ROI",ROI)

                cv2.waitKey(1)

                #self.global_variable_names.append("OCR" + len(global_variables))

            if x["type"] == "scroll":
                import time
                time.sleep(10)
                import pyautogui
                pyautogui.scroll(int(graph["nodes"][node_id]["custom"]["Distance"]))
                self.thread_signals.notifications.emit("Scrolling", "Scrolling", "icon.png", None, "No")

                print("scroll")

            if x["type"] == "getData":
                print(graph["nodes"][node_id]["custom"]["URL"])
                if ":\\" in graph["nodes"][node_id]["custom"]["URL"]:
                    #open the file and read the contents 
                    #check if the file is a word doc, then load it appropriately
                    if ".docx" in graph["nodes"][node_id]["custom"]["URL"]:
                        import docx
                        doc = docx.Document(graph["nodes"][node_id]["custom"]["URL"])
                        data = ""
                        for para in doc.paragraphs:
                            data += para.text
                        self.global_variables["GetData"  +  graph["nodes"][node_id]["name"].split("GetData ")[1]] = data
                    elif "xlsx" in graph["nodes"][node_id]["custom"]["URL"]:
                        import openpyxl
                        wb = openpyxl.load_workbook(graph["nodes"][node_id]["custom"]["URL"])
                        sheet = wb.active
                        data = ""
                        for row in sheet.iter_rows(values_only=True):
                            data += str(row)
                        self.global_variables["GetData"  +  graph["nodes"][node_id]["name"].split("GetData ")[1]] = data
                    else:
                        with open(graph["nodes"][node_id]["custom"]["URL"], "r") as file:
                            data = file.read()
                            self.global_variables["GetData"  +  graph["nodes"][node_id]["name"].split("GetData ")[1]] = data
                elif "docs.google" in graph["nodes"][node_id]["custom"]["URL"]:
                    print("google sheets")
                    data = self.sendMessageRTCAsync('google_sheets_read:' + graph["nodes"][node_id]["custom"]["URL"])
                    print("get data output")
                    print(data)

                    self.global_variables["GetData"  +  graph["nodes"][node_id]["name"].split("GetData ")[1]] = data
                elif len(graph["nodes"][node_id]["custom"]["Headers"]) != 0:
                    response = requests.get(url=graph["nodes"][node_id]["custom"]["URL"],headers=json.loads(graph["nodes"][node_id]["custom"]["Headers"]))
                    self.global_variables["GetData"  + str(len(self.global_variables))] = response.text
                    print(self.global_variables)
                else:
                    response = requests.get(url=graph["nodes"][node_id]["custom"]["URL"])
                    print(response)
                    print(response.text)
                    import csv

                    csv_bytestream = response.content.decode('utf-8')
                    cr = csv.reader(csv_bytestream.splitlines(), delimiter=',')
                    my_list = list(cr)
                    my_list = list(cr)
                    counter = 0
                    for row in my_list:
                        print(row)
                        self.global_variables["get_data_"  + counter + str(len(self.global_variables))] = row
                        counter += 1
                        #self.global_variable_names.append(key)


            if x["type"] == "screenshot":
                pyautogui.screenshot(x["download_location"])
            if x["type"] == "Approval":
                r = requests.post("https://cheatlayer.com/triggers/extension", data={"key":graph["nodes"][node_id]["custom"]["key"]  ,"script":"sendGmail(`" + graph["nodes"][node_id]["custom"]["To"] + "`,`Cheat Layer Desktop Agent Approval [Action Required]`,`" + graph["nodes"][node_id]["custom"]["body"] + var + "`)","start":""})
                agent_approval = 0
                while agent_approval == 0:
                    time.sleep(5)
                    print("check email")
                    print (graph["nodes"][node_id]["custom"]["To"])
                if agent_approval == 1:
                    print("approved")
                    
                    self.global_variables["Approval"] = True
                elif agent_approval == -1:
                    print("denied")
                    self.global_variables["Approval"] = False

            if x["type"] == "download":
                #check if URL is a file or URL and download it if it's a URL or load the file from a pdf, doc, text, or csv
                if True:
                    #open the file and read the contents
                    #check if the file is a word doc, then load it appropriately
                    if ".docx" in graph["nodes"][node_id]["custom"]["URL or File"]:
                        import docx
                        doc = docx.Document(graph["nodes"][node_id]["custom"]["URL or File"])
                        data = ""
                        for para in doc.paragraphs:
                            data += para.text
                        self.global_variables["Download" + graph["nodes"][node_id]["name"].split("Download ")[1]] = data
                    elif "pdf" in graph["nodes"][node_id]["custom"]["URL or File"]:
                        import PyPDF2
                        pdfFileObj = open(graph["nodes"][node_id]["custom"]["URL or File"], 'rb')
                        pdfReader = PyPDF2.PdfFileReader(pdfFileObj)
                        data = ""
                        for page_num in range(pdfReader.numPages):
                            pageObj = pdfReader.getPage(page_num)
                            data += pageObj.extractText()
                        self.global_variables["Download" + graph["nodes"][node_id]["name"].split("Download ")[1]] = data

                    elif "xlsx" in graph["nodes"][node_id]["custom"]["URL or File"]:
                        import openpyxl
                        wb = openpyxl.load_workbook(graph["nodes"][node_id]["custom"]["URL or File"])
                        sheet = wb.active
                        data = ""
                        for row in sheet.iter_rows(values_only=True):
                            data += str(row)
                        self.global_variables["Download" + graph["nodes"][node_id]["name"].split("Download ")[1]] = data
                    else:
                        with open(graph["nodes"][node_id]["custom"]["URL or File"], "r") as file:
                            data = file.read()
                            self.global_variables["Download" + graph["nodes"][node_id]["name"].split("Download ")[1]] = data
                    print(self.global_variables)
                elif "docs.google" in graph["nodes"][node_id]["custom"]["URL or File"]:
                    print("google sheets")
                    data = self.sendMessageRTCAsync('google_sheets_read:' + graph["nodes"][node_id]["custom"]["URL or File"])
                    print("get data output")
                    print(data)

                    self.global_variables["Download" + graph["nodes"][node_id]["name"].split("Download ")[1]] = data
                elif len(graph["nodes"][node_id]["custom"]["Headers"]) != 0:
                    response = requests.get(url=graph["nodes"][node_id]["custom"]["URL or File"],headers=json.loads(graph["nodes"][node_id]["custom"]["Headers"]))
                    self.global_variables["Download" + graph["nodes"][node_id]["name"].split("Download ")[1]] = response.text
                    print(self.global_variables)
                else:
                    response = requests.get(graph["nodes"][node_id]["custom"]["URL or File"])
                    open(graph["nodes"][node_id]["custom"]["Download Location"], "wb").write(response.content)
                    #store the download data in the global variables
                    self.global_variables["Download" + graph["nodes"][node_id]["name"].split("Download ")[1]] = response.content


            if x["type"] == "Move Mouse" or ("Type" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Type"] == "Move Mouse"):
                print("Move Mouse")
                if "recording" in graph["nodes"][node_id]["custom"]:
                    print(graph["nodes"][node_id]["custom"]["recording"])
                    import ast

                    list_moves = ast.literal_eval(graph["nodes"][node_id]["custom"]["recording"])
                    print(list_moves)
                    for move in list_moves:
                        pyautogui.moveTo(move[0], move[1])
                else:
                    pyautogui.moveTo((int(graph["nodes"][node_id]["custom"]["X Coordinate"])), (int(graph["nodes"][node_id]["custom"]["Y Coordinate"])))
                    last_mouse_x = int(graph["nodes"][node_id]["custom"]["X Coordinate"])
                    last_mouse_y = int(graph["nodes"][node_id]["custom"]["Y Coordinate"])
            if x["type"] == "getRecording":
                #record the microphone and save it to a file
                print("getRecording")
                #use python to record the microphone in time seconds 
                if "Screenshot File" in graph["nodes"][node_id]["custom"]:
                    import pyautogui
                    print("Screenshot File")
                    print(graph["nodes"][node_id]["custom"]["Screenshot File"])
                    pyautogui.screenshot(resource_path(graph["nodes"][node_id]["custom"]["Screenshot File"] + ".png"))
                audio_file = "test"
                if "Audio File" in graph["nodes"][node_id]["custom"] and len(graph["nodes"][node_id]["custom"]["Audio File"]) > 0:
                    audio_file = graph["nodes"][node_id]["custom"]["Audio File"]
                
                if "File" in graph["nodes"][node_id]["custom"] and len(graph["nodes"][node_id]["custom"]["File"]) > 0:
                    audio_file = graph["nodes"][node_id]["custom"]["File"]
                
                FORMAT = pyaudio.paInt16
                CHANNELS = 2
                RATE = 44100
                CHUNK = 1024
                RECORD_SECONDS = int(graph["nodes"][node_id]["custom"]["Time"])
                WAVE_OUTPUT_FILENAME = resource_path(audio_file+ ".wav")

                audio = pyaudio.PyAudio()

                # start Recording
                stream = audio.open(format=FORMAT, channels=CHANNELS,
                                    rate=RATE, input=True,
                                    frames_per_buffer=CHUNK)
                print("recording...")
                frames = []

                for i in range(0, int(RATE / CHUNK * RECORD_SECONDS)):
                    data = stream.read(CHUNK)
                    frames.append(data)
                print("finished recording")
                # stop Recording
                stream.stop_stream()
                stream.close()
                audio.terminate()
                waveFile = wave.open(WAVE_OUTPUT_FILENAME, 'wb')
                #overwrite the file and replace it

                waveFile.setnchannels(CHANNELS)
                waveFile.setsampwidth(audio.get_sample_size(FORMAT))
                waveFile.setframerate(RATE)
                waveFile.writeframes(b''.join(frames))
                waveFile.close()
            
                import whisper

                # Load the pre-trained Whisper model
                model = whisper.load_model("base")  # You can choose 'tiny', 'base', 'small', 'medium', or 'large'

                # Transcribe the WAV file
                result = model.transcribe(WAVE_OUTPUT_FILENAME)

                # Output the recognized text
                print(result['text'])
    #cleanip the file
                self.global_variables["Recording" + graph["nodes"][node_id]["name"].split("Recording ")[1]] = result['text']

                os.remove(WAVE_OUTPUT_FILENAME)             
            if x["type"] == "Left Mouse Lift":
                pyautogui.mouseUp()
            if x["type"] == "StableVideo":
                print("Synthetic video generator")
                print("RUNNING DALLE 3")
                prompt = graph["nodes"][node_id]["custom"]["prompt"]
                if "{{" in prompt:
                    #replace any strings that include {{name}} with self.global_variables["name"]
                    for key, value in self.global_variables.items():
                        if key in prompt:
                            prompt = prompt.replace("{{" + key + "}}", self.global_variables[key])
                automation_input = ""
                if "Automation Input" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Automation Input"] != "None" and len(graph["nodes"][node_id]["custom"]["Automation Input"]) > 0:
                    automation_input = self.global_variables[graph["nodes"][node_id]["custom"]["Automation Input"]]
                response = requests.post("https://cheatlayer.com/user/generateDalleU", data={'id': self.user_key, "input":  prompt + automation_input },  verify=False)

                print(response.json())
                if response.status_code != 200:
                    print(response.json())

                    raise ValueError("Failed to start image generation. Status code:", response.json())

                data = response.json()
                URL = data["data"][0]["url"]
                self.thread_signals.notifications.emit("Stable Video Generating", "Stable Video Generating", "icon.png", None, "No")

                #"sdGenerationJob":{"generationId":
                #download and convert the image to base64
                response = requests.get(URL)
                img = Image.open(io.BytesIO(response.content))
                img.save(resource_path("image.png"))
                
                self.upload_image(resource_path("image.png"))
                print("RUNNING SYNTHETIC VIDEO GENERATOR")
                response_URL = call_replicate3(URL, graph["nodes"][node_id]["custom"]["resolution"], graph["nodes"][node_id]["custom"]["frames"], graph["nodes"][node_id]["custom"]["motion bucket"])
                #download and open the .mp4
                response = requests.get(response_URL)
                with open(resource_path(graph["nodes"][node_id]["custom"]["file"]), 'wb') as f:
                    f.write(response.content)
                #open the mp4
                self.upload_image(resource_path(graph["nodes"][node_id]["custom"]["file"]))

               # os.system("start " + resource_path(graph["nodes"][node_id]["custom"]["file"]))

                print("FINISED SYNTHETIC VIDEO GENERATOR")
                print("saved video to " + resource_path(graph["nodes"][node_id]["custom"]["file"]))

            if x["type"] == "cheatlayer":
                print("SEND Cheat Layer")
                print(graph["nodes"][node_id]["custom"]["URL"])
                print(graph["nodes"][node_id]["custom"]["code"])
                r = requests.post("https://cheatlayer.com/triggers/extension", data={"key":graph["nodes"][node_id]["custom"]["key"]  ,"script":graph["nodes"][node_id]["custom"]["code"],"start":graph["nodes"][node_id]["custom"]["URL"]})
                print(r)

            if x["type"] == "tracker":
                # Assuming some condition here; currently always True
                # if True:
                import time, pyautogui, io, base64
                time.sleep(3)
                
                marketing_data, cookies, user_id = load_data(os.path.expanduser("~/marketingData.json"))
                #self.thread_signals.scraper.emit(user_id, graph["nodes"][node_id]["custom"]["Prompts Used"])
            if x["type"] == "SemanticDescribe":
                if "Scrape Browser" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Scrape Browser"] != "" and len(graph["nodes"][node_id]["custom"]["Scrape Browser"]) > 2:
                    #open chrome 
                    import subprocess
                    print(graph["nodes"][node_id]["custom"]["Scrape Browser"])
                    print("SCRAPE BROWSER")
                    #subprocess.Popen("google-chrome")
                    if graph["nodes"][node_id]["custom"]["Scrape Browser"] == "Links":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'a')
                    elif graph["nodes"][node_id]["custom"]["Scrape Browser"] == "Images":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'img')
                    elif graph["nodes"][node_id]["custom"]["Scrape Browser"] == "Text": 
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'p')
                    elif graph["nodes"][node_id]["custom"]["Scrape Browser"] == "Headings":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'h1')
                    elif graph["nodes"][node_id]["custom"]["Scrape Browser"] == "Buttons":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'button')
                    elif graph["nodes"][node_id]["custom"]["Scrape Browser"] == "Inputs":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'input')
                    elif graph["nodes"][node_id]["custom"]["Scrape Browser"] == "Forms":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'form')
                    elif graph["nodes"][node_id]["custom"]["Scrape Browser"] == "Divs":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'div')
                    elif graph["nodes"][node_id]["custom"]["Scrape Browser"] == "Spans":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'span')
                    elif graph["nodes"][node_id]["custom"]["Scrape Browser"] == "Lists":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'ul')
                    elif graph["nodes"][node_id]["custom"]["Scrape Browser"] == "Tables":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'table')
                    elif graph["nodes"][node_id]["custom"]["Scrape Browser"] == "Selects":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'select')
                    elif graph["nodes"][node_id]["custom"]["Scrape Browser"] == "Options":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'option')
                    elif graph["nodes"][node_id]["custom"]["Scrape Browser"] == "Textareas":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'textarea')
                    elif graph["nodes"][node_id]["custom"]["Scrape Browser"] == "Images":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'img')
                    elif graph["nodes"][node_id]["custom"]["Scrape Browser"] == "Videos":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'video')
                    elif graph["nodes"][node_id]["custom"]["Scrape Browser"] == "Audios":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'audio')
                    elif graph["nodes"][node_id]["custom"]["Scrape Browser"] == "Canvas":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'canvas')
                    

                    websocket_data = ""
                    print(browser_elements)
                    description = ""
                    automation_input = ""
                    if "Automation Input" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Automation Input"] != "None" and len(graph["nodes"][node_id]["custom"]["Automation Input"]) > 0:
                        automation_input = self.global_variables[graph["nodes"][node_id]["custom"]["Automation Input"]]

                    description = graph["nodes"][node_id]["custom"]["Target In English"]
                    if "Add Node Runs To Prompt" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Add Node Runs To Prompt"] == "true":
                        description += " " + str(node_runs[graph["nodes"][node_id]["name"]])
                    if "Target In English" in graph["nodes"][node_id]["custom"]:
                        description = graph["nodes"][node_id]["custom"]["Target In English"]
                    log = [{"role": "system", "content": "You are a helpfull assistant that takes the list of input elements and returns only the array of elements which matches the user's intent. Return a properly formatted array in a json like this:'{\"data\":[data1,data2,data3...]}'.  Only return the json directly and don't add quotes or a prefix like ```json."}]

                    print(browser_elements)
                    print("sending to semantic search")
                    log.append({"role": "user", "content": "The intended target is " + description  + ". Generate a valid JSON only with double quotes for strings. Only return the json directly and don't add quotes or a prefix like ```json. The list of input elements is:" + json.dumps(browser_elements)[:100000]})
                    if not hasattr(self, 'llama_server_process') or self.llama_server_process is None:
                        self.setup_local_llama_server()
    
                    # Convert OpenAI vision format to llama.cpp format
                    def convert_message_for_llama(message):
                        """Convert OpenAI vision API format to llama.cpp format"""
                        if isinstance(message.get('content'), list):
                            # Extract text and image from the content array
                            text_content = ""
                            image_data = None

                            for item in message['content']:
                                if item['type'] == 'text':
                                    text_content += item['text']
                                elif item['type'] == 'image_url':
                                    # Extract base64 data from data URL
                                    image_url = item['image_url']['url']
                                    if image_url.startswith('data:image/'):
                                        # Remove the data:image/jpeg;base64, prefix
                                        image_data = image_url.split(',', 1)[1]

                            return {
                                "role": message['role'],
                                "content": text_content,
                                "images": [image_data] if image_data else []
                            }
                        else:
                            # Regular text message
                            return message

                    # Convert all messages to llama.cpp format
                    converted_log = [convert_message_for_llama(msg) for msg in log]

                    data = {
                        "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
                        "messages": converted_log,
                        "temperature": 0.7,
                        "max_tokens": 5000,
                        "stream": True  # Enable streaming to match original behavior
                    }

                    print(data)

                    # Call local server instead of remote API
                    response = requests.post("http://localhost:8000/v1/chat/completions", 
                                            headers=headers, 
                                            data=json.dumps(data), 
                                            stream=True)

                    import pyautogui
                    if response.status_code == 200:
                        items = []
                        total = ""
                        for chunk in response.iter_lines():
                            try:
                                if chunk:
                                    # Parse SSE format from llama.cpp server
                                    chunk_str = chunk.decode('utf-8')
                                    if chunk_str.startswith('data: '):
                                        chunk_data = chunk_str[6:]  # Remove 'data: ' prefix
                                        if chunk_data.strip() == '[DONE]':
                                            break
                                        
                                        chunk_json = json.loads(chunk_data)
                                        if 'choices' in chunk_json and len(chunk_json['choices']) > 0:
                                            delta = chunk_json['choices'][0].get('delta', {})
                                            if 'content' in delta:
                                                content = delta['content']
                                                total += content
                                                print(content, end='', flush=True)
                            except Exception as e:
                                # If streaming fails, fall back to non-streaming
                                print(f"Streaming error: {e}")
                                break
                            
                        # If streaming didn't work or total is empty, try non-streaming
                        if not total:
                            data["stream"] = False
                            response = requests.post("http://localhost:8000/v1/chat/completions", 
                                                    headers=headers, 
                                                    data=json.dumps(data))
                            if response.status_code == 200:
                                result = response.json()
                                total = result["choices"][0]["message"]["content"]
                                print(total)
                            else:
                                print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                                total = "Error communicating with local LLM server"

                        print()  # New line after streaming output
                        print(total)
                    else:
                        print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                        total = "Error communicating with local LLM server"
                        import re
                        def is_valid_json(json_string):
                            try:
                                json.loads(json_string)  # Try parsing the JSON
                                return True  # Return True if parsing is successful
                            except json.JSONDecodeError:
                                return False  # Return False if an error occurs

         

                        def fix_json(json_string):
                            # Check and fix missing quotes, brackets, or trailing commas
                            if not is_valid_json(json_string):
                                try_fixes = [
                                    json_string + ']}',             # Close with nested structures
                                    json_string + '}',              # Close with object
                                    json_string + ']',              # Close with array
                                    json_string.replace('}{', '},{') # Fix missing commas between objects
                                ]
                                for fix in try_fixes:
                                    if is_valid_json(fix):
                                        return fix
                                return None
                            return json_string
                        if is_valid_json(total):
                            print("valid json")
                        else:
                            print("invalid json")
                            total = fix_json(total)
                            print(total)
                        
                        links = json.loads(total)["data"]
                        link_dests = []
                        link_text = []
                        for link in links:
                            print(link)
                            if ":;:" in link:
                                link_dests.append(link.split(":;:")[1])
                                link_text.append(link.split(":;:")[0])
                            else:
                                link_dests.append(link)
                        self.global_variables["Browser Elements"  +  graph["nodes"][node_id]["name"].split("Magic Scraper ")[1]] = link_dests
                        self.global_variables["Browser Texts"  +  graph["nodes"][node_id]["name"].split("Magic Scraper ")[1]] = link_text

                        self.global_variables["Links"  +  graph["nodes"][node_id]["name"].split("Magic Scraper ")[1]] = link_dests
                        self.global_variables["Link Texts"  +  graph["nodes"][node_id]["name"].split("Magic Scraper ")[1]] = link_text
                        all_magic_scrapers = ""
                        for key, value in self.global_variables.items():
                            if "Magic Scraper" in key:
                                all_magic_scrapers += key + " : " + value + "\n"
                        self.global_variables["Magic Scraper All Outputs"] = all_magic_scrapers
                        print(self.global_variables)
                elif "Scrape Links" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Scrape Links"] == "True":
                    #open chrome 
                    import subprocess
                    #subprocess.Popen("google-chrome")
                    browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'a')
                    websocket_data = ""
                    print(browser_elements)
                    description = ""
                    automation_input = ""
                    if "Automation Input" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Automation Input"] != "None" and len(graph["nodes"][node_id]["custom"]["Automation Input"]) > 0:
                        automation_input = self.global_variables[graph["nodes"][node_id]["custom"]["Automation Input"]]

                    description = graph["nodes"][node_id]["custom"]["Target In English"]
                    if "Target In English" in graph["nodes"][node_id]["custom"]:
                        description = graph["nodes"][node_id]["custom"]["Target In English"]
                    log = [{"role": "system", "content": "You are a helpfull assistant that takes the list of input elements and returns only the array of elements which matches the user's intent. Return a properly formatted array in a json like this:'{\"data\":[data1,data2,data3...]}'.  Only return the json directly and don't add quotes or a prefix like ```json."}]

                    print("sending to semantic search")
                    log.append({"role": "user", "content": "The intended target is " + description  + ". Only return the json directly and don't add quotes or a prefix like ```json. The list of input elements is:" + json.dumps(browser_elements)[:50000]})
                    data = {
                        "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
                        "messages": log,
                        "temperature": 0.7,
                        "max_tokens": 5000,
                        "stream": True  # Enable streaming to match original behavior
                    }

                    print(data)

                    # Call local server instead of remote API
                    response = requests.post("http://localhost:8000/v1/chat/completions", 
                                            headers=headers, 
                                            data=json.dumps(data), 
                                            stream=True)

                    import pyautogui
                    if response.status_code == 200:
                        items = []
                        total = ""
                        for chunk in response.iter_lines():
                            try:
                                if chunk:
                                    # Parse SSE format from llama.cpp server
                                    chunk_str = chunk.decode('utf-8')
                                    if chunk_str.startswith('data: '):
                                        chunk_data = chunk_str[6:]  # Remove 'data: ' prefix
                                        if chunk_data.strip() == '[DONE]':
                                            break
                                        
                                        chunk_json = json.loads(chunk_data)
                                        if 'choices' in chunk_json and len(chunk_json['choices']) > 0:
                                            delta = chunk_json['choices'][0].get('delta', {})
                                            if 'content' in delta:
                                                content = delta['content']
                                                total += content
                                                print(content, end='', flush=True)
                            except Exception as e:
                                # If streaming fails, fall back to non-streaming
                                print(f"Streaming error: {e}")
                                break
                            
                        # If streaming didn't work or total is empty, try non-streaming
                        if not total:
                            data["stream"] = False
                            response = requests.post("http://localhost:8000/v1/chat/completions", 
                                                    headers=headers, 
                                                    data=json.dumps(data))
                            if response.status_code == 200:
                                result = response.json()
                                total = result["choices"][0]["message"]["content"]
                                print(total)
                            else:
                                print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                                total = "Error communicating with local LLM server"

                        print()  # New line after streaming output
                        print(total)
                    else:
                        print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                        total = "Error communicating with local LLM server"
                        import re
                        def is_valid_json(json_string):
                            try:
                                json.loads(json_string)  # Try parsing the JSON
                                return True  # Return True if parsing is successful
                            except json.JSONDecodeError:
                                return False  # Return False if an error occurs

         

                        def fix_json(json_string):
                            # Check and fix missing quotes, brackets, or trailing commas
                            if not is_valid_json(json_string):
                                try_fixes = [
                                    json_string + ']}',             # Close with nested structures
                                    json_string + '}',              # Close with object
                                    json_string + ']',              # Close with array
                                    json_string.replace('}{', '},{') # Fix missing commas between objects
                                ]
                                for fix in try_fixes:
                                    if is_valid_json(fix):
                                        return fix
                                return None
                            return json_string
                        if is_valid_json(total):
                            print("valid json")
                        else:
                            print("invalid json")
                            total = fix_json(total)
                            print(total)
                        
                        links = json.loads(total)["data"]
                        link_dests = []
                        link_text = []
                        for link in links:
                            print(link)
                            if ":;:" in link:
                                link_dests.append(link.split(":;:")[1])
                                link_text.append(link.split(":;:")[0])
                            else:
                                link_dests.append(link)
                        self.global_variables["Links"  +  graph["nodes"][node_id]["name"].split("Magic Scraper ")[1]] = link_dests
                        self.global_variables["Link Texts"  +  graph["nodes"][node_id]["name"].split("Magic Scraper ")[1]] = link_text

                        print(self.global_variables)
                else:
                    import pyautogui
                    import time
                    import io
                    time.sleep(3)
                    screenshot = pyautogui.screenshot()
                    
                    buf = io.BytesIO()
                
                    rgb_screenshot = screenshot.convert("RGB")
                    rgb_screenshot.save(buf, format='JPEG')
                    byte_im = buf.getvalue()
                    import base64
                    encoded_jpeg = base64.b64encode(io.BytesIO(byte_im).read()).decode('utf-8')
                    automation_input = ""
                    if "Automation Input" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Automation Input"] != "None" and len(graph["nodes"][node_id]["custom"]["Automation Input"]) > 0:
                        automation_input = self.global_variables[graph["nodes"][node_id]["custom"]["Automation Input"]]

                    description = graph["nodes"][node_id]["custom"]["Target In English"]
                    if "Target In English" in graph["nodes"][node_id]["custom"]:
                        description = graph["nodes"][node_id]["custom"]["Target In English"]
                    if type(automation_input) == dict:
                        automation_input = json.dumps(automation_input)
                
                    if "Add Node Runs To Prompt" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Add Node Runs To Prompt"] == "true":
                        description += " " + str(node_runs[graph["nodes"][node_id]["name"]])
                    print(description)
                    caption = blip_caption_describe(encoded_jpeg,description + " " + automation_input, self.user_key, self.user_plan)
                    print(caption)
                    self.global_variables["Magic Scraper Output"  +  graph["nodes"][node_id]["name"].split("Magic Scraper ")[1]] = caption
                    print(self.global_variables)
                    self.thread_signals.notifications.emit("Magic Scraper ", caption, "icon.png", None, "No")
                    import time
                    time.sleep(2)
               #     self.notification_manager.add_notification("Magic Scraper ", caption, "icon.png", lambda: print("Notification closed"))
              #      self.notification_manager.toggle_notifications_visibility()
                    print("Magic Scraper  ran")
            if x["type"] == "semanticScrape":
                time.sleep(2)
                description = graph["nodes"][node_id]["custom"]["Semantic Description"]

                browser_max = 0
                browser_new_max = {}
                total_matches = []
                browser_matches = []
                self.chatInstance.showMinimized()
                self.label.hide()
                self.label2.hide()
                
                screen_width = pyautogui.size()[0]
                screen_height = pyautogui.size()[1]
                raw_image = pyautogui.screenshot(region=(0,0,screen_width,screen_height))
           #     self.chatInstance.toggle_chat_display()

                raw_image.save(r"screen.png")
                elements = UIModel.predict("screen.png", confidence=1, overlap=1).json()
                UIModel.predict("screen.png", confidence=1, overlap=1).save("prediction.jpg")
                #add elements to the array that include boxes of varying sizes around x and y that resemble buttons or other elements
                  #add elements to the array that include boxes of varying sizes around x and y that resemble buttons or other elements
                x1 = graph["nodes"][node_id]["custom"]["X"]
                y = graph["nodes"][node_id]["custom"]["Y"]

                if x1 != 0 and y != 0 and graph["nodes"][node_id]["custom"]["Target"] != "Link Destination" and graph["nodes"][node_id]["custom"]["Target"] != "article" and graph["nodes"][node_id]["custom"]["Target"] != "Twitter posts":

                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 100, "height": 100, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 200, "height": 200, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 300, "height": 300, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 50, "height": 50, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 200, "height": 50, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 200, "height": 100, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 300, "height": 300, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 400, "height": 400, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 500, "height": 200, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 500, "height": 500, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 700, "height": 700, "class": "Button"})

                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 50, "height": 50, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 200, "height": 50, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 200, "height": 50, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 300, "height": 50, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 400, "height": 50, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 500, "height": 200, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 50, "height": 200, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 200, "height": 200, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 200, "height": 200, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 300, "height": 200, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 400, "height": 200, "class": "Button"})
        
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 50, "height": 50, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 200, "height": 50, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 200, "height": 50, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 300, "height": 50, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 400, "height": 50, "class": "Button"})
                    elements['predictions'].insert(0,{"x": x1, "y": y, "width": 500, "height": 200, "class": "Button"})
                print("Semantic Click")
                print(elements)

                model_url = 'https://storage.googleapis.com/sfr-vision-language-research/BLIP/models/model_base_capfilt_large.pth'
                image_size = 384

                scores = []
                self.label.show()
                self.label2.show()
                def process_element(element):
                    gc.collect()
                    score_data = {}
                    if element["class"] != "TextView":
                        caption = description
                        print('loop text: %s' %caption)
                        print('completed chat')
                        print(element)
                        image = raw_image.crop((float(element['x']) - element['width']/2, float(element['y']) - element['height']/2, float(element['x']) + element['width']/2, float(element['y']) + element['height']/2))
                        #generate a unique string id for each element based on its coordinates
                         #this is used to save the cropped image to disk
                        id_element = str(element['x']) + str(element['y']) + str(element['width']) + str(element['height'])

                        qpixmap = pil2pixmap(image)

                        image.save(resource_path(id_element + "crop.png"))
                        pixmap = qpixmap
                        self.label.setPixmap(pixmap)
                        self.label.setFixedSize(element['width'], element['height'])
                        print('completed output 1')
                        buf = io.BytesIO()
                        image.save(buf, format='JPEG')
                        byte_im = buf.getvalue()
                        
                
                        encoded_string = ""
                        with open(resource_path(id_element + "crop.png"), "rb") as image_file:
                            encoded_string = base64.b64encode(image_file.read())

                        itm_output = blip_caption(encoded_string, caption, self.user_key, self.user_plan)
                        print(itm_output)
                        score = float(itm_output.split("probability of ")[1].split(".\n")[0])
                        print('completed output 3')
                        print(itm_output)
                        self.label2.setText("Semantic Step:" + caption + " Probability:" + str(score))
                        self.label2.setFixedSize(1280,48)
                        print('The image and text is matched with a probability of %.4f'%score)
                        score_data = {"score": float(score), "x": element['x'], "y": element['y'], "width": element['width'], "height": element['height'], "cropped": open(resource_path(id_element + "crop.png"), "rb"), "browser": "false"}
                    return score_data

                # Function to split a list into chunks
                def chunks(lst, n):
                    """Yield successive n-sized chunks from lst."""
                    for i in range(0, len(lst), n):
                        yield lst[i:i + n]
        
                # Process chunks of elements
                all_scores = []
                found = False
                total_runs = 0
                if graph["nodes"][node_id]["custom"]["Target"] != "link destination" and graph["nodes"][node_id]["custom"]["Target"] != "article" and graph["nodes"][node_id]["custom"]["Target"] != "Twitter posts":
                    for chunk in chunks(elements['predictions'], 5):
             
                        with concurrent.futures.ThreadPoolExecutor() as executor:
                            scores_chunk = list(executor.map(process_element, chunk))
                            if found == False:
                                all_scores.extend(scores_chunk)
                            total_runs += 1
                            for score in scores_chunk:
                                print(score)
                                if float(score.get("score", 0)) > 0.8:
                                    print("Semantic Click FOUND MATCH")
                                    found = True

                # Continue with the rest of your code
                max_score = 0
                thresh = 0.8
                if "Threshold" in graph["nodes"][node_id]["custom"]:
                    thresh = float(graph["nodes"][node_id]["custom"]["Threshold"])
                print("done append")
                new_max = {}
                found_match = False
                for score in all_scores:
                    itm_score = score.get("score", 0)
                    if itm_score > max_score:
                        print("Semantic scrape VALID")
                        max_score = itm_score
                        new_max = score
                        print(score)
                    if itm_score > thresh:
                        found_match = True
                        #use OCR to get the text of the element
                        print("OCR BLIP")
                        #cv2.namedWindow("OCR")
                        data = blip_caption2(score["cropped"], description)
                        print("results")
                        print(data)
                        self.global_variables["OCR" + str(len(self.global_variables))] = data
                        score["text"] = data
      
                        #self.global_variable_names.append("OCR" + len(global_variables))
                        
                        
                        total_matches.append(score)
                all_scores = []

                print(total_matches)
                print("total matches")
                if found_match == True:
                    total_matches.sort(key=lambda x: x["score"], reverse=True)
                    index = 0
                    for match in total_matches:
                        all_scores.append(match["text"])
                        self.global_variables["Scrape_" + str(index)  + "_" + str(len(self.global_variables))] = match["text"]
                    #sort the matches by the score
                        index += 1
                    print(all_scores)
                    print("all scores")
                    if graph["nodes"][node_id]["custom"]["Mode"] == "First":
                        print("first")
                        print(all_scores[0])
                        self.global_variables["Scrape"  + graph["nodes"][node_id]["name"].split("Semantic Scrape ")[1]] = all_scores[0]
                    if graph["nodes"][node_id]["custom"]["Mode"] == "Last":
                        self.global_variables["Scrape"   + graph["nodes"][node_id]["name"].split("Semantic Scrape ")[1]] = all_scores[len(all_scores)-1]
                    if graph["nodes"][node_id]["custom"]["Mode"] == "All":
                        self.global_variables["Scrape"   + graph["nodes"][node_id]["name"].split("Semantic Scrape ")[1]] = ",".join(all_scores)
                else:
                    if graph["nodes"][node_id]["custom"]["Target"] == "span":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'span, h1, h2, h3, h4, h5, h6')
                    elif graph["nodes"][node_id]["custom"]["Target"] == "link text":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'a')
                    elif graph["nodes"][node_id]["custom"]["Target"] == "link destination":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'a')
                    elif graph["nodes"][node_id]["custom"]["Target"] == "article":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'article')
                    elif graph["nodes"][node_id]["custom"]["Target"] == "Twitter posts":
                        browser_elements = self.sendMessageRTCAsync('semanticSearch:' + 'article')
                    print(browser_elements)
                    print("fished loop sent")
                    images_data_urls = browser_elements["screenshots"]  # Parse JSON into list of data URLs
                    print(images_data_urls)
                    print(len(images_data_urls))
                    images = []

                    # Loop through list of data URLs
                    for data_url in images_data_urls:
                        # Strip header from data URL ('data:image/png;base64,')
                        if data_url != None and "img" in data_url and "data:image/png;base64," in data_url["img"]:
                            print(data_url)
                            base64_image = data_url["img"].split(',')[1]
                            # Decode base64
                            decoded_image = base64.b64decode(base64_image)
                            # Open image and append to list
                            temp_image = Image.open(io.BytesIO(decoded_image))
                            # Convert transparent pixels to white
                            temp_image = temp_image.convert("RGBA")
                            href = ""
                            if "href" in data_url:
                                href = data_url["href"]
                            images.append({"img":temp_image, "css": data_url["css"], "text": data_url["text"], "href": href})

                        else:
                            images.append(None)

                    # Now you have a list of PIL Image objects to process...
                    browser_scores = []
                    index = 0
                    for image in images:
                        if image != None and image["img"] != None:
                            qpixmap = pil2pixmap(image["img"])

                            image["img"].save(resource_path(str(index) + "browser.png"))
                            pixmap = qpixmap
                            self.label.setPixmap(pixmap)
                            #set the label size based on the image width and height 
                            self.label.setFixedSize(image["img"].width, image["img"].height)
                            
# Apply a stylesheet to the QLabel to add a border
# You can adjust the border width, style, and color as preferred
                            self.label.setStyleSheet("QLabel { border: 2px solid black; }")
#                            self.label.setFixedSize(element['width'], element['height'])
                            print('completed browser output 1')
                            buf = io.BytesIO()
                            image["img"].save(buf, format='PNG')
                            byte_im = buf.getvalue()
                            
                
                            encoded_string = ""
                            with open(resource_path(str(index) + "browser.png"), "rb") as image_file:
                                encoded_string = base64.b64encode(image_file.read())

                            itm_output = blip_caption(encoded_string, description, self.user_key, self.user_plan)
                            print(itm_output)
                            score = float(itm_output.split("probability of ")[1].split(".\n")[0])
                            print('completed browser output 3')
                            print(itm_output)
                            self.label2.setText("Semantic Step browser:" + description + " Probability:" + str(score))
                            self.label2.setFixedSize(1280,48)
                            print('The browser image and text is matched with a probability of %.4f'%score)
                            score_data = {"href": image["href"], "text": image["text"], "css": image["css"], "score": float(score), "x": -1, "y": -1, "width": -1 , "height": -1, "cropped": open(resource_path(str(index) + "browser.png"), "rb"), "index": index, "browser": "true"}
                            browser_scores.append(score_data)

                        index += 1

                        # Process image...
     #                   image.show()

                    for score in browser_scores:
                        itm_score = score.get("score", 0)
                        if itm_score > browser_max:
                            print("Semantic scrape VALID")
                            browser_max = itm_score
                            browser_new_max = score
                            print(score)
                        if itm_score > 0.90:
                            browser_matches.append(score)
                            total_matches.append(score)
                    print(browser_new_max)
                    print(browser_new_max["index"])
                    print("browser max")
                    all_scores = []
                    index = 0
                    for match in total_matches:
                        if graph["nodes"][node_id]["custom"]["Target"] == "span":
                            all_scores.append(match["text"])
                        if graph["nodes"][node_id]["custom"]["Target"] == "link text":
                            all_scores.append(match["text"])
                        if graph["nodes"][node_id]["custom"]["Target"] == "link destination":
                            all_scores.append(match["href"])
                        if graph["nodes"][node_id]["custom"]["Target"] == "article":
                            all_scores.append(match["text"])
                        if graph["nodes"][node_id]["custom"]["Target"] == "Twitter posts":
                            all_scores.append(match["text"])
                            
                        #self.global_variables["Scrape_" + index  + "_" + str(len(self.global_variables))] = match["text"]
                    if graph["nodes"][node_id]["custom"]["Mode"] == "First":

                        self.global_variables["Scrape"  + graph["nodes"][node_id]["name"].split("Semantic Scrape ")[1]] = all_scores[0]
                    if graph["nodes"][node_id]["custom"]["Mode"] == "Last":
                        self.global_variables["Scrape"   + graph["nodes"][node_id]["name"].split("Semantic Scrape ")[1]] = all_scores[len(all_scores)-1]
                    if graph["nodes"][node_id]["custom"]["Mode"] == "All":
                        self.global_variables["Scrape"   + graph["nodes"][node_id]["name"].split("Semantic Scrape ")[1]] = ",".join(all_scores)
                    self.global_variables["Scrape_all"  + str(len(self.global_variables))] = all_scores
                    print(all_scores)
                    print("Global Variables")
                    print(self.global_variables)
                    print(graph["nodes"][node_id]["custom"]["Target"] )

            if x["type"] == "Left Mouse Click":
                img = []
                image = []
               # self.chatInstance.showMinimized()
                print("Left Mouse Click WHAT")
                print(graph["nodes"][node_id]["custom"])
                import time
                #self.thread_signals.notifications.emit("", "Clicking " +graph["nodes"][node_id]["custom"]["semanticTarget"], "icon.png", None, "No")

                target = ""
                if "Target In English"  in graph["nodes"][node_id]["custom"]:
                    print("Target In English")
                    print(graph["nodes"][node_id]["custom"]["Target In English"])

                    target = graph["nodes"][node_id]["custom"]["Target In English"]
                if "semanticTarget" in graph["nodes"][node_id]["custom"]:
                    print(graph["nodes"][node_id]["custom"]["semanticTarget"])
                    print("Semantic Click")

                    target = graph["nodes"][node_id]["custom"]["semanticTarget"]
                    print(self.global_variables)
                #replace any strings that include {{name}} with self.global_variables["name"] in body 
                for key, value in self.global_variables.items():
                    print(self.global_variables[key])
                    print(key)
                    if key in target:
                        target = target.replace("{{" + key + "}}", str(value))
                if "Mode" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Mode"] == "Browser":
                    print("browser mode")
                    self.browserClick(target)
                else:
                    print('finished globals')
                    automation_target = ""
                    if "Automation Input" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Automation Input"] != "" and len(graph["nodes"][node_id]["custom"]["Automation Input"]) > 0:
                        automation_target = self.global_variables[graph["nodes"][node_id]["custom"]["Automation Input"]]
                    print('finished automation')
                    new_max = self.semanticSearch(target +  automation_target, 100,100)
                    print(new_max)
                    if 'cropped' in new_max:
                      #  new_max['cropped'].save(resource_path("combined_image.png"))

                        self.label.show()
                        self.label2.show()

                        self.label.setPixmap(pil2pixmap(new_max['cropped']))
#                        self.label2.setText("Winning Semantic Step:" + x['semanticTarget'][0] + " Probability:" +  '%.4f'%new_max["score"])
                        #save the pixmap to disk
                        self.label.setFixedSize(new_max['width'], new_max['height'])
                        time.sleep(1)
                        self.label.hide()
                        self.label2.hide()

       #                 self.label2.setFixedSize(960,48)
      #  cropped     = raw_image.crop((new_max['x'], new_max['y'], new_max['x'] + new_max['width'], new_max['y'] + new_max['height']))
       # cropped    .show()
                    print(new_max)
                    print("winning click")
          #          self.chatInstance.showMinimized()
                    import pyautogui
                    if new_max["browser"] == "false":
                        if "Click Type" in graph["nodes"][node_id]["custom"]:
                            print("CUSTOM CLICK")
                            if "Double" in graph["nodes"][node_id]["custom"]["Click Type"]:

                                pyautogui.click(new_max['x'], new_max['y'])
                            elif graph["nodes"][node_id]["custom"]["Click Type"] == "Single Left Click":
                                pyautogui.click(new_max['x'], new_max['y'] )
                            elif graph["nodes"][node_id]["custom"]["Click Type"] == "Single Right Click":
                                pyautogui.click(new_max['x'], new_max['y'], button='right')
                            elif graph["nodes"][node_id]["custom"]["Click Type"] == "Drag":
                                pyautogui.mouseDown(new_max['x'], new_max['y'])
                                pyautogui.moveTo(new_max['x'] + 10, new_max['y'])
                        else:
                            pyautogui.click(new_max['x'], new_max['y'])

                    else:
                        self.sendMessageRTC("semanticClick" + str(new_max['index']))
               #     time.sleep(2)
                    #self.chatInstance.toggle_chat_display()

                    #self.chatInstance.showNormal()
           #         self.label.hide()
            #        self.label2.hide()
                    #print("try click")
                    ##test = match(image, x["image"])
             #      # cv2.imshow('test2',cv2.cvtColor(x["image"], cv2.COLOR_BGR2RGB))
             #      # cv2.imshow('test1',cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
              #     # cv2.waitKey(1
                    #in_image = graph["nodes"][node_id]["custom"]["image"]
                    #match_res = cv2.matchTemplate(cv2.cvtColor(graph["nodes"][node_id]["custom"]["image"], cv2.COLOR_BGR2RGB), image, cv2.TM_SQDIFF_NORMED)
                    #min_val, max_val, min_loc, max_loc = cv2.minMaxLoc(match_res)
                    #print(min_loc)
                    #print(min_val)
                    #if min_val < 0.01:
                    #    print("Success")
                    #    #self.mouseclick(x["x"], x["y"])
                    #    pyautogui.mouseDown(graph["nodes"][node_id]["custom"]["Bounding Box X1"] + 25, graph["nodes"][node_id]["custom"]["Bounding Box Y1"] + 25)
                    #    print("Finish click")
                    #else:
                    #    img = pyautogui.screenshot()
                    #    image = cv2.cvtColor(np.array(img), cv2.COLOR_BGR2RGB)
                    #    best_loc = None
                    #    best_val = 1000
                    #    best_scale = 1.0
                    #    print("FAILED")
                    #	# loop over the scales of the image
                    #    for scale in np.linspace(0.5, 1.3, 25)[::-1]:
                    #        resized = cv2.resize(cv2.cvtColor(in_image, cv2.COLOR_BGR2RGB), (int(in_image.shape[1] * scale), int(in_image.shape[0] * scale)), interpolation = cv2.INTER_AREA)
                    #        ##cv2.imshow('image',image)
                    #        #cv2.waitKey(1)
                    #        match_res = cv2.matchTemplate(image, resized,  cv2.TM_SQDIFF_NORMED)
                    #        min_val, max_val, min_loc, max_loc = cv2.minMaxLoc(match_res)
                    #        if min_val < best_val:
                    #            best_val = min_val
                    #            best_loc = min_loc
                    #            best_scale = scale
                    #        print(str(best_val) + ":" +str(best_scale))
                    #        print(best_loc)
                    #
                    #    #cv2.rectangle(image, (best_loc[0], best_loc[1]), (best_loc[0] + 50, best_loc[1] + 50), (255,0,0), 2)
                    #    #cv2.imshow('resized',image
                    #    #cv2.waitKey(1)
                    ##print("SUPER MATCH")
                    #    if best_val  < .05:
                    #        #self.mouseclick(best_loc[0] + self.half_small*best_scale, best_loc[1] + self.half_small*best_scale)
                    #        print("SUPER MATCH" + str(best_val))
                    #        print(str(best_loc[0]) )
                    #        print(str(best_loc[1]) )
                    #        if "Windows" in self.platform:
                    #            pyautogui.mouseDown(math.floor(best_loc[0] + 25*best_scale), math.floor(best_loc[1] + 25*best_scale))
                    #        else:
                    #            pyautogui.mouseDown(math.floor(best_loc[0]/2 + 25*best_scale/2), math.floor(best_loc[1]/2 + 25*best_scale/2))
                    #     #  offsetX = x["x"] - max_loc[0] + 25
                    #     #  offsetY = x["y"] - max_loc[1] + 25
            if x["type"] == "keypressup":
                if "shift" in str(x["key"]).split("KEY_")[1].lower():
                    pyautogui.keyUp("shfift")
                    self.shift_down = False
            if x["type"] == "general":
                thread_signals.chat.emit(graph["nodes"][node_id]["custom"]["prompt"])
            if x["type"] == "keypress" or x["type"] == "keypres_manual":
                print("keypress")
                if "String" in graph["nodes"][node_id]["custom"]:
                    x["key"] = graph["nodes"][node_id]["custom"]["String"]
                    x["saved"] = ""
                
                    self.thread_signals.notifications.emit("Keypress", "Typing " +  x["key"], "icon.png", None, "No")
                if "recording" in graph["nodes"][node_id]["custom"]:
                    import ast
                    events = ast.literal_eval(graph["nodes"][node_id]["custom"]["events"])
                    keys = ast.literal_eval(graph["nodes"][node_id]["custom"]["recording"])
                    
                    if "GPT-4 Mode" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["GPT-4 Mode"] == "True":
                        print("GPT-4 Mode")
                        #convert the array of keys like 'KeyboardKey.KEY_S', 'KeyboardKey.KEY_S', 'KeyboardKey.KEY_O', 'KeyboardKey.KEY_O' into a string prompt
                        prompt = ""

                        for key, event in zip(keys, events):
                            if "DOWN" in str(event):
                                prompt += key.replace("KeyboardKey.KEY_","").lower()

                        log = [{"role": "system", "content": "You are a helpfull assistant that only generates exactly the text the user requests. Do not explain the text. SImply generate exactly what the user requests. For example, if a user asks to generate a social post about ai automation, you would write 'Our business keeps growing faster thanks to AI automation #AI #aiautomation'"}]
                        Automation_input = ""
                        Webhook_input = ""
                        
                        if "Automation Input" in graph["nodes"][node_id]["custom"] and graph["nodes"][node_id]["custom"]["Automation Input"] != "None" and len(graph["nodes"][node_id]["custom"]["Automation Input"]) > 0:
                            Automation_input = self.global_variables[graph["nodes"][node_id]["custom"]["Automation Input"]]
                        log.append({"role": "user", "content": prompt.replace("space", " ") + Automation_input + Webhook_input})
                        data = {
                            "model": "gemma-3-4b-it",  # This is what the llama.cpp server expects
                            "messages": log,
                            "temperature": 0.7,
                            "max_tokens": 5000,
                            "stream": True  # Enable streaming to match original behavior
                        }
                        
                        print(data)
                        
                        # Call local server instead of remote API
                        response = requests.post("http://localhost:8000/v1/chat/completions", 
                                                headers=headers, 
                                                data=json.dumps(data), 
                                                stream=True)
                        
                        import pyautogui
                        if response.status_code == 200:
                            items = []
                            total = ""
                            for chunk in response.iter_lines():
                                try:
                                    if chunk:
                                        # Parse SSE format from llama.cpp server
                                        chunk_str = chunk.decode('utf-8')
                                        if chunk_str.startswith('data: '):
                                            chunk_data = chunk_str[6:]  # Remove 'data: ' prefix
                                            if chunk_data.strip() == '[DONE]':
                                                break
                                            
                                            chunk_json = json.loads(chunk_data)
                                            if 'choices' in chunk_json and len(chunk_json['choices']) > 0:
                                                delta = chunk_json['choices'][0].get('delta', {})
                                                if 'content' in delta:
                                                    content = delta['content']
                                                    total += content
                                                    print(content, end='', flush=True)
                                except Exception as e:
                                    # If streaming fails, fall back to non-streaming
                                    print(f"Streaming error: {e}")
                                    break
                                
                            # If streaming didn't work or total is empty, try non-streaming
                            if not total:
                                data["stream"] = False
                                response = requests.post("http://localhost:8000/v1/chat/completions", 
                                                        headers=headers, 
                                                        data=json.dumps(data))
                                if response.status_code == 200:
                                    result = response.json()
                                    total = result["choices"][0]["message"]["content"]
                                    print(total)
                                else:
                                    print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                                    total = "Error communicating with local LLM server"
                            
                            print()  # New line after streaming output
                            print(total)
                        else:
                            print(f"Error response from llama.cpp server: {response.status_code} {response.text}")
                            total = "Error communicating with local LLM server"
                            pyautogui.write(total, interval=0.01)
                            self.thread_signals.notifications.emit("GPT-4", total, "icon.png", None, "No")
                            print(self.global_variables)
                    else:

                        for key, events in zip(keys, events):
                            x["key"] = key.replace("KeyboardKey.KEY_","")
                            x["event"] = events
                            if "Current Directory" in str(x["saved"]):
                                pyautogui.write(resource_path(''), interval=0.05)
                            if "resource_path" in str(x["key"]).lower():
                                if "DOWN" in str(x["event"]):

                                    pyautogui.write(resource_path(''), interval=0.05)
                            if "backspace" in str(x["key"]).lower():
                                if "DOWN" in str(x["event"]):

                                    pyautogui.press('backspace')
                            elif "space" in str(x["key"]).lower():
                                if "DOWN" in str(x["event"]):

                                    pyautogui.write(" ", interval=0.05)
                            elif "period" in str(x["key"]).lower():
                                if "DOWN" in str(x["event"]):

                                    pyautogui.write(".", interval=0.05)
                            elif "slash" in str(x["key"]).lower():
                                if "DOWN" in str(x["event"]):

                                    pyautogui.write("/", interval=0.05)

                            elif "return" in str(x["key"]).lower():
                                if "DOWN" in str(x["event"]):
                                    pyautogui.press('enter')
                            elif "shift" in str(x["key"]).lower():
                                if "DOWN" in str(x["event"]):
                                   pyautogui.keyDown("shift")

                                   self.shift_down = True
                                else:   
                                    pyautogui.keyUp("shift")
                                    self.shift_down = False
                            elif "ctrl" in str(x["key"]).lower():
                                if "DOWN" in str(x["event"]):
                                    pyautogui.keyDown("ctrl")
                                    self.ctrl_down = True
                                else:
                                    pyautogui.keyUp("ctrl")
                                    self.ctrl_down = False
                            elif "alt" in str(x["key"]).lower():
                                if "DOWN" in str(x["event"]):
                                    pyautogui.keyDown("alt")
                                    self.alt_down = True
                                else:
                                    pyautogui.keyUp("alt")
                                    self.alt_down = False
                            elif "tab" in str(x["key"]).lower():
                                if "DOWN" in str(x["event"]):
                                    pyautogui.keyDown("tab")
                                else:
                                    pyautogui.keyUp("tab")
                            elif "escape" in str(x["key"]).lower():
                                print("escape")
                            elif "enter" in str(x["key"]).lower():
                                if "DOWN" in str(x["event"]):
                                    pyautogui.press('enter')

                            else:
                                if self.shift_down:
                                    if "{{folder}}" in str(x["key"]):
                                        if "DOWN" in str(x["event"]):
                                            pyautogui.write(folder, interval=0.01)
                                            pyautogui.write("\\", interval=0.01)
                                            pyautogui.write(str(x["key"]).upper().replace("{{FOLDER}}", ""), interval=0.01)

                                    else:
                                        if "DOWN" in str(x["event"]):

                                            clipboard.copy(str(x["key"]).upper())
                                            pyautogui.hotkey('ctrl', 'v')
                                        #pyautogui.write(str(x["key"]).upper(), interval=0.01)                        
                                else:
                                    if "{{folder}}" in str(x["key"]):
                                        if "DOWN" in str(x["event"]):
                                            pyautogui.write(folder, interval=0.01)
                                            pyautogui.write("\\", interval=0.01)
                                            pyautogui.write(str(x["key"]).lower().replace("{{folder}}", ""), interval=0.01)

                                    else:
                                        if "DOWN" in str(x["event"]):

                                            clipboard.copy(str(x["key"].lower()))
                                            pyautogui.hotkey('ctrl', 'v')

                                 #   pyautogui.write(str(x["key"]).lower(), interval=0.01)

                if "event" not in x or "DOWN" in str(x["event"]):
                    if "Current Directory" in str(x["saved"]):
                        pyautogui.write(resource_path(''), interval=0.05)
                    if "resource_path" in str(x["key"]).lower():
                        pyautogui.write(resource_path(''), interval=0.05)
                    if "backspace" in str(x["key"]).lower():
                        pyautogui.press('backspace')
                    elif "space" in str(x["key"]).lower():
                        pyautogui.write(" ", interval=0.05)
                    elif "period" in str(x["key"]).lower():
                        pyautogui.write(".", interval=0.05)
                    elif "slash" in str(x["key"]).lower():
                        pyautogui.write("/", interval=0.05)
                    elif "return" in str(x["key"]).lower():
                        pyautogui.press('enter')
                    elif "shift" in str(x["key"]).lower():
                        pyautogui.keyDown("shift")
                        self.shift_down = True
                    elif "ctrl" in str(x["key"]).lower():
                        pyautogui.keyDown("ctrl")
                        self.ctrl_down = True
                    elif "alt" in str(x["key"]).lower():
                        pyautogui.keyDown("alt")
                        self.alt_down = True
                    elif "tab" in str(x["key"]).lower():
                        pyautogui.keyDown("tab")
                    elif "enter" in str(x["key"]).lower():
                        pyautogui.press('enter')
                    elif "escape" in str(x["key"]).lower():
                        pyautogui.press('escape')
                    else:
                        if self.shift_down:
                            if "{{folder}}" in str(x["key"]):
                                pyautogui.write(folder, interval=0.01)
                                pyautogui.write("\\", interval=0.01)
                                pyautogui.write(str(x["key"]).upper().replace("{{FOLDER}}", ""), interval=0.01)
                                
                            else:
                                
                                clipboard.copy(str(x["key"]).upper())
                                pyautogui.hotkey('ctrl', 'v')
                                #pyautogui.write(str(x["key"]).upper(), interval=0.01)                        
                        else:
                            if "{{folder}}" in str(x["key"]):
                                pyautogui.write(folder, interval=0.01)
                                pyautogui.write("\\", interval=0.01)
                                pyautogui.write(str(x["key"]).lower().replace("{{folder}}", ""), interval=0.01)
                                
                            else:
                                
                                clipboard.copy(str(x["key"]))
                                pyautogui.hotkey('ctrl', 'v')
                                
                             #   pyautogui.write(str(x["key"]).lower(), interval=0.01)
            found_node = False
            print("PATH")
            print(path)
            print("what?")
            for key, node in graph["nodes"].items():
                if node_id == key:
                    if "connections" in graph:
                        for connections in graph["connections"]:
                            print("connections")
                            print(connections)
                            print(connections["out"][1])
                            if key == connections["out"][0] and ( path in connections["out"][1] or "out"  == connections["out"][1]):
                                found_node = True
                                print(connections["out"][1])
                                try:
                                    print("new Node in graph found")
                                    print(connections["in"][0])
                                    self.runNode(graph, connections["in"][0], thread_signals, direct, manual, folder)
                                except Exception as e:
                                    import traceback
                                    print(traceback.format_exc())
            if not found_node:
                print("No Node Found")
                thread_signals.recording.emit()
            if not found_node and direct == True:
                thread_signals.recording.emit()
                return

            
           # virtual_desktop_accessor = ctypes.WinDLL("VirtualDesktopAccessor.dll")
           # virtual_desktop_accessor.GoToDesktopNumber(0)


    def runRecording(self, thread_signals):
        graph_nodes = self.serialize_session()
        time.sleep(2)
        thread_signals.recording_start.emit()
        for key, node in graph_nodes["nodes"].items():
            if "Start" in node["name"]:
                #self.runNode(graph_nodes, key)
                import threading
                t0 = threading.Thread(target = self.runNode, kwargs={"graph":graph_nodes, "node_id": key, "thread_signals": thread_signals, "direct": False, "manual": True})
                t0.start()
                break

        if False:
      #  print history
            for index,y in enumerate(self.history):
                x = y
                if isinstance(x, str):
                     x = json.loads(y)
            #print(y)s

                for key, value in x.items()  :
                    if key == "image":
                        x["image"] = np.asarray(x["image"],dtype = "uint8")
                if x["type"] == "Move Mouse":
                    pyautogui.moveTo(x["x"], x["y"])
                if x["type"] == "Left Mouse Lift":
                    pyautogui.mouseUp()
                if x["type"] == "Left Mouse Click":
                    img = []
                    image = []
                    if "Windows" in self.platform:
                        img = pyautogui.screenshot(region=(x["x"]-self.half_small,x["y"]-self.half_small, self.size_small, self.size_small))
                        image =  cv2.cvtColor(np.array(img), cv2.COLOR_BGR2RGB)
                    else:
                        img = pyautogui.screenshot(region=(x["x"]*2-self.half_small,x["y"]*2-self.half_small, self.size_small, self.size_small))
                        image =  cv2.cvtColor(np.array(img), cv2.COLOR_BGR2RGB)

                    print(x)
               # test = match(image, x["image"])
                 #   cv2.imshow('test2',cv2.cvtColor(x["image"], cv2.COLOR_BGR2RGB))
                 #   cv2.imshow('test1',cv2.cvtColor(image, cv2.COLOR_BGR2RGB))

                  #  cv2.waitKey(1)

                    match_res = cv2.matchTemplate(cv2.cvtColor(x["image"], cv2.COLOR_BGR2RGB), image, cv2.TM_SQDIFF_NORMED)
                    min_val, max_val, min_loc, max_loc = cv2.minMaxLoc(match_res)
                    print(min_loc)
                    print(min_val)
                    if min_val < 0.01:
                        print("Success")
                        print(str(x["x"]) + "," + str(x["y"]))
                        #self.mouseclick(x["x"], x["y"])
                        pyautogui.mouseDown(x["x"], x["y"])
                        print("Finish click")
                    else:
                        img = pyautogui.screenshot()
                        image = cv2.cvtColor(np.array(img), cv2.COLOR_BGR2RGB)
                        best_loc = None
                        best_val = 1000
                        best_scale = 1.0
                        print("FAILED")
                    	# loop over the scales of the image
                        for scale in np.linspace(0.5, 1.3, 25)[::-1]:
                            resized = cv2.resize(cv2.cvtColor(x["image"], cv2.COLOR_BGR2RGB), (int(x["image"].shape[1] * scale), int(x["image"].shape[0] * scale)), interpolation = cv2.INTER_AREA)

                            ##cv2.imshow('image',image)
                            #cv2.waitKey(1)
                            match_res = cv2.matchTemplate(image, resized,  cv2.TM_SQDIFF_NORMED)
                            min_val, max_val, min_loc, max_loc = cv2.minMaxLoc(match_res)
                            if min_val < best_val:
                                best_val = min_val
                                best_loc = min_loc
                                best_scale = scale
                            print(str(best_val) + ":" +str(best_scale))
                            print(best_loc)

                        #cv2.rectangle(image, (best_loc[0], best_loc[1]), (best_loc[0] + 50, best_loc[1] + 50), (255,0,0), 2)
                        #cv2.imshow('resized',image)

                        #cv2.waitKey(1)
                    #print("SUPER MATCH")
                        if best_val  < .05:
                            #self.mouseclick(best_loc[0] + self.half_small*best_scale, best_loc[1] + self.half_small*best_scale)
                            print("SUPER MATCH" + str(best_val))
                            print(str(best_loc[0]) )
                            print(str(best_loc[1]) )
                            if "Windows" in self.platform:
                                pyautogui.mouseDown(math.floor(best_loc[0] + 25*best_scale), math.floor(best_loc[1] + 25*best_scale))
                            else:
                                pyautogui.mouseDown(math.floor(best_loc[0]/2 + 25*best_scale/2), math.floor(best_loc[1]/2 + 25*best_scale/2))
                         #  offsetX = x["x"] - max_loc[0] + 25
                         #  offsetY = x["y"] - max_loc[1] + 25
                if x["type"] == "keypressup":
                    if "shift" in str(x["key"]).split("KEY_")[1].lower():
                        pyautogui.keyUp("shfift")
                        self.shift_down = False
                if x["type"] == "keypress":
                    if True:
                        if "backspace" in str(x["key"]).split("KEY_")[1].lower():
                            pyautogui.press('backspace')
                        elif "space" in str(x["key"]).split("KEY_")[1].lower():
                            pyautogui.write(" ", interval=0.05)
                        elif "period" in str(x["key"]).split("KEY_")[1].lower():
                            pyautogui.write(".", interval=0.05)
                        elif "slash" in str(x["key"]).split("KEY_")[1].lower():
                            pyautogui.write("/", interval=0.05)
                        elif "return" in str(x["key"]).split("KEY_")[1].lower():
                            pyautogui.press('enter')
                        elif "shift" in str(x["key"]).split("KEY_")[1].lower():
                            pyautogui.keyDown("shift")
                            self.shift_down = True
                        else:
                            if self.shift_down:
                                pyautogui.write(str(x["key"]).split("KEY_")[1].upper(), interval=0.05)
                            else:
                                pyautogui.write(str(x["key"]).split("KEY_")[1].lower(), interval=0.05)

    def shape_selection(self, event, x, y, flags, param):
      # grab references to the global variables
      global ref_point, cropping, screenshot, global_variables, nodes

      # if the left mouse button was clicked, record the starting
      # (x, y) coordinates and indicate that cropping is being
      # performed
      if event == cv2.EVENT_LBUTTONDOWN:
        ref_point = [(x, y)]
        cropping = True

      # check to see if the left mouse button was released
      elif event == cv2.EVENT_LBUTTONUP:
        # record the ending (x, y) coordinates and indicate that
        # the cropping operation is finished
        ref_point.append((x, y))
        cropping = False

        # draw a rectangle around the region of interest
        cv2.rectangle(screenshot, ref_point[0], ref_point[1], (0, 255, 0), 2)
        cv2.imshow("OCR", screenshot)
        x = {"type":"OCR", "bounding_box":ref_point}
        nodes.append(self.create_node('nodes.basic.BasicNodeA', name="OCR " + str(len(nodes)), data=x))#, color= "#FFFFFF"
        nodes[len(nodes)-1].create_property('Bounding Box X1', ref_point[0][0], widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Bounding Box Y1', ref_point[0][1], widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Bounding Box X2', ref_point[1][0], widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Bounding Box Y2', ref_point[1][1], widget_type=NODE_PROP_QLINEEDIT)

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        this_path = os.path.dirname(os.path.abspath(__file__))
        icon = os.path.join(this_path, 'examples', 'OCRScraper.png')


        screenshot = cv2.cvtColor(np.array(screenshot), cv2.COLOR_BGR2RGB)

        # draw a rectangle around the region of interest

        thresh = 255 - cv2.threshold(cv2.cvtColor(screenshot, cv2.COLOR_BGR2GRAY), 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)[1]

        xp,yp,w,h = ref_point[0][0], ref_point[0][1], ref_point[1][0]-ref_point[0][0], ref_point[1][1]-ref_points[0][1]
        ROI = thresh[ref_point[0][1]:ref_point[1][1],ref_point[0][0]:ref_point[1][0]]

        data = pytesseract.image_to_string(ROI, lang='eng',config='--psm 6')
        print("results")
        print(data)
        self.global_variables["OCR" + str(len(self.global_variables))] = (data)
        cv2.imshow("thresh",thresh)
        cv2.imshow("ROI",ROI)
        #global_variables.append("OCR_" + str(len(global_variables)))

        nodes[len(nodes)-1].set_icon(icon)

        self.auto_layout_nodes()

        cv2.waitKey(1)

    # crate a backdrop node and wrap it around
    # "custom port node" and "group node".
    # fit node selection to the viewer.
        self.fit_to_selection()
        time.sleep(1)
        
        self.QMainWindow.showMaximized()

    def click_coordinates(self, event,x,y,flags,params):
      # grab references to the global variables
      global ref_point, screenshot, global_variables

      self.label2.hide()
      if event == cv2.EVENT_LBUTTONDOWN:
        ref_point = [(x-25, y-25),(x+25,y+25)]

        # draw a rectangle around the region of interest
        cut_image=screenshot[y-25:y+25,x-25:x+25]
        x = {"type":"Left Mouse Click",  "x": x, "y": y,"image":np.array(cut_image)}
        nodes.append(graph.create_node('nodes.basic.BasicNodeA', name="CLICK " + str(len(nodes)), data=x))#, color= "#FFFFFF"
        nodes[len(nodes)-1].create_property('Bounding Box X1', ref_point[0][0], widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Bounding Box Y1', ref_point[0][1], widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Bounding Box X2', ref_point[1][0], widget_type=NODE_PROP_QLINEEDIT)
        nodes[len(nodes)-1].create_property('Bounding Box Y2', ref_point[1][1], widget_type=NODE_PROP_QLINEEDIT)

        nodes[len(nodes)-1].create_property('Data', json.dumps(x, cls=NumpyEncoder))
        this_path = os.path.dirname(os.path.abspath(__file__))
        icon = os.path.join(this_path, 'examples', 'OCRScraper.png')
        global_variables.append("CLICK_" + str(len(global_variables)))
        nodes[len(nodes)-1].set_icon(icon)
        self.auto_layout_nodes()
        cv2.waitKey(1)
        self.fit_to_selection()
        time.sleep(1)
        
        self.QMainWindow.showMaximized()

    def eventRecord(self, event):
   # print(len(history))
        print("EVENT")
        print(event)
        if "MOVE" in str(event.event):
            self.mouse_counter += 1
            if self.mouse_counter%6 == 0:
                self.history.append({"type":"Move Mouse", "x": int(event.x), "y": int(event.y)})
        if "CLICK" in str(event.event) and "DOWN" in str(event.direction):
            

            element = {'x':event.x, 'y':event.y, 'height':200, 'width':200}
            raw_image = pyautogui.screenshot()
            cut_image  = raw_image
            
            draw = ImageDraw.Draw(raw_image)
            draw.rectangle([int(event.x)-100, int(event.y)-50, int(event.x)+100, int(event.y)+50], outline ="red", width=10)
                #self.chatWindow(self.chatInstance,"Semantic Step:" + caption + " Probability:" +  '%.4f'%itm_score, image
            if "Window" in self.platform :
                self.history.append({"type":"Left Mouse Click", "image": cut_image, "Bounding Box X1": int(event.x) - 25, "Bounding Box Y1": int(event.y) -25, "Bounding Box X2": int(event.x) + 25, "Bounding Box Y2": int(event.y) + 25, "semanticTarget": "caption"})
            else:
                self.history.append({"type":"Left Mouse Click", "image": cut_image, "Bounding Box X1": int(event.x) - 25, "Bounding Box Y1": int(event.y) - 25, "Bounding Box X2": int(event.x) + 25, "Bounding Box Y2": int(event.y) + 25, "semanticTarget": "caption"})

        if "CLICK" in str(event.event) and "UP" in str(event.direction):
            self.history.append({"type":"Left Mouse Lift", "x": int(event.x), "y": int(event.y)})
    #    if stop == False:
        if "KeyboardEvent" in str(event.event) and "UP" in str(event.event):
            self.history.append({"type":"keypressup", "event": event.event, "key": str(event.keyboard_key)})
        if "KeyboardEvent" in str(event.event) and "DOWN" in str(event.event):
            if "PRINT" in str(event.keyboard_key):
                img = pyautogui.screenshot()
                self.history.append({"type":"scrape",  "image": img})

                #np.array(pyautogui.screenshot(region=(event.x*2 - 25 ,event.y*2 - 25, 50, 50)))
               # self.history.append(json.dumps({"type":"keypressup", "key": str(event.keyboard_key)}, cls=NumpyEncoder))
            if "KEY_ESCAPE" in str(event.keyboard_key):
            #stopRecording()
            #reco

                print("STOP")

                self.stopAction.trigger()
            #record_tasks.put("stop")
            #redrawHistory()             #stopRecording()
            #reco
                #newWindow = tk.Toplevel(window)

                #newWindow.title("Edit Event" + n)
                #step = json.loads(history[int(n)])
                #newWindow.geometry("1000x800")
                #img = pyautogui.screenshot()
                #image = cv2.cvtColor(np.array(img), cv2.COLOR_BGR2RGB)

            #window.event_generate("<<stop>>", when="tail", state=123)
            #record_tasks.put("stop")
            #redrawHistory()
            else:
                self.history.append({"type":"keypress",  "event": event.event, "key": str(event.keyboard_key)})

    def openOneCheat(self,f):
        self.load_session(f)

    def stopRecording(self):
        print("start stop 1")
        

        print(self.history)
        self.drawHistory(self.history)

        self.started = False
        self.QMainWindow.showMaximized()
        self.label2.hide()
        self.keyboard_listener.stop()
        self.mouse_listener.stop()
#        self.keyboard_listener.join()
        self.mouse_listener.join()
        
        
        self.auto_layout_nodes()
        self.fit_to_selection()
        
        self.viewer().clear_key_state()
        self.viewer().clearFocus()
        self.QMainWindow.showMaximized()

        #import threading
        #t0 = threading.Thread(target = self.stopRecording_thread, kwargs={"history":self.history})
        #t0.start()
    def stopRecording_thread(self, history):

        self.drawHistory(history)
        self.QMainWindow.showMaximized()

    def trainModel(self):
       # import torch
       # print(torch.cuda.is_available())sdfsdf
       # model = core.Model.load( "model_weights.pth",  ["links","buttons","sliders","dropdowns","switches","menus","textareas","textfields","checkboxes","radioboxes","images","text"])

        #custom_transforms = transforms.Compose([transforms.ToPILImage(),transforms.Resize(900), transforms.RandomHorizontalFlip(0.5), transforms.ColorJitter(saturation=0.2),transforms.ToTensor(),utils.normalize_transform(), ])

        utils.xml_to_csv('train/', 'train_labels.csv')
        utils.xml_to_csv('test/', 'val_labels.csv')
        import csv
       # with open('train_labels.csv', 'r') as file:
       #     # open the file in the write mode
       #
       #     reader = csv.reader(file)
       #     counter = 0
       #     for row in reader:
       #         counter += 1
       #         if "xmin" != row[4] and (int(row[4]) >= int(row[6]) or int(row[5]) >= int(row[7])):
       #             print(row[4], row[6], row[5], row[7])
       #             print(row[0].split(".jpg")[0] + ".xml")
       #             print(row[0])
       #             import os
       #             try:
       #                 os.remove('train/' + row[0].split(".jpg")[0] + ".xml")
       #
       #                 os.remove( 'train/' + row[0])
       #             except:
       #                 print("error")
#
       # with open('val_labels.csv', 'r') as file:
       #     # open the file in the write mode
       #
       #     reader = csv.reader(file)
       #     counter = 0
       #     for row in reader:
       #         counter += 1
       #         if "xmin" != row[4] and (int(row[4]) >= int(row[6]) or int(row[5]) >= int(row[7])):
       #             print(row[4], row[6], row[5], row[7])
       #             print(row[0].split(".jpg")[0] + ".xml")
       #             print(row[0])
       #             import os
       #             try:
       #                 os.remove('test/' + row[0].split(".jpg")[0] + ".xml")
       #
       #                 os.remove('test/' + row[0])
       #             except:
       #                 print("error")
# Defin#e custom transforms to apply to your datasetcd
#
# Pass in a CSV file instead of XML files for faster Dataset initialization speeds
        dataset = core.Dataset('train_labels.csv', 'train/')
        valid_dataset = []
        valid_valid = []
        removes = []

        val_dataset = core.Dataset('val_labels.csv', 'test/')  # Validation dataset for training

# Create your own DataLoader with custom options
        loader = core.DataLoader(val_dataset, batch_size=2, shuffle=True)

# Create your own DataLoader with custom options

# Use MobileNet instead of the default ResNet
        model = core.Model(["button", "field", "heading", "iframe", "image", "label", "link", "text"], model_name='fasterrcnn_mobilenet_v3_large_fpn')
        losses = model.fit(loader, val_dataset, epochs=500, learning_rate=0.005, verbose=True)

   #     plt.plot(losses)  # Visualize loss throughout training
 #       plt.show()

        model.save('model_weights.pth')  # Save model to a file
        image = utils.read_image("test_image.png")
        predictions = model.predict(image)
        labels, boxes, scores = predictions
        #show_labeled_image(image, boxes, labels)
        #Train_dataset=core.Dataset( folder +"/train/", transform=custom_transforms)#L1
        #Test_dataset = core.Dataset( folder + "/test/")#L2
        #loader=core.DataLoader(Train_dataset, batch_size=1, shuffle=True)#L3
        #model = core.Model(["links","buttons","sliders","dropdowns","switches","menus","textareas","textfields","checkboxes","radioboxes","images","text"])
        #losses = model.fit(loader, Test_dataset, epochs=25, lr_step_size=3, learning_rate=0.001, verbose=True)#L5
        #plt.plot(losses)
        #plt.show()
        #model.save("cheat_model.pth")
       # model = core.Model.load("cheat_model.pth",  ["buttons","links","sliders","dropdowns","switches","menus","textareas","textfields","checkboxes","radioboxes","images","text"])
       # image = utils.read_image("test_image2.png")
       # predictions = model.predict(image)
       # labels, boxes, scores = predictions

    def white_label(self):
        # Save the session by moving the form.cheat file
        if "Agency" in self.user_plan or "Business" in self.user_plan or len(self.referral.split(",")) > 9:
            self.save_session(resource_path("logs/_internal/form.cheat"))

            # Create the application (required by PySide2)

            # Display input dialog to get the new filename from the user
            new_filename, ok = QInputDialog.getText(None, "Input", "Enter new filename for whitelabel.exe:")
            if "exe" not in new_filename:
                new_filename += ".exe"
            import shutil
            if ok and new_filename:
                new_file_path = os.path.join('logs', new_filename)
                shutil.move(resource_path("logs/whitelabel.exe"), resource_path(new_file_path))
                print(f'Renamed {"logs/whitelabel.exe"} to {new_file_path}')
            else:
                print('No filename entered. Whitelabel.exe not renamed.')
        else:
            #messagbox alert
            print(self.referral)
            self.thread_signals.notifications.emit("Error", "You do not have permission to white label. Please upgrade to the Agency or Business plan.", "icon.png", None, "No")
        # Properly close the application

    def playRecording(self):
        self.stop = False
        print('play')

        self.QMainWindow.showMinimized()
        print("finish minimize")
        self.runRecording(self.thread_signals)
    def stopActions(self):
        print("stop")
        #self.QMainWindow.showMinimized()
        self.stop = True


    def on_press(self, key):
        print(key)
        if self.started == True:
            if "PRINT" in str(key):
                img = pyautogui.screenshot()
                self.history.append({"type":"scrape",  "image": img})
                #np.array(pyautogui.screenshot(region=(event.x*2 - 25 ,event.y*2 - 25, 50, 50)))
               # self.history.append(json.dumps({"type":"keypressup", "key": str(key)}, cls=NumpyEncoder))
            if "esc" in str(key):
            #stopRecording()
            #reco
                print("STOP")
                #self.stopAction.trigger()
                self.stopRecording()
                #self.keyboard_listener.stop()
                #self.mouse_listener.stop()

                #self.keyboard_listener.join()
                #self.mouse_listener.join()
        
            #record_tasks.put("stop")
            #redrawHistory()             #stopRecording()
            #reco
                #newWindow = tk.Toplevel(window)
                #newWindow.title("Edit Event" + n)
                #step = json.loads(history[int(n)])
                #newWindow.geometry("1000x800")
                #img = pyautogui.screenshot()
                #image = cv2.cvtColor(np.array(img), cv2.COLOR_BGR2RGB)
            #window.event_generate("<<stop>>", when="tail", state=123)
            #record_tasks.put("stop")
            #redrawHistory()
            else:
                self.history.append({"type":"keypress",  "event": "DOWN", "key": str(key)})


    def on_release(self, key):
        if self.started == True:
            print('{0} released'.format(key))


    def on_move(self, x, y):
        if self.started == True:
            self.mouse_counter += 1
            if self.mouse_counter%6 == 0:
                self.history.append({"type":"Move Mouse", "x": int(x), "y": int(y)})


    def on_click(self, x, y, button, pressed):
        print("CLICK")
        if self.started == True:
            element = {'x':x, 'y':y, 'height':200, 'width':200}
            raw_image = pyautogui.screenshot()
            cut_image  = raw_image

            draw = ImageDraw.Draw(raw_image)
            draw.rectangle([int(x)-100, int(y)-50, int(x)+100, int(y)+50], outline ="red", width=10)
                #self.chatWindow(self.chatInstance,"Semantic Step:" + caption + " Probability:" +  '%.4f'%itm_score, image
            if "Window" in self.platform :
                self.history.append({"type":"Left Mouse Click", "image": cut_image, "Bounding Box X1": int(x) - 25, "Bounding Box Y1": int(y) -25, "Bounding Box X2": int(x) + 25, "Bounding Box Y2": int(y) + 25, "semanticTarget": "caption"})
            else:
                self.history.append({"type":"Left Mouse Click", "image": cut_image, "Bounding Box X1": int(x) - 25, "Bounding Box Y1": int(y) - 25, "Bounding Box X2": int(x) + 25, "Bounding Box Y2": int(y) + 25, "semanticTarget": "caption"})




    def on_scroll(self, x, y, dx, dy):
        if self.started == True:
            print('Scrolled {0}'.format((x, y)))



    def startRecording(self):
        global ref_point, cropping, screenshot, global_variables, nodes

        if self.started == False:
            self.started = True
            self.QMainWindow.showMinimized()
            self.counter = 0
            self.bnt = 0
            print('start recording')

          #  self.recorder = Recorder.record(self.eventRecord)
            print('started recording')
            time.sleep(2)
            screenshot = pyautogui.screenshot()

            buf = io.BytesIO()
            
            rgb_screenshot = screenshot.convert("RGB")
            rgb_screenshot.save(buf, format='JPEG')
            byte_im = buf.getvalue()
            print("CHECK URL")
            encoded_jpeg = base64.b64encode(io.BytesIO(byte_im).read()).decode('utf-8')
            caption = blip_caption_describe(encoded_jpeg, "What is the text of the URL of the browser exactly? Only return the URL exactly without describing it. Add https if it is not included", self.user_key, self.user_plan)
            print(caption)
            if "http" in caption:
                print("FOUND URL")
                self.global_variables["Current URL"] = caption
                graph_nodes = self.serialize_session()
                for key, node in graph_nodes["nodes"].items():
                    if "Start" in node["name"]:
                        node_id = key
                        start_node = self.get_node_by_id(node_id)
                        start_node.set_property('Initial Program', 'google-chrome'+  caption)
            print(self.global_variables)
            self.label2.setText("Agent Recorder Enabled. Use your keyboard and mouse to record actions. Hit ESC to stop recording.")
            self.label2.setFixedSize(1920,48)
            self.label2.move(0, 0)
            self.label2.show()
            from pynput import mouse, keyboard

            self.keyboard_listener = keyboard.Listener(
            on_press=self.on_press,
            on_release=self.on_release)
            self.mouse_listener = mouse.Listener(
            on_click=self.on_click,
            on_scroll=self.on_scroll,
            on_move=self.on_move)
            self.keyboard_listener.start()
            self.mouse_listener.start()

    def defExit(self):
        sys.exit()
    def saveCheat(self):
        name, save = QFileDialog.getSaveFileName(self.QMainWindow, 'Save Cheat',"", "CHEAT (*.cheat)")
        self.save_session(name)

    def defineMenu(self):
        self.myQMenuBar = self.QMainWindow.menuBar()
        exitMenu = self.myQMenuBar.addMenu('File')

        screen = QApplication.primaryScreen()
        print('Screen: %s' % screen.name())
        size = screen.size()
        print('Size: %d x %d' % (size.width(), size.height()))
        rect = screen.availableGeometry()
        print('Available: %d x %d' % (rect.width(), rect.height()))
        self.width = size.width()


        modelMenu = self.myQMenuBar.addMenu('Project Atlas')


        # scheduleAction = QAction('Schedule Recording', self)
        # scheduleAction.triggered.connect(self.startRecording)


        saveAction = QAction('Save', self)
        saveAction.triggered.connect(self.saveCheat)
        saveAction.setShortcut(QKeySequence("Ctrl+s"))


#loopAction = QAction('Loop', self)
#        loopAction.triggered.connect(self.loopRecording)
#        loopAction.setShortcut(QKeySequence("Ctrl+l"))

        newAction = QAction('New', self)
        newAction.triggered.connect(self.newCheat)
        newAction.setShortcut(QKeySequence("Ctrl+n"))



        scheduleAction = QAction('Schedule', self)
        scheduleAction.triggered.connect(self.scheduleCheat)
        scheduleAction.setShortcut(QKeySequence("Ctrl+c"))

       # trainAction = QAction('Train Custom UI Model', self)
       # trainAction.triggered.connect(self.launchAtlas)
       # trainAction.setShortcut(QKeySequence("Ctrl+t"))


        openAtlas2 = QAction('Launch Project Atlas', self)
        openAtlas2.triggered.connect(self.openAtlas)
        openAtlas2.setShortcut(QKeySequence("Alt+s"))


        playAction = QAction('Run', self)
        playAction.triggered.connect(self.playRecording)
        playAction.setShortcut(QKeySequence("Ctrl+r"))


        stopActions = QAction('Stop', self)
        stopActions.triggered.connect(self.stopActions)
        stopActions.setShortcut(QKeySequence("Ctrl+q"))

        openAction = QAction('Open', self)
        openAction.triggered.connect(self.openCheat)
        openAction.setShortcut(QKeySequence("Ctrl+o"))
     # File toolbar
        # Edit toolbar


        exitAction = QAction('Exit', self)
        exitAction.triggered.connect(self.defExit)
        exitAction.setShortcut(QKeySequence("Ctrl+x"))
     # File toolbar
        # Edit toolbar
        modelMenu.addAction(openAtlas2)

        exitMenu.addAction(newAction)

        exitMenu.addAction(openAction)
        exitMenu.addAction(stopActions)
        exitMenu.addAction(playAction)
        exitMenu.addAction(scheduleAction)
        exitMenu.addAction(saveAction)

        exitMenu.setStyleSheet("""
        QMenuBar {
            background-color: rgb(49,49,49);
            color: rgb(255,255,255);
            border: 1px solid #000;
        }

        QMenuBar::item {
            background-color: rgb(49,49,49);
            color: rgb(255,255,255);
        }

        QMenuBar::item::selected {
            background-color: rgb(30,30,30);
        }

        QMenu {
            background-color: rgb(49,49,49);
            color: rgb(255,255,255);
            border: 1px solid #000;
        }

        QMenu::item::selected {
            background-color: rgb(30,30,30);
        }
    """)

        exitAction = QAction('Exit', self)
        exitAction.triggered.connect(self.defExit)

        exitMenu.addAction(exitAction)
        self.myQMenuBar.show()
        self.pic = ['']*5

        self.pic[0] = QLabel(self.QMainWindow)
        #self.pic.setPixmap(QtGui.QPixmap("icon.png"))

        width = self.QMainWindow.frameGeometry().width()
        height =self.QMainWindow.frameGeometry().height()
        self.pic[0].resize(size.width()/8, size.height()/8)
        self.pic[0].setScaledContents(True)

        self.pic[0].move(size.width()-size.width()/8- 100,  size.height() - 3*size.height()/8)
        self.pic[0].show() # You were missing this.


        self.pic[1] = QLabel(self.QMainWindow)
        #self.pic.setPixmap(QtGui.QPixmap("icon.png"))

        width = self.QMainWindow.frameGeometry().width()
        height =self.QMainWindow.frameGeometry().height()
        self.pic[1].resize(size.width()/8, size.height()/8)
        self.pic[1].setScaledContents(True)

        self.pic[1].move(size.width()-size.width()/8 - 100,  size.height() - 4*size.height()/8)
        self.pic[1].show() # You were missing this.


        self.pic[2] = QLabel(self.QMainWindow)
        #self.pic.setPixmap(QtGui.QPixmap("icon.png"))

        width = self.QMainWindow.frameGeometry().width()
        height =self.QMainWindow.frameGeometry().height()
        self.pic[2].resize(size.width()/8, size.height()/8)
        self.pic[2].setScaledContents(True)

        self.pic[2].move(size.width()-size.width()/8 - 100,  size.height() - 5*size.height()/8)
        self.pic[2].show() # You were missing this.


        self.pic[3] = QLabel(self.QMainWindow)
        #self.pic.setPixmap(QtGui.QPixmap("icon.png"))

        width = self.QMainWindow.frameGeometry().width()
        height =self.QMainWindow.frameGeometry().height()
        self.pic[3].resize(size.width()/8, size.height()/8)
        self.pic[3].setScaledContents(True)

        self.pic[3].move(size.width()-size.width()/8 - 100,  size.height() - 6*size.height()/8)
        self.pic[3].show() # You were missing this.

        self.pic[4] = QLabel(self.QMainWindow)
        #self.pic.setPixmap(QtGui.QPixmap("icon.png"))

        width = self.QMainWindow.frameGeometry().width()
        height =self.QMainWindow.frameGeometry().height()
        self.pic[4].resize(size.width()/8, size.height()/8)
        self.pic[4].setScaledContents(True)

        self.pic[4].move(size.width()-size.width()/8 - 100,  size.height() - 7*size.height()/8)
        self.pic[4].show() # You were missing this.



    def getIcon(self,url):
        data = urllib.request.urlopen(url).read()
        icon = QtGui.QPixmap()
        icon.loadFromData(data)
        return icon

    def _createToolBars(self):
        # File toolbar
        current_directory = str(pathlib.Path(__file__).parent.absolute())
        path = current_directory + '/OCR.png'\


# Create object LeftToolbar in your main window
        self.LeftToolbar = LeftToolbar( self.addTracker, self.openPhone, self.addGeneral, self.addWebcam, self.addElevenLabs,self.addSynthesia, self.addLlama, self.getRecording, self.AddStableVideo, self.addScrape, self.addOCRSemantic, self.addDescribe, self.addOCR, self.addPrint, self.addScroll, self.addSendData, self.addIfElse, self.addKeypress, self.addMove, self.addClick, self.getData, self.getScreenshot, self.download_file, self.openCheat, self.scheduleCheat, self.addAIDetector, self.runOnCheatLayer, self.trainModel, self.launchAtlas, self.AddStableDiffusion, self.addCustomCode, self.addEmail,self.addRiku,self.addMath,self.addGsheets,self.addOpen,self.addDelay, self.newCheat, self.addGetFiles, self.AddGPT4)
        self.QMainWindow.addToolBar(QtCore.Qt.LeftToolBarArea,self.LeftToolbar)

 
        self.BottomToolbar = BottomToolbar(self.scheduled_jobs, self.scheduleCheat, self.user_key, self.openCheat, self, self.thread_signals, self.workflows)
        self.QMainWindow.addToolBar(QtCore.Qt.RightToolBarArea,self.BottomToolbar)

       # self.scroll_area = QScrollArea()
       # self.scroll_area.setWidget(fileToolBar)


# Add test notifications
        
      #  self.notification_manager.add_notification("Title 1", "This is a message with some overflow text that is intended to be cut off if it's too long.", "image.png", lambda: print("Notification 1 clicked!"))
     #   self.notification_manager.add_notification("Title 2", "Another message", "image.png", lambda: print("Notification 2 clicked!"))

        #self.QMainWindow.addToolBar(QtCore.Qt.LeftToolBarArea, fileToolBar)
        #self.QMainWindow.setToolButtonStyle(QtCore.Qt.ToolButtonTextUnderIcon)
    def setkey(self, key):
        self.user_key = key
        self.thread_signals.setkey.emit(key)
    def send_complex_message(self, to, subject, text, file_name):

        return requests.post(
            "https://api.mailgun.net/v3/YOUR_DOMAIN_NAME/messages",
            auth=("api", "YOUR_API_KEY"),
            files=[("attachment", open(resource_path(file_name)))],
            data={"from": "Desktop Cloud Agent <support@mg.cheatlayer.com>",
                  "to": to,
                  "subject": subject,
                  "text": text})
    def trigger(self, job):
        global nodes, current_desktop
        print("trigger")
        print(job)
        
        for key, value in job.items():
            self.global_variables[key] = value
        if "script" in job:
           params = job["script"][1:][:-1]
        #params includes url encoded parameters. Split them each and store them in self.global_variables
        params = params.split("&")
        print(params)
        print("PARAMS")
        for param in params:
            if len(param) > 0:
                self.global_variables[param.split("=")[0]] = param.split("=")[1].replace("%20"," ")
        print(self.global_variables)
        self.job_runner.run(job)

                    
    def __init__(self, addTracker, referral, openPhone, user_email, addGeneral, addWebcam, addElevenLabs, addSynthesia, addLlama, getRecording, AddStableVideo, RoboflowKey, workflows, addScrape, scheduled_jobs, atlasWindow, user_key, drawHistory, verified, addOCR, addPrint, addScroll, addSendData, addIfElse, addKeypress, addMove, addClick,getData,getScreenshot,download_file, openCheat, scheduleCheat, addAIDetector, runOnCheatLayer, trainModel, launchAtlas, AddStableDiffusion, addCustomCode, addEmail,addRiku,addMath,addGsheets,addOpen,addDelay, newCheat, addGetFiles,openAtlas, chatWindow, chatInstance, blip_decoder, addDescribe, addOCRSemantic, openaikey, user_plan, thread_signals, AddGPT4, parent=None, **kwargs):
        """
        Args:
            parent (object): object parent.
            **kwargs (dict): Used for overriding internal objects at init time.
        """
        super(NodeGraph, self).__init__(parent)
        self.setObjectName('NodeGraph')
        self.notification_manager = NotificationManager()
        self.addTracker = addTracker
        self.getRecording = getRecording
        self.verified = verified
        self.verified = verified
        self.openPhone = openPhone
        self.keyboard_listener = None
        self.mouse_listener = None
        self.referral = referral
        self.addWebcam = addWebcam
        self.AddGPT4 = AddGPT4
        self.lat = 0
        self.long = 0

        self.addGeneral = addGeneral
        self.thread_signals = thread_signals
        self.addElevenLabs = addElevenLabs
        self.addSynthesia = addSynthesia
        self.workflows = workflows
        self.webhook_jobs = []
        self.addLlama = addLlama
        self.addGetFiles = addGetFiles
        self.addAIDetector = addAIDetector
        self.roboflow = RoboflowKey
        self.total_errors = 0
        self.addOCRSemantic = addOCRSemantic
        self.AddStableVideo = AddStableVideo
        self.AddStableDiffusion = AddStableDiffusion
        self.scheduled_jobs = scheduled_jobs
        self.drawHistory = drawHistory
        self.atlasWindow = atlasWindow
        self.addSendData = addSendData
        self.addCustomCode = addCustomCode
        self.addEmail = addEmail
        self.addScrape = addScrape
        self.addRiku = addRiku
        self.addMath = addMath
        self.openAtlas = openAtlas
        self.addGsheets = addGsheets
        self.addOpen = addOpen
        self.chatWindow = chatWindow
        self.addDelay = addDelay
        self.addScroll = addScroll
        self.newCheat = newCheat
        self.addClick = addClick
        self.openCheat = openCheat
        self.blip_decoder = blip_decoder
        self.chatInstance = chatInstance
        self.window = QWidget()
        #made self.window transparnt and always on top over other windows without a title bar and make the size 500x500 pixels
        self.window.setWindowFlags(QtCore.Qt.WindowStaysOnTopHint | QtCore.Qt.FramelessWindowHint | QtCore.Qt.WindowTransparentForInput | QtCore.Qt.WindowDoesNotAcceptFocus)
        self.window.setAttribute(QtCore.Qt.WA_TranslucentBackground)
        self.window.setFixedSize(500, 500)
        #move the window to the top left corner and make it change size dynamically based on the content
        self.window.move(0, 0)
 #       self.window.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)

        # Set the window title
        self.window.setWindowTitle('Image Viewer')

        # Create a QVBoxLayout instance
        self.layout = QVBoxLayout()

        # Create a QLabel instance
        self.label = QLabel(self.window)
        self.layout.addWidget(self.label)



        self.label2 = QLabel(self.window)
        self.layout.addWidget(self.label2)
        self.layout.addStretch()
        self.label2.setStyleSheet("QLabel { background-color: white; }")

        #move label2 down by 100 pixels
        self.label.move(0, 48)
        self.window.show()
        #self.trainModel = trainModel
        self.getData = getData
        self.launchAtlas = launchAtlas
        self.runOnCheatLayer = runOnCheatLayer
        self.scheduleCheat = scheduleCheat
        self.getScreenshot = getScreenshot
        self.download_file = download_file
        self.addMove = addMove
        self.addOCR = addOCR
        self.addDescribe = addDescribe
        self.addPrint = addPrint
        self.addKeypress = addKeypress
        self.addIfElse = addIfElse
        self.stop = False
        self.openaikey = openaikey
        self.user_plan = user_plan
        self.last_cheat = "base"
        self.user_key = user_key
        self.user_email = user_email
        self.data_graph = []
        #screenshot_taker_thread = ScreenshotTaker(workflows, self.user_key, self.user_plan, self, self.thread_signals)
   #     screenshot_taker_thread.daemon = True  # Set as a daemon so it will be killed once the main thread is dead.
   #     screenshot_taker_thread.start()
        self.job_runner = JobRunner(self, thread_signals, self.user_key)
        self.creator_stats = ""
        self.global_variables = {"creator_stats" :"test;;;test;;;test;;test","Total Runs": 0, "Total Stable Diffusion Runs":0,"user" :"rohan@cheatlayer.com", "topic":"gpt-4", "URL":"https://storage.googleapis.com/cheatlayer/Image_To_Video.png", 
                "caption" :"ai automation built by cheatlayer.com","voiceover" :"ai automation built by cheatlayer.com","prompt" : "ai automation built by cheatlayer.com"}
        
        self.agent_data = {}
        self.totalSDIRuns = 0
        self.totalSDRuns = 0
        self.totalRuns = 0
        self.Processing = False
        self.mouse_counter = 0
        self.variables = []
        self.history = []
        self.half_small = 25
        self.shift_down = False
        self.started = False
        self.size_small = 50
        self.half = 25
        self.platform = platform.platform()
        #self.history.append(json.dumps({"type":"Start Node", "x": 0, "y": 0, "Application":"chrome"}, cls=NumpyEncoder))

        self.size = 50
        self.drawHistory = drawHistory
        self._model = (
            kwargs.get('model') or NodeGraphModel())
        self._node_factory = (
            kwargs.get('node_factory') or NodeFactory())

        self._undo_view = None
        self._undo_stack = (
            kwargs.get('undo_stack') or QtWidgets.QUndoStack(self))
        self.QMainWindow = QMainWindow()

        global UIModel, rf, project
        
        #rf = Roboflow(api_key=RoboflowKey)
        #project = rf.workspace().project("cheats")
        #UIModel = project.version(1).model
        self._widget = None

        self._sub_graphs = {}

        self._viewer = (
            kwargs.get('viewer') or NodeViewer(undo_stack=self._undo_stack))

        self._build_context_menu()
        self._register_builtin_nodes()
        self._wire_signals()
        self.restartServer()

    def restartServer(self):
        self.server = Server(self.trigger, self.setkey, self.data_graph, self.runNode, self.thread_signals, self.sendMessageRTCAsync, self)
        server_thread = threading.Thread(target=self.server.start)
        server_thread.start()

    def updateSchedule(self, workflow):
        self.scheduled_jobs = workflow
        self.BottomToolbar.updateSchedule(workflow)
        
    def __repr__(self):
        return '<{}("root") object at {}>'.format(
            self.__class__.__name__, hex(id(self)))

    def _build_context_menu(self):
        """
        build the essential menus commands for the graph context menu.
        """
        from NodeGraphQt.base.graph_actions import build_context_menu
        build_context_menu(self)

    def _register_builtin_nodes(self):
        """
        Register the default builtin nodes to the :meth:`NodeGraph.node_factory`
        """
        self.register_node(BackdropNode, alias='Backdrop')

    def _wire_signals(self):
        """
        Connect up all the signals and slots here.
        """
        # TODO: refactor hard coded tab search logic into
        #       "graph_actions.py" module.
        # hard coded tab search.
        tab = QtWidgets.QShortcut(
            QtGui.QKeySequence(QtCore.Qt.Key_Tab), self._viewer)
        tab.activated.connect(self._toggle_tab_search)
        self._viewer.show_tab_search.connect(self._toggle_tab_search)

        # internal signals.
        self._viewer.search_triggered.connect(self._on_search_triggered)
        self._viewer.connection_sliced.connect(self._on_connection_sliced)
        self._viewer.connection_changed.connect(self._on_connection_changed)
        self._viewer.moved_nodes.connect(self._on_nodes_moved)
        self._viewer.node_double_clicked.connect(self._on_node_double_clicked)
        self._viewer.node_name_changed.connect(self._on_node_name_changed)
        self._viewer.node_backdrop_updated.connect(
            self._on_node_backdrop_updated)
        self._viewer.insert_node.connect(self._on_insert_node)

        # pass through translated signals.
        self._viewer.node_selected.connect(self._on_node_selected)
        self._viewer.node_selection_changed.connect(
            self._on_node_selection_changed)
        self._viewer.data_dropped.connect(self._on_node_data_dropped)

    def _on_insert_node(self, pipe, node_id, prev_node_pos):
        """
        Slot function triggered when a selected node has collided with a pipe.

        Args:
            pipe (Pipe): collided pipe item.
            node_id (str): selected node id to insert.
            prev_node_pos (dict): previous node position. {NodeItem: [prev_x, prev_y]}
        """
        node = self.get_node_by_id(node_id)

        # exclude if not a BaseNode
        if not isinstance(node, BaseNode):
            return

        disconnected = [(pipe.input_port, pipe.output_port)]
        connected = []

        if node.input_ports():
            connected.append(
                (pipe.output_port, node.input_ports()[0].view)
            )
        if node.output_ports():
            connected.append(
                (node.output_ports()[0].view, pipe.input_port)
            )

        self._undo_stack.beginMacro('inserted node')
        self._on_connection_changed(disconnected, connected)
        self._on_nodes_moved(prev_node_pos)
        self._undo_stack.endMacro()

    def _toggle_tab_search(self):
        """
        toggle the tab search widget.
        """
        if self._viewer.underMouse():
            self._viewer.tab_search_set_nodes(self._node_factory.names)
            self._viewer.tab_search_toggle()

    def _on_property_bin_changed(self, node_id, prop_name, prop_value):
        """
        called when a property widget has changed in a properties bin.
        (emits the node object, property name, property value)

        Args:
            node_id (str): node id.
            prop_name (str): node property name.
            prop_value (object): python built in types.
        """
        node = self.get_node_by_id(node_id)

        # prevent signals from causing a infinite loop.
        if node.get_property(prop_name) != prop_value:
            node.set_property(prop_name, prop_value)

    def _on_node_name_changed(self, node_id, name):
        """
        called when a node text qgraphics item in the viewer is edited.
        (sets the name through the node object so undo commands are registered.)

        Args:
            node_id (str): node id emitted by the viewer.
            name (str): new node name.
        """
        node = self.get_node_by_id(node_id)
        node.set_name(name)

        # TODO: not sure about redrawing the node here.
        node.view.draw_node()

    def _on_node_double_clicked(self, node_id):
        """
        called when a node in the viewer is double click.
        (emits the node object when the node is clicked)

        Args:
            node_id (str): node id emitted by the viewer.
        """
        node = self.get_node_by_id(node_id)
        self.node_double_clicked.emit(node)

    def _on_node_selected(self, node_id):
        """
        called when a node in the viewer is selected on left click.
        (emits the node object when the node is clicked)

        Args:
            node_id (str): node id emitted by the viewer.
        """
        node = self.get_node_by_id(node_id)
        self.node_selected.emit(node)

    def _on_node_selection_changed(self, sel_ids, desel_ids):
        """
        called when the node selection changes in the viewer.
        (emits node objects <selected nodes>, <deselected nodes>)

        Args:
            sel_ids (list[str]): new selected node ids.
            desel_ids (list[str]): deselected node ids.
        """
        sel_nodes = [self.get_node_by_id(nid) for nid in sel_ids]
        unsel_nodes = [self.get_node_by_id(nid) for nid in desel_ids]
        self.node_selection_changed.emit(sel_nodes, unsel_nodes)

    def _on_node_data_dropped(self, data, pos):
        """
        called when data has been dropped on the viewer.

        Example Identifiers:
            URI = ngqt://path/to/node/session.graph
            URN = ngqt::node:com.nodes.MyNode1;node:com.nodes.MyNode2

        Args:
            data (QtCore.QMimeData): mime data.
            pos (QtCore.QPoint): scene position relative to the drop.
        """
        uri_regex = re.compile(r'{}(?:/*)([\w/]+)(\.\w+)'.format(URI_SCHEME))
        urn_regex = re.compile(r'{}([\w\.:;]+)'.format(URN_SCHEME))
        if data.hasFormat('text/uri-list'):
            for url in data.urls():
                local_file = url.toLocalFile()
                if local_file:
                    try:
                        self.import_session(local_file)
                        continue
                    except Exception as e:
                        pass

                url_str = url.toString()
                uri_search = uri_regex.search(url_str)
                urn_search = urn_regex.search(url_str)
                if uri_search:
                    path = uri_search.group(1)
                    ext = uri_search.group(2)
                    self.import_session('{}{}'.format(path, ext))
                elif urn_search:
                    search_str = urn_search.group(1)
                    node_ids = sorted(re.findall('node:([\w\\.]+)', search_str))
                    x, y = pos.x(), pos.y()
                    for node_id in node_ids:
                        self.create_node(node_id, pos=[x, y])
                        x += 80
                        y += 80

    def _on_nodes_moved(self, node_data):
        """
        called when selected nodes in the viewer has changed position.

        Args:
            node_data (dict): {<node_view>: <previous_pos>}
        """
        self._undo_stack.beginMacro('move nodes')
        for node_view, prev_pos in node_data.items():
            node = self._model.nodes[node_view.id]
            self._undo_stack.push(NodeMovedCmd(node, node.pos(), prev_pos))
        self._undo_stack.endMacro()

    def _on_node_backdrop_updated(self, node_id, update_property, value):
        """
        called when a BackdropNode is updated.

        Args:
            node_id (str): backdrop node id.
            value (str): update type.
        """
        backdrop = self.get_node_by_id(node_id)
        if backdrop and isinstance(backdrop, BackdropNode):
            backdrop.on_backdrop_updated(update_property, value)

    def _on_search_triggered(self, node_type, pos):
        """
        called when the tab search widget is triggered in the viewer.

        Args:
            node_type (str): node identifier.
            pos (tuple or list): x, y position for the node.
        """
        self.create_node(node_type, pos=pos)

    def _on_connection_changed(self, disconnected, connected):
        """
        called when a pipe connection has been changed in the viewer.

        Args:
            disconnected (list[list[widgets.port.PortItem]):
                pair list of port view items.
            connected (list[list[widgets.port.PortItem]]):
                pair list of port view items.
        """
        if not (disconnected or connected):
            return

        label = 'connect node(s)' if connected else 'disconnect node(s)'
        ptypes = {IN_PORT: 'inputs', OUT_PORT: 'outputs'}

        self._undo_stack.beginMacro(label)
        for p1_view, p2_view in disconnected:
            node1 = self._model.nodes[p1_view.node.id]
            node2 = self._model.nodes[p2_view.node.id]
            port1 = getattr(node1, ptypes[p1_view.port_type])()[p1_view.name]
            port2 = getattr(node2, ptypes[p2_view.port_type])()[p2_view.name]
            port1.disconnect_from(port2)
        for p1_view, p2_view in connected:
            node1 = self._model.nodes[p1_view.node.id]
            node2 = self._model.nodes[p2_view.node.id]
            port1 = getattr(node1, ptypes[p1_view.port_type])()[p1_view.name]
            port2 = getattr(node2, ptypes[p2_view.port_type])()[p2_view.name]
            port1.connect_to(port2)
        self._undo_stack.endMacro()
    def startMenu(self):
        self.defineMenu()
        self._createToolBars()


    def _on_connection_sliced(self, ports):
        """
        slot when connection pipes have been sliced.

        Args:
            ports (list[list[widgets.port.PortItem]]):
                pair list of port connections (in port, out port)
        """
        if not ports:
            return
        ptypes = {IN_PORT: 'inputs', OUT_PORT: 'outputs'}
        self._undo_stack.beginMacro('slice connections')
        for p1_view, p2_view in ports:
            node1 = self._model.nodes[p1_view.node.id]
            node2 = self._model.nodes[p2_view.node.id]
            port1 = getattr(node1, ptypes[p1_view.port_type])()[p1_view.name]
            port2 = getattr(node2, ptypes[p2_view.port_type])()[p2_view.name]
            port1.disconnect_from(port2)
        self._undo_stack.endMacro()

    @property
    def model(self):
        """
        The model used for storing the node graph data.

        Returns:
            NodeGraphQt.base.model.NodeGraphModel: node graph model.
        """
        return self._model

    @property
    def node_factory(self):
        """
        Return the node factory object used by the node graph.

        Returns:
            NodeFactory: node factory.
        """
        return self._node_factory

    @property
    def widget(self):
        """
        The node graph widget for adding into a layout.

        Returns:
            NodeGraphWidget: node graph widget.
        """
        if self._widget is None:
            self._widget = NodeGraphWidget(self.drawHistory, self.verified, self)
            self.QMainWindow.setCentralWidget(self._widget)
            self.startMenu()
            self._widget.addTab(self._viewer, 'Node Graph')
            # hide the close button on the first tab.
            tab_bar = self._widget.tabBar()
            for btn_flag in [tab_bar.RightSide, tab_bar.LeftSide]:
                tab_btn = tab_bar.tabButton(0, btn_flag)
                if tab_btn:
                    tab_btn.deleteLater()
                    tab_bar.setTabButton(0, btn_flag, None)
            self._widget.tabCloseRequested.connect(
                self._on_close_sub_graph_tab
            )
        return self._widget

    @property
    def undo_view(self):
        """
        Returns node graph undo history list widget.

        Returns:
            PySide2.QtWidgets.QUndoView: node graph undo view.
        """
        if self._undo_view is None:
            self._undo_view = QtWidgets.QUndoView(self._undo_stack)
            self._undo_view.setWindowTitle('Undo History')
        return self._undo_view

    def show(self):
        """
        Show node graph widget this is just a convenience
        function to :meth:`NodeGraph.widget.show()`.
        """
        self.widget.show()

    def close(self):
        """
        Close node graph NodeViewer widget this is just a convenience
        function to :meth:`NodeGraph.widget.close()`.
        """
        self.widget.close()

    def viewer(self):
        """
        Returns the internal view interface used by the node graph.

        Warnings:
            Methods in the ``NodeViewer`` are used internally
            by ``NodeGraphQt`` components to get the widget use
            :attr:`NodeGraph.widget`.

        See Also:
            :attr:`NodeGraph.widget` to add the node graph widget into a
            :class:`PySide2.QtWidgets.QLayout`.

        Returns:
            NodeGraphQt.widgets.viewer.NodeViewer: viewer interface.
        """
        return self._viewer

    def scene(self):
        """
        Returns the ``QGraphicsScene`` object used in the node graph.

        Returns:
            NodeGraphQt.widgets.scene.NodeScene: node scene.
        """
        return self._viewer.scene()

    def background_color(self):
        """
        Return the node graph background color.

        Returns:
            tuple: r, g ,b
        """
        return self.scene().background_color

    def set_background_color(self, r, g, b):
        """
        Set node graph background color.

        Args:
            r (int): red value.
            g (int): green value.
            b (int): blue value.
        """
        self.scene().background_color = (r, g, b)
        self._viewer.force_update()

    def grid_color(self):
        """
        Return the node graph grid color.

        Returns:
            tuple: r, g ,b
        """
        return self.scene().grid_color

    def set_grid_color(self, r, g, b):
        """
        Set node graph grid color.

        Args:
            r (int): red value.
            g (int): green value.
            b (int): blue value.
        """
        self.scene().grid_color = (r, g, b)
        self._viewer.force_update()

    def set_grid_mode(self, mode=VIEWER_GRID_DOTS):
        """
        Set node graph grid mode.

        Note:
            By default grid mode is set to "VIEWER_GRID_LINES".

            Node graph background types:

            * :attr:`NodeGraphQt.constants.VIEWER_GRID_NONE`
            * :attr:`NodeGraphQt.constants.VIEWER_GRID_DOTS`
            * :attr:`NodeGraphQt.constants.VIEWER_GRID_LINES`

        Args:
            mode (int): background style.
        """
        self.scene().grid_mode = VIEWER_GRID_DOTS
        self._viewer.force_update()

    def add_properties_bin(self, prop_bin):
        """
        Wire up a properties bin widget to the node graph.

        Args:
            prop_bin (NodeGraphQt.PropertiesBinWidget): properties widget.
        """
        prop_bin.property_changed.connect(self._on_property_bin_changed)

    def undo_stack(self):
        """
        Returns the undo stack used in the node graph.

        See Also:
            :meth:`NodeGraph.begin_undo()`,
            :meth:`NodeGraph.end_undo()`

        Returns:
            QtWidgets.QUndoStack: undo stack.
        """
        return self._undo_stack

    def clear_undo_stack(self):
        """
        Clears the undo stack.

        Note:
            Convenience function to
            :meth:`NodeGraph.undo_stack().clear()`

        See Also:
            :meth:`NodeGraph.begin_undo()`,
            :meth:`NodeGraph.end_undo()`,
            :meth:`NodeGraph.undo_stack()`
        """
        self._undo_stack.clear()

    def begin_undo(self, name):
        """
        Start of an undo block followed by a
        :meth:`NodeGraph.end_undo()`.

        Args:
            name (str): name for the undo block.
        """
        self._undo_stack.beginMacro(name)

    def end_undo(self):
        """
        End of an undo block started by
        :meth:`NodeGraph.begin_undo()`.
        """
        self._undo_stack.endMacro()

    def context_menu(self):
        """
        Returns the context menu for the node graph.

        Note:
            This is a convenience function to
            :meth:`NodeGraph.get_context_menu`
            with the arg ``menu="graph"``

        Returns:
            NodeGraphQt.NodeGraphMenu: context menu object.
        """
        return self.get_context_menu('graph')

    def context_nodes_menu(self):
        """
        Returns the context menu for the nodes.

        Note:
            This is a convenience function to
            :meth:`NodeGraph.get_context_menu`
            with the arg ``menu="nodes"``

        Returns:
            NodeGraphQt.NodesMenu: context menu object.
        """
        return self.get_context_menu('nodes')

    def get_context_menu(self, menu):
        """
        Returns the context menu specified by the name.

        Menu Types:
            - ``"graph"`` context menu from the node graph.
            - ``"nodes"`` context menu for the nodes.

        Args:
            menu (str): menu name.

        Returns:
            NodeGraphQt.NodeGraphMenu or NodeGraphQt.NodesMenu: context menu object.
        """
        menus = self._viewer.context_menus()
        if menus.get(menu):
            if menu == 'graph':
                return NodeGraphMenu(self, menus[menu])
            elif menu == 'nodes':
                return NodesMenu(self, menus[menu])

    def disable_context_menu(self, disabled=True, name='all'):
        """
        Disable/Enable context menus from the node graph.

        Menu Types:
            - ``"all"`` all context menus from the node graph.
            - ``"graph"`` context menu from the node graph.
            - ``"nodes"`` context menu for the nodes.

        Args:
            disabled (bool): true to enable context menu.
            name (str): menu name. (default: ``"all"``)
        """
        if name == 'all':
            for k, menu in self._viewer.context_menus().items():
                menu.setDisabled(disabled)
                menu.setVisible(not disabled)
            return
        menus = self._viewer.context_menus()
        if menus.get(name):
            menus[name].setDisabled(disabled)
            menus[name].setVisible(not disabled)

    def acyclic(self):
        """
        Returns true if the current node graph is acyclic.

        See Also:
            :meth:`NodeGraph.set_acyclic`

        Returns:
            bool: true if acyclic (default: ``True``).
        """
        return self._model.acyclic

    def set_acyclic(self, mode=False):
        """
        Enable the node graph to be a acyclic graph. (default: ``False``)

        See Also:
            :meth:`NodeGraph.acyclic`

        Args:
            mode (bool): true to enable acyclic.
        """
        self._model.acyclic = mode
        self._viewer.acyclic = mode

    def pipe_collision(self):
        """
        Returns if pipe collision is enabled.

        See Also:
            To enable/disable pipe collision
            :meth:`NodeGraph.set_pipe_collision_enabled`

        Returns:
            bool: True if pipe collision is enabled.
        """
        return self._model.pipe_collision

    def set_pipe_collision(self, mode=True):
        """
        Enable/Disable pipe collision.

        When enabled dragging a node over a pipe will allow the node to be
        inserted as a new connection between the pipe.

        See Also:
            :meth:`NodeGraph.pipe_collision`

        Args:
            mode (bool): False to disable pipe collision.
        """
        self._model.pipe_collision = mode
        self._viewer.pipe_collision = mode

    def set_pipe_style(self, style=PIPE_LAYOUT_CURVED):
        """
        Set node graph pipes to be drawn as straight, curved or angled.

        .. image:: _images/pipe_layout_types.gif
            :width: 80%

        Note:
            By default pipe layout is set to "PIPE_LAYOUT_CURVED".

            Pipe Layout Styles:

            * :attr:`NodeGraphQt.constants.PIPE_LAYOUT_CURVED`
            * :attr:`NodeGraphQt.constants.PIPE_LAYOUT_STRAIGHT`
            * :attr:`NodeGraphQt.constants.PIPE_LAYOUT_ANGLE`

        Args:
            style (int): pipe layout style.
        """
        pipe_max = max([PIPE_LAYOUT_CURVED,
                        PIPE_LAYOUT_STRAIGHT,
                        PIPE_LAYOUT_ANGLE])
        style = style if 0 <= style <= pipe_max else PIPE_LAYOUT_CURVED
        self._viewer.set_pipe_layout(style)

    def fit_to_selection(self):
        """
        Sets the zoom level to fit selected nodes.
        If no nodes are selected then all nodes in the graph will be framed.
        """
        nodes = self.selected_nodes() or self.all_nodes()
        if not nodes:
            return
        self._viewer.zoom_to_nodes([n.view for n in nodes])
        
        self._viewer.clear_key_state()
        self._viewer.clearFocus()

    def reset_zoom(self):
        """
        Reset the zoom level
        """
        self._viewer.reset_zoom()

    def set_zoom(self, zoom=0):
        """
        Set the zoom factor of the Node Graph the default is ``0.0``

        Args:
            zoom (float): zoom factor (max zoom out ``-0.9`` / max zoom in ``2.0``)
        """
        self._viewer.set_zoom(zoom)

    def get_zoom(self):
        """
        Get the current zoom level of the node graph.

        Returns:
            float: the current zoom level.
        """
        return self._viewer.get_zoom()

    def center_on(self, nodes=None):
        """
        Center the node graph on the given nodes or all nodes by default.

        Args:
            nodes (list[NodeGraphQt.BaseNode]): a list of nodes.
        """
        self._viewer.center_selection([n.view for n in nodes])

    def center_selection(self):
        """
        Centers on the current selected nodes.
        """
        nodes = self._viewer.selected_nodes()
        self._viewer.center_selection(nodes)

    def registered_nodes(self):
        """
        Return a list of all node types that have been registered.

        See Also:
            To register a node :meth:`NodeGraph.register_node`

        Returns:
            list[str]: list of node type identifiers.
        """
        return sorted(self._node_factory.nodes.keys())

    def register_node(self, node, alias=None):
        """
        Register the node to the :meth:`NodeGraph.node_factory`

        Args:
            node (_NodeGraphQt.NodeObject): node object.
            alias (str): custom alias name for the node type.
        """
        self._node_factory.register_node(node, alias)
        self._viewer.rebuild_tab_search()

    def register_nodes(self, nodes):
        """
        Register the nodes to the :meth:`NodeGraph.node_factory`

        Args:
            nodes (list): list of nodes.
        """
        [self._node_factory.register_node(n) for n in nodes]
        self._viewer.rebuild_tab_search()

    def create_node(self, node_type, data=None, name=None, selected=True, color=None,
                    text_color=None, pos=None, push_undo=True):
        """
        Create a new node in the node graph.

        See Also:
            To list all node types :meth:`NodeGraph.registered_nodes`

        Args:
            node_type (str): node instance type.
            name (str): set name of the node.
            selected (bool): set created node to be selected.
            color (tuple or str): node color ``(255, 255, 255)`` or ``"#FFFFFF"``.
            text_color (tuple or str): text color ``(255, 255, 255)`` or ``"#FFFFFF"``.
            pos (list[int, int]): initial x, y position for the node (default: ``(0, 0)``).
            push_undo (bool): register the command to the undo stack. (default: True)

        Returns:
            BaseNode: the created instance of the node.
        """
        node = self._node_factory.create_node_instance(node_type)
        if node:
            node._graph = self
            node.model._graph_model = self.model

            wid_types = node.model.__dict__.pop('_TEMP_property_widget_types')
            prop_attrs = node.model.__dict__.pop('_TEMP_property_attrs')

            if self.model.get_node_common_properties(node.type_) is None:
                node_attrs = {node.type_: {
                    n: {'widget_type': wt} for n, wt in wid_types.items()
                }}
                for pname, pattrs in prop_attrs.items():
                    node_attrs[node.type_][pname].update(pattrs)
                self.model.set_node_common_properties(node_attrs)

            node.NODE_NAME = self.get_unique_name(name or node.NODE_NAME)
            node.data = data
            node.model.name = node.NODE_NAME
            node.model.selected = selected

            def format_color(clr):
                if isinstance(clr, str):
                    clr = clr.strip('#')
                    return tuple(int(clr[i:i + 2], 16) for i in (0, 2, 4))
                return clr

            if color:
                node.model.color = format_color(color)
            if text_color:
                node.model.text_color = format_color(text_color)
            if pos:
                node.model.pos = [float(pos[0]), float(pos[1])]

            node.update()

            if push_undo:
                undo_cmd = NodeAddedCmd(self, node, node.model.pos)
                undo_cmd.setText('create node: "{}"'.format(node.NODE_NAME))
                self._undo_stack.push(undo_cmd)
            else:
                NodeAddedCmd(self, node, node.model.pos).redo()

            self.node_created.emit(node)
            return node
        raise TypeError('\n\n>> Cannot find node:\t"{}"\n'.format(node_type))

    def add_node(self, node, pos=None, selected=True, push_undo=True):
        """
        Add a node into the node graph.
        unlike the :meth:`NodeGraph.create_node` function this will not
        trigger the :attr:`NodeGraph.node_created` signal.

        Args:
            node (NodeGraphQt.BaseNode): node object.
            pos (list[float]): node x,y position. (optional)
            selected (bool): node selected state. (optional)
            push_undo (bool): register the command to the undo stack. (default: True)
        """
        assert isinstance(node, NodeObject), 'node must be a Node instance.'

        wid_types = node.model.__dict__.pop('_TEMP_property_widget_types')
        prop_attrs = node.model.__dict__.pop('_TEMP_property_attrs')

        if self.model.get_node_common_properties(node.type_) is None:
            node_attrs = {node.type_: {
                n: {'widget_type': wt} for n, wt in wid_types.items()
            }}
            for pname, pattrs in prop_attrs.items():
                node_attrs[node.type_][pname].update(pattrs)
            self.model.set_node_common_properties(node_attrs)

        node._graph = self
        node.NODE_NAME = self.get_unique_name(node.NODE_NAME)
        node.model._graph_model = self.model
        node.model.name = node.NODE_NAME
        node.update()

        if push_undo:
            self._undo_stack.beginMacro('add node: "{}"'.format(node.name()))
            self._undo_stack.push(NodeAddedCmd(self, node, pos))
            if selected:
                node.set_selected(True)
            self._undo_stack.endMacro()
        else:
            NodeAddedCmd(self, node, pos).redo()

    def delete_node(self, node, push_undo=True):
        """
        Remove the node from the node graph.

        Args:
            node (NodeGraphQt.BaseNode): node object.
            push_undo (bool): register the command to the undo stack. (default: True)
        """
        assert isinstance(node, NodeObject), \
            'node must be a instance of a NodeObject.'
        node_id = node.id
        if push_undo:
            self._undo_stack.beginMacro('delete node: "{}"'.format(node.name()))

        if isinstance(node, BaseNode):
            for p in node.input_ports():
                if p.locked():
                    p.set_locked(False,
                                 connected_ports=False,
                                 push_undo=push_undo)
                p.clear_connections()
            for p in node.output_ports():
                if p.locked():
                    p.set_locked(False,
                                 connected_ports=False,
                                 push_undo=push_undo)
                p.clear_connections()

        if push_undo:
            self._undo_stack.push(NodeRemovedCmd(self, node))
            self._undo_stack.endMacro()
        else:
            NodeRemovedCmd(self, node).redo()

        self.nodes_deleted.emit([node_id])

    def remove_node(self, node, push_undo=True):
        """
        Remove the node from the node graph.

        unlike the :meth:`NodeGraph.delete_node` function this will not
        trigger the :attr:`NodeGraph.nodes_deleted` signal.

        Args:
            node (NodeGraphQt.BaseNode): node object.
            push_undo (bool): register the command to the undo stack. (default: True)

        """
        assert isinstance(node, NodeObject), 'node must be a Node instance.'

        if push_undo:
            self._undo_stack.beginMacro('delete node: "{}"'.format(node.name()))

        if isinstance(node, BaseNode):
            for p in node.input_ports():
                if p.locked():
                    p.set_locked(False,
                                 connected_ports=False,
                                 push_undo=push_undo)
                p.clear_connections()
            for p in node.output_ports():
                if p.locked():
                    p.set_locked(False,
                                 connected_ports=False,
                                 push_undo=push_undo)
                p.clear_connections()

        if push_undo:
            self._undo_stack.push(NodeRemovedCmd(self, node))
            self._undo_stack.endMacro()
        else:
            NodeRemovedCmd(self, node).redo()

    def delete_nodes(self, nodes, push_undo=True):
        """
        Remove a list of specified nodes from the node graph.

        Args:
            nodes (list[NodeGraphQt.BaseNode]): list of node instances.
            push_undo (bool): register the command to the undo stack. (default: True)
        """
        if not nodes:
            return
        if len(nodes) == 1:
            self.delete_node(nodes[0], push_undo=push_undo)
            return
        node_ids = [n.id for n in nodes]
        if push_undo:
            self._undo_stack.beginMacro('deleted "{}" nodes'.format(len(nodes)))
        for node in nodes:
            if isinstance(node, BaseNode):
                for p in node.input_ports():
                    if p.locked():
                        p.set_locked(False,
                                     connected_ports=False,
                                     push_undo=push_undo)
                    p.clear_connections(push_undo=push_undo)
                for p in node.output_ports():
                    if p.locked():
                        p.set_locked(False,
                                     connected_ports=False,
                                     push_undo=push_undo)
                    p.clear_connections(push_undo=push_undo)
            if push_undo:
                self._undo_stack.push(NodeRemovedCmd(self, node))
            else:
                NodeRemovedCmd(self, node).redo()
        if push_undo:
            self._undo_stack.endMacro()
        self.nodes_deleted.emit(node_ids)

    def all_nodes(self):
        """
        Return all nodes in the node graph.

        Returns:
            list[NodeGraphQt.BaseNode]: list of nodes.
        """
        return list(self._model.nodes.values())

    def selected_nodes(self):
        """
        Return all selected nodes that are in the node graph.

        Returns:
            list[NodeGraphQt.BaseNode]: list of nodes.
        """
        nodes = []
        for item in self._viewer.selected_nodes():
            node = self._model.nodes[item.id]
            nodes.append(node)
        return nodes

    def select_all(self):
        """
        Select all nodes in the node graph.
        """
        self._undo_stack.beginMacro('select all')
        [node.set_selected(True) for node in self.all_nodes()]
        self._undo_stack.endMacro()

    def clear_selection(self):
        """
        Clears the selection in the node graph.
        """
        self._undo_stack.beginMacro('clear selection')
        [node.set_selected(False) for node in self.all_nodes()]
        self._undo_stack.endMacro()

    def get_node_by_id(self, node_id=None):
        """
        Returns the node from the node id string.

        Args:
            node_id (str): node id (:attr:`NodeObject.id`)

        Returns:
            NodeGraphQt.NodeObject: node object.
        """
        return self._model.nodes.get(node_id, None)

    def get_node_by_name(self, name):
        """
        Returns node that matches the name.

        Args:
            name (str): name of the node.
        Returns:
            NodeGraphQt.NodeObject: node object.
        """
        for node_id, node in self._model.nodes.items():
            if node.name() == name:
                return node

    def get_nodes_by_type(self, node_type):
        """
        Return all nodes by their node type identifier.
        (see: :attr:`NodeGraphQt.NodeObject.type_`)

        Args:
            node_type (str): node type identifier.

        Returns:
            list[NodeGraphQt.NodeObject]: list of nodes.
        """
        return [n for n in self._model.nodes.values() if n.type_ == node_type]

    def get_unique_name(self, name):
        """
        Creates a unique node name to avoid having nodes with the same name.

        Args:
            name (str): node name.

        Returns:
            str: unique node name.
        """
        name = ' '.join(name.split())
        node_names = [n.name() for n in self.all_nodes()]
        if name not in node_names:
            return name

        regex = re.compile(r'[\w ]+(?: )*(\d+)')
        search = regex.search(name)
        if not search:
            for x in range(1, len(node_names) + 2):
                new_name = '{} {}'.format(name, x)
                if new_name not in node_names:
                    return new_name

        version = search.group(1)
        name = name[:len(version) * -1].strip()
        for x in range(1, len(node_names) + 2):
            new_name = '{} {}'.format(name, x)
            if new_name not in node_names:
                return new_name

    def current_session(self):
        """
        Returns the file path to the currently loaded session.

        Returns:
            str: path to the currently loaded session
        """
        return self._model.session

    def clear_session(self):
        """
        Clears the current node graph session.
        """
        for n in self.all_nodes():
            if isinstance(n, BaseNode):
                for p in n.input_ports():
                    if p.locked():
                        p.set_locked(False, connected_ports=False)
                    p.clear_connections()
                for p in n.output_ports():
                    if p.locked():
                        p.set_locked(False, connected_ports=False)
                    p.clear_connections()
            self._undo_stack.push(NodeRemovedCmd(self, n))
        self._undo_stack.clear()
        self._model.session = ''

    def _serialize(self, nodes):
        """
        serialize nodes to a dict.
        (used internally by the node graph)

        Args:
            nodes (list[NodeGraphQt.Nodes]): list of node instances.

        Returns:
            dict: serialized data.
        """
        serial_data = {'graph': {}, 'nodes': {}, 'connections': []}
        nodes_data = {}

        # serialize graph session.
        serial_data['graph']['acyclic'] = self.acyclic()
        serial_data['graph']['pipe_collision'] = self.pipe_collision()

        # serialize nodes.
        for n in nodes:
            # update the node model.
            n.update_model()

            node_dict = n.model.to_dict
            nodes_data.update(node_dict)

        for n_id, n_data in nodes_data.items():
            serial_data['nodes'][n_id] = n_data

            # serialize connections
            inputs = n_data.pop('inputs') if n_data.get('inputs') else {}
            outputs = n_data.pop('outputs') if n_data.get('outputs') else {}

            for pname, conn_data in inputs.items():
                for conn_id, prt_names in conn_data.items():
                    for conn_prt in prt_names:
                        pipe = {IN_PORT: [n_id, pname],
                                OUT_PORT: [conn_id, conn_prt]}
                        if pipe not in serial_data['connections']:
                            serial_data['connections'].append(pipe)

            for pname, conn_data in outputs.items():
                for conn_id, prt_names in conn_data.items():
                    for conn_prt in prt_names:
                        pipe = {OUT_PORT: [n_id, pname],
                                IN_PORT: [conn_id, conn_prt]}
                        if pipe not in serial_data['connections']:
                            serial_data['connections'].append(pipe)

        if not serial_data['connections']:
            serial_data.pop('connections')

        return serial_data

    def _deserialize(self, data, relative_pos=False, pos=None):
        """
        deserialize node data.
        (used internally by the node graph)

        Args:
            data (dict): node data.
            relative_pos (bool): position node relative to the cursor.
            pos (tuple or list): custom x, y position.

        Returns:
            list[NodeGraphQt.Nodes]: list of node instances.
        """
        # update node graph properties.
        for attr_name, attr_value in data.get('graph', {}).items():
            if attr_name == 'acyclic':
                self.set_acyclic(attr_value)
            elif attr_name == 'pipe_collision':
                self.set_pipe_collision(attr_value)

        # build the nodes.
        nodes = {}
        for n_id, n_data in data.get('nodes', {}).items():
            identifier = n_data['type_']
            node = self._node_factory.create_node_instance(identifier)
            if node:
                node.NODE_NAME = n_data.get('name', node.NODE_NAME)
                # set properties.
                for prop in node.model.properties.keys():
                    if prop in n_data.keys():
                        node.model.set_property(prop, n_data[prop])
                # set custom properties.
                for prop, val in n_data.get('custom', {}).items():
                    node.model.set_property(prop, val)
                    if prop in node.view.widgets:
                        node.view.widgets[prop].set_value(val)

                nodes[n_id] = node
                self.add_node(node, n_data.get('pos'))

                if n_data.get('port_deletion_allowed', None):
                    node.set_ports({
                        'input_ports': n_data['input_ports'],
                        'output_ports': n_data['output_ports']
                    })

        # build the connections.
        for connection in data.get('connections', []):
            nid, pname = connection.get('in', ('', ''))
            in_node = nodes.get(nid) or self.get_node_by_id(nid)
            if not in_node:
                continue
            in_port = in_node.inputs().get(pname) if in_node else None

            nid, pname = connection.get('out', ('', ''))
            out_node = nodes.get(nid) or self.get_node_by_id(nid)
            if not out_node:
                continue
            out_port = out_node.outputs().get(pname) if out_node else None

            if in_port and out_port:
                # only connect if input port is not connected yet or input port
                # can have multiple connections.
                # important when duplicating nodes.
                allow_connection = any([not in_port.model.connected_ports,
                                        in_port.model.multi_connection])
                if allow_connection:
                    self._undo_stack.push(PortConnectedCmd(in_port, out_port))

        node_objs = nodes.values()
        if relative_pos:
            self._viewer.move_nodes([n.view for n in node_objs])
            [setattr(n.model, 'pos', n.view.xy_pos) for n in node_objs]
        elif pos:
            self._viewer.move_nodes([n.view for n in node_objs], pos=pos)
            [setattr(n.model, 'pos', n.view.xy_pos) for n in node_objs]

        return node_objs

    def serialize_session(self):
        """
        Serializes the current node graph layout to a dictionary.

        See Also:
            :meth:`NodeGraph.deserialize_session`,
            :meth:`NodeGraph.save_session`,
            :meth:`NodeGraph.load_session`

        Returns:
            dict: serialized session of the current node layout.
        """
        return self._serialize(self.all_nodes())

    def deserialize_session(self, layout_data):
        """
        Load node graph session from a dictionary object.

        See Also:
            :meth:`NodeGraph.serialize_session`,
            :meth:`NodeGraph.load_session`,
            :meth:`NodeGraph.save_session`

        Args:
            layout_data (dict): dictionary object containing a node session.
        """
        self.clear_session()
        self._deserialize(layout_data)
        self.clear_selection()
        self._undo_stack.clear()

    def save_session(self, file_path):
        """
        Saves the current node graph session layout to a `JSON` formatted file.

        See Also:
            :meth:`NodeGraph.serialize_session`,
            :meth:`NodeGraph.deserialize_session`,
            :meth:`NodeGraph.load_session`,

        Args:
            file_path (str): path to the saved node layout.
        """
        serialized_data = self._serialize(self.all_nodes())
        file_path = file_path.strip()
        with open(file_path, 'w') as file_out:
            json.dump(
                serialized_data,
                file_out,
                indent=2,
                separators=(',', ':')
            )

    def load_session(self, file_path):
        """
        Load node graph session layout file.

        See Also:
            :meth:`NodeGraph.deserialize_session`,
            :meth:`NodeGraph.serialize_session`,
            :meth:`NodeGraph.save_session`

        Args:
            file_path (str): path to the serialized layout file.
        """
        file_path = file_path.strip()
        if not os.path.isfile(file_path):
            raise IOError('file does not exist: {}'.format(file_path))

        self.clear_session()
        self.import_session(file_path)

    def import_session(self, file_path):
        """
        Import node graph session layout file.

        Args:
            file_path (str): path to the serialized layout file.
        """
        file_path = file_path.strip()
        if not os.path.isfile(file_path):
            raise IOError('file does not exist: {}'.format(file_path))

        try:
            with open(file_path) as data_file:
                layout_data = json.load(data_file)
        except Exception as e:
            layout_data = None
            print('Cannot read data from file.\n{}'.format(e))

        if not layout_data:
            return

        self._deserialize(layout_data)
        self._undo_stack.clear()
        self._model.session = file_path

        self.session_changed.emit(file_path)

    def copy_nodes(self, nodes=None):
        """
        Copy nodes to the clipboard.

        See Also:
            :meth:`NodeGraph.cut_nodes`

        Args:
            nodes (list[NodeGraphQt.BaseNode]):
                list of nodes (default: selected nodes).
        """
        nodes = nodes or self.selected_nodes()
        if not nodes:
            return False
        clipboard = QtWidgets.QApplication.clipboard()
        serial_data = self._serialize(nodes)
        serial_str = json.dumps(serial_data)
        if serial_str:
            clipboard.setText(serial_str)
            return True
        return False

    def cut_nodes(self, nodes=None):
        """
        Cut nodes to the clipboard.

        See Also:
            :meth:`NodeGraph.copy_nodes`

        Args:
            nodes (list[NodeGraphQt.BaseNode]):
                list of nodes (default: selected nodes).
        """
        nodes = nodes or self.selected_nodes()
        self.copy_nodes(nodes)
        self._undo_stack.beginMacro('cut nodes')
        [self._undo_stack.push(NodeRemovedCmd(self, n)) for n in nodes]
        self._undo_stack.endMacro()

    def paste_nodes(self):
        """
        Pastes nodes copied from the clipboard.
        """
        clipboard = QtWidgets.QApplication.clipboard()
        cb_text = clipboard.text()
        if not cb_text:
            return

        try:
            serial_data = json.loads(cb_text)
        except json.decoder.JSONDecodeError as e:
            print('ERROR: Can\'t Decode Clipboard Data:\n'
                  '"{}"'.format(cb_text))
            return

        self._undo_stack.beginMacro('pasted nodes')
        self.clear_selection()
        nodes = self._deserialize(serial_data, relative_pos=True)
        [n.set_selected(True) for n in nodes]
        self._undo_stack.endMacro()

    def duplicate_nodes(self, nodes):
        """
        Create duplicate copy from the list of nodes.

        Args:
            nodes (list[NodeGraphQt.BaseNode]): list of nodes.
        Returns:
            list[NodeGraphQt.BaseNode]: list of duplicated node instances.
        """
        if not nodes:
            return

        self._undo_stack.beginMacro('duplicate nodes')

        self.clear_selection()
        serial = self._serialize(nodes)
        new_nodes = self._deserialize(serial)
        offset = 50
        for n in new_nodes:
            x, y = n.pos()
            n.set_pos(x + offset, y + offset)
            n.set_property('selected', True)

        self._undo_stack.endMacro()
        return new_nodes

    def disable_nodes(self, nodes, mode=None):
        """
        Set weather to Disable or Enable specified nodes.

        See Also:
            :meth:`NodeObject.set_disabled`

        Args:
            nodes (list[NodeGraphQt.BaseNode]): list of node instances.
            mode (bool): (optional) disable state of the nodes.
        """
        if not nodes:
            return
        if mode is None:
            mode = not nodes[0].disabled()
        if len(nodes) > 1:
            text = {False: 'enable', True: 'disable'}[mode]
            text = '{} ({}) nodes'.format(text, len(nodes))
            self._undo_stack.beginMacro(text)
            [n.set_disabled(mode) for n in nodes]
            self._undo_stack.endMacro()
            return
        nodes[0].set_disabled(mode)

    def use_OpenGL(self):
        """
        Set the viewport to use QOpenGLWidget widget to draw the graph.
        """
        self._viewer.use_OpenGL()

    # auto layout node functions.
    # --------------------------------------------------------------------------

    @staticmethod
    def _update_node_rank(node, nodes_rank, down_stream=True):
        """
        Recursive function for updating the node ranking.

        Args:
            node (NodeGraphQt.BaseNode): node to start from.
            nodes_rank (dict): node ranking object to be updated.
            down_stream (bool): true to rank down stram.
        """
        if down_stream:
            node_values = node.connected_output_nodes().values()
        else:
            node_values = node.connected_input_nodes().values()

        connected_nodes = set()
        for nodes in node_values:
            connected_nodes.update(nodes)

        rank = nodes_rank[node] + 1
        for n in connected_nodes:
            if n in nodes_rank:
                nodes_rank[n] = max(nodes_rank[n], rank)
            else:
                nodes_rank[n] = rank
            NodeGraph._update_node_rank(n, nodes_rank, down_stream)

    @staticmethod
    def _compute_node_rank(nodes, down_stream=True):
        """
        Compute the ranking of nodes.

        Args:
            nodes (list[NodeGraphQt.BaseNode]): nodes to start ranking from.
            down_stream (bool): true to compute down stream.

        Returns:
            dict: {NodeGraphQt.BaseNode: node_rank, ...}
        """
        nodes_rank = {}
        for node in nodes:
            nodes_rank[node] = 0
            NodeGraph._update_node_rank(node, nodes_rank, down_stream)
        return nodes_rank

    def auto_layout_nodes(self, nodes=None, down_stream=True, start_nodes=None):
        """
        Auto layout the nodes in the node graph.

        Note:
            If the node graph is acyclic then the ``start_nodes`` will need
            to be specified.

        Args:
            nodes (list[NodeGraphQt.BaseNode]): list of nodes to auto layout
                if nodes is None then all nodes is layed out.
            down_stream (bool): false to layout up stream.
            start_nodes (list[NodeGraphQt.BaseNode]):
                list of nodes to start the auto layout from (Optional).
        """
        self.begin_undo('Auto Layout Nodes')

        nodes = nodes or self.all_nodes()

        # filter out the backdrops.
        backdrops = {
            n: n.nodes() for n in nodes if isinstance(n, BackdropNode)
        }
        filtered_nodes = [n for n in nodes if not isinstance(n, BackdropNode)]

        start_nodes = start_nodes or []
        if down_stream:
            start_nodes += [
                n for n in filtered_nodes
                if not any(n.connected_input_nodes().values())
            ]
        else:
            start_nodes += [
                n for n in filtered_nodes
                if not any(n.connected_output_nodes().values())
            ]

        if not start_nodes:
            return

        node_views = [n.view for n in nodes]
        nodes_center_0 = self.viewer().nodes_rect_center(node_views)

        nodes_rank = NodeGraph._compute_node_rank(start_nodes, down_stream)

        rank_map = {}
        for node, rank in nodes_rank.items():
            if rank in rank_map:
                rank_map[rank].append(node)
            else:
                rank_map[rank] = [node]

        if NODE_LAYOUT_DIRECTION is NODE_LAYOUT_HORIZONTAL:
            current_x = 0
            node_height = 120
            line_counter = 0
            for rank in sorted(range(len(rank_map)), reverse=not down_stream):
                ranked_nodes = rank_map[rank]
                line_counter += 1
                square_side = math.floor(math.sqrt(len(nodes)))
                max_width = max([node.view.width for node in ranked_nodes])
                #if math.floor(line_counter/5)%2 == 0:
                current_x = line_counter%(square_side)*(max_width + 100)

                current_y = math.floor(line_counter/square_side)*node.view.height*2
                for idx, node in enumerate(ranked_nodes):
                    dy = max(node_height, node.view.height)
                    current_y += 0 if idx == 0 else dy
                    node.set_pos(current_x, current_y)
                    current_y += dy * 0.5 + 10


        elif NODE_LAYOUT_DIRECTION is NODE_LAYOUT_VERTICAL:
            current_y = 0
            node_width = 250
            for rank in sorted(range(len(rank_map)), reverse=not down_stream):
                ranked_nodes = rank_map[rank]
                max_height = max([node.view.height for node in ranked_nodes])
                current_y += max_height
                current_x = 0
                for idx, node in enumerate(ranked_nodes):
                    dx = max(node_width, node.view.width)
                    current_x += 0 if idx == 0 else dx
                    node.set_pos(current_x, current_y)
                    current_x += dx * 0.5 + 10

                current_y += max_height * 0.5 + 100

        nodes_center_1 = self.viewer().nodes_rect_center(node_views)
        dx = nodes_center_0[0] - nodes_center_1[0]
        dy = nodes_center_0[1] - nodes_center_1[1]
        [n.set_pos(n.x_pos() + dx, n.y_pos() + dy) for n in nodes]

        # wrap the backdrop nodes.
        for backdrop, contained_nodes in backdrops.items():
            backdrop.wrap_nodes(contained_nodes)

        self.end_undo()

    # convenience dialog functions.
    # --------------------------------------------------------------------------

    def question_dialog(self, text, title='Node Graph'):
        """
        Prompts a question open dialog with ``"Yes"`` and ``"No"`` buttons in
        the node graph.

        Note:
            Convenience function to
            :meth:`NodeGraph.viewer().question_dialog`

        Args:
            text (str): question text.
            title (str): dialog window title.

        Returns:
            bool: true if user clicked yes.
        """
        return self._viewer.question_dialog(text, title)

    def message_dialog(self, text, title='Node Graph'):
        """
        Prompts a file open dialog in the node graph.

        Note:
            Convenience function to
            :meth:`NodeGraph.viewer().message_dialog`

        Args:
            text (str): message text.
            title (str): dialog window title.
        """
        self._viewer.message_dialog(text, title)

    def load_dialog(self, current_dir=None, ext=None):
        """
        Prompts a file open dialog in the node graph.

        Note:
            Convenience function to
            :meth:`NodeGraph.viewer().load_dialog`

        Args:
            current_dir (str): path to a directory.
            ext (str): custom file type extension (default: ``"json"``)

        Returns:
            str: selected file path.
        """
        return self._viewer.load_dialog(current_dir, ext)

    def save_dialog(self, current_dir=None, ext=None):
        """
        Prompts a file save dialog in the node graph.

        Note:
            Convenience function to
            :meth:`NodeGraph.viewer().save_dialog`

        Args:
            current_dir (str): path to a directory.
            ext (str): custom file type extension (default: ``"json"``)

        Returns:
            str: selected file path.
        """
        return self._viewer.save_dialog(current_dir, ext)

    # group node / sub graph.
    # --------------------------------------------------------------------------

    def _on_close_sub_graph_tab(self, index):
        """
        Called when the close button is clicked on a expanded sub graph tab.

        Args:
            index (int): tab index.
        """
        node_id = self.widget.tabToolTip(index)
        group_node = self.get_node_by_id(node_id)
        self.collapse_group_node(group_node)

    @property
    def is_root(self):
        """
        Returns if the node graph controller is the root graph.

        Returns:
            bool: true is the node graph is root.
        """
        return True

    @property
    def sub_graphs(self):
        """
        Returns expanded group node sub graphs.

        Returns:
            dict: {<node_id>: <sub_graph>}
        """
        return self._sub_graphs

    # def graph_rect(self):
    #     """
    #     Get the graph viewer range (scene size).
    #
    #     Returns:
    #         list[float]: [x, y, width, height].
    #     """
    #     return self._viewer.scene_rect()
    #
    # def set_graph_rect(self, rect):
    #     """
    #     Set the graph viewer range (scene size).
    #
    #     Args:
    #         rect (list[float]): [x, y, width, height].
    #     """
    #     self._viewer.set_scene_rect(rect)

    def expand_group_node(self, node):
        """
        Expands a group node session in a new tab.

        Args:
            node (NodeGraphQt.GroupNode): group node.

        Returns:
            SubGraph: sub node graph used to manage the group node session.
        """
        if not isinstance(node, GroupNode):
            return
        if self._widget is None:
            raise RuntimeError('NodeGraph.widget not initialized!')

        self.viewer().clear_key_state()
        self.viewer().clearFocus()

        if node.id in self._sub_graphs:
            sub_graph = self._sub_graphs[node.id]
            tab_index = self._widget.indexOf(sub_graph.widget)
            self._widget.setCurrentIndex(tab_index)
            return sub_graph

        # build new sub graph.
        node_factory = copy.deepcopy(self.node_factory)
        sub_graph = SubGraph(self, node=node, node_factory=node_factory)

        # populate the sub graph.
        session = node.get_sub_graph_session()
        sub_graph.deserialize_session(session)

        # store reference to expanded.
        self._sub_graphs[node.id] = sub_graph

        # open new tab at root level.
        self.widget.add_viewer(sub_graph.widget, node.name(), node.id)

        return sub_graph

    def collapse_group_node(self, node):
        """
        Collapse a group node session tab and it's expanded child sub graphs.

        Args:
            node (NodeGraphQt.GroupNode): group node.
        """
        assert isinstance(node, GroupNode), 'node must be a GroupNode instance.'
        if self._widget is None:
            return

        if node.id not in self._sub_graphs:
            err = '{} sub graph not initialized!'.format(node.name())
            raise RuntimeError(err)

        sub_graph = self._sub_graphs.pop(node.id)
        sub_graph.collapse_group_node(node)

        # remove the sub graph tab.
        self.widget.remove_viewer(sub_graph.widget)

        # TODO: delete sub graph hmm... not sure if I need this here.
        del sub_graph


class SubGraph(NodeGraph):
    """
    The ``SubGraph`` class is just like the ``NodeGraph`` but is the main
    controller for managing the expanded node graph for a group node.

    Inherited from: :class:`NodeGraphQt.NodeGraph`

    .. image:: _images/sub_graph.png
        :width: 70%

    -
    """

    def __init__(self, parent=None, node=None, node_factory=None):
        """
        Args:
            parent (object): object parent.
            node (GroupNode): group node related to this sub graph.
            node_factory (NodeFactory): override node factory.
        """
        super(SubGraph, self).__init__(parent, node_factory=node_factory)

        # sub graph attributes.
        self._node = node
        self._parent_graph = parent
        self._subviewer_widget = None

        if self._parent_graph.is_root:
            self._initialized_graphs = [self]
            self._sub_graphs[self._node.id] = self
        else:
            # delete attributes if not top level sub graph.
            del self._widget
            del self._sub_graphs

    def __repr__(self):
        return '<{}("{}") object at {}>'.format(
            self.__class__.__name__, self._node.name(), hex(id(self)))

    def _register_builtin_nodes(self):
        """
        Register the default builtin nodes to the :meth:`NodeGraph.node_factory`
        """
        return

    def _build_port_nodes(self):
        """
        Build the corresponding input & output nodes from the parent node ports
        and remove any port nodes that are outdated..

        Returns:
             tuple(dict, dict): input nodes, output nodes.
        """
        # build the parent input port nodes.
        input_nodes = {n.name(): n for n in
                       self.get_nodes_by_type(PortInputNode.type_)}
        for port in self.node.input_ports():
            if port.name() not in input_nodes:
                input_node = PortInputNode(parent_port=port)
                input_node.NODE_NAME = port.name()
                input_node.model.set_property('name', port.name())
                input_node.add_output(port.name())
                input_nodes[port.name()] = input_node
                self.add_node(input_node, selected=False, push_undo=False)
                x, y = input_node.pos()
                if NODE_LAYOUT_DIRECTION is NODE_LAYOUT_HORIZONTAL:
                    x -= 100
                elif NODE_LAYOUT_DIRECTION is NODE_LAYOUT_VERTICAL:
                    y -= 100
                input_node.set_property('pos', [x, y], push_undo=False)

        # build the parent output port nodes.
        output_nodes = {n.name(): n for n in
                        self.get_nodes_by_type(PortOutputNode.type_)}
        for port in self.node.output_ports():
            if port.name() not in output_nodes:
                output_node = PortOutputNode(parent_port=port)
                output_node.NODE_NAME = port.name()
                output_node.model.set_property('name', port.name())
                output_node.add_input(port.name())
                output_nodes[port.name()] = output_node
                self.add_node(output_node, selected=False, push_undo=False)
                x, y = output_node.pos()
                if NODE_LAYOUT_DIRECTION is NODE_LAYOUT_HORIZONTAL:
                    x += 100
                elif NODE_LAYOUT_DIRECTION is NODE_LAYOUT_VERTICAL:
                    y += 100
                output_node.set_property('pos', [x, y], push_undo=False)

        return input_nodes, output_nodes

    def _deserialize(self, data, relative_pos=False, pos=None):
        """
        deserialize node data.
        (used internally by the node graph)

        Args:
            data (dict): node data.
            relative_pos (bool): position node relative to the cursor.
            pos (tuple or list): custom x, y position.

        Returns:
            list[NodeGraphQt.Nodes]: list of node instances.
        """
        # update node graph properties.
        for attr_name, attr_value in data.get('graph', {}).items():
            if attr_name == 'acyclic':
                self.set_acyclic(attr_value)
            elif attr_name == 'pipe_collision':
                self.set_pipe_collision(attr_value)

        # build the port input & output nodes here.
        input_nodes, output_nodes = self._build_port_nodes()

        # build the nodes.
        nodes = {}
        for n_id, n_data in data.get('nodes', {}).items():
            identifier = n_data['type_']
            name = n_data.get('name')
            if identifier == PortInputNode.type_:
                nodes[n_id] = input_nodes[name]
                nodes[n_id].set_pos(*(n_data.get('pos') or [0, 0]))
                continue
            elif identifier == PortOutputNode.type_:
                nodes[n_id] = output_nodes[name]
                nodes[n_id].set_pos(*(n_data.get('pos') or [0, 0]))
                continue

            node = self._node_factory.create_node_instance(identifier)
            if not node:
                continue

            node.NODE_NAME = name or node.NODE_NAME
            # set properties.
            for prop in node.model.properties.keys():
                if prop in n_data.keys():
                    node.model.set_property(prop, n_data[prop])
            # set custom properties.
            for prop, val in n_data.get('custom', {}).items():
                node.model.set_property(prop, val)

            nodes[n_id] = node
            self.add_node(node, n_data.get('pos'))

            if n_data.get('port_deletion_allowed', None):
                node.set_ports({
                    'input_ports': n_data['input_ports'],
                    'output_ports': n_data['output_ports']
                })

        # build the connections.
        for connection in data.get('connections', []):
            nid, pname = connection.get('in', ('', ''))
            in_node = nodes.get(nid)
            if not in_node:
                continue
            in_port = in_node.inputs().get(pname) if in_node else None

            nid, pname = connection.get('out', ('', ''))
            out_node = nodes.get(nid)
            if not out_node:
                continue
            out_port = out_node.outputs().get(pname) if out_node else None

            if in_port and out_port:
                self._undo_stack.push(PortConnectedCmd(in_port, out_port))

        node_objs = list(nodes.values())
        if relative_pos:
            self._viewer.move_nodes([n.view for n in node_objs])
            [setattr(n.model, 'pos', n.view.xy_pos) for n in node_objs]
        elif pos:
            self._viewer.move_nodes([n.view for n in node_objs], pos=pos)
            [setattr(n.model, 'pos', n.view.xy_pos) for n in node_objs]

        return node_objs

    def _on_navigation_changed(self, node_id, rm_node_ids):
        """
        Slot when the node navigation widget has changed.

        Args:
            node_id (str): selected group node id.
            rm_node_ids (list[str]): list of group node id to remove.
        """
        # collapse child sub graphs.
        for rm_node_id in rm_node_ids:
            child_node = self.sub_graphs[rm_node_id].node
            self.collapse_group_node(child_node)

        # show the selected node id sub graph.
        sub_graph = self.sub_graphs.get(node_id)
        if sub_graph:
            self.widget.show_viewer(sub_graph.subviewer_widget)
            sub_graph.viewer().setFocus()

    @property
    def is_root(self):
        """
        Returns if the node graph controller is the main root graph.

        Returns:
            bool: true is the node graph is root.
        """
        return False

    @property
    def sub_graphs(self):
        """
        Returns expanded group node sub graphs.

        Returns:
            dict: {<node_id>: <sub_graph>}
        """
        if self.parent_graph.is_root:
            return self._sub_graphs
        return self.parent_graph.sub_graphs

    @property
    def initialized_graphs(self):
        """
        Returns a list of the sub graphs in the order they were initialized.

        Returns:
            list[NodeGraphQt.SubGraph]: list of sub graph objects.
        """
        if self._parent_graph.is_root:
            return self._initialized_graphs
        return self._parent_graph.initialized_graphs

    @property
    def widget(self):
        """
        The sub graph widget from the top most sub graph.

        Returns:
            SubGraphWidget: node graph widget.
        """
        if self.parent_graph.is_root:
            if self._widget is None:
                self._widget = SubGraphWidget()
                self._widget.add_viewer(self.subviewer_widget,
                                        self.node.name(),
                                        self.node.id)
                # connect the navigator widget signals.
                navigator = self._widget.navigator
                navigator.navigation_changed.connect(
                    self._on_navigation_changed
                )
            return self._widget
        return self.parent_graph.widget

    @property
    def navigation_widget(self):
        """
        The navigation widget from the top most sub graph.

        Returns:
            NodeNavigationWidget: navigation widget.
        """
        if self.parent_graph.is_root:
            return self.widget.navigator
        return self.parent_graph.navigation_widget

    @property
    def subviewer_widget(self):
        """
        The widget to the sub graph.

        Returns:
            PySide2.QtWidgets.QWidget: node graph widget.
        """
        if self._subviewer_widget is None:
            self._subviewer_widget = QtWidgets.QWidget()
            layout = QtWidgets.QVBoxLayout(self._subviewer_widget)
            layout.setContentsMargins(0, 0, 0, 0)
            layout.setSpacing(1)
            layout.addWidget(self._viewer)
        return self._subviewer_widget

    @property
    def parent_graph(self):
        """
        The parent node graph controller.

        Returns:
            NodeGraphQt.NodeGraph or NodeGraphQt.SubGraph: parent graph.
        """
        return self._parent_graph

    @property
    def node(self):
        """
        Returns the parent node to the sub graph.

        .. image:: _images/group_node.png
            :width: 250px

        Returns:
            NodeGraphQt.GroupNode: group node.
        """
        return self._node

    def delete_node(self, node, push_undo=True):
        """
        Remove the node from the node sub graph.

        Note:
            :class:`.PortInputNode` & :class:`.PortOutputNode` can't be deleted
            as they are connected to a :class:`.Port` to remove these port nodes
            see :meth:`BaseNode.delete_input`, :meth:`BaseNode.delete_output`.

        Args:
            node (NodeGraphQt.BaseNode): node object.
            push_undo (bool): register the command to the undo stack. (default: True)
        """
        port_nodes = self.get_input_port_nodes() + self.get_output_port_nodes()
        if node in port_nodes and node.parent_port is not None:
            # note: port nodes can only be deleted by deleting the parent
            #       port object.
            raise RuntimeError(
                'Can\'t delete node "{}" it is attached to port "{}"'
                .format(node, node.parent_port)
            )
        super(SubGraph, self).delete_node(node, push_undo=push_undo)

    def collapse_graph(self, clear_session=True):
        """
        Collapse the current sub graph and hide its widget.

        Args:
            clear_session (bool): clear the current session.
        """
        # update the group node.
        serialized_session = self.serialize_session()
        self.node.set_sub_graph_session(serialized_session)

        # close the visible widgets.
        if self._undo_view:
            self._undo_view.close()

        if self._subviewer_widget:
            self.widget.hide_viewer(self._subviewer_widget)

        if clear_session:
            self.clear_session()

    def expand_group_node(self, node):
        """
        Expands a group node session in current sub view.

        Args:
            node (NodeGraphQt.GroupNode): group node.

        Returns:
            SubGraph: sub node graph used to manage the group node session.
        """
        assert isinstance(node, GroupNode), 'node must be a GroupNode instance.'
        if self._subviewer_widget is None:
            raise RuntimeError('SubGraph.widget not initialized!')

        self.viewer().clear_key_state()
        self.viewer().clearFocus()

        if node.id in self.sub_graphs:
            sub_graph_viewer = self.sub_graphs[node.id].viewer()
            sub_graph_viewer.setFocus()
            return self.sub_graphs[node.id]

        # collapse expanded child sub graphs.
        group_ids = [n.id for n in self.all_nodes() if isinstance(n, GroupNode)]
        for grp_node_id, grp_sub_graph in self.sub_graphs.items():
            # collapse current group node.
            if grp_node_id in group_ids:
                grp_node = self.get_node_by_id(grp_node_id)
                self.collapse_group_node(grp_node)

            # close the widgets
            grp_sub_graph.collapse_graph(clear_session=False)

        # build new sub graph.
        node_factory = copy.deepcopy(self.node_factory)
        sub_graph = SubGraph(self, node=node, node_factory=node_factory)

        # populate the sub graph.
        serialized_session = node.get_sub_graph_session()
        sub_graph.deserialize_session(serialized_session)

        # open new sub graph view.
        self.widget.add_viewer(sub_graph.subviewer_widget,
                               node.name(),
                               node.id)

        # store the references.
        self.sub_graphs[node.id] = sub_graph
        self.initialized_graphs.append(sub_graph)

        return sub_graph

    def collapse_group_node(self, node):
        """
        Collapse a group node session and it's expanded child sub graphs.

        Args:
            node (NodeGraphQt.GroupNode): group node.
        """
        # update the references.
        sub_graph = self.sub_graphs.pop(node.id, None)
        if not sub_graph:
            return

        init_idx = self.initialized_graphs.index(sub_graph) + 1
        for sgraph in reversed(self.initialized_graphs[init_idx:]):
            self.initialized_graphs.remove(sgraph)

        # collapse child sub graphs here.
        child_ids = [
            n.id for n in sub_graph.all_nodes() if isinstance(n, GroupNode)
        ]
        for child_id in child_ids:
            if self.sub_graphs.get(child_id):
                child_graph = self.sub_graphs.pop(child_id)
                child_graph.collapse_graph(clear_session=True)
                # remove child viewer widget.
                self.widget.remove_viewer(child_graph.subviewer_widget)

        sub_graph.collapse_graph(clear_session=True)
        self.widget.remove_viewer(sub_graph.subviewer_widget)

    def get_input_port_nodes(self):
        """
        Return all the port nodes related to the group node input ports.

        .. image:: _images/port_in_node.png
            :width: 150px

        -

        See Also:
            :meth:`NodeGraph.get_nodes_by_type`,
            :meth:`SubGraph.get_output_port_nodes`

        Returns:
            list[NodeGraphQt.PortInputNode]: input nodes.
        """
        return self.get_nodes_by_type(PortInputNode.type_)

    def get_output_port_nodes(self):
        """
        Return all the port nodes related to the group node output ports.

        .. image:: _images/port_out_node.png
            :width: 150px

        -

        See Also:
            :meth:`NodeGraph.get_nodes_by_type`,
            :meth:`SubGraph.get_input_port_nodes`

        Returns:
            list[NodeGraphQt.PortOutputNode]: output nodes.
        """
        return self.get_nodes_by_type(PortOutputNode.type_)

    def get_node_by_port(self, port):
        """
        Returns the node related to the parent group node port object.

        Args:
            port (NodeGraphQt.Port): parent node port object.

        Returns:
            PortInputNode or PortOutputNode: port node object.
        """
        func_type = {IN_PORT: self.get_input_port_nodes,
                     OUT_PORT: self.get_output_port_nodes}
        for n in func_type.get(port.type_(), []):
            if port == n.parent_port:
                return n
